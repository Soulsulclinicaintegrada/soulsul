import os
import re
import calendar
import sqlite3
import hashlib
import hmac
import secrets
import unicodedata
from io import BytesIO
from datetime import date, datetime, timedelta

import pandas as pd
import requests
import streamlit as st
import altair as alt
from docx import Document
from lxml import etree
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    REPORTLAB_DISPONIVEL = True
except ImportError:
    REPORTLAB_DISPONIVEL = False

try:
    import openpyxl  # noqa: F401
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_DISPONIVEL = True
except ImportError:
    OPENPYXL_DISPONIVEL = False

try:
    from pypdf import PdfReader

    PDF_LEITURA_DISPONIVEL = True
except ImportError:
    try:
        from PyPDF2 import PdfReader

        PDF_LEITURA_DISPONIVEL = True
    except ImportError:
        PDF_LEITURA_DISPONIVEL = False


st.set_page_config(page_title="Soul Sul Sistema", layout="wide")

DB_PATH = "clinica.db"
DOCS_DIR = "documentos"
TEMPLATE_PATH = "modelo_documento.docx"
FORMAS_PAGAMENTO = ["Pix", "Dinheiro", "Debito", "Credito", "Boleto"]
STATUS_CONTAS_PAGAR = ["A vencer", "Pago", "Atrasado", "Suspenso", "Cancelado"]
DESCRICAO_CONTRATO = "Contrato odontologico"
USUARIO_ADMIN_INICIAL = "admin"
SENHA_ADMIN_INICIAL = "admin123"
MODULOS_SISTEMA = {
    "Dashboard": "acesso_dashboard",
    "Pacientes": "acesso_pacientes",
    "Editar Paciente": "acesso_pacientes",
    "Contratos": "acesso_contratos",
    "Editar Contrato": "acesso_contratos",
    "Importacoes": "acesso_contratos",
    "Financeiro": "acesso_financeiro",
    "Usuarios": "acesso_usuarios",
}


def aplicar_tema_visual():
    st.markdown(
        """
        <style>
        :root {
            --ss-bg: #f6f8f7;
            --ss-surface: #ffffff;
            --ss-sidebar: linear-gradient(180deg, #173b33 0%, #102c26 100%);
            --ss-border: #dfe7e3;
            --ss-shadow: 0 10px 30px rgba(16, 44, 38, 0.08);
            --ss-text: #18312b;
            --ss-muted: #667a73;
            --ss-accent: #2fa36b;
            --ss-accent-soft: #e7f6ee;
        }

        .stApp {
            background:
                radial-gradient(circle at top left, rgba(47, 163, 107, 0.08), transparent 24%),
                radial-gradient(circle at top right, rgba(47, 163, 107, 0.06), transparent 18%),
                var(--ss-bg);
        }

        section[data-testid="stSidebar"] {
            background: var(--ss-sidebar);
            border-right: 1px solid rgba(255,255,255,0.08);
        }

        section[data-testid="stSidebar"] * {
            color: #f4f7f5;
        }

        .ss-brand {
            background: linear-gradient(180deg, rgba(255,255,255,0.10), rgba(255,255,255,0.04));
            border: 1px solid rgba(255,255,255,0.12);
            box-shadow: 0 18px 35px rgba(0,0,0,0.18);
            border-radius: 22px;
            padding: 18px 16px;
            margin: 0 0 14px 0;
        }

        .ss-brand-title {
            font-size: 1.55rem;
            font-weight: 700;
            letter-spacing: 0.02em;
            margin: 0;
        }

        .ss-brand-subtitle {
            color: rgba(244,247,245,0.78);
            font-size: 0.92rem;
            margin-top: 2px;
        }

        div[data-testid="stRadio"] label {
            background: transparent;
            border-radius: 14px;
            border: 1px solid transparent;
            padding: 8px 12px;
            margin-bottom: 6px;
        }

        div[data-testid="stRadio"] label:hover {
            background: rgba(255,255,255,0.08);
            border-color: rgba(255,255,255,0.10);
        }

        div[data-testid="stRadio"] label:has(input:checked) {
            background: linear-gradient(90deg, rgba(255,255,255,0.96), rgba(233, 248, 240, 0.96));
            box-shadow: 0 8px 24px rgba(0,0,0,0.18);
            border-color: rgba(47, 163, 107, 0.18);
        }

        div[data-testid="stRadio"] label:has(input:checked) p {
            color: var(--ss-text) !important;
            font-weight: 700;
        }

        [data-testid="stMetric"] {
            background: var(--ss-surface);
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            padding: 14px 16px;
            box-shadow: var(--ss-shadow);
        }

        [data-testid="stMetricLabel"], [data-testid="stMetricValue"] {
            color: var(--ss-text);
        }

        div[data-testid="stHorizontalBlock"] > div:has(> [data-testid="stMetric"]) {
            align-self: stretch;
        }

        div[data-testid="stTextInput"] > div > div,
        div[data-testid="stNumberInput"] > div > div,
        div[data-testid="stDateInput"] > div > div,
        div[data-testid="stSelectbox"] > div > div,
        div[data-testid="stMultiSelect"] > div > div,
        div[data-testid="stTextArea"] > div > div {
            border-radius: 16px !important;
            border: 1px solid var(--ss-border) !important;
            box-shadow: none !important;
            background: rgba(255,255,255,0.96) !important;
        }

        div[data-testid="stDataFrame"] {
            border: 1px solid var(--ss-border);
            border-radius: 20px;
            overflow: hidden;
            box-shadow: var(--ss-shadow);
            background: var(--ss-surface);
        }

        .stButton > button,
        .stDownloadButton > button {
            border-radius: 14px;
            border: 1px solid rgba(47, 163, 107, 0.18);
            background: linear-gradient(180deg, #ffffff 0%, #edf7f1 100%);
            color: var(--ss-text);
            font-weight: 600;
            box-shadow: 0 8px 20px rgba(47, 163, 107, 0.10);
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            border-color: rgba(47, 163, 107, 0.32);
            color: var(--ss-accent);
        }

        .stTabs [data-baseweb="tab-list"] {
            gap: 10px;
            background: transparent;
        }

        .stTabs [data-baseweb="tab"] {
            background: rgba(255,255,255,0.88);
            border-radius: 14px;
            border: 1px solid var(--ss-border);
            padding: 10px 16px;
            box-shadow: 0 6px 18px rgba(16, 44, 38, 0.06);
        }

        .stTabs [aria-selected="true"] {
            background: linear-gradient(180deg, #ffffff 0%, #eaf7f0 100%) !important;
            border-color: rgba(47, 163, 107, 0.22) !important;
        }

        h1, h2, h3 {
            color: var(--ss-text);
            letter-spacing: -0.02em;
        }

        p, label, .stMarkdown, .stCaption {
            color: var(--ss-text);
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def agora_str():
    return datetime.now().isoformat(sep=" ", timespec="seconds")


def formatar_data_br(data_obj):
    return data_obj.strftime("%d/%m/%Y")


def formatar_data_br_valor(valor):
    texto_bruto = str(valor or "").strip()
    if texto_bruto.lower() in {"nat", "nan", "none"}:
        return ""
    data_convertida = parse_data_contrato(valor)
    if data_convertida:
        return formatar_data_br(data_convertida)
    return texto_bruto


def formatar_data_hora_br_valor(valor):
    texto = str(valor or "").strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return ""
    try:
        return datetime.fromisoformat(texto).strftime("%d/%m/%Y %H:%M:%S")
    except ValueError:
        pass
    return formatar_data_br_valor(valor)


def formatar_prontuario_valor(valor):
    texto = texto_importacao(valor)
    if not texto:
        return ""
    if re.fullmatch(r"-?\d+\.0+", texto):
        return texto.split(".")[0]
    return texto


def formatar_parcela_valor(valor):
    texto = str(valor or "").strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return "-"
    try:
        numero = float(str(valor).replace(",", "."))
        if numero.is_integer():
            return str(int(numero))
        return texto
    except (TypeError, ValueError):
        texto_limpo = texto_importacao(valor)
        return texto_limpo or "-"


def forma_pagamento_a_vista(forma_pagamento):
    return normalizar_forma_pagamento(forma_pagamento) in {"Pix", "Dinheiro", "Debito", "Credito"}


def parse_data_contrato(valor):
    texto = str(valor or "").strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return None

    try:
        return datetime.fromisoformat(texto).date()
    except ValueError:
        pass

    formatos = ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M:%S.%f"]
    for formato in formatos:
        try:
            return datetime.strptime(texto, formato).date()
        except ValueError:
            continue
    try:
        data_pandas = pd.to_datetime(texto, dayfirst=True, errors="coerce")
        if not pd.isna(data_pandas):
            return data_pandas.date()
    except Exception:
        pass
    return None


def adicionar_meses(data_base, quantidade_meses):
    mes = data_base.month - 1 + quantidade_meses
    ano = data_base.year + mes // 12
    mes = mes % 12 + 1
    ultimo_dia = [31, 29 if ano % 4 == 0 and (ano % 100 != 0 or ano % 400 == 0) else 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31][mes - 1]
    dia = min(data_base.day, ultimo_dia)
    return date(ano, mes, dia)


def extrair_texto_pdf_upload(arquivo_pdf):
    if not PDF_LEITURA_DISPONIVEL:
        raise ValueError("Leitura de PDF indisponivel. Instale pypdf ou PyPDF2 para habilitar essa funcao.")

    arquivo_pdf.seek(0)
    reader = PdfReader(BytesIO(arquivo_pdf.read()))
    textos = []
    for pagina in reader.pages:
        texto = pagina.extract_text() or ""
        if texto.strip():
            textos.append(texto)
    return "\n".join(textos)


def extrair_por_rotulo(texto, rotulos):
    for rotulo in rotulos:
        padrao = rf"{rotulo}\s*[:\-]?\s*(.+)"
        match = re.search(padrao, texto, flags=re.IGNORECASE)
        if match:
            valor = match.group(1).strip()
            valor = valor.split("\n")[0].strip()
            if valor:
                return valor
    return ""


def limpar_trecho_endereco(valor):
    texto = str(valor or "").strip(" -,:;")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def separar_endereco_campos(endereco, numero, bairro, cidade, estado):
    endereco = limpar_trecho_endereco(endereco)
    numero = limpar_trecho_endereco(numero)
    bairro = limpar_trecho_endereco(bairro)
    cidade = limpar_trecho_endereco(cidade)
    estado = limpar_trecho_endereco(estado)

    texto_completo = " ".join([parte for parte in [endereco, numero, bairro, cidade, estado] if parte]).strip()

    if not numero:
        match_numero = re.search(
            r"\b(\d+[A-Za-z]?(?:\s*complemento\s*[:\-]?\s*[^,]+)?)",
            texto_completo,
            flags=re.IGNORECASE,
        )
        if match_numero:
            numero = limpar_trecho_endereco(match_numero.group(1))

    if not estado:
        match_estado = re.search(r"\b([A-Z]{2})\b", texto_completo)
        if match_estado:
            estado = limpar_trecho_endereco(match_estado.group(1))

    if not cidade:
        match_cidade = re.search(r"cidade\s*[:\-]?\s*([A-Za-zÀ-ÿ\s]+)", texto_completo, flags=re.IGNORECASE)
        if match_cidade:
            cidade = limpar_trecho_endereco(match_cidade.group(1))
        elif " - " in texto_completo:
            partes = [limpar_trecho_endereco(parte) for parte in texto_completo.split(" - ") if limpar_trecho_endereco(parte)]
            if partes:
                cidade = partes[-1]

    if not bairro:
        match_bairro = re.search(r"bairro\s*[:\-]?\s*([A-Za-zÀ-ÿ\s]+)", texto_completo, flags=re.IGNORECASE)
        if match_bairro:
            bairro = limpar_trecho_endereco(match_bairro.group(1))

    endereco = re.sub(r"\bcidade\s*[:\-]?\s*.*$", "", endereco, flags=re.IGNORECASE)
    endereco = re.sub(r"\bestado\s*[:\-]?\s*.*$", "", endereco, flags=re.IGNORECASE)
    endereco = re.sub(r"\bcep\s*[:\-]?\s*.*$", "", endereco, flags=re.IGNORECASE)
    endereco = limpar_trecho_endereco(endereco)

    cidade = re.sub(r"estado\s*[:\-]?\s*.*$", "", cidade, flags=re.IGNORECASE)
    bairro = re.sub(r"cidade\s*[:\-]?\s*.*$", "", bairro, flags=re.IGNORECASE)
    bairro = re.sub(r"estado\s*[:\-]?\s*.*$", "", bairro, flags=re.IGNORECASE)
    bairro = re.sub(r"\bcidade\b$", "", bairro, flags=re.IGNORECASE)
    cidade = limpar_trecho_endereco(cidade)
    bairro = limpar_trecho_endereco(bairro)

    if bairro and cidade and cidade in bairro:
        bairro = limpar_trecho_endereco(bairro.replace(cidade, ""))
    if cidade and estado and estado in cidade:
        cidade = limpar_trecho_endereco(cidade.replace(estado, ""))

    return {
        "endereco": endereco,
        "numero": numero,
        "bairro": bairro,
        "cidade": cidade,
        "estado": estado,
    }


def extrair_dados_paciente_pdf(texto):
    if not texto.strip():
        return {}

    linhas = [linha.strip() for linha in texto.splitlines() if linha.strip()]
    texto_unico = "\n".join(linhas)

    cpf = ""
    match_cpf = re.search(r"\b\d{3}\.?\d{3}\.?\d{3}-?\d{2}\b", texto_unico)
    if match_cpf:
        cpf = limpar_cpf(match_cpf.group(0))

    telefone = ""
    match_tel = re.search(r"(\(?\d{2}\)?\s?\d{4,5}-?\d{4})", texto_unico)
    if match_tel:
        telefone = match_tel.group(1)

    cep = ""
    match_cep = re.search(r"\b\d{5}-?\d{3}\b", texto_unico)
    if match_cep:
        cep = match_cep.group(0)

    nascimento = ""
    match_data = re.search(r"\b\d{2}/\d{2}/\d{4}\b", texto_unico)
    if match_data:
        nascimento = match_data.group(0)

    nome = extrair_por_rotulo(texto_unico, [r"nome do paciente", r"paciente", r"nome"])
    prontuario = extrair_por_rotulo(texto_unico, [r"prontuario", r"prontu[aá]rio"])
    endereco = extrair_por_rotulo(texto_unico, [r"endereco", r"endere[cç]o", r"rua", r"logradouro"])
    numero = extrair_por_rotulo(texto_unico, [r"numero", r"n[uú]mero"])
    bairro = extrair_por_rotulo(texto_unico, [r"bairro"])
    cidade = extrair_por_rotulo(texto_unico, [r"cidade"])
    estado = extrair_por_rotulo(texto_unico, [r"estado", r"uf"])
    responsavel = extrair_por_rotulo(texto_unico, [r"responsavel legal", r"respons[aá]vel"])

    cpf_responsavel = ""
    match_cpf_resp = re.search(r"cpf do respons[aá]vel\s*[:\-]?\s*(\d{3}\.?\d{3}\.?\d{3}-?\d{2})", texto_unico, flags=re.IGNORECASE)
    if match_cpf_resp:
        cpf_responsavel = limpar_cpf(match_cpf_resp.group(1))

    if not nome:
        for linha in linhas[:8]:
            texto_linha = normalizar_texto(linha)
            if any(ch.isdigit() for ch in linha):
                continue
            if any(rotulo in texto_linha for rotulo in ["cpf", "telefone", "endereco", "bairro", "cidade", "estado", "responsavel", "prontuario"]):
                continue
            if len(linha.split()) >= 2:
                nome = linha
                break

    endereco_campos = separar_endereco_campos(endereco, numero, bairro, cidade, estado)

    return {
        "nome": nome,
        "prontuario": prontuario,
        "cpf": cpf,
        "data_nascimento": formatar_data_br_valor(nascimento),
        "telefone": telefone,
        "cep": cep,
        "endereco": endereco_campos["endereco"],
        "numero": endereco_campos["numero"],
        "bairro": endereco_campos["bairro"],
        "cidade": endereco_campos["cidade"],
        "estado": endereco_campos["estado"],
        "menor_idade": bool(responsavel.strip()),
        "responsavel": responsavel,
        "cpf_responsavel": cpf_responsavel,
    }


def aplicar_dados_extraidos_paciente(dados):
    mapa = {
        "paciente_nome_input": dados.get("nome", ""),
        "paciente_prontuario_input": dados.get("prontuario", ""),
        "paciente_cpf_input": dados.get("cpf", ""),
        "paciente_nascimento_input": dados.get("data_nascimento", ""),
        "paciente_telefone_input": dados.get("telefone", ""),
        "paciente_cep_input": dados.get("cep", ""),
        "paciente_endereco_input": dados.get("endereco", ""),
        "paciente_numero_input": dados.get("numero", ""),
        "paciente_bairro_input": dados.get("bairro", ""),
        "paciente_cidade_input": dados.get("cidade", ""),
        "paciente_estado_input": dados.get("estado", ""),
        "paciente_menor_input": dados.get("menor_idade", False),
        "paciente_responsavel_input": dados.get("responsavel", ""),
        "paciente_cpf_responsavel_input": dados.get("cpf_responsavel", ""),
    }
    for chave, valor in mapa.items():
        st.session_state[chave] = valor


def limpar_nome(nome):
    nome = unicodedata.normalize("NFKD", nome or "")
    nome = nome.encode("ASCII", "ignore").decode("ASCII")
    nome = nome.replace(" ", "_")
    return nome.upper()


def normalizar_texto(valor):
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    texto = texto.encode("ASCII", "ignore").decode("ASCII")
    return texto.strip().lower()


def normalizar_forma_pagamento(valor):
    valor_normalizado = normalizar_texto(valor)
    mapa = {
        "pix": "Pix",
        "dinheiro": "Dinheiro",
        "debito": "Debito",
        "credito": "Credito",
        "boleto": "Boleto",
    }
    return mapa.get(valor_normalizado, FORMAS_PAGAMENTO[0])


def limpar_cpf(cpf):
    return "".join(ch for ch in str(cpf or "") if ch.isdigit())


def cpf_valido(cpf):
    cpf_limpo = limpar_cpf(cpf)
    if not cpf_limpo:
        return True
    if len(cpf_limpo) != 11:
        return False
    if cpf_limpo == cpf_limpo[0] * 11:
        return False

    soma = sum(int(cpf_limpo[i]) * (10 - i) for i in range(9))
    digito_1 = (soma * 10 % 11) % 10
    if digito_1 != int(cpf_limpo[9]):
        return False

    soma = sum(int(cpf_limpo[i]) * (11 - i) for i in range(10))
    digito_2 = (soma * 10 % 11) % 10
    return digito_2 == int(cpf_limpo[10])


def conectar_banco():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


conn = conectar_banco()
cursor = conn.cursor()


def tabela_existe(nome_tabela):
    row = cursor.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",
        (nome_tabela,),
    ).fetchone()
    return row is not None


def colunas_tabela(nome_tabela):
    if not tabela_existe(nome_tabela):
        return set()
    return {row["name"] for row in conn.execute(f"PRAGMA table_info({nome_tabela})")}


def gerar_hash_senha(senha, salt=None):
    salt_bytes = salt or secrets.token_bytes(16)
    hash_bytes = hashlib.pbkdf2_hmac("sha256", str(senha or "").encode("utf-8"), salt_bytes, 100000)
    return f"{salt_bytes.hex()}${hash_bytes.hex()}"


def verificar_senha(senha, senha_hash):
    if not senha_hash or "$" not in str(senha_hash):
        return False
    salt_hex, hash_hex = str(senha_hash).split("$", 1)
    try:
        salt_bytes = bytes.fromhex(salt_hex)
    except ValueError:
        return False
    hash_calculado = gerar_hash_senha(senha, salt_bytes).split("$", 1)[1]
    return hmac.compare_digest(hash_calculado, hash_hex)


def garantir_usuario_admin_inicial():
    usuario_existente = cursor.execute(
        "SELECT id FROM usuarios WHERE ativo=1 ORDER BY id LIMIT 1"
    ).fetchone()
    if usuario_existente:
        return
    cursor.execute(
        """
        INSERT INTO usuarios
        (nome, usuario, senha_hash, perfil, ativo, acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios, data_criacao)
        VALUES (?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?)
        """,
        (
            "Administrador",
            USUARIO_ADMIN_INICIAL,
            gerar_hash_senha(SENHA_ADMIN_INICIAL),
            "Administrador",
            1,
            1,
            1,
            1,
            1,
            agora_str(),
        ),
    )
    conn.commit()


def autenticar_usuario(usuario, senha):
    row = cursor.execute(
        """
        SELECT * FROM usuarios
        WHERE lower(usuario)=lower(?) AND ativo=1
        LIMIT 1
        """,
        (str(usuario or "").strip(),),
    ).fetchone()
    if not row:
        return None
    if not verificar_senha(senha, row["senha_hash"]):
        return None
    return row


def carregar_usuarios():
    return pd.read_sql(
        """
        SELECT id, nome, usuario, perfil, ativo, data_criacao,
               acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios
        FROM usuarios
        ORDER BY nome, usuario
        """,
        conn,
    )


def usuario_existe(usuario, ignorar_id=None):
    if ignorar_id is None:
        row = cursor.execute("SELECT id FROM usuarios WHERE lower(usuario)=lower(?) LIMIT 1", (str(usuario or "").strip(),)).fetchone()
    else:
        row = cursor.execute(
            "SELECT id FROM usuarios WHERE lower(usuario)=lower(?) AND id<>? LIMIT 1",
            (str(usuario or "").strip(), int(ignorar_id)),
        ).fetchone()
    return row is not None


def criar_usuario(nome, usuario, senha, perfil="Usuario"):
    acesso_dashboard = 1
    acesso_pacientes = 1
    acesso_contratos = 1
    acesso_financeiro = 1 if perfil == "Administrador" else 0
    acesso_usuarios = 1 if perfil == "Administrador" else 0
    cursor.execute(
        """
        INSERT INTO usuarios
        (nome, usuario, senha_hash, perfil, ativo, acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios, data_criacao)
        VALUES (?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?)
        """,
        (
            nome.strip(),
            usuario.strip(),
            gerar_hash_senha(senha),
            perfil.strip() or "Usuario",
            acesso_dashboard,
            acesso_pacientes,
            acesso_contratos,
            acesso_financeiro,
            acesso_usuarios,
            agora_str(),
        ),
    )


def atualizar_usuario_admin(usuario_id, nome, usuario, perfil, ativo, acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios):
    cursor.execute(
        """
        UPDATE usuarios
        SET nome=?, usuario=?, perfil=?, ativo=?, acesso_dashboard=?, acesso_pacientes=?, acesso_contratos=?, acesso_financeiro=?, acesso_usuarios=?
        WHERE id=?
        """,
        (
            nome.strip(),
            usuario.strip(),
            perfil.strip() or "Usuario",
            int(bool(ativo)),
            int(bool(acesso_dashboard)),
            int(bool(acesso_pacientes)),
            int(bool(acesso_contratos)),
            int(bool(acesso_financeiro)),
            int(bool(acesso_usuarios)),
            int(usuario_id),
        ),
    )


def redefinir_senha_usuario(usuario_id, nova_senha):
    cursor.execute(
        "UPDATE usuarios SET senha_hash=? WHERE id=?",
        (gerar_hash_senha(nova_senha), int(usuario_id)),
    )


def registrar_log_acesso(usuario_id, usuario_login, evento):
    cursor.execute(
        """
        INSERT INTO logs_acesso
        (usuario_id, usuario, evento, data_hora)
        VALUES (?, ?, ?, ?)
        """,
        (int(usuario_id), str(usuario_login or "").strip(), str(evento or "").strip(), agora_str()),
    )
    conn.commit()


def usuario_tem_acesso(usuario_sessao, menu_nome):
    if not usuario_sessao:
        return False
    if usuario_sessao.get("perfil") == "Administrador":
        return True
    chave = MODULOS_SISTEMA.get(menu_nome)
    if not chave:
        return False
    return bool(usuario_sessao.get(chave, 0))


def renderizar_login():
    st.title("Acesso ao Sistema")
    st.caption("Entre com usuario e senha para acessar o sistema.")

    with st.form("login_form"):
        usuario = st.text_input("Usuario")
        senha = st.text_input("Senha", type="password")
        entrar = st.form_submit_button("Entrar")

    if entrar:
        usuario_row = autenticar_usuario(usuario, senha)
        if usuario_row is None:
            st.error("Usuario ou senha invalidos.")
        else:
            st.session_state["usuario_logado"] = {
                "id": int(usuario_row["id"]),
                "nome": usuario_row["nome"],
                "usuario": usuario_row["usuario"],
                "perfil": usuario_row["perfil"],
                "acesso_dashboard": int(usuario_row["acesso_dashboard"] or 0),
                "acesso_pacientes": int(usuario_row["acesso_pacientes"] or 0),
                "acesso_contratos": int(usuario_row["acesso_contratos"] or 0),
                "acesso_financeiro": int(usuario_row["acesso_financeiro"] or 0),
                "acesso_usuarios": int(usuario_row["acesso_usuarios"] or 0),
            }
            registrar_log_acesso(usuario_row["id"], usuario_row["usuario"], "LOGIN")
            st.rerun()

def garantir_coluna(nome_tabela, definicao_coluna):
    nome_coluna = definicao_coluna.split()[0]
    if nome_coluna not in colunas_tabela(nome_tabela):
        cursor.execute(f"ALTER TABLE {nome_tabela} ADD COLUMN {definicao_coluna}")


def inicializar_banco():
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS pacientes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            prontuario TEXT,
            cpf TEXT,
            data_nascimento TEXT,
            telefone TEXT,
            cep TEXT,
            endereco TEXT,
            numero TEXT,
            bairro TEXT,
            cidade TEXT,
            estado TEXT,
            menor_idade INTEGER DEFAULT 0,
            responsavel TEXT,
            cpf_responsavel TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS contratos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            paciente_id INTEGER,
            valor_total REAL DEFAULT 0,
            entrada REAL DEFAULT 0,
            parcelas INTEGER DEFAULT 1,
            primeiro_vencimento TEXT,
            data_pagamento_entrada TEXT,
            forma_pagamento TEXT,
            hash_importacao TEXT,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS procedimentos_contrato (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contrato_id INTEGER,
            procedimento TEXT,
            valor REAL DEFAULT 0
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS financeiro (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            origem TEXT,
            descricao TEXT,
            valor REAL DEFAULT 0,
            data TEXT,
            tipo TEXT,
            contrato_id INTEGER,
            recebivel_id INTEGER,
            prontuario TEXT,
            forma_pagamento TEXT,
            conta_caixa TEXT,
            observacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS saldos_conta (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data TEXT,
            conta TEXT,
            saldo REAL DEFAULT 0,
            observacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS recebiveis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contrato_id INTEGER,
            paciente_id INTEGER,
            paciente_nome TEXT,
            prontuario TEXT,
            parcela_numero INTEGER,
            vencimento TEXT,
            valor REAL DEFAULT 0,
            forma_pagamento TEXT,
            status TEXT DEFAULT 'Aberto',
            observacao TEXT,
            data_criacao TEXT,
            hash_importacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            usuario TEXT,
            senha_hash TEXT,
            perfil TEXT DEFAULT 'Administrador',
            ativo INTEGER DEFAULT 1,
            acesso_dashboard INTEGER DEFAULT 1,
            acesso_pacientes INTEGER DEFAULT 1,
            acesso_contratos INTEGER DEFAULT 1,
            acesso_financeiro INTEGER DEFAULT 1,
            acesso_usuarios INTEGER DEFAULT 0,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS contas_pagar (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data_vencimento TEXT,
            descricao TEXT,
            fornecedor TEXT,
            valor REAL DEFAULT 0,
            pago TEXT,
            valor_pago REAL DEFAULT 0,
            status TEXT DEFAULT 'A vencer',
            observacao TEXT,
            data_criacao TEXT,
            hash_importacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS vendas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data_venda TEXT,
            paciente_nome TEXT,
            valor_total REAL DEFAULT 0,
            valor_a_vista REAL DEFAULT 0,
            valor_cartao REAL DEFAULT 0,
            valor_boleto REAL DEFAULT 0,
            saldo REAL DEFAULT 0,
            data_a_pagar TEXT,
            avaliador TEXT,
            vendedor TEXT,
            nf TEXT,
            contrato_id INTEGER,
            hash_importacao TEXT,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS metas_vendas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ano INTEGER,
            meta REAL DEFAULT 100000,
            supermeta REAL DEFAULT 150000,
            hipermeta REAL DEFAULT 200000,
            data_atualizacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS logs_acesso (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER,
            usuario TEXT,
            evento TEXT,
            data_hora TEXT
        )
        """
    )

    garantir_coluna("financeiro", "origem TEXT")
    garantir_coluna("financeiro", "contrato_id INTEGER")
    garantir_coluna("financeiro", "recebivel_id INTEGER")
    garantir_coluna("financeiro", "prontuario TEXT")
    garantir_coluna("financeiro", "forma_pagamento TEXT")
    garantir_coluna("financeiro", "conta_caixa TEXT")
    garantir_coluna("financeiro", "observacao TEXT")
    garantir_coluna("saldos_conta", "data TEXT")
    garantir_coluna("saldos_conta", "conta TEXT")
    garantir_coluna("saldos_conta", "saldo REAL DEFAULT 0")
    garantir_coluna("saldos_conta", "observacao TEXT")
    garantir_coluna("recebiveis", "paciente_id INTEGER")
    garantir_coluna("recebiveis", "paciente_nome TEXT")
    garantir_coluna("recebiveis", "prontuario TEXT")
    garantir_coluna("recebiveis", "parcela_numero INTEGER")
    garantir_coluna("recebiveis", "vencimento TEXT")
    garantir_coluna("recebiveis", "valor REAL DEFAULT 0")
    garantir_coluna("recebiveis", "forma_pagamento TEXT")
    garantir_coluna("recebiveis", "status TEXT DEFAULT 'Aberto'")
    garantir_coluna("recebiveis", "observacao TEXT")
    garantir_coluna("recebiveis", "data_criacao TEXT")
    garantir_coluna("recebiveis", "data_pagamento TEXT")
    garantir_coluna("recebiveis", "hash_importacao TEXT")
    garantir_coluna("contas_pagar", "data_vencimento TEXT")
    garantir_coluna("contas_pagar", "descricao TEXT")
    garantir_coluna("contas_pagar", "fornecedor TEXT")
    garantir_coluna("contas_pagar", "valor REAL DEFAULT 0")
    garantir_coluna("contas_pagar", "pago TEXT")
    garantir_coluna("contas_pagar", "valor_pago REAL DEFAULT 0")
    garantir_coluna("contas_pagar", "status TEXT DEFAULT 'A vencer'")
    garantir_coluna("contas_pagar", "observacao TEXT")
    garantir_coluna("contas_pagar", "data_criacao TEXT")
    garantir_coluna("contas_pagar", "hash_importacao TEXT")
    garantir_coluna("vendas", "data_venda TEXT")
    garantir_coluna("vendas", "paciente_nome TEXT")
    garantir_coluna("vendas", "valor_total REAL DEFAULT 0")
    garantir_coluna("vendas", "valor_a_vista REAL DEFAULT 0")
    garantir_coluna("vendas", "valor_cartao REAL DEFAULT 0")
    garantir_coluna("vendas", "valor_boleto REAL DEFAULT 0")
    garantir_coluna("vendas", "saldo REAL DEFAULT 0")
    garantir_coluna("vendas", "data_a_pagar TEXT")
    garantir_coluna("vendas", "avaliador TEXT")
    garantir_coluna("vendas", "vendedor TEXT")
    garantir_coluna("vendas", "nf TEXT")
    garantir_coluna("vendas", "contrato_id INTEGER")
    garantir_coluna("vendas", "hash_importacao TEXT")
    garantir_coluna("vendas", "data_criacao TEXT")
    garantir_coluna("metas_vendas", "ano INTEGER")
    garantir_coluna("metas_vendas", "meta REAL DEFAULT 100000")
    garantir_coluna("metas_vendas", "supermeta REAL DEFAULT 150000")
    garantir_coluna("metas_vendas", "hipermeta REAL DEFAULT 200000")
    garantir_coluna("metas_vendas", "data_atualizacao TEXT")
    garantir_coluna("usuarios", "nome TEXT")
    garantir_coluna("usuarios", "usuario TEXT")
    garantir_coluna("usuarios", "senha_hash TEXT")
    garantir_coluna("usuarios", "perfil TEXT DEFAULT 'Administrador'")
    garantir_coluna("usuarios", "ativo INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_dashboard INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_pacientes INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_contratos INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_financeiro INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_usuarios INTEGER DEFAULT 0")
    garantir_coluna("usuarios", "data_criacao TEXT")
    garantir_coluna("logs_acesso", "usuario_id INTEGER")
    garantir_coluna("logs_acesso", "usuario TEXT")
    garantir_coluna("logs_acesso", "evento TEXT")
    garantir_coluna("logs_acesso", "data_hora TEXT")
    garantir_coluna("contratos", "data_criacao TEXT")
    garantir_coluna("contratos", "data_pagamento_entrada TEXT")
    garantir_coluna("contratos", "hash_importacao TEXT")
    garantir_coluna("pacientes", "numero TEXT")
    garantir_coluna("pacientes", "bairro TEXT")
    garantir_coluna("pacientes", "estado TEXT")
    garantir_coluna("pacientes", "responsavel TEXT")
    garantir_coluna("pacientes", "cpf_responsavel TEXT")

    conn.commit()
    garantir_usuario_admin_inicial()


def garantir_pasta_documentos():
    if not os.path.exists(DOCS_DIR):
        os.makedirs(DOCS_DIR)


def proximo_arquivo_documento(nome_base):
    arquivo_inicial = os.path.join(DOCS_DIR, f"{nome_base}.docx")
    if not os.path.exists(arquivo_inicial):
        return arquivo_inicial

    versao = 2
    while True:
        candidato = os.path.join(DOCS_DIR, f"{nome_base}_v{versao}.docx")
        if not os.path.exists(candidato):
            return candidato
        versao += 1


def substituir_runs(doc, dados):
    for p in doc.paragraphs:
        substituir_placeholders_paragrafo(p, dados)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    substituir_placeholders_paragrafo(p, dados)


def substituir_placeholders_paragrafo(paragrafo, dados):
    for run in paragrafo.runs:
        for chave, valor in dados.items():
            if chave in run.text:
                run.text = run.text.replace(chave, str(valor))

    texto = paragrafo.text
    texto_atualizado = texto
    for chave, valor in dados.items():
        texto_atualizado = texto_atualizado.replace(chave, str(valor))

    if texto_atualizado != texto:
        paragrafo.text = texto_atualizado


def valor_paciente(paciente, coluna):
    valor = paciente[coluna].iloc[0] if coluna in paciente.columns else ""
    if pd.isna(valor):
        return ""
    return str(valor).strip()


def montar_endereco_paciente(paciente):
    endereco = valor_paciente(paciente, "endereco")
    numero = valor_paciente(paciente, "numero")
    bairro = valor_paciente(paciente, "bairro")
    cidade = valor_paciente(paciente, "cidade")
    estado = valor_paciente(paciente, "estado")
    cep = valor_paciente(paciente, "cep")

    partes = []
    if endereco:
        if numero:
            partes.append(f"{endereco}, {numero}")
        else:
            partes.append(endereco)
    elif numero:
        partes.append(numero)

    for item in [bairro, cidade, estado]:
        if item:
            partes.append(item)

    endereco_completo = " - ".join(partes)
    if cep:
        endereco_completo = f"{endereco_completo} - CEP {cep}" if endereco_completo else f"CEP {cep}"
    return endereco_completo


def montar_qualificacao(paciente):
    nome = valor_paciente(paciente, "nome")
    cpf = valor_paciente(paciente, "cpf")
    telefone = valor_paciente(paciente, "telefone")
    nascimento = valor_paciente(paciente, "data_nascimento")
    endereco = montar_endereco_paciente(paciente)
    menor_idade = valor_paciente(paciente, "menor_idade") in {"1", "True", "true"}
    responsavel = valor_paciente(paciente, "responsavel")
    cpf_responsavel = valor_paciente(paciente, "cpf_responsavel")

    if menor_idade and responsavel:
        partes = [responsavel]
        if cpf_responsavel:
            partes.append(f"CPF {cpf_responsavel}")
        if telefone:
            partes.append(f"telefone {telefone}")
        if endereco:
            partes.append(f"endereco {endereco}")
        if nome:
            partes.append(f"responsavel legal pelo(a) paciente {nome}")
        if cpf:
            partes.append(f"CPF do paciente {cpf}")
        if nascimento:
            partes.append(f"nascido(a) em {nascimento}")
        return ", ".join(partes)

    partes = [nome]
    if cpf:
        partes.append(f"CPF {cpf}")
    if telefone:
        partes.append(f"telefone {telefone}")
    if nascimento:
        partes.append(f"nascido(a) em {nascimento}")
    if endereco:
        partes.append(f"endereco {endereco}")
    return ", ".join([parte for parte in partes if parte])


def dados_assinatura(paciente):
    nome = valor_paciente(paciente, "nome")
    cpf = valor_paciente(paciente, "cpf")
    menor_idade = valor_paciente(paciente, "menor_idade") in {"1", "True", "true"}
    responsavel = valor_paciente(paciente, "responsavel")
    cpf_responsavel = valor_paciente(paciente, "cpf_responsavel")

    if menor_idade and responsavel:
        return {
            "titulo_assinatura": "",
            "nome_assinatura": responsavel,
            "cpf_assinatura": f"CPF do responsavel: {cpf_responsavel}" if cpf_responsavel else "CPF do responsavel:",
            "assinatura_menor": True,
            "cpf_original": cpf,
        }

    return {
        "titulo_assinatura": "",
        "nome_assinatura": nome,
        "cpf_assinatura": f"CPF: {cpf}" if cpf else "CPF:",
        "assinatura_menor": False,
        "cpf_original": cpf,
    }


def montar_termo_cirurgia(procedimentos):
    itens = []
    for proc in procedimentos["procedimento"].tolist():
        texto = str(proc).strip()
        if not texto:
            continue
        itens.append(texto)

    if not itens:
        return ""
    if len(itens) == 1:
        return itens[0]
    return "; ".join(itens)


def ajustar_assinatura_capa(doc, nome_paciente, cpf_paciente, nome_assinatura, cpf_assinatura):
    paragrafos = doc.paragraphs
    titulo_contrato = "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO"
    limite = len(paragrafos)

    for indice, paragrafo in enumerate(paragrafos):
        if titulo_contrato in paragrafo.text:
            limite = indice
            break

    for indice in range(limite - 1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip() if indice + 1 < limite else ""

        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            paragrafos[indice].text = nome_assinatura
            paragrafos[indice + 1].text = cpf_assinatura
            return

    for indice in range(limite):
        paragrafo = paragrafos[indice]
        if "___________________________________________________________" not in paragrafo.text:
            continue
        if indice + 2 >= limite:
            continue
        proximo = paragrafos[indice + 1]
        seguinte = paragrafos[indice + 2]
        proximo.text = nome_assinatura
        seguinte.text = cpf_assinatura
        return


def inserir_paragrafo_apos(paragrafo, texto=""):
    novo_p = OxmlElement("w:p")
    paragrafo._p.addnext(novo_p)
    novo_paragrafo = Paragraph(novo_p, paragrafo._parent)
    novo_paragrafo.text = texto
    novo_paragrafo.alignment = paragrafo.alignment
    if paragrafo.style is not None:
        novo_paragrafo.style = paragrafo.style
    return novo_paragrafo


def substituir_par_assinatura(paragrafo_nome, paragrafo_cpf, assinatura):
    if assinatura["assinatura_menor"]:
        paragrafo_nome.text = assinatura["nome_assinatura"]
        paragrafo_cpf.text = assinatura["cpf_assinatura"]
        return

    paragrafo_nome.text = assinatura["nome_assinatura"]
    paragrafo_cpf.text = assinatura["cpf_assinatura"]


def ajustar_assinaturas_menor(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip()

        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            substituir_par_assinatura(paragrafos[indice], paragrafos[indice + 1], assinatura)


def ajustar_assinaturas_por_texto(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    for paragrafo in paragrafos:
        texto = paragrafo.text
        if not texto:
            continue

        texto_ajustado = texto
        if nome_paciente and nome_paciente in texto_ajustado:
            texto_ajustado = texto_ajustado.replace(nome_paciente, assinatura["nome_assinatura"])

        if cpf_paciente:
            texto_ajustado = texto_ajustado.replace(f"CPF: {cpf_paciente}", assinatura["cpf_assinatura"])
            texto_ajustado = texto_ajustado.replace(f"CPF {cpf_paciente}", assinatura["cpf_assinatura"])
            texto_ajustado = texto_ajustado.replace(cpf_paciente, assinatura["cpf_original"] if not assinatura["assinatura_menor"] else cpf_paciente)

        if texto_ajustado != texto:
            if assinatura["assinatura_menor"]:
                texto_ajustado = texto_ajustado.replace(cpf_paciente, assinatura["cpf_assinatura"].replace("CPF do responsavel: ", "").strip())
            paragrafo.text = texto_ajustado


def ajustar_ultima_assinatura_menor(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 2, -1, -1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip()

        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            substituir_par_assinatura(paragrafos[indice], paragrafos[indice + 1], assinatura)
            return


def ajustar_bloco_final_assinatura(doc, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, 1, -1):
        if "SOUL SUL CLÍNICA INTEGRADA" not in paragrafos[indice].text:
            continue

        indice_linha = None
        for cursor in range(indice - 1, -1, -1):
            texto = paragrafos[cursor].text.strip()
            if "____" in texto or "___" in texto:
                indice_linha = cursor
                break

        if indice_linha is None:
            return

        candidatos = []
        for cursor in range(indice_linha + 1, indice):
            texto = paragrafos[cursor].text.strip()
            if texto:
                candidatos.append(cursor)

        if len(candidatos) < 2:
            return

        indice_nome = candidatos[0]
        indice_cpf = candidatos[1]

        paragrafos[indice_nome].text = assinatura["nome_assinatura"]
        paragrafos[indice_cpf].text = assinatura["cpf_assinatura"]
        paragrafos[indice_nome].alignment = 1
        paragrafos[indice_cpf].alignment = 1

        for cursor in candidatos[2:]:
            paragrafos[cursor].text = ""
        return


def ajustar_assinatura_termo_cirurgico_final(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    ultimo_indice_data = None
    ultimo_indice_clinica = None

    for indice, paragrafo in enumerate(paragrafos):
        texto = paragrafo.text.strip()
        if "Campos dos Goytacazes" in texto:
            ultimo_indice_data = indice
        if "SOUL SUL CLÍNICA INTEGRADA" in texto:
            ultimo_indice_clinica = indice

    if ultimo_indice_data is None or ultimo_indice_clinica is None:
        return
    if ultimo_indice_data >= ultimo_indice_clinica:
        return

    candidatos = []
    for indice in range(ultimo_indice_data + 1, ultimo_indice_clinica):
        texto = paragrafos[indice].text.strip()
        if texto:
            candidatos.append(indice)

    if len(candidatos) < 2:
        return

    for posicao in range(len(candidatos) - 1):
        indice_nome = candidatos[posicao]
        indice_cpf = candidatos[posicao + 1]
        texto_nome = paragrafos[indice_nome].text.strip()
        texto_cpf = paragrafos[indice_cpf].text.strip()

        if texto_nome == nome_paciente and cpf_paciente and cpf_paciente in texto_cpf:
            substituir_par_assinatura(paragrafos[indice_nome], paragrafos[indice_cpf], assinatura)
            return


def ajustar_assinatura_contrato(doc, nome_assinatura, cpf_assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, 1, -1):
        if "SOUL SUL CLÍNICA INTEGRADA" not in paragrafos[indice].text:
            continue

        candidatos = []
        cursor = indice - 1
        while cursor >= 0 and len(candidatos) < 2:
            texto = paragrafos[cursor].text.strip()
            if texto:
                candidatos.append(cursor)
            cursor -= 1

        if len(candidatos) < 2:
            return

        indice_cpf = candidatos[0]
        indice_nome = candidatos[1]
        paragrafos[indice_nome].text = nome_assinatura
        paragrafos[indice_cpf].text = cpf_assinatura
        return


def reescrever_bloco_final_termo(doc, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, -1, -1):
        texto = paragrafos[indice].text
        if "SOUL SUL" not in texto or "TESTEMUNHAS" not in texto:
            continue

        candidatos_limpar = []
        for cursor in range(indice - 1, -1, -1):
            texto_cursor = paragrafos[cursor].text.strip()
            if not texto_cursor:
                continue
            candidatos_limpar.append(cursor)
            if len(candidatos_limpar) == 2:
                break

        for cursor in candidatos_limpar:
            paragrafos[cursor].text = ""

        paragrafos[indice].text = (
            "\n"
            "___________________________________________________________\n"
            f"{assinatura['nome_assinatura']}\n"
            f"{assinatura['cpf_assinatura']}\n\n"
            "\n___________________________________________________________\n"
            "SOUL SUL CLINICA INTEGRADA\n\n"
            "TESTEMUNHAS:\n\n"
            "________________________________                    _________________________________\n"
            "NOME:                                                          NOME:\n"
            "CPF:                                                             CPF:"
        )
        paragrafos[indice].alignment = 1
        return


def remover_elemento(elemento):
    parent = elemento.getparent()
    if parent is not None:
        parent.remove(elemento)


def texto_elemento(elemento):
    return "".join(elemento.itertext())


def remover_secao_termo_cirurgico(doc):
    corpo = doc._element.body
    removendo = False
    elementos_para_remover = []

    for elemento in corpo:
        texto = texto_elemento(elemento)
        if "{TERMO_CIRURGIA}" in texto:
            removendo = True

        if removendo:
            if "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO" in texto:
                removendo = False
                continue
            elementos_para_remover.append(elemento)

    for elemento in elementos_para_remover:
        remover_elemento(elemento)

    for elemento in list(corpo):
        texto = texto_elemento(elemento).strip()
        xml = etree.tostring(elemento, encoding="unicode")
        if "lastRenderedPageBreak" in xml and not texto:
            remover_elemento(elemento)


def normalizar_quebra_antes_contrato(doc):
    titulo_contrato = "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO"
    paragrafos = doc.paragraphs
    indice_titulo = None

    for indice, paragrafo in enumerate(paragrafos):
        if titulo_contrato in paragrafo.text:
            indice_titulo = indice
            break

    if indice_titulo is None or indice_titulo == 0:
        return

    for indice in range(indice_titulo):
        paragrafo = paragrafos[indice]
        for run in paragrafo.runs:
            elementos_remover = []
            for child in run._element:
                if child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
                    elementos_remover.append(child)
            for child in elementos_remover:
                run._element.remove(child)

    paragrafo_anterior = paragrafos[indice_titulo - 1]
    if not any(
        child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
        for run in paragrafo_anterior.runs
        for child in run._element
    ):
        paragrafo_anterior.add_run().add_break(WD_BREAK.PAGE)


def normalizar_quebra_antes_termo(doc):
    marcadores_termo = [
        "TERMO DE CONSENTIMENTO ESCLARECIDO",
        "CONSENTIMENTO ESCLARECIDO",
        "DO CONSENTIMENTO ESCLARECIDO",
    ]
    paragrafos = doc.paragraphs
    indice_titulo = None

    for indice, paragrafo in enumerate(paragrafos):
        texto_normalizado = normalizar_texto(paragrafo.text).upper()
        if any(normalizar_texto(marcador).upper() in texto_normalizado for marcador in marcadores_termo):
            indice_titulo = indice
            break

    if indice_titulo is None or indice_titulo == 0:
        return

    paragrafo_anterior = paragrafos[indice_titulo - 1]
    if not any(
        child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
        for run in paragrafo_anterior.runs
        for child in run._element
    ):
        paragrafo_anterior.add_run().add_break(WD_BREAK.PAGE)


def garantir_logo_no_cabecalho(doc):
    if not doc.sections:
        return

    doc.sections[0].different_first_page_header_footer = False
    for secao in doc.sections[1:]:
        secao.different_first_page_header_footer = False
        secao.header.is_linked_to_previous = True


def aplicar_fonte_times_new_roman(doc):
    estilos = doc.styles
    for nome_estilo in ["Normal", "Header", "Footer"]:
        if nome_estilo in estilos:
            estilo = estilos[nome_estilo]
            estilo.font.name = "Times New Roman"
            estilo._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            estilo._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            estilo._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    def aplicar_paragrafo(paragrafo):
        for run in paragrafo.runs:
            run.font.name = "Times New Roman"
            if run._element.rPr is None:
                run._element.get_or_add_rPr()
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    for paragrafo in doc.paragraphs:
        aplicar_paragrafo(paragrafo)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    aplicar_paragrafo(paragrafo)

    for secao in doc.sections:
        for paragrafo in secao.header.paragraphs:
            aplicar_paragrafo(paragrafo)
        for tabela in secao.header.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        aplicar_paragrafo(paragrafo)
        for paragrafo in secao.footer.paragraphs:
            aplicar_paragrafo(paragrafo)


def compactar_bloco_final(doc):
    paragrafos = doc.paragraphs
    indice_data = None
    indice_testemunhas = None

    for indice in range(len(paragrafos) - 1, -1, -1):
        texto = paragrafos[indice].text.strip()
        if indice_testemunhas is None and "TESTEMUNHAS" in texto:
            indice_testemunhas = indice
        if indice_data is None and "Campos dos Goytacazes" in texto:
            indice_data = indice
        if indice_data is not None and indice_testemunhas is not None:
            break

    if indice_data is None or indice_testemunhas is None or indice_data >= indice_testemunhas:
        return

    for indice in range(indice_data, len(paragrafos)):
        texto = paragrafos[indice].text.strip()
        if not texto:
            continue

        formato = paragrafos[indice].paragraph_format
        formato.space_before = Pt(0)
        formato.space_after = Pt(0)
        formato.line_spacing = 1.0

        for run in paragrafos[indice].runs:
            run.font.size = Pt(10)

        if "TESTEMUNHAS" in texto:
            for cursor in range(indice + 1, min(indice + 6, len(paragrafos))):
                formato_cursor = paragrafos[cursor].paragraph_format
                formato_cursor.space_before = Pt(0)
                formato_cursor.space_after = Pt(0)
                formato_cursor.line_spacing = 1.0
                for run in paragrafos[cursor].runs:
                    run.font.size = Pt(10)
            break


def preencher_tabela_checklist(tabela, procedimentos):
    for row_index, row in enumerate(tabela.rows):
        for cell in row.cells:
            if "{PROCEDIMENTO}" not in cell.text:
                continue

            if len(tabela.columns) == 1:
                tabela.add_column(Inches(1.0))

            cabecalho = tabela.rows[row_index]
            cabecalho.cells[0].text = "Procedimento"
            cabecalho.cells[1].text = "OK"

            total_procedimentos = len(procedimentos.index)
            while len(tabela.rows) < row_index + 1 + total_procedimentos:
                tabela.add_row()

            for indice_proc, proc_row in enumerate(procedimentos.itertuples(index=False), start=1):
                linha = tabela.rows[row_index + indice_proc]
                linha.cells[0].text = str(proc_row.procedimento)
                linha.cells[1].text = "[   ]"
            return True

    return False


def preencher_tabela_contrato(tabela, procedimentos):
    for row_index, row in enumerate(tabela.rows):
        if len(row.cells) < 2:
            continue

        tem_procedimento = any("{PROCEDIMENTO}" in cell.text for cell in row.cells)
        tem_valor = any("{VALOR}" in cell.text for cell in row.cells)
        if not (tem_procedimento or tem_valor):
            continue

        total_procedimentos = len(procedimentos.index)
        while len(tabela.rows) < row_index + total_procedimentos:
            tabela.add_row()

        for indice_proc, proc_row in enumerate(procedimentos.itertuples(index=False)):
            linha = tabela.rows[row_index + indice_proc]
            linha.cells[0].text = str(proc_row.procedimento)
            linha.cells[1].text = f"R$ {float(proc_row.valor):.2f}"
        return True

    return False


def gerar_documento(conn_local, contrato_id):
    contrato = pd.read_sql(
        "SELECT * FROM contratos WHERE id=?",
        conn_local,
        params=(contrato_id,),
    )
    if contrato.empty:
        raise ValueError("Contrato nao encontrado para gerar documento.")

    paciente_id = int(contrato["paciente_id"].iloc[0])
    paciente = pd.read_sql(
        "SELECT * FROM pacientes WHERE id=?",
        conn_local,
        params=(paciente_id,),
    )
    procedimentos = pd.read_sql(
        "SELECT * FROM procedimentos_contrato WHERE contrato_id=?",
        conn_local,
        params=(contrato_id,),
    )

    if paciente.empty:
        raise ValueError("Paciente vinculado ao contrato nao encontrado.")

    nome = paciente["nome"].iloc[0]
    prontuario = paciente["prontuario"].iloc[0]
    pagamento = montar_texto_pagamento(contrato.iloc[0])

    endereco_completo = montar_endereco_paciente(paciente)
    qualificacao = montar_qualificacao(paciente)
    termo_cirurgia = montar_termo_cirurgia(procedimentos)
    assinatura = dados_assinatura(paciente)
    doc = Document(TEMPLATE_PATH)
    garantir_logo_no_cabecalho(doc)
    if not termo_cirurgia:
        remover_secao_termo_cirurgico(doc)
    dados = {
        "{PACIENTE}": nome,
        "{paciente}": nome,
        "{CPF}": valor_paciente(paciente, "cpf"),
        "{cpf}": valor_paciente(paciente, "cpf"),
        "{PRONTUARIO}": prontuario,
        "{prontuario}": prontuario,
        "{PAGAMENTO}": pagamento,
        "{pagamento}": pagamento,
        "{QUALIFICACAO}": qualificacao,
        "{qualificacao}": qualificacao,
        "{ENDERECO}": valor_paciente(paciente, "endereco"),
        "{endereco}": valor_paciente(paciente, "endereco"),
        "{NUMERO}": valor_paciente(paciente, "numero"),
        "{numero}": valor_paciente(paciente, "numero"),
        "{BAIRRO}": valor_paciente(paciente, "bairro"),
        "{bairro}": valor_paciente(paciente, "bairro"),
        "{CIDADE}": valor_paciente(paciente, "cidade"),
        "{cidade}": valor_paciente(paciente, "cidade"),
        "{ESTADO}": valor_paciente(paciente, "estado"),
        "{estado}": valor_paciente(paciente, "estado"),
        "{CEP}": valor_paciente(paciente, "cep"),
        "{cep}": valor_paciente(paciente, "cep"),
        "{TELEFONE}": valor_paciente(paciente, "telefone"),
        "{telefone}": valor_paciente(paciente, "telefone"),
        "{DATA_NASCIMENTO}": valor_paciente(paciente, "data_nascimento"),
        "{data_nascimento}": valor_paciente(paciente, "data_nascimento"),
        "{TERMO_CIRURGIA}": "",
        "{termo_cirurgia}": "",
    }
    substituir_runs(doc, dados)

    checklist_preenchido = False
    tabela_contrato_preenchida = False
    for tabela in doc.tables:
        if not checklist_preenchido and preencher_tabela_checklist(tabela, procedimentos):
            checklist_preenchido = True
            continue
        if not tabela_contrato_preenchida and preencher_tabela_contrato(tabela, procedimentos):
            tabela_contrato_preenchida = True

    if assinatura["assinatura_menor"]:
        ajustar_assinaturas_menor(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_assinaturas_por_texto(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_assinatura_termo_cirurgico_final(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_ultima_assinatura_menor(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_bloco_final_assinatura(doc, assinatura)
    else:
        ajustar_assinatura_capa(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura["nome_assinatura"],
            assinatura["cpf_assinatura"],
        )
        ajustar_assinatura_contrato(doc, assinatura["nome_assinatura"], assinatura["cpf_assinatura"])
    reescrever_bloco_final_termo(doc, assinatura)
    normalizar_quebra_antes_contrato(doc)
    normalizar_quebra_antes_termo(doc)
    compactar_bloco_final(doc)
    aplicar_fonte_times_new_roman(doc)

    nome_base = f"{limpar_nome(nome)}_{prontuario}"
    arquivo = proximo_arquivo_documento(nome_base)
    doc.save(arquivo)
    return arquivo


def carregar_pacientes():
    return pd.read_sql("SELECT * FROM pacientes ORDER BY nome", conn)


def carregar_contratos():
    return pd.read_sql("SELECT * FROM contratos ORDER BY id DESC", conn)


def carregar_procedimentos(contrato_id):
    return pd.read_sql(
        "SELECT * FROM procedimentos_contrato WHERE contrato_id=? ORDER BY id",
        conn,
        params=(contrato_id,),
    )


def validar_paciente(nome, prontuario):
    erros = []
    if not nome.strip():
        erros.append("Informe o nome do paciente.")
    if not prontuario.strip():
        erros.append("Informe o prontuario.")
    return erros


def validar_dados_paciente(nome, prontuario, cpf, menor, responsavel, cpf_responsavel):
    erros = validar_paciente(nome, prontuario)
    if cpf and not cpf_valido(cpf):
        erros.append("O CPF do paciente e invalido.")
    if menor and not responsavel.strip():
        erros.append("Informe o responsavel legal para paciente menor de idade.")
    if cpf_responsavel and not cpf_valido(cpf_responsavel):
        erros.append("O CPF do responsavel e invalido.")
    return erros


def validar_contrato(procedimentos, valores, entrada, parcelas, vencimento, forma_pagamento, data_pagamento_entrada):
    erros = []
    if not procedimentos:
        erros.append("Informe ao menos um procedimento no contrato.")
    total = sum(valores)
    if entrada < 0:
        erros.append("A entrada nao pode ser negativa.")
    if entrada > total:
        erros.append("A entrada nao pode ser maior que o valor total.")
    if forma_pagamento_a_vista(forma_pagamento):
        if not parse_data_contrato(data_pagamento_entrada):
            erros.append("Informe a data do pagamento no formato DD/MM/AAAA.")
    elif entrada > 0 and not parse_data_contrato(data_pagamento_entrada):
        erros.append("Informe a data do pagamento da entrada no formato DD/MM/AAAA.")
    if not forma_pagamento_a_vista(forma_pagamento) and total - entrada > 0 and int(parcelas or 1) > 0 and not parse_data_contrato(vencimento):
        erros.append("Informe o primeiro vencimento no formato DD/MM/AAAA.")
    return erros, total


def buscar_endereco_por_cep(cep):
    if not cep.strip():
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}

    cep_limpo = "".join(ch for ch in cep if ch.isdigit())
    if len(cep_limpo) != 8:
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}

    try:
        resposta = requests.get(
            f"https://viacep.com.br/ws/{cep_limpo}/json/",
            timeout=5,
        )
        resposta.raise_for_status()
        dados = resposta.json()
        if dados.get("erro"):
            return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}
        return dados
    except requests.RequestException:
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}


def salvar_procedimentos_contrato(contrato_id, procedimentos, valores):
    cursor.execute("DELETE FROM procedimentos_contrato WHERE contrato_id=?", (contrato_id,))
    for procedimento, valor in zip(procedimentos, valores):
        cursor.execute(
            """
            INSERT INTO procedimentos_contrato
            (contrato_id, procedimento, valor)
            VALUES (?, ?, ?)
            """,
            (contrato_id, procedimento, valor),
        )


def sincronizar_financeiro_contrato(contrato_id, paciente_nome, valor_total):
    registro = cursor.execute(
        """
        SELECT id FROM financeiro
        WHERE contrato_id=? AND tipo='Entrada'
        ORDER BY id DESC
        LIMIT 1
        """,
        (contrato_id,),
    ).fetchone()

    if not registro:
        candidatos = cursor.execute(
            """
            SELECT id FROM financeiro
            WHERE contrato_id IS NULL
              AND tipo='Entrada'
              AND descricao=?
              AND origem=?
            ORDER BY id DESC
            """,
            (DESCRICAO_CONTRATO, paciente_nome),
        ).fetchall()
        if len(candidatos) == 1:
            registro = candidatos[0]

    if registro:
        cursor.execute(
            """
            UPDATE financeiro
            SET origem=?, descricao=?, valor=?, data=?, tipo='Entrada', contrato_id=?
            WHERE id=?
            """,
            (
                paciente_nome,
                DESCRICAO_CONTRATO,
                valor_total,
                agora_str(),
                contrato_id,
                registro["id"],
            ),
        )
        return

    cursor.execute(
        """
        INSERT INTO financeiro
        (origem, descricao, valor, data, tipo, contrato_id)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (paciente_nome, DESCRICAO_CONTRATO, valor_total, agora_str(), "Entrada", contrato_id),
    )


def calcular_valores_parcelas(valor_total, entrada, parcelas):
    restante = round(float(valor_total or 0) - float(entrada or 0), 2)
    quantidade = max(int(parcelas or 1), 1)
    if restante <= 0:
        return []

    valor_base = round(restante / quantidade, 2)
    valores = [valor_base] * quantidade
    diferenca = round(restante - sum(valores), 2)
    if valores:
        valores[-1] = round(valores[-1] + diferenca, 2)
    return valores


def montar_texto_pagamento(contrato_row):
    parcelas = int(contrato_row["parcelas"] or 1)
    forma = contrato_row["forma_pagamento"] or ""
    valor_total = float(contrato_row["valor_total"] or 0)
    entrada = float(contrato_row["entrada"] or 0)
    restante = round(valor_total - entrada, 2)
    valor_parcela = round(restante / parcelas, 2) if parcelas > 0 else restante
    data_pagamento = formatar_data_br_valor(
        contrato_row["data_pagamento_entrada"] or contrato_row["data_criacao"] or ""
    )
    vencimento = formatar_data_br_valor(contrato_row["primeiro_vencimento"] or "")

    if forma_pagamento_a_vista(forma):
        return (
            f"VALOR TOTAL R$ {valor_total:.2f}. "
            f"PAGO NO {str(forma).upper()} NO DIA {data_pagamento}."
        )

    partes = [f"VALOR TOTAL R$ {valor_total:.2f}."]
    if entrada > 0:
        partes.append(f"ENTRADA R$ {entrada:.2f} PAGA EM {data_pagamento}.")
    if restante > 0:
        partes.append(
            f"RESTANTE R$ {restante:.2f} EM {parcelas} PARCELAS DE R$ {valor_parcela:.2f} "
            f"NO {str(forma).upper()} COM PRIMEIRO VENCIMENTO EM {vencimento}."
        )
    return " ".join(partes)


def montar_recebiveis_planejados(contrato_id, paciente_id, paciente_nome, prontuario, valor_total, entrada, parcelas, primeiro_vencimento, forma_pagamento):
    data_inicial = parse_data_contrato(primeiro_vencimento)
    if not data_inicial:
        raise ValueError("Informe o primeiro vencimento no formato DD/MM/AAAA para gerar os recebiveis.")

    valores = calcular_valores_parcelas(valor_total, entrada, parcelas)
    recebiveis = []
    for indice, valor_parcela in enumerate(valores, start=1):
        vencimento = adicionar_meses(data_inicial, indice - 1)
        recebiveis.append(
            {
                "contrato_id": int(contrato_id),
                "paciente_id": int(paciente_id),
                "paciente_nome": paciente_nome,
                "prontuario": prontuario,
                "parcela_numero": indice,
                "vencimento": formatar_data_br(vencimento),
                "valor": round(valor_parcela, 2),
                "forma_pagamento": forma_pagamento,
                "status": "Aberto",
                "observacao": "",
            }
        )
    return recebiveis


def carregar_recebiveis_contrato(contrato_id):
    return pd.read_sql(
        "SELECT * FROM recebiveis WHERE contrato_id=? ORDER BY parcela_numero, id",
        conn,
        params=(contrato_id,),
    )


def atualizar_recebivel(recebivel_id, paciente_nome, prontuario, vencimento, valor, forma_pagamento, status, observacao):
    cursor.execute(
        """
        UPDATE recebiveis
        SET paciente_nome=?, prontuario=?, vencimento=?, valor=?, forma_pagamento=?, status=?, observacao=?
        WHERE id=?
        """,
        (
            paciente_nome.strip(),
            prontuario.strip(),
            vencimento.strip(),
            float(valor),
            forma_pagamento,
            status,
            observacao.strip(),
            int(recebivel_id),
        ),
    )


def atualizar_recebiveis_lote_contrato(contrato_id, paciente_nome, prontuario, forma_pagamento, status, observacao, primeiro_vencimento=""):
    recebiveis = cursor.execute(
        """
        SELECT id, parcela_numero, vencimento
        FROM recebiveis
        WHERE contrato_id=?
        ORDER BY parcela_numero, id
        """,
        (int(contrato_id),),
    ).fetchall()

    if not recebiveis:
        return

    data_base = parse_data_contrato(primeiro_vencimento) if primeiro_vencimento.strip() else None
    for row in recebiveis:
        vencimento = row["vencimento"]
        if data_base:
            vencimento = formatar_data_br(adicionar_meses(data_base, int(row["parcela_numero"] or 1) - 1))

        cursor.execute(
            """
            UPDATE recebiveis
            SET paciente_nome=?, prontuario=?, vencimento=?, forma_pagamento=?, status=?, observacao=?
            WHERE id=?
            """,
            (
                paciente_nome.strip(),
                prontuario.strip(),
                vencimento,
                forma_pagamento,
                status,
                observacao.strip(),
                int(row["id"]),
            ),
        )


def atualizar_status_contas_pagar_automaticamente():
    hoje = date.today()
    contas = cursor.execute(
        """
        SELECT id, data_vencimento, status
        FROM contas_pagar
        """
    ).fetchall()

    houve_atualizacao = False
    for conta in contas:
        status_atual = str(conta["status"] or "A vencer").strip() or "A vencer"
        if status_atual in {"Pago", "Suspenso", "Cancelado"}:
            continue

        data_vencimento = parse_data_contrato(conta["data_vencimento"])
        if not data_vencimento:
            continue

        novo_status = "Atrasado" if data_vencimento < hoje else "A vencer"
        if status_atual != novo_status:
            cursor.execute(
                "UPDATE contas_pagar SET status=? WHERE id=?",
                (novo_status, int(conta["id"])),
            )
            houve_atualizacao = True

    if houve_atualizacao:
        conn.commit()


def registrar_movimento_caixa(origem, descricao, valor, tipo, data_movimento, prontuario="", forma_pagamento="", conta_caixa="", observacao="", contrato_id=None, recebivel_id=None):
    data_registro = data_movimento.isoformat() if hasattr(data_movimento, "isoformat") else str(data_movimento or agora_str())
    cursor.execute(
        """
        INSERT INTO financeiro
        (origem, descricao, valor, data, tipo, contrato_id, recebivel_id, prontuario, forma_pagamento, conta_caixa, observacao)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            origem.strip(),
            descricao.strip(),
            float(valor),
            data_registro,
            tipo,
            contrato_id,
            recebivel_id,
            prontuario.strip(),
            forma_pagamento.strip(),
            (conta_caixa or "").strip(),
            observacao.strip(),
        ),
    )
    return cursor.lastrowid


def registrar_saldo_conta(data_saldo, conta, saldo, observacao=""):
    data_registro = data_saldo.isoformat() if hasattr(data_saldo, "isoformat") else str(data_saldo or agora_str())
    cursor.execute(
        """
        INSERT INTO saldos_conta
        (data, conta, saldo, observacao)
        VALUES (?, ?, ?, ?)
        """,
        (
            data_registro,
            conta.strip(),
            float(saldo),
            observacao.strip(),
        ),
    )
    return cursor.lastrowid


def baixar_recebivel_no_caixa(recebivel_id, origem, data_pagamento, observacao="", forma_recebimento="", conta_caixa=""):
    recebivel = cursor.execute(
        """
        SELECT *
        FROM recebiveis
        WHERE id=?
        """,
        (int(recebivel_id),),
    ).fetchone()

    if not recebivel:
        raise ValueError("Recebivel nao encontrado.")
    if (recebivel["status"] or "") == "Pago":
        raise ValueError("Este recebivel ja esta pago.")

    forma_baixa = forma_recebimento.strip() or (recebivel["forma_pagamento"] or "")
    descricao = f"Baixa recebivel parcela {int(recebivel['parcela_numero'] or 0)} - {recebivel['paciente_nome']}"
    financeiro_id = registrar_movimento_caixa(
        origem=origem or "Pagamento",
        descricao=descricao,
        valor=float(recebivel["valor"] or 0),
        tipo="Entrada",
        data_movimento=data_pagamento,
        prontuario=str(recebivel["prontuario"] or ""),
        forma_pagamento=forma_baixa,
        conta_caixa=conta_caixa or forma_baixa,
        observacao=observacao,
        contrato_id=recebivel["contrato_id"],
        recebivel_id=recebivel["id"],
    )

    observacao_atual = str(recebivel["observacao"] or "").strip()
    observacao_nova = observacao.strip()
    if observacao_atual and observacao_nova:
        observacao_final = f"{observacao_atual} | {observacao_nova}"
    else:
        observacao_final = observacao_nova or observacao_atual

    cursor.execute(
        """
        UPDATE recebiveis
        SET status='Pago', data_pagamento=?, observacao=?
        WHERE id=?
        """,
        (
            data_pagamento.strftime("%d/%m/%Y") if hasattr(data_pagamento, "strftime") else str(data_pagamento),
            observacao_final,
            int(recebivel_id),
        ),
    )
    return financeiro_id


def dataframe_para_excel_bytes(df, nome_aba="Dados"):
    if not OPENPYXL_DISPONIVEL:
        return None
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=nome_aba[:31])
    buffer.seek(0)
    return buffer.getvalue()


def formatar_moeda_br(valor):
    try:
        valor_float = float(valor or 0)
    except (TypeError, ValueError):
        valor_float = 0.0
    texto = f"{valor_float:,.2f}"
    return f"R$ {texto.replace(',', 'X').replace('.', ',').replace('X', '.')}"


CONTAS_CAIXA_MODELO = ["CAIXA", "SICOOB", "INFINITEPAY", "C6", "SANTANDER", "PAGSEGURO"]


def identificar_conta_caixa(origem, descricao="", observacao=""):
    texto = normalizar_texto(f"{origem} {descricao} {observacao}")
    if "infinitepay" in texto:
        return "INFINITEPAY"
    if "sicoob" in texto:
        return "SICOOB"
    if re.search(r"\bc6\b", texto):
        return "C6"
    if "santander" in texto:
        return "SANTANDER"
    if "pagseguro" in texto:
        return "PAGSEGURO"
    return "CAIXA"


def saldos_informados_por_conta_ate(data_limite):
    saldos = {conta: 0.0 for conta in CONTAS_CAIXA_MODELO}
    registros = pd.read_sql(
        "SELECT * FROM saldos_conta ORDER BY data ASC, id ASC",
        conn,
    )
    if registros.empty:
        return saldos

    registros["data_ref"] = registros["data"].apply(parse_data_contrato)
    for _, row in registros.iterrows():
        data_ref = row["data_ref"]
        conta = str(row["conta"] or "").strip().upper()
        if not data_ref or conta not in saldos:
            continue
        if data_ref <= data_limite:
            saldos[conta] = float(row["saldo"] or 0)
    return saldos


def montar_resumo_saldos_conta(saldos_dict, data_referencia):
    linhas = []
    data_texto = formatar_data_br(data_referencia) if hasattr(data_referencia, "strftime") else formatar_data_br_valor(data_referencia)
    for conta in CONTAS_CAIXA_MODELO:
        linhas.append(
            {
                "Data": data_texto,
                "Descricao": f"SALDO {conta}",
                "Saldo": formatar_moeda_br(saldos_dict.get(conta, 0)),
                "Banco": conta,
            }
        )
    return pd.DataFrame(linhas)


def estilizar_resumo_saldos(df):
    return (
        df.style
        .apply(lambda _: ["background-color: #C6E0B4"] * len(df.columns), axis=1)
        .set_properties(**{"color": "black"})
    )


def montar_caixa_diario(financeiro_df):
    if financeiro_df.empty:
        return pd.DataFrame(), []

    caixa = financeiro_df.copy()
    caixa["data_date"] = caixa["data"].apply(parse_data_contrato)
    caixa["data_grupo"] = caixa["data_date"].apply(lambda valor: formatar_data_br(valor) if valor else "")
    caixa["data_grupo"] = caixa["data_grupo"].replace("", "Data nao informada")
    caixa["ordem_data"] = caixa["data_date"].apply(lambda valor: valor.toordinal() if valor else -1)
    caixa = caixa.sort_values(["ordem_data", "id"], ascending=[False, False])

    grupos = []
    resumo_linhas = []
    for data_grupo, grupo in caixa.groupby("data_grupo", sort=False):
        entradas = float(grupo.loc[grupo["tipo"] == "Entrada", "valor"].sum())
        saidas = float(grupo.loc[grupo["tipo"] == "Saida", "valor"].sum())
        saldo = entradas - saidas

        detalhamento = grupo[
            ["tipo", "origem", "conta_caixa", "prontuario", "descricao", "forma_pagamento", "observacao", "valor"]
        ].copy()
        detalhamento = detalhamento.rename(
            columns={
                "tipo": "Tipo",
                "origem": "Origem",
                "conta_caixa": "Conta/Banco",
                "prontuario": "Prontuario",
                "descricao": "Descricao",
                "forma_pagamento": "Forma de pagamento",
                "observacao": "Observacao",
                "valor": "Valor",
            }
        )
        detalhamento["Conta/Banco"] = detalhamento["Conta/Banco"].fillna("").replace("None", "")
        detalhamento["Prontuario"] = detalhamento["Prontuario"].fillna("").replace("None", "")
        detalhamento["Forma de pagamento"] = detalhamento["Forma de pagamento"].fillna("").replace("None", "")
        detalhamento["Observacao"] = detalhamento["Observacao"].fillna("").replace("None", "")

        grupos.append(
            {
                "data": data_grupo,
                "entradas": entradas,
                "saidas": saidas,
                "saldo": saldo,
                "detalhamento": detalhamento.reset_index(drop=True),
            }
        )

        resumo_linhas.append(
            {
                "Data": data_grupo,
                "Entradas": entradas,
                "Saidas": saidas,
                "Saldo do dia": saldo,
                "Lancamentos": len(grupo),
            }
        )

    resumo_df = pd.DataFrame(resumo_linhas)
    return resumo_df, grupos


def caixa_diario_para_excel_bytes(financeiro_df):
    if not OPENPYXL_DISPONIVEL:
        return None
    caixa = financeiro_df.copy()
    if caixa.empty:
        return dataframe_para_excel_bytes(pd.DataFrame(columns=["Data", "Descricao", "Entradas", "Saidas", "Banco"]), nome_aba="Caixa")

    caixa["data_date"] = caixa["data"].apply(parse_data_contrato)
    caixa = caixa[caixa["data_date"].notna()].copy()
    if caixa.empty:
        return dataframe_para_excel_bytes(pd.DataFrame(columns=["Data", "Descricao", "Entradas", "Saidas", "Banco"]), nome_aba="Caixa")

    caixa["conta_caixa"] = caixa.apply(
        lambda row: str(row.get("conta_caixa") or "").strip().upper()
        or identificar_conta_caixa(row.get("origem", ""), row.get("descricao", ""), row.get("observacao", "")),
        axis=1,
    )
    caixa = caixa.sort_values(["data_date", "id"], ascending=[True, True])

    wb = Workbook()
    ws = wb.active
    ws.title = "Caixa"

    fonte_titulo = Font(name="Times New Roman", size=20, bold=False, color="000000")
    fonte_cabecalho = Font(name="Times New Roman", size=12, bold=False, color="000000")
    fonte_normal = Font(name="Times New Roman", size=12, color="000000")
    fonte_vermelha = Font(name="Times New Roman", size=12, color="FF0000")
    fonte_preta = Font(name="Times New Roman", size=12, color="000000")
    preenchimento_verde = PatternFill(fill_type="solid", fgColor="C6E0B4")
    borda_fina = Border(
        left=Side(style="thin", color="000000"),
        right=Side(style="thin", color="000000"),
        top=Side(style="thin", color="000000"),
        bottom=Side(style="thin", color="000000"),
    )

    ws.merge_cells("A1:E1")
    ws["A1"] = "CAIXA"
    ws["A1"].font = fonte_titulo
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")

    larguras = {"A": 14, "B": 58, "C": 18, "D": 16, "E": 18}
    for coluna, largura in larguras.items():
        ws.column_dimensions[coluna].width = largura

    linha_atual = 3
    saldo_anterior = {conta: 0.0 for conta in CONTAS_CAIXA_MODELO}

    for data_atual in sorted(caixa["data_date"].dropna().unique().tolist()):
        data_texto = formatar_data_br(data_atual)
        grupo = caixa[caixa["data_date"] == data_atual].copy()
        saldo_anterior = saldos_informados_por_conta_ate(data_atual)

        ws.cell(row=linha_atual, column=1, value="Data").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=2, value="Descrição").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=3, value="Entradas").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=4, value="Saídas").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=5, value="Banco").font = fonte_cabecalho
        for coluna in range(1, 6):
            ws.cell(row=linha_atual, column=coluna).border = borda_fina
        linha_atual += 1

        for conta in CONTAS_CAIXA_MODELO:
            saldo_valor = round(saldo_anterior[conta], 2)
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=f"SALDO {conta}")
            ws.cell(row=linha_atual, column=3, value=saldo_valor)
            for coluna in range(1, 6):
                ws.cell(row=linha_atual, column=coluna).fill = preenchimento_verde
                ws.cell(row=linha_atual, column=coluna).border = borda_fina
                ws.cell(row=linha_atual, column=coluna).font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=3).font = fonte_vermelha if saldo_valor < 0 else fonte_preta
            linha_atual += 1

        linha_atual += 3

        for _, row in grupo.iterrows():
            valor = round(float(row["valor"] or 0), 2)
            conta = row["conta_caixa"]
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=str(row["descricao"] or row["origem"] or "").strip())
            if str(row["tipo"]) == "Entrada":
                ws.cell(row=linha_atual, column=3, value=valor)
                ws.cell(row=linha_atual, column=4, value=None)
                saldo_anterior[conta] += valor
            else:
                ws.cell(row=linha_atual, column=3, value=None)
                ws.cell(row=linha_atual, column=4, value=valor)
                saldo_anterior[conta] -= valor
            ws.cell(row=linha_atual, column=5, value=conta)

            for coluna in range(1, 6):
                ws.cell(row=linha_atual, column=coluna).border = borda_fina
                ws.cell(row=linha_atual, column=coluna).font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=4).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=4).font = fonte_vermelha if ws.cell(row=linha_atual, column=4).value not in (None, 0, 0.0) else fonte_normal
            linha_atual += 1

        linha_atual += 3

        for conta in CONTAS_CAIXA_MODELO:
            saldo_valor = round(saldo_anterior[conta], 2)
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=f"SALDO {conta}")
            ws.cell(row=linha_atual, column=3, value=saldo_valor)
            for coluna in range(1, 6):
                ws.cell(row=linha_atual, column=coluna).fill = preenchimento_verde
                ws.cell(row=linha_atual, column=coluna).border = borda_fina
                ws.cell(row=linha_atual, column=coluna).font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=3).font = fonte_vermelha if saldo_valor < 0 else fonte_preta
            linha_atual += 1

        linha_atual += 3

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def dataframe_para_pdf_bytes(df, titulo="Relatorio"):
    if not REPORTLAB_DISPONIVEL:
        return None

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    largura, altura = A4
    y = altura - 40

    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(40, y, titulo)
    y -= 24

    pdf.setFont("Helvetica", 8)
    colunas = [str(coluna) for coluna in df.columns]
    linhas = [colunas] + df.astype(str).values.tolist()

    for linha in linhas:
        texto = " | ".join(linha)
        while len(texto) > 140:
            pdf.drawString(40, y, texto[:140])
            texto = texto[140:]
            y -= 12
            if y < 40:
                pdf.showPage()
                pdf.setFont("Helvetica", 8)
                y = altura - 40
        pdf.drawString(40, y, texto)
        y -= 12
        if y < 40:
            pdf.showPage()
            pdf.setFont("Helvetica", 8)
            y = altura - 40

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def identificar_coluna(colunas, candidatos):
    mapa = {normalizar_texto(coluna): coluna for coluna in colunas}
    for candidato in candidatos:
        chave = normalizar_texto(candidato)
        if chave in mapa:
            return mapa[chave]
    for coluna in colunas:
        texto = normalizar_texto(coluna)
        if any(normalizar_texto(candidato) in texto for candidato in candidatos):
            return coluna
    return None


def ler_extrato_arquivo(arquivo):
    nome = (arquivo.name or "").lower()
    if nome.endswith(".csv"):
        try:
            return pd.read_csv(arquivo, sep=None, engine="python")
        except Exception:
            arquivo.seek(0)
            return pd.read_csv(arquivo, sep=";", engine="python")
    return pd.read_excel(arquivo)


def reconhecer_extrato(df_origem):
    if df_origem is None or df_origem.empty:
        return pd.DataFrame()

    df = df_origem.copy()
    coluna_data = identificar_coluna(df.columns, ["data", "date", "lancamento", "lançamento", "movimento"])
    coluna_descricao = identificar_coluna(df.columns, ["descricao", "descrição", "historico", "histórico", "detalhe", "complemento"])
    coluna_valor = identificar_coluna(df.columns, ["valor", "amount", "valor (r$)", "lancamento", "lançamento"])
    coluna_credito = identificar_coluna(df.columns, ["credito", "crédito", "entrada"])
    coluna_debito = identificar_coluna(df.columns, ["debito", "débito", "saida", "saída"])

    if not coluna_data or not coluna_descricao or (not coluna_valor and not (coluna_credito or coluna_debito)):
        return pd.DataFrame()

    reconhecido = pd.DataFrame()
    reconhecido["data"] = pd.to_datetime(df[coluna_data], dayfirst=True, errors="coerce")
    reconhecido["descricao"] = df[coluna_descricao].fillna("").astype(str).str.strip()

    if coluna_valor:
        serie_valor = pd.to_numeric(
            df[coluna_valor]
            .astype(str)
            .str.replace(".", "", regex=False)
            .str.replace(",", ".", regex=False),
            errors="coerce",
        )
    else:
        credito = pd.to_numeric(
            df[coluna_credito].fillna(0).astype(str).str.replace(".", "", regex=False).str.replace(",", ".", regex=False),
            errors="coerce",
        ) if coluna_credito else 0
        debito = pd.to_numeric(
            df[coluna_debito].fillna(0).astype(str).str.replace(".", "", regex=False).str.replace(",", ".", regex=False),
            errors="coerce",
        ) if coluna_debito else 0
        serie_valor = credito.fillna(0) - debito.fillna(0)

    reconhecido["valor_original"] = serie_valor
    reconhecido = reconhecido.dropna(subset=["data", "valor_original"])
    reconhecido = reconhecido[reconhecido["descricao"] != ""]
    reconhecido["tipo"] = reconhecido["valor_original"].apply(lambda valor: "Entrada" if float(valor) >= 0 else "Saida")
    reconhecido["valor"] = reconhecido["valor_original"].abs().round(2)
    reconhecido["data_exibicao"] = reconhecido["data"].dt.strftime("%d/%m/%Y")
    return reconhecido[["data", "data_exibicao", "descricao", "valor", "tipo"]].reset_index(drop=True)


def preencher_prontuarios_recebiveis():
    pacientes = carregar_pacientes()
    if pacientes.empty:
        return

    prontuario_por_id = {}
    prontuario_por_nome = {}
    for _, row in pacientes.iterrows():
        prontuario = str(row["prontuario"] or "").strip()
        prontuario_por_id[int(row["id"])] = prontuario
        nome_normalizado = normalizar_texto(row["nome"])
        if nome_normalizado and prontuario and nome_normalizado not in prontuario_por_nome:
            prontuario_por_nome[nome_normalizado] = prontuario

    recebiveis_sem_prontuario = cursor.execute(
        """
        SELECT id, paciente_id, paciente_nome, prontuario
        FROM recebiveis
        """
    ).fetchall()

    houve_atualizacao = False
    for row in recebiveis_sem_prontuario:
        prontuario_atual = str(row["prontuario"] or "").strip()
        if prontuario_atual and prontuario_atual.lower() != "none":
            continue

        prontuario = ""
        if row["paciente_id"] is not None:
            prontuario = prontuario_por_id.get(int(row["paciente_id"]), "")
        if not prontuario:
            prontuario = prontuario_por_nome.get(normalizar_texto(row["paciente_nome"]), "")

        if prontuario:
            cursor.execute(
                "UPDATE recebiveis SET prontuario=? WHERE id=?",
                (prontuario, int(row["id"])),
            )
            houve_atualizacao = True

    if houve_atualizacao:
        conn.commit()


def sincronizar_recebiveis_contrato(contrato_id, paciente_id, paciente_nome, prontuario, valor_total, entrada, parcelas, primeiro_vencimento, forma_pagamento):
    planejados = montar_recebiveis_planejados(
        contrato_id,
        paciente_id,
        paciente_nome,
        prontuario,
        valor_total,
        entrada,
        parcelas,
        primeiro_vencimento,
        forma_pagamento,
    )

    existentes = cursor.execute(
        """
        SELECT paciente_id, paciente_nome, prontuario, parcela_numero, vencimento, valor, forma_pagamento, status, observacao
        FROM recebiveis
        WHERE contrato_id=?
        ORDER BY parcela_numero, id
        """,
        (contrato_id,),
    ).fetchall()

    existentes_normalizados = [
        (
            int(row["paciente_id"] or 0),
            row["paciente_nome"] or "",
            row["prontuario"] or "",
            int(row["parcela_numero"] or 0),
            row["vencimento"] or "",
            round(float(row["valor"] or 0), 2),
            row["forma_pagamento"] or "",
            row["status"] or "Aberto",
            row["observacao"] or "",
        )
        for row in existentes
    ]
    planejados_normalizados = [
        (
            item["paciente_id"],
            item["paciente_nome"],
            item["prontuario"],
            item["parcela_numero"],
            item["vencimento"],
            item["valor"],
            item["forma_pagamento"],
            item["status"],
            item["observacao"],
        )
        for item in planejados
    ]

    if existentes_normalizados == planejados_normalizados:
        return "inalterado"

    cursor.execute("DELETE FROM recebiveis WHERE contrato_id=?", (contrato_id,))
    for item in planejados:
        cursor.execute(
            """
            INSERT INTO recebiveis
            (contrato_id, paciente_id, paciente_nome, prontuario, parcela_numero, vencimento, valor, forma_pagamento, status, observacao, data_criacao)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                item["contrato_id"],
                item["paciente_id"],
                item["paciente_nome"],
                item["prontuario"],
                item["parcela_numero"],
                item["vencimento"],
                item["valor"],
                item["forma_pagamento"],
                item["status"],
                item["observacao"],
                agora_str(),
            ),
        )

    return "criado" if not existentes_normalizados else "atualizado"


def opcoes_pacientes(df_pacientes):
    opcoes = []
    for _, row in df_pacientes.iterrows():
        opcoes.append((int(row["id"]), f"{row['nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])}"))
    return opcoes


def opcoes_contratos(df_contratos, df_pacientes):
    pacientes_por_id = {}
    for _, row in df_pacientes.iterrows():
        pacientes_por_id[int(row["id"])] = {
            "nome": row["nome"],
            "prontuario": row["prontuario"],
        }

    opcoes = []
    for _, row in df_contratos.iterrows():
        contrato_id = int(row["id"])
        paciente_id = int(row["paciente_id"]) if not pd.isna(row["paciente_id"]) else None
        paciente = pacientes_por_id.get(paciente_id, {"nome": "Paciente nao encontrado", "prontuario": ""})
        valor_total = float(row["valor_total"] or 0)
        forma_pagamento = row["forma_pagamento"] or ""
        descricao = (
            f"Contrato {contrato_id} - {paciente['nome']} - "
            f"Prontuario {paciente['prontuario']} - "
            f"R$ {valor_total:.2f} - {forma_pagamento}"
        )
        opcoes.append((contrato_id, descricao))
    return opcoes


def filtrar_opcoes_contratos(contratos_opcoes, termo_busca):
    termo = normalizar_texto(termo_busca)
    if not termo:
        return contratos_opcoes
    return [
        (contrato_id, descricao)
        for contrato_id, descricao in contratos_opcoes
        if termo in normalizar_texto(descricao)
    ]


def normalizar_nome_coluna_importacao(coluna):
    texto = normalizar_texto(coluna)
    texto = texto.replace(" ", "_")
    return texto


def normalizar_forma_pagamento_importacao(valor):
    texto = normalizar_texto(valor)
    if "boleto" in texto:
        return "Boleto"
    if "pix" in texto:
        return "Pix"
    if "debito" in texto:
        return "Debito"
    if "credito" in texto or "cartao" in texto or "cartao de credito" in texto or "a vista" in texto:
        return "Credito" if "credito" in texto or "cartao" in texto else "Dinheiro"
    if "dinheiro" in texto:
        return "Dinheiro"
    return "Boleto" if "parcela" in texto else "Credito"


def valor_float_importacao(valor):
    if pd.isna(valor):
        return 0.0
    if isinstance(valor, (int, float)):
        return round(float(valor), 2)
    texto = str(valor).strip()
    if not texto:
        return 0.0
    texto = texto.replace("R$", "").replace(".", "").replace(",", ".").strip()
    try:
        return round(float(texto), 2)
    except ValueError:
        return 0.0


def texto_importacao(valor):
    if pd.isna(valor):
        return ""
    texto = str(valor).strip()
    return "" if texto.lower() in {"nat", "nan", "none"} else texto


def data_importacao_para_br(valor):
    return formatar_data_br_valor(valor)


def montar_hash_importacao_contrato(prontuario, cpf, data_contrato, valor_total, forma_pagamento, procedimentos):
    base = "|".join(
        [
            str(prontuario or "").strip(),
            limpar_cpf(cpf),
            data_importacao_para_br(data_contrato),
            f"{valor_float_importacao(valor_total):.2f}",
            normalizar_forma_pagamento_importacao(forma_pagamento),
            "|".join([normalizar_texto(proc.get("procedimento", "")) + f":{float(proc.get('valor', 0)):.2f}" for proc in procedimentos]),
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def localizar_ou_criar_paciente_importacao(row):
    prontuario = texto_importacao(row.get("cdg"))
    cpf = limpar_cpf(row.get("cpf"))
    paciente_row = None

    if prontuario:
        paciente_row = cursor.execute("SELECT * FROM pacientes WHERE prontuario=? LIMIT 1", (prontuario,)).fetchone()
    if paciente_row is None and cpf:
        paciente_row = cursor.execute("SELECT * FROM pacientes WHERE cpf=? LIMIT 1", (cpf,)).fetchone()

    nome = texto_importacao(row.get("nome"))
    cep = texto_importacao(row.get("cep"))
    endereco = texto_importacao(row.get("rua"))
    numero = texto_importacao(row.get("numero"))
    complemento = texto_importacao(row.get("complemento"))
    if complemento:
        numero = f"{numero} {complemento}".strip()
    bairro = texto_importacao(row.get("bairro"))
    cidade = texto_importacao(row.get("cidade"))
    estado = texto_importacao(row.get("estado"))
    telefone = texto_importacao(row.get("telefone"))
    nascimento = data_importacao_para_br(row.get("nascimento"))
    menor_nome = texto_importacao(row.get("menor_nome"))
    menor_cpf = limpar_cpf(row.get("menor_cpf"))
    menor_nascimento = data_importacao_para_br(row.get("menor_nascimento"))

    if menor_nome:
        responsavel = nome
        cpf_responsavel = cpf
        nome_paciente = menor_nome
        cpf_paciente = menor_cpf
        nascimento_paciente = menor_nascimento
        menor_idade = 1
    else:
        responsavel = ""
        cpf_responsavel = ""
        nome_paciente = nome
        cpf_paciente = cpf
        nascimento_paciente = nascimento
        menor_idade = 0

    if paciente_row is None:
        cursor.execute(
            """
            INSERT INTO pacientes
            (nome, prontuario, cpf, data_nascimento, telefone, cep, endereco, numero, bairro, cidade, estado, menor_idade, responsavel, cpf_responsavel)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                nome_paciente,
                prontuario,
                cpf_paciente,
                nascimento_paciente,
                telefone,
                cep,
                endereco,
                numero,
                bairro,
                cidade,
                estado,
                menor_idade,
                responsavel,
                cpf_responsavel,
            ),
        )
        return cursor.lastrowid

    cursor.execute(
        """
        UPDATE pacientes
        SET nome=?, prontuario=?, cpf=?, data_nascimento=?, telefone=?, cep=?, endereco=?, numero=?, bairro=?, cidade=?, estado=?, menor_idade=?, responsavel=?, cpf_responsavel=?
        WHERE id=?
        """,
        (
            nome_paciente or paciente_row["nome"],
            prontuario or paciente_row["prontuario"],
            cpf_paciente or paciente_row["cpf"],
            nascimento_paciente or paciente_row["data_nascimento"],
            telefone or paciente_row["telefone"],
            cep or paciente_row["cep"],
            endereco or paciente_row["endereco"],
            numero or paciente_row["numero"],
            bairro or paciente_row["bairro"],
            cidade or paciente_row["cidade"],
            estado or paciente_row["estado"],
            menor_idade,
            responsavel or paciente_row["responsavel"],
            cpf_responsavel or paciente_row["cpf_responsavel"],
            int(paciente_row["id"]),
        ),
    )
    return int(paciente_row["id"])


def extrair_procedimentos_importacao(row):
    procedimentos = []
    for indice in range(1, 12):
        chave_proc = f"procedimento_{indice}"
        chave_valor = f"valor_{indice}"
        procedimento = texto_importacao(row.get(chave_proc))
        valor = valor_float_importacao(row.get(chave_valor))
        if procedimento:
            procedimentos.append({"procedimento": procedimento, "valor": valor})
    return procedimentos


def preparar_planilha_importacao_contratos(arquivo):
    df = pd.read_excel(arquivo, sheet_name="Planilha1")
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    contratos = []
    for _, row in df.iterrows():
        nome = texto_importacao(row.get("nome"))
        prontuario = texto_importacao(row.get("cdg"))
        if not nome or not prontuario:
            continue
        procedimentos = extrair_procedimentos_importacao(row)
        if not procedimentos:
            continue
        forma_pagamento = normalizar_forma_pagamento_importacao(row.get("forma_de_pagamento"))
        valor_total = valor_float_importacao(row.get("total") or row.get("valor"))
        entrada = valor_float_importacao(row.get("valor_de_entrada"))
        parcelas = int(valor_float_importacao(row.get("parcelas")) or 1)
        contratos.append(
            {
                "row": row.to_dict(),
                "nome": nome,
                "prontuario": prontuario,
                "cpf": limpar_cpf(row.get("cpf")),
                "data": data_importacao_para_br(row.get("data")),
                "forma_pagamento": forma_pagamento,
                "valor_total": valor_total,
                "entrada": entrada,
                "parcelas": parcelas if parcelas > 0 else 1,
                "data_entrada": data_importacao_para_br(row.get("data_da_entrada") or row.get("data")),
                "primeiro_boleto": data_importacao_para_br(row.get("vencimento_do_1_boleto")),
                "procedimentos": procedimentos,
                "hash_importacao": montar_hash_importacao_contrato(
                    prontuario,
                    row.get("cpf"),
                    row.get("data"),
                    valor_total,
                    forma_pagamento,
                    procedimentos,
                ),
            }
        )
    return contratos


def importar_contratos_preparados(contratos_preparados):
    resultado = {"importados": 0, "ignorados": 0, "atualizados": 0, "erros": [], "recebiveis_pendentes": []}
    for item in contratos_preparados:
        contrato_existente = cursor.execute(
            "SELECT * FROM contratos WHERE hash_importacao=? LIMIT 1",
            (item["hash_importacao"],),
        ).fetchone()
        if contrato_existente is not None:
            try:
                houve_atualizacao = False
                primeiro_boleto_atual = formatar_data_br_valor(contrato_existente["primeiro_vencimento"] or "")
                data_entrada_atual = formatar_data_br_valor(contrato_existente["data_pagamento_entrada"] or "")
                primeiro_boleto_novo = formatar_data_br_valor(item["primeiro_boleto"])
                data_entrada_nova = formatar_data_br_valor(item["data_entrada"])

                if (
                    (not primeiro_boleto_atual and primeiro_boleto_novo)
                    or (not data_entrada_atual and data_entrada_nova)
                ):
                    cursor.execute(
                        """
                        UPDATE contratos
                        SET primeiro_vencimento=?, data_pagamento_entrada=?
                        WHERE id=?
                        """,
                        (
                            primeiro_boleto_novo or contrato_existente["primeiro_vencimento"],
                            data_entrada_nova or contrato_existente["data_pagamento_entrada"],
                            int(contrato_existente["id"]),
                        ),
                    )
                    houve_atualizacao = True

                if not forma_pagamento_a_vista(item["forma_pagamento"]):
                    recebiveis_existentes = cursor.execute(
                        "SELECT COUNT(1) AS total FROM recebiveis WHERE contrato_id=?",
                        (int(contrato_existente["id"]),),
                    ).fetchone()
                    if int(recebiveis_existentes["total"] or 0) == 0:
                        if parse_data_contrato(primeiro_boleto_novo):
                            sincronizar_recebiveis_contrato(
                                int(contrato_existente["id"]),
                                int(contrato_existente["paciente_id"]),
                                item["nome"],
                                item["prontuario"],
                                item["valor_total"],
                                item["entrada"],
                                item["parcelas"],
                                primeiro_boleto_novo,
                                item["forma_pagamento"],
                            )
                            houve_atualizacao = True
                        else:
                            resultado["recebiveis_pendentes"].append(
                                f"{item['nome']} / prontuario {item['prontuario']}"
                            )

                if houve_atualizacao:
                    resultado["atualizados"] += 1
                else:
                    resultado["ignorados"] += 1
                continue
            except Exception as exc:
                resultado["erros"].append(f"{item['nome']} / prontuario {item['prontuario']}: {exc}")
                continue

        try:
            paciente_id = localizar_ou_criar_paciente_importacao(item["row"])
            cursor.execute(
                """
                INSERT INTO contratos
                (paciente_id, valor_total, entrada, parcelas, primeiro_vencimento, data_pagamento_entrada, forma_pagamento, hash_importacao, data_criacao)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    paciente_id,
                    item["valor_total"],
                    item["entrada"],
                    item["parcelas"],
                    item["primeiro_boleto"],
                    item["data_entrada"],
                    item["forma_pagamento"],
                    item["hash_importacao"],
                    item["data"],
                ),
            )
            contrato_id = cursor.lastrowid
            salvar_procedimentos_contrato(
                contrato_id,
                [proc["procedimento"] for proc in item["procedimentos"]],
                [proc["valor"] for proc in item["procedimentos"]],
            )
            if not forma_pagamento_a_vista(item["forma_pagamento"]):
                if parse_data_contrato(item["primeiro_boleto"]):
                    sincronizar_recebiveis_contrato(
                        contrato_id,
                        paciente_id,
                        item["nome"],
                        item["prontuario"],
                        item["valor_total"],
                        item["entrada"],
                        item["parcelas"],
                        item["primeiro_boleto"],
                        item["forma_pagamento"],
                    )
                else:
                    resultado["recebiveis_pendentes"].append(
                        f"{item['nome']} / prontuario {item['prontuario']}"
                    )
            resultado["importados"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['nome']} / prontuario {item['prontuario']}: {exc}")
    conn.commit()
    return resultado


def garantir_metas_vendas_iniciais(ano=None):
    ano_referencia = int(ano or date.today().year)
    row = cursor.execute("SELECT id FROM metas_vendas WHERE ano=? LIMIT 1", (ano_referencia,)).fetchone()
    if row is None:
        cursor.execute(
            """
            INSERT INTO metas_vendas
            (ano, meta, supermeta, hipermeta, data_atualizacao)
            VALUES (?, ?, ?, ?, ?)
            """,
            (ano_referencia, 100000.0, 150000.0, 200000.0, agora_str()),
        )
        conn.commit()


def carregar_metas_vendas(ano=None):
    ano_referencia = int(ano or date.today().year)
    garantir_metas_vendas_iniciais(ano_referencia)
    row = cursor.execute("SELECT * FROM metas_vendas WHERE ano=? LIMIT 1", (ano_referencia,)).fetchone()
    return dict(row) if row else {
        "ano": ano_referencia,
        "meta": 100000.0,
        "supermeta": 150000.0,
        "hipermeta": 200000.0,
    }


def salvar_metas_vendas(ano, meta, supermeta, hipermeta):
    garantir_metas_vendas_iniciais(ano)
    cursor.execute(
        """
        UPDATE metas_vendas
        SET meta=?, supermeta=?, hipermeta=?, data_atualizacao=?
        WHERE ano=?
        """,
        (float(meta), float(supermeta), float(hipermeta), agora_str(), int(ano)),
    )


def localizar_paciente_por_nome_importacao(nome, pacientes_df=None):
    nome_texto = texto_importacao(nome)
    if not nome_texto:
        return None
    pacientes = pacientes_df.copy() if pacientes_df is not None else carregar_pacientes()
    if pacientes.empty:
        return None
    if "_nome_norm" not in pacientes.columns:
        pacientes["_nome_norm"] = pacientes["nome"].apply(normalizar_texto)
    nome_norm = normalizar_texto(nome_texto)
    encontrados = pacientes[pacientes["_nome_norm"] == nome_norm]
    if encontrados.empty:
        return None
    return encontrados.iloc[0]


def montar_hash_importacao_recebivel(nome, vencimento, valor):
    base = "|".join(
        [
            normalizar_texto(nome),
            data_importacao_para_br(vencimento),
            f"{valor_float_importacao(valor):.2f}",
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def normalizar_status_recebivel_importacao(status, pago):
    status_texto = normalizar_texto(status)
    pago_texto = normalizar_texto(pago)
    data_pagamento = data_importacao_para_br(pago) if parse_data_contrato(pago) else ""
    marcadores_pago = {"pago", "sim", "ok", "x", "quitado", "quitada"}

    if data_pagamento or status_texto in marcadores_pago or pago_texto in marcadores_pago:
        return "Pago", data_pagamento
    if "atras" in status_texto:
        return "Atrasado", ""
    if any(termo in status_texto for termo in ["suspens", "cancel"]):
        return "Suspenso", ""
    return "Aberto", ""


def montar_observacao_recebivel_importacao(row, status_final):
    partes = []
    cobranca = texto_importacao(row.get("cobranca"))
    telefone = texto_importacao(row.get("telefone_da_cobranca"))
    status_original = texto_importacao(row.get("status"))
    pago_original = texto_importacao(row.get("pago"))

    if cobranca:
        partes.append(f"Cobranca: {cobranca}")
    if telefone:
        partes.append(f"Telefone cobranca: {telefone}")
    if status_original and normalizar_texto(status_original) not in {normalizar_texto(status_final), ""}:
        partes.append(f"Status original: {status_original}")
    if pago_original and not parse_data_contrato(pago_original) and normalizar_texto(pago_original) not in {"", "pago", "sim", "ok", "x"}:
        partes.append(f"Pago original: {pago_original}")
    return " | ".join(partes)


def preparar_planilha_importacao_recebiveis(arquivo):
    df = pd.read_excel(arquivo)
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    pacientes_df = carregar_pacientes()
    if not pacientes_df.empty:
        pacientes_df["_nome_norm"] = pacientes_df["nome"].apply(normalizar_texto)

    recebiveis_preparados = []
    for _, row in df.iterrows():
        nome = texto_importacao(row.get("paciente"))
        vencimento = data_importacao_para_br(row.get("data_venc"))
        valor = valor_float_importacao(row.get("valor"))
        if not nome or not parse_data_contrato(vencimento) or valor <= 0:
            continue

        paciente = localizar_paciente_por_nome_importacao(nome, pacientes_df)
        status_final, data_pagamento = normalizar_status_recebivel_importacao(row.get("status"), row.get("pago"))
        recebiveis_preparados.append(
            {
                "paciente_id": int(paciente["id"]) if paciente is not None else None,
                "paciente_nome": str(paciente["nome"]) if paciente is not None else nome,
                "prontuario": str(paciente["prontuario"] or "") if paciente is not None else "",
                "vencimento": vencimento,
                "valor": valor,
                "forma_pagamento": "Boleto",
                "status": status_final,
                "data_pagamento": data_pagamento,
                "observacao": montar_observacao_recebivel_importacao(row, status_final),
                "hash_importacao": montar_hash_importacao_recebivel(nome, vencimento, valor),
            }
        )
    return recebiveis_preparados


def importar_recebiveis_preparados(recebiveis_preparados):
    resultado = {"inseridos": 0, "atualizados": 0, "erros": []}
    for item in recebiveis_preparados:
        try:
            existente = cursor.execute(
                "SELECT id FROM recebiveis WHERE hash_importacao=? LIMIT 1",
                (item["hash_importacao"],),
            ).fetchone()
            if existente is None:
                existente = cursor.execute(
                    """
                    SELECT id
                    FROM recebiveis
                    WHERE lower(trim(paciente_nome)) = lower(trim(?))
                      AND vencimento = ?
                      AND abs(valor - ?) < 0.01
                    LIMIT 1
                    """,
                    (item["paciente_nome"], item["vencimento"], float(item["valor"])),
                ).fetchone()

            if existente is None:
                cursor.execute(
                    """
                    INSERT INTO recebiveis
                    (contrato_id, paciente_id, paciente_nome, prontuario, parcela_numero, vencimento, valor, forma_pagamento, status, observacao, data_criacao, data_pagamento, hash_importacao)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        None,
                        item["paciente_id"],
                        item["paciente_nome"],
                        item["prontuario"],
                        None,
                        item["vencimento"],
                        item["valor"],
                        item["forma_pagamento"],
                        item["status"],
                        item["observacao"],
                        agora_str(),
                        item["data_pagamento"],
                        item["hash_importacao"],
                    ),
                )
                resultado["inseridos"] += 1
            else:
                cursor.execute(
                    """
                    UPDATE recebiveis
                    SET paciente_id=?, paciente_nome=?, prontuario=?, vencimento=?, valor=?, forma_pagamento=?, status=?, observacao=?, data_pagamento=?, hash_importacao=?
                    WHERE id=?
                    """,
                    (
                        item["paciente_id"],
                        item["paciente_nome"],
                        item["prontuario"],
                        item["vencimento"],
                        item["valor"],
                        item["forma_pagamento"],
                        item["status"],
                        item["observacao"],
                        item["data_pagamento"],
                        item["hash_importacao"],
                        int(existente["id"]),
                    ),
                )
                resultado["atualizados"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['paciente_nome']} / {item['vencimento']}: {exc}")
    conn.commit()
    return resultado


def montar_hash_importacao_conta_pagar(data_vencimento, descricao, fornecedor, valor):
    base = "|".join(
        [
            data_importacao_para_br(data_vencimento),
            normalizar_texto(descricao),
            normalizar_texto(fornecedor),
            f"{valor_float_importacao(valor):.2f}",
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def normalizar_status_conta_pagar_importacao(status, pago):
    status_texto = normalizar_texto(status)
    pago_texto = normalizar_texto(pago)
    if parse_data_contrato(pago) or pago_texto in {"sim", "pago", "ok", "x"} or status_texto == "pago":
        return "Pago", data_importacao_para_br(pago) if parse_data_contrato(pago) else ""
    if "atras" in status_texto:
        return "Atrasado", ""
    if any(termo in status_texto for termo in ["cancel", "suspens"]):
        return "Suspenso", ""
    return "A vencer", ""


def preparar_planilha_importacao_contas_pagar(arquivo):
    colunas_esperadas = {"data_de_vencimento", "descricao", "fornecedor", "valor"}
    df = pd.read_excel(arquivo, header=1)
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    if not colunas_esperadas.issubset(set(df.columns)):
        arquivo.seek(0)
        df = pd.read_excel(arquivo)
        df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    contas_preparadas = []
    for _, row in df.iterrows():
        data_vencimento = data_importacao_para_br(row.get("data_de_vencimento"))
        descricao = texto_importacao(row.get("descricao"))
        fornecedor = texto_importacao(row.get("fornecedor"))
        valor = valor_float_importacao(row.get("valor"))
        if not parse_data_contrato(data_vencimento) or (not descricao and not fornecedor) or valor <= 0:
            continue
        status_final, data_pagamento = normalizar_status_conta_pagar_importacao(row.get("status"), row.get("pago"))
        valor_pago = valor_float_importacao(row.get("valor_pago")) or (valor if status_final == "Pago" else 0.0)
        observacao = ""
        if texto_importacao(row.get("status")) and normalizar_texto(row.get("status")) not in {normalizar_texto(status_final), ""}:
            observacao = f"Status original: {texto_importacao(row.get('status'))}"
        contas_preparadas.append(
            {
                "data_vencimento": data_vencimento,
                "descricao": descricao,
                "fornecedor": fornecedor,
                "valor": valor,
                "pago": data_pagamento,
                "valor_pago": valor_pago,
                "status": status_final,
                "observacao": observacao,
                "hash_importacao": montar_hash_importacao_conta_pagar(data_vencimento, descricao, fornecedor, valor),
            }
        )
    return contas_preparadas


def importar_contas_pagar_preparadas(contas_preparadas):
    resultado = {"inseridos": 0, "atualizados": 0, "erros": []}
    for item in contas_preparadas:
        try:
            existente = cursor.execute(
                "SELECT id FROM contas_pagar WHERE hash_importacao=? LIMIT 1",
                (item["hash_importacao"],),
            ).fetchone()
            if existente is None:
                cursor.execute(
                    """
                    INSERT INTO contas_pagar
                    (data_vencimento, descricao, fornecedor, valor, pago, valor_pago, status, observacao, data_criacao, hash_importacao)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        item["data_vencimento"],
                        item["descricao"],
                        item["fornecedor"],
                        item["valor"],
                        item["pago"],
                        item["valor_pago"],
                        item["status"],
                        item["observacao"],
                        agora_str(),
                        item["hash_importacao"],
                    ),
                )
                resultado["inseridos"] += 1
            else:
                cursor.execute(
                    """
                    UPDATE contas_pagar
                    SET data_vencimento=?, descricao=?, fornecedor=?, valor=?, pago=?, valor_pago=?, status=?, observacao=?, hash_importacao=?
                    WHERE id=?
                    """,
                    (
                        item["data_vencimento"],
                        item["descricao"],
                        item["fornecedor"],
                        item["valor"],
                        item["pago"],
                        item["valor_pago"],
                        item["status"],
                        item["observacao"],
                        item["hash_importacao"],
                        int(existente["id"]),
                    ),
                )
                resultado["atualizados"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['fornecedor']} / {item['descricao']}: {exc}")
    conn.commit()
    return resultado


def montar_hash_importacao_venda(data_venda, paciente_nome, valor_total, nf):
    base = "|".join(
        [
            data_importacao_para_br(data_venda),
            normalizar_texto(paciente_nome),
            f"{valor_float_importacao(valor_total):.2f}",
            texto_importacao(nf),
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def preparar_planilha_importacao_vendas(arquivo):
    planilhas = pd.ExcelFile(arquivo)
    nome_planilha = "Planilha2" if "Planilha2" in planilhas.sheet_names else planilhas.sheet_names[0]
    arquivo.seek(0)
    df = pd.read_excel(arquivo, sheet_name=nome_planilha)
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    vendas_preparadas = []
    for _, row in df.iterrows():
        data_venda = data_importacao_para_br(row.get("data"))
        paciente_nome = texto_importacao(row.get("nome_paciente"))
        valor_total = valor_float_importacao(row.get("valor_total"))
        if not parse_data_contrato(data_venda) or not paciente_nome or valor_total <= 0:
            continue
        valor_a_vista = valor_float_importacao(row.get("a_vista"))
        valor_cartao = valor_float_importacao(row.get("cartao"))
        valor_boleto = valor_float_importacao(row.get("boleto"))
        saldo = valor_float_importacao(row.get("saldo"))
        data_a_pagar = data_importacao_para_br(row.get("data_a_pagar"))
        vendas_preparadas.append(
            {
                "data_venda": data_venda,
                "paciente_nome": paciente_nome,
                "valor_total": valor_total,
                "valor_a_vista": valor_a_vista,
                "valor_cartao": valor_cartao,
                "valor_boleto": valor_boleto,
                "saldo": saldo,
                "data_a_pagar": data_a_pagar,
                "avaliador": texto_importacao(row.get("avaliador")),
                "vendedor": texto_importacao(row.get("vendedor")),
                "nf": texto_importacao(row.get("nf")),
                "hash_importacao": montar_hash_importacao_venda(data_venda, paciente_nome, valor_total, row.get("nf")),
            }
        )
    return vendas_preparadas


def criar_contrato_automatico_venda(item, pacientes_df=None):
    paciente = localizar_paciente_por_nome_importacao(item["paciente_nome"], pacientes_df)
    if paciente is None:
        return None

    forma_pagamento = "Boleto" if item["saldo"] > 0 or item["valor_boleto"] > 0 else ("Credito" if item["valor_cartao"] > 0 else "Dinheiro")
    entrada = round(float(item["valor_a_vista"] or 0) + float(item["valor_cartao"] or 0), 2)
    primeiro_vencimento = item["data_a_pagar"] if not forma_pagamento_a_vista(forma_pagamento) else ""
    data_pagamento_entrada = item["data_venda"]

    contrato_hash = item["hash_importacao"]
    existente = cursor.execute("SELECT id FROM contratos WHERE hash_importacao=? LIMIT 1", (contrato_hash,)).fetchone()
    if existente is not None:
        return int(existente["id"])

    cursor.execute(
        """
        INSERT INTO contratos
        (paciente_id, valor_total, entrada, parcelas, primeiro_vencimento, data_pagamento_entrada, forma_pagamento, hash_importacao, data_criacao)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            int(paciente["id"]),
            item["valor_total"],
            entrada,
            1,
            primeiro_vencimento,
            data_pagamento_entrada,
            forma_pagamento,
            contrato_hash,
            item["data_venda"],
        ),
    )
    contrato_id = cursor.lastrowid
    salvar_procedimentos_contrato(contrato_id, ["Venda importada"], [float(item["valor_total"])])
    sincronizar_financeiro_contrato(contrato_id, item["paciente_nome"], item["valor_total"])
    if not forma_pagamento_a_vista(forma_pagamento) and item["saldo"] > 0:
        sincronizar_recebiveis_contrato(
            contrato_id,
            int(paciente["id"]),
            item["paciente_nome"],
            str(paciente["prontuario"] or ""),
            item["valor_total"],
            entrada,
            1,
            primeiro_vencimento,
            forma_pagamento,
        )
    return contrato_id


def importar_vendas_preparadas(vendas_preparadas, criar_contratos=True):
    resultado = {"inseridas": 0, "ignoradas": 0, "contratos_criados": 0, "erros": []}
    pacientes_df = carregar_pacientes()
    if not pacientes_df.empty:
        pacientes_df["_nome_norm"] = pacientes_df["nome"].apply(normalizar_texto)

    for item in vendas_preparadas:
        if cursor.execute("SELECT id FROM vendas WHERE hash_importacao=? LIMIT 1", (item["hash_importacao"],)).fetchone():
            resultado["ignoradas"] += 1
            continue
        try:
            contrato_id = None
            if criar_contratos:
                contrato_id = criar_contrato_automatico_venda(item, pacientes_df)
                if contrato_id:
                    resultado["contratos_criados"] += 1
            cursor.execute(
                """
                INSERT INTO vendas
                (data_venda, paciente_nome, valor_total, valor_a_vista, valor_cartao, valor_boleto, saldo, data_a_pagar, avaliador, vendedor, nf, contrato_id, hash_importacao, data_criacao)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    item["data_venda"],
                    item["paciente_nome"],
                    item["valor_total"],
                    item["valor_a_vista"],
                    item["valor_cartao"],
                    item["valor_boleto"],
                    item["saldo"],
                    item["data_a_pagar"],
                    item["avaliador"],
                    item["vendedor"],
                    item["nf"],
                    contrato_id,
                    item["hash_importacao"],
                    agora_str(),
                ),
            )
            resultado["inseridas"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['paciente_nome']} / {item['data_venda']}: {exc}")
    conn.commit()
    return resultado


def inicializar_banco_e_arquivos():
    inicializar_banco()
    garantir_metas_vendas_iniciais()
    garantir_pasta_documentos()


inicializar_banco_e_arquivos()
aplicar_tema_visual()

if "usuario_logado" not in st.session_state:
    st.session_state["usuario_logado"] = None

if not st.session_state["usuario_logado"]:
    renderizar_login()
    st.stop()

usuario_logado = st.session_state["usuario_logado"]
menus_disponiveis = [menu_nome for menu_nome in MODULOS_SISTEMA if usuario_tem_acesso(usuario_logado, menu_nome)]
st.sidebar.markdown(
    """
    <div class="ss-brand">
        <div class="ss-brand-title">SoulSul</div>
        <div class="ss-brand-subtitle">clínica integrada</div>
    </div>
    """,
    unsafe_allow_html=True,
)
st.sidebar.markdown("### Menu")
menu = st.sidebar.radio(
    "Menu",
    menus_disponiveis,
    label_visibility="collapsed",
)

st.sidebar.markdown("<div style='height: 12rem;'></div>", unsafe_allow_html=True)
st.sidebar.markdown("---")
st.sidebar.caption(f"Usuario: {usuario_logado['nome']}")
st.sidebar.caption(f"Perfil: {usuario_logado['perfil']}")

with st.sidebar.expander("Minha senha"):
    senha_atual = st.text_input("Senha atual", type="password", key="senha_atual_sidebar")
    nova_senha = st.text_input("Nova senha", type="password", key="nova_senha_sidebar")
    confirmar_nova_senha = st.text_input("Confirmar nova senha", type="password", key="confirmar_nova_senha_sidebar")
    if st.button("Alterar minha senha"):
        usuario_db = cursor.execute("SELECT * FROM usuarios WHERE id=?", (int(usuario_logado["id"]),)).fetchone()
        if usuario_db is None:
            st.sidebar.error("Usuario nao encontrado.")
        elif not verificar_senha(senha_atual, usuario_db["senha_hash"]):
            st.sidebar.error("Senha atual invalida.")
        elif len(nova_senha) < 4:
            st.sidebar.error("A nova senha deve ter pelo menos 4 caracteres.")
        elif nova_senha != confirmar_nova_senha:
            st.sidebar.error("A confirmacao da nova senha nao confere.")
        else:
            redefinir_senha_usuario(usuario_logado["id"], nova_senha)
            conn.commit()
            st.sidebar.success("Senha alterada com sucesso.")

if st.sidebar.button("Sair", use_container_width=True):
    registrar_log_acesso(usuario_logado["id"], usuario_logado["usuario"], "LOGOUT")
    st.session_state["usuario_logado"] = None
    st.rerun()

if menu == "Dashboard":
    st.title("Dashboard de Vendas")
    ano_dashboard = st.selectbox(
        "Ano das vendas",
        options=list(range(date.today().year - 2, date.today().year + 3)),
        index=2,
    )
    metas = carregar_metas_vendas(ano_dashboard)
    vendas = pd.read_sql("SELECT * FROM vendas ORDER BY data_venda DESC", conn)
    atualizar_status_contas_pagar_automaticamente()
    contas_pagar_dashboard = pd.read_sql("SELECT * FROM contas_pagar ORDER BY data_vencimento, fornecedor, descricao", conn)

    with st.expander("Metas de vendas", expanded=False):
        m1, m2, m3 = st.columns(3)
        meta = m1.number_input("Meta", min_value=0.0, value=float(metas["meta"] or 100000))
        supermeta = m2.number_input("Supermeta", min_value=0.0, value=float(metas["supermeta"] or 150000))
        hipermeta = m3.number_input("Hipermeta", min_value=0.0, value=float(metas["hipermeta"] or 200000))
        if st.button("Salvar metas de vendas"):
            salvar_metas_vendas(ano_dashboard, meta, supermeta, hipermeta)
            conn.commit()
            st.success("Metas atualizadas.")
            st.rerun()

    if vendas.empty:
        st.info("Nenhuma venda importada ainda.")
    else:
        vendas["data_ref"] = vendas["data_venda"].apply(parse_data_contrato)
        vendas = vendas[vendas["data_ref"].notna()].copy()
        vendas["ano"] = vendas["data_ref"].apply(lambda valor: valor.year)
        vendas = vendas[vendas["ano"] == ano_dashboard].copy()
        if vendas.empty:
            st.info("Nao ha vendas para o ano selecionado.")
        else:
            vendas["mes"] = vendas["data_ref"].apply(lambda valor: valor.month)
            vendas["mes_nome"] = vendas["mes"].map(
                {
                    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                    7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                }
            )
            fv1, fv2, fv3, fv4 = st.columns(4)
            filtro_pacientes_venda = fv1.multiselect(
                "Pacientes",
                options=sorted([valor for valor in vendas["paciente_nome"].fillna("").unique().tolist() if valor]),
            )
            filtro_avaliadores_venda = fv2.multiselect(
                "Avaliadores",
                options=sorted([valor for valor in vendas["avaliador"].fillna("").unique().tolist() if valor]),
            )
            filtro_vendedores_venda = fv3.multiselect(
                "Vendedores",
                options=sorted([valor for valor in vendas["vendedor"].fillna("").unique().tolist() if valor]),
            )
            filtro_meses_venda = fv4.multiselect(
                "Meses",
                options=list(range(1, 13)),
                format_func=lambda valor: {
                    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                    7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                }[valor],
            )

            vendas_filtradas = vendas.copy()
            if filtro_pacientes_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["paciente_nome"].isin(filtro_pacientes_venda)]
            if filtro_avaliadores_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["avaliador"].isin(filtro_avaliadores_venda)]
            if filtro_vendedores_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["vendedor"].isin(filtro_vendedores_venda)]
            if filtro_meses_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["mes"].isin(filtro_meses_venda)]

            if vendas_filtradas.empty:
                st.info("Nao ha vendas para os filtros selecionados.")
                st.stop()

            st.subheader("Calendário de pagamentos")
            if contas_pagar_dashboard.empty:
                st.info("Nenhum pagamento cadastrado para exibir no calendário.")
            else:
                contas_pagar_dashboard["valor"] = contas_pagar_dashboard["valor"].fillna(0).astype(float)
                contas_pagar_dashboard["_ordem_vencimento"] = pd.to_datetime(
                    contas_pagar_dashboard["data_vencimento"],
                    format="%d/%m/%Y",
                    errors="coerce",
                )
                contas_calendario = contas_pagar_dashboard[contas_pagar_dashboard["_ordem_vencimento"].notna()].copy()

                if contas_calendario.empty:
                    st.info("Nenhum vencimento válido para exibir no calendário.")
                else:
                    meses_pt_br = {
                        1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                        7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                    }
                    anos_calendario = sorted(set(contas_calendario["_ordem_vencimento"].dt.year.tolist()) | {date.today().year})
                    mes_atual = date.today().month
                    ano_atual = date.today().year

                    cal1, cal2 = st.columns(2)
                    ano_pagamentos = cal1.selectbox(
                        "Ano do calendário",
                        options=anos_calendario,
                        index=anos_calendario.index(ano_atual) if ano_atual in anos_calendario else 0,
                        key="dashboard_ano_pagamentos",
                    )
                    mes_pagamentos = cal2.selectbox(
                        "Mês do calendário",
                        options=list(range(1, 13)),
                        index=mes_atual - 1,
                        format_func=lambda valor: meses_pt_br[valor],
                        key="dashboard_mes_pagamentos",
                    )

                    contas_mes = contas_calendario[
                        (contas_calendario["_ordem_vencimento"].dt.year == ano_pagamentos)
                        & (contas_calendario["_ordem_vencimento"].dt.month == mes_pagamentos)
                    ].copy()

                    cabecalho = st.columns(7)
                    for coluna, nome_dia in zip(cabecalho, ["Seg", "Ter", "Qua", "Qui", "Sex", "Sáb", "Dom"]):
                        coluna.markdown(f"**{nome_dia}**")

                    hoje = date.today()
                    for semana in calendar.monthcalendar(ano_pagamentos, mes_pagamentos):
                        cols = st.columns(7)
                        for idx, dia in enumerate(semana):
                            with cols[idx]:
                                if dia == 0:
                                    st.write("")
                                else:
                                    dia_contas = contas_mes[contas_mes["_ordem_vencimento"].dt.day == dia].copy()
                                    total_dia = float(dia_contas["valor"].sum()) if not dia_contas.empty else 0.0
                                    qtd_dia = len(dia_contas)
                                    qtd_pagos = int((dia_contas["status"] == "Pago").sum()) if not dia_contas.empty else 0
                                    qtd_atrasados = int((dia_contas["status"] == "Atrasado").sum()) if not dia_contas.empty else 0

                                    eh_hoje = hoje.year == ano_pagamentos and hoje.month == mes_pagamentos and hoje.day == dia
                                    borda = "2px solid #111827" if eh_hoje else "1px solid #d1d5db"
                                    fundo = "#ffffff"
                                    if qtd_atrasados > 0:
                                        fundo = "#fef2f2"
                                    elif qtd_dia > 0 and qtd_pagos == qtd_dia:
                                        fundo = "#f0fdf4"
                                    elif qtd_dia > 0:
                                        fundo = "#fffbeb"

                                    resumo_status = ""
                                    if qtd_atrasados > 0:
                                        resumo_status = f"<div style='font-size:11px;color:#b91c1c;'>Atrasados: {qtd_atrasados}</div>"
                                    elif qtd_dia > 0 and qtd_pagos == qtd_dia:
                                        resumo_status = f"<div style='font-size:11px;color:#15803d;'>Todos pagos</div>"
                                    elif qtd_dia > 0:
                                        resumo_status = f"<div style='font-size:11px;color:#92400e;'>Títulos: {qtd_dia}</div>"

                                    st.markdown(
                                        f"""
                                        <div style="border:{borda};border-radius:10px;padding:6px;min-height:88px;background:{fundo};">
                                            <div style="font-weight:700;font-size:13px;margin-bottom:4px;">{dia:02d}</div>
                                            <div style="font-size:11px;color:#374151;">{formatar_moeda_br(total_dia) if qtd_dia else "Sem títulos"}</div>
                                            {resumo_status}
                                        </div>
                                        """,
                                        unsafe_allow_html=True,
                                    )

            hoje = date.today()
            vendas_mes_atual = vendas_filtradas[(vendas_filtradas["mes"] == hoje.month) & (vendas_filtradas["ano"] == hoje.year)]
            total_ano = float(vendas_filtradas["valor_total"].sum())
            total_mes = float(vendas_mes_atual["valor_total"].sum()) if not vendas_mes_atual.empty else 0.0
            ticket_medio = float(vendas_filtradas["valor_total"].mean()) if not vendas_filtradas.empty else 0.0

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Vendas do ano", formatar_moeda_br(total_ano))
            c2.metric("Vendas do mes", formatar_moeda_br(total_mes))
            c3.metric("Quantidade", str(len(vendas_filtradas)))
            c4.metric("Ticket medio", formatar_moeda_br(ticket_medio))

            c5, c6, c7 = st.columns(3)
            c5.metric("Meta", formatar_moeda_br(float(metas["meta"] or 0)))
            c6.metric("Supermeta", formatar_moeda_br(float(metas["supermeta"] or 0)))
            c7.metric("Hipermeta", formatar_moeda_br(float(metas["hipermeta"] or 0)))

            st.subheader("Progresso das metas")
            metas_progresso = [
                ("Meta", float(metas["meta"] or 0)),
                ("Supermeta", float(metas["supermeta"] or 0)),
                ("Hipermeta", float(metas["hipermeta"] or 0)),
            ]
            for nome_meta, valor_meta in metas_progresso:
                percentual = (total_mes / valor_meta) if valor_meta > 0 else 0
                percentual_exibicao = max(0.0, min(percentual, 1.0))
                diferenca = total_mes - valor_meta
                if diferenca >= 0:
                    status_meta = f"Meta batida. Excedente: {formatar_moeda_br(diferenca)}"
                else:
                    status_meta = f"Faltam {formatar_moeda_br(abs(diferenca))}"

                st.markdown(f"**{nome_meta}**: {formatar_moeda_br(valor_meta)}")
                st.progress(percentual_exibicao)
                st.caption(f"Progresso mensal: {percentual * 100:.1f}% | {status_meta}")

            resumo_mensal = vendas_filtradas.groupby(["mes", "mes_nome"], as_index=False)["valor_total"].sum().sort_values("mes")
            resumo_mensal = resumo_mensal.rename(columns={"mes_nome": "Mes", "valor_total": "Valor total"})
            ordem_meses = [
                "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
            ]
            resumo_mensal["Mes"] = pd.Categorical(resumo_mensal["Mes"], categories=ordem_meses, ordered=True)
            resumo_mensal = resumo_mensal.sort_values("Mes")
            st.subheader("Vendas por mes")
            grafico_vendas = resumo_mensal.copy()
            metas_linhas = pd.DataFrame(
                [
                    {"Meta": "Meta", "Valor": float(metas["meta"] or 0)},
                    {"Meta": "Supermeta", "Valor": float(metas["supermeta"] or 0)},
                    {"Meta": "Hipermeta", "Valor": float(metas["hipermeta"] or 0)},
                ]
            )
            barras = alt.Chart(grafico_vendas).mark_bar(color="#c5a77a").encode(
                x=alt.X("Mes:N", sort=ordem_meses, title="Mês"),
                y=alt.Y("Valor total:Q", title="Valor"),
                tooltip=[
                    alt.Tooltip("Mes:N", title="Mês"),
                    alt.Tooltip("Valor total:Q", title="Valor", format=",.2f"),
                ],
            )
            linhas = alt.Chart(metas_linhas).mark_rule(strokeWidth=2).encode(
                y=alt.Y("Valor:Q"),
                color=alt.Color(
                    "Meta:N",
                    scale=alt.Scale(
                        domain=["Meta", "Supermeta", "Hipermeta"],
                        range=["#2563eb", "#ea580c", "#15803d"],
                    ),
                    legend=alt.Legend(title="Metas"),
                ),
                tooltip=[
                    alt.Tooltip("Meta:N", title="Meta"),
                    alt.Tooltip("Valor:Q", title="Valor", format=",.2f"),
                ],
            )
            st.altair_chart((barras + linhas).properties(height=320), use_container_width=True)
            resumo_exibicao = resumo_mensal.copy()
            resumo_exibicao["Mes"] = resumo_exibicao["Mes"].astype(str)
            resumo_exibicao["Valor total"] = resumo_exibicao["Valor total"].map(formatar_moeda_br)
            st.dataframe(resumo_exibicao[["Mes", "Valor total"]], use_container_width=True, hide_index=True)

            detalhe_vendas = vendas_filtradas[
                ["data_venda", "paciente_nome", "valor_total", "valor_a_vista", "valor_cartao", "valor_boleto", "saldo", "avaliador", "vendedor", "nf"]
            ].copy()
            detalhe_vendas = detalhe_vendas.rename(
                columns={
                    "data_venda": "Data",
                    "paciente_nome": "Paciente",
                    "valor_total": "Valor total",
                    "valor_a_vista": "A vista",
                    "valor_cartao": "Cartao",
                    "valor_boleto": "Boleto",
                    "saldo": "Saldo",
                    "avaliador": "Avaliador",
                    "vendedor": "Vendedor",
                    "nf": "NF",
                }
            )
            for coluna in ["Valor total", "A vista", "Cartao", "Boleto", "Saldo"]:
                detalhe_vendas[coluna] = detalhe_vendas[coluna].map(formatar_moeda_br)
            st.subheader("Vendas importadas")
            st.dataframe(detalhe_vendas, use_container_width=True, hide_index=True)

            st.markdown("**Editar venda importada**")
            opcoes_vendas = [
                (
                    int(row["id"]),
                    f"{formatar_data_br_valor(row['data_venda'])} - {row['paciente_nome']} - {formatar_moeda_br(float(row['valor_total'] or 0))}"
                )
                for _, row in vendas_filtradas.sort_values(["data_ref", "paciente_nome"], ascending=[False, True]).iterrows()
            ]
            venda_id = st.selectbox(
                "Venda para editar",
                options=[opcao[0] for opcao in opcoes_vendas],
                format_func=lambda valor: next(rotulo for chave, rotulo in opcoes_vendas if chave == valor),
                key="venda_edicao_id",
            )
            venda_row = vendas_filtradas[vendas_filtradas["id"] == venda_id].iloc[0]

            ve1, ve2 = st.columns(2)
            edit_data_venda = ve1.text_input(
                "Data da venda",
                formatar_data_br_valor(venda_row["data_venda"] or ""),
                key=f"edit_data_venda_{venda_id}",
            )
            edit_paciente_venda = ve2.text_input(
                "Paciente",
                venda_row["paciente_nome"] or "",
                key=f"edit_paciente_venda_{venda_id}",
            )

            ve3, ve4, ve5, ve6 = st.columns(4)
            edit_valor_total = ve3.number_input(
                "Valor total",
                min_value=0.0,
                value=float(venda_row["valor_total"] or 0),
                key=f"edit_valor_total_venda_{venda_id}",
            )
            edit_valor_a_vista = ve4.number_input(
                "A vista",
                min_value=0.0,
                value=float(venda_row["valor_a_vista"] or 0),
                key=f"edit_valor_avista_venda_{venda_id}",
            )
            edit_valor_cartao = ve5.number_input(
                "Cartao",
                min_value=0.0,
                value=float(venda_row["valor_cartao"] or 0),
                key=f"edit_valor_cartao_venda_{venda_id}",
            )
            edit_valor_boleto = ve6.number_input(
                "Boleto",
                min_value=0.0,
                value=float(venda_row["valor_boleto"] or 0),
                key=f"edit_valor_boleto_venda_{venda_id}",
            )

            ve7, ve8, ve9, ve10 = st.columns(4)
            edit_saldo_venda = ve7.number_input(
                "Saldo",
                min_value=0.0,
                value=float(venda_row["saldo"] or 0),
                key=f"edit_saldo_venda_{venda_id}",
            )
            edit_data_a_pagar = ve8.text_input(
                "Data a pagar",
                formatar_data_br_valor(venda_row["data_a_pagar"] or ""),
                key=f"edit_data_a_pagar_venda_{venda_id}",
            )
            edit_avaliador = ve9.text_input(
                "Avaliador",
                venda_row["avaliador"] or "",
                key=f"edit_avaliador_venda_{venda_id}",
            )
            edit_vendedor = ve10.text_input(
                "Vendedor",
                venda_row["vendedor"] or "",
                key=f"edit_vendedor_venda_{venda_id}",
            )

            edit_nf = st.text_input(
                "NF",
                venda_row["nf"] or "",
                key=f"edit_nf_venda_{venda_id}",
            )

            if st.button("Salvar alteracoes da venda", key=f"salvar_venda_{venda_id}"):
                if not parse_data_contrato(edit_data_venda):
                    st.error("Informe a data da venda no formato DD/MM/AAAA.")
                elif not edit_paciente_venda.strip():
                    st.error("Informe o nome do paciente.")
                elif float(edit_valor_total) <= 0:
                    st.error("Informe um valor total maior que zero.")
                elif edit_data_a_pagar.strip() and not parse_data_contrato(edit_data_a_pagar):
                    st.error("Informe a data a pagar no formato DD/MM/AAAA.")
                else:
                    cursor.execute(
                        """
                        UPDATE vendas
                        SET data_venda=?, paciente_nome=?, valor_total=?, valor_a_vista=?, valor_cartao=?, valor_boleto=?, saldo=?, data_a_pagar=?, avaliador=?, vendedor=?, nf=?
                        WHERE id=?
                        """,
                        (
                            formatar_data_br_valor(edit_data_venda),
                            edit_paciente_venda.strip(),
                            float(edit_valor_total),
                            float(edit_valor_a_vista),
                            float(edit_valor_cartao),
                            float(edit_valor_boleto),
                            float(edit_saldo_venda),
                            formatar_data_br_valor(edit_data_a_pagar),
                            edit_avaliador.strip(),
                            edit_vendedor.strip(),
                            edit_nf.strip(),
                            int(venda_id),
                        ),
                    )
                    conn.commit()
                    st.success("Venda atualizada com sucesso.")
                    st.rerun()

if menu == "Pacientes":
    st.title("Cadastro de Pacientes")
    st.subheader("Importar paciente por PDF")
    pdf_paciente = st.file_uploader(
        "Enviar PDF com os dados do paciente",
        type=["pdf"],
        key="pdf_paciente_upload",
    )
    if pdf_paciente is not None:
        try:
            texto_pdf_paciente = extrair_texto_pdf_upload(pdf_paciente)
            dados_extraidos = extrair_dados_paciente_pdf(texto_pdf_paciente)
            if any(str(valor).strip() for valor in dados_extraidos.values()):
                st.write("Dados reconhecidos no PDF")
                st.json(dados_extraidos)
                if st.button("Usar dados extraidos no formulario"):
                    aplicar_dados_extraidos_paciente(dados_extraidos)
                    st.success("Formulario preenchido com os dados extraidos do PDF.")
            else:
                st.warning("Nao consegui reconhecer dados suficientes nesse PDF.")
        except Exception as exc:
            st.error(f"Falha ao ler o PDF do paciente: {exc}")

    nome = st.text_input("Nome", key="paciente_nome_input")
    prontuario = st.text_input("Prontuario", key="paciente_prontuario_input")
    cpf = st.text_input("CPF", key="paciente_cpf_input")
    nascimento = st.text_input("Data nascimento", key="paciente_nascimento_input")
    telefone = st.text_input("Telefone", key="paciente_telefone_input")
    cep = st.text_input("CEP", key="paciente_cep_input")

    dados_cep = buscar_endereco_por_cep(cep) if cep else {}
    if dados_cep:
        if not st.session_state.get("paciente_endereco_input"):
            st.session_state["paciente_endereco_input"] = dados_cep.get("logradouro", "")
        if not st.session_state.get("paciente_bairro_input"):
            st.session_state["paciente_bairro_input"] = dados_cep.get("bairro", "")
        if not st.session_state.get("paciente_cidade_input"):
            st.session_state["paciente_cidade_input"] = dados_cep.get("localidade", "")
        if not st.session_state.get("paciente_estado_input"):
            st.session_state["paciente_estado_input"] = dados_cep.get("uf", "")

    endereco = st.text_input("Rua", key="paciente_endereco_input")
    numero = st.text_input("Numero", key="paciente_numero_input")
    bairro = st.text_input("Bairro", key="paciente_bairro_input")
    cidade = st.text_input("Cidade", key="paciente_cidade_input")
    estado = st.text_input("Estado", key="paciente_estado_input")

    menor = st.checkbox("Paciente menor de idade", key="paciente_menor_input")
    responsavel = ""
    cpf_responsavel = ""

    if menor:
        responsavel = st.text_input("Responsavel", key="paciente_responsavel_input")
        cpf_responsavel = st.text_input("CPF responsavel", key="paciente_cpf_responsavel_input")

    if st.button("Salvar paciente"):
        erros = validar_dados_paciente(nome, prontuario, cpf, menor, responsavel, cpf_responsavel)
        if erros:
            for erro in erros:
                st.error(erro)
        else:
            cursor.execute(
                """
                INSERT INTO pacientes
                (
                    nome, prontuario, cpf, data_nascimento, telefone, cep,
                    endereco, numero, bairro, cidade, estado, menor_idade,
                    responsavel, cpf_responsavel
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    nome.strip(),
                    prontuario.strip(),
                    limpar_cpf(cpf),
                    nascimento.strip(),
                    telefone.strip(),
                    cep.strip(),
                    endereco.strip(),
                    numero.strip(),
                    bairro.strip(),
                    cidade.strip(),
                    estado.strip(),
                    int(menor),
                    responsavel.strip(),
                    limpar_cpf(cpf_responsavel),
                ),
            )
            conn.commit()
            st.success("Paciente cadastrado.")

    st.subheader("Pacientes cadastrados")
    pacientes = carregar_pacientes()
    if not pacientes.empty:
        st.dataframe(pacientes)

if menu == "Editar Paciente":
    st.title("Editar paciente")
    pacientes = carregar_pacientes()

    if pacientes.empty:
        st.warning("Nenhum paciente cadastrado.")
        st.stop()

    pacientes_opcoes = opcoes_pacientes(pacientes)
    paciente_id = st.selectbox(
        "Paciente",
        options=[opcao[0] for opcao in pacientes_opcoes],
        format_func=lambda paciente_id: next(
            label for valor, label in pacientes_opcoes if valor == paciente_id
        ),
    )

    paciente = pacientes[pacientes["id"] == paciente_id].iloc[0]

    nome = st.text_input("Nome", paciente["nome"] or "")
    prontuario = st.text_input("Prontuario", paciente["prontuario"] or "")
    cpf = st.text_input("CPF", paciente["cpf"] or "")
    nascimento = st.text_input("Data nascimento", paciente["data_nascimento"] or "")
    telefone = st.text_input("Telefone", paciente["telefone"] or "")
    cep = st.text_input("CEP", paciente["cep"] or "")

    endereco = st.text_input("Rua", paciente["endereco"] or "")
    numero = st.text_input("Numero", paciente["numero"] or "")
    bairro = st.text_input("Bairro", paciente["bairro"] or "")
    cidade = st.text_input("Cidade", paciente["cidade"] or "")
    estado = st.text_input("Estado", paciente["estado"] or "")

    menor = st.checkbox(
        "Paciente menor de idade",
        value=str(paciente["menor_idade"]) in {"1", "True", "true"},
    )

    responsavel_inicial = paciente["responsavel"] or ""
    cpf_responsavel_inicial = paciente["cpf_responsavel"] or ""
    responsavel = ""
    cpf_responsavel = ""

    if menor:
        responsavel = st.text_input("Responsavel", responsavel_inicial)
        cpf_responsavel = st.text_input("CPF responsavel", cpf_responsavel_inicial)

    if st.button("Salvar alteracoes do paciente"):
        erros = validar_dados_paciente(nome, prontuario, cpf, menor, responsavel, cpf_responsavel)
        if erros:
            for erro in erros:
                st.error(erro)
        else:
            cursor.execute(
                """
                UPDATE pacientes
                SET nome=?, prontuario=?, cpf=?, data_nascimento=?, telefone=?, cep=?,
                    endereco=?, numero=?, bairro=?, cidade=?, estado=?, menor_idade=?,
                    responsavel=?, cpf_responsavel=?
                WHERE id=?
                """,
                (
                    nome.strip(),
                    prontuario.strip(),
                    limpar_cpf(cpf),
                    nascimento.strip(),
                    telefone.strip(),
                    cep.strip(),
                    endereco.strip(),
                    numero.strip(),
                    bairro.strip(),
                    cidade.strip(),
                    estado.strip(),
                    int(menor),
                    responsavel.strip() if menor else "",
                    limpar_cpf(cpf_responsavel) if menor else "",
                    paciente_id,
                ),
            )
            conn.commit()
            st.success("Paciente atualizado.")

    st.subheader("Pacientes cadastrados")
    st.dataframe(pacientes)

if menu == "Contratos":
    st.title("Contratos")
    tab_contrato_novo, tab_contrato_lista = st.tabs(["Novo contrato", "Lista de contratos"])

    with tab_contrato_novo:
        pacientes = carregar_pacientes()

        if pacientes.empty:
            st.warning("Cadastre paciente primeiro.")
            st.stop()

        pacientes_opcoes = opcoes_pacientes(pacientes)
        paciente_id = st.selectbox(
            "Paciente",
            options=[opcao[0] for opcao in pacientes_opcoes],
            format_func=lambda paciente_id: next(
                label for valor, label in pacientes_opcoes if valor == paciente_id
            ),
        )
        paciente_nome = pacientes.loc[pacientes["id"] == paciente_id, "nome"].iloc[0]
        prontuario_paciente = pacientes.loc[pacientes["id"] == paciente_id, "prontuario"].iloc[0]

        forma = st.selectbox("Forma pagamento", FORMAS_PAGAMENTO)
        entrada = st.number_input("Entrada", min_value=0.0, value=0.0)
        data_pagamento_entrada = st.text_input(
            "Data do pagamento" if forma_pagamento_a_vista(forma) else "Data do pagamento da entrada",
            value=formatar_data_br(date.today()),
        )
        if forma_pagamento_a_vista(forma):
            parcelas = 1
            vencimento = ""
            st.caption("Para Pix, Dinheiro, Debito e Credito o sistema considera pagamento a vista.")
        else:
            parcelas = st.number_input("Parcelas", min_value=1, max_value=36, value=1)
            vencimento = st.text_input("Primeiro vencimento")

        procedimentos = []
        valores = []
        for i in range(10):
            col1, col2 = st.columns(2)
            proc = col1.text_input(f"Procedimento {i + 1}")
            valor = col2.number_input(f"Valor {i + 1}", min_value=0.0, value=0.0, key=f"v{i}")
            if proc.strip():
                procedimentos.append(proc.strip())
                valores.append(float(valor))

        total = sum(valores)
        st.write("Valor total:", total)

        if st.button("Salvar contrato"):
            erros, total = validar_contrato(procedimentos, valores, entrada, parcelas, vencimento, forma, data_pagamento_entrada)
            if erros:
                for erro in erros:
                    st.error(erro)
            else:
                cursor.execute(
                    """
                    INSERT INTO contratos
                    (
                        paciente_id, valor_total, entrada, parcelas,
                        primeiro_vencimento, data_pagamento_entrada, forma_pagamento, data_criacao
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        paciente_id,
                        total,
                        entrada,
                        parcelas,
                        vencimento.strip(),
                        formatar_data_br_valor(data_pagamento_entrada),
                        forma,
                        agora_str(),
                    ),
                )
                contrato_id = cursor.lastrowid
                salvar_procedimentos_contrato(contrato_id, procedimentos, valores)
                sincronizar_financeiro_contrato(contrato_id, paciente_nome, total)
                if not forma_pagamento_a_vista(forma):
                    sincronizar_recebiveis_contrato(
                        contrato_id,
                        paciente_id,
                        paciente_nome,
                        prontuario_paciente,
                        total,
                        entrada,
                        parcelas,
                        vencimento.strip(),
                        forma,
                    )
                conn.commit()

                try:
                    arquivo = gerar_documento(conn, contrato_id)
                    st.success(f"Contrato salvo em: {os.path.abspath(arquivo)}")
                except Exception as exc:
                    st.error(f"Contrato salvo, mas houve falha ao gerar o documento: {exc}")

    with tab_contrato_lista:
        contratos_lista = carregar_contratos()
        pacientes_lista = carregar_pacientes()
        if contratos_lista.empty:
            st.info("Nenhum contrato cadastrado ainda.")
        else:
            pacientes_base = pacientes_lista[["id", "nome", "prontuario"]].copy() if not pacientes_lista.empty else pd.DataFrame(columns=["id", "nome", "prontuario"])
            pacientes_base = pacientes_base.rename(columns={"id": "paciente_id", "nome": "paciente_nome", "prontuario": "prontuario"})
            contratos_exibicao = contratos_lista.merge(pacientes_base, on="paciente_id", how="left")
            procedimentos_df = pd.read_sql(
                "SELECT contrato_id, procedimento FROM procedimentos_contrato ORDER BY id",
                conn,
            )
            if not procedimentos_df.empty:
                resumo_procedimentos = (
                    procedimentos_df.groupby("contrato_id")["procedimento"]
                    .apply(lambda serie: ", ".join([str(valor).strip() for valor in serie.tolist() if str(valor).strip()]))
                    .reset_index()
                    .rename(columns={"procedimento": "procedimentos"})
                )
                contratos_exibicao = contratos_exibicao.merge(resumo_procedimentos, left_on="id", right_on="contrato_id", how="left")
            else:
                contratos_exibicao["procedimentos"] = ""

            contratos_exibicao["paciente_nome"] = contratos_exibicao["paciente_nome"].fillna("")
            contratos_exibicao["prontuario"] = contratos_exibicao["prontuario"].apply(formatar_prontuario_valor)
            contratos_exibicao["forma_pagamento"] = contratos_exibicao["forma_pagamento"].fillna("")
            contratos_exibicao["data_criacao_bruta"] = pd.to_datetime(contratos_exibicao["data_criacao"], errors="coerce")

            fc1, fc2, fc3, fc4 = st.columns(4)
            filtro_pacientes_contrato = fc1.multiselect(
                "Pacientes",
                options=sorted([valor for valor in contratos_exibicao["paciente_nome"].unique().tolist() if valor]),
            )
            filtro_prontuarios_contrato = fc2.multiselect(
                "Prontuarios",
                options=sorted([valor for valor in contratos_exibicao["prontuario"].unique().tolist() if valor]),
            )
            filtro_formas_contrato = fc3.multiselect(
                "Formas de pagamento",
                options=sorted([valor for valor in contratos_exibicao["forma_pagamento"].unique().tolist() if valor]),
            )
            datas_contrato = [
                valor.date()
                for valor in contratos_exibicao["data_criacao_bruta"].dropna().sort_values().unique().tolist()
            ]
            data_inicio_contrato = datas_contrato[0] if datas_contrato else None
            data_fim_contrato = datas_contrato[-1] if datas_contrato else None
            filtro_periodo_contrato = fc4.date_input(
                "Periodo de criacao",
                value=(data_inicio_contrato, data_fim_contrato) if data_inicio_contrato and data_fim_contrato else (),
            )

            contratos_filtrados = contratos_exibicao.copy()
            if filtro_pacientes_contrato:
                contratos_filtrados = contratos_filtrados[contratos_filtrados["paciente_nome"].isin(filtro_pacientes_contrato)]
            if filtro_prontuarios_contrato:
                contratos_filtrados = contratos_filtrados[contratos_filtrados["prontuario"].isin(filtro_prontuarios_contrato)]
            if filtro_formas_contrato:
                contratos_filtrados = contratos_filtrados[contratos_filtrados["forma_pagamento"].isin(filtro_formas_contrato)]
            if isinstance(filtro_periodo_contrato, tuple) and len(filtro_periodo_contrato) == 2 and filtro_periodo_contrato[0] and filtro_periodo_contrato[1]:
                inicio_contrato = pd.to_datetime(filtro_periodo_contrato[0])
                fim_contrato = pd.to_datetime(filtro_periodo_contrato[1])
                contratos_filtrados = contratos_filtrados[
                    (contratos_filtrados["data_criacao_bruta"] >= inicio_contrato) &
                    (contratos_filtrados["data_criacao_bruta"] <= fim_contrato)
                ]

            rc1, rc2, rc3 = st.columns(3)
            rc1.metric("Contratos filtrados", str(len(contratos_filtrados)))
            rc2.metric("Valor total", formatar_moeda_br(float(contratos_filtrados["valor_total"].fillna(0).sum())))
            rc3.metric("Entradas", formatar_moeda_br(float(contratos_filtrados["entrada"].fillna(0).sum())))

            lista_exportacao = contratos_filtrados[
                ["id", "paciente_nome", "prontuario", "valor_total", "entrada", "parcelas", "primeiro_vencimento", "data_pagamento_entrada", "forma_pagamento", "data_criacao", "procedimentos"]
            ].copy()
            lista_exportacao = lista_exportacao.rename(
                columns={
                    "id": "Contrato",
                    "paciente_nome": "Paciente",
                    "prontuario": "Prontuario",
                    "valor_total": "Valor total",
                    "entrada": "Entrada",
                    "parcelas": "Parcelas",
                    "primeiro_vencimento": "Primeiro vencimento",
                    "data_pagamento_entrada": "Data pagamento",
                    "forma_pagamento": "Forma pagamento",
                    "data_criacao": "Criado em",
                    "procedimentos": "Procedimentos",
                }
            )
            lista_exportacao["Valor total"] = lista_exportacao["Valor total"].map(formatar_moeda_br)
            lista_exportacao["Entrada"] = lista_exportacao["Entrada"].map(formatar_moeda_br)
            lista_exportacao["Primeiro vencimento"] = lista_exportacao["Primeiro vencimento"].apply(formatar_data_br_valor)
            lista_exportacao["Data pagamento"] = lista_exportacao["Data pagamento"].apply(formatar_data_br_valor)
            lista_exportacao["Criado em"] = lista_exportacao["Criado em"].apply(formatar_data_hora_br_valor)

            if OPENPYXL_DISPONIVEL:
                excel_contratos = dataframe_para_excel_bytes(lista_exportacao, nome_aba="Contratos")
                st.download_button(
                    "Baixar contratos em Excel",
                    data=excel_contratos,
                    file_name="contratos_filtrados.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            else:
                st.info("Excel indisponivel: instale openpyxl para habilitar esta exportacao.")

            st.dataframe(lista_exportacao, use_container_width=True, hide_index=True)

if menu == "Editar Contrato":
    st.title("Editar contrato")
    contratos = carregar_contratos()

    if contratos.empty:
        st.warning("Nenhum contrato encontrado.")
        st.stop()

    pacientes = carregar_pacientes()
    contratos_opcoes = opcoes_contratos(contratos, pacientes)
    busca_contrato = st.text_input(
        "Buscar contrato",
        help="Digite nome do paciente, prontuario, forma de pagamento ou numero do contrato.",
    )
    contratos_filtrados = filtrar_opcoes_contratos(contratos_opcoes, busca_contrato)

    if not contratos_filtrados:
        st.warning("Nenhum contrato encontrado para a busca informada.")
        st.stop()

    contrato_id = st.selectbox(
        "Contrato",
        options=[opcao[0] for opcao in contratos_filtrados],
        format_func=lambda contrato_id: next(
            label for valor, label in contratos_filtrados if valor == contrato_id
        ),
        help="Pesquise pelo nome do paciente, prontuario ou identificacao do contrato.",
    )
    contrato = contratos[contratos["id"] == contrato_id].iloc[0]

    if pacientes.empty:
        st.warning("Nenhum paciente cadastrado.")
        st.stop()

    pacientes_opcoes = opcoes_pacientes(pacientes)
    paciente_id_atual = int(contrato["paciente_id"])
    if paciente_id_atual not in [opcao[0] for opcao in pacientes_opcoes]:
        st.error("O paciente vinculado a este contrato nao foi encontrado.")
        st.stop()

    paciente_id = st.selectbox(
        "Paciente",
        options=[opcao[0] for opcao in pacientes_opcoes],
        index=[opcao[0] for opcao in pacientes_opcoes].index(paciente_id_atual),
        format_func=lambda paciente_id: next(
            label for valor, label in pacientes_opcoes if valor == paciente_id
        ),
    )
    paciente_nome = pacientes.loc[pacientes["id"] == paciente_id, "nome"].iloc[0]
    prontuario_paciente = pacientes.loc[pacientes["id"] == paciente_id, "prontuario"].iloc[0]

    forma_atual = normalizar_forma_pagamento(contrato["forma_pagamento"])
    forma = st.selectbox(
        "Forma pagamento",
        FORMAS_PAGAMENTO,
        index=FORMAS_PAGAMENTO.index(forma_atual),
    )
    entrada = st.number_input("Entrada", min_value=0.0, value=float(contrato["entrada"] or 0))
    data_pagamento_entrada = st.text_input(
        "Data do pagamento" if forma_pagamento_a_vista(forma) else "Data do pagamento da entrada",
        formatar_data_br_valor(contrato["data_pagamento_entrada"] or contrato["data_criacao"] or ""),
    )
    if forma_pagamento_a_vista(forma):
        parcelas = 1
        vencimento = ""
        st.caption("Para Pix, Dinheiro, Debito e Credito o sistema considera pagamento a vista.")
    else:
        parcelas = st.number_input("Parcelas", min_value=1, max_value=36, value=int(contrato["parcelas"] or 1))
        vencimento = st.text_input("Primeiro vencimento", contrato["primeiro_vencimento"] or "")

    procedimentos_db = carregar_procedimentos(contrato_id)
    procedimentos = []
    valores = []

    for i in range(10):
        proc_nome = ""
        proc_valor = 0.0
        if i < len(procedimentos_db):
            proc_nome = procedimentos_db.iloc[i]["procedimento"]
            proc_valor = float(procedimentos_db.iloc[i]["valor"] or 0)

        col1, col2 = st.columns(2)
        proc = col1.text_input(
            f"Procedimento {i + 1}",
            proc_nome,
            key=f"ep_{contrato_id}_{i}",
        )
        valor = col2.number_input(
            f"Valor {i + 1}",
            min_value=0.0,
            value=proc_valor,
            key=f"ev_{contrato_id}_{i}",
        )
        if proc.strip():
            procedimentos.append(proc.strip())
            valores.append(float(valor))

    total = sum(valores)
    st.write("Valor total:", total)

    if st.button("Salvar alteracoes"):
        erros, total = validar_contrato(procedimentos, valores, entrada, parcelas, vencimento, forma, data_pagamento_entrada)
        if erros:
            for erro in erros:
                st.error(erro)
        else:
            cursor.execute(
                """
                UPDATE contratos
                SET paciente_id=?, valor_total=?, entrada=?, parcelas=?,
                    primeiro_vencimento=?, data_pagamento_entrada=?, forma_pagamento=?
                WHERE id=?
                """,
                (
                    paciente_id,
                    total,
                    entrada,
                    parcelas,
                    vencimento.strip(),
                    formatar_data_br_valor(data_pagamento_entrada),
                    forma,
                    contrato_id,
                ),
            )
            salvar_procedimentos_contrato(contrato_id, procedimentos, valores)
            sincronizar_financeiro_contrato(contrato_id, paciente_nome, total)
            if forma_pagamento_a_vista(forma):
                cursor.execute("DELETE FROM recebiveis WHERE contrato_id=?", (contrato_id,))
                status_recebiveis = "removido"
            else:
                status_recebiveis = sincronizar_recebiveis_contrato(
                    contrato_id,
                    paciente_id,
                    paciente_nome,
                    prontuario_paciente,
                    total,
                    entrada,
                    parcelas,
                    vencimento.strip(),
                    forma,
                )
            conn.commit()
            st.success("Contrato atualizado.")
            if status_recebiveis == "atualizado":
                st.warning("Os recebiveis do contrato foram atualizados porque houve mudanca no cronograma.")

            try:
                arquivo = gerar_documento(conn, contrato_id)
                st.success(f"Contrato atualizado e salvo em: {os.path.abspath(arquivo)}")
            except Exception as exc:
                st.error(f"Contrato atualizado, mas houve falha ao gerar o documento: {exc}")

if menu == "Importacoes":
    st.title("Importacoes")
    tab_imp_contratos, tab_imp_recebiveis, tab_imp_pagar, tab_imp_vendas = st.tabs(
        ["Contratos", "Recebiveis", "A pagar", "Vendas"]
    )

    with tab_imp_contratos:
        st.subheader("Importar contratos por Excel")
        st.caption("Use a planilha de contratos para trazer pacientes, contratos, procedimentos e recebiveis para o sistema.")

        arquivo_importacao = st.file_uploader(
            "Enviar planilha de contratos",
            type=["xlsx", "xls"],
            key="arquivo_importacao_contratos",
        )

        if arquivo_importacao is not None:
            try:
                contratos_preparados = preparar_planilha_importacao_contratos(arquivo_importacao)
                if not contratos_preparados:
                    st.warning("Nao encontrei contratos validos para importar nessa planilha.")
                else:
                    preview_df = pd.DataFrame(
                        [
                            {
                                "Paciente": item["nome"],
                                "Prontuario": item["prontuario"],
                                "Data": item["data"],
                                "Forma": item["forma_pagamento"],
                                "Valor total": formatar_moeda_br(item["valor_total"]),
                                "Entrada": formatar_moeda_br(item["entrada"]),
                                "Parcelas": item["parcelas"],
                                "Procedimentos": len(item["procedimentos"]),
                            }
                            for item in contratos_preparados
                        ]
                    )
                    st.success(f"{len(contratos_preparados)} contratos encontrados para importacao.")
                    st.dataframe(preview_df, use_container_width=True, hide_index=True)

                    if st.button("Importar contratos da planilha"):
                        resultado_importacao = importar_contratos_preparados(contratos_preparados)
                        st.success(
                            f"Importacao concluida. Importados: {resultado_importacao['importados']} | Atualizados: {resultado_importacao['atualizados']} | Ignorados: {resultado_importacao['ignorados']}"
                        )
                        if resultado_importacao["recebiveis_pendentes"]:
                            st.warning("Alguns contratos foram importados sem gerar recebiveis porque nao havia primeiro vencimento informado:")
                            for pendencia in resultado_importacao["recebiveis_pendentes"]:
                                st.write(f"- {pendencia}")
                        if resultado_importacao["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado_importacao["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha: {exc}")

    with tab_imp_recebiveis:
        st.subheader("Importar recebiveis por Excel")
        st.caption("Atualiza recebiveis existentes por paciente + vencimento + valor, ou cria recebiveis avulsos quando nao encontrar correspondencia.")
        arquivo_recebiveis = st.file_uploader(
            "Enviar planilha de recebiveis",
            type=["xlsx", "xls"],
            key="arquivo_importacao_recebiveis",
        )
        if arquivo_recebiveis is not None:
            try:
                recebiveis_preparados = preparar_planilha_importacao_recebiveis(arquivo_recebiveis)
                if not recebiveis_preparados:
                    st.warning("Nao encontrei recebiveis validos para importar nessa planilha.")
                else:
                    preview_recebiveis = pd.DataFrame(recebiveis_preparados).rename(
                        columns={
                            "paciente_nome": "Paciente",
                            "prontuario": "Prontuario",
                            "vencimento": "Vencimento",
                            "valor": "Valor",
                            "forma_pagamento": "Forma pagamento",
                            "status": "Status",
                            "data_pagamento": "Data pagamento",
                            "observacao": "Observacao",
                        }
                    )
                    preview_recebiveis["Valor"] = preview_recebiveis["Valor"].map(formatar_moeda_br)
                    st.success(f"{len(recebiveis_preparados)} recebiveis encontrados para importacao.")
                    st.dataframe(
                        preview_recebiveis[["Paciente", "Prontuario", "Vencimento", "Valor", "Forma pagamento", "Status", "Data pagamento", "Observacao"]],
                        use_container_width=True,
                        hide_index=True,
                    )
                    if st.button("Importar recebiveis da planilha"):
                        resultado = importar_recebiveis_preparados(recebiveis_preparados)
                        st.success(
                            f"Importacao concluida. Inseridos: {resultado['inseridos']} | Atualizados: {resultado['atualizados']}"
                        )
                        if resultado["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha de recebiveis: {exc}")

    with tab_imp_pagar:
        st.subheader("Importar contas a pagar por Excel")
        st.caption("Use a planilha do contas a pagar para trazer vencimentos futuros e titulos ja pagos para o sistema.")
        arquivo_pagar = st.file_uploader(
            "Enviar planilha de contas a pagar",
            type=["xlsx", "xls"],
            key="arquivo_importacao_pagar",
        )
        if arquivo_pagar is not None:
            try:
                contas_preparadas = preparar_planilha_importacao_contas_pagar(arquivo_pagar)
                if not contas_preparadas:
                    st.warning("Nao encontrei contas validas para importar nessa planilha.")
                else:
                    preview_pagar = pd.DataFrame(contas_preparadas).rename(
                        columns={
                            "data_vencimento": "Vencimento",
                            "descricao": "Descricao",
                            "fornecedor": "Fornecedor",
                            "valor": "Valor",
                            "pago": "Data pagamento",
                            "valor_pago": "Valor pago",
                            "status": "Status",
                            "observacao": "Observacao",
                        }
                    )
                    preview_pagar["Valor"] = preview_pagar["Valor"].map(formatar_moeda_br)
                    preview_pagar["Valor pago"] = preview_pagar["Valor pago"].map(formatar_moeda_br)
                    st.success(f"{len(contas_preparadas)} contas encontradas para importacao.")
                    st.dataframe(
                        preview_pagar[["Vencimento", "Descricao", "Fornecedor", "Valor", "Data pagamento", "Valor pago", "Status", "Observacao"]],
                        use_container_width=True,
                        hide_index=True,
                    )
                    if st.button("Importar contas a pagar da planilha"):
                        resultado = importar_contas_pagar_preparadas(contas_preparadas)
                        st.success(
                            f"Importacao concluida. Inseridas: {resultado['inseridos']} | Atualizadas: {resultado['atualizados']}"
                        )
                        if resultado["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha do contas a pagar: {exc}")

    with tab_imp_vendas:
        st.subheader("Importar vendas por Excel")
        st.caption("Registra as vendas no dashboard e tenta criar contratos automaticamente quando o paciente ja existe no cadastro.")
        arquivo_vendas = st.file_uploader(
            "Enviar planilha de vendas",
            type=["xlsx", "xls"],
            key="arquivo_importacao_vendas",
        )
        criar_contratos_automaticamente = st.checkbox(
            "Tentar criar contratos automaticamente para pacientes ja cadastrados",
            value=True,
            key="criar_contratos_vendas",
        )
        if arquivo_vendas is not None:
            try:
                vendas_preparadas = preparar_planilha_importacao_vendas(arquivo_vendas)
                if not vendas_preparadas:
                    st.warning("Nao encontrei vendas validas para importar nessa planilha.")
                else:
                    preview_vendas = pd.DataFrame(vendas_preparadas).rename(
                        columns={
                            "data_venda": "Data",
                            "paciente_nome": "Paciente",
                            "valor_total": "Valor total",
                            "valor_a_vista": "A vista",
                            "valor_cartao": "Cartao",
                            "valor_boleto": "Boleto",
                            "saldo": "Saldo",
                            "data_a_pagar": "Data a pagar",
                            "avaliador": "Avaliador",
                            "vendedor": "Vendedor",
                            "nf": "NF",
                        }
                    )
                    for coluna in ["Valor total", "A vista", "Cartao", "Boleto", "Saldo"]:
                        preview_vendas[coluna] = preview_vendas[coluna].map(formatar_moeda_br)
                    st.success(f"{len(vendas_preparadas)} vendas encontradas para importacao.")
                    st.dataframe(
                        preview_vendas[["Data", "Paciente", "Valor total", "A vista", "Cartao", "Boleto", "Saldo", "Data a pagar", "Avaliador", "Vendedor", "NF"]],
                        use_container_width=True,
                        hide_index=True,
                    )
                    if st.button("Importar vendas da planilha"):
                        resultado = importar_vendas_preparadas(vendas_preparadas, criar_contratos_automaticamente)
                        st.success(
                            f"Importacao concluida. Vendas inseridas: {resultado['inseridas']} | Ignoradas: {resultado['ignoradas']} | Contratos criados: {resultado['contratos_criados']}"
                        )
                        if resultado["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha de vendas: {exc}")

if menu == "Usuarios":
    st.title("Usuarios e Acessos")

    if usuario_logado["perfil"] != "Administrador":
        st.error("Apenas administradores podem gerenciar usuarios.")
        st.stop()

    tab_novo_usuario, tab_gerenciar_usuario, tab_logs_acesso = st.tabs(["Novo usuario", "Gerenciar usuarios", "Logs de acesso"])

    with tab_novo_usuario:
        st.subheader("Criar usuario")
        nu1, nu2 = st.columns(2)
        novo_nome = nu1.text_input("Nome do usuario", key="novo_usuario_nome")
        novo_usuario = nu2.text_input("Login", key="novo_usuario_login")
        nu3, nu4 = st.columns(2)
        novo_perfil = nu3.selectbox("Perfil", ["Administrador", "Usuario"], key="novo_usuario_perfil")
        nova_senha = nu4.text_input("Senha inicial", type="password", key="novo_usuario_senha")
        st.markdown("**Permissoes do novo usuario**")
        np1, np2, np3 = st.columns(3)
        novo_acesso_dashboard = np1.checkbox("Dashboard", value=True, key="novo_acesso_dashboard")
        novo_acesso_pacientes = np1.checkbox("Pacientes", value=True, key="novo_acesso_pacientes")
        novo_acesso_contratos = np2.checkbox("Contratos", value=True, key="novo_acesso_contratos")
        novo_acesso_financeiro = np2.checkbox("Financeiro", value=(novo_perfil == "Administrador"), key="novo_acesso_financeiro")
        novo_acesso_usuarios = np3.checkbox("Usuarios", value=(novo_perfil == "Administrador"), key="novo_acesso_usuarios")

        if st.button("Criar usuario"):
            if not novo_nome.strip():
                st.error("Informe o nome do usuario.")
            elif not novo_usuario.strip():
                st.error("Informe o login do usuario.")
            elif len(nova_senha) < 4:
                st.error("A senha inicial deve ter pelo menos 4 caracteres.")
            elif usuario_existe(novo_usuario):
                st.error("Ja existe um usuario com esse login.")
            else:
                criar_usuario(novo_nome, novo_usuario, nova_senha, novo_perfil)
                cursor.execute(
                    """
                    UPDATE usuarios
                    SET acesso_dashboard=?, acesso_pacientes=?, acesso_contratos=?, acesso_financeiro=?, acesso_usuarios=?
                    WHERE lower(usuario)=lower(?)
                    """,
                    (
                        int(bool(novo_acesso_dashboard)),
                        int(bool(novo_acesso_pacientes)),
                        int(bool(novo_acesso_contratos)),
                        int(bool(novo_acesso_financeiro)),
                        int(bool(novo_acesso_usuarios)),
                        novo_usuario.strip(),
                    ),
                )
                conn.commit()
                st.success("Usuario criado com sucesso.")
                st.rerun()

    with tab_gerenciar_usuario:
        st.subheader("Gerenciar usuarios")
        usuarios_df = carregar_usuarios()
        if usuarios_df.empty:
            st.info("Nenhum usuario cadastrado.")
        else:
            opcoes_usuarios = [
                (
                    int(row["id"]),
                    f"{row['nome']} - {row['usuario']} - {row['perfil']} - {'Ativo' if int(row['ativo']) == 1 else 'Inativo'}",
                )
                for _, row in usuarios_df.iterrows()
            ]
            usuario_id_edicao = st.selectbox(
                "Usuario",
                options=[opcao[0] for opcao in opcoes_usuarios],
                format_func=lambda valor: next(label for chave, label in opcoes_usuarios if chave == valor),
                key="usuario_id_edicao",
            )
            usuario_edicao = usuarios_df[usuarios_df["id"] == usuario_id_edicao].iloc[0]

            gu1, gu2 = st.columns(2)
            editar_nome = gu1.text_input("Nome", value=usuario_edicao["nome"] or "", key="editar_usuario_nome")
            editar_login = gu2.text_input("Login", value=usuario_edicao["usuario"] or "", key="editar_usuario_login")
            gu3, gu4 = st.columns(2)
            editar_perfil = gu3.selectbox(
                "Perfil",
                ["Administrador", "Usuario"],
                index=["Administrador", "Usuario"].index(usuario_edicao["perfil"] or "Usuario"),
                key="editar_usuario_perfil",
            )
            editar_ativo = gu4.checkbox(
                "Usuario ativo",
                value=int(usuario_edicao["ativo"] or 0) == 1,
                key="editar_usuario_ativo",
            )
            st.markdown("**Permissoes por modulo**")
            ep1, ep2, ep3 = st.columns(3)
            editar_acesso_dashboard = ep1.checkbox("Dashboard", value=int(usuario_edicao["acesso_dashboard"] or 0) == 1, key="editar_acesso_dashboard")
            editar_acesso_pacientes = ep1.checkbox("Pacientes", value=int(usuario_edicao["acesso_pacientes"] or 0) == 1, key="editar_acesso_pacientes")
            editar_acesso_contratos = ep2.checkbox("Contratos", value=int(usuario_edicao["acesso_contratos"] or 0) == 1, key="editar_acesso_contratos")
            editar_acesso_financeiro = ep2.checkbox("Financeiro", value=int(usuario_edicao["acesso_financeiro"] or 0) == 1, key="editar_acesso_financeiro")
            editar_acesso_usuarios = ep3.checkbox("Usuarios", value=int(usuario_edicao["acesso_usuarios"] or 0) == 1, key="editar_acesso_usuarios")

            if st.button("Salvar dados do usuario"):
                if not editar_nome.strip():
                    st.error("Informe o nome do usuario.")
                elif not editar_login.strip():
                    st.error("Informe o login do usuario.")
                elif usuario_existe(editar_login, ignorar_id=usuario_id_edicao):
                    st.error("Ja existe outro usuario com esse login.")
                else:
                    atualizar_usuario_admin(
                        usuario_id_edicao,
                        editar_nome,
                        editar_login,
                        editar_perfil,
                        editar_ativo,
                        editar_acesso_dashboard,
                        editar_acesso_pacientes,
                        editar_acesso_contratos,
                        editar_acesso_financeiro,
                        editar_acesso_usuarios,
                    )
                    conn.commit()
                    st.success("Usuario atualizado.")
                    st.rerun()

            st.markdown("**Redefinir senha**")
            nova_senha_usuario = st.text_input("Nova senha", type="password", key="redefinir_senha_usuario")
            if st.button("Redefinir senha do usuario"):
                if len(nova_senha_usuario) < 4:
                    st.error("A nova senha deve ter pelo menos 4 caracteres.")
                else:
                    redefinir_senha_usuario(usuario_id_edicao, nova_senha_usuario)
                    conn.commit()
                    st.success("Senha redefinida com sucesso.")

            usuarios_exibicao = usuarios_df.copy()
            usuarios_exibicao["ativo"] = usuarios_exibicao["ativo"].apply(lambda valor: "Ativo" if int(valor or 0) == 1 else "Inativo")
            usuarios_exibicao["data_criacao"] = usuarios_exibicao["data_criacao"].apply(formatar_data_br_valor)
            usuarios_exibicao = usuarios_exibicao.rename(
                columns={
                    "nome": "Nome",
                    "usuario": "Login",
                    "perfil": "Perfil",
                    "ativo": "Status",
                    "data_criacao": "Criado em",
                }
            )[["Nome", "Login", "Perfil", "Status", "Criado em"]]
            st.markdown("**Usuarios cadastrados**")
            st.dataframe(usuarios_exibicao, use_container_width=True, hide_index=True)

    with tab_logs_acesso:
        st.subheader("Logs de acesso")
        logs_df = pd.read_sql("SELECT usuario, evento, data_hora FROM logs_acesso ORDER BY id DESC LIMIT 200", conn)
        if logs_df.empty:
            st.info("Nenhum log de acesso registrado ainda.")
        else:
            logs_df["data_hora"] = logs_df["data_hora"].apply(formatar_data_hora_br_valor)
            logs_df = logs_df.rename(columns={"usuario": "Usuario", "evento": "Evento", "data_hora": "Data"})
            st.dataframe(logs_df, use_container_width=True, hide_index=True)

if menu == "Financeiro":
    st.title("Controle Financeiro")
    preencher_prontuarios_recebiveis()
    atualizar_status_contas_pagar_automaticamente()
    financeiro = pd.read_sql("SELECT * FROM financeiro ORDER BY data DESC", conn)
    saldos_conta = pd.read_sql("SELECT * FROM saldos_conta ORDER BY data DESC, id DESC", conn)
    contas_pagar = pd.read_sql("SELECT * FROM contas_pagar ORDER BY data_vencimento, fornecedor, descricao", conn)
    recebiveis = pd.read_sql(
        "SELECT * FROM recebiveis ORDER BY vencimento, paciente_nome, parcela_numero",
        conn,
    )
    tab_caixa, tab_visao, tab_individual, tab_lote, tab_pagar, tab_calendario_pagar, tab_novo_pagar = st.tabs(
        ["Caixa", "Visao de Recebiveis", "Editar Individual", "Editar em Lote", "Contas a Pagar", "Calendario", "Nova Divida"]
    )

    if not financeiro.empty:
        financeiro["prontuario"] = financeiro["prontuario"].fillna("").replace("None", "")
        financeiro["forma_pagamento"] = financeiro["forma_pagamento"].fillna("").replace("None", "")
        financeiro["conta_caixa"] = financeiro["conta_caixa"].fillna("").replace("None", "")
        financeiro["observacao"] = financeiro["observacao"].fillna("").replace("None", "")
        financeiro["data_exibicao"] = financeiro["data"].apply(formatar_data_br_valor)

    caixa_real = financeiro.copy()
    if not caixa_real.empty:
        caixa_real = caixa_real[
            ~(
                (caixa_real["descricao"].fillna("") == DESCRICAO_CONTRATO)
                & (caixa_real["recebivel_id"].isna())
            )
        ].copy()

    if not recebiveis.empty:
        recebiveis["valor"] = recebiveis["valor"].astype(float)
        recebiveis["prontuario"] = recebiveis["prontuario"].fillna("").replace("None", "")
        recebiveis["observacao"] = recebiveis["observacao"].fillna("").replace("None", "")
        recebiveis["data_pagamento"] = recebiveis["data_pagamento"].fillna("").replace("None", "")
        recebiveis["data_pagamento"] = recebiveis["data_pagamento"].apply(formatar_data_br_valor)
        recebiveis["_ordem_vencimento"] = pd.to_datetime(recebiveis["vencimento"], format="%d/%m/%Y", errors="coerce")
        recebiveis = recebiveis.sort_values(["_ordem_vencimento", "paciente_nome", "parcela_numero"], na_position="last")

    if not contas_pagar.empty:
        contas_pagar["valor"] = contas_pagar["valor"].astype(float)
        contas_pagar["valor_pago"] = contas_pagar["valor_pago"].fillna(0).astype(float)
        contas_pagar["pago"] = contas_pagar["pago"].fillna("").replace("None", "").apply(formatar_data_br_valor)
        contas_pagar["observacao"] = contas_pagar["observacao"].fillna("").replace("None", "")
        contas_pagar["_ordem_vencimento"] = pd.to_datetime(contas_pagar["data_vencimento"], format="%d/%m/%Y", errors="coerce")
        contas_pagar = contas_pagar.sort_values(["_ordem_vencimento", "fornecedor", "descricao"], na_position="last")

    with tab_caixa:
        st.subheader("Lancamento manual")
        cx1, cx2, cx3, cx4 = st.columns(4)
        caixa_data = cx1.date_input("Data do movimento", value=date.today(), key="caixa_data")
        caixa_tipo = cx2.selectbox("Tipo", ["Entrada", "Saida"], key="caixa_tipo")
        caixa_forma = cx3.selectbox("Forma pagamento", ["", *FORMAS_PAGAMENTO], key="caixa_forma")
        caixa_conta = cx4.selectbox("Conta/Banco", CONTAS_CAIXA_MODELO, key="caixa_conta")

        cx5, cx6 = st.columns(2)
        caixa_origem = cx5.text_input("Origem", help="Ex.: banco, pagamento, fornecedor.", key="caixa_origem")
        caixa_prontuario = cx6.text_input("Prontuario", key="caixa_prontuario")

        cx7, cx8 = st.columns(2)
        caixa_descricao = cx7.text_input("Descricao", key="caixa_descricao")
        caixa_valor = cx8.number_input("Valor", min_value=0.0, value=0.0, key="caixa_valor")
        caixa_observacao = st.text_area("Observacao", key="caixa_observacao")

        if st.button("Registrar no caixa"):
            if not caixa_descricao.strip():
                st.error("Informe a descricao do movimento.")
            elif caixa_valor <= 0:
                st.error("Informe um valor maior que zero.")
            else:
                registrar_movimento_caixa(
                    origem=caixa_origem or "Caixa manual",
                    descricao=caixa_descricao,
                    valor=caixa_valor,
                    tipo=caixa_tipo,
                    data_movimento=caixa_data,
                    prontuario=caixa_prontuario,
                    forma_pagamento=caixa_forma,
                    conta_caixa=caixa_conta,
                    observacao=caixa_observacao,
                )
                conn.commit()
                st.success("Movimento registrado no caixa.")
                st.rerun()

        st.divider()
        st.subheader("Saldos do dia anterior")
        st.caption("Use este bloco para registrar os saldos de abertura do dia anterior em cada conta.")
        saldo_data = st.date_input(
            "Data de referencia dos saldos",
            value=date.today() - timedelta(days=1),
            key="saldo_conta_data_anterior",
        )
        saldo_observacao = st.text_input("Observacao dos saldos", key="saldo_conta_observacao")

        col_saldos_1, col_saldos_2, col_saldos_3 = st.columns(3)
        colunas_saldos = [col_saldos_1, col_saldos_2, col_saldos_3]
        saldos_para_registro = {}
        for indice, conta in enumerate(CONTAS_CAIXA_MODELO):
            coluna = colunas_saldos[indice % len(colunas_saldos)]
            saldos_para_registro[conta] = coluna.number_input(
                f"Saldo {conta}",
                value=0.0,
                key=f"saldo_{conta.lower()}_anterior",
            )

        if st.button("Registrar saldos do dia anterior"):
            for conta, saldo_valor in saldos_para_registro.items():
                registrar_saldo_conta(
                    data_saldo=saldo_data,
                    conta=conta,
                    saldo=saldo_valor,
                    observacao=saldo_observacao,
                )
            conn.commit()
            st.success("Saldos do dia anterior registrados.")
            st.rerun()

        if not saldos_conta.empty:
            st.markdown("**Ultimos saldos informados**")
            saldos_exibicao = saldos_conta.copy()
            saldos_exibicao["data"] = saldos_exibicao["data"].apply(formatar_data_br_valor)
            saldos_exibicao["saldo"] = saldos_exibicao["saldo"].map(formatar_moeda_br)
            saldos_exibicao = saldos_exibicao.rename(
                columns={
                    "data": "Data",
                    "conta": "Conta",
                    "saldo": "Saldo",
                    "observacao": "Observacao",
                }
            )[["Data", "Conta", "Saldo", "Observacao"]]
            st.dataframe(saldos_exibicao.head(20), use_container_width=True, hide_index=True)

        st.divider()
        st.subheader("Baixar recebivel no caixa")
        if recebiveis.empty:
            st.info("Nao ha recebiveis cadastrados.")
        else:
            recebiveis_abertos = recebiveis[recebiveis["status"].isin(["Aberto", "Atrasado"])].copy()
            if recebiveis_abertos.empty:
                st.info("Nao ha recebiveis em aberto para baixa.")
            else:
                filtro_baixa = st.text_input(
                    "Buscar recebivel para baixa",
                    help="Pesquise por paciente, prontuario, parcela ou vencimento.",
                    key="busca_baixa_recebivel",
                )
                if filtro_baixa.strip():
                    termo_baixa = normalizar_texto(filtro_baixa)
                    recebiveis_abertos = recebiveis_abertos[
                        recebiveis_abertos.apply(
                            lambda row: termo_baixa in normalizar_texto(
                                f"{row['paciente_nome']} {row['prontuario']} {row['parcela_numero']} {row['vencimento']}"
                            ),
                            axis=1,
                        )
                    ]

                if recebiveis_abertos.empty:
                    st.warning("Nenhum recebivel encontrado para a busca informada.")
                else:
                    opcoes_baixa = [
                        (
                            int(row["id"]),
                            f"{row['paciente_nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])} - Parcela {formatar_parcela_valor(row['parcela_numero'])} - {row['vencimento']} - R$ {float(row['valor']):.2f}"
                        )
                        for _, row in recebiveis_abertos.iterrows()
                    ]
                    recebivel_baixa_id = st.selectbox(
                        "Recebivel para dar baixa",
                        options=[opcao[0] for opcao in opcoes_baixa],
                        format_func=lambda recebivel_id: next(
                            label for valor, label in opcoes_baixa if valor == recebivel_id
                        ),
                        key="recebivel_baixa_id",
                    )
                    recebivel_baixa = recebiveis_abertos[recebiveis_abertos["id"] == recebivel_baixa_id].iloc[0]

                    bx1, bx2, bx3, bx4 = st.columns(4)
                    baixa_data = bx1.date_input("Data do pagamento", value=date.today(), key="baixa_data")
                    baixa_origem = bx2.text_input("Origem da entrada", value="Pagamento", key="baixa_origem")
                    baixa_forma = bx3.selectbox(
                        "Forma de recebimento",
                        FORMAS_PAGAMENTO,
                        index=FORMAS_PAGAMENTO.index(normalizar_forma_pagamento(recebivel_baixa["forma_pagamento"])),
                        key="baixa_forma",
                    )
                    baixa_conta = bx4.selectbox("Conta/Banco da entrada", CONTAS_CAIXA_MODELO, key="baixa_conta")
                    baixa_observacao = st.text_area("Observacao da baixa", key="baixa_observacao")

                    if st.button("Dar baixa no recebivel"):
                        try:
                            baixar_recebivel_no_caixa(
                                recebivel_id=recebivel_baixa_id,
                                origem=baixa_origem,
                                data_pagamento=baixa_data,
                                observacao=baixa_observacao,
                                forma_recebimento=baixa_forma,
                                conta_caixa=baixa_conta,
                            )
                            conn.commit()
                            st.success("Recebivel baixado no caixa e marcado como pago.")
                            st.rerun()
                        except ValueError as exc:
                            st.error(str(exc))

        st.divider()
        st.subheader("Livro-caixa")
        st.caption("O caixa mostra apenas movimentacoes efetivas que passaram pela conta ou foram registradas manualmente. Previsoes e contratos futuros nao entram aqui.")
        if caixa_real.empty:
            st.info("Nenhum lancamento no caixa ainda.")
        else:
            resumo_caixa_df, grupos_caixa = montar_caixa_diario(caixa_real)

            st.markdown("**Resumo diario**")
            resumo_exibicao = resumo_caixa_df.copy()
            for coluna in ["Entradas", "Saidas", "Saldo do dia"]:
                resumo_exibicao[coluna] = resumo_exibicao[coluna].map(formatar_moeda_br)
            st.dataframe(resumo_exibicao, use_container_width=True, hide_index=True)

            st.markdown("**Lancamentos por dia**")
            for grupo in grupos_caixa:
                data_referencia = parse_data_contrato(grupo["data"])
                saldos_abertura = saldos_informados_por_conta_ate(data_referencia) if data_referencia else {conta: 0.0 for conta in CONTAS_CAIXA_MODELO}
                saldos_fechamento = saldos_abertura.copy()
                detalhamento_para_saldo = caixa_real.copy()
                detalhamento_para_saldo["data_date"] = detalhamento_para_saldo["data"].apply(parse_data_contrato)
                detalhamento_para_saldo["conta_resolvida"] = detalhamento_para_saldo.apply(
                    lambda row: str(row.get("conta_caixa") or "").strip().upper()
                    or identificar_conta_caixa(row.get("origem", ""), row.get("descricao", ""), row.get("observacao", "")),
                    axis=1,
                )
                if data_referencia:
                    movimentos_dia = detalhamento_para_saldo[detalhamento_para_saldo["data_date"] == data_referencia]
                    for _, mov in movimentos_dia.iterrows():
                        conta = mov["conta_resolvida"] if mov["conta_resolvida"] in CONTAS_CAIXA_MODELO else "CAIXA"
                        valor = float(mov["valor"] or 0)
                        if str(mov["tipo"]) == "Entrada":
                            saldos_fechamento[conta] += valor
                        else:
                            saldos_fechamento[conta] -= valor

                st.markdown(f"### {grupo['data']}")
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("Entradas do dia", formatar_moeda_br(grupo["entradas"]))
                m2.metric("Saidas do dia", formatar_moeda_br(grupo["saidas"]))
                m3.metric("Saldo do dia", formatar_moeda_br(grupo["saldo"]))
                m4.metric("Lancamentos", len(grupo["detalhamento"]))

                st.markdown("**Saldos de abertura**")
                resumo_abertura_df = montar_resumo_saldos_conta(saldos_abertura, data_referencia)
                st.dataframe(
                    estilizar_resumo_saldos(resumo_abertura_df),
                    use_container_width=True,
                    hide_index=True,
                )

                detalhamento_exibicao = grupo["detalhamento"].copy()
                detalhamento_exibicao["Valor"] = detalhamento_exibicao["Valor"].map(formatar_moeda_br)
                st.dataframe(detalhamento_exibicao, use_container_width=True, hide_index=True)

                st.markdown("**Saldos de fechamento**")
                resumo_fechamento_df = montar_resumo_saldos_conta(saldos_fechamento, data_referencia)
                st.dataframe(
                    estilizar_resumo_saldos(resumo_fechamento_df),
                    use_container_width=True,
                    hide_index=True,
                )

            exp1, exp2 = st.columns(2)
            if OPENPYXL_DISPONIVEL:
                exp1.download_button(
                    "Exportar caixa para Excel",
                    data=caixa_diario_para_excel_bytes(caixa_real),
                    file_name=f"caixa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            else:
                exp1.info("Excel indisponivel: instale openpyxl para habilitar esta exportacao.")
            if REPORTLAB_DISPONIVEL:
                caixa_pdf_df = caixa_real[
                    ["data_exibicao", "tipo", "origem", "conta_caixa", "prontuario", "descricao", "valor", "forma_pagamento", "observacao"]
                ].copy()
                caixa_pdf_df = caixa_pdf_df.rename(
                    columns={
                        "data_exibicao": "Data",
                        "tipo": "Tipo",
                        "origem": "Origem",
                        "conta_caixa": "Conta/Banco",
                        "prontuario": "Prontuario",
                        "descricao": "Descricao",
                        "valor": "Valor",
                        "forma_pagamento": "Forma de pagamento",
                        "observacao": "Observacao",
                    }
                )
                caixa_pdf_df["Valor"] = caixa_pdf_df["Valor"].map(formatar_moeda_br)
                exp2.download_button(
                    "Exportar caixa para PDF",
                    data=dataframe_para_pdf_bytes(caixa_pdf_df, titulo="Livro-caixa"),
                    file_name=f"caixa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                )
            else:
                exp2.info("PDF indisponivel: instale reportlab para habilitar esta exportacao.")

        st.divider()
        st.subheader("Importar extrato bancario")
        st.caption("Reconhecimento automatico disponivel para CSV e Excel (.xlsx/.xls). PDF bancario fica como proxima etapa.")
        arquivo_extrato = st.file_uploader(
            "Enviar extrato",
            type=["csv", "xlsx", "xls"],
            key="arquivo_extrato_caixa",
        )
        if arquivo_extrato is not None:
            try:
                extrato_origem = ler_extrato_arquivo(arquivo_extrato)
                extrato_reconhecido = reconhecer_extrato(extrato_origem)
                if extrato_reconhecido.empty:
                    st.error("Nao foi possivel reconhecer automaticamente esse extrato. Verifique se o arquivo tem colunas de data, descricao e valor.")
                else:
                    st.success(f"{len(extrato_reconhecido)} transacoes reconhecidas.")
                    preview_extrato = extrato_reconhecido.rename(
                        columns={
                            "data_exibicao": "Data",
                            "descricao": "Descricao",
                            "valor": "Valor",
                            "tipo": "Tipo",
                        }
                    )[["Data", "Descricao", "Valor", "Tipo"]].copy()
                    preview_extrato["Valor"] = preview_extrato["Valor"].map(lambda valor: f"R$ {float(valor):.2f}")
                    st.dataframe(preview_extrato, use_container_width=True, hide_index=True)

                    ie1, ie2, ie3 = st.columns(3)
                    origem_extrato = ie1.text_input("Origem para importacao", value="Extrato bancario", key="origem_import_extrato")
                    prontuario_extrato = ie2.text_input("Prontuario padrao (opcional)", key="prontuario_import_extrato")
                    forma_extrato = ie3.selectbox("Forma pagamento padrao", ["", *FORMAS_PAGAMENTO], key="forma_import_extrato")
                    observacao_extrato = st.text_area("Observacao padrao da importacao", key="obs_import_extrato")

                    if st.button("Importar extrato para o caixa"):
                        for _, row in extrato_reconhecido.iterrows():
                            registrar_movimento_caixa(
                                origem=origem_extrato or "Extrato bancario",
                                descricao=row["descricao"],
                                valor=float(row["valor"]),
                                tipo=row["tipo"],
                                data_movimento=row["data"].date() if hasattr(row["data"], "date") else row["data"],
                                prontuario=prontuario_extrato,
                                forma_pagamento=forma_extrato,
                                observacao=observacao_extrato,
                            )
                        conn.commit()
                        st.success("Extrato importado para o caixa.")
                        st.rerun()
            except Exception as exc:
                st.error(f"Falha ao ler o extrato: {exc}")

    with tab_visao:
        st.subheader("Recebiveis")
        col_filtro_1, col_filtro_2 = st.columns(2)
        filtro_nomes = col_filtro_1.multiselect(
            "Pacientes",
            options=sorted(recebiveis["paciente_nome"].fillna("").unique().tolist()),
        )
        filtro_prontuarios = col_filtro_2.multiselect(
            "Prontuarios",
            options=sorted([valor for valor in recebiveis["prontuario"].fillna("").unique().tolist() if valor]),
        )

        col_filtro_4, col_filtro_5, col_filtro_6 = st.columns(3)
        opcoes_status = sorted(recebiveis["status"].fillna("Aberto").unique().tolist())
        filtro_status = col_filtro_4.multiselect("Status", options=opcoes_status)
        opcoes_forma = sorted(recebiveis["forma_pagamento"].fillna("").unique().tolist())
        filtro_formas = col_filtro_5.multiselect("Formas de pagamento", options=opcoes_forma)
        datas_disponiveis = [
            valor.date()
            for valor in recebiveis["_ordem_vencimento"].dropna().sort_values().unique().tolist()
        ]
        data_inicio = datas_disponiveis[0] if datas_disponiveis else None
        data_fim = datas_disponiveis[-1] if datas_disponiveis else None
        filtro_periodo = col_filtro_6.date_input(
            "Periodo de vencimento",
            value=(data_inicio, data_fim) if data_inicio and data_fim else (),
        )
        mapa_meses = {
            1: "Janeiro", 2: "Fevereiro", 3: "Marco", 4: "Abril", 5: "Maio", 6: "Junho",
            7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
        }
        col_filtro_7, col_filtro_8 = st.columns(2)
        anos_disponiveis_recebiveis = sorted(
            recebiveis["_ordem_vencimento"].dropna().dt.year.unique().tolist()
        ) if not recebiveis.empty else []
        filtro_anos_recebiveis = col_filtro_7.multiselect(
            "Anos",
            options=anos_disponiveis_recebiveis,
        )
        filtro_meses_recebiveis = col_filtro_8.multiselect(
            "Meses",
            options=list(mapa_meses.keys()),
            format_func=lambda valor: mapa_meses.get(valor, str(valor)),
        )

        recebiveis_filtrados = recebiveis.copy()
        if filtro_nomes:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["paciente_nome"].isin(filtro_nomes)]
        if filtro_prontuarios:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["prontuario"].isin(filtro_prontuarios)]
        if filtro_status:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["status"].isin(filtro_status)]
        if filtro_formas:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["forma_pagamento"].isin(filtro_formas)]
        if isinstance(filtro_periodo, tuple) and len(filtro_periodo) == 2 and filtro_periodo[0] and filtro_periodo[1]:
            inicio = pd.to_datetime(filtro_periodo[0])
            fim = pd.to_datetime(filtro_periodo[1])
            recebiveis_filtrados = recebiveis_filtrados[
                (recebiveis_filtrados["_ordem_vencimento"] >= inicio) &
                (recebiveis_filtrados["_ordem_vencimento"] <= fim)
            ]
        if filtro_anos_recebiveis:
            recebiveis_filtrados = recebiveis_filtrados[
                recebiveis_filtrados["_ordem_vencimento"].dt.year.isin(filtro_anos_recebiveis)
            ]
        if filtro_meses_recebiveis:
            recebiveis_filtrados = recebiveis_filtrados[
                recebiveis_filtrados["_ordem_vencimento"].dt.month.isin(filtro_meses_recebiveis)
            ]

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Total filtrado", f"R$ {recebiveis_filtrados['valor'].sum():.2f}")
        c2.metric("Parcelas filtradas", str(len(recebiveis_filtrados)))
        c3.metric("Vencimentos unicos", str(recebiveis_filtrados['vencimento'].nunique()))
        c4.metric("Pacientes unicos", str(recebiveis_filtrados['paciente_nome'].nunique()))

        if not recebiveis_filtrados.empty:
            agenda_mensal = recebiveis_filtrados.copy()
            agenda_mensal["ano"] = agenda_mensal["_ordem_vencimento"].dt.year
            agenda_mensal["mes"] = agenda_mensal["_ordem_vencimento"].dt.month
            agenda_mensal["mes_nome"] = agenda_mensal["mes"].map(mapa_meses)
            resumo_por_mes = (
                agenda_mensal.groupby(["ano", "mes", "mes_nome"], as_index=False)
                .agg(
                    total_previsto=("valor", "sum"),
                    titulos=("id", "count"),
                )
                .sort_values(["ano", "mes"], ascending=[True, True])
            )
            st.write("Agenda mensal de vencimentos")
            for _, resumo_mes in resumo_por_mes.iterrows():
                ano_agenda = int(resumo_mes["ano"])
                mes_agenda = int(resumo_mes["mes"])
                mes_nome = resumo_mes["mes_nome"]
                total_mes = float(resumo_mes["total_previsto"] or 0)
                titulos_mes = int(resumo_mes["titulos"] or 0)
                itens_mes = agenda_mensal[
                    (agenda_mensal["ano"] == ano_agenda) & (agenda_mensal["mes"] == mes_agenda)
                ].copy()
                resumo_dias = (
                    itens_mes.groupby("vencimento", as_index=False)
                    .agg(
                        total_previsto=("valor", "sum"),
                        titulos=("id", "count"),
                    )
                )
                resumo_dias["_ordem_vencimento"] = pd.to_datetime(
                    resumo_dias["vencimento"],
                    format="%d/%m/%Y",
                    errors="coerce",
                )
                resumo_dias = resumo_dias.sort_values("_ordem_vencimento", na_position="last").drop(columns=["_ordem_vencimento"])
                resumo_dias["total_previsto"] = resumo_dias["total_previsto"].map(formatar_moeda_br)
                resumo_dias = resumo_dias.rename(
                    columns={
                        "vencimento": "Vencimento",
                        "total_previsto": "Total previsto",
                        "titulos": "Titulos",
                    }
                )

                with st.expander(f"{mes_nome}/{ano_agenda}  |  Total previsto: {formatar_moeda_br(total_mes)}  |  Titulos: {titulos_mes}", expanded=False):
                    st.dataframe(resumo_dias, use_container_width=True, hide_index=True)

            detalhe = recebiveis_filtrados[
                ["id", "paciente_nome", "prontuario", "parcela_numero", "vencimento", "data_pagamento", "valor", "forma_pagamento", "status", "observacao"]
            ].copy()
            detalhe = detalhe.rename(
                columns={
                    "id": "ID",
                    "paciente_nome": "Paciente",
                    "prontuario": "Prontuario",
                    "parcela_numero": "Parcela",
                    "vencimento": "Vencimento",
                    "data_pagamento": "Data do pagamento",
                    "valor": "Valor",
                    "forma_pagamento": "Forma pagamento",
                    "status": "Status",
                    "observacao": "Observacao",
                }
            )
            detalhe = detalhe.reset_index(drop=True)
            detalhe["Valor"] = detalhe["Valor"].map(lambda valor: f"R$ {float(valor):.2f}")
            st.write("Detalhamento dos recebiveis")
            st.dataframe(detalhe, use_container_width=True, hide_index=True)

    with tab_individual:
        st.subheader("Editar recebivel individual")
        if recebiveis.empty:
            st.info("Nao ha recebiveis para editar.")
        else:
            opcoes_recebiveis = [
                (
                    int(row["id"]),
                    f"{row['paciente_nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])} - Parcela {formatar_parcela_valor(row['parcela_numero'])} - {row['vencimento']}"
                )
                for _, row in recebiveis.iterrows()
            ]
            recebivel_id = st.selectbox(
                "Recebivel",
                options=[opcao[0] for opcao in opcoes_recebiveis],
                format_func=lambda recebivel_id: next(
                    label for valor, label in opcoes_recebiveis if valor == recebivel_id
                ),
            )
            recebivel = recebiveis[recebiveis["id"] == recebivel_id].iloc[0]

            edit_col_1, edit_col_2 = st.columns(2)
            edit_nome = edit_col_1.text_input("Nome do paciente", recebivel["paciente_nome"] or "", key=f"rec_nome_{recebivel_id}")
            edit_prontuario = edit_col_2.text_input("Prontuario", recebivel["prontuario"] or "", key=f"rec_pront_{recebivel_id}")

            edit_col_3, edit_col_4, edit_col_5 = st.columns(3)
            edit_vencimento = edit_col_3.text_input("Vencimento", recebivel["vencimento"] or "", key=f"rec_venc_{recebivel_id}")
            edit_valor = edit_col_4.number_input(
                "Valor",
                min_value=0.0,
                value=float(recebivel["valor"] or 0),
                key=f"rec_valor_{recebivel_id}",
            )
            edit_forma = edit_col_5.selectbox(
                "Forma pagamento",
                FORMAS_PAGAMENTO,
                index=FORMAS_PAGAMENTO.index(normalizar_forma_pagamento(recebivel["forma_pagamento"])),
                key=f"rec_forma_{recebivel_id}",
            )
            edit_observacao = st.text_area(
                "Observacao",
                recebivel["observacao"] or "",
                key=f"rec_obs_{recebivel_id}",
            )
            edit_status = st.selectbox(
                "Status do recebivel",
                ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"],
                index=["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"].index((recebivel["status"] or "Aberto") if (recebivel["status"] or "Aberto") in ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"] else "Aberto"),
                key=f"rec_status_{recebivel_id}",
            )

            if st.button("Salvar alteracoes do recebivel"):
                if not edit_nome.strip():
                    st.error("Informe o nome do paciente.")
                elif not edit_prontuario.strip():
                    st.error("Informe o prontuario.")
                elif not parse_data_contrato(edit_vencimento):
                    st.error("Informe o vencimento no formato DD/MM/AAAA.")
                else:
                    atualizar_recebivel(
                        recebivel_id,
                        edit_nome,
                        edit_prontuario,
                        edit_vencimento,
                        edit_valor,
                        edit_forma,
                        edit_status,
                        edit_observacao,
                    )
                    conn.commit()
                    st.success("Recebivel atualizado.")

    with tab_lote:
        st.subheader("Editar recebiveis em lote")
        if recebiveis.empty:
            st.info("Nao ha recebiveis para editar em lote.")
        else:
            recebiveis_com_contrato = recebiveis[recebiveis["contrato_id"].notna()].copy()
            recebiveis_avulsos = recebiveis[recebiveis["contrato_id"].isna()].copy()

            if not recebiveis_avulsos.empty:
                st.caption("Recebiveis sem contrato vinculado continuam no controle financeiro e podem ser alterados na aba Editar Individual.")

            if recebiveis_com_contrato.empty:
                st.info("Nao ha lotes vinculados a contrato para editar em lote.")
            else:
                lotes_recebiveis = (
                    recebiveis_com_contrato.groupby(["contrato_id", "paciente_nome", "prontuario"], as_index=False)
                    .agg(
                        primeiro_vencimento=("vencimento", "first"),
                        quantidade_parcelas=("id", "count"),
                    )
                    .sort_values(["paciente_nome", "prontuario", "primeiro_vencimento"])
                )
                opcoes_lotes_recebiveis = [
                    (
                        int(row["contrato_id"]),
                        f"{row['paciente_nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])} - {int(row['quantidade_parcelas'])} parcelas - inicio {row['primeiro_vencimento']}"
                    )
                    for _, row in lotes_recebiveis.iterrows()
                ]
                contrato_lote_id = st.selectbox(
                    "Lote para editar",
                    options=[opcao[0] for opcao in opcoes_lotes_recebiveis],
                    format_func=lambda contrato_id: next(
                        label for valor, label in opcoes_lotes_recebiveis if valor == contrato_id
                    ),
                )
                recebiveis_lote = recebiveis_com_contrato[recebiveis_com_contrato["contrato_id"] == contrato_lote_id].copy()
                if recebiveis_lote.empty:
                    st.warning("Nao encontrei recebiveis para este lote. Selecione outro contrato.")
                else:
                    recebivel_base = recebiveis_lote.iloc[0]

                    lote_col_1, lote_col_2 = st.columns(2)
                    lote_nome = lote_col_1.text_input("Nome do paciente", recebivel_base["paciente_nome"] or "", key=f"lote_nome_{contrato_lote_id}")
                    lote_prontuario = lote_col_2.text_input("Prontuario", recebivel_base["prontuario"] or "", key=f"lote_pront_{contrato_lote_id}")

                    lote_col_3, lote_col_4, lote_col_5 = st.columns(3)
                    lote_primeiro_vencimento = lote_col_3.text_input(
                        "Novo primeiro vencimento (opcional)",
                        recebivel_base["vencimento"] or "",
                        key=f"lote_venc_{contrato_lote_id}",
                    )
                    lote_forma = lote_col_4.selectbox(
                        "Forma pagamento",
                        FORMAS_PAGAMENTO,
                        index=FORMAS_PAGAMENTO.index(normalizar_forma_pagamento(recebivel_base["forma_pagamento"])),
                        key=f"lote_forma_{contrato_lote_id}",
                    )
                    lote_status = lote_col_5.selectbox(
                        "Status",
                        ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"],
                        index=["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"].index((recebivel_base["status"] or "Aberto") if (recebivel_base["status"] or "Aberto") in ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"] else "Aberto"),
                        key=f"lote_status_{contrato_lote_id}",
                    )
                    lote_observacao = st.text_area(
                        "Observacao para todos os recebiveis deste contrato",
                        recebivel_base["observacao"] or "",
                        key=f"lote_obs_{contrato_lote_id}",
                    )

                    st.write("Recebiveis deste contrato")
                    st.dataframe(
                        recebiveis_lote[
                            ["paciente_nome", "prontuario", "parcela_numero", "vencimento", "valor", "forma_pagamento", "status"]
                        ].rename(
                            columns={
                                "paciente_nome": "Paciente",
                                "prontuario": "Prontuario",
                                "parcela_numero": "Parcela",
                                "vencimento": "Vencimento",
                                "valor": "Valor",
                                "forma_pagamento": "Forma pagamento",
                                "status": "Status",
                            }
                        ),
                        use_container_width=True,
                    )

                    if st.button("Salvar alteracoes em lote"):
                        if not lote_nome.strip():
                            st.error("Informe o nome do paciente.")
                        elif not lote_prontuario.strip():
                            st.error("Informe o prontuario.")
                        elif lote_primeiro_vencimento.strip() and not parse_data_contrato(lote_primeiro_vencimento):
                            st.error("Informe o primeiro vencimento no formato DD/MM/AAAA.")
                        else:
                            atualizar_recebiveis_lote_contrato(
                                contrato_lote_id,
                                lote_nome,
                                lote_prontuario,
                                lote_forma,
                                lote_status,
                                lote_observacao,
                                lote_primeiro_vencimento,
                            )
                            conn.commit()
                            st.success("Recebiveis do contrato atualizados em lote.")

    with tab_pagar:
        st.subheader("Contas a pagar")
        if contas_pagar.empty:
            st.info("Nao ha contas a pagar cadastradas.")
        else:
            fp1, fp2, fp3 = st.columns(3)
            filtro_fornecedor = fp1.multiselect(
                "Fornecedores",
                options=sorted([valor for valor in contas_pagar["fornecedor"].fillna("").unique().tolist() if valor]),
            )
            filtro_status_pagar = fp2.multiselect(
                "Status",
                options=sorted([valor for valor in contas_pagar["status"].fillna("A vencer").unique().tolist() if valor]),
            )
            datas_pagar = [
                valor.date()
                for valor in contas_pagar["_ordem_vencimento"].dropna().sort_values().unique().tolist()
            ]
            pagar_inicio = datas_pagar[0] if datas_pagar else None
            pagar_fim = datas_pagar[-1] if datas_pagar else None
            filtro_periodo_pagar = fp3.date_input(
                "Periodo de vencimento",
                value=(pagar_inicio, pagar_fim) if pagar_inicio and pagar_fim else (),
            )

            pagar_filtrado = contas_pagar.copy()
            if filtro_fornecedor:
                pagar_filtrado = pagar_filtrado[pagar_filtrado["fornecedor"].isin(filtro_fornecedor)]
            if filtro_status_pagar:
                pagar_filtrado = pagar_filtrado[pagar_filtrado["status"].isin(filtro_status_pagar)]
            if isinstance(filtro_periodo_pagar, tuple) and len(filtro_periodo_pagar) == 2 and filtro_periodo_pagar[0] and filtro_periodo_pagar[1]:
                inicio = pd.to_datetime(filtro_periodo_pagar[0])
                fim = pd.to_datetime(filtro_periodo_pagar[1])
                pagar_filtrado = pagar_filtrado[
                    (pagar_filtrado["_ordem_vencimento"] >= inicio) &
                    (pagar_filtrado["_ordem_vencimento"] <= fim)
                ]

            p1, p2, p3, p4 = st.columns(4)
            p1.metric("Total filtrado", formatar_moeda_br(float(pagar_filtrado["valor"].sum())))
            p2.metric("Titulos", str(len(pagar_filtrado)))
            p3.metric("Fornecedores", str(pagar_filtrado["fornecedor"].nunique()))
            p4.metric("Pagos", formatar_moeda_br(float(pagar_filtrado["valor_pago"].sum())))

            detalhe_pagar = pagar_filtrado[
                ["data_vencimento", "descricao", "fornecedor", "valor", "pago", "valor_pago", "status", "observacao"]
            ].copy()
            detalhe_pagar = detalhe_pagar.rename(
                columns={
                    "data_vencimento": "Vencimento",
                    "descricao": "Descricao",
                    "fornecedor": "Fornecedor",
                    "valor": "Valor",
                    "pago": "Data pagamento",
                    "valor_pago": "Valor pago",
                    "status": "Status",
                    "observacao": "Observacao",
                }
            )
            detalhe_pagar["Valor"] = detalhe_pagar["Valor"].map(formatar_moeda_br)
            detalhe_pagar["Valor pago"] = detalhe_pagar["Valor pago"].map(formatar_moeda_br)
            st.dataframe(detalhe_pagar, use_container_width=True, hide_index=True)

            st.markdown("**Atualizacao rapida**")
            opcoes_pagar = [
                (int(row["id"]), f"{row['data_vencimento']} - {row['fornecedor']} - {row['descricao']} - {formatar_moeda_br(row['valor'])}")
                for _, row in pagar_filtrado.iterrows()
            ]
            if opcoes_pagar:
                conta_pagar_id = st.selectbox(
                    "Titulo para atualizar",
                    options=[opcao[0] for opcao in opcoes_pagar],
                    format_func=lambda valor: next(rotulo for chave, rotulo in opcoes_pagar if chave == valor),
                )
                conta_pagar_row = contas_pagar[contas_pagar["id"] == conta_pagar_id].iloc[0]
                ap1, ap2 = st.columns(2)
                novo_titulo_pagar = ap1.text_input(
                    "Titulo / descricao",
                    conta_pagar_row["descricao"] or "",
                    key=f"descricao_pagar_{conta_pagar_id}",
                )
                novo_fornecedor_pagar = ap2.text_input(
                    "Fornecedor",
                    conta_pagar_row["fornecedor"] or "",
                    key=f"fornecedor_pagar_{conta_pagar_id}",
                )

                ap3, ap4, ap5 = st.columns(3)
                nova_data_vencimento = ap3.text_input(
                    "Vencimento",
                    formatar_data_br_valor(conta_pagar_row["data_vencimento"] or ""),
                    key=f"vencimento_pagar_{conta_pagar_id}",
                )
                novo_valor_titulo = ap4.number_input(
                    "Valor do titulo",
                    min_value=0.0,
                    value=float(conta_pagar_row["valor"] or 0),
                    key=f"valor_titulo_pagar_{conta_pagar_id}",
                )
                novo_status_pagar = ap5.selectbox(
                    "Status",
                    STATUS_CONTAS_PAGAR,
                    index=STATUS_CONTAS_PAGAR.index(conta_pagar_row["status"]) if conta_pagar_row["status"] in STATUS_CONTAS_PAGAR else 0,
                    key=f"status_pagar_{conta_pagar_id}",
                )

                ap6, ap7 = st.columns(2)
                nova_data_pagamento = ap6.text_input(
                    "Data do pagamento",
                    formatar_data_br_valor(conta_pagar_row["pago"] or ""),
                    key=f"data_pagamento_pagar_{conta_pagar_id}",
                )
                novo_valor_pago = ap7.number_input(
                    "Valor pago",
                    min_value=0.0,
                    value=float(conta_pagar_row["valor_pago"] or 0),
                    key=f"valor_pago_pagar_{conta_pagar_id}",
                )

                nova_obs_pagar = st.text_area(
                    "Observacao",
                    conta_pagar_row["observacao"] or "",
                    key=f"obs_pagar_{conta_pagar_id}",
                )
                if st.button("Salvar conta a pagar"):
                    if not novo_titulo_pagar.strip():
                        st.error("Informe o titulo da conta.")
                    elif not novo_fornecedor_pagar.strip():
                        st.error("Informe o fornecedor.")
                    elif not parse_data_contrato(nova_data_vencimento):
                        st.error("Informe o vencimento no formato DD/MM/AAAA.")
                    elif nova_data_pagamento.strip() and not parse_data_contrato(nova_data_pagamento):
                        st.error("Informe a data do pagamento no formato DD/MM/AAAA.")
                    else:
                        cursor.execute(
                            """
                            UPDATE contas_pagar
                            SET data_vencimento=?, descricao=?, fornecedor=?, valor=?, status=?, pago=?, valor_pago=?, observacao=?
                            WHERE id=?
                            """,
                            (
                                formatar_data_br_valor(nova_data_vencimento),
                                novo_titulo_pagar.strip(),
                                novo_fornecedor_pagar.strip(),
                                float(novo_valor_titulo),
                                novo_status_pagar,
                                formatar_data_br_valor(nova_data_pagamento),
                                float(novo_valor_pago),
                                nova_obs_pagar.strip(),
                                int(conta_pagar_id),
                            ),
                        )
                        conn.commit()
                        st.success("Conta a pagar atualizada.")
                        st.rerun()

    with tab_calendario_pagar:
        st.subheader("Calendario de vencimentos")
        if contas_pagar.empty:
            st.info("Nao ha contas a pagar para exibir no calendario.")
        else:
            contas_validas = contas_pagar[contas_pagar["_ordem_vencimento"].notna()].copy()
            if contas_validas.empty:
                st.info("Nao ha vencimentos validos para exibir no calendario.")
            else:
                ano_cal = st.selectbox(
                    "Ano",
                    options=sorted(contas_validas["_ordem_vencimento"].dt.year.unique().tolist()),
                    key="ano_calendario_pagar",
                )
                mes_cal = st.selectbox(
                    "Mes",
                    options=list(range(1, 13)),
                    format_func=lambda valor: {
                        1: "Janeiro", 2: "Fevereiro", 3: "Marco", 4: "Abril", 5: "Maio", 6: "Junho",
                        7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                    }[valor],
                    key="mes_calendario_pagar",
                )
                contas_mes = contas_validas[
                    (contas_validas["_ordem_vencimento"].dt.year == ano_cal) &
                    (contas_validas["_ordem_vencimento"].dt.month == mes_cal)
                ].copy()
                st.caption("Pagamentos agrupados por dia de vencimento.")
                cabecalho = st.columns(7)
                for coluna, nome_dia in zip(cabecalho, ["Seg", "Ter", "Qua", "Qui", "Sex", "Sab", "Dom"]):
                    coluna.markdown(f"**{nome_dia}**")

                hoje = date.today()
                for semana in calendar.monthcalendar(ano_cal, mes_cal):
                    cols = st.columns(7)
                    for idx, dia in enumerate(semana):
                        with cols[idx]:
                            if dia == 0:
                                st.write("")
                            else:
                                eh_hoje = hoje.year == ano_cal and hoje.month == mes_cal and hoje.day == dia
                                contas_dia = contas_mes[contas_mes["_ordem_vencimento"].dt.day == dia]
                                statuses_dia = set(contas_dia["status"].fillna("").tolist()) if not contas_dia.empty else set()
                                borda_dia = "2px solid #1f2937" if eh_hoje else "1px solid #e5e7eb"
                                if "Atrasado" in statuses_dia:
                                    fundo_dia = "#fef2f2"
                                elif "Pago" in statuses_dia and statuses_dia.issubset({"Pago"}):
                                    fundo_dia = "#f0fdf4"
                                else:
                                    fundo_dia = "#f8fafc" if eh_hoje else "#ffffff"
                                st.markdown(
                                    f"<div style='border:{borda_dia};border-radius:10px;padding:8px;background:{fundo_dia};'>"
                                    f"<div style='font-weight:700;margin-bottom:8px;'>{dia:02d}/{mes_cal:02d}/{ano_cal}</div>",
                                    unsafe_allow_html=True,
                                )
                                if contas_dia.empty:
                                    st.caption("Sem vencimentos")
                                else:
                                    for _, conta in contas_dia.iterrows():
                                        status_conta = str(conta["status"] or "")
                                        if status_conta == "Pago":
                                            cor = "#15803d"
                                            fundo = "#dcfce7"
                                            borda = "#86efac"
                                        elif status_conta == "Atrasado":
                                            cor = "#b91c1c"
                                            fundo = "#fee2e2"
                                            borda = "#fca5a5"
                                        else:
                                            cor = "#111827"
                                            fundo = "#ffffff"
                                            borda = "#d1d5db"
                                        st.markdown(
                                            f"<div style='border:1px solid {borda};background:{fundo};padding:6px;margin-bottom:6px;border-radius:6px;'>"
                                            f"<div style='font-size:12px;color:{cor};'><strong>{conta['fornecedor']}</strong></div>"
                                            f"<div style='font-size:12px;'>{conta['descricao']}</div>"
                                            f"<div style='font-size:12px;'><strong>{formatar_moeda_br(conta['valor'])}</strong></div>"
                                            f"<div style='font-size:11px;color:#666;'>{conta['status']}</div>"
                                            f"</div>",
                                            unsafe_allow_html=True,
                                        )
                                st.markdown("</div>", unsafe_allow_html=True)

    with tab_novo_pagar:
        st.subheader("Nova divida")
        np1, np2 = st.columns(2)
        nova_data_venc = np1.text_input("Data de vencimento", value=formatar_data_br(date.today()))
        novo_fornecedor = np2.text_input("Fornecedor")
        nova_descricao = st.text_input("Descricao")
        np3, np4, np5 = st.columns(3)
        novo_valor_pagar = np3.number_input("Valor", min_value=0.0, value=0.0)
        novo_status_pagar = np4.selectbox("Status", STATUS_CONTAS_PAGAR, index=0)
        nova_data_pago = np5.text_input("Data de pagamento", value="")
        novo_valor_pago = st.number_input("Valor pago", min_value=0.0, value=0.0, key="novo_valor_pago_manual")
        nova_obs_divida = st.text_area("Observacao")
        if st.button("Adicionar divida"):
            if not parse_data_contrato(nova_data_venc):
                st.error("Informe o vencimento no formato DD/MM/AAAA.")
            elif nova_data_pago.strip() and not parse_data_contrato(nova_data_pago):
                st.error("Informe a data de pagamento no formato DD/MM/AAAA.")
            elif not nova_descricao.strip() and not novo_fornecedor.strip():
                st.error("Informe pelo menos a descricao ou o fornecedor.")
            elif float(novo_valor_pagar) <= 0:
                st.error("Informe um valor maior que zero.")
            else:
                cursor.execute(
                    """
                    INSERT INTO contas_pagar
                    (data_vencimento, descricao, fornecedor, valor, pago, valor_pago, status, observacao, data_criacao, hash_importacao)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        formatar_data_br_valor(nova_data_venc),
                        nova_descricao.strip(),
                        novo_fornecedor.strip(),
                        float(novo_valor_pagar),
                        formatar_data_br_valor(nova_data_pago),
                        float(novo_valor_pago),
                        novo_status_pagar,
                        nova_obs_divida.strip(),
                        agora_str(),
                        montar_hash_importacao_conta_pagar(nova_data_venc, nova_descricao, novo_fornecedor, novo_valor_pagar),
                    ),
                )
                conn.commit()
                st.success("Divida adicionada.")
                st.rerun()
