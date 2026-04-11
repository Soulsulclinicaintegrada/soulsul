from __future__ import annotations

import os
import re
import sqlite3
import unicodedata
from io import BytesIO
from urllib import error as urllib_error
from urllib import request as urllib_request
import json
from datetime import date, datetime
from shutil import copyfile
import base64
try:
    from zoneinfo import ZoneInfo
except ImportError:  # pragma: no cover
    ZoneInfo = None

from fastapi import FastAPI, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel, Field

try:
    from lxml import etree
except ImportError:  # pragma: no cover - ambiente pode nao ter lxml
    etree = None

try:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.text.paragraph import Paragraph
except ImportError:  # pragma: no cover - ambiente pode nao ter python-docx
    Document = None
    WD_BREAK = None
    OxmlElement = None
    qn = None
    Inches = None
    Pt = None
    Paragraph = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
except ImportError:  # pragma: no cover
    Workbook = None
    Alignment = None
    Border = None
    Font = None
    PatternFill = None
    Side = None

from database import (
    SENHA_PADRAO_USUARIOS,
    conectar,
    corrigir_texto_importado,
    garantir_coluna,
    gerar_hash_senha,
    inicializar_banco,
    verificar_senha,
)


DOCS_DIR = "documentos"
EXAMES_DIR = os.path.join("dados_pacientes", "exames")
FOTOS_DIR = os.path.join("dados_pacientes", "fotos")
TEMPLATE_PATH = "modelo_documento.docx"
TEMPLATE_ORDEM_SERVICO_PATH = "ORDEM DE SERVIÇO PROTÉTICO.docx"
CONTAS_CAIXA_MODELO = ["CAIXA", "SICOOB", "INFINITEPAY", "PAGBANK", "C6"]
TIMEZONE_LOCAL = ZoneInfo("America/Sao_Paulo") if ZoneInfo is not None else None


def agora_str() -> str:
    return datetime.now().isoformat(sep=" ", timespec="seconds")


def agora_local() -> datetime:
    if TIMEZONE_LOCAL is not None:
        return datetime.now(TIMEZONE_LOCAL)
    return datetime.now()


def formatar_data_br(data_obj: date | None) -> str:
    if not data_obj:
        return ""
    return data_obj.strftime("%d/%m/%Y")


def parse_data_contrato(valor) -> date | None:
    if valor is None:
        return None
    texto = str(valor).strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return None

    formatos = ("%d/%m/%Y", "%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M")
    for formato in formatos:
        try:
            return datetime.strptime(texto, formato).date()
        except ValueError:
            continue
    if "T" in texto:
        try:
            return datetime.fromisoformat(texto).date()
        except ValueError:
            pass
    return None


def formatar_data_br_valor(valor) -> str:
    return formatar_data_br(parse_data_contrato(valor))


def formatar_prontuario_valor(valor) -> str:
    if valor is None:
        return ""
    texto = str(valor).strip()
    if texto.endswith(".0"):
        texto = texto[:-2]
    return texto


def limpar_nome(nome: str) -> str:
    texto = str(nome or "").strip().upper()
    texto = re.sub(r"[^A-Z0-9 ]+", " ", texto)
    return re.sub(r"\s+", " ", texto).strip()


def normalizar_texto(valor: str) -> str:
    texto = str(valor or "")
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    texto = texto.lower().strip()
    return re.sub(r"\s+", " ", texto)


def limpar_cpf(cpf: str) -> str:
    return "".join(ch for ch in str(cpf or "") if ch.isdigit())


def cpf_valido(cpf: str) -> bool:
    numero = limpar_cpf(cpf)
    if len(numero) != 11 or numero == numero[0] * 11:
        return False

    soma = sum(int(numero[i]) * (10 - i) for i in range(9))
    dig1 = (soma * 10 % 11) % 10
    if dig1 != int(numero[9]):
        return False

    soma = sum(int(numero[i]) * (11 - i) for i in range(10))
    dig2 = (soma * 10 % 11) % 10
    return dig2 == int(numero[10])


def formatar_moeda_br(valor) -> str:
    try:
        valor_float = float(valor or 0)
    except (TypeError, ValueError):
        valor_float = 0.0
    texto = f"{valor_float:,.2f}"
    return f"R$ {texto.replace(',', 'X').replace('.', ',').replace('X', '.')}"


def normalizar_texto_maiusculo(valor: str) -> str:
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    texto = texto.encode("ascii", "ignore").decode("ascii")
    return texto.upper().strip()


def formatar_titulo(valor: str) -> str:
    texto = str(valor or "").strip()
    if not texto:
        return ""
    return " ".join(parte.capitalize() for parte in texto.split())


def colunas_tabela(conn: sqlite3.Connection, nome_tabela: str) -> set[str]:
    return {row["name"] for row in conn.execute(f"PRAGMA table_info({nome_tabela})")}


def garantir_colunas_pacientes_api() -> None:
    conn = conectar()
    try:
        garantir_coluna(conn, "pacientes", "apelido TEXT")
        garantir_coluna(conn, "pacientes", "sexo TEXT")
        garantir_coluna(conn, "pacientes", "rg TEXT")
        garantir_coluna(conn, "pacientes", "email TEXT")
        garantir_coluna(conn, "pacientes", "estado_civil TEXT")
        garantir_coluna(conn, "pacientes", "observacoes TEXT")
        garantir_coluna(conn, "pacientes", "complemento TEXT")
        garantir_coluna(conn, "pacientes", "profissao TEXT")
        garantir_coluna(conn, "pacientes", "origem TEXT")
        garantir_coluna(conn, "pacientes", "foto_path TEXT")
        garantir_coluna(conn, "contratos", "status TEXT DEFAULT 'EM_ABERTO'")
        garantir_coluna(conn, "contratos", "aprovado_por TEXT")
        garantir_coluna(conn, "contratos", "data_aprovacao TEXT")
        garantir_coluna(conn, "contratos", "observacoes TEXT")
        garantir_coluna(conn, "contratos", "clinica_snapshot TEXT")
        garantir_coluna(conn, "contratos", "criado_por_snapshot TEXT")
        garantir_coluna(conn, "contratos", "tabela_snapshot TEXT")
        garantir_coluna(conn, "contratos", "plano_pagamento_json TEXT")
        garantir_coluna(conn, "contratos", "desconto_percentual REAL DEFAULT 0")
        garantir_coluna(conn, "contratos", "desconto_valor REAL DEFAULT 0")
        garantir_coluna(conn, "contratos", "validade_orcamento TEXT")
        garantir_coluna(conn, "procedimentos_contrato", "profissional_snapshot TEXT")
        garantir_coluna(conn, "procedimentos_contrato", "denticao_snapshot TEXT")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_pacientes_nome ON pacientes(nome)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_pacientes_prontuario ON pacientes(prontuario)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_pacientes_cpf ON pacientes(cpf)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_recebiveis_paciente_id ON recebiveis(paciente_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_contratos_paciente_id ON contratos(paciente_id)")
        garantir_coluna(conn, "recebiveis", "data_pagamento TEXT")
        garantir_coluna(conn, "financeiro", "conta_caixa TEXT")
        garantir_coluna(conn, "contas_pagar", "categoria TEXT")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS procedimentos_dente (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                paciente_id INTEGER,
                contrato_id INTEGER,
                dente INTEGER,
                regiao TEXT,
                procedimento TEXT,
                status TEXT,
                faces TEXT,
                valor REAL DEFAULT 0,
                data TEXT
            )
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_procedimentos_dente_paciente ON procedimentos_dente(paciente_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_procedimentos_dente_status ON procedimentos_dente(status)")
        conn.commit()
    finally:
        conn.close()


def coluna_data_agenda() -> str:
    conn = conectar()
    try:
        cols = colunas_tabela(conn, "agendamentos")
    finally:
        conn.close()
    return "data_agendamento" if "data_agendamento" in cols else "data"


DATA_COLUNA_AGENDA = coluna_data_agenda()


def validar_dados_paciente(
    nome: str,
    prontuario: str,
    cpf: str,
    menor: bool,
    responsavel: str,
    cpf_responsavel: str,
) -> list[str]:
    erros: list[str] = []
    if not str(nome or "").strip():
        erros.append("Informe o nome do paciente.")
    if not str(prontuario or "").strip():
        erros.append("Informe o prontuario.")
    if cpf and not cpf_valido(cpf):
        erros.append("O CPF do paciente e invalido.")
    if menor and not str(responsavel or "").strip():
        erros.append("Informe o responsavel legal para paciente menor de idade.")
    if cpf_responsavel and not cpf_valido(cpf_responsavel):
        erros.append("O CPF do responsavel e invalido.")
    return erros


def validar_duplicidade_paciente(
    conn: sqlite3.Connection,
    prontuario: str,
    cpf: str,
    paciente_id_atual: int | None = None,
) -> list[str]:
    erros: list[str] = []
    filtro_id = ""
    params_exclusao: list[object] = []
    if paciente_id_atual is not None:
        filtro_id = " AND id <> ?"
        params_exclusao.append(int(paciente_id_atual))

    prontuario_limpo = formatar_prontuario_valor(prontuario)
    if prontuario_limpo:
        row = conn.execute(
            f"SELECT id FROM pacientes WHERE trim(COALESCE(prontuario, '')) = ?{filtro_id} LIMIT 1",
            [prontuario_limpo, *params_exclusao],
        ).fetchone()
        if row is not None:
            erros.append("Ja existe um paciente com este prontuario.")

    cpf_limpo = limpar_cpf(cpf)
    if cpf_limpo:
        row = conn.execute(
            f"SELECT id FROM pacientes WHERE trim(COALESCE(cpf, '')) = ?{filtro_id} LIMIT 1",
            [cpf_limpo, *params_exclusao],
        ).fetchone()
        if row is not None:
            erros.append("Ja existe um paciente com este CPF.")

    return erros


def buscar_endereco_por_cep(cep: str) -> dict[str, str]:
    cep_limpo = "".join(ch for ch in str(cep or "") if ch.isdigit())
    if len(cep_limpo) != 8:
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}
    try:
        with urllib_request.urlopen(f"https://viacep.com.br/ws/{cep_limpo}/json/", timeout=5) as resposta:
            dados = json.loads(resposta.read().decode("utf-8"))
            if dados.get("erro"):
                return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}
            return {
                "logradouro": str(dados.get("logradouro") or ""),
                "bairro": str(dados.get("bairro") or ""),
                "localidade": str(dados.get("localidade") or ""),
                "uf": str(dados.get("uf") or ""),
            }
    except (urllib_error.URLError, TimeoutError, json.JSONDecodeError):
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}


def proximo_prontuario(conn: sqlite3.Connection) -> str:
    rows = conn.execute("SELECT prontuario FROM pacientes WHERE COALESCE(prontuario, '') <> ''").fetchall()
    maior = 0
    for row in rows:
        texto = formatar_prontuario_valor(row["prontuario"])
        digitos = "".join(ch for ch in texto if ch.isdigit())
        if digitos:
            maior = max(maior, int(digitos))
    return str(maior + 1 if maior else 1)


def slug_paciente_arquivos(paciente_row: sqlite3.Row | dict) -> str:
    prontuario = formatar_prontuario_valor(paciente_row.get("prontuario") if isinstance(paciente_row, dict) else paciente_row["prontuario"])
    nome = limpar_nome(paciente_row.get("nome", "PACIENTE") if isinstance(paciente_row, dict) else paciente_row["nome"])
    base = f"{prontuario}_{nome}".strip("_")
    return base or f"paciente_{int(paciente_row['id'])}"


def pasta_exames_paciente(paciente_row: sqlite3.Row | dict) -> str:
    pasta = os.path.join(EXAMES_DIR, slug_paciente_arquivos(paciente_row))
    os.makedirs(pasta, exist_ok=True)
    return pasta


def pasta_fotos_paciente(paciente_row: sqlite3.Row | dict) -> str:
    pasta = os.path.join(FOTOS_DIR, slug_paciente_arquivos(paciente_row))
    os.makedirs(pasta, exist_ok=True)
    return pasta


def url_foto_paciente(row: sqlite3.Row) -> str:
    if not str(row["foto_path"] or "").strip():
        return ""
    return f"/api/pacientes/{int(row['id'])}/foto"


class PacientePayload(BaseModel):
    nome: str
    apelido: str = ""
    sexo: str = ""
    prontuario: str | None = None
    cpf: str = ""
    rg: str = ""
    data_nascimento: str = ""
    telefone: str = ""
    email: str = ""
    cep: str = ""
    endereco: str = ""
    complemento: str = ""
    numero: str = ""
    bairro: str = ""
    cidade: str = ""
    estado: str = ""
    estado_civil: str = ""
    profissao: str = ""
    origem: str = ""
    observacoes: str = ""
    menor_idade: bool = False
    responsavel: str = ""
    cpf_responsavel: str = ""


class PacienteResumo(BaseModel):
    id: int
    nome: str
    apelido: str = ""
    prontuario: str
    cpf: str = ""
    telefone: str = ""
    email: str = ""
    dataNascimento: str = ""
    fotoUrl: str = ""


class PacienteDetalhe(BaseModel):
    id: int
    nome: str
    apelido: str = ""
    sexo: str = ""
    prontuario: str
    cpf: str = ""
    rg: str = ""
    dataNascimento: str = ""
    telefone: str = ""
    email: str = ""
    cep: str = ""
    endereco: str = ""
    complemento: str = ""
    numero: str = ""
    bairro: str = ""
    cidade: str = ""
    estado: str = ""
    estadoCivil: str = ""
    profissao: str = ""
    origem: str = ""
    observacoes: str = ""
    menorIdade: bool = False
    responsavel: str = ""
    cpfResponsavel: str = ""
    fotoUrl: str = ""


class ContratoResumo(BaseModel):
    id: int
    valorTotal: str
    entrada: str
    parcelas: int
    primeiroVencimento: str = ""
    dataCriacao: str = ""
    formaPagamento: str = ""
    status: str = "EM_ABERTO"
    aprovadoPor: str = ""
    dataAprovacao: str = ""
    procedimentos: list[str] = Field(default_factory=list)


class RecebivelResumo(BaseModel):
    id: int
    pacienteId: int | None = None
    pacienteNome: str = ""
    prontuario: str = ""
    contratoId: int | None = None
    parcela: int | None = None
    vencimento: str = ""
    valor: str
    formaPagamento: str = ""
    status: str = ""
    dataPagamento: str = ""
    observacao: str = ""


class RecebivelAtualizacaoPayload(BaseModel):
    paciente_nome: str = ""
    prontuario: str = ""
    vencimento: str = ""
    valor: float = 0
    forma_pagamento: str = ""
    status: str = ""
    data_pagamento: str = ""
    observacao: str = ""


class RecebivelLotePayload(BaseModel):
    paciente_nome: str = ""
    prontuario: str = ""
    forma_pagamento: str = ""
    status: str = ""
    observacao: str = ""
    primeiro_vencimento: str = ""


class AgendamentoResumo(BaseModel):
    id: int
    data: str = ""
    horario: str = ""
    profissional: str = ""
    status: str = ""
    procedimento: str = ""
    observacao: str = ""


class ArquivoPacienteItem(BaseModel):
    nome: str
    caminho: str
    modificadoEm: str = ""
    extensao: str = ""


class FinanceiroResumo(BaseModel):
    total: str
    emAberto: str
    atrasado: str
    pagos: str
    quantidadeAtrasados: int
    indicador: str


class MovimentoCaixaResumo(BaseModel):
    id: int
    data: str = ""
    origem: str = ""
    descricao: str = ""
    valor: str = ""
    tipo: str = ""
    prontuario: str = ""
    formaPagamento: str = ""
    contaCaixa: str = ""
    observacao: str = ""
    contratoId: int | None = None
    recebivelId: int | None = None


class ContaPagarResumo(BaseModel):
    id: int
    vencimento: str = ""
    descricao: str = ""
    fornecedor: str = ""
    categoria: str = ""
    valor: str = ""
    valorPago: str = ""
    pagoEm: str = ""
    status: str = ""
    observacao: str = ""


class SaldoContaResumo(BaseModel):
    id: int
    data: str = ""
    conta: str = ""
    saldo: str = ""
    observacao: str = ""


class SaldoContaPayload(BaseModel):
    data: str = ""
    conta: str = ""
    saldo: float = 0
    observacao: str = ""


class ContaPagarPayload(BaseModel):
    vencimento: str = ""
    descricao: str = ""
    fornecedor: str = ""
    categoria: str = ""
    valor: float = 0
    valor_pago: float = 0
    pago_em: str = ""
    status: str = ""
    observacao: str = ""


class BaixaRecebivelPayload(BaseModel):
    data_pagamento: str = ""
    forma_pagamento: str = ""
    conta_caixa: str = ""
    desconto_valor: float = 0
    observacao: str = ""


class MovimentoCaixaPayload(BaseModel):
    origem: str = ""
    descricao: str = ""
    valor: float = 0
    tipo: str = "Entrada"
    data_movimento: str = ""
    prontuario: str = ""
    forma_pagamento: str = ""
    conta_caixa: str = ""
    observacao: str = ""
    contrato_id: int | None = None
    recebivel_id: int | None = None


class MovimentoCaixaAtualizacaoPayload(BaseModel):
    origem: str = ""
    descricao: str = ""
    valor: float = 0
    tipo: str = ""
    data_movimento: str = ""
    prontuario: str = ""
    forma_pagamento: str = ""
    conta_caixa: str = ""
    observacao: str = ""


class FinanceiroPainelResposta(BaseModel):
    resumo: FinanceiroResumo
    recebiveis: list[RecebivelResumo]
    caixa: list[MovimentoCaixaResumo]
    contasPagar: list[ContaPagarResumo]
    saldosConta: list[SaldoContaResumo]


class ReciboManualPayload(BaseModel):
    valor: float = 0
    pagador: str = ""
    recebedor: str = ""
    data_pagamento: str = ""
    referente: str = ""
    observacao: str = ""
    cidade: str = ""


class ReciboManualResumo(BaseModel):
    id: int
    valor: str = ""
    pagador: str = ""
    recebedor: str = ""
    dataPagamento: str = ""
    referente: str = ""
    observacao: str = ""
    cidade: str = ""
    criadoEm: str = ""


class MetaMensalPayload(BaseModel):
    meta: float = 0
    supermeta: float = 0
    hipermeta: float = 0


class MetaMensalResumo(BaseModel):
    ano: int
    mes: int
    mesNome: str = ""
    meta: float = 0
    supermeta: float = 0
    hipermeta: float = 0
    dataAtualizacao: str = ""


class NotaFiscalEmitidaPayload(BaseModel):
    competencia: str = ""
    data_emissao: str = ""
    data_recebimento: str = ""
    numero_nf: str = ""
    serie: str = ""
    cliente: str = ""
    descricao: str = ""
    conta_destino: str = ""
    valor_nf: float = 0
    valor_recebido: float = 0
    status: str = "Pendente"
    observacao: str = ""


class NotaFiscalEmitidaResumo(BaseModel):
    id: int
    competencia: str = ""
    dataEmissao: str = ""
    dataRecebimento: str = ""
    numeroNf: str = ""
    serie: str = ""
    cliente: str = ""
    descricao: str = ""
    contaDestino: str = ""
    valorNf: str = ""
    valorRecebido: str = ""
    valorNfNumero: float = 0
    valorRecebidoNumero: float = 0
    diferenca: str = ""
    diferencaNumero: float = 0
    status: str = ""
    observacao: str = ""
    conciliado: bool = False
    criadoEm: str = ""
    atualizadoEm: str = ""


class DashboardIndicadorResposta(BaseModel):
    chave: str
    titulo: str
    valor: str
    detalhe: str = ""


class DashboardResumoHojeResposta(BaseModel):
    entradasConfirmadas: str
    saidasPrevistas: str
    saldoProjetado: str


class DashboardAgendaHojeItemResposta(BaseModel):
    horario: str = ""
    titulo: str = ""
    subtitulo: str = ""


class DashboardAlertaItemResposta(BaseModel):
    titulo: str
    detalhe: str


class DashboardAtividadeItemResposta(BaseModel):
    paciente: str
    evento: str
    valor: str
    status: str


class DashboardMetasResposta(BaseModel):
    vendidoMes: str
    vendidoAno: str
    metaMes: str
    supermetaMes: str
    hipermetaMes: str
    faltaMetaMes: str
    faltaMetaAno: str
    percentualMetaMes: float = 0
    percentualMetaAno: float = 0


class DashboardPainelResposta(BaseModel):
    indicadores: list[DashboardIndicadorResposta]
    meses: list[str]
    serieVendas: list[float]
    resumoHoje: DashboardResumoHojeResposta
    metas: DashboardMetasResposta
    agendaHoje: list[DashboardAgendaHojeItemResposta]
    alertas: list[DashboardAlertaItemResposta]
    atividades: list[DashboardAtividadeItemResposta]


class FichaPacienteResposta(BaseModel):
    paciente: PacienteDetalhe
    contratos: list[ContratoResumo]
    recebiveis: list[RecebivelResumo]
    financeiro: FinanceiroResumo
    agendamentos: list[AgendamentoResumo]
    proximoAgendamento: AgendamentoResumo | None = None
    documentos: list[ArquivoPacienteItem]
    exames: list[ArquivoPacienteItem]
    recibos: list[RecebivelResumo]
    crm: "CrmPacienteResumoResposta | None" = None


class CrmPacienteResumoResposta(BaseModel):
    crmId: int | None = None
    finalizado: bool = False
    avaliacao: bool = False
    etapaFunil: str = ""
    campanha: str = ""
    canal: str = "Facebook"
    ultimaAvaliacaoEm: str = ""
    finalizadoEm: str = ""


class CrmPacienteItemResposta(BaseModel):
    id: int
    pacienteId: int
    nome: str
    prontuario: str = ""
    telefone: str = ""
    origemFinalizado: bool = False
    origemAvaliacao: bool = False
    etapaFunil: str = "Novo lead"
    canal: str = "Facebook"
    campanha: str = ""
    conjuntoAnuncio: str = ""
    anuncio: str = ""
    responsavel: str = ""
    proximoContato: str = ""
    observacao: str = ""
    ultimaInteracao: str = ""
    ultimaAvaliacaoEm: str = ""
    finalizadoEm: str = ""
    atualizadoEm: str = ""


class CrmAvaliacaoItemResposta(BaseModel):
    pacienteId: int
    nome: str
    prontuario: str = ""
    telefone: str = ""
    dataAvaliacao: str = ""
    profissional: str = ""
    status: str = ""
    procedimento: str = ""
    jaNoCrm: bool = False


class CrmPainelResposta(BaseModel):
    pipeline: list[CrmPacienteItemResposta] = Field(default_factory=list)
    finalizados: list[CrmPacienteItemResposta] = Field(default_factory=list)
    avaliacoes: list[CrmAvaliacaoItemResposta] = Field(default_factory=list)


class CrmAtualizacaoPayload(BaseModel):
    etapa_funil: str = ""
    canal: str = "Facebook"
    campanha: str = ""
    conjunto_anuncio: str = ""
    anuncio: str = ""
    responsavel: str = ""
    proximo_contato: str = ""
    observacao: str = ""
    ultima_interacao: str = ""


class OdontogramaResposta(BaseModel):
    dentes_contratados: list[int] = Field(default_factory=list)
    elementos: list[dict] = Field(default_factory=list)


class OrcamentoRegiaoPayload(BaseModel):
    regiao: str
    dente: int | None = None
    valor: float = 0
    ativo: bool = True
    faces: list[str] = Field(default_factory=list)


class OrcamentoItemPayload(BaseModel):
    procedimento: str
    profissional: str = ""
    denticao: str = ""
    valor_unitario: float = 0
    regioes: list[OrcamentoRegiaoPayload] = Field(default_factory=list)


class ParcelaPagamentoPayload(BaseModel):
    indice: int = 0
    descricao: str = ""
    data: str = ""
    forma: str = ""
    valor: float = 0
    parcelas_cartao: int = 1


class OrcamentoPacientePayload(BaseModel):
    clinica: str = ""
    criado_por: str = ""
    data: str = ""
    observacoes: str = ""
    tabela: str = ""
    desconto_percentual: float = 0
    desconto_valor: float = 0
    validade_orcamento: str = ""
    forma_pagamento: str = ""
    parcelas: int = 1
    entrada: bool = False
    plano_pagamento: list[ParcelaPagamentoPayload] = Field(default_factory=list)
    itens: list[OrcamentoItemPayload] = Field(default_factory=list)


class OrcamentoCriadoResposta(BaseModel):
    contrato_id: int


class OrcamentoStatusPayload(BaseModel):
    status: str
    aprovado_por: str = ""


class OrcamentoDetalheResposta(BaseModel):
    contrato_id: int
    status: str = "EM_ABERTO"
    aprovadoPor: str = ""
    dataAprovacao: str = ""
    clinica: str = ""
    criadoPor: str = ""
    data: str = ""
    observacoes: str = ""
    tabela: str = ""
    descontoPercentual: float = 0
    descontoValor: float = 0
    validadeOrcamento: str = ""
    formaPagamento: str = ""
    parcelas: int = 1
    entrada: bool = False
    planoPagamento: list[ParcelaPagamentoPayload] = Field(default_factory=list)
    itens: list[OrcamentoItemPayload] = Field(default_factory=list)


class ProcedimentoResumoResposta(BaseModel):
    id: int
    nome: str
    categoria: str = ""
    valorPadrao: float = 0
    duracaoPadraoMinutos: int = 60
    descricao: str = ""
    etapasPadrao: list[str] = Field(default_factory=list)
    materiaisPadrao: list[str] = Field(default_factory=list)
    ativo: bool = True


class ProcedimentoPayload(BaseModel):
    nome: str = ""
    categoria: str = ""
    valor_padrao: float = 0
    duracao_padrao_minutos: int = 60
    descricao: str = ""
    etapas_padrao: list[str] = Field(default_factory=list)
    materiais_padrao: list[str] = Field(default_factory=list)
    ativo: bool = True


class OrdemServicoEtapaPayload(BaseModel):
    etapa: str = ""
    descricao_outro: str = ""


class OrdemServicoPayload(BaseModel):
    procedimento_id: int
    material: str = ""
    material_outro: str = ""
    cor: str = ""
    escala: str = ""
    elemento_arcada: str = ""
    carga_imediata: bool = False
    retorno_solicitado: str = ""
    observacao: str = ""
    etapas: list[OrdemServicoEtapaPayload] = Field(default_factory=list)


class OrdemServicoResumoResposta(BaseModel):
    id: int
    procedimentoId: int | None = None
    procedimentoNome: str = ""
    material: str = ""
    materialOutro: str = ""
    cor: str = ""
    escala: str = ""
    elementoArcada: str = ""
    cargaImediata: bool = False
    retornoSolicitado: str = ""
    documentoNome: str = ""
    observacao: str = ""
    criadoEm: str = ""
    etapas: list[OrdemServicoEtapaPayload] = Field(default_factory=list)


class LoginPayload(BaseModel):
    usuario: str = ""
    senha: str = ""


class TrocaSenhaPayload(BaseModel):
    usuario: str = ""
    senha_atual: str = ""
    nova_senha: str = ""


class RedefinirSenhaPayload(BaseModel):
    nova_senha: str = ""


class UsuarioAtualizacaoPayload(BaseModel):
    nome: str = ""
    nome_agenda: str = ""
    cargo: str = ""
    agenda_escopo: str = ""
    agenda_disponivel: bool | None = None
    perfil: str = ""
    usuario: str = ""
    ativo: bool | None = None
    modulos: dict[str, str] = Field(default_factory=dict)
    pacientes_abas: dict[str, str] = Field(default_factory=dict)


class UsuarioCriacaoPayload(BaseModel):
    nome: str = ""
    nome_agenda: str = ""
    cargo: str = ""
    agenda_escopo: str = ""
    agenda_disponivel: bool | None = None
    perfil: str = ""
    ativo: bool | None = True
    modulos: dict[str, str] = Field(default_factory=dict)
    pacientes_abas: dict[str, str] = Field(default_factory=dict)


class UsuarioResumoResposta(BaseModel):
    id: int
    nome: str
    usuario: str
    nomeAgenda: str = ""
    perfil: str
    cargo: str = ""
    agendaEscopo: str = ""
    agendaDisponivel: bool = False
    status: str = "Ativo"
    ultimoAcesso: str = "-"
    modulos: dict[str, str] = Field(default_factory=dict)
    pacientesAbas: dict[str, str] = Field(default_factory=dict)


class LoginResposta(BaseModel):
    id: int
    nome: str
    usuario: str
    perfil: str
    cargo: str = ""
    agendaEscopo: str = ""
    agendaDisponivel: bool = False
    nomeAgenda: str = ""
    modulos: dict[str, str] = Field(default_factory=dict)
    pacientesAbas: dict[str, str] = Field(default_factory=dict)
    precisaTrocarSenha: bool = False


app = FastAPI(title="SoulSul Pacientes API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_origin_regex=r"https?://.*",
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def startup_event() -> None:
    inicializar_banco()
    garantir_colunas_pacientes_api()


def usuario_request(request: Request) -> str:
    return str(request.headers.get("x-usuario") or "").strip()


def mapear_acao_http(method: str) -> str:
    metodo = (method or "").upper()
    if metodo == "POST":
        return "Criacao"
    if metodo in {"PUT", "PATCH"}:
        return "Edicao"
    if metodo == "DELETE":
        return "Exclusao"
    return metodo.title()


def mapear_tipo_rota(path: str) -> str:
    rota = normalizar_texto(path)
    if "agenda" in rota:
        return "Agendamento"
    if "orcamento" in rota or "contrato" in rota:
        return "Plano de Tratamento"
    if "recebivel" in rota:
        return "Parcela"
    if "caixa" in rota:
        return "Caixa"
    if "contas_pagar" in rota or "contas-pagar" in rota:
        return "Conta a Pagar"
    if "paciente" in rota:
        return "Paciente"
    if "usuario" in rota:
        return "Usuario"
    return "Sistema"


def extrair_info_acao(payload: object) -> str:
    if isinstance(payload, dict):
      for chave in ("paciente_nome", "pacienteNome", "nomePaciente", "nome", "descricao", "fornecedor", "procedimento", "titulo"):
        valor = payload.get(chave)
        if str(valor or "").strip():
          return str(valor).strip()
      if "itens" in payload and isinstance(payload.get("itens"), list) and payload["itens"]:
        primeiro = payload["itens"][0]
        if isinstance(primeiro, dict):
          for chave in ("procedimento", "nome"):
            valor = primeiro.get(chave)
            if str(valor or "").strip():
              return str(valor).strip()
    return ""


def registrar_acao_usuario(
    usuario: str,
    *,
    acao: str,
    tipo: str,
    info: str = "",
    metodo_http: str = "",
    rota: str = "",
) -> None:
    usuario_limpo = str(usuario or "").strip()
    if not usuario_limpo:
        return
    conn = conectar()
    try:
        usuario_row = conn.execute(
            "SELECT id FROM usuarios WHERE lower(usuario)=lower(?) OR lower(nome)=lower(?) LIMIT 1",
            (usuario_limpo, usuario_limpo),
        ).fetchone()
        conn.execute(
            """
            INSERT INTO acoes_usuario (usuario_id, usuario, acao, tipo, info, metodo_http, rota, data_hora)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                int(usuario_row["id"]) if usuario_row else None,
                usuario_limpo.upper(),
                str(acao or "").strip() or "Edicao",
                str(tipo or "").strip() or "Sistema",
                str(info or "").strip(),
                str(metodo_http or "").strip().upper(),
                str(rota or "").strip(),
                agora_str(),
            ),
        )
        conn.commit()
    finally:
        conn.close()


def bytes_relatorio_acoes_por_dia(data_referencia: str) -> bytes:
    if Workbook is None:
        raise HTTPException(status_code=500, detail="openpyxl nao disponivel")
    data_alvo = parse_data_contrato(data_referencia) or date.today()
    data_br = formatar_data_br(data_alvo)
    conn = conectar()
    try:
        rows = conn.execute(
            """
            SELECT data_hora, acao, tipo, info, usuario
            FROM acoes_usuario
            WHERE substr(COALESCE(data_hora, ''), 1, 10)=?
            ORDER BY COALESCE(data_hora, '') DESC, id DESC
            """,
            (data_alvo.isoformat(),),
        ).fetchall()
    finally:
        conn.close()

    wb = Workbook()
    ws = wb.active
    ws.title = "Relatorio"
    ws.append(["Data", "Hora", "Acao", "Tipo", "Info", "Usuario"])
    for row in rows:
        data_hora = parse_data_contrato(row["data_hora"])
        data_txt = formatar_data_br(data_hora)
        hora_txt = ""
        texto_data_hora = str(row["data_hora"] or "").strip()
        if " " in texto_data_hora:
            hora_txt = texto_data_hora.split(" ", 1)[1][:5]
        ws.append([
            data_txt or data_br,
            hora_txt,
            str(row["acao"] or ""),
            str(row["tipo"] or ""),
            str(row["info"] or ""),
            str(row["usuario"] or ""),
        ])
    largura = {"A": 14, "B": 10, "C": 16, "D": 24, "E": 42, "F": 20}
    for coluna, valor in largura.items():
        ws.column_dimensions[coluna].width = valor
    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


@app.middleware("http")
async def auditoria_middleware(request: Request, call_next):
    body_bytes = await request.body()

    async def receive():
        return {"type": "http.request", "body": body_bytes, "more_body": False}

    payload_dict: dict[str, object] = {}
    if body_bytes:
        try:
            payload_dict = json.loads(body_bytes.decode("utf-8"))
        except Exception:
            payload_dict = {}

    response = await call_next(Request(request.scope, receive))

    if request.url.path.startswith("/api/") and request.method.upper() in {"POST", "PUT", "PATCH", "DELETE"} and response.status_code < 500:
        usuario = usuario_request(request)
        if usuario and not request.url.path.startswith("/api/auth/"):
            registrar_acao_usuario(
                usuario,
                acao=mapear_acao_http(request.method),
                tipo=mapear_tipo_rota(request.url.path),
                info=extrair_info_acao(payload_dict),
                metodo_http=request.method,
                rota=request.url.path,
            )
    return response


def detalhar_usuario_login(usuario_row: sqlite3.Row) -> LoginResposta:
    cargo = str(usuario_row["cargo"] or "Profissional")
    perfil = str(usuario_row["perfil"] or "Usuario")
    modulos_padrao, abas_padrao, _ = permissoes_padrao_backend(cargo, perfil)
    return LoginResposta(
        id=int(usuario_row["id"]),
        nome=str(usuario_row["nome"] or ""),
        usuario=str(usuario_row["usuario"] or ""),
        perfil=perfil,
        cargo=cargo,
        agendaEscopo=str(usuario_row["agenda_escopo"] or ""),
        agendaDisponivel=bool(int(usuario_row["agenda_disponivel"] or 0)),
        nomeAgenda=str(usuario_row["nome_agenda"] or usuario_row["nome"] or "").upper(),
        modulos=json_dict_seguro(usuario_row["modulos_json"], modulos_padrao),
        pacientesAbas=json_dict_seguro(usuario_row["pacientes_abas_json"], abas_padrao),
        precisaTrocarSenha=bool(int(usuario_row["precisa_trocar_senha"] or 0)),
    )


MODULOS_USUARIOS = ["Dashboard", "Pacientes", "Agenda", "Financeiro", "Tabelas", "Usuarios"]
ABAS_PACIENTES_USUARIOS = ["Cadastro", "Orcamentos", "Financeiro", "Documentos", "Plano e Ficha Clinica", "Odontograma", "Agendamentos"]


def permissoes_padrao_backend(cargo: str, perfil: str) -> tuple[dict[str, str], dict[str, str], str]:
    cargo_limpo = corrigir_texto_importado(str(cargo or "").strip()) or "Profissional"
    perfil_limpo = corrigir_texto_importado(str(perfil or "").strip()) or "Usuario"
    modulos = {modulo: "Sem acesso" for modulo in MODULOS_USUARIOS}
    pacientes_abas = {aba: "Sem acesso" for aba in ABAS_PACIENTES_USUARIOS}
    agenda_escopo = "TODA_CLINICA"

    if perfil_limpo == "Administrador" or cargo_limpo == "Administrador":
        modulos = {modulo: "Edicao" for modulo in MODULOS_USUARIOS}
        pacientes_abas = {aba: "Edicao" for aba in ABAS_PACIENTES_USUARIOS}
        agenda_escopo = "TODA_CLINICA"
    elif cargo_limpo == "Profissional":
        modulos["Pacientes"] = "Edicao"
        modulos["Agenda"] = "Visualizacao"
        pacientes_abas["Documentos"] = "Edicao"
        pacientes_abas["Plano e Ficha Clinica"] = "Visualizacao"
        pacientes_abas["Odontograma"] = "Visualizacao"
        pacientes_abas["Agendamentos"] = "Visualizacao"
        agenda_escopo = "SOMENTE_PROPRIA"
    else:
        modulos["Dashboard"] = "Visualizacao"
        modulos["Pacientes"] = "Edicao"
        modulos["Agenda"] = "Edicao"
        modulos["Financeiro"] = "Visualizacao"
        pacientes_abas["Cadastro"] = "Edicao"
        pacientes_abas["Orcamentos"] = "Visualizacao"
        pacientes_abas["Financeiro"] = "Visualizacao"
        pacientes_abas["Documentos"] = "Visualizacao"
        pacientes_abas["Agendamentos"] = "Edicao"
        agenda_escopo = "TODA_CLINICA"

    return modulos, pacientes_abas, agenda_escopo


def json_dict_seguro(texto: object, padrao: dict[str, str]) -> dict[str, str]:
    bruto = str(texto or "").strip()
    if not bruto:
        return dict(padrao)
    try:
        valor = json.loads(bruto)
        if isinstance(valor, dict):
            return {str(chave): str(item) for chave, item in valor.items()}
    except Exception:
        pass
    return dict(padrao)


def normalizar_payload_permissoes(
    cargo: str,
    perfil: str,
    modulos_payload: dict[str, str] | None,
    pacientes_abas_payload: dict[str, str] | None,
) -> tuple[dict[str, str], dict[str, str], str]:
    modulos_padrao, abas_padrao, agenda_padrao = permissoes_padrao_backend(cargo, perfil)
    modulos = dict(modulos_padrao)
    for modulo, nivel in (modulos_payload or {}).items():
        if modulo in modulos and str(nivel or "").strip():
            modulos[modulo] = str(nivel)
    abas = dict(abas_padrao)
    for aba, nivel in (pacientes_abas_payload or {}).items():
        if aba in abas and str(nivel or "").strip():
            abas[aba] = str(nivel)
    return modulos, abas, agenda_padrao


def mapear_usuario_resumo(usuario_row: sqlite3.Row) -> UsuarioResumoResposta:
    cargo = str(usuario_row["cargo"] or "Profissional")
    perfil = str(usuario_row["perfil"] or "Usuario")
    modulos_padrao, abas_padrao, _ = permissoes_padrao_backend(cargo, perfil)
    return UsuarioResumoResposta(
        id=int(usuario_row["id"]),
        nome=str(usuario_row["nome"] or ""),
        usuario=str(usuario_row["usuario"] or ""),
        nomeAgenda=str(usuario_row["nome_agenda"] or usuario_row["nome"] or "").upper(),
        perfil=perfil,
        cargo=cargo,
        agendaEscopo=str(usuario_row["agenda_escopo"] or ""),
        agendaDisponivel=bool(int(usuario_row["agenda_disponivel"] or 0)),
        status="Ativo" if bool(int(usuario_row["ativo"] or 0)) else "Inativo",
        ultimoAcesso=str(usuario_row["ultimo_login"] or "-"),
        modulos=json_dict_seguro(usuario_row["modulos_json"], modulos_padrao),
        pacientesAbas=json_dict_seguro(usuario_row["pacientes_abas_json"], abas_padrao),
    )


def mapear_procedimento_resumo(row: sqlite3.Row) -> ProcedimentoResumoResposta:
    etapas: list[str] = []
    materiais: list[str] = []
    texto_etapas = str(row["etapas_json"] or "").strip()
    if texto_etapas:
        try:
            bruto = json.loads(texto_etapas)
            if isinstance(bruto, list):
                etapas = [corrigir_texto_importado(str(item).strip()) for item in bruto if str(item).strip()]
        except Exception:
            etapas = []
    texto_materiais = str(row["materiais_json"] or "").strip()
    if texto_materiais:
        try:
            bruto = json.loads(texto_materiais)
            if isinstance(bruto, list):
                materiais = [corrigir_texto_importado(str(item).strip()) for item in bruto if str(item).strip()]
        except Exception:
            materiais = []
    return ProcedimentoResumoResposta(
        id=int(row["id"]),
        nome=corrigir_texto_importado(str(row["nome"] or "")),
        categoria=corrigir_texto_importado(str(row["categoria"] or "")),
        valorPadrao=float(row["valor_padrao"] or 0),
        duracaoPadraoMinutos=int(row["duracao_padrao_minutos"] or 60),
        descricao=corrigir_texto_importado(str(row["descricao"] or "")),
        etapasPadrao=etapas,
        materiaisPadrao=materiais,
        ativo=bool(int(row["ativo"] or 0)),
    )


@app.post("/api/auth/login", response_model=LoginResposta)
def login_usuario(payload: LoginPayload):
    usuario_digitado = str(payload.usuario or "").strip()
    senha = str(payload.senha or "")
    if not usuario_digitado or not senha:
        raise HTTPException(status_code=400, detail="Informe usuario e senha.")

    conn = conectar()
    try:
        usuario_row = conn.execute(
            """
            SELECT *
            FROM usuarios
            WHERE ativo=1 AND (lower(usuario)=lower(?) OR lower(nome)=lower(?))
            LIMIT 1
            """,
            (usuario_digitado, usuario_digitado),
        ).fetchone()
        if usuario_row is None or not verificar_senha(senha, usuario_row["senha_hash"]):
            raise HTTPException(status_code=401, detail="Usuario ou senha invalidos.")

        conn.execute("UPDATE usuarios SET ultimo_login=? WHERE id=?", (agora_str(), int(usuario_row["id"])))
        conn.execute(
            "INSERT INTO logs_acesso (usuario_id, usuario, evento, data_hora) VALUES (?, ?, ?, ?)",
            (int(usuario_row["id"]), str(usuario_row["usuario"] or usuario_row["nome"] or "").upper(), "LOGIN", agora_str()),
        )
        conn.commit()
    finally:
        conn.close()

    registrar_acao_usuario(
        str(usuario_row["usuario"] or usuario_row["nome"] or ""),
        acao="Login",
        tipo="Acesso",
        info=str(usuario_row["nome"] or ""),
        metodo_http="POST",
        rota="/api/auth/login",
    )
    return detalhar_usuario_login(usuario_row)


@app.post("/api/auth/trocar-senha", response_model=LoginResposta)
def trocar_senha_primeiro_acesso(payload: TrocaSenhaPayload):
    usuario_digitado = str(payload.usuario or "").strip()
    senha_atual = str(payload.senha_atual or "")
    nova_senha = str(payload.nova_senha or "").strip()
    if not usuario_digitado or not senha_atual or not nova_senha:
        raise HTTPException(status_code=400, detail="Preencha usuario, senha atual e nova senha.")
    if len(nova_senha) < 4:
        raise HTTPException(status_code=400, detail="A nova senha precisa ter pelo menos 4 caracteres.")

    conn = conectar()
    try:
        usuario_row = conn.execute(
            """
            SELECT *
            FROM usuarios
            WHERE ativo=1 AND (lower(usuario)=lower(?) OR lower(nome)=lower(?))
            LIMIT 1
            """,
            (usuario_digitado, usuario_digitado),
        ).fetchone()
        if usuario_row is None or not verificar_senha(senha_atual, usuario_row["senha_hash"]):
            raise HTTPException(status_code=401, detail="Usuario ou senha atual invalidos.")

        novo_hash = gerar_hash_senha(nova_senha)
        conn.execute(
            """
            UPDATE usuarios
            SET senha_hash=?, senha_temporaria=0, precisa_trocar_senha=0, ultimo_login=?
            WHERE id=?
            """,
            (novo_hash, agora_str(), int(usuario_row["id"])),
        )
        conn.commit()
        usuario_row = conn.execute("SELECT * FROM usuarios WHERE id=?", (int(usuario_row["id"]),)).fetchone()
    finally:
        conn.close()

    registrar_acao_usuario(
        str(usuario_row["usuario"] or usuario_row["nome"] or ""),
        acao="Troca de senha",
        tipo="Acesso",
        info=str(usuario_row["nome"] or ""),
        metodo_http="POST",
        rota="/api/auth/trocar-senha",
    )
    return detalhar_usuario_login(usuario_row)


@app.post("/api/usuarios/{usuario_id}/redefinir-senha", response_model=LoginResposta)
def redefinir_senha_usuario(usuario_id: int, payload: RedefinirSenhaPayload, request: Request):
    nova_senha = str(payload.nova_senha or "").strip()
    if len(nova_senha) < 4:
        raise HTTPException(status_code=400, detail="A nova senha precisa ter pelo menos 4 caracteres.")

    conn = conectar()
    try:
        usuario_row = conn.execute(
            """
            SELECT *
            FROM usuarios
            WHERE id=? AND ativo=1
            LIMIT 1
            """,
            (int(usuario_id),),
        ).fetchone()
        if usuario_row is None:
            raise HTTPException(status_code=404, detail="Usuario nao encontrado.")

        novo_hash = gerar_hash_senha(nova_senha)
        conn.execute(
            """
            UPDATE usuarios
            SET senha_hash=?, senha_temporaria=0, precisa_trocar_senha=0
            WHERE id=?
            """,
            (novo_hash, int(usuario_id)),
        )
        conn.commit()
        usuario_row = conn.execute("SELECT * FROM usuarios WHERE id=?", (int(usuario_id),)).fetchone()
    finally:
        conn.close()

    registrar_acao_usuario(
        usuario_request(request) or str(usuario_row["usuario"] or usuario_row["nome"] or ""),
        acao="Redefinicao de senha",
        tipo="Usuario",
        info=str(usuario_row["nome"] or ""),
        metodo_http="POST",
        rota=f"/api/usuarios/{usuario_id}/redefinir-senha",
    )
    return detalhar_usuario_login(usuario_row)


@app.get("/api/usuarios", response_model=list[UsuarioResumoResposta])
def listar_usuarios():
    conn = conectar()
    try:
        rows = conn.execute(
            """
            SELECT *
            FROM usuarios
            ORDER BY lower(COALESCE(nome, usuario, '')), id
            """
        ).fetchall()
        return [mapear_usuario_resumo(row) for row in rows]
    finally:
        conn.close()


@app.post("/api/usuarios", response_model=UsuarioResumoResposta)
def criar_usuario(payload: UsuarioCriacaoPayload, request: Request):
    nome = corrigir_texto_importado(str(payload.nome or "").strip())
    cargo = corrigir_texto_importado(str(payload.cargo or "").strip()) or "Profissional"
    perfil = corrigir_texto_importado(str(payload.perfil or "").strip()) or ("Administrador" if cargo == "Administrador" else "Usuario")
    nome_agenda = corrigir_texto_importado(str(payload.nome_agenda or "").strip()) or nome.upper()
    if not nome:
        raise HTTPException(status_code=400, detail="Informe o nome do usuario.")

    usuario_login = normalizar_texto(str(payload.nome or "").strip()).replace(" ", "")
    if not usuario_login:
        raise HTTPException(status_code=400, detail="Nao foi possivel gerar o usuario de login.")

    modulos, pacientes_abas, agenda_padrao = normalizar_payload_permissoes(cargo, perfil, payload.modulos, payload.pacientes_abas)
    agenda_escopo = corrigir_texto_importado(str(payload.agenda_escopo or "").strip()) or agenda_padrao
    agenda_disponivel = bool(payload.agenda_disponivel if payload.agenda_disponivel is not None else cargo in {"Profissional", "Administrador"})
    ativo = 1 if (payload.ativo if payload.ativo is not None else True) else 0

    conn = conectar()
    try:
        existente = conn.execute("SELECT id FROM usuarios WHERE lower(usuario)=lower(?) LIMIT 1", (usuario_login,)).fetchone()
        if existente is not None:
            raise HTTPException(status_code=409, detail="Ja existe um usuario com esse login.")
        cursor = conn.execute(
            """
            INSERT INTO usuarios
            (nome, usuario, nome_agenda, senha_hash, perfil, cargo, agenda_escopo, agenda_disponivel, senha_temporaria, precisa_trocar_senha, ultimo_login, ativo, modulos_json, pacientes_abas_json, data_criacao)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, 1, '', ?, ?, ?, ?)
            """,
            (
                nome,
                usuario_login,
                nome_agenda.upper(),
                gerar_hash_senha(SENHA_PADRAO_USUARIOS),
                perfil,
                cargo,
                agenda_escopo,
                1 if agenda_disponivel else 0,
                ativo,
                json.dumps(modulos, ensure_ascii=False),
                json.dumps(pacientes_abas, ensure_ascii=False),
                agora_str(),
            ),
        )
        usuario_id = int(cursor.lastrowid)
        conn.commit()
        usuario_row = conn.execute("SELECT * FROM usuarios WHERE id=?", (usuario_id,)).fetchone()
    except HTTPException:
        conn.rollback()
        raise
    except sqlite3.IntegrityError as exc:
        conn.rollback()
        raise HTTPException(status_code=400, detail=f"Falha ao criar o usuario: {str(exc)}")
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"Falha ao criar o usuario: {str(exc)}")
    finally:
        conn.close()

    registrar_acao_usuario(
        usuario_request(request),
        acao="Criacao",
        tipo="Usuario",
        info=nome,
        metodo_http="POST",
        rota="/api/usuarios",
    )
    return mapear_usuario_resumo(usuario_row)


@app.put("/api/usuarios/{usuario_id}", response_model=UsuarioResumoResposta)
def atualizar_usuario(usuario_id: int, payload: UsuarioAtualizacaoPayload, request: Request):
    nome = corrigir_texto_importado(str(payload.nome or "").strip())
    nome_agenda = corrigir_texto_importado(str(payload.nome_agenda or "").strip())
    usuario_login = normalizar_texto(str(payload.usuario or "").strip()).replace(" ", "")
    cargo = corrigir_texto_importado(str(payload.cargo or "").strip())
    agenda_escopo = corrigir_texto_importado(str(payload.agenda_escopo or "").strip())
    perfil = corrigir_texto_importado(str(payload.perfil or "").strip())

    if not nome:
        raise HTTPException(status_code=400, detail="Informe o nome do usuario.")

    conn = conectar()
    try:
        usuario_row = conn.execute(
            """
            SELECT *
            FROM usuarios
            WHERE id=?
            LIMIT 1
            """,
            (int(usuario_id),),
        ).fetchone()
        if usuario_row is None:
            raise HTTPException(status_code=404, detail="Usuario nao encontrado.")

        if usuario_login:
            existente = conn.execute(
                "SELECT id FROM usuarios WHERE lower(usuario)=lower(?) AND id<>? LIMIT 1",
                (usuario_login, int(usuario_id)),
            ).fetchone()
            if existente is not None:
                raise HTTPException(status_code=409, detail="Ja existe um usuario com esse login.")

        cargo_final = cargo or str(usuario_row["cargo"] or "Profissional")
        perfil_final = perfil or str(usuario_row["perfil"] or ("Administrador" if cargo_final == "Administrador" else "Usuario"))
        modulos, pacientes_abas, agenda_padrao = normalizar_payload_permissoes(
            cargo_final,
            perfil_final,
            payload.modulos,
            payload.pacientes_abas,
        )

        conn.execute(
            """
            UPDATE usuarios
            SET nome=?,
                usuario=?,
                nome_agenda=?,
                perfil=?,
                cargo=?,
                agenda_escopo=?,
                agenda_disponivel=COALESCE(?, agenda_disponivel),
                ativo=COALESCE(?, ativo),
                modulos_json=?,
                pacientes_abas_json=?
            WHERE id=?
            """,
            (
                nome,
                usuario_login or str(usuario_row["usuario"] or ""),
                (nome_agenda or str(usuario_row["nome_agenda"] or "") or nome).upper(),
                perfil_final,
                cargo_final,
                agenda_escopo or str(usuario_row["agenda_escopo"] or "") or agenda_padrao,
                1 if payload.agenda_disponivel else 0 if payload.agenda_disponivel is not None else None,
                1 if payload.ativo else 0 if payload.ativo is not None else None,
                json.dumps(modulos, ensure_ascii=False),
                json.dumps(pacientes_abas, ensure_ascii=False),
                int(usuario_id),
            ),
        )
        conn.commit()
        usuario_row = conn.execute("SELECT * FROM usuarios WHERE id=?", (int(usuario_id),)).fetchone()
    finally:
        conn.close()

    registrar_acao_usuario(
        usuario_request(request),
        acao="Atualizacao",
        tipo="Usuario",
        info=nome,
        metodo_http="PUT",
        rota=f"/api/usuarios/{usuario_id}",
    )
    return mapear_usuario_resumo(usuario_row)


@app.delete("/api/usuarios/{usuario_id}")
def excluir_usuario(usuario_id: int, request: Request):
    conn = conectar()
    try:
        usuario_row = conn.execute(
            """
            SELECT *
            FROM usuarios
            WHERE id=?
            LIMIT 1
            """,
            (int(usuario_id),),
        ).fetchone()
        if usuario_row is None:
            raise HTTPException(status_code=404, detail="Usuario nao encontrado.")

        if str(usuario_row["perfil"] or "") == "Administrador" and str(usuario_row["usuario"] or "").strip().lower() == "admin":
            raise HTTPException(status_code=400, detail="Nao e permitido excluir o administrador padrao.")

        conn.execute("DELETE FROM usuarios WHERE id=?", (int(usuario_id),))
        conn.commit()
    finally:
        conn.close()

    registrar_acao_usuario(
        usuario_request(request),
        acao="Exclusao",
        tipo="Usuario",
        info=str(usuario_row["nome"] or usuario_row["usuario"] or f"Usuario {usuario_id}"),
        metodo_http="DELETE",
        rota=f"/api/usuarios/{usuario_id}",
    )
    return {"ok": True}


@app.get("/api/usuarios/acoes/export.xlsx")
def exportar_acoes_usuarios_xlsx(data: str = Query(default="")):
    conteudo = bytes_relatorio_acoes_por_dia(data)
    nome = f"acoes_usuarios_{(parse_data_contrato(data) or date.today()).strftime('%Y%m%d')}.xlsx"
    headers = {"Content-Disposition": f'attachment; filename="{nome}"'}
    return Response(
        content=conteudo,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers,
    )


@app.get("/api/procedimentos", response_model=list[ProcedimentoResumoResposta])
def listar_procedimentos(q: str = Query(default=""), ativos_apenas: bool = Query(default=True)):
    termo = normalizar_texto(q)
    conn = conectar()
    try:
        rows = conn.execute(
            """
            SELECT *
            FROM procedimentos
            WHERE (? = '' OR lower(COALESCE(nome, '')) LIKE ? OR lower(COALESCE(categoria, '')) LIKE ?)
              AND (? = 0 OR COALESCE(ativo, 1) = 1)
            ORDER BY lower(COALESCE(nome, '')), id
            """,
            (termo, f"%{termo}%", f"%{termo}%", 1 if ativos_apenas else 0),
        ).fetchall()
    finally:
        conn.close()
    return [mapear_procedimento_resumo(row) for row in rows]


@app.post("/api/procedimentos", response_model=ProcedimentoResumoResposta)
def criar_procedimento(payload: ProcedimentoPayload, request: Request):
    nome = corrigir_texto_importado(str(payload.nome or "").strip())
    categoria = corrigir_texto_importado(str(payload.categoria or "").strip())
    descricao = corrigir_texto_importado(str(payload.descricao or "").strip())
    materiais_padrao = [corrigir_texto_importado(str(item).strip()) for item in payload.materiais_padrao if str(item).strip()]
    if not nome:
        raise HTTPException(status_code=400, detail="Informe o nome do procedimento.")

    conn = conectar()
    try:
        existente = conn.execute(
            "SELECT id FROM procedimentos WHERE lower(COALESCE(nome, ''))=lower(?) LIMIT 1",
            (nome,),
        ).fetchone()
        if existente is not None:
            raise HTTPException(status_code=409, detail="Ja existe um procedimento com esse nome.")
        cursor = conn.execute(
            """
            INSERT INTO procedimentos
            (nome, categoria, valor_padrao, duracao_padrao_minutos, descricao, etapas_json, materiais_json, cor_opcional, ativo, criado_em, atualizado_em)
            VALUES (?, ?, ?, ?, ?, ?, ?, '', ?, ?, ?)
            """,
            (
                nome,
                categoria,
                float(payload.valor_padrao or 0),
                max(5, int(payload.duracao_padrao_minutos or 60)),
                descricao,
                json.dumps([corrigir_texto_importado(str(item).strip()) for item in payload.etapas_padrao if str(item).strip()], ensure_ascii=False),
                json.dumps(materiais_padrao, ensure_ascii=False),
                1 if payload.ativo else 0,
                agora_str(),
                agora_str(),
            ),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM procedimentos WHERE id=?", (int(cursor.lastrowid),)).fetchone()
    finally:
        conn.close()
    registrar_acao_usuario(
        usuario_request(request),
        acao="Criacao",
        tipo="Procedimento",
        info=nome,
        metodo_http="POST",
        rota="/api/procedimentos",
    )
    return mapear_procedimento_resumo(row)


@app.put("/api/procedimentos/{procedimento_id}", response_model=ProcedimentoResumoResposta)
def atualizar_procedimento(procedimento_id: int, payload: ProcedimentoPayload, request: Request):
    nome = corrigir_texto_importado(str(payload.nome or "").strip())
    categoria = corrigir_texto_importado(str(payload.categoria or "").strip())
    descricao = corrigir_texto_importado(str(payload.descricao or "").strip())
    materiais_padrao = [corrigir_texto_importado(str(item).strip()) for item in payload.materiais_padrao if str(item).strip()]
    if not nome:
        raise HTTPException(status_code=400, detail="Informe o nome do procedimento.")

    conn = conectar()
    try:
        existente = conn.execute("SELECT id FROM procedimentos WHERE id=?", (int(procedimento_id),)).fetchone()
        if existente is None:
            raise HTTPException(status_code=404, detail="Procedimento nao encontrado.")
        duplicado = conn.execute(
            "SELECT id FROM procedimentos WHERE lower(COALESCE(nome, ''))=lower(?) AND id<>? LIMIT 1",
            (nome, int(procedimento_id)),
        ).fetchone()
        if duplicado is not None:
            raise HTTPException(status_code=409, detail="Ja existe outro procedimento com esse nome.")
        conn.execute(
            """
            UPDATE procedimentos
            SET nome=?, categoria=?, valor_padrao=?, duracao_padrao_minutos=?, descricao=?, etapas_json=?, materiais_json=?, ativo=?, atualizado_em=?
            WHERE id=?
            """,
            (
                nome,
                categoria,
                float(payload.valor_padrao or 0),
                max(5, int(payload.duracao_padrao_minutos or 60)),
                descricao,
                json.dumps([corrigir_texto_importado(str(item).strip()) for item in payload.etapas_padrao if str(item).strip()], ensure_ascii=False),
                json.dumps(materiais_padrao, ensure_ascii=False),
                1 if payload.ativo else 0,
                agora_str(),
                int(procedimento_id),
            ),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM procedimentos WHERE id=?", (int(procedimento_id),)).fetchone()
    finally:
        conn.close()
    registrar_acao_usuario(
        usuario_request(request),
        acao="Edicao",
        tipo="Procedimento",
        info=nome,
        metodo_http="PUT",
        rota=f"/api/procedimentos/{procedimento_id}",
    )
    return mapear_procedimento_resumo(row)


@app.delete("/api/procedimentos/{procedimento_id}")
def excluir_procedimento(procedimento_id: int, request: Request):
    conn = conectar()
    try:
        row = conn.execute("SELECT nome FROM procedimentos WHERE id=?", (int(procedimento_id),)).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="Procedimento nao encontrado.")
        conn.execute("DELETE FROM procedimentos WHERE id=?", (int(procedimento_id),))
        conn.commit()
        nome = str(row["nome"] or "")
    finally:
        conn.close()
    registrar_acao_usuario(
        usuario_request(request),
        acao="Exclusao",
        tipo="Procedimento",
        info=nome,
        metodo_http="DELETE",
        rota=f"/api/procedimentos/{procedimento_id}",
    )
    return {"ok": True}


@app.get("/api/pacientes/{paciente_id}/ordens-servico", response_model=list[OrdemServicoResumoResposta])
def listar_ordens_servico_paciente(paciente_id: int):
    conn = conectar()
    try:
        ordens = conn.execute(
            """
            SELECT *
            FROM ordens_servico_protetico
            WHERE paciente_id=?
            ORDER BY COALESCE(criado_em, '') DESC, id DESC
            """,
            (int(paciente_id),),
        ).fetchall()
        respostas: list[OrdemServicoResumoResposta] = []
        for ordem in ordens:
            etapas_rows = conn.execute(
                """
                SELECT etapa, descricao_outro
                FROM ordem_servico_protetico_etapas
                WHERE ordem_servico_id=?
                ORDER BY id
                """,
                (int(ordem["id"]),),
            ).fetchall()
            respostas.append(
                OrdemServicoResumoResposta(
                    id=int(ordem["id"]),
                    procedimentoId=int(ordem["procedimento_id"]) if ordem["procedimento_id"] is not None else None,
                    procedimentoNome=corrigir_texto_importado(str(ordem["procedimento_nome_snapshot"] or "")),
                    material=corrigir_texto_importado(str(ordem["material"] or "")),
                    materialOutro=corrigir_texto_importado(str(ordem["material_outro"] or "")),
                    cor=corrigir_texto_importado(str(ordem["cor"] or "")),
                    escala=corrigir_texto_importado(str(ordem["escala"] or "")),
                    elementoArcada=corrigir_texto_importado(str(ordem["elemento_arcada"] or "")),
                    cargaImediata=bool(ordem["carga_imediata"] or 0),
                    retornoSolicitado=str(ordem["retorno_solicitado"] or ""),
                    documentoNome=str(ordem["documento_nome"] or ""),
                    observacao=corrigir_texto_importado(str(ordem["observacao"] or "")),
                    criadoEm=str(ordem["criado_em"] or ""),
                    etapas=[
                        OrdemServicoEtapaPayload(
                            etapa=corrigir_texto_importado(str(item["etapa"] or "")),
                            descricao_outro=corrigir_texto_importado(str(item["descricao_outro"] or "")),
                        )
                        for item in etapas_rows
                    ],
                )
            )
    finally:
        conn.close()
    return respostas


@app.post("/api/pacientes/{paciente_id}/ordens-servico", response_model=OrdemServicoResumoResposta)
def criar_ordem_servico_paciente(paciente_id: int, payload: OrdemServicoPayload, request: Request):
    conn = conectar()
    try:
        paciente = conn.execute("SELECT id, nome, prontuario FROM pacientes WHERE id=?", (int(paciente_id),)).fetchone()
        if paciente is None:
            raise HTTPException(status_code=404, detail="Paciente nao encontrado.")
        procedimento = conn.execute("SELECT * FROM procedimentos WHERE id=? AND COALESCE(ativo,1)=1", (int(payload.procedimento_id),)).fetchone()
        if procedimento is None:
            raise HTTPException(status_code=404, detail="Procedimento nao encontrado.")
        procedimentos_contratados_rows = conn.execute(
            """
            SELECT pc.procedimento
            FROM procedimentos_contrato pc
            JOIN contratos c ON c.id = pc.contrato_id
            WHERE c.paciente_id=? AND COALESCE(c.status, 'EM_ABERTO')='APROVADO'
            """,
            (int(paciente_id),),
        ).fetchall()
        procedimentos_contratados = {
            normalizar_texto(str(item["procedimento"] or ""))
            for item in procedimentos_contratados_rows
            if str(item["procedimento"] or "").strip()
        }
        if not procedimentos_contratados:
            raise HTTPException(status_code=400, detail="Este paciente nao possui procedimentos contratados para ordem de servico.")
        if normalizar_texto(str(procedimento["nome"] or "")) not in procedimentos_contratados:
            raise HTTPException(status_code=400, detail="Selecione apenas um procedimento contratado para este paciente.")
        material = corrigir_texto_importado(str(payload.material or "").strip())
        material_outro = corrigir_texto_importado(str(payload.material_outro or "").strip())
        cor = corrigir_texto_importado(str(payload.cor or "").strip())
        escala = corrigir_texto_importado(str(payload.escala or "").strip())
        elemento_arcada = corrigir_texto_importado(str(payload.elemento_arcada or "").strip())
        retorno_solicitado = str(payload.retorno_solicitado or "").strip()
        carga_imediata = bool(payload.carga_imediata)
        if not material:
            raise HTTPException(status_code=400, detail="Selecione o material.")
        materiais_validos = mapear_procedimento_resumo(procedimento).materiaisPadrao
        if materiais_validos and material not in materiais_validos:
            raise HTTPException(status_code=400, detail="Selecione um material permitido para o procedimento.")
        if normalizar_texto(material) == "outro" and not material_outro:
            raise HTTPException(status_code=400, detail="Descreva o material quando selecionar Outro.")
        if not elemento_arcada:
            raise HTTPException(status_code=400, detail="Informe o elemento ou arcada.")
        if not retorno_solicitado:
            raise HTTPException(status_code=400, detail="Informe a data de retorno solicitada.")
        etapas_validas = mapear_procedimento_resumo(procedimento).etapasPadrao
        etapas_limpas: list[tuple[str, str]] = []
        for item in payload.etapas:
            etapa = corrigir_texto_importado(str(item.etapa or "").strip())
            descricao_outro = corrigir_texto_importado(str(item.descricao_outro or "").strip())
            if not etapa:
                continue
            if etapa.lower() == "outro":
                if not descricao_outro:
                    raise HTTPException(status_code=400, detail="Descreva a etapa quando selecionar Outro.")
            elif etapas_validas and etapa not in etapas_validas:
                raise HTTPException(status_code=400, detail="Etapa invalida para o procedimento selecionado.")
            etapas_limpas.append((etapa, descricao_outro))
        if not etapas_limpas:
            raise HTTPException(status_code=400, detail="Selecione ao menos uma etapa.")

        cursor = conn.execute(
            """
            INSERT INTO ordens_servico_protetico
            (paciente_id, procedimento_id, procedimento_nome_snapshot, material, material_outro, cor, escala, elemento_arcada, carga_imediata, retorno_solicitado, observacao, criado_em, atualizado_em)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                int(paciente_id),
                int(payload.procedimento_id),
                str(procedimento["nome"] or ""),
                material,
                material_outro,
                cor,
                escala,
                elemento_arcada,
                1 if carga_imediata else 0,
                retorno_solicitado,
                corrigir_texto_importado(str(payload.observacao or "").strip()),
                agora_str(),
                agora_str(),
            ),
        )
        ordem_id = int(cursor.lastrowid)
        for etapa, descricao_outro in etapas_limpas:
            conn.execute(
                """
                INSERT INTO ordem_servico_protetico_etapas
                (ordem_servico_id, etapa, descricao_outro, criado_em)
                VALUES (?, ?, ?, ?)
                """,
                (ordem_id, etapa, descricao_outro, agora_str()),
            )
        documento_nome = gerar_documento_ordem_servico(
            paciente_row=paciente,
            ordem_id=ordem_id,
            procedimento_nome=corrigir_texto_importado(str(procedimento["nome"] or "")),
            material=material,
            material_outro=material_outro,
            elemento_arcada=elemento_arcada,
            escala=escala,
            cor=cor,
            carga_imediata=carga_imediata,
            retorno_solicitado=retorno_solicitado,
            etapas=etapas_limpas,
            observacao=corrigir_texto_importado(str(payload.observacao or "").strip()),
        )
        conn.execute("UPDATE ordens_servico_protetico SET documento_nome=?, atualizado_em=? WHERE id=?", (documento_nome, agora_str(), ordem_id))
        conn.commit()
        ordem = conn.execute("SELECT * FROM ordens_servico_protetico WHERE id=?", (ordem_id,)).fetchone()
    finally:
        conn.close()

    registrar_acao_usuario(
        usuario_request(request),
        acao="Criacao",
        tipo="Ordem de Servico",
        info=str(procedimento["nome"] or ""),
        metodo_http="POST",
        rota=f"/api/pacientes/{paciente_id}/ordens-servico",
    )

    return OrdemServicoResumoResposta(
        id=int(ordem["id"]),
        procedimentoId=int(ordem["procedimento_id"]) if ordem["procedimento_id"] is not None else None,
        procedimentoNome=corrigir_texto_importado(str(ordem["procedimento_nome_snapshot"] or "")),
        material=corrigir_texto_importado(str(ordem["material"] or "")),
        materialOutro=corrigir_texto_importado(str(ordem["material_outro"] or "")),
        cor=corrigir_texto_importado(str(ordem["cor"] or "")),
        escala=corrigir_texto_importado(str(ordem["escala"] or "")),
        elementoArcada=corrigir_texto_importado(str(ordem["elemento_arcada"] or "")),
        cargaImediata=bool(ordem["carga_imediata"] or 0),
        retornoSolicitado=str(ordem["retorno_solicitado"] or ""),
        documentoNome=str(ordem["documento_nome"] or ""),
        observacao=corrigir_texto_importado(str(ordem["observacao"] or "")),
        criadoEm=str(ordem["criado_em"] or ""),
        etapas=[OrdemServicoEtapaPayload(etapa=etapa, descricao_outro=descricao_outro) for etapa, descricao_outro in etapas_limpas],
    )


def mapear_paciente_resumo(row: sqlite3.Row) -> PacienteResumo:
    return PacienteResumo(
        id=int(row["id"]),
        nome=str(row["nome"] or ""),
        apelido=str(row["apelido"] or ""),
        prontuario=formatar_prontuario_valor(row["prontuario"]),
        cpf=str(row["cpf"] or ""),
        telefone=str(row["telefone"] or ""),
        email=str(row["email"] or ""),
        dataNascimento=formatar_data_br_valor(row["data_nascimento"]),
        fotoUrl=url_foto_paciente(row),
    )


def mapear_paciente_detalhe(row: sqlite3.Row) -> PacienteDetalhe:
    return PacienteDetalhe(
        id=int(row["id"]),
        nome=str(row["nome"] or ""),
        apelido=str(row["apelido"] or ""),
        sexo=str(row["sexo"] or ""),
        prontuario=formatar_prontuario_valor(row["prontuario"]),
        cpf=str(row["cpf"] or ""),
        rg=str(row["rg"] or ""),
        dataNascimento=formatar_data_br_valor(row["data_nascimento"]),
        telefone=str(row["telefone"] or ""),
        email=str(row["email"] or ""),
        cep=str(row["cep"] or ""),
        endereco=str(row["endereco"] or ""),
        complemento=str(row["complemento"] or ""),
        numero=str(row["numero"] or ""),
        bairro=str(row["bairro"] or ""),
        cidade=str(row["cidade"] or ""),
        estado=str(row["estado"] or ""),
        estadoCivil=str(row["estado_civil"] or ""),
        profissao=str(row["profissao"] or ""),
        origem=str(row["origem"] or ""),
        observacoes=str(row["observacoes"] or ""),
        menorIdade=bool(int(row["menor_idade"] or 0)),
        responsavel=str(row["responsavel"] or ""),
        cpfResponsavel=str(row["cpf_responsavel"] or ""),
        fotoUrl=url_foto_paciente(row),
    )


def carregar_paciente_por_id(conn: sqlite3.Connection, paciente_id: int) -> sqlite3.Row:
    row = conn.execute("SELECT * FROM pacientes WHERE id=?", (paciente_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Paciente nao encontrado.")
    return row


def filtrar_pacientes(rows: list[sqlite3.Row], termo: str) -> list[sqlite3.Row]:
    termo_norm = normalizar_texto(termo)
    if not termo_norm:
        return rows

    def combina(row: sqlite3.Row) -> bool:
        campos = [
            row["nome"],
            row["apelido"],
            formatar_prontuario_valor(row["prontuario"]),
            row["cpf"],
            row["telefone"],
            row["email"],
        ]
        return termo_norm in normalizar_texto(" ".join(str(campo or "") for campo in campos))

    return [row for row in rows if combina(row)]


def carregar_procedimentos_contrato(conn: sqlite3.Connection, contrato_id: int) -> list[str]:
    rows = conn.execute(
        "SELECT procedimento FROM procedimentos_contrato WHERE contrato_id=? ORDER BY id",
        (contrato_id,),
    ).fetchall()
    return [str(row["procedimento"] or "") for row in rows if str(row["procedimento"] or "").strip()]


def carregar_dentes_contratados(conn: sqlite3.Connection, paciente_id: int) -> list[int]:
    rows = conn.execute(
        """
        SELECT DISTINCT dente
        FROM procedimentos_dente
        WHERE paciente_id=? AND upper(trim(COALESCE(status, '')))='CONTRATADO' AND dente IS NOT NULL
        ORDER BY dente
        """,
        (paciente_id,),
    ).fetchall()
    return [int(row["dente"]) for row in rows if row["dente"] is not None]


def denticao_por_dente(dente: int | None) -> str:
    if dente is None:
        return ""
    return "Decidua" if 50 < int(dente) < 90 else "Permanente"


def carregar_elementos_odontograma(conn: sqlite3.Connection, paciente_id: int) -> list[dict]:
    rows = conn.execute(
        """
        SELECT dente, regiao, procedimento
        FROM procedimentos_dente
        WHERE paciente_id=? AND upper(trim(COALESCE(status, '')))='CONTRATADO'
        ORDER BY COALESCE(dente, 999), regiao, procedimento
        """,
        (paciente_id,),
    ).fetchall()
    agrupado: dict[str, dict] = {}
    for row in rows:
        dente = int(row["dente"]) if row["dente"] is not None else None
        elemento = str(dente) if dente is not None else str(row["regiao"] or "").strip()
        if not elemento:
            continue
        item = agrupado.setdefault(
            elemento,
            {
                "elemento": elemento,
                "dente": dente,
                "denticao": denticao_por_dente(dente),
                "procedimentos": [],
            },
        )
        procedimento = str(row["procedimento"] or "").strip()
        if procedimento and procedimento not in item["procedimentos"]:
            item["procedimentos"].append(procedimento)

    return sorted(
        agrupado.values(),
        key=lambda item: (
            0 if item["dente"] is not None else 1,
            item["dente"] if item["dente"] is not None else item["elemento"],
        ),
    )


def carregar_orcamento_detalhe(conn: sqlite3.Connection, paciente_id: int, contrato_id: int) -> OrcamentoDetalheResposta:
    contrato = conn.execute(
        "SELECT * FROM contratos WHERE id=? AND paciente_id=? LIMIT 1",
        (contrato_id, paciente_id),
    ).fetchone()
    if contrato is None:
        raise HTTPException(status_code=404, detail="Orcamento nao encontrado.")

    procedimentos_rows = conn.execute(
        "SELECT * FROM procedimentos_contrato WHERE contrato_id=? ORDER BY id",
        (contrato_id,),
    ).fetchall()
    dentes_rows = conn.execute(
        "SELECT * FROM procedimentos_dente WHERE contrato_id=? ORDER BY id",
        (contrato_id,),
    ).fetchall()

    dentes_por_grupo: dict[int, list[sqlite3.Row]] = {}
    dentes_sem_grupo_por_procedimento: dict[str, list[sqlite3.Row]] = {}
    for row in dentes_rows:
        if row["grupo_item"] is not None:
            dentes_por_grupo.setdefault(int(row["grupo_item"]), []).append(row)
        else:
            dentes_sem_grupo_por_procedimento.setdefault(normalizar_texto(row["procedimento"]), []).append(row)

    itens: list[OrcamentoItemPayload] = []
    for indice_proc, proc_row in enumerate(procedimentos_rows, start=1):
        procedimento_nome = str(proc_row["procedimento"] or "").strip()
        grupo_rows = dentes_por_grupo.get(indice_proc)
        if grupo_rows is None:
            chave = normalizar_texto(procedimento_nome)
            restantes = dentes_sem_grupo_por_procedimento.get(chave, [])
            alvo = float(proc_row["valor"] or 0)
            if not restantes:
                grupo_rows = []
            elif alvo <= 0:
                grupo_rows = restantes
                dentes_sem_grupo_por_procedimento[chave] = []
            else:
                grupo_rows = []
                acumulado = 0.0
                while restantes:
                    row = restantes.pop(0)
                    grupo_rows.append(row)
                    acumulado += float(row["valor"] or 0)
                    if acumulado + 0.009 >= alvo:
                        break
                dentes_sem_grupo_por_procedimento[chave] = restantes
        regioes = [
            OrcamentoRegiaoPayload(
                regiao=str(row["regiao"] or row["dente"] or "").strip(),
                dente=int(row["dente"]) if row["dente"] is not None else None,
                valor=float(row["valor"] or 0),
                ativo=normalizar_texto(row["status"]) != "excluido",
                faces=[face for face in str(row["faces"] or "").split(",") if face],
            )
            for row in grupo_rows
        ]
        itens.append(
            OrcamentoItemPayload(
                procedimento=procedimento_nome,
                profissional=str(proc_row["profissional_snapshot"] or ""),
                denticao=str(proc_row["denticao_snapshot"] or ""),
                valor_unitario=float(proc_row["valor"] or 0),
                regioes=regioes,
            )
        )

    plano_pagamento_bruto = []
    try:
        plano_pagamento_bruto = json.loads(str(contrato["plano_pagamento_json"] or "[]"))
    except (TypeError, ValueError, json.JSONDecodeError):
        plano_pagamento_bruto = []
    plano_pagamento = [
        ParcelaPagamentoPayload(
            indice=int(item.get("indice", 0) or 0),
            descricao=str(item.get("descricao", "") or ""),
            data=str(item.get("data", "") or ""),
            forma=str(item.get("forma", "") or ""),
            valor=float(item.get("valor", 0) or 0),
            parcelas_cartao=int(item.get("parcelas_cartao", 1) or 1),
        )
        for item in plano_pagamento_bruto
        if isinstance(item, dict)
    ]

    return OrcamentoDetalheResposta(
        contrato_id=contrato_id,
        status=str(contrato["status"] or "EM_ABERTO"),
        aprovadoPor=str(contrato["aprovado_por"] or ""),
        dataAprovacao=formatar_data_br_valor(contrato["data_aprovacao"]),
        clinica=str(contrato["clinica_snapshot"] or ""),
        criadoPor=str(contrato["criado_por_snapshot"] or ""),
        data=formatar_data_br_valor(contrato["data_criacao"]),
        observacoes=str(contrato["observacoes"] or ""),
        tabela=str(contrato["tabela_snapshot"] or ""),
        descontoPercentual=float(contrato["desconto_percentual"] or 0),
        descontoValor=float(contrato["desconto_valor"] or 0),
        validadeOrcamento=formatar_data_br_valor(contrato["validade_orcamento"]),
        formaPagamento=str(contrato["forma_pagamento"] or ""),
        parcelas=int(contrato["parcelas"] or 1),
        entrada=float(contrato["entrada"] or 0) > 0,
        planoPagamento=plano_pagamento,
        itens=itens,
    )


def salvar_orcamento_paciente(
    conn: sqlite3.Connection,
    paciente_id: int,
    payload: OrcamentoPacientePayload,
    contrato_id: int | None = None,
) -> int:
    itens_validos = []
    for item in payload.itens:
        regioes_validas = [regiao for regiao in item.regioes if str(regiao.regiao or "").strip()]
        if regioes_validas and str(item.procedimento or "").strip():
            itens_validos.append((item, regioes_validas))

    if not any(regiao.ativo for item, regioes in itens_validos for regiao in regioes):
        raise HTTPException(status_code=400, detail="Informe ao menos um procedimento com dentes/regioes ativos.")

    valor_bruto = sum(sum(float(regiao.valor or 0) for regiao in regioes if regiao.ativo) for _, regioes in itens_validos)
    desconto_percentual = max(0.0, float(payload.desconto_percentual or 0))
    desconto_valor = max(0.0, float(payload.desconto_valor or 0))
    valor_total = max(0.0, valor_bruto - (valor_bruto * desconto_percentual / 100.0) - desconto_valor)
    data_base = payload.data.strip() or agora_str()
    parcelas_plano = [
        parcela
        for parcela in payload.plano_pagamento
        if str(parcela.forma or "").strip() and float(parcela.valor or 0) >= 0
    ]
    forma_pagamento = payload.forma_pagamento.strip() or "A Definir"
    quantidade_parcelas = max(int(payload.parcelas or 1), 1)
    entrada_valor = sum(float(parcela.valor or 0) for parcela in parcelas_plano if normalizar_texto(parcela.descricao) == "entrada")
    primeiro_vencimento = next(
        (parcela.data.strip() for parcela in parcelas_plano if normalizar_texto(parcela.descricao) != "entrada" and parcela.data.strip()),
        data_base,
    )
    data_pagamento_entrada = next(
        (parcela.data.strip() for parcela in parcelas_plano if normalizar_texto(parcela.descricao) == "entrada" and parcela.data.strip()),
        None,
    )
    plano_pagamento_json = json.dumps([parcela.model_dump() for parcela in parcelas_plano], ensure_ascii=False)

    if contrato_id is None:
        cursor = conn.execute(
            """
            INSERT INTO contratos (
                paciente_id, valor_total, entrada, parcelas, primeiro_vencimento, data_pagamento_entrada,
                forma_pagamento, hash_importacao, data_criacao, status, observacoes, clinica_snapshot, criado_por_snapshot, tabela_snapshot, desconto_percentual, desconto_valor, validade_orcamento
            )
            VALUES (?, ?, 0, 1, ?, NULL, ?, NULL, ?, 'EM_ABERTO', ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                paciente_id,
                valor_total,
                data_base,
                "OrÃ§amento",
                data_base,
                payload.observacoes.strip(),
                payload.clinica.strip(),
                payload.criado_por.strip(),
                payload.tabela.strip(),
                desconto_percentual,
                desconto_valor,
                payload.validade_orcamento.strip(),
            ),
        )
        contrato_id = int(cursor.lastrowid)
    else:
        contrato = conn.execute("SELECT * FROM contratos WHERE id=? AND paciente_id=? LIMIT 1", (contrato_id, paciente_id)).fetchone()
        if contrato is None:
            raise HTTPException(status_code=404, detail="Orcamento nao encontrado.")
        if normalizar_texto(contrato["status"]) == "aprovado":
            raise HTTPException(status_code=400, detail="Orcamento aprovado nao pode ser editado.")
        conn.execute(
            """
            UPDATE contratos
            SET valor_total=?, primeiro_vencimento=?, forma_pagamento=?, data_criacao=?, observacoes=?, clinica_snapshot=?, criado_por_snapshot=?, tabela_snapshot=?, desconto_percentual=?, desconto_valor=?, validade_orcamento=?
            WHERE id=? AND paciente_id=?
            """,
            (
                valor_total,
                data_base,
                "OrÃ§amento",
                data_base,
                payload.observacoes.strip(),
                payload.clinica.strip(),
                payload.criado_por.strip(),
                payload.tabela.strip(),
                desconto_percentual,
                desconto_valor,
                payload.validade_orcamento.strip(),
                contrato_id,
                paciente_id,
            ),
        )
        conn.execute("DELETE FROM procedimentos_contrato WHERE contrato_id=?", (contrato_id,))
        conn.execute("DELETE FROM procedimentos_dente WHERE contrato_id=?", (contrato_id,))

    for indice_item, (item, regioes) in enumerate(itens_validos, start=1):
        subtotal = sum(float(regiao.valor or 0) for regiao in regioes if regiao.ativo)
        conn.execute(
            """
            INSERT INTO procedimentos_contrato (contrato_id, procedimento, valor, profissional_snapshot, denticao_snapshot)
            VALUES (?, ?, ?, ?, ?)
            """,
            (contrato_id, item.procedimento.strip(), subtotal, item.profissional.strip(), item.denticao.strip()),
        )
        for regiao in regioes:
            dente = regiao.dente if regiao.dente is not None else None
            conn.execute(
                """
                INSERT INTO procedimentos_dente (
                    paciente_id, contrato_id, grupo_item, dente, regiao, procedimento, status, faces, valor, data
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    paciente_id,
                    contrato_id,
                    indice_item,
                    dente,
                    regiao.regiao.strip(),
                    item.procedimento.strip(),
                    "ORCAMENTO" if regiao.ativo else "EXCLUIDO",
                    ",".join(face.strip() for face in regiao.faces if face.strip()),
                    float(regiao.valor or 0),
                    data_base,
                ),
            )
    return contrato_id


def salvar_orcamento_paciente_com_pagamento(
    conn: sqlite3.Connection,
    paciente_id: int,
    payload: OrcamentoPacientePayload,
    contrato_id: int | None = None,
) -> int:
    itens_validos = []
    for item in payload.itens:
        regioes_validas = [regiao for regiao in item.regioes if str(regiao.regiao or "").strip()]
        if regioes_validas and str(item.procedimento or "").strip():
            itens_validos.append((item, regioes_validas))

    if not any(regiao.ativo for item, regioes in itens_validos for regiao in regioes):
        raise HTTPException(status_code=400, detail="Informe ao menos um procedimento com dentes/regioes ativos.")

    valor_bruto = sum(sum(float(regiao.valor or 0) for regiao in regioes if regiao.ativo) for _, regioes in itens_validos)
    desconto_percentual = max(0.0, float(payload.desconto_percentual or 0))
    desconto_valor = max(0.0, float(payload.desconto_valor or 0))
    valor_total = max(0.0, valor_bruto - (valor_bruto * desconto_percentual / 100.0) - desconto_valor)
    data_base = payload.data.strip() or agora_str()
    parcelas_plano = [
        parcela
        for parcela in payload.plano_pagamento
        if str(parcela.forma or "").strip() and float(parcela.valor or 0) >= 0
    ]
    forma_pagamento = payload.forma_pagamento.strip() or "A Definir"
    quantidade_parcelas = max(int(payload.parcelas or 1), 1)
    entrada_valor = sum(float(parcela.valor or 0) for parcela in parcelas_plano if normalizar_texto(parcela.descricao) == "entrada")
    primeiro_vencimento = next(
        (parcela.data.strip() for parcela in parcelas_plano if normalizar_texto(parcela.descricao) != "entrada" and parcela.data.strip()),
        data_base,
    )
    data_pagamento_entrada = next(
        (parcela.data.strip() for parcela in parcelas_plano if normalizar_texto(parcela.descricao) == "entrada" and parcela.data.strip()),
        None,
    )
    plano_pagamento_json = json.dumps([parcela.model_dump() for parcela in parcelas_plano], ensure_ascii=False)

    if contrato_id is None:
        cursor = conn.execute(
            """
            INSERT INTO contratos (
                paciente_id, valor_total, entrada, parcelas, primeiro_vencimento, data_pagamento_entrada,
                forma_pagamento, hash_importacao, data_criacao, status, observacoes, clinica_snapshot, criado_por_snapshot, tabela_snapshot, plano_pagamento_json, desconto_percentual, desconto_valor, validade_orcamento
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, NULL, ?, 'EM_ABERTO', ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                paciente_id,
                valor_total,
                entrada_valor,
                quantidade_parcelas,
                primeiro_vencimento,
                data_pagamento_entrada,
                forma_pagamento,
                data_base,
                payload.observacoes.strip(),
                payload.clinica.strip(),
                payload.criado_por.strip(),
                payload.tabela.strip(),
                plano_pagamento_json,
                desconto_percentual,
                desconto_valor,
                payload.validade_orcamento.strip(),
            ),
        )
        contrato_id = int(cursor.lastrowid)
    else:
        contrato = conn.execute("SELECT * FROM contratos WHERE id=? AND paciente_id=? LIMIT 1", (contrato_id, paciente_id)).fetchone()
        if contrato is None:
            raise HTTPException(status_code=404, detail="Orcamento nao encontrado.")
        if normalizar_texto(contrato["status"]) == "aprovado":
            raise HTTPException(status_code=400, detail="Orcamento aprovado nao pode ser editado.")
        conn.execute(
            """
            UPDATE contratos
            SET valor_total=?, entrada=?, parcelas=?, primeiro_vencimento=?, data_pagamento_entrada=?, forma_pagamento=?, data_criacao=?, observacoes=?, clinica_snapshot=?, criado_por_snapshot=?, tabela_snapshot=?, plano_pagamento_json=?, desconto_percentual=?, desconto_valor=?, validade_orcamento=?
            WHERE id=? AND paciente_id=?
            """,
            (
                valor_total,
                entrada_valor,
                quantidade_parcelas,
                primeiro_vencimento,
                data_pagamento_entrada,
                forma_pagamento,
                data_base,
                payload.observacoes.strip(),
                payload.clinica.strip(),
                payload.criado_por.strip(),
                payload.tabela.strip(),
                plano_pagamento_json,
                desconto_percentual,
                desconto_valor,
                payload.validade_orcamento.strip(),
                contrato_id,
                paciente_id,
            ),
        )
        conn.execute("DELETE FROM procedimentos_contrato WHERE contrato_id=?", (contrato_id,))
        conn.execute("DELETE FROM procedimentos_dente WHERE contrato_id=?", (contrato_id,))

    for indice_item, (item, regioes) in enumerate(itens_validos, start=1):
        subtotal = sum(float(regiao.valor or 0) for regiao in regioes if regiao.ativo)
        conn.execute(
            """
            INSERT INTO procedimentos_contrato (contrato_id, procedimento, valor, profissional_snapshot, denticao_snapshot)
            VALUES (?, ?, ?, ?, ?)
            """,
            (contrato_id, item.procedimento.strip(), subtotal, item.profissional.strip(), item.denticao.strip()),
        )
        for regiao in regioes:
            dente = regiao.dente if regiao.dente is not None else None
            conn.execute(
                """
                INSERT INTO procedimentos_dente (
                    paciente_id, contrato_id, grupo_item, dente, regiao, procedimento, status, faces, valor, data
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    paciente_id,
                    contrato_id,
                    indice_item,
                    dente,
                    regiao.regiao.strip(),
                    item.procedimento.strip(),
                    "ORCAMENTO" if regiao.ativo else "EXCLUIDO",
                    ",".join(face.strip() for face in regiao.faces if face.strip()),
                    float(regiao.valor or 0),
                    data_base,
                ),
            )

    return contrato_id


def carregar_recebiveis_paciente(conn: sqlite3.Connection, paciente_row: sqlite3.Row) -> list[sqlite3.Row]:
    return conn.execute(
        """
        SELECT *
        FROM recebiveis
        WHERE paciente_id = ?
           OR lower(trim(paciente_nome)) = lower(trim(?))
        ORDER BY vencimento, id
        """,
        (int(paciente_row["id"]), str(paciente_row["nome"] or "")),
    ).fetchall()


def atualizar_status_recebiveis_automaticamente(conn: sqlite3.Connection) -> None:
    hoje = date.today()
    rows = conn.execute(
        "SELECT id, vencimento, status FROM recebiveis"
    ).fetchall()
    houve_atualizacao = False
    for row in rows:
        status_atual = str(row["status"] or "Aberto").strip() or "Aberto"
        if normalizar_texto(status_atual) in {"pago", "cancelado", "suspenso"}:
            continue
        vencimento = parse_data_contrato(row["vencimento"])
        if not vencimento:
            continue
        novo_status = "Atrasado" if vencimento < hoje else "Aberto"
        if status_atual != novo_status:
            conn.execute("UPDATE recebiveis SET status=? WHERE id=?", (novo_status, int(row["id"])))
            houve_atualizacao = True
    if houve_atualizacao:
        conn.commit()


def atualizar_status_contas_pagar_automaticamente(conn: sqlite3.Connection) -> None:
    hoje = date.today()
    rows = conn.execute(
        "SELECT id, data_vencimento, status FROM contas_pagar"
    ).fetchall()
    houve_atualizacao = False
    for row in rows:
        status_atual = str(row["status"] or "A vencer").strip() or "A vencer"
        if normalizar_texto(status_atual) in {"pago", "cancelado", "suspenso"}:
            continue
        vencimento = parse_data_contrato(row["data_vencimento"])
        if not vencimento:
            continue
        novo_status = "Atrasado" if vencimento < hoje else "A vencer"
        if status_atual != novo_status:
            conn.execute("UPDATE contas_pagar SET status=? WHERE id=?", (novo_status, int(row["id"])))
            houve_atualizacao = True
    if houve_atualizacao:
        conn.commit()


def adicionar_meses(data_base: date, quantidade: int) -> date:
    mes = data_base.month - 1 + quantidade
    ano = data_base.year + mes // 12
    mes = mes % 12 + 1
    ultimo_dia = 31
    while True:
      try:
          date(ano, mes, ultimo_dia)
          break
      except ValueError:
          ultimo_dia -= 1
    dia = min(data_base.day, ultimo_dia)
    return date(ano, mes, dia)


def atualizar_recebiveis_lote_contrato(
    conn: sqlite3.Connection,
    contrato_id: int,
    paciente_nome: str,
    prontuario: str,
    forma_pagamento: str,
    status: str,
    observacao: str,
    primeiro_vencimento: str = "",
) -> None:
    rows = conn.execute(
        """
        SELECT id, parcela_numero, vencimento
        FROM recebiveis
        WHERE contrato_id=?
        ORDER BY parcela_numero, id
        """,
        (contrato_id,),
    ).fetchall()
    if not rows:
        return

    data_base = parse_data_contrato(primeiro_vencimento) if primeiro_vencimento.strip() else None
    for row in rows:
        vencimento = str(row["vencimento"] or "")
        if data_base is not None:
            vencimento = formatar_data_br(adicionar_meses(data_base, max(int(row["parcela_numero"] or 1) - 1, 0)))
        conn.execute(
            """
            UPDATE recebiveis
            SET paciente_nome=?, prontuario=?, vencimento=?, forma_pagamento=?, status=?, observacao=?
            WHERE id=?
            """,
            (
                paciente_nome.strip(),
                prontuario.strip(),
                vencimento,
                forma_pagamento.strip(),
                status.strip(),
                observacao.strip(),
                int(row["id"]),
            ),
        )


def carregar_agendamentos_paciente(conn: sqlite3.Connection, paciente_row: sqlite3.Row) -> list[sqlite3.Row]:
    return conn.execute(
        f"""
        SELECT *
        FROM agendamentos
        WHERE paciente_id = ?
           OR lower(trim(COALESCE(nome_paciente_snapshot, paciente_nome, ''))) = lower(trim(?))
        ORDER BY COALESCE({DATA_COLUNA_AGENDA}, data), hora_inicio, id
        """,
        (int(paciente_row["id"]), str(paciente_row["nome"] or "")),
    ).fetchall()


def proximo_agendamento_paciente(agendamentos: list[sqlite3.Row]) -> sqlite3.Row | None:
    agora = datetime.now()
    candidatos: list[tuple[datetime, sqlite3.Row]] = []
    for row in agendamentos:
        data_ref = parse_data_contrato(row[DATA_COLUNA_AGENDA] if DATA_COLUNA_AGENDA in row.keys() else row["data"])
        hora_ref = str(row["hora_inicio"] or "").strip()
        if not data_ref or not re.fullmatch(r"\d{2}:\d{2}", hora_ref):
            continue
        hora_obj = datetime.strptime(hora_ref, "%H:%M").time()
        data_hora = datetime.combine(data_ref, hora_obj)
        if data_hora >= agora and normalizar_texto(row["status"]) not in {"cancelado", "faltou"}:
            candidatos.append((data_hora, row))
    if not candidatos:
        return None
    candidatos.sort(key=lambda item: item[0])
    return candidatos[0][1]


def resumo_financeiro_paciente(recebiveis: list[sqlite3.Row]) -> FinanceiroResumo:
    total = sum(float(row["valor"] or 0) for row in recebiveis)
    pagos = sum(float(row["valor"] or 0) for row in recebiveis if str(row["status"] or "") == "Pago")
    atrasado = sum(float(row["valor"] or 0) for row in recebiveis if str(row["status"] or "") == "Atrasado")
    em_aberto = sum(
        float(row["valor"] or 0)
        for row in recebiveis
        if str(row["status"] or "") in {"Aberto", "A vencer", "Atrasado"}
    )
    qtd_atrasados = sum(1 for row in recebiveis if str(row["status"] or "") == "Atrasado")
    indicador = "pendente" if qtd_atrasados > 0 or em_aberto > 0 else "ok"
    return FinanceiroResumo(
        total=formatar_moeda_br(total),
        emAberto=formatar_moeda_br(em_aberto),
        atrasado=formatar_moeda_br(atrasado),
        pagos=formatar_moeda_br(pagos),
        quantidadeAtrasados=qtd_atrasados,
        indicador=indicador,
    )


def mapear_movimento_caixa(row: sqlite3.Row) -> MovimentoCaixaResumo:
    return MovimentoCaixaResumo(
        id=int(row["id"]),
        data=formatar_data_br_valor(row["data"]),
        origem=str(row["origem"] or ""),
        descricao=str(row["descricao"] or ""),
        valor=formatar_moeda_br(row["valor"]),
        tipo=str(row["tipo"] or ""),
        prontuario=formatar_prontuario_valor(row["prontuario"]),
        formaPagamento=str(row["forma_pagamento"] or ""),
        contaCaixa=str(row["conta_caixa"] or ""),
        observacao=str(row["observacao"] or ""),
        contratoId=int(row["contrato_id"]) if row["contrato_id"] is not None else None,
        recebivelId=int(row["recebivel_id"]) if row["recebivel_id"] is not None else None,
    )


def mapear_conta_pagar(row: sqlite3.Row) -> ContaPagarResumo:
    return ContaPagarResumo(
        id=int(row["id"]),
        vencimento=formatar_data_br_valor(row["data_vencimento"]),
        descricao=str(row["descricao"] or ""),
        fornecedor=str(row["fornecedor"] or ""),
        categoria=str(row["categoria"] or ""),
        valor=formatar_moeda_br(row["valor"]),
        valorPago=formatar_moeda_br(row["valor_pago"]),
        pagoEm=formatar_data_br_valor(row["pago"]),
        status=str(row["status"] or ""),
        observacao=str(row["observacao"] or ""),
    )


def mapear_recibo_manual(row: sqlite3.Row) -> ReciboManualResumo:
    return ReciboManualResumo(
        id=int(row["id"]),
        valor=formatar_moeda_br(row["valor"]),
        pagador=str(row["pagador"] or ""),
        recebedor=str(row["recebedor"] or ""),
        dataPagamento=formatar_data_br_valor(row["data_pagamento"]),
        referente=str(row["referente"] or ""),
        observacao=str(row["observacao"] or ""),
        cidade=str(row["cidade"] or ""),
        criadoEm=str(row["criado_em"] or ""),
    )


def mapear_saldo_conta(row: sqlite3.Row) -> SaldoContaResumo:
    return SaldoContaResumo(
        id=int(row["id"]),
        data=formatar_data_br_valor(row["data"]),
        conta=str(row["conta"] or ""),
        saldo=formatar_moeda_br(row["saldo"]),
        observacao=str(row["observacao"] or ""),
    )


def nome_mes_portugues(mes: int) -> str:
    meses = ["Janeiro", "Fevereiro", "Marco", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    if 1 <= int(mes) <= 12:
        return meses[int(mes) - 1]
    return ""


def carregar_metas_mensais(conn: sqlite3.Connection, ano: int) -> list[MetaMensalResumo]:
    ano_referencia = int(ano or date.today().year)
    base_row = conn.execute(
        """
        SELECT meta, supermeta, hipermeta
        FROM metas_vendas
        WHERE ano=?
        LIMIT 1
        """,
        (ano_referencia,),
    ).fetchone()
    meta_padrao = float(base_row["meta"] or 100000.0) if base_row else 100000.0
    supermeta_padrao = float(base_row["supermeta"] or 150000.0) if base_row else 150000.0
    hipermeta_padrao = float(base_row["hipermeta"] or 200000.0) if base_row else 200000.0
    rows = conn.execute(
        """
        SELECT *
        FROM metas_mensais
        WHERE ano=?
        ORDER BY mes ASC
        """,
        (ano_referencia,),
    ).fetchall()
    por_mes = {int(row["mes"]): row for row in rows if row["mes"] is not None}
    metas: list[MetaMensalResumo] = []
    alterou = False
    for mes in range(1, 13):
        row = por_mes.get(mes)
        if row is None:
            conn.execute(
                """
                INSERT INTO metas_mensais
                (ano, mes, meta, supermeta, hipermeta, data_atualizacao)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (ano_referencia, mes, meta_padrao, supermeta_padrao, hipermeta_padrao, agora_str()),
            )
            alterou = True
            meta = meta_padrao
            supermeta = supermeta_padrao
            hipermeta = hipermeta_padrao
            atualizado = agora_str()
        else:
            meta = float(row["meta"] or meta_padrao)
            supermeta = float(row["supermeta"] or supermeta_padrao)
            hipermeta = float(row["hipermeta"] or hipermeta_padrao)
            atualizado = str(row["data_atualizacao"] or "")
        metas.append(
            MetaMensalResumo(
                ano=ano_referencia,
                mes=mes,
                mesNome=nome_mes_portugues(mes),
                meta=meta,
                supermeta=supermeta,
                hipermeta=hipermeta,
                dataAtualizacao=atualizado,
            )
        )
    if alterou:
        conn.commit()
    return metas


def obter_meta_mensal(conn: sqlite3.Connection, ano: int, mes: int) -> MetaMensalResumo:
    metas = carregar_metas_mensais(conn, ano)
    for item in metas:
        if item.mes == int(mes):
            return item
    raise HTTPException(status_code=404, detail="Meta mensal nao encontrada.")


def carregar_notas_fiscais_emitidas(conn: sqlite3.Connection) -> list[sqlite3.Row]:
    return conn.execute(
        """
        SELECT *
        FROM notas_fiscais_emitidas
        ORDER BY COALESCE(competencia, '') DESC, COALESCE(data_emissao, '') DESC, id DESC
        """
    ).fetchall()


def mapear_nota_fiscal_emitida(row: sqlite3.Row) -> NotaFiscalEmitidaResumo:
    valor_nf = float(row["valor_nf"] or 0)
    valor_recebido = float(row["valor_recebido"] or 0)
    diferenca = valor_recebido - valor_nf
    conciliado = abs(diferenca) < 0.01
    return NotaFiscalEmitidaResumo(
        id=int(row["id"]),
        competencia=str(row["competencia"] or ""),
        dataEmissao=formatar_data_br_valor(row["data_emissao"]),
        dataRecebimento=formatar_data_br_valor(row["data_recebimento"]),
        numeroNf=str(row["numero_nf"] or ""),
        serie=str(row["serie"] or ""),
        cliente=str(row["cliente"] or ""),
        descricao=str(row["descricao"] or ""),
        contaDestino=str(row["conta_destino"] or ""),
        valorNf=formatar_moeda_br(valor_nf),
        valorRecebido=formatar_moeda_br(valor_recebido),
        valorNfNumero=valor_nf,
        valorRecebidoNumero=valor_recebido,
        diferenca=formatar_moeda_br(diferenca),
        diferencaNumero=diferenca,
        status=str(row["status"] or ""),
        observacao=str(row["observacao"] or ""),
        conciliado=conciliado,
        criadoEm=str(row["criado_em"] or ""),
        atualizadoEm=str(row["atualizado_em"] or ""),
    )


def resumo_financeiro_global(recebiveis: list[sqlite3.Row]) -> FinanceiroResumo:
    return resumo_financeiro_paciente(recebiveis)


def carregar_recebiveis_financeiro(conn: sqlite3.Connection) -> list[sqlite3.Row]:
    atualizar_status_recebiveis_automaticamente(conn)
    return conn.execute(
        """
        SELECT *
        FROM recebiveis
        ORDER BY COALESCE(vencimento, ''), paciente_nome, parcela_numero, id
        """
    ).fetchall()


def carregar_caixa_financeiro(conn: sqlite3.Connection) -> list[sqlite3.Row]:
    return conn.execute(
        """
        SELECT f.*
        FROM financeiro f
        LEFT JOIN recebiveis r ON r.id = f.recebivel_id
        WHERE f.recebivel_id IS NULL OR COALESCE(r.status, '') = 'Pago'
        ORDER BY COALESCE(f.data, '') DESC, f.id DESC
        """
    ).fetchall()


def carregar_contas_pagar_financeiro(conn: sqlite3.Connection) -> list[sqlite3.Row]:
    atualizar_status_contas_pagar_automaticamente(conn)
    return conn.execute(
        """
        SELECT *
        FROM contas_pagar
        ORDER BY COALESCE(data_vencimento, ''), fornecedor, descricao, id
        """
    ).fetchall()


def carregar_saldos_conta_financeiro(conn: sqlite3.Connection) -> list[sqlite3.Row]:
    return conn.execute(
        """
        SELECT *
        FROM saldos_conta
        ORDER BY COALESCE(data, '') DESC, conta, id DESC
        """
    ).fetchall()


def dados_dashboard(conn: sqlite3.Connection) -> DashboardPainelResposta:
    hoje = date.today()
    hoje_iso = hoje.isoformat()
    ano_atual = hoje.year
    mes_atual = hoje.month
    meses = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun", "Jul", "Ago", "Set", "Out", "Nov", "Dez"]

    atualizar_status_recebiveis_automaticamente(conn)
    atualizar_status_contas_pagar_automaticamente(conn)

    recebiveis_rows = carregar_recebiveis_financeiro(conn)
    caixa_rows = carregar_caixa_financeiro(conn)

    total_em_aberto = sum(
        float(row["valor"] or 0)
        for row in recebiveis_rows
        if normalizar_texto(row["status"]) in {"aberto", "atrasado"}
    )
    entradas_hoje = sum(
        float(row["valor"] or 0)
        for row in caixa_rows
        if parse_data_contrato(row["data"]) == hoje and normalizar_texto(row["tipo"]) == "entrada"
    )
    saidas_hoje = sum(
        float(row["valor"] or 0)
        for row in caixa_rows
        if parse_data_contrato(row["data"]) == hoje and normalizar_texto(row["tipo"]) == "saida"
    )
    saldo_hoje = entradas_hoje - saidas_hoje

    meta_atual = obter_meta_mensal(conn, ano_atual, mes_atual)
    meta_mes = float(meta_atual.meta or 0)
    supermeta_mes = float(meta_atual.supermeta or 0)
    hipermeta_mes = float(meta_atual.hipermeta or 0)

    vendas_rows = conn.execute(
        """
        SELECT data_venda, valor_total
        FROM vendas
        WHERE COALESCE(data_venda, '') <> ''
        """
    ).fetchall()

    serie_vendas: list[float] = []
    for mes_indice in range(1, 13):
        total_mes = 0.0
        for row in vendas_rows:
            data_ref = parse_data_contrato(str(row["data_venda"] or ""))
            if not data_ref or data_ref.year != ano_atual or data_ref.month != mes_indice:
                continue
            total_mes += float(row["valor_total"] or 0)
        serie_vendas.append(total_mes)

    vendido_mes = serie_vendas[mes_atual - 1] if 0 < mes_atual <= len(serie_vendas) else 0.0
    vendido_ano = sum(serie_vendas)
    falta_meta_mes = max(meta_mes - vendido_mes, 0.0)
    falta_meta_ano = max((meta_mes * 12) - vendido_ano, 0.0) if meta_mes > 0 else 0.0
    percentual_meta_mes = min((vendido_mes / meta_mes) * 100, 999) if meta_mes > 0 else 0.0
    percentual_meta_ano = min((vendido_ano / (meta_mes * 12)) * 100, 999) if meta_mes > 0 else 0.0

    indicadores = [
        DashboardIndicadorResposta(
            chave="vendido_mes",
            titulo="Vendido no mês",
            valor=formatar_moeda_br(vendido_mes),
            detalhe=f"{meses[mes_atual - 1]} de {ano_atual}",
        ),
        DashboardIndicadorResposta(
            chave="falta_meta_mes",
            titulo="Falta para meta do mês",
            valor=formatar_moeda_br(falta_meta_mes),
            detalhe=f"Meta mensal: {formatar_moeda_br(meta_mes)}",
        ),
        DashboardIndicadorResposta(
            chave="vendido_ano",
            titulo="Vendido no ano",
            valor=formatar_moeda_br(vendido_ano),
            detalhe=f"Acumulado de {ano_atual}",
        ),
        DashboardIndicadorResposta(
            chave="falta_meta_ano",
            titulo="Falta para meta do ano",
            valor=formatar_moeda_br(falta_meta_ano),
            detalhe=f"Meta anual: {formatar_moeda_br(meta_mes * 12)}",
        ),
    ]

    return DashboardPainelResposta(
        indicadores=indicadores,
        meses=meses,
        serieVendas=serie_vendas,
        resumoHoje=DashboardResumoHojeResposta(
            entradasConfirmadas=formatar_moeda_br(entradas_hoje),
            saidasPrevistas=formatar_moeda_br(saidas_hoje),
            saldoProjetado=formatar_moeda_br(saldo_hoje),
        ),
        metas=DashboardMetasResposta(
            vendidoMes=formatar_moeda_br(vendido_mes),
            vendidoAno=formatar_moeda_br(vendido_ano),
            metaMes=formatar_moeda_br(meta_mes),
            supermetaMes=formatar_moeda_br(supermeta_mes),
            hipermetaMes=formatar_moeda_br(hipermeta_mes),
            faltaMetaMes=formatar_moeda_br(falta_meta_mes),
            faltaMetaAno=formatar_moeda_br(falta_meta_ano),
            percentualMetaMes=percentual_meta_mes,
            percentualMetaAno=percentual_meta_ano,
        ),
        agendaHoje=[],
        alertas=[],
        atividades=[],
    )


def registrar_saldo_conta(conn: sqlite3.Connection, *, data_saldo: str, conta: str, saldo: float, observacao: str = "") -> int:
    cursor = conn.execute(
        """
        INSERT INTO saldos_conta (data, conta, saldo, observacao)
        VALUES (?, ?, ?, ?)
        """,
        (data_saldo.strip(), conta.strip().upper(), float(saldo or 0), observacao.strip()),
    )
    return int(cursor.lastrowid)


def saldos_informados_por_conta_ate(conn: sqlite3.Connection, data_limite: date) -> dict[str, float]:
    saldos = {conta: 0.0 for conta in CONTAS_CAIXA_MODELO}
    rows = conn.execute("SELECT * FROM saldos_conta ORDER BY data ASC, id ASC").fetchall()
    for row in rows:
        data_ref = parse_data_contrato(row["data"])
        conta = str(row["conta"] or "").strip().upper()
        if not data_ref or conta not in saldos:
            continue
        if data_ref <= data_limite:
            saldos[conta] = float(row["saldo"] or 0)
    return saldos


def caixa_diario_para_excel_bytes(conn: sqlite3.Connection, financeiro_rows: list[sqlite3.Row]) -> bytes | None:
    if Workbook is None or Alignment is None or Border is None or Font is None or PatternFill is None or Side is None:
        return None
    if not financeiro_rows:
        return None

    movimentos = []
    for row in financeiro_rows:
        data_ref = parse_data_contrato(row["data"])
        if not data_ref:
            continue
        movimentos.append(
            {
                "id": int(row["id"]),
                "data": data_ref,
                "descricao": str(row["descricao"] or row["origem"] or "").strip(),
                "valor": float(row["valor"] or 0),
                "tipo": str(row["tipo"] or ""),
                "conta": str(row["conta_caixa"] or "CAIXA").strip().upper() or "CAIXA",
            }
        )
    if not movimentos:
        return None

    movimentos.sort(key=lambda item: (item["data"], item["id"]))
    wb = Workbook()
    ws = wb.active
    ws.title = "Caixa"

    fonte_titulo = Font(name="Times New Roman", size=20, color="000000")
    fonte_cabecalho = Font(name="Times New Roman", size=12, color="000000")
    fonte_normal = Font(name="Times New Roman", size=12, color="000000")
    fonte_vermelha = Font(name="Times New Roman", size=12, color="FF0000")
    preenchimento_saldo = PatternFill(fill_type="solid", fgColor="D0CECE")
    borda_fina = Border(
        left=Side(style="thin", color="000000"),
        right=Side(style="thin", color="000000"),
        top=Side(style="thin", color="000000"),
        bottom=Side(style="thin", color="000000"),
    )

    ws.merge_cells("A1:E1")
    ws["A1"] = "CAIXA"
    ws["A1"].font = fonte_titulo
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    for coluna, largura in {"A": 14, "B": 58, "C": 18, "D": 16, "E": 18}.items():
        ws.column_dimensions[coluna].width = largura

    linha_atual = 3
    datas = sorted({item["data"] for item in movimentos})
    for data_atual in datas:
        data_texto = formatar_data_br(data_atual)
        grupo = [item for item in movimentos if item["data"] == data_atual]
        saldo_atual = saldos_informados_por_conta_ate(conn, data_atual)

        cabecalhos = ["Data", "Descrição", "Entradas", "Saídas", "Banco"]
        for indice, titulo in enumerate(cabecalhos, start=1):
            celula = ws.cell(row=linha_atual, column=indice, value=titulo)
            celula.font = fonte_cabecalho
            celula.border = borda_fina
        linha_atual += 1

        for conta in CONTAS_CAIXA_MODELO:
            valor_saldo = round(saldo_atual.get(conta, 0), 2)
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=f"SALDO DO DIA ANTERIOR - {conta}")
            ws.cell(row=linha_atual, column=3, value=valor_saldo)
            ws.cell(row=linha_atual, column=5, value=conta)
            for coluna in range(1, 6):
                celula = ws.cell(row=linha_atual, column=coluna)
                celula.fill = preenchimento_saldo
                celula.border = borda_fina
                celula.font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=3).font = fonte_vermelha if valor_saldo < 0 else fonte_normal
            linha_atual += 1

        linha_atual += 3

        for item in grupo:
            valor = round(item["valor"], 2)
            conta = item["conta"] if item["conta"] in CONTAS_CAIXA_MODELO else "CAIXA"
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=item["descricao"])
            if item["tipo"] == "Entrada":
                ws.cell(row=linha_atual, column=3, value=valor)
                saldo_atual[conta] += valor
            else:
                ws.cell(row=linha_atual, column=4, value=valor)
                saldo_atual[conta] -= valor
            ws.cell(row=linha_atual, column=5, value=conta)
            for coluna in range(1, 6):
                celula = ws.cell(row=linha_atual, column=coluna)
                celula.border = borda_fina
                celula.font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=4).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=4).font = fonte_vermelha if ws.cell(row=linha_atual, column=4).value not in (None, 0, 0.0) else fonte_normal
            linha_atual += 1

        linha_atual += 3

        for conta in CONTAS_CAIXA_MODELO:
            valor_saldo = round(saldo_atual.get(conta, 0), 2)
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=f"SALDO DO DIA - {conta}")
            ws.cell(row=linha_atual, column=3, value=valor_saldo)
            ws.cell(row=linha_atual, column=5, value=conta)
            for coluna in range(1, 6):
                celula = ws.cell(row=linha_atual, column=coluna)
                celula.fill = preenchimento_saldo
                celula.border = borda_fina
                celula.font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=3).font = fonte_vermelha if valor_saldo < 0 else fonte_normal
            linha_atual += 1

        linha_atual += 3

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def openpyxl_column_name(index: int) -> str:
    resultado = ""
    atual = index
    while atual > 0:
        atual, resto = divmod(atual - 1, 26)
        resultado = chr(65 + resto) + resultado
    return resultado


def exportar_todas_tabelas_excel_bytes(conn: sqlite3.Connection) -> bytes:
    if Workbook is None or Font is None or PatternFill is None:
        raise HTTPException(status_code=500, detail="openpyxl nao disponivel")

    wb = Workbook()
    ws_inicial = wb.active
    wb.remove(ws_inicial)

    tabelas = conn.execute(
        """
        SELECT name
        FROM sqlite_master
        WHERE type='table'
          AND name NOT LIKE 'sqlite_%'
        ORDER BY name
        """
    ).fetchall()

    for item in tabelas:
        tabela = str(item["name"] or "").strip()
        if not tabela:
            continue
        sheet = wb.create_sheet(title=tabela[:31])
        colunas_info = conn.execute(f"PRAGMA table_info({tabela})").fetchall()
        colunas = [str(col["name"]) for col in colunas_info]
        if not colunas:
            sheet["A1"] = "Sem dados"
            continue

        for idx, coluna in enumerate(colunas, start=1):
            celula = sheet.cell(row=1, column=idx, value=coluna)
            celula.font = Font(bold=True)
            celula.fill = PatternFill(fill_type="solid", fgColor="E8E1D2")

        rows = conn.execute(f"SELECT * FROM {tabela}").fetchall()
        for row_idx, row in enumerate(rows, start=2):
            for col_idx, coluna in enumerate(colunas, start=1):
                valor = row[coluna]
                if valor is None:
                    valor = ""
                elif isinstance(valor, bytes):
                    try:
                        valor = valor.decode("utf-8", errors="ignore")
                    except Exception:
                        valor = str(valor)
                elif isinstance(valor, (dict, list)):
                    valor = json.dumps(valor, ensure_ascii=False)
                sheet.cell(row=row_idx, column=col_idx, value=valor)

        for idx, coluna in enumerate(colunas, start=1):
            largura = max(14, min(len(coluna) + 2, 40))
            if rows:
                exemplo = str(rows[0][coluna] or "")
                largura = max(largura, min(len(exemplo) + 2, 40))
            sheet.column_dimensions[openpyxl_column_name(idx)].width = largura
        sheet.freeze_panes = "A2"

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def registrar_movimento_caixa(
    conn: sqlite3.Connection,
    *,
    origem: str,
    descricao: str,
    valor: float,
    tipo: str,
    data_movimento: str,
    prontuario: str = "",
    forma_pagamento: str = "",
    conta_caixa: str = "",
    observacao: str = "",
    contrato_id: int | None = None,
    recebivel_id: int | None = None,
) -> int:
    cursor = conn.execute(
        """
        INSERT INTO financeiro (
            origem, descricao, valor, data, tipo, contrato_id, recebivel_id, prontuario,
            forma_pagamento, conta_caixa, observacao
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            origem.strip(),
            descricao.strip(),
            float(valor or 0),
            data_movimento.strip(),
            tipo.strip(),
            contrato_id,
            recebivel_id,
            prontuario.strip(),
            forma_pagamento.strip(),
            conta_caixa.strip(),
            observacao.strip(),
        ),
    )
    return int(cursor.lastrowid)


def baixar_recebivel_no_caixa(
    conn: sqlite3.Connection,
    *,
    paciente_id: int,
    recebivel_id: int,
    data_pagamento: str,
    forma_pagamento: str = "",
    conta_caixa: str = "",
    desconto_valor: float = 0,
    observacao: str = "",
) -> sqlite3.Row:
    recebivel = carregar_recebivel_paciente(conn, paciente_id, recebivel_id)
    if normalizar_texto(recebivel["status"]) == "pago":
        raise HTTPException(status_code=400, detail="Este recebivel ja esta pago.")

    data_ref = parse_data_contrato(data_pagamento) or date.today()
    data_iso = data_ref.isoformat()
    data_br = formatar_data_br(data_ref)
    forma = (forma_pagamento or str(recebivel["forma_pagamento"] or "")).strip()
    conta = (conta_caixa or forma).strip()
    valor_original = float(recebivel["valor"] or 0)
    desconto = max(0.0, float(desconto_valor or 0))
    valor_recebido = max(0.0, valor_original - desconto)
    paciente_nome = str(recebivel["paciente_nome"] or "").strip()
    parcela_atual = max(int(recebivel["parcela_numero"] or 0), 0)
    total_parcelas = parcela_atual
    if recebivel["contrato_id"] is not None:
        contrato_row = conn.execute("SELECT parcelas FROM contratos WHERE id=?", (int(recebivel["contrato_id"]),)).fetchone()
        if contrato_row is not None:
            total_parcelas = max(int(contrato_row["parcelas"] or parcela_atual or 1), 1)
    if parcela_atual == 0:
        descricao = f"{paciente_nome or 'Recebimento'} - ENTRADA"
    else:
        descricao = f"{paciente_nome or 'Recebimento'} - {parcela_atual}/{total_parcelas}"
    if desconto > 0:
        descricao = f"{descricao} (desc. {formatar_moeda_br(desconto)})"

    observacao_atual = str(recebivel["observacao"] or "").strip()
    observacao_desconto = f"Desconto na baixa: {formatar_moeda_br(desconto)}" if desconto > 0 else ""
    observacao_final = " | ".join(parte for parte in [observacao_atual, observacao_desconto, observacao.strip()] if parte)

    registrar_movimento_caixa(
        conn,
        origem=paciente_nome or "Recebimento",
        descricao=descricao,
        valor=valor_recebido,
        tipo="Entrada",
        data_movimento=data_iso,
        prontuario=formatar_prontuario_valor(recebivel["prontuario"]),
        forma_pagamento=forma,
        conta_caixa=conta,
        observacao=observacao_final,
        contrato_id=int(recebivel["contrato_id"]) if recebivel["contrato_id"] is not None else None,
        recebivel_id=int(recebivel["id"]),
    )

    conn.execute(
        """
        UPDATE recebiveis
        SET status='Pago', data_pagamento=?, forma_pagamento=?, observacao=?
        WHERE id=? AND paciente_id=?
        """,
        (data_br, forma, observacao_final, recebivel_id, paciente_id),
    )
    return carregar_recebivel_paciente(conn, paciente_id, recebivel_id)


def listar_exames_paciente(paciente_row: sqlite3.Row) -> list[ArquivoPacienteItem]:
    pasta = pasta_exames_paciente(paciente_row)
    itens: list[ArquivoPacienteItem] = []
    for nome_arquivo in sorted(os.listdir(pasta), reverse=True):
        caminho = os.path.join(pasta, nome_arquivo)
        if os.path.isfile(caminho):
            itens.append(
                ArquivoPacienteItem(
                    nome=nome_arquivo,
                    caminho=caminho,
                    modificadoEm=formatar_data_br(datetime.fromtimestamp(os.path.getmtime(caminho)).date()),
                    extensao=os.path.splitext(nome_arquivo)[1].lower(),
                )
            )
    return itens


def listar_documentos_paciente(paciente_row: sqlite3.Row) -> list[ArquivoPacienteItem]:
    if not os.path.isdir(DOCS_DIR):
        return []
    prontuario = normalizar_texto(formatar_prontuario_valor(paciente_row["prontuario"]))
    nome = normalizar_texto(limpar_nome(paciente_row["nome"]))
    itens: list[ArquivoPacienteItem] = []
    for nome_arquivo in sorted(os.listdir(DOCS_DIR), reverse=True):
        caminho = os.path.join(DOCS_DIR, nome_arquivo)
        if not os.path.isfile(caminho):
            continue
        nome_norm = normalizar_texto(nome_arquivo)
        if (prontuario and prontuario in nome_norm) or (nome and nome in nome_norm):
            itens.append(
                ArquivoPacienteItem(
                    nome=nome_arquivo,
                    caminho=caminho,
                    modificadoEm=formatar_data_br(datetime.fromtimestamp(os.path.getmtime(caminho)).date()),
                    extensao=os.path.splitext(nome_arquivo)[1].lower(),
                )
            )
    return itens


def extrair_contrato_id_documento(nome_arquivo: str) -> int | None:
    nome = os.path.basename(str(nome_arquivo or "").strip())
    if not nome.upper().startswith("CONTRATO_"):
        return None
    sem_extensao, _ = os.path.splitext(nome)
    partes = sem_extensao.split("_")
    for indice, parte in enumerate(partes):
        if parte.isdigit() and indice + 1 < len(partes) and partes[indice + 1].isdigit():
            return int(parte)
    return None


def limpar_documentos_contrato_variantes(paciente_row: sqlite3.Row, contrato_id: int) -> None:
    if not os.path.isdir(DOCS_DIR):
        return
    prefixo = nome_base_contrato(paciente_row, contrato_id).upper()
    for nome_arquivo in os.listdir(DOCS_DIR):
        caminho = os.path.join(DOCS_DIR, nome_arquivo)
        if not os.path.isfile(caminho):
            continue
        nome_upper = str(nome_arquivo or "").upper()
        if nome_upper.startswith(prefixo) and nome_upper.endswith((".DOCX", ".HTML")):
            os.remove(caminho)


def resolver_documento_contrato_atual(
    conn: sqlite3.Connection,
    paciente_row: sqlite3.Row,
    nome_arquivo: str,
) -> str | None:
    contrato_id = extrair_contrato_id_documento(nome_arquivo)
    if contrato_id is None:
        return None
    contrato = conn.execute(
        "SELECT * FROM contratos WHERE id=? AND paciente_id=? LIMIT 1",
        (contrato_id, int(paciente_row["id"])),
    ).fetchone()
    if contrato is None:
        return None
    limpar_documentos_contrato_variantes(paciente_row, contrato_id)
    return gerar_documento_contrato(conn, paciente_row, contrato, contrato_id)


def buscar_documento_paciente(paciente_row: sqlite3.Row, nome_arquivo: str) -> str:
    nome_busca = os.path.basename(str(nome_arquivo or "").strip())
    if not nome_busca:
        raise HTTPException(status_code=404, detail="Documento nao encontrado.")
    for item in listar_documentos_paciente(paciente_row):
        if item.nome == nome_busca and os.path.isfile(item.caminho):
            return item.caminho
    raise HTTPException(status_code=404, detail="Documento nao encontrado.")


def buscar_exame_paciente(paciente_row: sqlite3.Row, nome_arquivo: str) -> str:
    nome_busca = os.path.basename(str(nome_arquivo or "").strip())
    if not nome_busca:
        raise HTTPException(status_code=404, detail="Arquivo nao encontrado.")
    for item in listar_exames_paciente(paciente_row):
        if item.nome == nome_busca and os.path.isfile(item.caminho):
            return item.caminho
    raise HTTPException(status_code=404, detail="Arquivo nao encontrado.")


def gerar_html_recibo_especie(recebivel: sqlite3.Row) -> str:
    paciente = formatar_titulo(str(recebivel["paciente_nome"] or ""))
    prontuario = formatar_prontuario_valor(recebivel["prontuario"])
    parcela = int(recebivel["parcela_numero"] or 0)
    descricao = "ENTRADA" if parcela == 0 else f"PARCELA {parcela}"
    valor = formatar_moeda_br(recebivel["valor"])
    data_pagamento = formatar_data_br_valor(recebivel["data_pagamento"]) or formatar_data_br(date.today())
    observacao = str(recebivel["observacao"] or "").strip()
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8" />
  <title>Recibo</title>
  <style>
    body {{ font-family: Arial, sans-serif; padding: 32px; color: #222; }}
    .sheet {{ max-width: 760px; margin: 0 auto; border: 1px solid #d7d7d7; padding: 36px; }}
    h1 {{ margin: 0 0 24px; font-size: 24px; }}
    .meta {{ display: grid; gap: 8px; margin-bottom: 24px; }}
    .valor {{ font-size: 28px; font-weight: 700; margin: 20px 0; }}
    .assinatura {{ margin-top: 56px; padding-top: 16px; border-top: 1px solid #444; width: 320px; }}
    .small {{ color: #666; font-size: 12px; }}
    @media print {{ body {{ padding: 0; }} .sheet {{ border: 0; }} }}
  </style>
</head>
<body>
  <div class="sheet">
    <h1>RECIBO</h1>
    <div class="meta">
      <div><strong>CLÍNICA:</strong> SOUL SUL CLÍNICA INTEGRADA</div>
      <div><strong>PACIENTE:</strong> {paciente}</div>
      <div><strong>PRONTUÁRIO:</strong> {prontuario}</div>
      <div><strong>REFERÊNCIA:</strong> {descricao}</div>
      <div><strong>DATA:</strong> {data_pagamento}</div>
      <div><strong>FORMA:</strong> DINHEIRO</div>
    </div>
    <div>Recebemos de <strong>{paciente}</strong> a quantia de <strong>{valor}</strong>, referente a {descricao.lower()}.</div>
    <div class="valor">{valor}</div>
    {f'<div><strong>OBSERVAÇÃO:</strong> {observacao}</div>' if observacao else ''}
    <div class="assinatura">SOUL SUL CLÍNICA INTEGRADA</div>
    <div class="small">Documento gerado pelo sistema.</div>
  </div>
</body>
</html>"""


LOGO_RECIBO_BASE64: str | None = None


def obter_logo_recibo_base64() -> str:
    global LOGO_RECIBO_BASE64
    if LOGO_RECIBO_BASE64:
        return LOGO_RECIBO_BASE64
    caminho = os.path.join(os.path.dirname(__file__), "assets", "sou sul marca preta fundo.png")
    try:
        with open(caminho, "rb") as arquivo:
            LOGO_RECIBO_BASE64 = base64.b64encode(arquivo.read()).decode("ascii")
    except FileNotFoundError:
        LOGO_RECIBO_BASE64 = ""
    return LOGO_RECIBO_BASE64


def gerar_html_recibo_manual(recibo: sqlite3.Row) -> str:
    valor = formatar_moeda_br(recibo["valor"])
    data_pagamento = formatar_data_br_valor(recibo["data_pagamento"]) or formatar_data_br(date.today())
    pagador = formatar_titulo(str(recibo["pagador"] or "")) or "________________________________"
    recebedor = formatar_titulo(str(recibo["recebedor"] or "")) or "________________________________"
    referente = str(recibo["referente"] or "").strip() or "pagamento realizado"
    observacao = str(recibo["observacao"] or "").strip()
    cidade = str(recibo["cidade"] or "").strip() or "Campos dos Goytacazes/RJ"
    numero = int(recibo["numero"] or 0)
    logo = obter_logo_recibo_base64()
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8" />
  <title>Recibo</title>
  <style>
    body {{ font-family: Arial, sans-serif; color: #222; padding: 28px; }}
    .sheet {{ max-width: 820px; margin: 0 auto; border: 1px solid #d7d7d7; padding: 40px; }}
    h1 {{ margin: 0 0 24px; font-size: 30px; }}
    .topline {{ display: flex; justify-content: space-between; gap: 12px; margin-bottom: 28px; font-size: 14px; }}
    .body-text {{ font-size: 18px; line-height: 1.8; }}
    .valor {{ font-size: 34px; font-weight: 700; margin: 24px 0; }}
    .obs {{ margin-top: 20px; }}
    .assinaturas {{ display: grid; grid-template-columns: 1fr 1fr; gap: 28px; margin-top: 72px; }}
    .assinatura {{ padding-top: 14px; border-top: 1px solid #444; text-align: center; min-height: 48px; }}
    .small {{ color: #666; font-size: 12px; margin-top: 10px; }}
    @media print {{ body {{ padding: 0; }} .sheet {{ border: 0; }} }}
  </style>
</head>
<body>
  <div class="sheet">
    <div class="topline">
      <img class="logo-recibo" src="data:image/png;base64,{logo}" alt="Logo" />
      <div>
        <strong>RECIBO</strong>
        <div class="numero">Nº {numero:04d}</div>
      </div>
      <span>{cidade}, {data_pagamento}</span>
    </div>
    <div class="body-text">
      Recebi de <strong>{pagador}</strong> a quantia de <strong>{valor}</strong>,
      paga em {data_pagamento}, destinada a <strong>{recebedor}</strong>, referente a <strong>{referente}</strong>.
    </div>
    <div class="valor">{valor}</div>
    {f'<div class="obs"><strong>Observação:</strong> {observacao}</div>' if observacao else ''}
    <div class="assinaturas">
      <div class="assinatura">{pagador}<div class="small">Pagador(a)</div></div>
      <div class="assinatura">{recebedor}<div class="small">Recebedor(a)</div></div>
    </div>
  </div>
</body>
</html>"""


def buscar_foto_paciente(paciente_row: sqlite3.Row) -> str:
    caminho = str(paciente_row["foto_path"] or "").strip()
    if not caminho or not os.path.isfile(caminho):
        raise HTTPException(status_code=404, detail="Foto nao encontrada.")
    return caminho


def valor_row(row: sqlite3.Row | dict, coluna: str) -> str:
    if isinstance(row, sqlite3.Row):
        valor = row[coluna] if coluna in row.keys() else ""
    else:
        valor = row.get(coluna, "")
    if valor is None:
        return ""
    return str(valor).strip()


def nome_base_contrato(paciente_row: sqlite3.Row, contrato_id: int) -> str:
    nome = limpar_nome(str(paciente_row["nome"] or "")).replace(" ", "_")
    prontuario = formatar_prontuario_valor(paciente_row["prontuario"])
    return f"CONTRATO_{nome}_{prontuario}_{contrato_id}"


def nome_arquivo_contrato(paciente_row: sqlite3.Row, contrato_id: int) -> str:
    timestamp = agora_local().strftime("%Y%m%d_%H%M%S")
    return f"{nome_base_contrato(paciente_row, contrato_id)}_{timestamp}.docx"


def nome_arquivo_ordem_servico(paciente_row: sqlite3.Row, ordem_id: int) -> str:
    nome = limpar_nome(str(paciente_row["nome"] or "")).replace(" ", "_")
    prontuario = formatar_prontuario_valor(paciente_row["prontuario"])
    timestamp = agora_local().strftime("%Y%m%d_%H%M%S")
    return f"ORDEM_SERVICO_{nome}_{prontuario}_{ordem_id}_{timestamp}.docx"


def texto_ordem_servico_carga_imediata(valor: bool) -> str:
    return "SIM" if valor else "NÃO"


def preencher_campo_rotulado(celula, prefixo: str, valor: str) -> None:
    texto = prefixo if not valor else f"{prefixo} {valor}"
    if celula.paragraphs:
        celula.paragraphs[0].text = texto
    else:
        celula.text = texto


def limpar_celula_docx(celula) -> None:
    celula.text = ""


def preencher_celula_multilinha(celula, linhas: list[str]) -> None:
    limpar_celula_docx(celula)
    linhas_validas = [str(linha or "") for linha in linhas]
    if not linhas_validas:
        return
    primeiro = True
    for linha in linhas_validas:
        if primeiro:
            celula.paragraphs[0].text = linha
            primeiro = False
        else:
            celula.add_paragraph(linha)


def remover_linhas_extras_tabela_docx(tabela, manter_total: int) -> None:
    while len(tabela.rows) > manter_total:
        tabela._tbl.remove(tabela.rows[-1]._tr)


def gerar_documento_ordem_servico(
    paciente_row: sqlite3.Row,
    ordem_id: int,
    procedimento_nome: str,
    material: str,
    material_outro: str,
    elemento_arcada: str,
    escala: str,
    cor: str,
    carga_imediata: bool,
    retorno_solicitado: str,
    etapas: list[tuple[str, str]],
    observacao: str,
) -> str:
    if Document is None or not os.path.isfile(TEMPLATE_ORDEM_SERVICO_PATH):
        raise HTTPException(status_code=500, detail="Modelo da ordem de servico nao encontrado.")

    doc = Document(TEMPLATE_ORDEM_SERVICO_PATH)
    tabelas = doc.tables
    if not tabelas:
        raise HTTPException(status_code=500, detail="Modelo da ordem de servico invalido.")

    data_emissao = agora_local().strftime("%d/%m/%Y %H:%M")
    material_final = material if normalizar_texto(material) != "outro" else f"OUTRO - {material_outro}"
    etapas_texto = " | ".join(
        f"{etapa}: {descricao}" if normalizar_texto(etapa) == "outro" and descricao else etapa
        for etapa, descricao in etapas
        if str(etapa or "").strip()
    )
    linhas_topo = [
        f"PACIENTE: {formatar_titulo(valor_row(paciente_row, 'nome'))}",
        f"PRONTUÁRIO: {formatar_prontuario_valor(paciente_row['prontuario'])}",
        f"DATA E HORA DE EMISSÃO: {data_emissao}",
        f"DATA DE RETORNO SOLICITADA: {formatar_data_br_valor(retorno_solicitado) or retorno_solicitado}",
    ]
    linhas_servico = [
        f"MATERIAL: {material_final}",
        f"SERVIÇO: {procedimento_nome}",
        f"ELEMENTO (S): {elemento_arcada}",
        f"ESCALA: {escala}" if str(escala or "").strip() else "ESCALA:",
        f"COR: {cor}" if str(cor or "").strip() else "COR:",
    ]
    linhas_carga = [
        f"CARGA IMEDIATA: {texto_ordem_servico_carga_imediata(carga_imediata)}",
    ]
    linhas_final = [
        "FORA DE PRAZO:",
        f"ETAPAS: {etapas_texto}" if etapas_texto else "ETAPAS:",
    ]
    if observacao.strip():
        linhas_final.append(f"OBSERVAÇÃO: {observacao.strip()}")

    try:
        for tabela in tabelas:
            total_colunas = len(tabela.columns)
            colunas_alvo = range(total_colunas) if total_colunas else [0]
            for coluna in colunas_alvo:
                preencher_celula_multilinha(tabela.cell(0, coluna), linhas_topo)
                preencher_celula_multilinha(tabela.cell(1, coluna), linhas_servico)
                preencher_celula_multilinha(tabela.cell(2, coluna), linhas_carga)
                preencher_celula_multilinha(tabela.cell(3, coluna), linhas_final)
            remover_linhas_extras_tabela_docx(tabela, 4)
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Falha ao preencher ordem de servico: {exc}") from exc

    os.makedirs(DOCS_DIR, exist_ok=True)
    nome_arquivo = nome_arquivo_ordem_servico(paciente_row, ordem_id)
    caminho = os.path.join(DOCS_DIR, nome_arquivo)
    doc.save(caminho)
    return nome_arquivo


def substituir_placeholders_paragrafo(paragrafo, dados: dict[str, str]) -> None:
    for run in paragrafo.runs:
        for chave, valor in dados.items():
            if chave in run.text:
                run.text = run.text.replace(chave, str(valor))

    texto_original = paragrafo.text
    texto_atualizado = texto_original
    for chave, valor in dados.items():
        texto_atualizado = texto_atualizado.replace(chave, str(valor))
    if texto_atualizado != texto_original:
        paragrafo.text = texto_atualizado


def substituir_runs_doc(doc, dados: dict[str, str]) -> None:
    for paragrafo in doc.paragraphs:
        substituir_placeholders_paragrafo(paragrafo, dados)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_placeholders_paragrafo(paragrafo, dados)


def montar_endereco_paciente_contrato(paciente_row: sqlite3.Row) -> str:
    endereco = formatar_titulo(valor_row(paciente_row, "endereco"))
    numero = valor_row(paciente_row, "numero")
    complemento = formatar_titulo(valor_row(paciente_row, "complemento"))
    bairro = formatar_titulo(valor_row(paciente_row, "bairro"))
    cidade = formatar_titulo(valor_row(paciente_row, "cidade"))
    estado = valor_row(paciente_row, "estado")
    cep = valor_row(paciente_row, "cep")

    partes: list[str] = []
    if endereco:
        linha = endereco
        if numero:
            linha = f"{linha}, {numero}"
        if complemento:
            linha = f"{linha}, {complemento}"
        partes.append(linha)
    elif numero:
        partes.append(numero)

    for item in [bairro, cidade, estado]:
        if item:
            partes.append(item)

    texto = " - ".join(partes)
    if cep:
        texto = f"{texto} - CEP {cep}" if texto else f"CEP {cep}"
    return texto


def montar_qualificacao_contrato(paciente_row: sqlite3.Row) -> str:
    nome = formatar_titulo(valor_row(paciente_row, "nome"))
    cpf = valor_row(paciente_row, "cpf")
    telefone = valor_row(paciente_row, "telefone")
    nascimento = formatar_data_br_valor(valor_row(paciente_row, "data_nascimento"))
    endereco = montar_endereco_paciente_contrato(paciente_row)
    menor_idade = valor_row(paciente_row, "menor_idade") in {"1", "TRUE", "True", "true"}
    responsavel = formatar_titulo(valor_row(paciente_row, "responsavel"))
    cpf_responsavel = valor_row(paciente_row, "cpf_responsavel")

    if menor_idade and responsavel:
        partes = [responsavel]
        if cpf_responsavel:
            partes.append(f"CPF {cpf_responsavel}")
        if telefone:
            partes.append(f"telefone {telefone}")
        if endereco:
            partes.append(f"endereco {endereco}")
        if nome:
            partes.append(f"responsavel legal pelo(a) paciente {nome}")
        if cpf:
            partes.append(f"CPF do paciente {cpf}")
        if nascimento:
            partes.append(f"nascido(a) em {nascimento}")
        return ", ".join(partes)

    partes = [nome]
    if cpf:
        partes.append(f"CPF {cpf}")
    if telefone:
        partes.append(f"telefone {telefone}")
    if nascimento:
        partes.append(f"nascido(a) em {nascimento}")
    if endereco:
        partes.append(f"endereco {endereco}")
    return ", ".join([parte for parte in partes if parte])


def dados_assinatura_contrato(paciente_row: sqlite3.Row) -> dict[str, str | bool]:
    nome = formatar_titulo(valor_row(paciente_row, "nome"))
    cpf = valor_row(paciente_row, "cpf")
    menor_idade = valor_row(paciente_row, "menor_idade") in {"1", "TRUE", "True", "true"}
    responsavel = formatar_titulo(valor_row(paciente_row, "responsavel"))
    cpf_responsavel = valor_row(paciente_row, "cpf_responsavel")

    if menor_idade and responsavel:
        return {
            "nome_assinatura": responsavel,
            "cpf_assinatura": f"CPF do responsavel: {cpf_responsavel}" if cpf_responsavel else "CPF do responsavel:",
            "assinatura_menor": True,
            "cpf_original": cpf,
        }
    return {
        "nome_assinatura": nome,
        "cpf_assinatura": f"CPF: {cpf}" if cpf else "CPF:",
        "assinatura_menor": False,
        "cpf_original": cpf,
    }


def texto_elemento(elemento) -> str:
    return "".join(elemento.itertext())


def remover_elemento(elemento) -> None:
    parent = elemento.getparent()
    if parent is not None:
        parent.remove(elemento)


def remover_secao_termo_cirurgico(doc) -> None:
    if etree is None:
        return
    corpo = doc._element.body
    removendo = False
    elementos_para_remover = []

    for elemento in corpo:
        texto = texto_elemento(elemento)
        if "{TERMO_CIRURGIA}" in texto:
            removendo = True

        if removendo:
            if "CONTRATO DE PRESTACAO DE SERVICOS ODONTOLOGICO" in normalizar_texto_maiusculo(texto):
                removendo = False
                continue
            elementos_para_remover.append(elemento)

    for elemento in elementos_para_remover:
        remover_elemento(elemento)

    for elemento in list(corpo):
        texto = texto_elemento(elemento).strip()
        xml = etree.tostring(elemento, encoding="unicode")
        if "lastRenderedPageBreak" in xml and not texto:
            remover_elemento(elemento)


def preencher_tabela_checklist_contrato(tabela, procedimentos: list[sqlite3.Row]) -> bool:
    if Inches is None:
        return False
    for row_index, row in enumerate(tabela.rows):
        for cell in row.cells:
            if "{PROCEDIMENTO}" not in cell.text:
                continue

            if len(tabela.columns) == 1:
                tabela.add_column(Inches(1.0))

            cabecalho = tabela.rows[row_index]
            cabecalho.cells[0].text = "Procedimento"
            cabecalho.cells[1].text = "OK"

            while len(tabela.rows) < row_index + 1 + len(procedimentos):
                tabela.add_row()

            for indice_proc, proc_row in enumerate(procedimentos, start=1):
                linha = tabela.rows[row_index + indice_proc]
                linha.cells[0].text = str(proc_row["procedimento"] or "")
                linha.cells[1].text = "[   ]"
            return True
    return False


def preencher_tabela_contrato_docx(tabela, procedimentos: list[sqlite3.Row]) -> bool:
    for row_index, row in enumerate(tabela.rows):
        if len(row.cells) < 2:
            continue
        tem_procedimento = any("{PROCEDIMENTO}" in cell.text for cell in row.cells)
        tem_valor = any("{VALOR}" in cell.text for cell in row.cells)
        if not (tem_procedimento or tem_valor):
            continue

        while len(tabela.rows) < row_index + len(procedimentos):
            tabela.add_row()

        for indice_proc, proc_row in enumerate(procedimentos):
            linha = tabela.rows[row_index + indice_proc]
            linha.cells[0].text = str(proc_row["procedimento"] or "")
            linha.cells[1].text = formatar_moeda_br(proc_row["valor"])
        return True
    return False


def carregar_procedimentos_documento_contrato(conn: sqlite3.Connection, contrato_id: int) -> list[dict]:
    procedimentos = conn.execute(
        "SELECT id, procedimento, valor, profissional_snapshot, denticao_snapshot FROM procedimentos_contrato WHERE contrato_id=? ORDER BY id",
        (contrato_id,),
    ).fetchall()
    dentes = conn.execute(
        """
        SELECT procedimento, dente, regiao, status
        FROM procedimentos_dente
        WHERE contrato_id=? AND upper(trim(COALESCE(status, ''))) <> 'EXCLUIDO'
        ORDER BY id
        """,
        (contrato_id,),
    ).fetchall()

    regioes_por_procedimento: dict[str, list[str]] = {}
    for row in dentes:
        chave = normalizar_texto(row["procedimento"])
        valor = str(row["dente"]) if row["dente"] is not None else str(row["regiao"] or "").strip()
        if not valor:
            continue
        regioes_por_procedimento.setdefault(chave, [])
        if valor not in regioes_por_procedimento[chave]:
            regioes_por_procedimento[chave].append(valor)

    itens: list[dict] = []
    for row in procedimentos:
        procedimento = str(row["procedimento"] or "").strip()
        chave = normalizar_texto(procedimento)
        regioes = regioes_por_procedimento.get(chave, [])
        sufixo = ""
        if regioes:
            rotulo = "Elemento" if len(regioes) == 1 and regioes[0].isdigit() else "Elementos"
            sufixo = f" - {rotulo}: {', '.join(regioes)}"
        itens.append(
            {
                "procedimento": f"{procedimento}{sufixo}",
                "valor": float(row["valor"] or 0),
                "profissional_snapshot": str(row["profissional_snapshot"] or ""),
                "denticao_snapshot": str(row["denticao_snapshot"] or ""),
                "regioes": regioes,
            }
        )
    return itens


def montar_termo_cirurgia_contrato(procedimentos: list[sqlite3.Row]) -> str:
    itens = [str(row["procedimento"] or "").strip() for row in procedimentos if str(row["procedimento"] or "").strip()]
    if not itens:
        return ""
    if len(itens) == 1:
        return itens[0]
    return "; ".join(itens)


def aplicar_fonte_times_new_roman(doc) -> None:
    if qn is None or Pt is None:
        return

    estilos = doc.styles
    for nome_estilo in ["Normal", "Header", "Footer"]:
        if nome_estilo in estilos:
            estilo = estilos[nome_estilo]
            estilo.font.name = "Times New Roman"
            estilo._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            estilo._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            estilo._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    def aplicar_paragrafo(paragrafo) -> None:
        for run in paragrafo.runs:
            run.font.name = "Times New Roman"
            if run._element.rPr is None:
                run._element.get_or_add_rPr()
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    for paragrafo in doc.paragraphs:
        aplicar_paragrafo(paragrafo)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    aplicar_paragrafo(paragrafo)


def garantir_logo_no_cabecalho(doc) -> None:
    if not doc.sections:
        return
    doc.sections[0].different_first_page_header_footer = False
    for secao in doc.sections[1:]:
        secao.different_first_page_header_footer = False
        secao.header.is_linked_to_previous = True


def substituir_par_assinatura(paragrafo_nome, paragrafo_cpf, assinatura: dict[str, str | bool]) -> None:
    paragrafo_nome.text = str(assinatura["nome_assinatura"])
    paragrafo_cpf.text = str(assinatura["cpf_assinatura"])


def ajustar_assinatura_capa(doc, nome_paciente: str, cpf_paciente: str, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    titulo_contrato = "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO"
    limite = len(paragrafos)

    for indice, paragrafo in enumerate(paragrafos):
        if titulo_contrato in normalizar_texto_maiusculo(paragrafo.text):
            limite = indice
            break

    for indice in range(limite - 1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip() if indice + 1 < limite else ""
        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            substituir_par_assinatura(paragrafos[indice], paragrafos[indice + 1], assinatura)
            return


def ajustar_assinaturas_menor(doc, nome_paciente: str, cpf_paciente: str, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip()
        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            substituir_par_assinatura(paragrafos[indice], paragrafos[indice + 1], assinatura)


def ajustar_assinaturas_por_texto(doc, nome_paciente: str, cpf_paciente: str, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    for paragrafo in paragrafos:
        texto = paragrafo.text
        if not texto:
            continue
        texto_ajustado = texto
        if nome_paciente and nome_paciente in texto_ajustado:
            texto_ajustado = texto_ajustado.replace(nome_paciente, str(assinatura["nome_assinatura"]))
        if cpf_paciente:
            texto_ajustado = texto_ajustado.replace(f"CPF: {cpf_paciente}", str(assinatura["cpf_assinatura"]))
            texto_ajustado = texto_ajustado.replace(f"CPF {cpf_paciente}", str(assinatura["cpf_assinatura"]))
        if texto_ajustado != texto:
            paragrafo.text = texto_ajustado


def ajustar_ultima_assinatura_menor(doc, nome_paciente: str, cpf_paciente: str, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 2, -1, -1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip()
        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            substituir_par_assinatura(paragrafos[indice], paragrafos[indice + 1], assinatura)
            return


def ajustar_bloco_final_assinatura(doc, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, 1, -1):
        if "SOUL SUL CLINICA INTEGRADA" not in normalizar_texto_maiusculo(paragrafos[indice].text):
            continue
        indice_linha = None
        for cursor in range(indice - 1, -1, -1):
            texto = paragrafos[cursor].text.strip()
            if "____" in texto or "___" in texto:
                indice_linha = cursor
                break
        if indice_linha is None:
            return
        candidatos = []
        for cursor in range(indice_linha + 1, indice):
            texto = paragrafos[cursor].text.strip()
            if texto:
                candidatos.append(cursor)
        if len(candidatos) < 2:
            return
        indice_nome = candidatos[0]
        indice_cpf = candidatos[1]
        paragrafos[indice_nome].text = str(assinatura["nome_assinatura"])
        paragrafos[indice_cpf].text = str(assinatura["cpf_assinatura"])
        paragrafos[indice_nome].alignment = 1
        paragrafos[indice_cpf].alignment = 1
        for cursor in candidatos[2:]:
            paragrafos[cursor].text = ""
        return


def ajustar_assinatura_termo_cirurgico_final(doc, nome_paciente: str, cpf_paciente: str, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    ultimo_indice_data = None
    ultimo_indice_clinica = None
    for indice, paragrafo in enumerate(paragrafos):
        texto = paragrafo.text.strip()
        if "Campos dos Goytacazes" in texto:
            ultimo_indice_data = indice
        if "SOUL SUL CLINICA INTEGRADA" in normalizar_texto_maiusculo(texto):
            ultimo_indice_clinica = indice
    if ultimo_indice_data is None or ultimo_indice_clinica is None or ultimo_indice_data >= ultimo_indice_clinica:
        return
    candidatos = []
    for indice in range(ultimo_indice_data + 1, ultimo_indice_clinica):
        texto = paragrafos[indice].text.strip()
        if texto:
            candidatos.append(indice)
    if len(candidatos) < 2:
        return
    for posicao in range(len(candidatos) - 1):
        indice_nome = candidatos[posicao]
        indice_cpf = candidatos[posicao + 1]
        texto_nome = paragrafos[indice_nome].text.strip()
        texto_cpf = paragrafos[indice_cpf].text.strip()
        if texto_nome == nome_paciente and cpf_paciente and cpf_paciente in texto_cpf:
            substituir_par_assinatura(paragrafos[indice_nome], paragrafos[indice_cpf], assinatura)
            return


def ajustar_assinatura_contrato(doc, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, 1, -1):
        if "SOUL SUL CLINICA INTEGRADA" not in normalizar_texto_maiusculo(paragrafos[indice].text):
            continue
        candidatos = []
        cursor = indice - 1
        while cursor >= 0 and len(candidatos) < 2:
            texto = paragrafos[cursor].text.strip()
            if texto:
                candidatos.append(cursor)
            cursor -= 1
        if len(candidatos) < 2:
            return
        indice_cpf = candidatos[0]
        indice_nome = candidatos[1]
        paragrafos[indice_nome].text = str(assinatura["nome_assinatura"])
        paragrafos[indice_cpf].text = str(assinatura["cpf_assinatura"])
        return


def reescrever_bloco_final_termo(doc, assinatura: dict[str, str | bool]) -> None:
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, -1, -1):
        texto = paragrafo_texto = paragrafos[indice].text
        if "SOUL SUL" not in normalizar_texto_maiusculo(texto) or "TESTEMUNHAS" not in normalizar_texto_maiusculo(texto):
            continue
        candidatos_limpar = []
        for cursor in range(indice - 1, -1, -1):
            texto_cursor = paragrafos[cursor].text.strip()
            if not texto_cursor:
                continue
            candidatos_limpar.append(cursor)
            if len(candidatos_limpar) == 2:
                break
        for cursor in candidatos_limpar:
            paragrafos[cursor].text = ""
        paragrafos[indice].text = (
            "\n"
            "___________________________________________________________\n"
            f"{assinatura['nome_assinatura']}\n"
            f"{assinatura['cpf_assinatura']}\n\n"
            "___________________________________________________________\n"
            "SOUL SUL CLINICA INTEGRADA\n\n"
            "TESTEMUNHAS:\n\n"
            "________________________________                    _________________________________\n"
            "NOME:                                                          NOME:\n"
            "CPF:                                                             CPF:"
        )
        paragrafos[indice].alignment = 1
        return


def normalizar_quebra_antes_titulo(doc, marcador: str) -> None:
    if WD_BREAK is None:
        return
    paragrafos = doc.paragraphs
    indice_titulo = None
    marcador_norm = normalizar_texto_maiusculo(marcador)

    for indice, paragrafo in enumerate(paragrafos):
        if marcador_norm in normalizar_texto_maiusculo(paragrafo.text):
            indice_titulo = indice
            break
    if indice_titulo is None or indice_titulo == 0:
        return

    paragrafo_anterior = paragrafos[indice_titulo - 1]
    tem_quebra = any(
        child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
        for run in paragrafo_anterior.runs
        for child in run._element
    )
    if not tem_quebra:
        paragrafo_anterior.add_run().add_break(WD_BREAK.PAGE)


def compactar_bloco_final(doc) -> None:
    if Pt is None:
        return
    paragrafos = doc.paragraphs
    indice_data = None
    indice_testemunhas = None

    for indice in range(len(paragrafos) - 1, -1, -1):
        texto = paragrafos[indice].text.strip()
        if indice_testemunhas is None and "TESTEMUNHAS" in normalizar_texto_maiusculo(texto):
            indice_testemunhas = indice
        if indice_data is None and "CAMPOS DOS GOYTACAZES" in normalizar_texto_maiusculo(texto):
            indice_data = indice
        if indice_data is not None and indice_testemunhas is not None:
            break

    if indice_data is None or indice_testemunhas is None or indice_data >= indice_testemunhas:
        return

    for indice in range(indice_data, len(paragrafos)):
        texto = paragrafos[indice].text.strip()
        if not texto:
            continue
        formato = paragrafos[indice].paragraph_format
        formato.space_before = Pt(0)
        formato.space_after = Pt(0)
        formato.line_spacing = 1.0
        for run in paragrafos[indice].runs:
            run.font.size = Pt(10)
        if "TESTEMUNHAS" in normalizar_texto_maiusculo(texto):
            break


def montar_texto_pagamento_contrato(contrato: sqlite3.Row, plano: list[dict]) -> str:
    valor_total = float(contrato["valor_total"] or 0)
    entrada_valor = float(contrato["entrada"] or 0)
    forma = str(contrato["forma_pagamento"] or "").replace("_", " ").upper()
    formas_no_dia = {"CARTAO CREDITO", "CARTAO DE CREDITO", "PIX", "CARTAO DE DEBITO", "CARTAO DEBITO", "DINHEIRO"}

    itens_ativos = [item for item in plano if float(item.get("valor", 0) or 0) > 0]
    entrada = next((item for item in itens_ativos if normalizar_texto(item.get("descricao", "")) == "entrada"), None)
    posteriores = [item for item in itens_ativos if normalizar_texto(item.get("descricao", "")) != "entrada"]

    if not posteriores and entrada:
        return f"VALOR TOTAL {formatar_moeda_br(valor_total)}. PAGO NO {str(entrada.get('forma', forma)).replace('_', ' ').upper()} NO DIA {formatar_data_br_valor(entrada.get('data'))}."

    if not posteriores:
        return f"VALOR TOTAL {formatar_moeda_br(valor_total)}."

    valor_restante = sum(float(item.get("valor", 0) or 0) for item in posteriores)
    primeira = posteriores[0]
    valor_parcela = float(primeira.get("valor", 0) or 0)
    forma_primeira = str(primeira.get("forma", forma)).replace("_", " ").upper()
    sufixo_primeira = (
        f"NO DIA {formatar_data_br_valor(primeira.get('data'))}."
        if forma_primeira in formas_no_dia
        else f"COM PRIMEIRO VENCIMENTO EM {formatar_data_br_valor(primeira.get('data'))}."
    )
    if entrada and entrada_valor > 0:
        forma_entrada = str(entrada.get("forma", forma)).replace("_", " ").upper()
        return (
            f"VALOR TOTAL {formatar_moeda_br(valor_total)}. "
            f"ENTRADA {formatar_moeda_br(entrada_valor)} PAGA NO {forma_entrada} NO DIA {formatar_data_br_valor(entrada.get('data'))}. "
            f"RESTANTE {formatar_moeda_br(valor_restante)} EM {len(posteriores)} PARCELAS DE {formatar_moeda_br(valor_parcela)} NO {forma_primeira} "
            f"{sufixo_primeira}"
        )

    return (
        f"VALOR TOTAL {formatar_moeda_br(valor_total)}. "
        f"PAGO EM {len(posteriores)} PARCELAS DE {formatar_moeda_br(valor_parcela)} NO {forma_primeira} "
        f"{sufixo_primeira}"
    )


def carregar_plano_pagamento_contrato(contrato: sqlite3.Row) -> list[dict]:
    try:
        bruto = json.loads(str(contrato["plano_pagamento_json"] or "[]"))
    except (TypeError, ValueError, json.JSONDecodeError):
        bruto = []
    plano = [item for item in bruto if isinstance(item, dict)]
    if plano:
        return plano

    valor_total = float(contrato["valor_total"] or 0)
    if valor_total <= 0:
        return []

    entrada = float(contrato["entrada"] or 0)
    parcelas = max(int(contrato["parcelas"] or 1), 1)
    forma = str(contrato["forma_pagamento"] or "A Definir").strip() or "A Definir"
    data_base = parse_data_contrato(contrato["primeiro_vencimento"]) or parse_data_contrato(contrato["data_criacao"]) or date.today()
    plano_legacy: list[dict] = []

    if entrada > 0:
        data_entrada = parse_data_contrato(contrato["data_pagamento_entrada"]) or data_base
        plano_legacy.append(
            {
                "indice": 0,
                "descricao": "Entrada",
                "data": data_entrada.isoformat(),
                "forma": forma,
                "valor": round(entrada, 2),
                "parcelas_cartao": 1,
            }
        )

    restante = max(valor_total - entrada, 0)
    if restante <= 0:
        return plano_legacy

    quantidade_parcelas = max(parcelas, 1)
    valor_base = round(restante / quantidade_parcelas, 2)
    acumulado = 0.0
    for indice in range(quantidade_parcelas):
        valor_parcela = valor_base
        if indice == quantidade_parcelas - 1:
            valor_parcela = round(restante - acumulado, 2)
        acumulado += valor_parcela
        plano_legacy.append(
            {
                "indice": indice + 1,
                "descricao": str(indice + 1),
                "data": adicionar_meses(data_base, indice).isoformat(),
                "forma": forma,
                "valor": valor_parcela,
                "parcelas_cartao": quantidade_parcelas if "CARTAO" in normalizar_texto(forma).upper() else 1,
            }
        )

    return plano_legacy


def gerar_html_contrato_fallback(
    paciente_row: sqlite3.Row,
    contrato_id: int,
    procedimentos: list[dict],
    plano: list[dict],
) -> str:
    caminho_html = os.path.join(DOCS_DIR, f"{nome_base_contrato(paciente_row, contrato_id)}.html")
    linhas_procedimentos = "".join(
        f"<tr><td>{row['procedimento']}</td><td>{row['profissional_snapshot'] or ''}</td><td>{row['denticao_snapshot'] or ''}</td><td style='text-align:right'>{formatar_moeda_br(row['valor'])}</td></tr>"
        for row in procedimentos
    )
    linhas_pagamento = "".join(
        f"<tr><td>{item.get('descricao', '')}</td><td>{formatar_data_br_valor(item.get('data'))}</td><td>{str(item.get('forma', '')).replace('_', ' ')}</td><td style='text-align:right'>{formatar_moeda_br(item.get('valor', 0))}</td></tr>"
        for item in plano
        if float(item.get('valor', 0) or 0) > 0
    )
    html = f"""<!doctype html>
<html lang="pt-BR"><head><meta charset="utf-8"><title>Contrato {contrato_id}</title></head>
<body><h1>Contrato #{contrato_id}</h1><p>{formatar_titulo(paciente_row['nome'] or '')}</p>
<table>{linhas_procedimentos}</table><table>{linhas_pagamento}</table></body></html>"""
    with open(caminho_html, "w", encoding="utf-8") as arquivo:
        arquivo.write(html)
    return caminho_html


def gerar_docx_contrato_fallback(
    paciente_row: sqlite3.Row,
    contrato_id: int,
    procedimentos: list[dict],
    plano: list[dict],
) -> str:
    if Document is None:
        return gerar_html_contrato_fallback(paciente_row, contrato_id, procedimentos, plano)

    nome_arquivo = nome_arquivo_contrato(paciente_row, contrato_id)
    caminho = os.path.join(DOCS_DIR, nome_arquivo)
    doc = Document()
    doc.add_heading(f"Contrato #{contrato_id}", level=1)
    doc.add_paragraph(f"Paciente: {formatar_titulo(valor_row(paciente_row, 'nome'))}")
    doc.add_paragraph(f"Prontuário: {formatar_prontuario_valor(paciente_row['prontuario'])}")
    doc.add_paragraph("")

    doc.add_paragraph("Procedimentos")
    tabela_proc = doc.add_table(rows=1, cols=4)
    header_proc = tabela_proc.rows[0].cells
    header_proc[0].text = "Procedimento"
    header_proc[1].text = "Profissional"
    header_proc[2].text = "Dentição"
    header_proc[3].text = "Valor"
    for row in procedimentos:
        cells = tabela_proc.add_row().cells
        cells[0].text = str(row.get("procedimento", "") or "")
        cells[1].text = str(row.get("profissional_snapshot", "") or "")
        cells[2].text = str(row.get("denticao_snapshot", "") or "")
        cells[3].text = formatar_moeda_br(row.get("valor", 0))

    doc.add_paragraph("")
    doc.add_paragraph("Plano de pagamento")
    tabela_pag = doc.add_table(rows=1, cols=4)
    header_pag = tabela_pag.rows[0].cells
    header_pag[0].text = "Descrição"
    header_pag[1].text = "Data"
    header_pag[2].text = "Forma"
    header_pag[3].text = "Valor"
    for item in plano:
        if float(item.get("valor", 0) or 0) <= 0:
            continue
        cells = tabela_pag.add_row().cells
        cells[0].text = str(item.get("descricao", "") or "")
        cells[1].text = formatar_data_br_valor(item.get("data"))
        cells[2].text = str(item.get("forma", "") or "").replace("_", " ")
        cells[3].text = formatar_moeda_br(item.get("valor", 0))

    aplicar_fonte_times_new_roman(doc)
    doc.save(caminho)
    return caminho


def gerar_documento_contrato(
    conn: sqlite3.Connection,
    paciente_row: sqlite3.Row,
    contrato: sqlite3.Row,
    contrato_id: int,
) -> str:
    os.makedirs(DOCS_DIR, exist_ok=True)
    procedimentos = carregar_procedimentos_documento_contrato(conn, contrato_id)
    plano = carregar_plano_pagamento_contrato(contrato)
    nome_arquivo = nome_arquivo_contrato(paciente_row, contrato_id)
    caminho = os.path.join(DOCS_DIR, nome_arquivo)

    if Document is None or not os.path.isfile(TEMPLATE_PATH):
        return gerar_docx_contrato_fallback(paciente_row, contrato_id, procedimentos, plano)

    try:
        doc = Document(TEMPLATE_PATH)
        garantir_logo_no_cabecalho(doc)
        termo_cirurgia = montar_termo_cirurgia_contrato(procedimentos)
        if not termo_cirurgia:
            remover_secao_termo_cirurgico(doc)
        nome_paciente = formatar_titulo(valor_row(paciente_row, "nome"))
        cpf_paciente = valor_row(paciente_row, "cpf")
        assinatura = dados_assinatura_contrato(paciente_row)
        qualificacao = montar_qualificacao_contrato(paciente_row)
        pagamento = montar_texto_pagamento_contrato(contrato, plano)

        dados = {
            "{PACIENTE}": nome_paciente,
            "{paciente}": nome_paciente,
            "{CPF}": cpf_paciente,
            "{cpf}": cpf_paciente,
            "{PRONTUARIO}": formatar_prontuario_valor(paciente_row["prontuario"]),
            "{prontuario}": formatar_prontuario_valor(paciente_row["prontuario"]),
            "{PAGAMENTO}": pagamento,
            "{pagamento}": pagamento,
            "{QUALIFICACAO}": qualificacao,
            "{qualificacao}": qualificacao,
            "{ENDERECO}": formatar_titulo(valor_row(paciente_row, "endereco")),
            "{endereco}": formatar_titulo(valor_row(paciente_row, "endereco")),
            "{NUMERO}": valor_row(paciente_row, "numero"),
            "{numero}": valor_row(paciente_row, "numero"),
            "{BAIRRO}": formatar_titulo(valor_row(paciente_row, "bairro")),
            "{bairro}": formatar_titulo(valor_row(paciente_row, "bairro")),
            "{CIDADE}": formatar_titulo(valor_row(paciente_row, "cidade")),
            "{cidade}": formatar_titulo(valor_row(paciente_row, "cidade")),
            "{ESTADO}": valor_row(paciente_row, "estado"),
            "{estado}": valor_row(paciente_row, "estado"),
            "{CEP}": valor_row(paciente_row, "cep"),
            "{cep}": valor_row(paciente_row, "cep"),
            "{TELEFONE}": valor_row(paciente_row, "telefone"),
            "{telefone}": valor_row(paciente_row, "telefone"),
            "{DATA_NASCIMENTO}": formatar_data_br_valor(valor_row(paciente_row, "data_nascimento")),
            "{data_nascimento}": formatar_data_br_valor(valor_row(paciente_row, "data_nascimento")),
            "{TERMO_CIRURGIA}": termo_cirurgia,
            "{termo_cirurgia}": termo_cirurgia,
        }
        substituir_runs_doc(doc, dados)

        checklist_preenchido = False
        tabela_contrato_preenchida = False
        for tabela in doc.tables:
            if not checklist_preenchido and preencher_tabela_checklist_contrato(tabela, procedimentos):
                checklist_preenchido = True
                continue
            if not tabela_contrato_preenchida and preencher_tabela_contrato_docx(tabela, procedimentos):
                tabela_contrato_preenchida = True

        if bool(assinatura["assinatura_menor"]):
            ajustar_assinaturas_menor(doc, nome_paciente, cpf_paciente, assinatura)
            ajustar_assinaturas_por_texto(doc, nome_paciente, cpf_paciente, assinatura)
            ajustar_assinatura_termo_cirurgico_final(doc, nome_paciente, cpf_paciente, assinatura)
            ajustar_ultima_assinatura_menor(doc, nome_paciente, cpf_paciente, assinatura)
            ajustar_bloco_final_assinatura(doc, assinatura)
        else:
            ajustar_assinatura_capa(doc, nome_paciente, cpf_paciente, assinatura)
            ajustar_assinatura_contrato(doc, assinatura)
        reescrever_bloco_final_termo(doc, assinatura)
        normalizar_quebra_antes_titulo(doc, "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO")
        normalizar_quebra_antes_titulo(doc, "TERMO DE CONSENTIMENTO ESCLARECIDO")
        compactar_bloco_final(doc)
        aplicar_fonte_times_new_roman(doc)
        doc.save(caminho)
        return caminho
    except Exception as exc:
        print(f"[contrato] fallback docx acionado contrato={contrato_id} erro={exc}", flush=True)
        return gerar_docx_contrato_fallback(paciente_row, contrato_id, procedimentos, plano)


def limpar_documento_contrato(paciente_row: sqlite3.Row, contrato_id: int) -> None:
    limpar_documentos_contrato_variantes(paciente_row, contrato_id)


def sincronizar_recebiveis_contrato(conn: sqlite3.Connection, paciente_row: sqlite3.Row, contrato: sqlite3.Row, contrato_id: int) -> None:
    conn.execute("DELETE FROM recebiveis WHERE contrato_id=?", (contrato_id,))
    plano = carregar_plano_pagamento_contrato(contrato)
    prontuario = formatar_prontuario_valor(paciente_row["prontuario"])
    for item in plano:
        forma = str(item.get("forma", "") or "").strip().upper()
        valor_item = float(item.get("valor", 0) or 0)
        if valor_item <= 0:
            continue
        descricao_item = normalizar_texto(item.get("descricao", ""))
        if descricao_item == "entrada":
            parcela_numero = 0
        else:
            try:
                parcela_numero = int(str(item.get("descricao", "")).strip())
            except (TypeError, ValueError):
                parcela_numero = int(item.get("indice", 0) or 0) + 1
        conn.execute(
            """
            INSERT INTO recebiveis (
                contrato_id, paciente_id, paciente_nome, prontuario, parcela_numero, vencimento, valor,
                forma_pagamento, status, observacao, data_criacao, data_pagamento
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'Aberto', ?, ?, NULL)
            """,
            (
                contrato_id,
                int(paciente_row["id"]),
                str(paciente_row["nome"] or ""),
                prontuario,
                parcela_numero,
                str(item.get("data", "") or ""),
                valor_item,
                forma.replace("_", " "),
                f"CONTRATO_APROVADO_{contrato_id}",
                agora_str(),
            ),
        )


def sincronizar_financeiro_contrato(conn: sqlite3.Connection, paciente_row: sqlite3.Row, contrato: sqlite3.Row, contrato_id: int) -> None:
    conn.execute("DELETE FROM financeiro WHERE contrato_id=?", (contrato_id,))
    # O financeiro registra apenas caixa real. O contrato aprovado gera
    # recebiveis planejados; a entrada no financeiro acontece somente quando
    # a parcela e baixada no caixa.


def mapear_recebivel(row: sqlite3.Row) -> RecebivelResumo:
    return RecebivelResumo(
        id=int(row["id"]),
        pacienteId=int(row["paciente_id"]) if row["paciente_id"] is not None else None,
        pacienteNome=str(row["paciente_nome"] or ""),
        prontuario=formatar_prontuario_valor(row["prontuario"]),
        contratoId=int(row["contrato_id"]) if row["contrato_id"] is not None else None,
        parcela=int(row["parcela_numero"]) if row["parcela_numero"] is not None else None,
        vencimento=formatar_data_br_valor(row["vencimento"]),
        valor=formatar_moeda_br(row["valor"]),
        formaPagamento=str(row["forma_pagamento"] or ""),
        status=str(row["status"] or ""),
        dataPagamento=formatar_data_br_valor(row["data_pagamento"]) if "data_pagamento" in row.keys() else "",
        observacao=str(row["observacao"] or ""),
    )


def carregar_recebivel_paciente(conn: sqlite3.Connection, paciente_id: int, recebivel_id: int) -> sqlite3.Row:
    row = conn.execute(
        """
        SELECT *
        FROM recebiveis
        WHERE id=? AND paciente_id=?
        LIMIT 1
        """,
        (recebivel_id, paciente_id),
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Recebível não encontrado.")
    return row


def mapear_agendamento(row: sqlite3.Row) -> AgendamentoResumo:
    procedimento = str(row["procedimento_nome_snapshot"] or row["procedimento"] or "")
    horario_inicio = str(row["hora_inicio"] or "")
    horario_fim = str(row["hora_fim"] or "")
    horario = f"{horario_inicio} - {horario_fim}".strip(" -")
    return AgendamentoResumo(
        id=int(row["id"]),
        data=formatar_data_br_valor(row[DATA_COLUNA_AGENDA] if DATA_COLUNA_AGENDA in row.keys() else row["data"]),
        horario=horario,
        profissional=str(row["profissional"] or ""),
        status=str(row["status"] or ""),
        procedimento=procedimento,
        observacao=str(row["observacao"] or row["observacoes"] or ""),
    )


CRM_ETAPAS_PADRAO = [
    "Novo lead",
    "Contato inicial",
    "Tentando contato",
    "Conversando",
    "Agendou avaliação",
    "Em negociação",
    "Convertido",
    "Perdido",
]


def mapear_crm_paciente_resumo(row: sqlite3.Row | None) -> CrmPacienteResumoResposta | None:
    if row is None:
        return None
    return CrmPacienteResumoResposta(
        crmId=int(row["id"]) if row["id"] is not None else None,
        finalizado=bool(int(row["origem_finalizado"] or 0)),
        avaliacao=bool(int(row["origem_avaliacao"] or 0)),
        etapaFunil=str(row["etapa_funil"] or "Novo lead"),
        campanha=str(row["campanha"] or ""),
        canal=str(row["canal"] or "Facebook"),
        ultimaAvaliacaoEm=formatar_data_br_valor(row["ultima_avaliacao_em"]),
        finalizadoEm=formatar_data_br_valor(row["finalizado_em"]),
    )


def mapear_crm_paciente_item(row: sqlite3.Row) -> CrmPacienteItemResposta:
    return CrmPacienteItemResposta(
        id=int(row["id"]),
        pacienteId=int(row["paciente_id"]),
        nome=str(row["nome"] or ""),
        prontuario=formatar_prontuario_valor(row["prontuario"]),
        telefone=str(row["telefone"] or ""),
        origemFinalizado=bool(int(row["origem_finalizado"] or 0)),
        origemAvaliacao=bool(int(row["origem_avaliacao"] or 0)),
        etapaFunil=str(row["etapa_funil"] or "Novo lead"),
        canal=str(row["canal"] or "Facebook"),
        campanha=str(row["campanha"] or ""),
        conjuntoAnuncio=str(row["conjunto_anuncio"] or ""),
        anuncio=str(row["anuncio"] or ""),
        responsavel=str(row["responsavel"] or ""),
        proximoContato=formatar_data_br_valor(row["proximo_contato"]),
        observacao=str(row["observacao"] or ""),
        ultimaInteracao=formatar_data_br_valor(row["ultima_interacao"]),
        ultimaAvaliacaoEm=formatar_data_br_valor(row["ultima_avaliacao_em"]),
        finalizadoEm=formatar_data_br_valor(row["finalizado_em"]),
        atualizadoEm=str(row["atualizado_em"] or ""),
    )


def crm_entry_por_paciente(conn: sqlite3.Connection, paciente_id: int) -> sqlite3.Row | None:
    return conn.execute("SELECT * FROM crm_pacientes WHERE paciente_id=? LIMIT 1", (int(paciente_id),)).fetchone()


def agendamento_eh_avaliacao(row: sqlite3.Row) -> bool:
    blocos = [
        str(row["tipo_atendimento_nome_snapshot"] or ""),
        str(row["procedimento_nome_snapshot"] or ""),
        str(row["procedimento"] or ""),
        str(row["observacoes"] or row["observacao"] or ""),
    ]
    texto = " ".join(blocos)
    normalizado = normalizar_texto(texto)
    return "avaliacao" in normalizado or re.search(r"\baval\b", normalizado) is not None


def listar_avaliacoes_crm(conn: sqlite3.Connection) -> list[CrmAvaliacaoItemResposta]:
    rows = conn.execute(
        """
        SELECT
            a.id,
            a.paciente_id,
            COALESCE(NULLIF(p.nome, ''), NULLIF(a.nome_paciente_snapshot, ''), NULLIF(a.paciente_nome, '')) AS nome,
            COALESCE(p.prontuario, '') AS prontuario,
            COALESCE(NULLIF(p.telefone, ''), NULLIF(a.telefone_snapshot, '')) AS telefone,
            COALESCE(a.data_agendamento, a.data) AS data_referencia,
            COALESCE(a.hora_inicio, '') AS hora_inicio,
            COALESCE(a.profissional, '') AS profissional,
            COALESCE(a.status, '') AS status,
            COALESCE(NULLIF(a.procedimento_nome_snapshot, ''), NULLIF(a.procedimento, '')) AS procedimento,
            COALESCE(a.tipo_atendimento_nome_snapshot, '') AS tipo_atendimento,
            COALESCE(crm.id, 0) AS crm_id,
            COALESCE(crm.origem_avaliacao, 0) AS origem_avaliacao
        FROM agendamentos a
        LEFT JOIN pacientes p ON p.id = a.paciente_id
        LEFT JOIN crm_pacientes crm ON crm.paciente_id = a.paciente_id
        WHERE a.paciente_id IS NOT NULL
        ORDER BY COALESCE(a.data_agendamento, a.data) DESC, COALESCE(a.hora_inicio, '') DESC, a.id DESC
        """
    ).fetchall()
    por_paciente: dict[int, CrmAvaliacaoItemResposta] = {}
    hoje = date.today()
    for row in rows:
        paciente_id = int(row["paciente_id"] or 0)
        if paciente_id <= 0 or paciente_id in por_paciente:
            continue
        if not agendamento_eh_avaliacao(row):
            continue
        status = normalizar_texto(str(row["status"] or ""))
        if status in {"cancelado", "desmarcado", "faltou"}:
            continue
        data_ref = parse_data_contrato(row["data_referencia"])
        if not data_ref or data_ref > hoje:
            continue
        data_avaliacao = formatar_data_br_valor(row["data_referencia"])
        if str(row["hora_inicio"] or "").strip():
            data_avaliacao = f"{data_avaliacao} {str(row['hora_inicio'] or '').strip()}".strip()
        por_paciente[paciente_id] = CrmAvaliacaoItemResposta(
            pacienteId=paciente_id,
            nome=str(row["nome"] or ""),
            prontuario=formatar_prontuario_valor(row["prontuario"]),
            telefone=str(row["telefone"] or ""),
            dataAvaliacao=data_avaliacao,
            profissional=str(row["profissional"] or ""),
            status=str(row["status"] or ""),
            procedimento=str(row["procedimento"] or row["tipo_atendimento"] or ""),
            jaNoCrm=bool(int(row["crm_id"] or 0)),
        )
    return list(por_paciente.values())


def upsert_crm_origem(
    conn: sqlite3.Connection,
    paciente_id: int,
    *,
    marcar_finalizado: bool = False,
    marcar_avaliacao: bool = False,
    usuario: str = "",
    ultima_avaliacao_em: str = "",
) -> sqlite3.Row:
    existente = crm_entry_por_paciente(conn, paciente_id)
    agora = agora_str()
    usuario_limpo = str(usuario or "").strip()
    if existente is None:
        conn.execute(
            """
            INSERT INTO crm_pacientes (
                paciente_id, origem_finalizado, origem_avaliacao, etapa_funil, canal,
                ultima_avaliacao_em, finalizado_em, criado_por, atualizado_por, criado_em, atualizado_em
            )
            VALUES (?, ?, ?, 'Novo lead', 'Facebook', ?, ?, ?, ?, ?, ?)
            """,
            (
                int(paciente_id),
                1 if marcar_finalizado else 0,
                1 if marcar_avaliacao else 0,
                ultima_avaliacao_em or "",
                agora if marcar_finalizado else "",
                usuario_limpo,
                usuario_limpo,
                agora,
                agora,
            ),
        )
    else:
        conn.execute(
            """
            UPDATE crm_pacientes
            SET
                origem_finalizado=?,
                origem_avaliacao=?,
                ultima_avaliacao_em=?,
                finalizado_em=?,
                atualizado_por=?,
                atualizado_em=?
            WHERE paciente_id=?
            """,
            (
                1 if (bool(int(existente["origem_finalizado"] or 0)) or marcar_finalizado) else 0,
                1 if (bool(int(existente["origem_avaliacao"] or 0)) or marcar_avaliacao) else 0,
                ultima_avaliacao_em or str(existente["ultima_avaliacao_em"] or ""),
                str(existente["finalizado_em"] or "") or (agora if marcar_finalizado else ""),
                usuario_limpo or str(existente["atualizado_por"] or ""),
                agora,
                int(paciente_id),
            ),
        )
    return crm_entry_por_paciente(conn, paciente_id)


def montar_ficha_paciente(conn: sqlite3.Connection, paciente_row: sqlite3.Row) -> FichaPacienteResposta:
    atualizar_status_recebiveis_automaticamente(conn)
    contratos_rows = conn.execute(
        "SELECT * FROM contratos WHERE paciente_id=? ORDER BY COALESCE(data_criacao, '') DESC, id DESC",
        (int(paciente_row["id"]),),
    ).fetchall()
    recebiveis_rows = carregar_recebiveis_paciente(conn, paciente_row)
    agendamentos_rows = carregar_agendamentos_paciente(conn, paciente_row)
    proximo = proximo_agendamento_paciente(agendamentos_rows)

    contratos = [
        ContratoResumo(
            id=int(row["id"]),
            valorTotal=formatar_moeda_br(row["valor_total"]),
            entrada=formatar_moeda_br(row["entrada"]),
            parcelas=int(row["parcelas"] or 1),
            primeiroVencimento=formatar_data_br_valor(row["primeiro_vencimento"]),
            dataCriacao=formatar_data_br_valor(row["data_criacao"]),
            formaPagamento=str(row["forma_pagamento"] or ""),
            status=str(row["status"] or "EM_ABERTO"),
            aprovadoPor=str(row["aprovado_por"] or ""),
            dataAprovacao=formatar_data_br_valor(row["data_aprovacao"]),
            procedimentos=carregar_procedimentos_contrato(conn, int(row["id"])),
        )
        for row in contratos_rows
    ]

    recebiveis = [mapear_recebivel(row) for row in recebiveis_rows]
    agendamentos = [mapear_agendamento(row) for row in agendamentos_rows]
    recibos = [mapear_recebivel(row) for row in recebiveis_rows if str(row["status"] or "") == "Pago"]

    return FichaPacienteResposta(
        paciente=mapear_paciente_detalhe(paciente_row),
        contratos=contratos,
        recebiveis=recebiveis,
        financeiro=resumo_financeiro_paciente(recebiveis_rows),
        agendamentos=agendamentos,
        proximoAgendamento=mapear_agendamento(proximo) if proximo else None,
        documentos=listar_documentos_paciente(paciente_row),
        exames=listar_exames_paciente(paciente_row),
        recibos=recibos,
        crm=mapear_crm_paciente_resumo(crm_entry_por_paciente(conn, int(paciente_row["id"]))),
    )


@app.get("/api/pacientes", response_model=list[PacienteResumo])
def listar_pacientes(
    q: str = Query("", description="Busca por nome, apelido, prontuario, CPF ou telefone"),
    limit: int = Query(50, ge=1, le=500),
):
    conn = conectar()
    try:
        rows = conn.execute("SELECT * FROM pacientes ORDER BY nome").fetchall()
        filtrados = filtrar_pacientes(rows, q)
        return [mapear_paciente_resumo(row) for row in filtrados[:limit]]
    finally:
        conn.close()


@app.get("/api/pacientes/recentes", response_model=list[PacienteResumo])
def listar_pacientes_recentes(limit: int = Query(6, ge=1, le=50)):
    conn = conectar()
    try:
        rows = conn.execute("SELECT * FROM pacientes ORDER BY id DESC LIMIT ?", (limit,)).fetchall()
        return [mapear_paciente_resumo(row) for row in rows]
    finally:
        conn.close()


@app.get("/api/pacientes/cep/{cep}")
def consultar_cep(cep: str):
    return buscar_endereco_por_cep(cep)


@app.get("/api/pacientes/{paciente_id}", response_model=PacienteDetalhe)
def detalhar_paciente(paciente_id: int):
    conn = conectar()
    try:
        row = carregar_paciente_por_id(conn, paciente_id)
        return mapear_paciente_detalhe(row)
    finally:
        conn.close()


@app.get("/api/pacientes/{paciente_id}/ficha", response_model=FichaPacienteResposta)
def ficha_paciente(paciente_id: int):
    conn = conectar()
    try:
        row = carregar_paciente_por_id(conn, paciente_id)
        return montar_ficha_paciente(conn, row)
    finally:
        conn.close()


@app.get("/api/crm", response_model=CrmPainelResposta)
def listar_crm():
    conn = conectar()
    try:
        pipeline_rows = conn.execute(
            """
            SELECT
                crm.*,
                COALESCE(p.nome, '') AS nome,
                COALESCE(p.prontuario, '') AS prontuario,
                COALESCE(p.telefone, '') AS telefone
            FROM crm_pacientes crm
            JOIN pacientes p ON p.id = crm.paciente_id
            ORDER BY
                CASE COALESCE(crm.etapa_funil, 'Novo lead')
                    WHEN 'Novo lead' THEN 1
                    WHEN 'Contato inicial' THEN 2
                    WHEN 'Tentando contato' THEN 3
                    WHEN 'Conversando' THEN 4
                    WHEN 'Agendou avaliação' THEN 5
                    WHEN 'Em negociação' THEN 6
                    WHEN 'Convertido' THEN 7
                    WHEN 'Perdido' THEN 8
                    ELSE 99
                END,
                lower(COALESCE(p.nome, '')),
                crm.id DESC
            """
        ).fetchall()
        pipeline = [mapear_crm_paciente_item(row) for row in pipeline_rows]
        finalizados = [item for item in pipeline if item.origemFinalizado]
        avaliacoes = listar_avaliacoes_crm(conn)
        return CrmPainelResposta(
            pipeline=pipeline,
            finalizados=finalizados,
            avaliacoes=avaliacoes,
        )
    finally:
        conn.close()


@app.post("/api/crm/pacientes/{paciente_id}/finalizado", response_model=CrmPacienteItemResposta)
def marcar_paciente_finalizado_crm(paciente_id: int, request: Request):
    conn = conectar()
    try:
        paciente = carregar_paciente_por_id(conn, paciente_id)
        crm = upsert_crm_origem(
            conn,
            paciente_id=int(paciente_id),
            marcar_finalizado=True,
            usuario=usuario_request(request),
        )
        conn.commit()
        row = conn.execute(
            """
            SELECT crm.*, p.nome, p.prontuario, p.telefone
            FROM crm_pacientes crm
            JOIN pacientes p ON p.id = crm.paciente_id
            WHERE crm.id=?
            LIMIT 1
            """,
            (int(crm["id"]),),
        ).fetchone()
    finally:
        conn.close()
    registrar_acao_usuario(
        usuario_request(request),
        acao="CRM",
        tipo="Paciente finalizado",
        info=str(paciente["nome"] or f"Paciente {paciente_id}"),
        metodo_http="POST",
        rota=f"/api/crm/pacientes/{paciente_id}/finalizado",
    )
    return mapear_crm_paciente_item(row)


@app.post("/api/crm/pacientes/{paciente_id}/avaliacao", response_model=CrmPacienteItemResposta)
def adicionar_paciente_avaliacao_crm(paciente_id: int, request: Request):
    conn = conectar()
    try:
        paciente = carregar_paciente_por_id(conn, paciente_id)
        avaliacao = next((item for item in listar_avaliacoes_crm(conn) if item.pacienteId == int(paciente_id)), None)
        ultima_avaliacao_iso = ""
        if avaliacao and str(avaliacao.dataAvaliacao or "").strip():
            data_avaliacao = parse_data_contrato(str(avaliacao.dataAvaliacao).split(" ")[0])
            ultima_avaliacao_iso = data_avaliacao.isoformat() if data_avaliacao else ""
        crm = upsert_crm_origem(
            conn,
            paciente_id=int(paciente_id),
            marcar_avaliacao=True,
            usuario=usuario_request(request),
            ultima_avaliacao_em=ultima_avaliacao_iso,
        )
        conn.commit()
        row = conn.execute(
            """
            SELECT crm.*, p.nome, p.prontuario, p.telefone
            FROM crm_pacientes crm
            JOIN pacientes p ON p.id = crm.paciente_id
            WHERE crm.id=?
            LIMIT 1
            """,
            (int(crm["id"]),),
        ).fetchone()
    finally:
        conn.close()
    registrar_acao_usuario(
        usuario_request(request),
        acao="CRM",
        tipo="Paciente avaliação",
        info=str(paciente["nome"] or f"Paciente {paciente_id}"),
        metodo_http="POST",
        rota=f"/api/crm/pacientes/{paciente_id}/avaliacao",
    )
    return mapear_crm_paciente_item(row)


@app.put("/api/crm/{crm_id}", response_model=CrmPacienteItemResposta)
def atualizar_item_crm(crm_id: int, payload: CrmAtualizacaoPayload, request: Request):
    conn = conectar()
    try:
        atual = conn.execute("SELECT * FROM crm_pacientes WHERE id=? LIMIT 1", (int(crm_id),)).fetchone()
        if atual is None:
            raise HTTPException(status_code=404, detail="Registro de CRM não encontrado.")
        etapa_funil = str(payload.etapa_funil or atual["etapa_funil"] or "Novo lead").strip() or "Novo lead"
        if etapa_funil not in CRM_ETAPAS_PADRAO:
            etapa_funil = "Novo lead"
        conn.execute(
            """
            UPDATE crm_pacientes
            SET etapa_funil=?, canal=?, campanha=?, conjunto_anuncio=?, anuncio=?,
                responsavel=?, proximo_contato=?, observacao=?, ultima_interacao=?,
                atualizado_por=?, atualizado_em=?
            WHERE id=?
            """,
            (
                etapa_funil,
                str(payload.canal or atual["canal"] or "Facebook").strip() or "Facebook",
                str(payload.campanha or "").strip(),
                str(payload.conjunto_anuncio or "").strip(),
                str(payload.anuncio or "").strip(),
                str(payload.responsavel or "").strip(),
                str(payload.proximo_contato or "").strip(),
                str(payload.observacao or "").strip(),
                str(payload.ultima_interacao or "").strip(),
                usuario_request(request),
                agora_str(),
                int(crm_id),
            ),
        )
        conn.commit()
        row = conn.execute(
            """
            SELECT crm.*, p.nome, p.prontuario, p.telefone
            FROM crm_pacientes crm
            JOIN pacientes p ON p.id = crm.paciente_id
            WHERE crm.id=?
            LIMIT 1
            """,
            (int(crm_id),),
        ).fetchone()
    finally:
        conn.close()
    registrar_acao_usuario(
        usuario_request(request),
        acao="CRM",
        tipo="Atualização do funil",
        info=str(row["nome"] or f"CRM {crm_id}"),
        metodo_http="PUT",
        rota=f"/api/crm/{crm_id}",
    )
    return mapear_crm_paciente_item(row)


@app.put("/api/pacientes/{paciente_id}/recebiveis/{recebivel_id}", response_model=RecebivelResumo)
def atualizar_recebivel_paciente(paciente_id: int, recebivel_id: int, payload: RecebivelAtualizacaoPayload):
    conn = conectar()
    try:
        carregar_paciente_por_id(conn, paciente_id)
        atual = carregar_recebivel_paciente(conn, paciente_id, recebivel_id)
        status = (payload.status or atual["status"] or "Aberto").strip() or "Aberto"
        data_pagamento = payload.data_pagamento.strip() if payload.data_pagamento else ""
        if normalizar_texto(status) == "pago" and not data_pagamento:
            data_pagamento = date.today().isoformat()
        if normalizar_texto(status) != "pago":
            data_pagamento = ""

        conn.execute(
            """
            UPDATE recebiveis
            SET paciente_nome=?, prontuario=?, vencimento=?, valor=?, forma_pagamento=?, status=?, data_pagamento=?, observacao=?
            WHERE id=? AND paciente_id=?
            """,
            (
                payload.paciente_nome.strip() or str(atual["paciente_nome"] or ""),
                payload.prontuario.strip() or str(atual["prontuario"] or ""),
                payload.vencimento.strip() or str(atual["vencimento"] or ""),
                float(payload.valor or 0),
                payload.forma_pagamento.strip() or str(atual["forma_pagamento"] or ""),
                status,
                data_pagamento,
                payload.observacao.strip(),
                recebivel_id,
                paciente_id,
            ),
        )
        conn.commit()
        atualizado = carregar_recebivel_paciente(conn, paciente_id, recebivel_id)
        return mapear_recebivel(atualizado)
    finally:
        conn.close()


@app.put("/api/financeiro/recebiveis/lote/{contrato_id}")
def atualizar_recebiveis_lote(contrato_id: int, payload: RecebivelLotePayload):
    conn = conectar()
    try:
        atualizar_recebiveis_lote_contrato(
            conn,
            contrato_id=contrato_id,
            paciente_nome=payload.paciente_nome,
            prontuario=payload.prontuario,
            forma_pagamento=payload.forma_pagamento,
            status=payload.status,
            observacao=payload.observacao,
            primeiro_vencimento=payload.primeiro_vencimento,
        )
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()


@app.post("/api/pacientes/{paciente_id}/recebiveis/{recebivel_id}/baixar", response_model=RecebivelResumo)
def baixar_recebivel_paciente(paciente_id: int, recebivel_id: int, payload: BaixaRecebivelPayload):
    conn = conectar()
    try:
        carregar_paciente_por_id(conn, paciente_id)
        atualizado = baixar_recebivel_no_caixa(
            conn,
            paciente_id=paciente_id,
            recebivel_id=recebivel_id,
            data_pagamento=payload.data_pagamento or date.today().isoformat(),
            forma_pagamento=payload.forma_pagamento,
            conta_caixa=payload.conta_caixa,
            desconto_valor=payload.desconto_valor,
            observacao=payload.observacao,
        )
        conn.commit()
        return mapear_recebivel(atualizado)
    finally:
        conn.close()


@app.get("/api/pacientes/{paciente_id}/odontograma", response_model=OdontogramaResposta)
@app.get("/paciente/{paciente_id}/odontograma", response_model=OdontogramaResposta)
def odontograma_paciente(paciente_id: int):
    conn = conectar()
    try:
        carregar_paciente_por_id(conn, paciente_id)
        return OdontogramaResposta(
            dentes_contratados=carregar_dentes_contratados(conn, paciente_id),
            elementos=carregar_elementos_odontograma(conn, paciente_id),
        )
    finally:
        conn.close()


@app.get("/api/financeiro/painel", response_model=FinanceiroPainelResposta)
def painel_financeiro():
    conn = conectar()
    try:
        recebiveis_rows = carregar_recebiveis_financeiro(conn)
        caixa_rows = carregar_caixa_financeiro(conn)
        contas_rows = carregar_contas_pagar_financeiro(conn)
        saldos_rows = carregar_saldos_conta_financeiro(conn)
        return FinanceiroPainelResposta(
            resumo=resumo_financeiro_global(recebiveis_rows),
            recebiveis=[mapear_recebivel(row) for row in recebiveis_rows],
            caixa=[mapear_movimento_caixa(row) for row in caixa_rows],
            contasPagar=[mapear_conta_pagar(row) for row in contas_rows],
            saldosConta=[mapear_saldo_conta(row) for row in saldos_rows],
        )
    finally:
        conn.close()


@app.get("/api/dashboard", response_model=DashboardPainelResposta)
def painel_dashboard():
    conn = conectar()
    try:
        return dados_dashboard(conn)
    finally:
        conn.close()


@app.get("/api/financeiro/metas", response_model=list[MetaMensalResumo])
def listar_metas_financeiras(ano: int = Query(default_factory=lambda: date.today().year)):
    conn = conectar()
    try:
        return carregar_metas_mensais(conn, ano)
    finally:
        conn.close()


@app.put("/api/financeiro/metas/{ano}/{mes}", response_model=MetaMensalResumo)
def atualizar_meta_financeira(ano: int, mes: int, payload: MetaMensalPayload):
    if mes < 1 or mes > 12:
        raise HTTPException(status_code=400, detail="Mes invalido.")
    conn = conectar()
    try:
        carregar_metas_mensais(conn, ano)
        conn.execute(
            """
            UPDATE metas_mensais
            SET meta=?, supermeta=?, hipermeta=?, data_atualizacao=?
            WHERE ano=? AND mes=?
            """,
            (
                float(payload.meta or 0),
                float(payload.supermeta or 0),
                float(payload.hipermeta or 0),
                agora_str(),
                int(ano),
                int(mes),
            ),
        )
        conn.commit()
        return obter_meta_mensal(conn, ano, mes)
    finally:
        conn.close()


@app.get("/api/financeiro/notas-fiscais", response_model=list[NotaFiscalEmitidaResumo])
def listar_notas_fiscais_emitidas():
    conn = conectar()
    try:
        return [mapear_nota_fiscal_emitida(row) for row in carregar_notas_fiscais_emitidas(conn)]
    finally:
        conn.close()


@app.post("/api/financeiro/notas-fiscais", response_model=NotaFiscalEmitidaResumo)
def criar_nota_fiscal_emitida(payload: NotaFiscalEmitidaPayload):
    conn = conectar()
    try:
        cursor = conn.execute(
            """
            INSERT INTO notas_fiscais_emitidas (
                competencia, data_emissao, data_recebimento, numero_nf, serie, cliente,
                descricao, conta_destino, valor_nf, valor_recebido, status, observacao,
                criado_em, atualizado_em
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.competencia.strip(),
                payload.data_emissao.strip(),
                payload.data_recebimento.strip(),
                payload.numero_nf.strip(),
                payload.serie.strip(),
                payload.cliente.strip(),
                payload.descricao.strip(),
                payload.conta_destino.strip(),
                float(payload.valor_nf or 0),
                float(payload.valor_recebido or 0),
                payload.status.strip() or "Pendente",
                payload.observacao.strip(),
                agora_str(),
                agora_str(),
            ),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM notas_fiscais_emitidas WHERE id=?", (int(cursor.lastrowid),)).fetchone()
        if row is None:
            raise HTTPException(status_code=500, detail="Falha ao salvar nota fiscal.")
        return mapear_nota_fiscal_emitida(row)
    finally:
        conn.close()


@app.put("/api/financeiro/notas-fiscais/{nota_id}", response_model=NotaFiscalEmitidaResumo)
def atualizar_nota_fiscal_emitida(nota_id: int, payload: NotaFiscalEmitidaPayload):
    conn = conectar()
    try:
        row = conn.execute("SELECT * FROM notas_fiscais_emitidas WHERE id=?", (nota_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="Nota fiscal nao encontrada.")
        conn.execute(
            """
            UPDATE notas_fiscais_emitidas
            SET competencia=?, data_emissao=?, data_recebimento=?, numero_nf=?, serie=?, cliente=?,
                descricao=?, conta_destino=?, valor_nf=?, valor_recebido=?, status=?, observacao=?, atualizado_em=?
            WHERE id=?
            """,
            (
                payload.competencia.strip(),
                payload.data_emissao.strip(),
                payload.data_recebimento.strip(),
                payload.numero_nf.strip(),
                payload.serie.strip(),
                payload.cliente.strip(),
                payload.descricao.strip(),
                payload.conta_destino.strip(),
                float(payload.valor_nf or 0),
                float(payload.valor_recebido or 0),
                payload.status.strip() or "Pendente",
                payload.observacao.strip(),
                agora_str(),
                int(nota_id),
            ),
        )
        conn.commit()
        atualizado = conn.execute("SELECT * FROM notas_fiscais_emitidas WHERE id=?", (nota_id,)).fetchone()
        if atualizado is None:
            raise HTTPException(status_code=500, detail="Falha ao atualizar nota fiscal.")
        return mapear_nota_fiscal_emitida(atualizado)
    finally:
        conn.close()


@app.post("/api/financeiro/contas-pagar", response_model=ContaPagarResumo)
def criar_conta_pagar(payload: ContaPagarPayload):
    conn = conectar()
    try:
        status = payload.status.strip() or ("Pago" if payload.pago_em.strip() else "A vencer")
        cursor = conn.execute(
            """
            INSERT INTO contas_pagar (
                data_vencimento, descricao, fornecedor, categoria, valor, pago, valor_pago, status, observacao, data_criacao
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.vencimento.strip(),
                payload.descricao.strip(),
                payload.fornecedor.strip(),
                payload.categoria.strip(),
                float(payload.valor or 0),
                payload.pago_em.strip(),
                float(payload.valor_pago or 0),
                status,
                payload.observacao.strip(),
                agora_str(),
            ),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM contas_pagar WHERE id=?", (int(cursor.lastrowid),)).fetchone()
        if row is None:
            raise HTTPException(status_code=500, detail="Falha ao criar conta a pagar.")
        return mapear_conta_pagar(row)
    finally:
        conn.close()


@app.post("/api/financeiro/caixa", response_model=MovimentoCaixaResumo)
def criar_movimento_caixa(payload: MovimentoCaixaPayload):
    conn = conectar()
    try:
        data_movimento = payload.data_movimento.strip() or date.today().isoformat()
        movimento_id = registrar_movimento_caixa(
            conn,
            origem=payload.origem,
            descricao=payload.descricao,
            valor=payload.valor,
            tipo=payload.tipo or "Entrada",
            data_movimento=data_movimento,
            prontuario=payload.prontuario,
            forma_pagamento=payload.forma_pagamento,
            conta_caixa=payload.conta_caixa,
            observacao=payload.observacao,
            contrato_id=payload.contrato_id,
            recebivel_id=payload.recebivel_id,
        )
        conn.commit()
        row = conn.execute("SELECT * FROM financeiro WHERE id=?", (movimento_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=500, detail="Falha ao registrar movimento de caixa.")
        return mapear_movimento_caixa(row)
    finally:
        conn.close()


@app.put("/api/financeiro/caixa/{movimento_id}", response_model=MovimentoCaixaResumo)
def atualizar_movimento_caixa(movimento_id: int, payload: MovimentoCaixaAtualizacaoPayload):
    conn = conectar()
    try:
        row = conn.execute("SELECT * FROM financeiro WHERE id=?", (movimento_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="Movimento de caixa nao encontrado.")
        conn.execute(
            """
            UPDATE financeiro
            SET origem=?, descricao=?, valor=?, data=?, tipo=?, prontuario=?, forma_pagamento=?, conta_caixa=?, observacao=?
            WHERE id=?
            """,
            (
                payload.origem.strip() or str(row["origem"] or ""),
                payload.descricao.strip() or str(row["descricao"] or ""),
                float(payload.valor or row["valor"] or 0),
                payload.data_movimento.strip() or str(row["data"] or ""),
                payload.tipo.strip() or str(row["tipo"] or ""),
                payload.prontuario.strip() or str(row["prontuario"] or ""),
                payload.forma_pagamento.strip() or str(row["forma_pagamento"] or ""),
                payload.conta_caixa.strip() or str(row["conta_caixa"] or ""),
                payload.observacao.strip() or str(row["observacao"] or ""),
                movimento_id,
            ),
        )
        conn.commit()
        atualizado = conn.execute("SELECT * FROM financeiro WHERE id=?", (movimento_id,)).fetchone()
        if atualizado is None:
            raise HTTPException(status_code=500, detail="Falha ao atualizar movimento.")
        return mapear_movimento_caixa(atualizado)
    finally:
        conn.close()


@app.delete("/api/financeiro/caixa/{movimento_id}")
def excluir_movimento_caixa(movimento_id: int):
    conn = conectar()
    try:
        row = conn.execute("SELECT * FROM financeiro WHERE id=?", (movimento_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="Movimento de caixa nao encontrado.")
        recebivel_id = int(row["recebivel_id"]) if row["recebivel_id"] is not None else None
        conn.execute("DELETE FROM financeiro WHERE id=?", (movimento_id,))
        if recebivel_id is not None:
            conn.execute(
                """
                UPDATE recebiveis
                SET status='Aberto', data_pagamento=NULL
                WHERE id=?
                """,
                (recebivel_id,),
            )
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()


@app.post("/api/financeiro/saldos-conta", response_model=SaldoContaResumo)
def criar_saldo_conta(payload: SaldoContaPayload):
    conn = conectar()
    try:
        saldo_id = registrar_saldo_conta(
            conn,
            data_saldo=payload.data,
            conta=payload.conta,
            saldo=payload.saldo,
            observacao=payload.observacao,
        )
        conn.commit()
        row = conn.execute("SELECT * FROM saldos_conta WHERE id=?", (saldo_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=500, detail="Falha ao registrar saldo.")
        return mapear_saldo_conta(row)
    finally:
        conn.close()


@app.get("/api/financeiro/caixa/export.xlsx")
def exportar_caixa_excel():
    conn = conectar()
    try:
        caixa_rows = carregar_caixa_financeiro(conn)
        conteudo = caixa_diario_para_excel_bytes(conn, caixa_rows)
        if not conteudo:
            raise HTTPException(status_code=400, detail="Nao ha dados de caixa para exportar.")
        nome = f"caixa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return Response(
            content=conteudo,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{nome}"'},
        )
    finally:
        conn.close()


@app.get("/api/financeiro/recibos", response_model=list[ReciboManualResumo])
def listar_recibos_manuais():
    conn = conectar()
    try:
        rows = conn.execute("SELECT * FROM recibos_manuais ORDER BY date(COALESCE(data_pagamento, '')) DESC, id DESC").fetchall()
        return [mapear_recibo_manual(row) for row in rows]
    finally:
        conn.close()


@app.post("/api/financeiro/recibos", response_model=ReciboManualResumo)
def criar_recibo_manual(payload: ReciboManualPayload):
    conn = conectar()
    try:
        cursor = conn.execute(
            """
            INSERT INTO recibos_manuais (
                valor, pagador, recebedor, data_pagamento, referente, observacao, cidade, criado_em
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                float(payload.valor or 0),
                payload.pagador.strip(),
                payload.recebedor.strip(),
                payload.data_pagamento.strip(),
                payload.referente.strip(),
                payload.observacao.strip(),
                payload.cidade.strip(),
                agora_str(),
            ),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM recibos_manuais WHERE id=?", (cursor.lastrowid,)).fetchone()
        if row is None:
            raise HTTPException(status_code=500, detail="Falha ao salvar recibo.")
        return mapear_recibo_manual(row)
    finally:
        conn.close()


@app.get("/api/financeiro/recibos/{recibo_id}")
def abrir_recibo_manual(recibo_id: int):
    conn = conectar()
    try:
        row = conn.execute("SELECT * FROM recibos_manuais WHERE id=?", (recibo_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="Recibo nao encontrado.")
        html = gerar_html_recibo_manual(row)
        return Response(content=html, media_type="text/html; charset=utf-8")
    finally:
        conn.close()


@app.get("/api/sistema/export.xlsx")
def exportar_sistema_excel():
    conn = conectar()
    try:
        conteudo = exportar_todas_tabelas_excel_bytes(conn)
        nome = f"sistema_completo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return Response(
            content=conteudo,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{nome}"'},
        )
    finally:
        conn.close()


@app.put("/api/financeiro/contas-pagar/{conta_id}", response_model=ContaPagarResumo)
def atualizar_conta_pagar(conta_id: int, payload: ContaPagarPayload):
    conn = conectar()
    try:
        row = conn.execute("SELECT * FROM contas_pagar WHERE id=?", (conta_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="Conta a pagar nao encontrada.")
        status = payload.status.strip() or ("Pago" if payload.pago_em.strip() else str(row["status"] or "A vencer"))
        conn.execute(
            """
            UPDATE contas_pagar
            SET data_vencimento=?, descricao=?, fornecedor=?, categoria=?, valor=?, pago=?, valor_pago=?, status=?, observacao=?
            WHERE id=?
            """,
            (
                payload.vencimento.strip(),
                payload.descricao.strip(),
                payload.fornecedor.strip(),
                payload.categoria.strip(),
                float(payload.valor or 0),
                payload.pago_em.strip(),
                float(payload.valor_pago or 0),
                status,
                payload.observacao.strip(),
                conta_id,
            ),
        )
        conn.commit()
        atualizado = conn.execute("SELECT * FROM contas_pagar WHERE id=?", (conta_id,)).fetchone()
        if atualizado is None:
            raise HTTPException(status_code=500, detail="Falha ao atualizar conta a pagar.")
        return mapear_conta_pagar(atualizado)
    finally:
        conn.close()


@app.post("/api/financeiro/contas-pagar/{conta_id}/baixar", response_model=ContaPagarResumo)
def baixar_conta_pagar(conta_id: int, payload: BaixaRecebivelPayload):
    conn = conectar()
    try:
        row = conn.execute("SELECT * FROM contas_pagar WHERE id=?", (conta_id,)).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="Conta a pagar nao encontrada.")
        if normalizar_texto(row["status"]) == "pago":
            raise HTTPException(status_code=400, detail="Esta conta a pagar ja esta paga.")

        data_ref = parse_data_contrato(payload.data_pagamento) or date.today()
        data_iso = data_ref.isoformat()
        data_br = formatar_data_br(data_ref)
        valor_pago = float(row["valor"] or 0)
        forma = payload.forma_pagamento.strip() or "PIX"
        conta_caixa = payload.conta_caixa.strip() or forma

        registrar_movimento_caixa(
            conn,
            origem=str(row["fornecedor"] or "Pagamento"),
            descricao=str(row["descricao"] or "Conta a pagar"),
            valor=valor_pago,
            tipo="Saida",
            data_movimento=data_iso,
            forma_pagamento=forma,
            conta_caixa=conta_caixa,
            observacao=payload.observacao.strip(),
        )

        observacao_atual = str(row["observacao"] or "").strip()
        observacao_final = " | ".join(parte for parte in [observacao_atual, payload.observacao.strip()] if parte)
        conn.execute(
            """
            UPDATE contas_pagar
            SET pago=?, valor_pago=?, status='Pago', observacao=?
            WHERE id=?
            """,
            (data_br, valor_pago, observacao_final, conta_id),
        )
        conn.commit()
        atualizado = conn.execute("SELECT * FROM contas_pagar WHERE id=?", (conta_id,)).fetchone()
        if atualizado is None:
            raise HTTPException(status_code=500, detail="Falha ao baixar conta a pagar.")
        return mapear_conta_pagar(atualizado)
    finally:
        conn.close()


@app.get("/api/pacientes/{paciente_id}/documentos/{nome_arquivo}")
def abrir_documento_paciente(paciente_id: int, nome_arquivo: str, download: int = Query(default=0)):
    conn = conectar()
    try:
        paciente_row = carregar_paciente_por_id(conn, paciente_id)
        caminho = resolver_documento_contrato_atual(conn, paciente_row, nome_arquivo) or buscar_documento_paciente(paciente_row, nome_arquivo)
        media_type = None
        extensao = os.path.splitext(caminho)[1].lower()
        if extensao == ".html":
            media_type = "text/html; charset=utf-8"
        elif extensao == ".pdf":
            media_type = "application/pdf"
        elif extensao == ".docx":
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return FileResponse(
            path=caminho,
            filename=os.path.basename(caminho),
            media_type=media_type,
            content_disposition_type="attachment" if int(download or 0) else "inline",
        )
    finally:
        conn.close()


@app.get("/api/pacientes/{paciente_id}/recibos/{recebivel_id}")
def abrir_recibo_paciente(paciente_id: int, recebivel_id: int):
    conn = conectar()
    try:
        carregar_paciente_por_id(conn, paciente_id)
        recebivel = carregar_recebivel_paciente(conn, paciente_id, recebivel_id)
        if normalizar_texto(recebivel["status"]) != "pago":
            raise HTTPException(status_code=400, detail="O recibo so pode ser emitido para recebiveis pagos.")
        forma = normalizar_texto(str(recebivel["forma_pagamento"] or ""))
        if forma != "dinheiro":
            raise HTTPException(status_code=400, detail="O recibo desta tela esta disponivel apenas para pagamentos em especie.")
        html = gerar_html_recibo_especie(recebivel)
        return Response(content=html, media_type="text/html; charset=utf-8")
    finally:
        conn.close()


@app.get("/api/pacientes/{paciente_id}/exames/{nome_arquivo}")
def abrir_exame_paciente(paciente_id: int, nome_arquivo: str, download: int = Query(default=0)):
    conn = conectar()
    try:
        paciente_row = carregar_paciente_por_id(conn, paciente_id)
        caminho = buscar_exame_paciente(paciente_row, nome_arquivo)
        return FileResponse(
            path=caminho,
            filename=os.path.basename(caminho),
            content_disposition_type="attachment" if int(download or 0) else "inline",
        )
    finally:
        conn.close()


@app.get("/api/pacientes/{paciente_id}/foto")
def abrir_foto_paciente(paciente_id: int):
    conn = conectar()
    try:
        paciente_row = carregar_paciente_por_id(conn, paciente_id)
        caminho = buscar_foto_paciente(paciente_row)
        extensao = os.path.splitext(caminho)[1].lower()
        media_type = "image/jpeg"
        if extensao == ".png":
            media_type = "image/png"
        elif extensao == ".webp":
            media_type = "image/webp"
        return FileResponse(
            path=caminho,
            filename=os.path.basename(caminho),
            media_type=media_type,
            content_disposition_type="inline",
        )
    finally:
        conn.close()


@app.post("/api/pacientes/{paciente_id}/foto", response_model=PacienteDetalhe)
async def enviar_foto_paciente(paciente_id: int, request: Request):
    conn = conectar()
    try:
        paciente_row = carregar_paciente_por_id(conn, paciente_id)
        nome_arquivo = request.headers.get("x-filename", "foto.jpg")
        extensao = os.path.splitext(nome_arquivo)[1].lower()
        if extensao not in {".png", ".jpg", ".jpeg", ".webp"}:
            raise HTTPException(status_code=400, detail="Formato de foto invalido.")

        pasta = pasta_fotos_paciente(paciente_row)
        caminho = os.path.join(pasta, f"foto{extensao}")
        conteudo = await request.body()
        if not conteudo:
            raise HTTPException(status_code=400, detail="Arquivo de foto vazio.")
        with open(caminho, "wb") as destino:
            destino.write(conteudo)

        conn.execute("UPDATE pacientes SET foto_path=? WHERE id=?", (caminho, paciente_id))
        conn.commit()
        atualizado = carregar_paciente_por_id(conn, paciente_id)
        return mapear_paciente_detalhe(atualizado)
    finally:
        conn.close()


@app.post("/api/pacientes", response_model=PacienteDetalhe)
def criar_paciente(payload: PacientePayload):
    conn = conectar()
    try:
        prontuario = formatar_prontuario_valor(payload.prontuario) or proximo_prontuario(conn)
        erros = validar_dados_paciente(
            payload.nome,
            prontuario,
            payload.cpf,
            payload.menor_idade,
            payload.responsavel,
            payload.cpf_responsavel,
        )
        if erros:
            raise HTTPException(status_code=400, detail=erros)
        erros_duplicidade = validar_duplicidade_paciente(conn, prontuario, payload.cpf)
        if erros_duplicidade:
            raise HTTPException(status_code=400, detail=erros_duplicidade)

        cursor = conn.execute(
            """
            INSERT INTO pacientes (
                nome, apelido, sexo, prontuario, cpf, rg, data_nascimento, telefone, email, cep,
                endereco, complemento, numero, bairro, cidade, estado, estado_civil, profissao, origem, observacoes,
                menor_idade, responsavel, cpf_responsavel
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.nome.strip(),
                payload.apelido.strip(),
                payload.sexo.strip(),
                prontuario,
                limpar_cpf(payload.cpf),
                payload.rg.strip(),
                payload.data_nascimento.strip(),
                payload.telefone.strip(),
                payload.email.strip(),
                payload.cep.strip(),
                payload.endereco.strip(),
                payload.complemento.strip(),
                payload.numero.strip(),
                payload.bairro.strip(),
                payload.cidade.strip(),
                payload.estado.strip(),
                payload.estado_civil.strip(),
                payload.profissao.strip(),
                payload.origem.strip(),
                payload.observacoes.strip(),
                int(bool(payload.menor_idade)),
                payload.responsavel.strip(),
                limpar_cpf(payload.cpf_responsavel),
            ),
        )
        conn.commit()
        row = carregar_paciente_por_id(conn, int(cursor.lastrowid))
        return mapear_paciente_detalhe(row)
    finally:
        conn.close()


@app.put("/api/pacientes/{paciente_id}", response_model=PacienteDetalhe)
def atualizar_paciente(paciente_id: int, payload: PacientePayload):
    conn = conectar()
    try:
        existente = carregar_paciente_por_id(conn, paciente_id)
        prontuario = formatar_prontuario_valor(payload.prontuario) or formatar_prontuario_valor(existente["prontuario"])
        erros = validar_dados_paciente(
            payload.nome,
            prontuario,
            payload.cpf,
            payload.menor_idade,
            payload.responsavel,
            payload.cpf_responsavel,
        )
        if erros:
            raise HTTPException(status_code=400, detail=erros)
        erros_duplicidade = validar_duplicidade_paciente(conn, prontuario, payload.cpf, paciente_id_atual=paciente_id)
        if erros_duplicidade:
            raise HTTPException(status_code=400, detail=erros_duplicidade)

        conn.execute(
            """
            UPDATE pacientes
            SET nome=?, apelido=?, sexo=?, prontuario=?, cpf=?, rg=?, data_nascimento=?, telefone=?, email=?, cep=?,
                endereco=?, complemento=?, numero=?, bairro=?, cidade=?, estado=?, estado_civil=?, profissao=?, origem=?, observacoes=?, menor_idade=?,
                responsavel=?, cpf_responsavel=?
            WHERE id=?
            """,
            (
                payload.nome.strip(),
                payload.apelido.strip(),
                payload.sexo.strip(),
                prontuario,
                limpar_cpf(payload.cpf),
                payload.rg.strip(),
                payload.data_nascimento.strip(),
                payload.telefone.strip(),
                payload.email.strip(),
                payload.cep.strip(),
                payload.endereco.strip(),
                payload.complemento.strip(),
                payload.numero.strip(),
                payload.bairro.strip(),
                payload.cidade.strip(),
                payload.estado.strip(),
                payload.estado_civil.strip(),
                payload.profissao.strip(),
                payload.origem.strip(),
                payload.observacoes.strip(),
                int(bool(payload.menor_idade)),
                payload.responsavel.strip() if payload.menor_idade else "",
                limpar_cpf(payload.cpf_responsavel) if payload.menor_idade else "",
                paciente_id,
            ),
        )
        conn.commit()
        row = carregar_paciente_por_id(conn, paciente_id)
        return mapear_paciente_detalhe(row)
    finally:
        conn.close()


@app.post("/api/pacientes/{paciente_id}/orcamentos", response_model=OrcamentoCriadoResposta)
def criar_orcamento_paciente(paciente_id: int, payload: OrcamentoPacientePayload):
    conn = conectar()
    try:
        carregar_paciente_por_id(conn, paciente_id)
        contrato_id = salvar_orcamento_paciente_com_pagamento(conn, paciente_id, payload)
        conn.commit()
        return OrcamentoCriadoResposta(contrato_id=contrato_id)
    finally:
        conn.close()

@app.get("/api/pacientes/{paciente_id}/orcamentos/{contrato_id}", response_model=OrcamentoDetalheResposta)
def detalhar_orcamento_paciente(paciente_id: int, contrato_id: int):
    conn = conectar()
    try:
        carregar_paciente_por_id(conn, paciente_id)
        return carregar_orcamento_detalhe(conn, paciente_id, contrato_id)
    finally:
        conn.close()


@app.put("/api/pacientes/{paciente_id}/orcamentos/{contrato_id}", response_model=OrcamentoCriadoResposta)
def atualizar_orcamento_paciente(paciente_id: int, contrato_id: int, payload: OrcamentoPacientePayload):
    conn = conectar()
    try:
        carregar_paciente_por_id(conn, paciente_id)
        contrato_id_salvo = salvar_orcamento_paciente_com_pagamento(conn, paciente_id, payload, contrato_id=contrato_id)
        conn.commit()
        return OrcamentoCriadoResposta(contrato_id=contrato_id_salvo)
    finally:
        conn.close()


@app.put("/api/pacientes/{paciente_id}/orcamentos/{contrato_id}/status")
def alterar_status_orcamento_paciente(paciente_id: int, contrato_id: int, payload: OrcamentoStatusPayload):
    conn = conectar()
    try:
        paciente = carregar_paciente_por_id(conn, paciente_id)
        contrato = conn.execute("SELECT * FROM contratos WHERE id=? AND paciente_id=? LIMIT 1", (contrato_id, paciente_id)).fetchone()
        if contrato is None:
            raise HTTPException(status_code=404, detail="Orcamento nao encontrado.")
        status = normalizar_texto(payload.status).upper().replace(" ", "_")
        if status not in {"APROVADO", "EM_ABERTO"}:
            raise HTTPException(status_code=400, detail="Status de orcamento invalido.")

        if status == "APROVADO":
            conn.execute(
                "UPDATE contratos SET status='APROVADO', aprovado_por=?, data_aprovacao=? WHERE id=? AND paciente_id=?",
                (payload.aprovado_por.strip() or "JULIANA", agora_str(), contrato_id, paciente_id),
            )
            conn.execute(
                "UPDATE procedimentos_dente SET status='CONTRATADO' WHERE contrato_id=? AND paciente_id=?",
                (contrato_id, paciente_id),
            )
            contrato_atualizado = conn.execute("SELECT * FROM contratos WHERE id=? AND paciente_id=? LIMIT 1", (contrato_id, paciente_id)).fetchone()
            if contrato_atualizado is not None:
                sincronizar_recebiveis_contrato(conn, paciente, contrato_atualizado, contrato_id)
                sincronizar_financeiro_contrato(conn, paciente, contrato_atualizado, contrato_id)
                try:
                    gerar_documento_contrato(conn, paciente, contrato_atualizado, contrato_id)
                except Exception as exc:
                    print(f"[orcamento] falha ao gerar documento do contrato {contrato_id}: {exc}", flush=True)
        else:
            conn.execute(
                "UPDATE contratos SET status='EM_ABERTO', aprovado_por='', data_aprovacao=NULL WHERE id=? AND paciente_id=?",
                (contrato_id, paciente_id),
            )
            conn.execute(
                "UPDATE procedimentos_dente SET status='ORCAMENTO' WHERE contrato_id=? AND paciente_id=?",
                (contrato_id, paciente_id),
            )
            conn.execute("DELETE FROM recebiveis WHERE contrato_id=?", (contrato_id,))
            conn.execute("DELETE FROM financeiro WHERE contrato_id=?", (contrato_id,))
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()
