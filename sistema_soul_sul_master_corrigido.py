import os
import re
import calendar
import sqlite3
import hashlib
import hmac
import secrets
import base64
import unicodedata
from difflib import SequenceMatcher
from io import BytesIO
from datetime import date, datetime, timedelta
from urllib.parse import quote

import pandas as pd
import requests
import streamlit as st
import altair as alt
from docx import Document
from lxml import etree
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    REPORTLAB_DISPONIVEL = True
except ImportError:
    REPORTLAB_DISPONIVEL = False

try:
    import openpyxl  # noqa: F401
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_DISPONIVEL = True
except ImportError:
    OPENPYXL_DISPONIVEL = False

try:
    from pypdf import PdfReader

    PDF_LEITURA_DISPONIVEL = True
except ImportError:
    try:
        from PyPDF2 import PdfReader

        PDF_LEITURA_DISPONIVEL = True
    except ImportError:
        PDF_LEITURA_DISPONIVEL = False


st.set_page_config(page_title="Soul Sul Sistema", layout="wide")

DB_PATH = "clinica.db"
DOCS_DIR = "documentos"
TEMPLATE_PATH = "modelo_documento.docx"
ASSETS_DIR = "assets"
LOGO_PATH = os.path.join(ASSETS_DIR, "logo.png")
FORMAS_PAGAMENTO = ["Pix", "Dinheiro", "Debito", "Credito", "Boleto"]
STATUS_CONTAS_PAGAR = ["A vencer", "Pago", "Atrasado", "Suspenso", "Cancelado"]
STATUS_AGENDAMENTO = ["Agendado", "Confirmado", "Em espera", "Em atendimento", "Atendido", "Atrasado", "Faltou", "Cancelado"]
CATEGORIAS_CONTAS_PAGAR = [
    "Funcionarios",
    "Impostos",
    "Custo fixo",
    "Materia-prima",
    "Laboratorio",
    "Aluguel",
    "Marketing",
    "Comissoes",
    "Fornecedores",
    "Servicos",
    "Equipamentos",
    "Manutencao",
    "Tarifas bancarias",
    "Outros",
]
DIAS_SEMANA_PT = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sab", "Dom"]
CORES_PROFISSIONAIS_PADRAO = [
    "#fde68a", "#bfdbfe", "#fecaca", "#ddd6fe", "#bbf7d0",
    "#fbcfe8", "#fdba74", "#a7f3d0", "#c7d2fe", "#f9a8d4",
]
DESCRICAO_CONTRATO = "Contrato odontologico"
USUARIO_ADMIN_INICIAL = "admin"
SENHA_ADMIN_INICIAL = "admin123"
MODULOS_SISTEMA = {
    "Dashboard": "acesso_dashboard",
    "Pacientes": "acesso_pacientes",
    "Editar Paciente": "acesso_pacientes",
    "Contratos": "acesso_contratos",
    "Editar Contrato": "acesso_contratos",
    "Importacoes": "acesso_contratos",
    "Financeiro": "acesso_financeiro",
    "Agenda": "acesso_dashboard",
    "Usuarios": "acesso_usuarios",
}
ORDEM_MENU_SIDEBAR = [
    "Dashboard",
    "Pacientes",
    "Editar Paciente",
    "Contratos",
    "Editar Contrato",
    "Importacoes",
    "Financeiro",
    "Agenda",
    "Usuarios",
]
ROTULOS_MENU = {
    "Dashboard": "Dashboard",
    "Pacientes": "Pacientes",
    "Editar Paciente": "Editar Paciente",
    "Contratos": "Contratos",
    "Editar Contrato": "Editar Contrato",
    "Importacoes": "Importações",
    "Financeiro": "Financeiro",
    "Agenda": "Agenda",
    "Usuarios": "Usuários",
}
SVG_ICONES_MENU = {
    "Dashboard": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="8" height="8" rx="2"></rect><rect x="13" y="3" width="8" height="5" rx="2"></rect><rect x="13" y="10" width="8" height="11" rx="2"></rect><rect x="3" y="13" width="8" height="8" rx="2"></rect></svg>""",
    "Pacientes": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="8" r="4"></circle><path d="M5.5 20a6.5 6.5 0 0 1 13 0"></path></svg>""",
    "Editar Paciente": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><path d="M12 20h9"></path><path d="M16.5 3.5a2.1 2.1 0 0 1 3 3L7 19l-4 1 1-4 12.5-12.5z"></path></svg>""",
    "Contratos": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><path d="M14 3H7a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h10a2 2 0 0 0 2-2V8z"></path><path d="M14 3v5h5"></path><path d="M9 13h6"></path><path d="M9 17h6"></path></svg>""",
    "Editar Contrato": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><path d="M14 3H7a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h10a2 2 0 0 0 2-2V8z"></path><path d="M14 3v5h5"></path><path d="M9 13h3"></path><path d="M8 18l5.3-5.3a1.8 1.8 0 0 1 2.5 0l.5.5a1.8 1.8 0 0 1 0 2.5L11 21H8z"></path></svg>""",
    "Importacoes": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><path d="M12 3v12"></path><path d="M7 10l5 5 5-5"></path><path d="M5 21h14"></path></svg>""",
    "Financeiro": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><path d="M12 2v20"></path><path d="M17 6.5a4 4 0 0 0-4-2.5H10a3 3 0 0 0 0 6h4a3 3 0 0 1 0 6h-3a4 4 0 0 1-4-2.5"></path></svg>""",
    "Usuarios": """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.9" stroke-linecap="round" stroke-linejoin="round"><path d="M12 2v3"></path><path d="M12 19v3"></path><path d="M4.93 4.93l2.12 2.12"></path><path d="M16.95 16.95l2.12 2.12"></path><path d="M2 12h3"></path><path d="M19 12h3"></path><path d="M4.93 19.07l2.12-2.12"></path><path d="M16.95 7.05l2.12-2.12"></path><circle cx="12" cy="12" r="4"></circle></svg>""",
}
ICONES_MENU = {
    "Dashboard": "⌂",
    "Pacientes": "◉",
    "Editar Paciente": "✎",
    "Contratos": "▣",
    "Editar Contrato": "✐",
    "Importacoes": "⇪",
    "Financeiro": "$",
    "Usuarios": "⚙",
}


def aplicar_tema_visual():
    st.markdown(
        """
        <style>
        :root {
            --ss-bg: #f7f6f3;
            --ss-surface: #ffffff;
            --ss-sidebar: linear-gradient(180deg, #111111 0%, #1b1b1b 100%);
            --ss-border: #e5dfd3;
            --ss-shadow: 0 12px 34px rgba(15, 15, 15, 0.08);
            --ss-text: #222222;
            --ss-muted: #6f6a60;
            --ss-accent: #c5a77a;
            --ss-accent-soft: #f5ecdf;
        }

        .stApp {
            background:
                radial-gradient(circle at top left, rgba(197, 167, 122, 0.12), transparent 24%),
                radial-gradient(circle at top right, rgba(17, 17, 17, 0.05), transparent 18%),
                var(--ss-bg);
        }

        @keyframes ssFadeLift {
            from {
                opacity: 0;
                transform: translateY(10px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }

        @keyframes ssSoftPulse {
            0% { box-shadow: 0 0 0 0 rgba(197,167,122,0.22); }
            100% { box-shadow: 0 0 0 10px rgba(197,167,122,0.0); }
        }

        section[data-testid="stSidebar"] {
            background: var(--ss-sidebar);
            border-right: 1px solid rgba(255,255,255,0.08);
        }

        section[data-testid="stSidebar"] > div {
            padding-top: 1.2rem;
        }

        section[data-testid="stSidebar"] * {
            color: #f4f7f5;
        }

        .ss-brand {
            background: linear-gradient(180deg, rgba(255,255,255,0.12), rgba(255,255,255,0.05));
            border: 1px solid rgba(255,255,255,0.10);
            box-shadow: 0 20px 42px rgba(0,0,0,0.26);
            border-radius: 26px;
            padding: 18px 18px 20px 18px;
            margin: 0 0 22px 0;
            position: relative;
            overflow: hidden;
        }

        .ss-brand::after {
            content: "";
            position: absolute;
            inset: auto -10% -45px -10%;
            height: 70px;
            background: radial-gradient(circle, rgba(197,167,122,0.22), transparent 72%);
            pointer-events: none;
        }

        .ss-brand-logo {
            width: 122px;
            max-width: 140px;
            min-width: 100px;
            margin: 0 0 14px 0;
            display: block;
            overflow: hidden;
        }

        .ss-brand-logo img {
            width: 100%;
            height: auto;
            display: block;
        }

        .ss-brand-copy {
            display: flex;
            flex-direction: column;
            justify-content: flex-start;
            align-items: flex-start;
        }

        .ss-brand-title {
            font-size: 1.7rem;
            font-weight: 700;
            letter-spacing: 0.02em;
            margin: 0;
        }

        .ss-brand-subtitle {
            color: rgba(244,247,245,0.78);
            font-size: 0.96rem;
            margin-top: 3px;
        }

        .ss-menu-title {
            font-size: 1.15rem;
            font-weight: 700;
            letter-spacing: 0.02em;
            margin: 0 0 0.75rem 0.2rem;
            color: #f3efe8;
        }

        div[data-testid="stRadio"] > div {
            gap: 0.65rem;
        }

        div[data-testid="stRadio"] label {
            position: relative;
            display: flex;
            align-items: center;
            min-height: 56px;
            padding: 0.95rem 1rem 0.95rem 3.25rem;
            border-radius: 18px;
            border: 1px solid transparent;
            background: transparent;
            transition: all 0.18s ease;
        }

        div[data-testid="stRadio"] label > div:first-child {
            display: none;
        }

        div[data-testid="stRadio"] label p {
            font-size: 1rem;
            line-height: 1.1;
            color: rgba(244,247,245,0.88) !important;
            font-weight: 500;
        }

        div[data-testid="stRadio"] label:hover {
            background: rgba(255,255,255,0.08);
            border-color: rgba(255,255,255,0.10);
            transform: translateX(2px);
        }

        div[data-testid="stRadio"] label:has(input:checked) {
            background: linear-gradient(180deg, #f5ecdf 0%, #ead9bd 100%);
            border-color: rgba(197,167,122,0.34);
            box-shadow: 0 14px 24px rgba(0,0,0,0.22);
            transform: none;
        }

        div[data-testid="stRadio"] label:has(input:checked) p {
            color: #181818 !important;
            font-weight: 700;
        }

        div[data-testid="stRadio"] label::before {
            content: "";
            position: absolute;
            left: 1rem;
            top: 50%;
            transform: translateY(-50%);
            width: 20px;
            height: 20px;
            background: currentColor;
            color: rgba(244,247,245,0.88);
            -webkit-mask-repeat: no-repeat;
            -webkit-mask-position: center;
            -webkit-mask-size: contain;
            mask-repeat: no-repeat;
            mask-position: center;
            mask-size: contain;
        }

        div[data-testid="stRadio"] label:has(input:checked)::before {
            color: #181818;
        }

        div[data-testid="stRadio"] label:nth-of-type(1)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><rect x='3' y='3' width='8' height='8' rx='2'/><rect x='13' y='3' width='8' height='5' rx='2'/><rect x='13' y='10' width='8' height='11' rx='2'/><rect x='3' y='13' width='8' height='8' rx='2'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><rect x='3' y='3' width='8' height='8' rx='2'/><rect x='13' y='3' width='8' height='5' rx='2'/><rect x='13' y='10' width='8' height='11' rx='2'/><rect x='3' y='13' width='8' height='8' rx='2'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(2)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><circle cx='12' cy='8' r='4'/><path d='M5.5 20a6.5 6.5 0 0 1 13 0'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><circle cx='12' cy='8' r='4'/><path d='M5.5 20a6.5 6.5 0 0 1 13 0'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(3)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 20h9'/><path d='M16.5 3.5a2.1 2.1 0 0 1 3 3L7 19l-4 1 1-4 12.5-12.5z'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 20h9'/><path d='M16.5 3.5a2.1 2.1 0 0 1 3 3L7 19l-4 1 1-4 12.5-12.5z'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(4)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M14 3H7a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h10a2 2 0 0 0 2-2V8z'/><path d='M14 3v5h5'/><path d='M9 13h6'/><path d='M9 17h6'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M14 3H7a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h10a2 2 0 0 0 2-2V8z'/><path d='M14 3v5h5'/><path d='M9 13h6'/><path d='M9 17h6'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(5)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M14 3H7a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h10a2 2 0 0 0 2-2V8z'/><path d='M14 3v5h5'/><path d='M9 13h3'/><path d='M8 18l5.3-5.3a1.8 1.8 0 0 1 2.5 0l.5.5a1.8 1.8 0 0 1 0 2.5L11 21H8z'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M14 3H7a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h10a2 2 0 0 0 2-2V8z'/><path d='M14 3v5h5'/><path d='M9 13h3'/><path d='M8 18l5.3-5.3a1.8 1.8 0 0 1 2.5 0l.5.5a1.8 1.8 0 0 1 0 2.5L11 21H8z'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(6)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 3v12'/><path d='M7 10l5 5 5-5'/><path d='M5 21h14'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 3v12'/><path d='M7 10l5 5 5-5'/><path d='M5 21h14'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(7)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 2v20'/><path d='M17 6.5a4 4 0 0 0-4-2.5H10a3 3 0 0 0 0 6h4a3 3 0 0 1 0 6h-3a4 4 0 0 1-4-2.5'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 2v20'/><path d='M17 6.5a4 4 0 0 0-4-2.5H10a3 3 0 0 0 0 6h4a3 3 0 0 1 0 6h-3a4 4 0 0 1-4-2.5'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(8)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M8 2v4'/><path d='M16 2v4'/><rect x='3' y='4' width='18' height='18' rx='2'/><path d='M3 10h18'/><path d='M8 14h.01'/><path d='M12 14h.01'/><path d='M16 14h.01'/><path d='M8 18h.01'/><path d='M12 18h.01'/><path d='M16 18h.01'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M8 2v4'/><path d='M16 2v4'/><rect x='3' y='4' width='18' height='18' rx='2'/><path d='M3 10h18'/><path d='M8 14h.01'/><path d='M12 14h.01'/><path d='M16 14h.01'/><path d='M8 18h.01'/><path d='M12 18h.01'/><path d='M16 18h.01'/></svg>");
        }

        div[data-testid="stRadio"] label:nth-of-type(9)::before {
            -webkit-mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 2v3'/><path d='M12 19v3'/><path d='M4.93 4.93l2.12 2.12'/><path d='M16.95 16.95l2.12 2.12'/><path d='M2 12h3'/><path d='M19 12h3'/><path d='M4.93 19.07l2.12-2.12'/><path d='M16.95 7.05l2.12-2.12'/><circle cx='12' cy='12' r='4'/></svg>");
            mask-image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='black' stroke-width='1.9' stroke-linecap='round' stroke-linejoin='round'><path d='M12 2v3'/><path d='M12 19v3'/><path d='M4.93 4.93l2.12 2.12'/><path d='M16.95 16.95l2.12 2.12'/><path d='M2 12h3'/><path d='M19 12h3'/><path d='M4.93 19.07l2.12-2.12'/><path d='M16.95 7.05l2.12-2.12'/><circle cx='12' cy='12' r='4'/></svg>");
        }

        .ss-nav {
            display: flex;
            flex-direction: column;
            gap: 10px;
        }

        .ss-nav-link {
            display: flex;
            align-items: center;
            gap: 14px;
            padding: 14px 16px;
            border-radius: 18px;
            border: 1px solid transparent;
            text-decoration: none;
            color: rgba(244,247,245,0.88) !important;
            background: transparent;
            transition: all 0.18s ease;
            font-weight: 500;
            letter-spacing: 0.01em;
        }

        .ss-nav-link:hover {
            background: rgba(255,255,255,0.08);
            border-color: rgba(255,255,255,0.10);
            color: #ffffff !important;
            transform: translateX(2px);
        }

        .ss-nav-link.active {
            background: linear-gradient(180deg, #f5ecdf 0%, #ead9bd 100%);
            border-color: rgba(197, 167, 122, 0.34);
            box-shadow: 0 14px 24px rgba(0,0,0,0.22);
            color: #181818 !important;
            font-weight: 700;
            transform: none;
        }

        .ss-nav-icon {
            width: 20px;
            height: 20px;
            color: currentColor;
            display: flex;
            align-items: center;
            justify-content: center;
            flex: 0 0 20px;
        }

        .ss-nav-icon svg {
            width: 20px;
            height: 20px;
            display: block;
        }

        .ss-nav-label {
            font-size: 1rem;
            line-height: 1.1;
        }

        .ss-sidebar-spacer {
            min-height: 5rem;
        }

        .ss-sidebar-footer {
            margin-top: 1rem;
            padding-top: 1rem;
            border-top: 1px solid rgba(255,255,255,0.08);
        }

        .ss-user-meta {
            color: rgba(244,247,245,0.74);
            font-size: 0.92rem;
            margin-bottom: 0.3rem;
        }

        section[data-testid="stSidebar"] .streamlit-expanderHeader {
            border-radius: 16px;
            border: 1px solid rgba(255,255,255,0.10);
            background: rgba(255,255,255,0.04);
        }

        section[data-testid="stSidebar"] .stButton > button {
            background: linear-gradient(180deg, rgba(255,255,255,0.10), rgba(255,255,255,0.04));
            color: #f4f7f5;
            border: 1px solid rgba(255,255,255,0.10);
            box-shadow: none;
        }

        section[data-testid="stSidebar"] .stButton > button:hover {
            border-color: rgba(197, 167, 122, 0.50);
            color: #ffffff;
        }

        .ss-ag-prof-list {
            display: flex;
            flex-direction: column;
            gap: 8px;
            margin: 0.35rem 0 1rem;
        }

        .ss-ag-prof-item {
            display: flex;
            align-items: center;
            gap: 10px;
            color: var(--ss-text);
            font-size: 0.96rem;
        }

        .ss-ag-prof-dot {
            width: 12px;
            height: 12px;
            border-radius: 999px;
            display: inline-block;
            flex: 0 0 12px;
        }

        .ss-ag-day-wrap {
            display: grid;
            grid-template-columns: 88px repeat(auto-fit, minmax(210px, 1fr));
            gap: 14px;
            align-items: start;
            width: 100%;
        }

        .ss-ag-time-col,
        .ss-ag-day-col {
            background: var(--ss-surface);
            border: 1px solid var(--ss-border);
            border-radius: 20px;
            overflow: hidden;
            box-shadow: var(--ss-shadow);
        }

        .ss-ag-day-col-head {
            min-height: 62px;
            padding: 14px 14px;
            border-bottom: 1px solid #efe8dc;
            display: flex;
            align-items: center;
            justify-content: center;
            background: linear-gradient(180deg, #fffdfa 0%, #faf7f1 100%);
            font-weight: 700;
        }

        .ss-ag-day-col-body {
            position: relative;
            background-image: linear-gradient(to bottom, rgba(197,167,122,0.10) 1px, transparent 1px);
            background-size: 100% 28px;
            background-position: top left;
        }

        .ss-ag-time-label {
            position: absolute;
            left: 0;
            right: 0;
            height: 28px;
            padding-top: 2px;
            text-align: center;
            font-size: 0.82rem;
            color: var(--ss-muted);
        }

        .ss-ag-prof-chip {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            padding: 8px 12px;
            border-radius: 999px;
            border: 1px solid transparent;
            font-size: 0.9rem;
            font-weight: 700;
        }

        .ss-ag-slot-link {
            position: absolute;
            left: 0;
            right: 0;
            display: block;
            text-decoration: none;
            z-index: 1;
        }

        .ss-ag-slot-link:hover {
            background: rgba(197,167,122,0.10);
        }

        .ss-ag-event {
            position: absolute;
            left: 8px;
            right: 8px;
            padding: 8px 10px;
            border-radius: 14px;
            border: 1px solid rgba(197,167,122,0.18);
            border-left-width: 4px;
            color: #1f2937;
            text-decoration: none;
            z-index: 3;
            overflow: hidden;
            box-shadow: 0 8px 20px rgba(17, 17, 17, 0.08);
        }

        .ss-ag-event-time {
            font-size: 0.72rem;
            color: #6b7280;
            margin-bottom: 4px;
        }

        .ss-ag-event-title {
            font-size: 0.92rem;
            font-weight: 700;
            line-height: 1.2;
        }

        .ss-ag-event-sub {
            font-size: 0.78rem;
            color: #4b5563;
            margin-top: 4px;
            line-height: 1.15;
        }

        .ss-ag-now-line {
            position: absolute;
            left: 0;
            right: 0;
            height: 2px;
            background: #bb3e3e;
            box-shadow: 0 0 0 1px rgba(255,255,255,0.55);
            z-index: 2;
        }

        .ss-ag-week-head {
            background: var(--ss-surface);
            border: 1px solid var(--ss-border);
            border-radius: 18px;
            box-shadow: var(--ss-shadow);
            padding: 12px 14px;
            margin-bottom: 0.4rem;
        }

        .ss-ag-week-day {
            font-size: 0.8rem;
            color: var(--ss-muted);
            text-transform: uppercase;
            letter-spacing: 0.08em;
        }

        .ss-ag-week-date {
            font-size: 0.98rem;
            font-weight: 700;
            margin-top: 2px;
        }

        .ss-ag-week-total {
            font-size: 0.78rem;
            color: var(--ss-muted);
            margin-top: 3px;
        }

        .ss-ag-week-event,
        .ss-ag-month-day {
            display: block;
            text-decoration: none;
            color: inherit;
        }

        .ss-ag-week-event {
            background: var(--ss-surface);
            border: 1px solid var(--ss-border);
            border-left-width: 4px;
            border-radius: 16px;
            padding: 10px 12px;
            margin-bottom: 10px;
            box-shadow: var(--ss-shadow);
        }

        .ss-ag-week-time {
            font-size: 0.74rem;
            color: var(--ss-muted);
            margin-bottom: 4px;
        }

        .ss-ag-week-title {
            font-size: 0.9rem;
            font-weight: 700;
        }

        .ss-ag-week-sub {
            font-size: 0.78rem;
            color: #5b6170;
            margin-top: 4px;
        }

        .ss-ag-week-empty {
            min-height: 88px;
            border: 1px dashed var(--ss-border);
            border-radius: 16px;
            background: rgba(255,255,255,0.55);
            color: var(--ss-muted);
            display: flex;
            align-items: center;
            justify-content: center;
        }

        .ss-ag-month-day {
            min-height: 96px;
            border-radius: 16px;
            padding: 10px 12px;
            box-shadow: var(--ss-shadow);
            margin-bottom: 10px;
        }

        .ss-ag-month-num {
            font-size: 1rem;
            font-weight: 700;
            margin-bottom: 8px;
        }

        .ss-ag-month-total {
            font-size: 0.8rem;
            color: var(--ss-muted);
        }

        .ss-ag-slot-picker-wrap {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(88px, 1fr));
            gap: 8px;
            margin: 0.5rem 0 1rem;
        }

        .ss-ag-slot-picker {
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 42px;
            border-radius: 12px;
            border: 1px solid #d8d0c2;
            text-decoration: none;
            font-size: 0.84rem;
            font-weight: 600;
            transition: transform 0.12s ease, box-shadow 0.12s ease, border-color 0.12s ease;
        }

        .ss-ag-slot-picker:hover {
            transform: translateY(-1px);
            box-shadow: 0 8px 18px rgba(17, 17, 17, 0.08);
        }

        .ss-ag-slot-picker.free {
            background: #fcfaf6;
            color: #4b5563;
        }

        .ss-ag-slot-picker.occupied {
            background: #fce7e7;
            border-color: #efb4b4;
            color: #8d3f3f;
        }

        .ss-ag-slot-picker.selected {
            background: #dbeafe;
            border-color: #7aa7e8;
            color: #1d4ed8;
        }

        .ss-ag-slot-picker.blocked {
            background: #ececec;
            border-color: #d6d6d6;
            color: #8c8c8c;
        }

        .ss-ag-detail-card {
            background: linear-gradient(180deg, #ffffff 0%, #fbfaf7 100%);
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            padding: 18px 20px;
            box-shadow: var(--ss-shadow);
            margin-bottom: 1rem;
        }

        .ss-ag-detail-top {
            display: flex;
            align-items: flex-start;
            justify-content: space-between;
            gap: 16px;
            margin-bottom: 14px;
        }

        .ss-ag-detail-title {
            font-size: 1.18rem;
            font-weight: 700;
            color: #1f2937;
        }

        .ss-ag-detail-sub {
            font-size: 0.92rem;
            color: var(--ss-muted);
            margin-top: 3px;
        }

        .ss-ag-status-pill {
            display: inline-flex;
            align-items: center;
            padding: 8px 12px;
            border-radius: 999px;
            border: 1px solid transparent;
            font-size: 0.82rem;
            font-weight: 700;
            white-space: nowrap;
        }

        .ss-ag-detail-grid {
            display: grid;
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 12px 18px;
            margin-bottom: 12px;
        }

        .ss-ag-detail-grid div,
        .ss-ag-detail-obs {
            background: rgba(245, 236, 223, 0.35);
            border: 1px solid rgba(197, 167, 122, 0.14);
            border-radius: 14px;
            padding: 12px 14px;
            font-size: 0.92rem;
            color: #374151;
        }

        .ss-wa-link {
            display: inline-flex;
            width: 100%;
            min-height: 40px;
            align-items: center;
            justify-content: center;
            text-decoration: none;
            border-radius: 12px;
            background: linear-gradient(180deg, #f5ecdf 0%, #ead9bd 100%);
            color: #181818 !important;
            font-weight: 700;
            border: 1px solid rgba(197, 167, 122, 0.34);
            box-shadow: 0 8px 18px rgba(17, 17, 17, 0.08);
        }

        [data-testid="stMetric"] {
            background: var(--ss-surface);
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            padding: 14px 16px;
            box-shadow: var(--ss-shadow);
        }

        [data-testid="stMetricLabel"], [data-testid="stMetricValue"] {
            color: var(--ss-text);
        }

        div[data-testid="stHorizontalBlock"] > div:has(> [data-testid="stMetric"]) {
            align-self: stretch;
        }

        div[data-testid="stTextInput"] > div > div,
        div[data-testid="stNumberInput"] > div > div,
        div[data-testid="stDateInput"] > div > div,
        div[data-testid="stSelectbox"] > div > div,
        div[data-testid="stMultiSelect"] > div > div,
        div[data-testid="stTextArea"] > div > div {
            border-radius: 16px !important;
            border: 1px solid var(--ss-border) !important;
            box-shadow: none !important;
            background: rgba(255,255,255,0.96) !important;
        }

        div[data-testid="stDataFrame"] {
            border: 1px solid var(--ss-border);
            border-radius: 20px;
            overflow: hidden;
            box-shadow: var(--ss-shadow);
            background: var(--ss-surface);
        }

        .stButton > button,
        .stDownloadButton > button {
            border-radius: 14px;
            border: 1px solid rgba(197, 167, 122, 0.24);
            background: linear-gradient(180deg, #ffffff 0%, #f8f2e8 100%);
            color: var(--ss-text);
            font-weight: 600;
            box-shadow: 0 8px 20px rgba(17, 17, 17, 0.08);
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            border-color: rgba(197, 167, 122, 0.42);
            color: var(--ss-accent);
        }

        .stTabs [data-baseweb="tab-list"] {
            gap: 10px;
            background: transparent;
        }

        .stTabs [data-baseweb="tab"] {
            background: rgba(255,255,255,0.88);
            border-radius: 14px;
            border: 1px solid var(--ss-border);
            padding: 10px 16px;
            box-shadow: 0 6px 18px rgba(16, 44, 38, 0.06);
        }

        .stTabs [aria-selected="true"] {
            background: linear-gradient(180deg, #ffffff 0%, #f8f2e8 100%) !important;
            border-color: rgba(197, 167, 122, 0.24) !important;
        }

        h1, h2, h3 {
            color: var(--ss-text);
            letter-spacing: -0.02em;
        }

        p, label, .stMarkdown, .stCaption {
            color: var(--ss-text);
        }

        .ss-dashboard-shell {
            background: rgba(255,255,255,0.82);
            border: 1px solid rgba(197, 167, 122, 0.16);
            border-radius: 28px;
            padding: 18px 20px;
            box-shadow: var(--ss-shadow);
            backdrop-filter: blur(10px);
            margin-bottom: 18px;
        }

        .ss-dashboard-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            gap: 16px;
            margin-bottom: 10px;
        }

        .ss-search {
            flex: 1;
            border: 1px solid var(--ss-border);
            background: rgba(255,255,255,0.92);
            border-radius: 18px;
            padding: 14px 18px;
            color: var(--ss-muted);
            box-shadow: inset 0 1px 2px rgba(0,0,0,0.03);
        }

        .ss-userbox {
            min-width: 240px;
            text-align: right;
            color: var(--ss-text);
            font-weight: 600;
        }

        .ss-panel {
            background: var(--ss-surface);
            border: 1px solid var(--ss-border);
            border-radius: 24px;
            box-shadow: var(--ss-shadow);
            padding: 18px;
            margin-bottom: 16px;
        }

        .ss-card {
            background:
                linear-gradient(180deg, rgba(255,255,255,0.98), rgba(249,244,236,0.96));
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            box-shadow: var(--ss-shadow);
            padding: 18px;
            min-height: 118px;
            position: relative;
            overflow: hidden;
        }

        .ss-card::after {
            content: "";
            position: absolute;
            inset: auto -10% -30px -10%;
            height: 58px;
            background: linear-gradient(90deg, rgba(197,167,122,0.55), rgba(17,17,17,0.12));
            opacity: 0.35;
            border-radius: 100% 100% 0 0;
        }

        .ss-card-label {
            color: var(--ss-muted);
            font-size: 0.95rem;
            margin-bottom: 14px;
        }

        .ss-card-icon {
            font-size: 1.15rem;
            margin-right: 8px;
            color: #7b5d34;
        }

        .ss-card-value {
            font-size: 2rem;
            font-weight: 700;
            color: var(--ss-text);
            position: relative;
            z-index: 1;
        }

        .ss-list-item {
            padding: 10px 0;
            border-bottom: 1px solid #eee6da;
        }

        div[role="dialog"],
        div[data-testid="stDialog"] > div {
            animation: ssFadeLift 0.22s ease-out;
        }

        .stButton > button,
        .stDownloadButton > button,
        .stLinkButton > a {
            transition:
                transform 0.18s ease,
                box-shadow 0.22s ease,
                border-color 0.22s ease,
                background 0.22s ease,
                color 0.22s ease;
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover,
        .stLinkButton > a:hover {
            transform: translateY(-1px);
            box-shadow: 0 10px 24px rgba(17, 17, 17, 0.10);
        }

        .stButton > button:active,
        .stDownloadButton > button:active,
        .stLinkButton > a:active {
            transform: translateY(0);
        }

        .stTextInput input,
        .stTextArea textarea,
        .stNumberInput input,
        .stDateInput input,
        .stSelectbox [data-baseweb="select"],
        .stMultiSelect [data-baseweb="select"] {
            transition:
                border-color 0.18s ease,
                box-shadow 0.18s ease,
                background 0.18s ease;
        }

        .stTextInput input:focus,
        .stTextArea textarea:focus,
        .stNumberInput input:focus,
        .stDateInput input:focus,
        .stSelectbox [data-baseweb="select"]:focus-within,
        .stMultiSelect [data-baseweb="select"]:focus-within {
            border-color: rgba(197,167,122,0.42) !important;
            box-shadow: 0 0 0 4px rgba(197,167,122,0.10) !important;
        }

        div[data-testid="stSpinner"] {
            border-radius: 18px;
            background: rgba(255,255,255,0.74);
            border: 1px solid rgba(197,167,122,0.18);
            padding: 0.75rem 1rem;
            box-shadow: 0 16px 32px rgba(17,17,17,0.08);
            animation: ssFadeLift 0.18s ease-out;
        }

        .ss-patient-hero-search {
            background: linear-gradient(180deg, rgba(255,255,255,0.96), rgba(251,248,243,0.96));
            border: 1px solid var(--ss-border);
            border-radius: 28px;
            box-shadow: var(--ss-shadow);
            padding: 28px;
            margin-bottom: 1.4rem;
        }

        .ss-patient-search-note {
            color: var(--ss-muted);
            font-size: 0.96rem;
            margin-top: 0.4rem;
        }

        .ss-patient-grid-card {
            background: linear-gradient(180deg, #ffffff 0%, #fcfaf6 100%);
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            box-shadow: var(--ss-shadow);
            padding: 18px;
            min-height: 168px;
            margin-bottom: 1rem;
            transition:
                transform 0.18s ease,
                box-shadow 0.22s ease,
                border-color 0.22s ease;
            animation: ssFadeLift 0.2s ease-out;
        }

        .ss-patient-grid-card:hover {
            transform: translateY(-4px);
            box-shadow: 0 22px 36px rgba(17,17,17,0.10);
            border-color: rgba(197,167,122,0.26);
        }

        .ss-patient-grid-title {
            font-size: 1.04rem;
            font-weight: 700;
            color: #1f2937;
            line-height: 1.25;
            margin-bottom: 10px;
        }

        .ss-patient-grid-meta {
            font-size: 0.88rem;
            color: var(--ss-muted);
            line-height: 1.55;
            margin-bottom: 12px;
        }

        .ss-patient-chip-row {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
            margin: 0.15rem 0 0.8rem;
        }

        .ss-patient-chip {
            display: inline-flex;
            align-items: center;
            padding: 6px 10px;
            border-radius: 999px;
            background: #f5ecdf;
            color: #7b5d34;
            font-size: 0.76rem;
            font-weight: 700;
            border: 1px solid rgba(197,167,122,0.18);
        }

        .ss-patient-hero {
            background: linear-gradient(180deg, rgba(255,255,255,0.98), rgba(250,246,239,0.96));
            border: 1px solid var(--ss-border);
            border-radius: 28px;
            box-shadow: var(--ss-shadow);
            padding: 24px;
            margin-bottom: 1.25rem;
            animation: ssFadeLift 0.22s ease-out;
        }

        .ss-patient-hero-top {
            display: flex;
            justify-content: space-between;
            gap: 18px;
            align-items: flex-start;
            flex-wrap: wrap;
        }

        .ss-patient-hero-left {
            display: flex;
            gap: 18px;
            align-items: center;
            min-width: 0;
        }

        .ss-patient-avatar {
            width: 74px;
            height: 74px;
            border-radius: 24px;
            background: linear-gradient(180deg, #d7bb92 0%, #b89466 100%);
            color: #ffffff;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.55rem;
            font-weight: 800;
            box-shadow: 0 14px 30px rgba(17,17,17,0.12);
            flex: 0 0 74px;
        }

        .ss-patient-name {
            font-size: 1.85rem;
            font-weight: 800;
            line-height: 1.05;
            color: #1f2937;
            margin-bottom: 6px;
        }

        .ss-patient-subline {
            display: flex;
            flex-wrap: wrap;
            gap: 10px 14px;
            color: var(--ss-muted);
            font-size: 0.92rem;
        }

        .ss-patient-status {
            display: inline-flex;
            align-items: center;
            padding: 7px 12px;
            border-radius: 999px;
            background: #eefbf2;
            color: #027a48;
            border: 1px solid rgba(2,122,72,0.15);
            font-size: 0.8rem;
            font-weight: 700;
        }

        .ss-patient-sidecard {
            background: linear-gradient(180deg, #ffffff 0%, #fcfaf6 100%);
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            box-shadow: var(--ss-shadow);
            padding: 18px;
            margin-bottom: 1rem;
            transition:
                transform 0.18s ease,
                box-shadow 0.22s ease,
                border-color 0.22s ease;
        }

        .ss-patient-sidecard:hover {
            transform: translateY(-2px);
            box-shadow: 0 18px 30px rgba(17,17,17,0.09);
            border-color: rgba(197,167,122,0.22);
        }

        .ss-patient-sidecard-title {
            font-size: 0.82rem;
            color: var(--ss-muted);
            text-transform: uppercase;
            letter-spacing: 0.08em;
            margin-bottom: 10px;
            font-weight: 700;
        }

        .ss-patient-sidecard-value {
            font-size: 1rem;
            line-height: 1.55;
            color: #1f2937;
        }

        .ss-patient-sidecard-value strong {
            font-size: 1.02rem;
        }

        .ss-patient-form-section {
            background: linear-gradient(180deg, rgba(255,255,255,0.96), rgba(251,248,243,0.92));
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            padding: 18px;
            box-shadow: var(--ss-shadow);
            margin-bottom: 1rem;
            transition:
                transform 0.18s ease,
                box-shadow 0.22s ease,
                border-color 0.22s ease;
        }

        .ss-patient-form-section:hover {
            transform: translateY(-1px);
            box-shadow: 0 18px 28px rgba(17,17,17,0.08);
            border-color: rgba(197,167,122,0.20);
        }

        .ss-patient-form-title {
            font-size: 0.85rem;
            color: var(--ss-muted);
            text-transform: uppercase;
            letter-spacing: 0.09em;
            font-weight: 700;
            margin-bottom: 12px;
        }

        .ss-patient-empty {
            border: 1px dashed var(--ss-border);
            border-radius: 18px;
            background: rgba(255,255,255,0.62);
            padding: 18px;
            color: var(--ss-muted);
        }

        .ss-finance-item {
            border: 1px solid var(--ss-border);
            border-left-width: 4px;
            border-radius: 18px;
            padding: 14px 16px;
            background: #ffffff;
            box-shadow: var(--ss-shadow);
            margin-bottom: 10px;
            transition:
                transform 0.18s ease,
                box-shadow 0.22s ease,
                border-color 0.22s ease;
        }

        .ss-finance-item:hover {
            transform: translateY(-2px);
            box-shadow: 0 18px 28px rgba(17,17,17,0.08);
        }

        .ss-finance-head {
            display: flex;
            justify-content: space-between;
            gap: 14px;
            align-items: flex-start;
            margin-bottom: 6px;
        }

        .ss-finance-value {
            font-size: 1.15rem;
            font-weight: 800;
            color: #1f2937;
        }

        .ss-finance-meta {
            font-size: 0.84rem;
            color: var(--ss-muted);
            line-height: 1.45;
        }

        .ss-status-pill {
            display: inline-flex;
            align-items: center;
            padding: 6px 10px;
            border-radius: 999px;
            font-size: 0.75rem;
            font-weight: 700;
            border: 1px solid transparent;
        }

        .ss-doc-card {
            border: 1px solid var(--ss-border);
            border-radius: 18px;
            padding: 14px 16px;
            background: #ffffff;
            box-shadow: var(--ss-shadow);
            margin-bottom: 10px;
            transition:
                transform 0.18s ease,
                box-shadow 0.22s ease,
                border-color 0.22s ease;
        }

        .ss-doc-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 18px 28px rgba(17,17,17,0.08);
        }

        .ss-contract-card {
            border: 1px solid var(--ss-border);
            border-radius: 22px;
            background: linear-gradient(180deg, #ffffff 0%, #fcfaf6 100%);
            box-shadow: var(--ss-shadow);
            padding: 18px;
            margin-bottom: 1rem;
            transition:
                transform 0.18s ease,
                box-shadow 0.22s ease,
                border-color 0.22s ease;
        }

        .ss-contract-card:hover {
            transform: translateY(-3px);
            box-shadow: 0 22px 34px rgba(17,17,17,0.10);
            border-color: rgba(197,167,122,0.24);
        }

        .ss-contract-status {
            display: inline-flex;
            align-items: center;
            padding: 7px 12px;
            border-radius: 999px;
            background: #eefbf2;
            color: #027a48;
            font-size: 0.76rem;
            font-weight: 800;
            border: 1px solid rgba(2,122,72,0.12);
            margin-bottom: 12px;
            animation: ssSoftPulse 1.6s ease-out 1;
        }

        .ss-contract-value {
            font-size: 1.6rem;
            font-weight: 800;
            color: #1f2937;
            margin: 6px 0 10px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def agora_str():
    return datetime.now().isoformat(sep=" ", timespec="seconds")


def registrar_feedback_visual(mensagem, tipo="success"):
    st.session_state["ss_feedback_visual"] = {
        "mensagem": str(mensagem or "").strip(),
        "tipo": str(tipo or "success").strip().lower(),
    }


def exibir_feedback_visual():
    feedback = st.session_state.pop("ss_feedback_visual", None)
    if not feedback or not feedback.get("mensagem"):
        return

    mensagem = feedback["mensagem"]
    tipo = feedback.get("tipo", "success")
    icones = {"success": "✅", "warning": "⚠️", "error": "❌", "info": "ℹ️"}

    if hasattr(st, "toast"):
        st.toast(mensagem, icon=icones.get(tipo, "✅"))
        return

    if tipo == "error":
        st.error(mensagem)
    elif tipo == "warning":
        st.warning(mensagem)
    elif tipo == "info":
        st.info(mensagem)
    else:
        st.success(mensagem)


def formatar_data_br(data_obj):
    return data_obj.strftime("%d/%m/%Y")


def formatar_data_br_valor(valor):
    texto_bruto = str(valor or "").strip()
    if texto_bruto.lower() in {"nat", "nan", "none"}:
        return ""
    data_convertida = parse_data_contrato(valor)
    if data_convertida:
        return formatar_data_br(data_convertida)
    return texto_bruto


def formatar_data_hora_br_valor(valor):
    texto = str(valor or "").strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return ""
    try:
        return datetime.fromisoformat(texto).strftime("%d/%m/%Y %H:%M:%S")
    except ValueError:
        pass
    return formatar_data_br_valor(valor)


def parse_hora_agendamento(valor):
    texto = str(valor or "").strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return None
    for formato in ("%H:%M", "%H:%M:%S"):
        try:
            return datetime.strptime(texto, formato).time()
        except ValueError:
            continue
    return None


def formatar_hora_agendamento(valor):
    hora = parse_hora_agendamento(valor)
    return hora.strftime("%H:%M") if hora else ""


def ordenar_hora_agendamento(valor):
    hora = parse_hora_agendamento(valor)
    return hora.hour * 60 + hora.minute if hora else 10**9


def gerar_horarios_intervalo(inicio="08:00", fim="20:00", intervalo_minutos=15):
    hora_inicio = parse_hora_agendamento(inicio)
    hora_fim = parse_hora_agendamento(fim)
    if hora_inicio is None or hora_fim is None:
        return []
    atual = datetime.combine(date.today(), hora_inicio)
    limite = datetime.combine(date.today(), hora_fim)
    horarios = []
    while atual <= limite:
        horarios.append(atual.strftime("%H:%M"))
        atual += timedelta(minutes=intervalo_minutos)
    return horarios


def normalizar_dias_atendimento(valor):
    if isinstance(valor, list):
        dias = valor
    else:
        texto = str(valor or "").strip()
        dias = [parte.strip() for parte in texto.split(",") if parte.strip()]
    return [dia for dia in dias if dia in DIAS_SEMANA_PT]


def serializar_dias_atendimento(dias):
    return ",".join(normalizar_dias_atendimento(dias))


def profissional_atende_no_horario(profissional_row, data_ref, hora_texto):
    if profissional_row is None:
        return True
    dias = normalizar_dias_atendimento(profissional_row.get("dias_atendimento", ""))
    if dias:
        dia_atual = DIAS_SEMANA_PT[data_ref.weekday()]
        if dia_atual not in dias:
            return False
    hora_inicio = formatar_hora_agendamento(profissional_row.get("hora_inicio", ""))
    hora_fim = formatar_hora_agendamento(profissional_row.get("hora_fim", ""))
    if hora_inicio and hora_texto < hora_inicio:
        return False
    if hora_fim and hora_texto >= hora_fim:
        return False
    return True


def formatar_prontuario_valor(valor):
    texto = texto_importacao(valor)
    if not texto:
        return ""
    if re.fullmatch(r"-?\d+\.0+", texto):
        return texto.split(".")[0]
    return texto


def formatar_parcela_valor(valor):
    texto = str(valor or "").strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return "-"
    try:
        numero = float(str(valor).replace(",", "."))
        if numero.is_integer():
            return str(int(numero))
        return texto
    except (TypeError, ValueError):
        texto_limpo = texto_importacao(valor)
        return texto_limpo or "-"


def forma_pagamento_a_vista(forma_pagamento):
    return normalizar_forma_pagamento(forma_pagamento) in {"Pix", "Dinheiro", "Debito", "Credito"}


def parse_data_contrato(valor):
    texto = str(valor or "").strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return None

    try:
        return datetime.fromisoformat(texto).date()
    except ValueError:
        pass

    texto_apenas_digitos = re.sub(r"\D", "", texto)
    if len(texto_apenas_digitos) == 8:
        try:
            return datetime.strptime(texto_apenas_digitos, "%d%m%Y").date()
        except ValueError:
            pass

    formatos = ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M:%S.%f"]
    for formato in formatos:
        try:
            return datetime.strptime(texto, formato).date()
        except ValueError:
            continue
    try:
        data_pandas = pd.to_datetime(texto, dayfirst=True, errors="coerce")
        if not pd.isna(data_pandas):
            return data_pandas.date()
    except Exception:
        pass
    return None


def adicionar_meses(data_base, quantidade_meses):
    mes = data_base.month - 1 + quantidade_meses
    ano = data_base.year + mes // 12
    mes = mes % 12 + 1
    ultimo_dia = [31, 29 if ano % 4 == 0 and (ano % 100 != 0 or ano % 400 == 0) else 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31][mes - 1]
    dia = min(data_base.day, ultimo_dia)
    return date(ano, mes, dia)


def extrair_texto_pdf_upload(arquivo_pdf):
    if not PDF_LEITURA_DISPONIVEL:
        raise ValueError("Leitura de PDF indisponivel. Instale pypdf ou PyPDF2 para habilitar essa funcao.")

    arquivo_pdf.seek(0)
    reader = PdfReader(BytesIO(arquivo_pdf.read()))
    textos = []
    for pagina in reader.pages:
        texto = pagina.extract_text() or ""
        if texto.strip():
            textos.append(texto)
    return "\n".join(textos)


def extrair_por_rotulo(texto, rotulos):
    for rotulo in rotulos:
        padrao = rf"{rotulo}\s*[:\-]?\s*(.+)"
        match = re.search(padrao, texto, flags=re.IGNORECASE)
        if match:
            valor = match.group(1).strip()
            valor = valor.split("\n")[0].strip()
            if valor:
                return valor
    return ""


def limpar_trecho_endereco(valor):
    texto = str(valor or "").strip(" -,:;")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def separar_endereco_campos(endereco, numero, bairro, cidade, estado):
    endereco = limpar_trecho_endereco(endereco)
    numero = limpar_trecho_endereco(numero)
    bairro = limpar_trecho_endereco(bairro)
    cidade = limpar_trecho_endereco(cidade)
    estado = limpar_trecho_endereco(estado)

    texto_completo = " ".join([parte for parte in [endereco, numero, bairro, cidade, estado] if parte]).strip()

    if not numero:
        match_numero = re.search(
            r"\b(\d+[A-Za-z]?(?:\s*complemento\s*[:\-]?\s*[^,]+)?)",
            texto_completo,
            flags=re.IGNORECASE,
        )
        if match_numero:
            numero = limpar_trecho_endereco(match_numero.group(1))

    if not estado:
        match_estado = re.search(r"\b([A-Z]{2})\b", texto_completo)
        if match_estado:
            estado = limpar_trecho_endereco(match_estado.group(1))

    if not cidade:
        match_cidade = re.search(r"cidade\s*[:\-]?\s*([A-Za-zÀ-ÿ\s]+)", texto_completo, flags=re.IGNORECASE)
        if match_cidade:
            cidade = limpar_trecho_endereco(match_cidade.group(1))
        elif " - " in texto_completo:
            partes = [limpar_trecho_endereco(parte) for parte in texto_completo.split(" - ") if limpar_trecho_endereco(parte)]
            if partes:
                cidade = partes[-1]

    if not bairro:
        match_bairro = re.search(r"bairro\s*[:\-]?\s*([A-Za-zÀ-ÿ\s]+)", texto_completo, flags=re.IGNORECASE)
        if match_bairro:
            bairro = limpar_trecho_endereco(match_bairro.group(1))

    endereco = re.sub(r"\bcidade\s*[:\-]?\s*.*$", "", endereco, flags=re.IGNORECASE)
    endereco = re.sub(r"\bestado\s*[:\-]?\s*.*$", "", endereco, flags=re.IGNORECASE)
    endereco = re.sub(r"\bcep\s*[:\-]?\s*.*$", "", endereco, flags=re.IGNORECASE)
    endereco = limpar_trecho_endereco(endereco)

    cidade = re.sub(r"estado\s*[:\-]?\s*.*$", "", cidade, flags=re.IGNORECASE)
    bairro = re.sub(r"cidade\s*[:\-]?\s*.*$", "", bairro, flags=re.IGNORECASE)
    bairro = re.sub(r"estado\s*[:\-]?\s*.*$", "", bairro, flags=re.IGNORECASE)
    bairro = re.sub(r"\bcidade\b$", "", bairro, flags=re.IGNORECASE)
    cidade = limpar_trecho_endereco(cidade)
    bairro = limpar_trecho_endereco(bairro)

    if bairro and cidade and cidade in bairro:
        bairro = limpar_trecho_endereco(bairro.replace(cidade, ""))
    if cidade and estado and estado in cidade:
        cidade = limpar_trecho_endereco(cidade.replace(estado, ""))

    return {
        "endereco": endereco,
        "numero": numero,
        "bairro": bairro,
        "cidade": cidade,
        "estado": estado,
    }


def extrair_dados_paciente_pdf(texto):
    if not texto.strip():
        return {}

    linhas = [linha.strip() for linha in texto.splitlines() if linha.strip()]
    texto_unico = "\n".join(linhas)

    cpf = ""
    match_cpf = re.search(r"\b\d{3}\.?\d{3}\.?\d{3}-?\d{2}\b", texto_unico)
    if match_cpf:
        cpf = limpar_cpf(match_cpf.group(0))

    telefone = ""
    match_tel = re.search(r"(\(?\d{2}\)?\s?\d{4,5}-?\d{4})", texto_unico)
    if match_tel:
        telefone = match_tel.group(1)

    cep = ""
    match_cep = re.search(r"\b\d{5}-?\d{3}\b", texto_unico)
    if match_cep:
        cep = match_cep.group(0)

    nascimento = ""
    match_data = re.search(r"\b\d{2}/\d{2}/\d{4}\b", texto_unico)
    if match_data:
        nascimento = match_data.group(0)

    nome = extrair_por_rotulo(texto_unico, [r"nome do paciente", r"paciente", r"nome"])
    prontuario = extrair_por_rotulo(texto_unico, [r"prontuario", r"prontu[aá]rio"])
    endereco = extrair_por_rotulo(texto_unico, [r"endereco", r"endere[cç]o", r"rua", r"logradouro"])
    numero = extrair_por_rotulo(texto_unico, [r"numero", r"n[uú]mero"])
    bairro = extrair_por_rotulo(texto_unico, [r"bairro"])
    cidade = extrair_por_rotulo(texto_unico, [r"cidade"])
    estado = extrair_por_rotulo(texto_unico, [r"estado", r"uf"])
    responsavel = extrair_por_rotulo(texto_unico, [r"responsavel legal", r"respons[aá]vel"])

    cpf_responsavel = ""
    match_cpf_resp = re.search(r"cpf do respons[aá]vel\s*[:\-]?\s*(\d{3}\.?\d{3}\.?\d{3}-?\d{2})", texto_unico, flags=re.IGNORECASE)
    if match_cpf_resp:
        cpf_responsavel = limpar_cpf(match_cpf_resp.group(1))

    if not nome:
        for linha in linhas[:8]:
            texto_linha = normalizar_texto(linha)
            if any(ch.isdigit() for ch in linha):
                continue
            if any(rotulo in texto_linha for rotulo in ["cpf", "telefone", "endereco", "bairro", "cidade", "estado", "responsavel", "prontuario"]):
                continue
            if len(linha.split()) >= 2:
                nome = linha
                break

    endereco_campos = separar_endereco_campos(endereco, numero, bairro, cidade, estado)

    return {
        "nome": nome,
        "prontuario": prontuario,
        "cpf": cpf,
        "data_nascimento": formatar_data_br_valor(nascimento),
        "telefone": telefone,
        "cep": cep,
        "endereco": endereco_campos["endereco"],
        "numero": endereco_campos["numero"],
        "bairro": endereco_campos["bairro"],
        "cidade": endereco_campos["cidade"],
        "estado": endereco_campos["estado"],
        "menor_idade": bool(responsavel.strip()),
        "responsavel": responsavel,
        "cpf_responsavel": cpf_responsavel,
    }


def aplicar_dados_extraidos_paciente(dados):
    mapa = {
        "paciente_nome_input": dados.get("nome", ""),
        "paciente_prontuario_input": dados.get("prontuario", ""),
        "paciente_cpf_input": dados.get("cpf", ""),
        "paciente_nascimento_input": dados.get("data_nascimento", ""),
        "paciente_telefone_input": dados.get("telefone", ""),
        "paciente_cep_input": dados.get("cep", ""),
        "paciente_endereco_input": dados.get("endereco", ""),
        "paciente_numero_input": dados.get("numero", ""),
        "paciente_bairro_input": dados.get("bairro", ""),
        "paciente_cidade_input": dados.get("cidade", ""),
        "paciente_estado_input": dados.get("estado", ""),
        "paciente_menor_input": dados.get("menor_idade", False),
        "paciente_responsavel_input": dados.get("responsavel", ""),
        "paciente_cpf_responsavel_input": dados.get("cpf_responsavel", ""),
    }
    for chave, valor in mapa.items():
        st.session_state[chave] = valor


def limpar_nome(nome):
    nome = unicodedata.normalize("NFKD", nome or "")
    nome = nome.encode("ASCII", "ignore").decode("ASCII")
    nome = nome.replace(" ", "_")
    return nome.upper()


def normalizar_texto(valor):
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    texto = texto.encode("ASCII", "ignore").decode("ASCII")
    return texto.strip().lower()


def normalizar_forma_pagamento(valor):
    valor_normalizado = normalizar_texto(valor)
    mapa = {
        "pix": "Pix",
        "dinheiro": "Dinheiro",
        "debito": "Debito",
        "credito": "Credito",
        "boleto": "Boleto",
    }
    return mapa.get(valor_normalizado, FORMAS_PAGAMENTO[0])


def limpar_cpf(cpf):
    return "".join(ch for ch in str(cpf or "") if ch.isdigit())


def cpf_valido(cpf):
    cpf_limpo = limpar_cpf(cpf)
    if not cpf_limpo:
        return True
    if len(cpf_limpo) != 11:
        return False
    if cpf_limpo == cpf_limpo[0] * 11:
        return False

    soma = sum(int(cpf_limpo[i]) * (10 - i) for i in range(9))
    digito_1 = (soma * 10 % 11) % 10
    if digito_1 != int(cpf_limpo[9]):
        return False

    soma = sum(int(cpf_limpo[i]) * (11 - i) for i in range(10))
    digito_2 = (soma * 10 % 11) % 10
    return digito_2 == int(cpf_limpo[10])


def conectar_banco():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False, timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("PRAGMA busy_timeout = 30000")
    try:
        conn.execute("PRAGMA journal_mode = WAL")
    except sqlite3.OperationalError as exc:
        if "locked" not in str(exc).lower():
            raise
    return conn


conn = conectar_banco()
cursor = conn.cursor()


def tabela_existe(nome_tabela):
    row = cursor.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",
        (nome_tabela,),
    ).fetchone()
    return row is not None


def colunas_tabela(nome_tabela):
    if not tabela_existe(nome_tabela):
        return set()
    return {row["name"] for row in conn.execute(f"PRAGMA table_info({nome_tabela})")}


def gerar_hash_senha(senha, salt=None):
    salt_bytes = salt or secrets.token_bytes(16)
    hash_bytes = hashlib.pbkdf2_hmac("sha256", str(senha or "").encode("utf-8"), salt_bytes, 100000)
    return f"{salt_bytes.hex()}${hash_bytes.hex()}"


def verificar_senha(senha, senha_hash):
    if not senha_hash or "$" not in str(senha_hash):
        return False
    salt_hex, hash_hex = str(senha_hash).split("$", 1)
    try:
        salt_bytes = bytes.fromhex(salt_hex)
    except ValueError:
        return False
    hash_calculado = gerar_hash_senha(senha, salt_bytes).split("$", 1)[1]
    return hmac.compare_digest(hash_calculado, hash_hex)


def garantir_usuario_admin_inicial():
    usuario_existente = cursor.execute(
        "SELECT id FROM usuarios WHERE ativo=1 ORDER BY id LIMIT 1"
    ).fetchone()
    if usuario_existente:
        return
    cursor.execute(
        """
        INSERT INTO usuarios
        (nome, usuario, senha_hash, perfil, ativo, acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios, data_criacao)
        VALUES (?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?)
        """,
        (
            "Administrador",
            USUARIO_ADMIN_INICIAL,
            gerar_hash_senha(SENHA_ADMIN_INICIAL),
            "Administrador",
            1,
            1,
            1,
            1,
            1,
            agora_str(),
        ),
    )
    conn.commit()


def autenticar_usuario(usuario, senha):
    row = cursor.execute(
        """
        SELECT * FROM usuarios
        WHERE lower(usuario)=lower(?) AND ativo=1
        LIMIT 1
        """,
        (str(usuario or "").strip(),),
    ).fetchone()
    if not row:
        return None
    if not verificar_senha(senha, row["senha_hash"]):
        return None
    return row


def carregar_usuarios():
    return pd.read_sql(
        """
        SELECT id, nome, usuario, perfil, ativo, data_criacao,
               acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios
        FROM usuarios
        ORDER BY nome, usuario
        """,
        conn,
    )


def usuario_existe(usuario, ignorar_id=None):
    if ignorar_id is None:
        row = cursor.execute("SELECT id FROM usuarios WHERE lower(usuario)=lower(?) LIMIT 1", (str(usuario or "").strip(),)).fetchone()
    else:
        row = cursor.execute(
            "SELECT id FROM usuarios WHERE lower(usuario)=lower(?) AND id<>? LIMIT 1",
            (str(usuario or "").strip(), int(ignorar_id)),
        ).fetchone()
    return row is not None


def criar_usuario(nome, usuario, senha, perfil="Usuario"):
    acesso_dashboard = 1
    acesso_pacientes = 1
    acesso_contratos = 1
    acesso_financeiro = 1 if perfil == "Administrador" else 0
    acesso_usuarios = 1 if perfil == "Administrador" else 0
    cursor.execute(
        """
        INSERT INTO usuarios
        (nome, usuario, senha_hash, perfil, ativo, acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios, data_criacao)
        VALUES (?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?)
        """,
        (
            nome.strip(),
            usuario.strip(),
            gerar_hash_senha(senha),
            perfil.strip() or "Usuario",
            acesso_dashboard,
            acesso_pacientes,
            acesso_contratos,
            acesso_financeiro,
            acesso_usuarios,
            agora_str(),
        ),
    )


def atualizar_usuario_admin(usuario_id, nome, usuario, perfil, ativo, acesso_dashboard, acesso_pacientes, acesso_contratos, acesso_financeiro, acesso_usuarios):
    cursor.execute(
        """
        UPDATE usuarios
        SET nome=?, usuario=?, perfil=?, ativo=?, acesso_dashboard=?, acesso_pacientes=?, acesso_contratos=?, acesso_financeiro=?, acesso_usuarios=?
        WHERE id=?
        """,
        (
            nome.strip(),
            usuario.strip(),
            perfil.strip() or "Usuario",
            int(bool(ativo)),
            int(bool(acesso_dashboard)),
            int(bool(acesso_pacientes)),
            int(bool(acesso_contratos)),
            int(bool(acesso_financeiro)),
            int(bool(acesso_usuarios)),
            int(usuario_id),
        ),
    )


def redefinir_senha_usuario(usuario_id, nova_senha):
    cursor.execute(
        "UPDATE usuarios SET senha_hash=? WHERE id=?",
        (gerar_hash_senha(nova_senha), int(usuario_id)),
    )


def registrar_log_acesso(usuario_id, usuario_login, evento):
    cursor.execute(
        """
        INSERT INTO logs_acesso
        (usuario_id, usuario, evento, data_hora)
        VALUES (?, ?, ?, ?)
        """,
        (int(usuario_id), str(usuario_login or "").strip(), str(evento or "").strip(), agora_str()),
    )
    conn.commit()


def usuario_tem_acesso(usuario_sessao, menu_nome):
    if not usuario_sessao:
        return False
    if usuario_sessao.get("perfil") == "Administrador":
        return True
    chave = MODULOS_SISTEMA.get(menu_nome)
    if not chave:
        return False
    return bool(usuario_sessao.get(chave, 0))


def renderizar_login():
    st.title("Acesso ao Sistema")
    st.caption("Entre com usuario e senha para acessar o sistema.")

    with st.form("login_form"):
        usuario = st.text_input("Usuario")
        senha = st.text_input("Senha", type="password")
        entrar = st.form_submit_button("Entrar")

    if entrar:
        usuario_row = autenticar_usuario(usuario, senha)
        if usuario_row is None:
            st.error("Usuario ou senha invalidos.")
        else:
            st.session_state["usuario_logado"] = {
                "id": int(usuario_row["id"]),
                "nome": usuario_row["nome"],
                "usuario": usuario_row["usuario"],
                "perfil": usuario_row["perfil"],
                "acesso_dashboard": int(usuario_row["acesso_dashboard"] or 0),
                "acesso_pacientes": int(usuario_row["acesso_pacientes"] or 0),
                "acesso_contratos": int(usuario_row["acesso_contratos"] or 0),
                "acesso_financeiro": int(usuario_row["acesso_financeiro"] or 0),
                "acesso_usuarios": int(usuario_row["acesso_usuarios"] or 0),
            }
            registrar_log_acesso(usuario_row["id"], usuario_row["usuario"], "LOGIN")
            st.rerun()

def garantir_coluna(nome_tabela, definicao_coluna):
    nome_coluna = definicao_coluna.split()[0]
    if nome_coluna not in colunas_tabela(nome_tabela):
        cursor.execute(f"ALTER TABLE {nome_tabela} ADD COLUMN {definicao_coluna}")


def garantir_indice(nome_indice, sql_criacao):
    cursor.execute(sql_criacao)


def migrar_base_agenda_legada():
    # Preserva a compatibilidade com os campos antigos de agenda/profissionais,
    # espelhando dados legados nas novas colunas quando elas estiverem vazias.
    agora = agora_str()
    try:
        cursor.execute(
            """
            UPDATE agendamentos
            SET
                data_agendamento = COALESCE(NULLIF(data_agendamento, ''), NULLIF(data, '')),
                nome_paciente_snapshot = COALESCE(NULLIF(nome_paciente_snapshot, ''), NULLIF(paciente_nome, '')),
                procedimento_nome_snapshot = COALESCE(NULLIF(procedimento_nome_snapshot, ''), NULLIF(procedimento, '')),
                observacoes = COALESCE(NULLIF(observacoes, ''), NULLIF(observacao, '')),
                criado_em = COALESCE(NULLIF(criado_em, ''), NULLIF(data_criacao, ''), ?),
                atualizado_em = COALESCE(NULLIF(atualizado_em, ''), NULLIF(criado_em, ''), NULLIF(data_criacao, ''), ?)
            """
            ,
            (agora, agora),
        )
        cursor.execute(
            """
            UPDATE profissionais
            SET
                ordem_exibicao = COALESCE(ordem_exibicao, id),
                criado_em = COALESCE(NULLIF(criado_em, ''), NULLIF(data_criacao, ''), ?),
                atualizado_em = COALESCE(NULLIF(atualizado_em, ''), NULLIF(criado_em, ''), NULLIF(data_criacao, ''), ?)
            """
            ,
            (agora, agora),
        )
        cursor.execute(
            """
            UPDATE tipos_atendimento
            SET
                ordem_exibicao = COALESCE(ordem_exibicao, id),
                criado_em = COALESCE(NULLIF(criado_em, ''), ?),
                atualizado_em = COALESCE(NULLIF(atualizado_em, ''), NULLIF(criado_em, ''), ?)
            """
            ,
            (agora, agora),
        )
        cursor.execute(
            """
            UPDATE procedimentos
            SET
                criado_em = COALESCE(NULLIF(criado_em, ''), ?),
                atualizado_em = COALESCE(NULLIF(atualizado_em, ''), NULLIF(criado_em, ''), ?)
            """
            ,
            (agora, agora),
        )
    except sqlite3.OperationalError as exc:
        if "locked" not in str(exc).lower():
            raise


def inicializar_banco():
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS pacientes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            apelido TEXT,
            sexo TEXT,
            prontuario TEXT,
            cpf TEXT,
            rg TEXT,
            data_nascimento TEXT,
            telefone TEXT,
            email TEXT,
            cep TEXT,
            endereco TEXT,
            numero TEXT,
            bairro TEXT,
            cidade TEXT,
            estado TEXT,
            estado_civil TEXT,
            observacoes TEXT,
            menor_idade INTEGER DEFAULT 0,
            responsavel TEXT,
            cpf_responsavel TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS contratos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            paciente_id INTEGER,
            valor_total REAL DEFAULT 0,
            entrada REAL DEFAULT 0,
            parcelas INTEGER DEFAULT 1,
            primeiro_vencimento TEXT,
            data_pagamento_entrada TEXT,
            forma_pagamento TEXT,
            hash_importacao TEXT,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS procedimentos_contrato (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contrato_id INTEGER,
            procedimento TEXT,
            valor REAL DEFAULT 0
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS financeiro (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            origem TEXT,
            descricao TEXT,
            valor REAL DEFAULT 0,
            data TEXT,
            tipo TEXT,
            contrato_id INTEGER,
            recebivel_id INTEGER,
            prontuario TEXT,
            forma_pagamento TEXT,
            conta_caixa TEXT,
            observacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS agendamentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data TEXT,
            hora_inicio TEXT,
            hora_fim TEXT,
            paciente_id INTEGER,
            paciente_nome TEXT,
            profissional TEXT,
            procedimento TEXT,
            status TEXT DEFAULT 'Agendado',
            observacao TEXT,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS profissionais (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            especialidade TEXT,
            cor TEXT,
            dias_atendimento TEXT,
            hora_inicio TEXT,
            hora_fim TEXT,
            ativo INTEGER DEFAULT 1,
            observacao TEXT,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS tipos_atendimento (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            cor TEXT,
            ativo INTEGER DEFAULT 1,
            ordem_exibicao INTEGER,
            criado_em TEXT,
            atualizado_em TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS procedimentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            categoria TEXT,
            valor_padrao REAL DEFAULT 0,
            duracao_padrao_minutos INTEGER DEFAULT 60,
            descricao TEXT,
            cor_opcional TEXT,
            ativo INTEGER DEFAULT 1,
            criado_em TEXT,
            atualizado_em TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS agendamento_procedimentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            agendamento_id INTEGER,
            procedimento_id INTEGER,
            procedimento_nome_snapshot TEXT,
            valor_snapshot REAL DEFAULT 0,
            duracao_snapshot_minutos INTEGER,
            origem_contrato INTEGER DEFAULT 0,
            contrato_id INTEGER
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS lembretes_agendamento (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            agendamento_id INTEGER,
            tipo_lembrete TEXT,
            canal TEXT,
            mensagem TEXT,
            status_envio TEXT,
            criado_em TEXT,
            enviado_em TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS modelos_mensagem_agendamento (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo_mensagem TEXT UNIQUE,
            titulo TEXT,
            conteudo TEXT,
            ativo INTEGER DEFAULT 1,
            criado_em TEXT,
            atualizado_em TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS pacientes_rapidos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            telefone TEXT,
            email TEXT,
            criado_em TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS saldos_conta (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data TEXT,
            conta TEXT,
            saldo REAL DEFAULT 0,
            observacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS recebiveis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contrato_id INTEGER,
            paciente_id INTEGER,
            paciente_nome TEXT,
            prontuario TEXT,
            parcela_numero INTEGER,
            vencimento TEXT,
            valor REAL DEFAULT 0,
            forma_pagamento TEXT,
            status TEXT DEFAULT 'Aberto',
            observacao TEXT,
            data_criacao TEXT,
            hash_importacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT,
            usuario TEXT,
            senha_hash TEXT,
            perfil TEXT DEFAULT 'Administrador',
            ativo INTEGER DEFAULT 1,
            acesso_dashboard INTEGER DEFAULT 1,
            acesso_pacientes INTEGER DEFAULT 1,
            acesso_contratos INTEGER DEFAULT 1,
            acesso_financeiro INTEGER DEFAULT 1,
            acesso_usuarios INTEGER DEFAULT 0,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS contas_pagar (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data_vencimento TEXT,
            descricao TEXT,
            fornecedor TEXT,
            categoria TEXT,
            valor REAL DEFAULT 0,
            pago TEXT,
            valor_pago REAL DEFAULT 0,
            status TEXT DEFAULT 'A vencer',
            observacao TEXT,
            data_criacao TEXT,
            hash_importacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS vendas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data_venda TEXT,
            paciente_nome TEXT,
            valor_total REAL DEFAULT 0,
            valor_a_vista REAL DEFAULT 0,
            valor_cartao REAL DEFAULT 0,
            valor_boleto REAL DEFAULT 0,
            saldo REAL DEFAULT 0,
            data_a_pagar TEXT,
            avaliador TEXT,
            vendedor TEXT,
            nf TEXT,
            contrato_id INTEGER,
            hash_importacao TEXT,
            data_criacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS metas_vendas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ano INTEGER,
            meta REAL DEFAULT 100000,
            supermeta REAL DEFAULT 150000,
            hipermeta REAL DEFAULT 200000,
            data_atualizacao TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS logs_acesso (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER,
            usuario TEXT,
            evento TEXT,
            data_hora TEXT
        )
        """
    )

    garantir_coluna("financeiro", "origem TEXT")
    garantir_coluna("financeiro", "contrato_id INTEGER")
    garantir_coluna("financeiro", "recebivel_id INTEGER")
    garantir_coluna("financeiro", "prontuario TEXT")
    garantir_coluna("financeiro", "forma_pagamento TEXT")
    garantir_coluna("financeiro", "conta_caixa TEXT")
    garantir_coluna("financeiro", "observacao TEXT")
    garantir_coluna("saldos_conta", "data TEXT")
    garantir_coluna("saldos_conta", "conta TEXT")
    garantir_coluna("saldos_conta", "saldo REAL DEFAULT 0")
    garantir_coluna("saldos_conta", "observacao TEXT")
    garantir_coluna("recebiveis", "paciente_id INTEGER")
    garantir_coluna("recebiveis", "paciente_nome TEXT")
    garantir_coluna("recebiveis", "prontuario TEXT")
    garantir_coluna("recebiveis", "parcela_numero INTEGER")
    garantir_coluna("recebiveis", "vencimento TEXT")
    garantir_coluna("recebiveis", "valor REAL DEFAULT 0")
    garantir_coluna("recebiveis", "forma_pagamento TEXT")
    garantir_coluna("recebiveis", "status TEXT DEFAULT 'Aberto'")
    garantir_coluna("recebiveis", "observacao TEXT")
    garantir_coluna("recebiveis", "data_criacao TEXT")
    garantir_coluna("recebiveis", "data_pagamento TEXT")
    garantir_coluna("recebiveis", "hash_importacao TEXT")
    garantir_coluna("contas_pagar", "data_vencimento TEXT")
    garantir_coluna("contas_pagar", "descricao TEXT")
    garantir_coluna("contas_pagar", "fornecedor TEXT")
    garantir_coluna("contas_pagar", "categoria TEXT")
    garantir_coluna("contas_pagar", "valor REAL DEFAULT 0")
    garantir_coluna("contas_pagar", "pago TEXT")
    garantir_coluna("contas_pagar", "valor_pago REAL DEFAULT 0")
    garantir_coluna("contas_pagar", "status TEXT DEFAULT 'A vencer'")
    garantir_coluna("contas_pagar", "observacao TEXT")
    garantir_coluna("contas_pagar", "data_criacao TEXT")
    garantir_coluna("contas_pagar", "hash_importacao TEXT")
    garantir_coluna("agendamentos", "data TEXT")
    garantir_coluna("agendamentos", "hora_inicio TEXT")
    garantir_coluna("agendamentos", "hora_fim TEXT")
    garantir_coluna("agendamentos", "paciente_id INTEGER")
    garantir_coluna("agendamentos", "paciente_nome TEXT")
    garantir_coluna("agendamentos", "profissional TEXT")
    garantir_coluna("agendamentos", "procedimento TEXT")
    garantir_coluna("agendamentos", "status TEXT DEFAULT 'Agendado'")
    garantir_coluna("agendamentos", "observacao TEXT")
    garantir_coluna("agendamentos", "data_criacao TEXT")
    garantir_coluna("agendamentos", "nome_paciente_snapshot TEXT")
    garantir_coluna("agendamentos", "telefone_snapshot TEXT")
    garantir_coluna("agendamentos", "email_snapshot TEXT")
    garantir_coluna("agendamentos", "profissional_id INTEGER")
    garantir_coluna("agendamentos", "tipo_atendimento_id INTEGER")
    garantir_coluna("agendamentos", "procedimento_id INTEGER")
    garantir_coluna("agendamentos", "procedimento_nome_snapshot TEXT")
    garantir_coluna("agendamentos", "contrato_id INTEGER")
    garantir_coluna("agendamentos", "origem_contrato INTEGER DEFAULT 0")
    garantir_coluna("agendamentos", "data_agendamento TEXT")
    garantir_coluna("agendamentos", "duracao_minutos INTEGER")
    garantir_coluna("agendamentos", "observacoes TEXT")
    garantir_coluna("agendamentos", "criado_por TEXT")
    garantir_coluna("agendamentos", "criado_em TEXT")
    garantir_coluna("agendamentos", "atualizado_em TEXT")
    garantir_coluna("profissionais", "nome TEXT")
    garantir_coluna("profissionais", "especialidade TEXT")
    garantir_coluna("profissionais", "cor TEXT")
    garantir_coluna("profissionais", "dias_atendimento TEXT")
    garantir_coluna("profissionais", "hora_inicio TEXT")
    garantir_coluna("profissionais", "hora_fim TEXT")
    garantir_coluna("profissionais", "ativo INTEGER DEFAULT 1")
    garantir_coluna("profissionais", "observacao TEXT")
    garantir_coluna("profissionais", "data_criacao TEXT")
    garantir_coluna("profissionais", "ordem_exibicao INTEGER")
    garantir_coluna("profissionais", "criado_em TEXT")
    garantir_coluna("profissionais", "atualizado_em TEXT")
    garantir_coluna("tipos_atendimento", "nome TEXT")
    garantir_coluna("tipos_atendimento", "cor TEXT")
    garantir_coluna("tipos_atendimento", "ativo INTEGER DEFAULT 1")
    garantir_coluna("tipos_atendimento", "ordem_exibicao INTEGER")
    garantir_coluna("tipos_atendimento", "criado_em TEXT")
    garantir_coluna("tipos_atendimento", "atualizado_em TEXT")
    garantir_coluna("procedimentos", "nome TEXT")
    garantir_coluna("procedimentos", "categoria TEXT")
    garantir_coluna("procedimentos", "valor_padrao REAL DEFAULT 0")
    garantir_coluna("procedimentos", "duracao_padrao_minutos INTEGER DEFAULT 60")
    garantir_coluna("procedimentos", "descricao TEXT")
    garantir_coluna("procedimentos", "cor_opcional TEXT")
    garantir_coluna("procedimentos", "ativo INTEGER DEFAULT 1")
    garantir_coluna("procedimentos", "criado_em TEXT")
    garantir_coluna("procedimentos", "atualizado_em TEXT")
    garantir_coluna("agendamento_procedimentos", "agendamento_id INTEGER")
    garantir_coluna("agendamento_procedimentos", "procedimento_id INTEGER")
    garantir_coluna("agendamento_procedimentos", "procedimento_nome_snapshot TEXT")
    garantir_coluna("agendamento_procedimentos", "valor_snapshot REAL DEFAULT 0")
    garantir_coluna("agendamento_procedimentos", "duracao_snapshot_minutos INTEGER")
    garantir_coluna("agendamento_procedimentos", "origem_contrato INTEGER DEFAULT 0")
    garantir_coluna("agendamento_procedimentos", "contrato_id INTEGER")
    garantir_coluna("lembretes_agendamento", "agendamento_id INTEGER")
    garantir_coluna("lembretes_agendamento", "tipo_lembrete TEXT")
    garantir_coluna("lembretes_agendamento", "canal TEXT")
    garantir_coluna("lembretes_agendamento", "mensagem TEXT")
    garantir_coluna("lembretes_agendamento", "status_envio TEXT")
    garantir_coluna("lembretes_agendamento", "criado_em TEXT")
    garantir_coluna("lembretes_agendamento", "enviado_em TEXT")
    garantir_coluna("modelos_mensagem_agendamento", "tipo_mensagem TEXT")
    garantir_coluna("modelos_mensagem_agendamento", "titulo TEXT")
    garantir_coluna("modelos_mensagem_agendamento", "conteudo TEXT")
    garantir_coluna("modelos_mensagem_agendamento", "ativo INTEGER DEFAULT 1")
    garantir_coluna("modelos_mensagem_agendamento", "criado_em TEXT")
    garantir_coluna("modelos_mensagem_agendamento", "atualizado_em TEXT")
    garantir_coluna("pacientes_rapidos", "nome TEXT")
    garantir_coluna("pacientes_rapidos", "telefone TEXT")
    garantir_coluna("pacientes_rapidos", "email TEXT")
    garantir_coluna("pacientes_rapidos", "criado_em TEXT")
    garantir_coluna("vendas", "data_venda TEXT")
    garantir_coluna("vendas", "paciente_nome TEXT")
    garantir_coluna("vendas", "valor_total REAL DEFAULT 0")
    garantir_coluna("vendas", "valor_a_vista REAL DEFAULT 0")
    garantir_coluna("vendas", "valor_cartao REAL DEFAULT 0")
    garantir_coluna("vendas", "valor_boleto REAL DEFAULT 0")
    garantir_coluna("vendas", "saldo REAL DEFAULT 0")
    garantir_coluna("vendas", "data_a_pagar TEXT")
    garantir_coluna("vendas", "avaliador TEXT")
    garantir_coluna("vendas", "vendedor TEXT")
    garantir_coluna("vendas", "nf TEXT")
    garantir_coluna("vendas", "contrato_id INTEGER")
    garantir_coluna("vendas", "hash_importacao TEXT")
    garantir_coluna("vendas", "data_criacao TEXT")
    garantir_coluna("metas_vendas", "ano INTEGER")
    garantir_coluna("metas_vendas", "meta REAL DEFAULT 100000")
    garantir_coluna("metas_vendas", "supermeta REAL DEFAULT 150000")
    garantir_coluna("metas_vendas", "hipermeta REAL DEFAULT 200000")
    garantir_coluna("metas_vendas", "data_atualizacao TEXT")
    garantir_coluna("usuarios", "nome TEXT")
    garantir_coluna("usuarios", "usuario TEXT")
    garantir_coluna("usuarios", "senha_hash TEXT")
    garantir_coluna("usuarios", "perfil TEXT DEFAULT 'Administrador'")
    garantir_coluna("usuarios", "ativo INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_dashboard INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_pacientes INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_contratos INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_financeiro INTEGER DEFAULT 1")
    garantir_coluna("usuarios", "acesso_usuarios INTEGER DEFAULT 0")
    garantir_coluna("usuarios", "data_criacao TEXT")
    garantir_coluna("logs_acesso", "usuario_id INTEGER")
    garantir_coluna("logs_acesso", "usuario TEXT")
    garantir_coluna("logs_acesso", "evento TEXT")
    garantir_coluna("logs_acesso", "data_hora TEXT")
    garantir_coluna("contratos", "data_criacao TEXT")
    garantir_coluna("contratos", "data_pagamento_entrada TEXT")
    garantir_coluna("contratos", "hash_importacao TEXT")
    garantir_coluna("pacientes", "numero TEXT")
    garantir_coluna("pacientes", "bairro TEXT")
    garantir_coluna("pacientes", "estado TEXT")
    garantir_coluna("pacientes", "responsavel TEXT")
    garantir_coluna("pacientes", "cpf_responsavel TEXT")
    garantir_coluna("pacientes", "apelido TEXT")
    garantir_coluna("pacientes", "sexo TEXT")
    garantir_coluna("pacientes", "rg TEXT")
    garantir_coluna("pacientes", "email TEXT")
    garantir_coluna("pacientes", "estado_civil TEXT")
    garantir_coluna("pacientes", "observacoes TEXT")

    migrar_base_agenda_legada()
    garantir_indice(
        "idx_agendamentos_data_agendamento",
        "CREATE INDEX IF NOT EXISTS idx_agendamentos_data_agendamento ON agendamentos(data_agendamento)",
    )
    garantir_indice(
        "idx_agendamentos_profissional_id",
        "CREATE INDEX IF NOT EXISTS idx_agendamentos_profissional_id ON agendamentos(profissional_id)",
    )
    garantir_indice(
        "idx_agendamentos_paciente_id",
        "CREATE INDEX IF NOT EXISTS idx_agendamentos_paciente_id ON agendamentos(paciente_id)",
    )
    garantir_indice(
        "idx_agendamento_procedimentos_agendamento_id",
        "CREATE INDEX IF NOT EXISTS idx_agendamento_procedimentos_agendamento_id ON agendamento_procedimentos(agendamento_id)",
    )
    garantir_indice(
        "idx_lembretes_agendamento_id",
        "CREATE INDEX IF NOT EXISTS idx_lembretes_agendamento_id ON lembretes_agendamento(agendamento_id)",
    )
    garantir_indice(
        "idx_modelos_mensagem_tipo",
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_modelos_mensagem_tipo ON modelos_mensagem_agendamento(tipo_mensagem)",
    )
    garantir_indice(
        "idx_profissionais_ativo_ordem",
        "CREATE INDEX IF NOT EXISTS idx_profissionais_ativo_ordem ON profissionais(ativo, ordem_exibicao)",
    )
    garantir_indice(
        "idx_tipos_atendimento_ativo_ordem",
        "CREATE INDEX IF NOT EXISTS idx_tipos_atendimento_ativo_ordem ON tipos_atendimento(ativo, ordem_exibicao)",
    )
    garantir_indice(
        "idx_procedimentos_ativo_categoria",
        "CREATE INDEX IF NOT EXISTS idx_procedimentos_ativo_categoria ON procedimentos(ativo, categoria)",
    )

    conn.commit()
    garantir_usuario_admin_inicial()


def garantir_pasta_documentos():
    if not os.path.exists(DOCS_DIR):
        os.makedirs(DOCS_DIR)


def proximo_arquivo_documento(nome_base):
    arquivo_inicial = os.path.join(DOCS_DIR, f"{nome_base}.docx")
    if not os.path.exists(arquivo_inicial):
        return arquivo_inicial

    versao = 2
    while True:
        candidato = os.path.join(DOCS_DIR, f"{nome_base}_v{versao}.docx")
        if not os.path.exists(candidato):
            return candidato
        versao += 1


def substituir_runs(doc, dados):
    for p in doc.paragraphs:
        substituir_placeholders_paragrafo(p, dados)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    substituir_placeholders_paragrafo(p, dados)


def substituir_placeholders_paragrafo(paragrafo, dados):
    for run in paragrafo.runs:
        for chave, valor in dados.items():
            if chave in run.text:
                run.text = run.text.replace(chave, str(valor))

    texto = paragrafo.text
    texto_atualizado = texto
    for chave, valor in dados.items():
        texto_atualizado = texto_atualizado.replace(chave, str(valor))

    if texto_atualizado != texto:
        paragrafo.text = texto_atualizado


def valor_paciente(paciente, coluna):
    valor = paciente[coluna].iloc[0] if coluna in paciente.columns else ""
    if pd.isna(valor):
        return ""
    return str(valor).strip()


def montar_endereco_paciente(paciente):
    endereco = valor_paciente(paciente, "endereco")
    numero = valor_paciente(paciente, "numero")
    bairro = valor_paciente(paciente, "bairro")
    cidade = valor_paciente(paciente, "cidade")
    estado = valor_paciente(paciente, "estado")
    cep = valor_paciente(paciente, "cep")

    partes = []
    if endereco:
        if numero:
            partes.append(f"{endereco}, {numero}")
        else:
            partes.append(endereco)
    elif numero:
        partes.append(numero)

    for item in [bairro, cidade, estado]:
        if item:
            partes.append(item)

    endereco_completo = " - ".join(partes)
    if cep:
        endereco_completo = f"{endereco_completo} - CEP {cep}" if endereco_completo else f"CEP {cep}"
    return endereco_completo


def montar_qualificacao(paciente):
    nome = valor_paciente(paciente, "nome")
    cpf = valor_paciente(paciente, "cpf")
    telefone = valor_paciente(paciente, "telefone")
    nascimento = valor_paciente(paciente, "data_nascimento")
    endereco = montar_endereco_paciente(paciente)
    menor_idade = valor_paciente(paciente, "menor_idade") in {"1", "True", "true"}
    responsavel = valor_paciente(paciente, "responsavel")
    cpf_responsavel = valor_paciente(paciente, "cpf_responsavel")

    if menor_idade and responsavel:
        partes = [responsavel]
        if cpf_responsavel:
            partes.append(f"CPF {cpf_responsavel}")
        if telefone:
            partes.append(f"telefone {telefone}")
        if endereco:
            partes.append(f"endereco {endereco}")
        if nome:
            partes.append(f"responsavel legal pelo(a) paciente {nome}")
        if cpf:
            partes.append(f"CPF do paciente {cpf}")
        if nascimento:
            partes.append(f"nascido(a) em {nascimento}")
        return ", ".join(partes)

    partes = [nome]
    if cpf:
        partes.append(f"CPF {cpf}")
    if telefone:
        partes.append(f"telefone {telefone}")
    if nascimento:
        partes.append(f"nascido(a) em {nascimento}")
    if endereco:
        partes.append(f"endereco {endereco}")
    return ", ".join([parte for parte in partes if parte])


def dados_assinatura(paciente):
    nome = valor_paciente(paciente, "nome")
    cpf = valor_paciente(paciente, "cpf")
    menor_idade = valor_paciente(paciente, "menor_idade") in {"1", "True", "true"}
    responsavel = valor_paciente(paciente, "responsavel")
    cpf_responsavel = valor_paciente(paciente, "cpf_responsavel")

    if menor_idade and responsavel:
        return {
            "titulo_assinatura": "",
            "nome_assinatura": responsavel,
            "cpf_assinatura": f"CPF do responsavel: {cpf_responsavel}" if cpf_responsavel else "CPF do responsavel:",
            "assinatura_menor": True,
            "cpf_original": cpf,
        }

    return {
        "titulo_assinatura": "",
        "nome_assinatura": nome,
        "cpf_assinatura": f"CPF: {cpf}" if cpf else "CPF:",
        "assinatura_menor": False,
        "cpf_original": cpf,
    }


def montar_termo_cirurgia(procedimentos):
    itens = []
    for proc in procedimentos["procedimento"].tolist():
        texto = str(proc).strip()
        if not texto:
            continue
        itens.append(texto)

    if not itens:
        return ""
    if len(itens) == 1:
        return itens[0]
    return "; ".join(itens)


def ajustar_assinatura_capa(doc, nome_paciente, cpf_paciente, nome_assinatura, cpf_assinatura):
    paragrafos = doc.paragraphs
    titulo_contrato = "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO"
    limite = len(paragrafos)

    for indice, paragrafo in enumerate(paragrafos):
        if titulo_contrato in paragrafo.text:
            limite = indice
            break

    for indice in range(limite - 1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip() if indice + 1 < limite else ""

        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            paragrafos[indice].text = nome_assinatura
            paragrafos[indice + 1].text = cpf_assinatura
            return

    for indice in range(limite):
        paragrafo = paragrafos[indice]
        if "___________________________________________________________" not in paragrafo.text:
            continue
        if indice + 2 >= limite:
            continue
        proximo = paragrafos[indice + 1]
        seguinte = paragrafos[indice + 2]
        proximo.text = nome_assinatura
        seguinte.text = cpf_assinatura
        return


def inserir_paragrafo_apos(paragrafo, texto=""):
    novo_p = OxmlElement("w:p")
    paragrafo._p.addnext(novo_p)
    novo_paragrafo = Paragraph(novo_p, paragrafo._parent)
    novo_paragrafo.text = texto
    novo_paragrafo.alignment = paragrafo.alignment
    if paragrafo.style is not None:
        novo_paragrafo.style = paragrafo.style
    return novo_paragrafo


def substituir_par_assinatura(paragrafo_nome, paragrafo_cpf, assinatura):
    if assinatura["assinatura_menor"]:
        paragrafo_nome.text = assinatura["nome_assinatura"]
        paragrafo_cpf.text = assinatura["cpf_assinatura"]
        return

    paragrafo_nome.text = assinatura["nome_assinatura"]
    paragrafo_cpf.text = assinatura["cpf_assinatura"]


def ajustar_assinaturas_menor(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip()

        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            substituir_par_assinatura(paragrafos[indice], paragrafos[indice + 1], assinatura)


def ajustar_assinaturas_por_texto(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    for paragrafo in paragrafos:
        texto = paragrafo.text
        if not texto:
            continue

        texto_ajustado = texto
        if nome_paciente and nome_paciente in texto_ajustado:
            texto_ajustado = texto_ajustado.replace(nome_paciente, assinatura["nome_assinatura"])

        if cpf_paciente:
            texto_ajustado = texto_ajustado.replace(f"CPF: {cpf_paciente}", assinatura["cpf_assinatura"])
            texto_ajustado = texto_ajustado.replace(f"CPF {cpf_paciente}", assinatura["cpf_assinatura"])
            texto_ajustado = texto_ajustado.replace(cpf_paciente, assinatura["cpf_original"] if not assinatura["assinatura_menor"] else cpf_paciente)

        if texto_ajustado != texto:
            if assinatura["assinatura_menor"]:
                texto_ajustado = texto_ajustado.replace(cpf_paciente, assinatura["cpf_assinatura"].replace("CPF do responsavel: ", "").strip())
            paragrafo.text = texto_ajustado


def ajustar_ultima_assinatura_menor(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 2, -1, -1):
        texto_atual = paragrafos[indice].text.strip()
        texto_proximo = paragrafos[indice + 1].text.strip()

        if texto_atual == nome_paciente and cpf_paciente and cpf_paciente in texto_proximo:
            substituir_par_assinatura(paragrafos[indice], paragrafos[indice + 1], assinatura)
            return


def ajustar_bloco_final_assinatura(doc, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, 1, -1):
        if "SOUL SUL CLÍNICA INTEGRADA" not in paragrafos[indice].text:
            continue

        indice_linha = None
        for cursor in range(indice - 1, -1, -1):
            texto = paragrafos[cursor].text.strip()
            if "____" in texto or "___" in texto:
                indice_linha = cursor
                break

        if indice_linha is None:
            return

        candidatos = []
        for cursor in range(indice_linha + 1, indice):
            texto = paragrafos[cursor].text.strip()
            if texto:
                candidatos.append(cursor)

        if len(candidatos) < 2:
            return

        indice_nome = candidatos[0]
        indice_cpf = candidatos[1]

        paragrafos[indice_nome].text = assinatura["nome_assinatura"]
        paragrafos[indice_cpf].text = assinatura["cpf_assinatura"]
        paragrafos[indice_nome].alignment = 1
        paragrafos[indice_cpf].alignment = 1

        for cursor in candidatos[2:]:
            paragrafos[cursor].text = ""
        return


def ajustar_assinatura_termo_cirurgico_final(doc, nome_paciente, cpf_paciente, assinatura):
    paragrafos = doc.paragraphs
    ultimo_indice_data = None
    ultimo_indice_clinica = None

    for indice, paragrafo in enumerate(paragrafos):
        texto = paragrafo.text.strip()
        if "Campos dos Goytacazes" in texto:
            ultimo_indice_data = indice
        if "SOUL SUL CLÍNICA INTEGRADA" in texto:
            ultimo_indice_clinica = indice

    if ultimo_indice_data is None or ultimo_indice_clinica is None:
        return
    if ultimo_indice_data >= ultimo_indice_clinica:
        return

    candidatos = []
    for indice in range(ultimo_indice_data + 1, ultimo_indice_clinica):
        texto = paragrafos[indice].text.strip()
        if texto:
            candidatos.append(indice)

    if len(candidatos) < 2:
        return

    for posicao in range(len(candidatos) - 1):
        indice_nome = candidatos[posicao]
        indice_cpf = candidatos[posicao + 1]
        texto_nome = paragrafos[indice_nome].text.strip()
        texto_cpf = paragrafos[indice_cpf].text.strip()

        if texto_nome == nome_paciente and cpf_paciente and cpf_paciente in texto_cpf:
            substituir_par_assinatura(paragrafos[indice_nome], paragrafos[indice_cpf], assinatura)
            return


def ajustar_assinatura_contrato(doc, nome_assinatura, cpf_assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, 1, -1):
        if "SOUL SUL CLÍNICA INTEGRADA" not in paragrafos[indice].text:
            continue

        candidatos = []
        cursor = indice - 1
        while cursor >= 0 and len(candidatos) < 2:
            texto = paragrafos[cursor].text.strip()
            if texto:
                candidatos.append(cursor)
            cursor -= 1

        if len(candidatos) < 2:
            return

        indice_cpf = candidatos[0]
        indice_nome = candidatos[1]
        paragrafos[indice_nome].text = nome_assinatura
        paragrafos[indice_cpf].text = cpf_assinatura
        return


def reescrever_bloco_final_termo(doc, assinatura):
    paragrafos = doc.paragraphs
    for indice in range(len(paragrafos) - 1, -1, -1):
        texto = paragrafos[indice].text
        if "SOUL SUL" not in texto or "TESTEMUNHAS" not in texto:
            continue

        candidatos_limpar = []
        for cursor in range(indice - 1, -1, -1):
            texto_cursor = paragrafos[cursor].text.strip()
            if not texto_cursor:
                continue
            candidatos_limpar.append(cursor)
            if len(candidatos_limpar) == 2:
                break

        for cursor in candidatos_limpar:
            paragrafos[cursor].text = ""

        paragrafos[indice].text = (
            "\n"
            "___________________________________________________________\n"
            f"{assinatura['nome_assinatura']}\n"
            f"{assinatura['cpf_assinatura']}\n\n"
            "\n___________________________________________________________\n"
            "SOUL SUL CLINICA INTEGRADA\n\n"
            "TESTEMUNHAS:\n\n"
            "________________________________                    _________________________________\n"
            "NOME:                                                          NOME:\n"
            "CPF:                                                             CPF:"
        )
        paragrafos[indice].alignment = 1
        return


def remover_elemento(elemento):
    parent = elemento.getparent()
    if parent is not None:
        parent.remove(elemento)


def texto_elemento(elemento):
    return "".join(elemento.itertext())


def remover_secao_termo_cirurgico(doc):
    corpo = doc._element.body
    removendo = False
    elementos_para_remover = []

    for elemento in corpo:
        texto = texto_elemento(elemento)
        if "{TERMO_CIRURGIA}" in texto:
            removendo = True

        if removendo:
            if "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO" in texto:
                removendo = False
                continue
            elementos_para_remover.append(elemento)

    for elemento in elementos_para_remover:
        remover_elemento(elemento)

    for elemento in list(corpo):
        texto = texto_elemento(elemento).strip()
        xml = etree.tostring(elemento, encoding="unicode")
        if "lastRenderedPageBreak" in xml and not texto:
            remover_elemento(elemento)


def normalizar_quebra_antes_contrato(doc):
    titulo_contrato = "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ODONTOLÓGICO"
    paragrafos = doc.paragraphs
    indice_titulo = None

    for indice, paragrafo in enumerate(paragrafos):
        if titulo_contrato in paragrafo.text:
            indice_titulo = indice
            break

    if indice_titulo is None or indice_titulo == 0:
        return

    for indice in range(indice_titulo):
        paragrafo = paragrafos[indice]
        for run in paragrafo.runs:
            elementos_remover = []
            for child in run._element:
                if child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
                    elementos_remover.append(child)
            for child in elementos_remover:
                run._element.remove(child)

    paragrafo_anterior = paragrafos[indice_titulo - 1]
    if not any(
        child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
        for run in paragrafo_anterior.runs
        for child in run._element
    ):
        paragrafo_anterior.add_run().add_break(WD_BREAK.PAGE)


def normalizar_quebra_antes_termo(doc):
    marcadores_termo = [
        "TERMO DE CONSENTIMENTO ESCLARECIDO",
        "CONSENTIMENTO ESCLARECIDO",
        "DO CONSENTIMENTO ESCLARECIDO",
    ]
    paragrafos = doc.paragraphs
    indice_titulo = None

    for indice, paragrafo in enumerate(paragrafos):
        texto_normalizado = normalizar_texto(paragrafo.text).upper()
        if any(normalizar_texto(marcador).upper() in texto_normalizado for marcador in marcadores_termo):
            indice_titulo = indice
            break

    if indice_titulo is None or indice_titulo == 0:
        return

    paragrafo_anterior = paragrafos[indice_titulo - 1]
    if not any(
        child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
        for run in paragrafo_anterior.runs
        for child in run._element
    ):
        paragrafo_anterior.add_run().add_break(WD_BREAK.PAGE)


def garantir_logo_no_cabecalho(doc):
    if not doc.sections:
        return

    doc.sections[0].different_first_page_header_footer = False
    for secao in doc.sections[1:]:
        secao.different_first_page_header_footer = False
        secao.header.is_linked_to_previous = True


def aplicar_fonte_times_new_roman(doc):
    estilos = doc.styles
    for nome_estilo in ["Normal", "Header", "Footer"]:
        if nome_estilo in estilos:
            estilo = estilos[nome_estilo]
            estilo.font.name = "Times New Roman"
            estilo._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            estilo._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            estilo._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    def aplicar_paragrafo(paragrafo):
        for run in paragrafo.runs:
            run.font.name = "Times New Roman"
            if run._element.rPr is None:
                run._element.get_or_add_rPr()
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    for paragrafo in doc.paragraphs:
        aplicar_paragrafo(paragrafo)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    aplicar_paragrafo(paragrafo)

    for secao in doc.sections:
        for paragrafo in secao.header.paragraphs:
            aplicar_paragrafo(paragrafo)
        for tabela in secao.header.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        aplicar_paragrafo(paragrafo)
        for paragrafo in secao.footer.paragraphs:
            aplicar_paragrafo(paragrafo)


def compactar_bloco_final(doc):
    paragrafos = doc.paragraphs
    indice_data = None
    indice_testemunhas = None

    for indice in range(len(paragrafos) - 1, -1, -1):
        texto = paragrafos[indice].text.strip()
        if indice_testemunhas is None and "TESTEMUNHAS" in texto:
            indice_testemunhas = indice
        if indice_data is None and "Campos dos Goytacazes" in texto:
            indice_data = indice
        if indice_data is not None and indice_testemunhas is not None:
            break

    if indice_data is None or indice_testemunhas is None or indice_data >= indice_testemunhas:
        return

    for indice in range(indice_data, len(paragrafos)):
        texto = paragrafos[indice].text.strip()
        if not texto:
            continue

        formato = paragrafos[indice].paragraph_format
        formato.space_before = Pt(0)
        formato.space_after = Pt(0)
        formato.line_spacing = 1.0

        for run in paragrafos[indice].runs:
            run.font.size = Pt(10)

        if "TESTEMUNHAS" in texto:
            for cursor in range(indice + 1, min(indice + 6, len(paragrafos))):
                formato_cursor = paragrafos[cursor].paragraph_format
                formato_cursor.space_before = Pt(0)
                formato_cursor.space_after = Pt(0)
                formato_cursor.line_spacing = 1.0
                for run in paragrafos[cursor].runs:
                    run.font.size = Pt(10)
            break


def preencher_tabela_checklist(tabela, procedimentos):
    for row_index, row in enumerate(tabela.rows):
        for cell in row.cells:
            if "{PROCEDIMENTO}" not in cell.text:
                continue

            if len(tabela.columns) == 1:
                tabela.add_column(Inches(1.0))

            cabecalho = tabela.rows[row_index]
            cabecalho.cells[0].text = "Procedimento"
            cabecalho.cells[1].text = "OK"

            total_procedimentos = len(procedimentos.index)
            while len(tabela.rows) < row_index + 1 + total_procedimentos:
                tabela.add_row()

            for indice_proc, proc_row in enumerate(procedimentos.itertuples(index=False), start=1):
                linha = tabela.rows[row_index + indice_proc]
                linha.cells[0].text = str(proc_row.procedimento)
                linha.cells[1].text = "[   ]"
            return True

    return False


def preencher_tabela_contrato(tabela, procedimentos):
    for row_index, row in enumerate(tabela.rows):
        if len(row.cells) < 2:
            continue

        tem_procedimento = any("{PROCEDIMENTO}" in cell.text for cell in row.cells)
        tem_valor = any("{VALOR}" in cell.text for cell in row.cells)
        if not (tem_procedimento or tem_valor):
            continue

        total_procedimentos = len(procedimentos.index)
        while len(tabela.rows) < row_index + total_procedimentos:
            tabela.add_row()

        for indice_proc, proc_row in enumerate(procedimentos.itertuples(index=False)):
            linha = tabela.rows[row_index + indice_proc]
            linha.cells[0].text = str(proc_row.procedimento)
            linha.cells[1].text = f"R$ {float(proc_row.valor):.2f}"
        return True

    return False


def gerar_documento(conn_local, contrato_id):
    contrato = pd.read_sql(
        "SELECT * FROM contratos WHERE id=?",
        conn_local,
        params=(contrato_id,),
    )
    if contrato.empty:
        raise ValueError("Contrato nao encontrado para gerar documento.")

    paciente_id = int(contrato["paciente_id"].iloc[0])
    paciente = pd.read_sql(
        "SELECT * FROM pacientes WHERE id=?",
        conn_local,
        params=(paciente_id,),
    )
    procedimentos = pd.read_sql(
        "SELECT * FROM procedimentos_contrato WHERE contrato_id=?",
        conn_local,
        params=(contrato_id,),
    )

    if paciente.empty:
        raise ValueError("Paciente vinculado ao contrato nao encontrado.")

    nome = paciente["nome"].iloc[0]
    prontuario = paciente["prontuario"].iloc[0]
    pagamento = montar_texto_pagamento(contrato.iloc[0])

    endereco_completo = montar_endereco_paciente(paciente)
    qualificacao = montar_qualificacao(paciente)
    termo_cirurgia = montar_termo_cirurgia(procedimentos)
    assinatura = dados_assinatura(paciente)
    doc = Document(TEMPLATE_PATH)
    garantir_logo_no_cabecalho(doc)
    if not termo_cirurgia:
        remover_secao_termo_cirurgico(doc)
    dados = {
        "{PACIENTE}": nome,
        "{paciente}": nome,
        "{CPF}": valor_paciente(paciente, "cpf"),
        "{cpf}": valor_paciente(paciente, "cpf"),
        "{PRONTUARIO}": prontuario,
        "{prontuario}": prontuario,
        "{PAGAMENTO}": pagamento,
        "{pagamento}": pagamento,
        "{QUALIFICACAO}": qualificacao,
        "{qualificacao}": qualificacao,
        "{ENDERECO}": valor_paciente(paciente, "endereco"),
        "{endereco}": valor_paciente(paciente, "endereco"),
        "{NUMERO}": valor_paciente(paciente, "numero"),
        "{numero}": valor_paciente(paciente, "numero"),
        "{BAIRRO}": valor_paciente(paciente, "bairro"),
        "{bairro}": valor_paciente(paciente, "bairro"),
        "{CIDADE}": valor_paciente(paciente, "cidade"),
        "{cidade}": valor_paciente(paciente, "cidade"),
        "{ESTADO}": valor_paciente(paciente, "estado"),
        "{estado}": valor_paciente(paciente, "estado"),
        "{CEP}": valor_paciente(paciente, "cep"),
        "{cep}": valor_paciente(paciente, "cep"),
        "{TELEFONE}": valor_paciente(paciente, "telefone"),
        "{telefone}": valor_paciente(paciente, "telefone"),
        "{DATA_NASCIMENTO}": valor_paciente(paciente, "data_nascimento"),
        "{data_nascimento}": valor_paciente(paciente, "data_nascimento"),
        "{TERMO_CIRURGIA}": "",
        "{termo_cirurgia}": "",
    }
    substituir_runs(doc, dados)

    checklist_preenchido = False
    tabela_contrato_preenchida = False
    for tabela in doc.tables:
        if not checklist_preenchido and preencher_tabela_checklist(tabela, procedimentos):
            checklist_preenchido = True
            continue
        if not tabela_contrato_preenchida and preencher_tabela_contrato(tabela, procedimentos):
            tabela_contrato_preenchida = True

    if assinatura["assinatura_menor"]:
        ajustar_assinaturas_menor(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_assinaturas_por_texto(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_assinatura_termo_cirurgico_final(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_ultima_assinatura_menor(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura,
        )
        ajustar_bloco_final_assinatura(doc, assinatura)
    else:
        ajustar_assinatura_capa(
            doc,
            nome,
            valor_paciente(paciente, "cpf"),
            assinatura["nome_assinatura"],
            assinatura["cpf_assinatura"],
        )
        ajustar_assinatura_contrato(doc, assinatura["nome_assinatura"], assinatura["cpf_assinatura"])
    reescrever_bloco_final_termo(doc, assinatura)
    normalizar_quebra_antes_contrato(doc)
    normalizar_quebra_antes_termo(doc)
    compactar_bloco_final(doc)
    aplicar_fonte_times_new_roman(doc)

    nome_base = f"{limpar_nome(nome)}_{prontuario}"
    arquivo = proximo_arquivo_documento(nome_base)
    doc.save(arquivo)
    return arquivo


def carregar_pacientes():
    return pd.read_sql("SELECT * FROM pacientes ORDER BY nome", conn)


def carregar_profissionais(somente_ativos=False):
    if not tabela_existe("profissionais"):
        return pd.DataFrame(columns=["id", "nome", "especialidade", "ativo", "observacao", "data_criacao"])
    profissionais = pd.read_sql("SELECT * FROM profissionais ORDER BY nome", conn)
    if somente_ativos and not profissionais.empty:
        profissionais = profissionais[profissionais["ativo"] == 1].copy()
    return profissionais


def carregar_contratos():
    return pd.read_sql("SELECT * FROM contratos ORDER BY id DESC", conn)


def carregar_procedimentos(contrato_id):
    return pd.read_sql(
        "SELECT * FROM procedimentos_contrato WHERE contrato_id=? ORDER BY id",
        conn,
        params=(contrato_id,),
    )


def carregar_tipos_atendimento(somente_ativos=False):
    if not tabela_existe("tipos_atendimento"):
        return pd.DataFrame(columns=["id", "nome", "cor", "ativo", "ordem_exibicao", "criado_em", "atualizado_em"])
    tipos = pd.read_sql(
        "SELECT * FROM tipos_atendimento ORDER BY COALESCE(ordem_exibicao, 9999), nome",
        conn,
    )
    if somente_ativos and not tipos.empty:
        tipos = tipos[tipos["ativo"].fillna(1).astype(int) == 1].copy()
    return tipos


def carregar_procedimentos_catalogo(somente_ativos=False):
    if not tabela_existe("procedimentos"):
        return pd.DataFrame(
            columns=[
                "id", "nome", "categoria", "valor_padrao", "duracao_padrao_minutos",
                "descricao", "cor_opcional", "ativo", "criado_em", "atualizado_em",
            ]
        )
    procedimentos = pd.read_sql(
        "SELECT * FROM procedimentos ORDER BY categoria, nome",
        conn,
    )
    if somente_ativos and not procedimentos.empty:
        procedimentos = procedimentos[procedimentos["ativo"].fillna(1).astype(int) == 1].copy()
    return procedimentos


def normalizar_cor_interface(cor, fallback="#c5a77a"):
    texto = str(cor or "").strip()
    if not texto:
        return fallback
    if re.fullmatch(r"#[0-9a-fA-F]{6}", texto):
        return texto
    if re.fullmatch(r"#[0-9a-fA-F]{3}", texto):
        return "#" + "".join(ch * 2 for ch in texto[1:])
    return fallback


def hex_para_rgba(cor_hex, alpha=1.0):
    cor = normalizar_cor_interface(cor_hex)
    r = int(cor[1:3], 16)
    g = int(cor[3:5], 16)
    b = int(cor[5:7], 16)
    return f"rgba({r}, {g}, {b}, {alpha})"


def hora_para_minutos(valor):
    hora = parse_hora_agendamento(valor)
    if hora is None:
        return None
    return hora.hour * 60 + hora.minute


def adicionar_minutos_hora(valor, minutos):
    hora = parse_hora_agendamento(valor)
    if hora is None:
        return ""
    instante = datetime.combine(date.today(), hora) + timedelta(minutes=int(minutos or 0))
    return instante.strftime("%H:%M")


def inicio_semana(data_base):
    return data_base - timedelta(days=data_base.weekday())


def fim_semana(data_base):
    return inicio_semana(data_base) + timedelta(days=6)


def carregar_agendamentos_clinica(data_inicial=None, data_final=None, profissionais_ids=None):
    agendamentos = pd.read_sql(
        """
        SELECT
            ag.*,
            pr.nome AS profissional_nome_cadastro,
            pr.cor AS profissional_cor,
            pr.ordem_exibicao AS profissional_ordem,
            ta.nome AS tipo_atendimento_nome,
            ta.cor AS tipo_atendimento_cor,
            pc.nome AS procedimento_catalogo_nome,
            pc.cor_opcional AS procedimento_cor,
            pc.duracao_padrao_minutos AS procedimento_duracao_padrao
        FROM agendamentos ag
        LEFT JOIN profissionais pr ON pr.id = ag.profissional_id
        LEFT JOIN tipos_atendimento ta ON ta.id = ag.tipo_atendimento_id
        LEFT JOIN procedimentos pc ON pc.id = ag.procedimento_id
        ORDER BY COALESCE(ag.data_agendamento, ag.data), ag.hora_inicio, ag.id
        """,
        conn,
    )

    if agendamentos.empty:
        return agendamentos

    agendamentos["data_ref"] = agendamentos["data_agendamento"].apply(parse_data_contrato)
    agendamentos.loc[agendamentos["data_ref"].isna(), "data_ref"] = agendamentos.loc[
        agendamentos["data_ref"].isna(), "data"
    ].apply(parse_data_contrato)
    agendamentos["hora_inicio_fmt"] = agendamentos["hora_inicio"].apply(formatar_hora_agendamento)
    agendamentos["hora_fim_fmt"] = agendamentos["hora_fim"].apply(formatar_hora_agendamento)
    agendamentos["paciente_exibicao"] = agendamentos["nome_paciente_snapshot"].replace("", pd.NA).fillna(
        agendamentos["paciente_nome"]
    ).fillna("Paciente")
    agendamentos["procedimento_exibicao"] = agendamentos["procedimento_nome_snapshot"].replace("", pd.NA).fillna(
        agendamentos["procedimento_catalogo_nome"]
    ).fillna(agendamentos["procedimento"]).fillna("Sem procedimento")
    agendamentos["profissional_exibicao"] = agendamentos["profissional_nome_cadastro"].replace("", pd.NA).fillna(
        agendamentos["profissional"]
    ).fillna("Sem profissional")
    agendamentos["duracao_real"] = pd.to_numeric(agendamentos["duracao_minutos"], errors="coerce").fillna(0).astype(int)
    agendamentos.loc[agendamentos["duracao_real"] <= 0, "duracao_real"] = 60
    agendamentos["cor_profissional"] = agendamentos["profissional_cor"].apply(lambda valor: normalizar_cor_interface(valor, "#c5a77a"))
    agendamentos["cor_tipo"] = agendamentos["tipo_atendimento_cor"].apply(lambda valor: normalizar_cor_interface(valor, "#c5a77a"))
    agendamentos["cor_dia"] = agendamentos["cor_tipo"]
    agendamentos["cor_semana"] = agendamentos["cor_profissional"]
    agendamentos["ordem_profissional"] = pd.to_numeric(agendamentos["profissional_ordem"], errors="coerce").fillna(9999)

    if data_inicial:
        agendamentos = agendamentos[agendamentos["data_ref"] >= data_inicial].copy()
    if data_final:
        agendamentos = agendamentos[agendamentos["data_ref"] <= data_final].copy()
    if profissionais_ids:
        ids = {int(valor) for valor in profissionais_ids}
        agendamentos = agendamentos[
            pd.to_numeric(agendamentos["profissional_id"], errors="coerce").fillna(-1).astype(int).isin(ids)
        ].copy()

    return agendamentos.sort_values(
        ["data_ref", "ordem_profissional", "profissional_exibicao", "hora_inicio_fmt", "id"]
    )


def data_para_estado_agenda(valor):
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, date):
        return valor
    data_convertida = parse_data_contrato(valor)
    return data_convertida or date.today()


def valor_query_param(nome, padrao=""):
    try:
        valor = st.query_params.get(nome, padrao)
    except Exception:
        return padrao
    if isinstance(valor, list):
        return valor[0] if valor else padrao
    return valor or padrao


def limpar_query_params_agenda():
    try:
        for chave in ["agenda_acao", "agenda_evento", "agenda_data", "agenda_hora", "agenda_hora_fim", "agenda_profissional", "agenda_view"]:
            if chave in st.query_params:
                del st.query_params[chave]
    except Exception:
        pass


def sincronizar_estado_agenda_por_query():
    data_query = valor_query_param("agenda_data", "")
    view_query = valor_query_param("agenda_view", "")
    if data_query:
        st.session_state["agenda_data_referencia"] = data_para_estado_agenda(data_query)
    if view_query in {"Dia", "Semana", "Mes"}:
        st.session_state["agenda_visualizacao"] = view_query


def obter_profissionais_agenda():
    profissionais = carregar_profissionais(somente_ativos=False)
    if profissionais.empty:
        return profissionais
    if "ordem_exibicao" not in profissionais.columns:
        profissionais["ordem_exibicao"] = 9999
    profissionais["ordem_exibicao"] = pd.to_numeric(profissionais["ordem_exibicao"], errors="coerce").fillna(9999)
    profissionais["cor"] = profissionais["cor"].apply(lambda valor: normalizar_cor_interface(valor, "#c5a77a"))
    return profissionais.sort_values(["ativo", "ordem_exibicao", "nome"], ascending=[False, True, True]).reset_index(drop=True)


def obter_tipos_agenda():
    tipos = carregar_tipos_atendimento(somente_ativos=False)
    if tipos.empty:
        return tipos
    tipos["cor"] = tipos["cor"].apply(lambda valor: normalizar_cor_interface(valor, "#c5a77a"))
    return tipos


def obter_procedimentos_agenda():
    procedimentos = carregar_procedimentos_catalogo(somente_ativos=False)
    if procedimentos.empty:
        return procedimentos
    procedimentos["cor_opcional"] = procedimentos["cor_opcional"].apply(lambda valor: normalizar_cor_interface(valor, "#c5a77a"))
    procedimentos["duracao_padrao_minutos"] = pd.to_numeric(
        procedimentos["duracao_padrao_minutos"], errors="coerce"
    ).fillna(60).astype(int)
    procedimentos["valor_padrao"] = pd.to_numeric(procedimentos["valor_padrao"], errors="coerce").fillna(0.0)
    return procedimentos


def garantir_tipos_atendimento_padrao():
    nomes_padrao = ["Avaliação", "Cirurgia", "Consulta", "Emergência", "Periódico", "Retorno"]
    existentes = set()
    if tabela_existe("tipos_atendimento"):
        rows = cursor.execute("SELECT nome FROM tipos_atendimento").fetchall()
        existentes = {str(row["nome"] or "").strip().lower() for row in rows}
    for ordem, nome in enumerate(nomes_padrao, start=1):
        if nome.lower() not in existentes:
            cursor.execute(
                """
                INSERT INTO tipos_atendimento (nome, cor, ativo, ordem_exibicao, criado_em, atualizado_em)
                VALUES (?, ?, 1, ?, ?, ?)
                """,
                (nome, "#c5a77a", ordem, agora_str(), agora_str()),
            )


def salvar_paciente_rapido(nome, telefone, email):
    nome = str(nome or "").strip()
    if not nome:
        raise ValueError("Informe o nome do paciente.")
    cursor.execute(
        """
        INSERT INTO pacientes_rapidos (nome, telefone, email, criado_em)
        VALUES (?, ?, ?, ?)
        """,
        (nome, str(telefone or "").strip(), str(email or "").strip(), agora_str()),
    )
    return cursor.lastrowid


def carregar_agendamentos_profissional_dia(data_ref, profissional_id, excluir_agendamento_id=None):
    agenda = carregar_agendamentos_clinica(data_ref, data_ref, [profissional_id] if profissional_id else None)
    if excluir_agendamento_id and not agenda.empty:
        agenda = agenda[agenda["id"] != int(excluir_agendamento_id)].copy()
    return agenda


def faixa_horarios_agendamento(hora_inicio, hora_fim):
    inicio = hora_para_minutos(hora_inicio)
    fim = hora_para_minutos(hora_fim)
    if inicio is None or fim is None or fim <= inicio:
        return []
    horarios = []
    atual = inicio
    while atual < fim:
        horarios.append(f"{atual // 60:02d}:{atual % 60:02d}")
        atual += 15
    return horarios


def tem_conflito_agendamento(data_ref, profissional_id, hora_inicio, hora_fim, excluir_agendamento_id=None):
    agenda = carregar_agendamentos_profissional_dia(data_ref, profissional_id, excluir_agendamento_id)
    if agenda.empty:
        return False
    inicio_novo = hora_para_minutos(hora_inicio)
    fim_novo = hora_para_minutos(hora_fim)
    for _, item in agenda.iterrows():
        inicio_existente = hora_para_minutos(item["hora_inicio_fmt"])
        fim_existente = hora_para_minutos(item["hora_fim_fmt"]) or ((inicio_existente or 0) + int(item["duracao_real"] or 60))
        if inicio_existente is None or fim_existente is None:
            continue
        if inicio_novo < fim_existente and fim_novo > inicio_existente:
            return True
    return False


def cor_status_agendamento(status):
    mapa = {
        "Agendado": ("#f6efe3", "#8c6a34"),
        "Confirmado": ("#eef8f0", "#2f6a3d"),
        "Em espera": ("#f5f3ff", "#5b3fa3"),
        "Em atendimento": ("#fff6e8", "#9a5b11"),
        "Atendido": ("#edf8f1", "#2f6a3d"),
        "Atrasado": ("#fff2f2", "#b33d3d"),
        "Faltou": ("#fcecec", "#a73a3a"),
        "Cancelado": ("#f1f1f1", "#6b7280"),
    }
    return mapa.get(str(status or "Agendado"), ("#f6efe3", "#8c6a34"))


def garantir_modelos_mensagem_iniciais():
    modelos_padrao = {
        "confirmacao_agendamento": {
            "titulo": "Confirmação de agendamento",
            "conteudo": "Olá, {paciente}. Aqui é da {clinica}. Seu agendamento está confirmado para {data} às {hora} com {profissional}.",
        },
        "lembrete_1_dia_antes": {
            "titulo": "Lembrete 1 dia antes",
            "conteudo": "Olá, {paciente}. Lembrando seu atendimento na {clinica} amanhã, {data}, às {hora} com {profissional}.",
        },
        "lembrete_mesmo_dia": {
            "titulo": "Lembrete no mesmo dia",
            "conteudo": "Olá, {paciente}. Seu atendimento na {clinica} é hoje, {data}, às {hora} com {profissional}.",
        },
    }
    for tipo, dados in modelos_padrao.items():
        existente = cursor.execute(
            "SELECT id FROM modelos_mensagem_agendamento WHERE tipo_mensagem=? LIMIT 1",
            (tipo,),
        ).fetchone()
        if existente:
            continue
        cursor.execute(
            """
            INSERT INTO modelos_mensagem_agendamento (tipo_mensagem, titulo, conteudo, ativo, criado_em, atualizado_em)
            VALUES (?, ?, ?, 1, ?, ?)
            """,
            (tipo, dados["titulo"], dados["conteudo"], agora_str(), agora_str()),
        )


def carregar_modelos_mensagem_agendamento():
    garantir_modelos_mensagem_iniciais()
    return pd.read_sql(
        """
        SELECT *
        FROM modelos_mensagem_agendamento
        WHERE COALESCE(ativo, 1)=1
        ORDER BY id
        """,
        conn,
    )


def salvar_modelo_mensagem_agendamento(tipo_mensagem, titulo, conteudo):
    existente = cursor.execute(
        "SELECT id FROM modelos_mensagem_agendamento WHERE tipo_mensagem=? LIMIT 1",
        (tipo_mensagem,),
    ).fetchone()
    if existente:
        cursor.execute(
            """
            UPDATE modelos_mensagem_agendamento
            SET titulo=?, conteudo=?, atualizado_em=?
            WHERE tipo_mensagem=?
            """,
            (str(titulo or "").strip(), str(conteudo or "").strip(), agora_str(), tipo_mensagem),
        )
    else:
        cursor.execute(
            """
            INSERT INTO modelos_mensagem_agendamento (tipo_mensagem, titulo, conteudo, ativo, criado_em, atualizado_em)
            VALUES (?, ?, ?, 1, ?, ?)
            """,
            (tipo_mensagem, str(titulo or "").strip(), str(conteudo or "").strip(), agora_str(), agora_str()),
        )


def telefone_para_wa(telefone):
    digitos = re.sub(r"\D", "", str(telefone or ""))
    if not digitos:
        return ""
    if digitos.startswith("55"):
        return digitos
    return f"55{digitos}"


def preencher_variaveis_mensagem(texto, agendamento_row):
    paciente = str(agendamento_row.get("paciente_exibicao", "") or "Paciente").strip()
    data_txt = formatar_data_br_valor(agendamento_row.get("data_ref") or agendamento_row.get("data_agendamento") or agendamento_row.get("data") or "")
    hora_txt = str(agendamento_row.get("hora_inicio_fmt", "") or formatar_hora_agendamento(agendamento_row.get("hora_inicio", ""))).strip()
    profissional = str(agendamento_row.get("profissional_exibicao", "") or "Profissional").strip()
    return (
        str(texto or "")
        .replace("{paciente}", paciente)
        .replace("{data}", data_txt)
        .replace("{hora}", hora_txt)
        .replace("{profissional}", profissional)
        .replace("{clinica}", "SoulSul Clínica Integrada")
    )


def registrar_lembrete_agendamento(agendamento_id, tipo_lembrete, canal, mensagem, status_envio, enviado=False):
    cursor.execute(
        """
        INSERT INTO lembretes_agendamento (agendamento_id, tipo_lembrete, canal, mensagem, status_envio, criado_em, enviado_em)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            int(agendamento_id),
            str(tipo_lembrete or "").strip(),
            str(canal or "").strip(),
            str(mensagem or "").strip(),
            str(status_envio or "").strip(),
            agora_str(),
            agora_str() if enviado else "",
        ),
    )


def historico_lembretes_agendamento(agendamento_id):
    return pd.read_sql(
        """
        SELECT *
        FROM lembretes_agendamento
        WHERE agendamento_id=?
        ORDER BY id DESC
        """,
        conn,
        params=(int(agendamento_id),),
    )


def renderizar_grade_selecao_horario(data_ref, profissional_row, agendamento_atual, hora_inicio, hora_fim):
    horarios = gerar_horarios_intervalo("07:00", "20:00", 15)
    if not horarios:
        return
    profissional_id = int(profissional_row["id"]) if profissional_row is not None else None
    ocupados = set()
    agenda = carregar_agendamentos_profissional_dia(
        data_ref,
        profissional_id,
        excluir_agendamento_id=int(agendamento_atual["id"]) if agendamento_atual is not None else None,
    )
    for _, item in agenda.iterrows():
        ocupados.update(faixa_horarios_agendamento(item["hora_inicio_fmt"], item["hora_fim_fmt"]))

    selecionados = set(faixa_horarios_agendamento(hora_inicio, hora_fim))
    caixas = []
    for horario in horarios[:-1]:
        if profissional_row is not None and not profissional_atende_no_horario(profissional_row, data_ref, horario):
            classe = "blocked"
        elif horario in selecionados:
            classe = "selected"
        elif horario in ocupados:
            classe = "occupied"
        else:
            classe = "free"

        href = (
            f"?agenda_acao=novo&agenda_data={quote(formatar_data_br(data_ref))}&agenda_hora={quote(hora_inicio)}"
            f"&agenda_hora_fim={quote(adicionar_minutos_hora(horario, 15))}&agenda_profissional={profissional_id or ''}&agenda_view=Dia"
        )
        if agendamento_atual is not None:
            href = (
                f"?agenda_evento={int(agendamento_atual['id'])}&agenda_data={quote(formatar_data_br(data_ref))}"
                f"&agenda_hora_fim={quote(adicionar_minutos_hora(horario, 15))}&agenda_view=Dia"
            )
        caixas.append(
            f"<a class='ss-ag-slot-picker {classe}' href='{href}'>{horario}</a>"
        )
    st.markdown(f"<div class='ss-ag-slot-picker-wrap'>{''.join(caixas)}</div>", unsafe_allow_html=True)


def buscar_contratos_paciente_agenda(paciente_id):
    if not paciente_id:
        return pd.DataFrame()
    contratos = pd.read_sql(
        """
        SELECT *
        FROM contratos
        WHERE paciente_id=?
        ORDER BY COALESCE(data_criacao, id) DESC, id DESC
        """,
        conn,
        params=(int(paciente_id),),
    )
    return contratos


def buscar_procedimentos_contratados_em_aberto(paciente_id):
    contratos = buscar_contratos_paciente_agenda(paciente_id)
    if contratos.empty:
        return pd.DataFrame()

    procedimentos = pd.read_sql(
        """
        SELECT pc.id, pc.contrato_id, pc.procedimento, pc.valor, c.paciente_id
        FROM procedimentos_contrato pc
        JOIN contratos c ON c.id = pc.contrato_id
        WHERE c.paciente_id=?
        ORDER BY pc.contrato_id DESC, pc.id
        """,
        conn,
        params=(int(paciente_id),),
    )
    if procedimentos.empty:
        return pd.DataFrame()

    usados = pd.read_sql(
        """
        SELECT contrato_id, procedimento_nome_snapshot, COUNT(*) AS usados
        FROM agendamento_procedimentos
        WHERE contrato_id IS NOT NULL AND COALESCE(origem_contrato, 0) <> 0
        GROUP BY contrato_id, procedimento_nome_snapshot
        """,
        conn,
    )
    usados_lookup = {}
    if not usados.empty:
        for _, row in usados.iterrows():
            chave = (int(row["contrato_id"]), normalizar_texto(row["procedimento_nome_snapshot"]))
            usados_lookup[chave] = int(row["usados"] or 0)

    grupos = []
    agrupado = (
        procedimentos.assign(_proc_norm=procedimentos["procedimento"].apply(normalizar_texto))
        .groupby(["contrato_id", "procedimento", "_proc_norm"], as_index=False)
        .agg(sessoes_total=("id", "count"), valor_total=("valor", "sum"))
    )
    procedimentos_catalogo = obter_procedimentos_agenda()
    for _, row in agrupado.iterrows():
        chave = (int(row["contrato_id"]), row["_proc_norm"])
        usados_qtd = usados_lookup.get(chave, 0)
        restantes = max(int(row["sessoes_total"]) - usados_qtd, 0)
        if restantes <= 0:
            continue
        duracao_padrao = 60
        procedimento_id = None
        if not procedimentos_catalogo.empty:
            match = procedimentos_catalogo[
                procedimentos_catalogo["nome"].apply(normalizar_texto) == row["_proc_norm"]
            ]
            if not match.empty:
                procedimento_id = int(match.iloc[0]["id"])
                duracao_padrao = int(match.iloc[0]["duracao_padrao_minutos"] or 60)
        grupos.append(
            {
                "contrato_id": int(row["contrato_id"]),
                "procedimento_nome": row["procedimento"],
                "sessoes_total": int(row["sessoes_total"]),
                "sessoes_restantes": restantes,
                "sessoes_usadas": usados_qtd,
                "procedimento_id": procedimento_id,
                "duracao_padrao": duracao_padrao,
                "origem_contrato": 1,
            }
        )
    return pd.DataFrame(grupos)


def salvar_profissional_agenda(nome, cor, ativo, especialidade, ordem_exibicao, dias_atendimento, hora_inicio, hora_fim, profissional_id=None):
    nome = str(nome or "").strip()
    if not nome:
        raise ValueError("Informe o nome do profissional.")
    cor_normalizada = normalizar_cor_interface(cor, "#c5a77a")
    dias_serializados = serializar_dias_atendimento(dias_atendimento)
    if profissional_id:
        cursor.execute(
            """
            UPDATE profissionais
            SET nome=?, cor=?, ativo=?, especialidade=?, ordem_exibicao=?, atualizado_em=?,
                dias_atendimento=?, hora_inicio=?, hora_fim=?
            WHERE id=?
            """,
            (
                nome,
                cor_normalizada,
                int(bool(ativo)),
                str(especialidade or "").strip(),
                int(ordem_exibicao or 0),
                agora_str(),
                dias_serializados,
                formatar_hora_agendamento(hora_inicio),
                formatar_hora_agendamento(hora_fim),
                int(profissional_id),
            ),
        )
    else:
        cursor.execute(
            """
            INSERT INTO profissionais
            (nome, cor, ativo, especialidade, ordem_exibicao, criado_em, atualizado_em, dias_atendimento, hora_inicio, hora_fim)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                nome,
                cor_normalizada,
                int(bool(ativo)),
                str(especialidade or "").strip(),
                int(ordem_exibicao or 0),
                agora_str(),
                agora_str(),
                dias_serializados,
                formatar_hora_agendamento(hora_inicio),
                formatar_hora_agendamento(hora_fim),
            ),
        )


def salvar_tipo_atendimento(nome, cor, ativo, ordem_exibicao, tipo_id=None):
    nome = str(nome or "").strip()
    if not nome:
        raise ValueError("Informe o nome do tipo de atendimento.")
    if tipo_id:
        cursor.execute(
            """
            UPDATE tipos_atendimento
            SET nome=?, cor=?, ativo=?, ordem_exibicao=?, atualizado_em=?
            WHERE id=?
            """,
            (nome, normalizar_cor_interface(cor, "#c5a77a"), int(bool(ativo)), int(ordem_exibicao or 0), agora_str(), int(tipo_id)),
        )
    else:
        cursor.execute(
            """
            INSERT INTO tipos_atendimento
            (nome, cor, ativo, ordem_exibicao, criado_em, atualizado_em)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (nome, normalizar_cor_interface(cor, "#c5a77a"), int(bool(ativo)), int(ordem_exibicao or 0), agora_str(), agora_str()),
        )


def salvar_procedimento_catalogo(nome, categoria, valor_padrao, duracao_padrao_minutos, descricao, cor_opcional, ativo, procedimento_id=None):
    nome = str(nome or "").strip()
    if not nome:
        raise ValueError("Informe o nome do procedimento.")
    if procedimento_id:
        cursor.execute(
            """
            UPDATE procedimentos
            SET nome=?, categoria=?, valor_padrao=?, duracao_padrao_minutos=?, descricao=?, cor_opcional=?, ativo=?, atualizado_em=?
            WHERE id=?
            """,
            (
                nome,
                str(categoria or "").strip(),
                float(valor_padrao or 0),
                int(duracao_padrao_minutos or 60),
                str(descricao or "").strip(),
                normalizar_cor_interface(cor_opcional, "#c5a77a"),
                int(bool(ativo)),
                agora_str(),
                int(procedimento_id),
            ),
        )
    else:
        cursor.execute(
            """
            INSERT INTO procedimentos
            (nome, categoria, valor_padrao, duracao_padrao_minutos, descricao, cor_opcional, ativo, criado_em, atualizado_em)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                nome,
                str(categoria or "").strip(),
                float(valor_padrao or 0),
                int(duracao_padrao_minutos or 60),
                str(descricao or "").strip(),
                normalizar_cor_interface(cor_opcional, "#c5a77a"),
                int(bool(ativo)),
                agora_str(),
                agora_str(),
            ),
        )


def salvar_agendamento_completo(
    paciente_id,
    nome_paciente,
    telefone_snapshot,
    email_snapshot,
    profissional_id,
    tipo_atendimento_id,
    procedimento_id,
    contrato_id,
    origem_contrato,
    data_agendamento,
    hora_inicio,
    duracao_minutos,
    observacoes,
    status,
    criado_por,
    agendamento_id=None,
    procedimento_nome_override="",
):
    data_ref = data_para_estado_agenda(data_agendamento)
    hora_inicio_fmt = formatar_hora_agendamento(hora_inicio)
    if not hora_inicio_fmt:
        raise ValueError("Informe um horário inicial válido.")

    duracao = max(int(duracao_minutos or 15), 15)
    hora_fim_fmt = adicionar_minutos_hora(hora_inicio_fmt, duracao)

    profissional_row = cursor.execute("SELECT * FROM profissionais WHERE id=?", (int(profissional_id),)).fetchone() if profissional_id else None
    procedimento_row = cursor.execute("SELECT * FROM procedimentos WHERE id=?", (int(procedimento_id),)).fetchone() if procedimento_id else None
    nome_paciente = str(nome_paciente or "").strip() or "Paciente"
    profissional_nome = profissional_row["nome"] if profissional_row and profissional_row["nome"] else ""
    procedimento_nome = (
        str(procedimento_nome_override or "").strip()
        or (procedimento_row["nome"] if procedimento_row and procedimento_row["nome"] else "")
    )

    dados_base = (
        int(paciente_id) if paciente_id else None,
        nome_paciente,
        str(telefone_snapshot or "").strip(),
        str(email_snapshot or "").strip(),
        int(profissional_id) if profissional_id else None,
        int(tipo_atendimento_id) if tipo_atendimento_id else None,
        int(procedimento_id) if procedimento_id else None,
        procedimento_nome,
        int(contrato_id) if contrato_id else None,
        str(origem_contrato or "").strip(),
        formatar_data_br(data_ref),
        hora_inicio_fmt,
        hora_fim_fmt,
        duracao,
        str(observacoes or "").strip(),
        str(status or "Agendado").strip() or "Agendado",
        str(criado_por or "").strip(),
    )

    if agendamento_id:
        cursor.execute(
            """
            UPDATE agendamentos
            SET paciente_id=?, nome_paciente_snapshot=?, telefone_snapshot=?, email_snapshot=?, profissional_id=?,
                tipo_atendimento_id=?, procedimento_id=?, procedimento_nome_snapshot=?, contrato_id=?, origem_contrato=?,
                data_agendamento=?, hora_inicio=?, hora_fim=?, duracao_minutos=?, observacoes=?, status=?, criado_por=?,
                atualizado_em=?, data=?, paciente_nome=?, profissional=?, procedimento=?, observacao=?
            WHERE id=?
            """,
            (
                *dados_base,
                agora_str(),
                formatar_data_br(data_ref),
                nome_paciente,
                profissional_nome,
                procedimento_nome,
                str(observacoes or "").strip(),
                int(agendamento_id),
            ),
        )
        cursor.execute("DELETE FROM agendamento_procedimentos WHERE agendamento_id=?", (int(agendamento_id),))
        agendamento_id_final = int(agendamento_id)
    else:
        cursor.execute(
            """
            INSERT INTO agendamentos
            (
                paciente_id, nome_paciente_snapshot, telefone_snapshot, email_snapshot, profissional_id,
                tipo_atendimento_id, procedimento_id, procedimento_nome_snapshot, contrato_id, origem_contrato,
                data_agendamento, hora_inicio, hora_fim, duracao_minutos, observacoes, status, criado_por,
                criado_em, atualizado_em, data, paciente_nome, profissional, procedimento, observacao, data_criacao
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                *dados_base,
                agora_str(),
                agora_str(),
                formatar_data_br(data_ref),
                nome_paciente,
                profissional_nome,
                procedimento_nome,
                str(observacoes or "").strip(),
                agora_str(),
            ),
        )
        agendamento_id_final = cursor.lastrowid

    if procedimento_id:
        cursor.execute(
            """
            INSERT INTO agendamento_procedimentos
            (agendamento_id, procedimento_id, procedimento_nome_snapshot, valor_snapshot, duracao_snapshot_minutos, origem_contrato, contrato_id)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                int(agendamento_id_final),
                int(procedimento_id),
                procedimento_nome,
                float(procedimento_row["valor_padrao"] or 0) if procedimento_row else 0.0,
                duracao,
                str(origem_contrato or "").strip(),
                int(contrato_id) if contrato_id else None,
            ),
        )
    return agendamento_id_final


def excluir_agendamento_completo(agendamento_id):
    cursor.execute("DELETE FROM agendamento_procedimentos WHERE agendamento_id=?", (int(agendamento_id),))
    cursor.execute("DELETE FROM lembretes_agendamento WHERE agendamento_id=?", (int(agendamento_id),))
    cursor.execute("DELETE FROM agendamentos WHERE id=?", (int(agendamento_id),))


def renderizar_html_interface(html):
    conteudo = str(html or "").strip()
    if not conteudo:
        return
    if hasattr(st, "html"):
        st.html(conteudo)
    else:
        st.markdown(conteudo, unsafe_allow_html=True)


def html_cores_profissionais(profissionais_df, profissionais_ids):
    if profissionais_df.empty:
        return
    ids = {int(valor) for valor in profissionais_ids}
    itens = []
    for _, row in profissionais_df.iterrows():
        if int(row.get("ativo", 1) or 1) != 1:
            continue
        itens.append(
            f"""
            <div class="ss-ag-prof-item" style="opacity:{'1' if int(row['id']) in ids else '0.45'};">
                <span class="ss-ag-prof-dot" style="background:{normalizar_cor_interface(row.get('cor'), '#c5a77a')};"></span>
                <span>{row.get('nome', 'Profissional')}</span>
            </div>
            """
        )
    if itens:
        renderizar_html_interface(f"<div class='ss-ag-prof-list'>{''.join(itens)}</div>")


def gerar_html_grade_dia(data_ref, profissionais_df, agendamentos_df):
    horarios = gerar_horarios_intervalo("07:00", "20:00", 15)
    if not horarios:
        return ""
    altura_slot = 28
    altura_grade = (len(horarios) - 1) * altura_slot
    hora_minima = hora_para_minutos(horarios[0])

    time_labels = "".join(
        f"<div class='ss-ag-time-label' style='top:{indice * altura_slot}px'>{horario}</div>"
        for indice, horario in enumerate(horarios[:-1])
    )
    colunas = [
        f"""
        <div class="ss-ag-time-col">
            <div class="ss-ag-day-col-head">Horário</div>
            <div class="ss-ag-day-col-body" style="height:{altura_grade}px;">
                {time_labels}
            </div>
        </div>
        """
    ]

    agora_local = datetime.now()
    exibir_agora = data_ref == agora_local.date()
    topo_agora = None
    if exibir_agora:
        minutos_agora = agora_local.hour * 60 + agora_local.minute
        if hora_minima <= minutos_agora <= hora_para_minutos(horarios[-1]):
            topo_agora = int(((minutos_agora - hora_minima) / 15) * altura_slot)

    for _, profissional in profissionais_df.iterrows():
        profissional_id = int(profissional["id"])
        cor_prof = normalizar_cor_interface(profissional.get("cor"), "#c5a77a")
        slots = []
        for indice, horario in enumerate(horarios[:-1]):
            if not profissional_atende_no_horario(profissional, data_ref, horario):
                continue
            slots.append(
                f"""
                <a class="ss-ag-slot-link"
                   href="?agenda_acao=novo&agenda_data={quote(formatar_data_br(data_ref))}&agenda_hora={quote(horario)}&agenda_profissional={profissional_id}&agenda_view=Dia"
                   style="top:{indice * altura_slot}px;height:{altura_slot}px;"
                   title="Novo agendamento às {horario}"></a>
                """
            )

        eventos = []
        agenda_prof = agendamentos_df[agendamentos_df["profissional_id"].fillna(-1).astype(int) == profissional_id].copy()
        for _, item in agenda_prof.iterrows():
            inicio = hora_para_minutos(item["hora_inicio_fmt"])
            fim = hora_para_minutos(item["hora_fim_fmt"]) or ((inicio or 0) + int(item["duracao_real"] or 60))
            if inicio is None:
                continue
            topo = max(int(((inicio - hora_minima) / 15) * altura_slot), 0)
            altura = max(int(max(fim - inicio, 15) / 15 * altura_slot), altura_slot)
            cor_bloco = normalizar_cor_interface(item.get("cor_dia"), cor_prof)
            eventos.append(
                f"""
                <a class="ss-ag-event"
                   href="?agenda_evento={int(item['id'])}&agenda_data={quote(formatar_data_br(data_ref))}&agenda_view=Dia"
                   style="top:{topo}px;height:{altura}px;background:{hex_para_rgba(cor_bloco, 0.22)};border-left-color:{cor_bloco};">
                    <div class="ss-ag-event-time">{item['hora_inicio_fmt']} - {item['hora_fim_fmt']}</div>
                    <div class="ss-ag-event-title">{item['paciente_exibicao']}</div>
                    <div class="ss-ag-event-sub">{item['procedimento_exibicao']}</div>
                </a>
                """
            )

        linha_agora = f"<div class='ss-ag-now-line' style='top:{topo_agora}px;'></div>" if topo_agora is not None else ""
        colunas.append(
            f"""
            <div class="ss-ag-day-col">
                <div class="ss-ag-day-col-head">
                    <span class="ss-ag-prof-chip" style="background:{hex_para_rgba(cor_prof, 0.16)};border-color:{hex_para_rgba(cor_prof, 0.4)};">
                        <span class="ss-ag-prof-dot" style="background:{cor_prof};"></span>{profissional.get('nome', 'Profissional')}
                    </span>
                </div>
                <div class="ss-ag-day-col-body" style="height:{altura_grade}px;">
                    {''.join(slots)}
                    {linha_agora}
                    {''.join(eventos)}
                </div>
            </div>
            """
        )

    return f"<div class='ss-ag-day-wrap'>{''.join(colunas)}</div>"


def renderizar_agenda_semana(data_base, agendamentos_df):
    inicio = inicio_semana(data_base)
    cabecalho = st.columns(7)
    for idx, coluna in enumerate(cabecalho):
        dia = inicio + timedelta(days=idx)
        total = len(agendamentos_df[agendamentos_df["data_ref"] == dia])
        coluna.markdown(
            f"""
            <div class="ss-ag-week-head">
                <div class="ss-ag-week-day">{DIAS_SEMANA_PT[idx]}</div>
                <div class="ss-ag-week-date">{formatar_data_br(dia)}</div>
                <div class="ss-ag-week-total">{total} atendimento(s)</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    colunas = st.columns(7)
    for idx, coluna in enumerate(colunas):
        dia = inicio + timedelta(days=idx)
        agenda_dia = agendamentos_df[agendamentos_df["data_ref"] == dia].copy()
        agenda_dia = agenda_dia.sort_values(["hora_inicio_fmt", "profissional_exibicao"])
        with coluna:
            if agenda_dia.empty:
                renderizar_html_interface("<div class='ss-ag-week-empty'>Sem atendimentos</div>")
            else:
                for _, item in agenda_dia.iterrows():
                    cor = normalizar_cor_interface(item.get("cor_semana"), "#c5a77a")
                    renderizar_html_interface(
                        f"""
                        <a class="ss-ag-week-event" href="?agenda_evento={int(item['id'])}&agenda_data={quote(formatar_data_br(dia))}&agenda_view=Semana"
                           style="border-left-color:{cor};background:{hex_para_rgba(cor, 0.18)};">
                            <div class="ss-ag-week-time">{item['hora_inicio_fmt']} - {item['hora_fim_fmt']}</div>
                            <div class="ss-ag-week-title">{item['paciente_exibicao']}</div>
                            <div class="ss-ag-week-sub">{item['profissional_exibicao']}</div>
                        </a>
                        """
                    )


def renderizar_agenda_mes(data_base, agendamentos_df):
    ano = data_base.year
    mes = data_base.month
    hoje = date.today()
    cab = st.columns(7)
    for coluna, nome_dia in zip(cab, DIAS_SEMANA_PT):
        coluna.markdown(f"**{nome_dia}**")
    for semana in calendar.monthcalendar(ano, mes):
        cols = st.columns(7)
        for idx, dia in enumerate(semana):
            with cols[idx]:
                if dia == 0:
                    st.write("")
                    continue
                data_dia = date(ano, mes, dia)
                agenda_dia = agendamentos_df[agendamentos_df["data_ref"] == data_dia].copy()
                borda = "2px solid #c5a77a" if data_dia == hoje else "1px solid #e5dfd3"
                fundo = "#ffffff" if agenda_dia.empty else "#f8f3eb"
                renderizar_html_interface(
                    f"""
                    <a class="ss-ag-month-day" href="?agenda_data={quote(formatar_data_br(data_dia))}&agenda_view=Dia"
                       style="border:{borda};background:{fundo};">
                        <div class="ss-ag-month-num">{dia}</div>
                        <div class="ss-ag-month-total">{len(agenda_dia)} atendimento(s)</div>
                    </a>
                    """
                )


def renderizar_formulario_agendamento(usuario_logado, data_ref, profissionais_df, tipos_df, procedimentos_df, pacientes_df):
    garantir_tipos_atendimento_padrao()
    if tipos_df.empty:
        tipos_df = obter_tipos_agenda()
    acao = valor_query_param("agenda_acao", "")
    evento_id = valor_query_param("agenda_evento", "")
    horario_query = valor_query_param("agenda_hora", "")
    hora_fim_query = valor_query_param("agenda_hora_fim", "")
    profissional_query = valor_query_param("agenda_profissional", "")
    if not acao and not evento_id:
        return

    agendamento_atual = None
    if str(evento_id).isdigit():
        agendamentos = carregar_agendamentos_clinica()
        filtro = agendamentos[agendamentos["id"] == int(evento_id)]
        if not filtro.empty:
            agendamento_atual = filtro.iloc[0]

    titulo_modal = "Detalhes do agendamento" if agendamento_atual is not None else "Novo agendamento"

    @st.dialog(titulo_modal, width="large")
    def modal_agendamento():
        if agendamento_atual is not None:
            fundo_status, cor_status = cor_status_agendamento(agendamento_atual["status"])
            marcadores = []
            procedimentos_agendados = pd.read_sql(
                """
                SELECT procedimento_nome_snapshot
                FROM agendamento_procedimentos
                WHERE agendamento_id=?
                ORDER BY id
                """,
                conn,
                params=(int(agendamento_atual["id"]),),
            )
            procedimentos_exibicao = agendamento_atual["procedimento_exibicao"]
            if not procedimentos_agendados.empty:
                lista_procs = [str(valor).strip() for valor in procedimentos_agendados["procedimento_nome_snapshot"].tolist() if str(valor).strip()]
                if lista_procs:
                    procedimentos_exibicao = ", ".join(lista_procs)
            if pd.notna(agendamento_atual.get("contrato_id")) and int(agendamento_atual.get("contrato_id") or 0) > 0:
                marcadores.append("Contrato")
            if str(agendamento_atual.get("origem_contrato", "")).strip() not in {"", "0", "False", "false"}:
                marcadores.append("Procedimento contratado")
            if not marcadores:
                marcadores.append("Manual")

            indicador_financeiro = "Sem vínculo financeiro"
            if pd.notna(agendamento_atual.get("contrato_id")) and int(agendamento_atual.get("contrato_id") or 0) > 0:
                recebiveis_vinculados = cursor.execute(
                    "SELECT COUNT(*) AS total FROM recebiveis WHERE contrato_id=?",
                    (int(agendamento_atual["contrato_id"]),),
                ).fetchone()
                indicador_financeiro = f"Recebíveis vinculados: {int(recebiveis_vinculados['total'] or 0)}"

            st.markdown(
                f"""
                <div class="ss-ag-detail-card">
                    <div class="ss-ag-detail-top">
                        <div>
                            <div class="ss-ag-detail-title">{agendamento_atual['paciente_exibicao']}</div>
                            <div class="ss-ag-detail-sub">{agendamento_atual['procedimento_exibicao']}</div>
                        </div>
                        <span class="ss-ag-status-pill" style="background:{fundo_status};color:{cor_status};border-color:{hex_para_rgba(cor_status, 0.18)};">
                            {agendamento_atual['status']}
                        </span>
                    </div>
                    <div class="ss-ag-detail-grid">
                        <div><strong>Agendado em</strong><br>{formatar_data_hora_br_valor(agendamento_atual.get('criado_em', ''))}</div>
                        <div><strong>Profissional</strong><br>{agendamento_atual['profissional_exibicao']}</div>
                        <div><strong>Paciente / prontuário</strong><br>{agendamento_atual['paciente_exibicao']}</div>
                        <div><strong>Telefone</strong><br>{agendamento_atual.get('telefone_snapshot', '') or '-'}</div>
                        <div><strong>Horário</strong><br>{agendamento_atual['hora_inicio_fmt']} - {agendamento_atual['hora_fim_fmt']}</div>
                        <div><strong>Procedimento(s)</strong><br>{procedimentos_exibicao}</div>
                        <div><strong>Marcadores</strong><br>{', '.join(marcadores)}</div>
                        <div><strong>Indicador financeiro</strong><br>{indicador_financeiro}</div>
                    </div>
                    <div class="ss-ag-detail-obs"><strong>Observações</strong><br>{agendamento_atual.get('observacoes', '') or 'Sem observações'}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            c_status, c_btn = st.columns([2, 1])
            novo_status = c_status.selectbox(
                "Alterar status",
                options=STATUS_AGENDAMENTO,
                index=STATUS_AGENDAMENTO.index(str(agendamento_atual["status"])) if str(agendamento_atual["status"]) in STATUS_AGENDAMENTO else 0,
                key=f"agenda_status_detalhe_{int(agendamento_atual['id'])}",
            )
            if c_btn.button("Atualizar status", use_container_width=True, key=f"agenda_status_btn_{int(agendamento_atual['id'])}"):
                cursor.execute(
                    "UPDATE agendamentos SET status=?, atualizado_em=? WHERE id=?",
                    (novo_status, agora_str(), int(agendamento_atual["id"])),
                )
                conn.commit()
                st.success("Status atualizado.")
                st.rerun()

            a1, a2, a3, a4, a5 = st.columns(5)
            if a1.button("Editar", use_container_width=True, key=f"agenda_edit_btn_{int(agendamento_atual['id'])}"):
                st.session_state[f"agenda_editar_{int(agendamento_atual['id'])}"] = True
                st.rerun()
            if a2.button("Cancelar", use_container_width=True, key=f"agenda_cancelar_btn_{int(agendamento_atual['id'])}"):
                cursor.execute(
                    "UPDATE agendamentos SET status='Cancelado', atualizado_em=? WHERE id=?",
                    (agora_str(), int(agendamento_atual["id"])),
                )
                conn.commit()
                st.success("Agendamento cancelado.")
                st.rerun()
            if a3.button("Remarcar", use_container_width=True, key=f"agenda_remarcar_btn_{int(agendamento_atual['id'])}"):
                st.session_state[f"agenda_editar_{int(agendamento_atual['id'])}"] = True
                st.rerun()
            if a4.button("Abrir paciente", use_container_width=True, key=f"agenda_paciente_btn_{int(agendamento_atual['id'])}") and pd.notna(agendamento_atual.get("paciente_id")):
                limpar_query_params_agenda()
                navegar_para_menu("Editar Paciente", editar_paciente_id=int(agendamento_atual["paciente_id"]))
            if a5.button("Abrir financeiro", use_container_width=True, key=f"agenda_financeiro_btn_{int(agendamento_atual['id'])}"):
                limpar_query_params_agenda()
                navegar_para_menu("Financeiro", financeiro_foco={
                    "paciente_nome": agendamento_atual["paciente_exibicao"],
                    "contrato_id": int(agendamento_atual["contrato_id"]) if pd.notna(agendamento_atual.get("contrato_id")) and int(agendamento_atual.get("contrato_id") or 0) > 0 else None,
                })

            telefone_wa = telefone_para_wa(agendamento_atual.get("telefone_snapshot", ""))
            modelos_msg = carregar_modelos_mensagem_agendamento()
            if not modelos_msg.empty:
                st.markdown("### Mensagens ao paciente")
                tipo_escolhido = st.selectbox(
                    "Modelo de mensagem",
                    options=modelos_msg["tipo_mensagem"].tolist(),
                    format_func=lambda valor: {
                        "confirmacao_agendamento": "Confirmação de agendamento",
                        "lembrete_1_dia_antes": "Lembrete 1 dia antes",
                        "lembrete_mesmo_dia": "Lembrete no mesmo dia",
                    }.get(valor, valor),
                    key=f"agenda_msg_tipo_{int(agendamento_atual['id'])}",
                )
                modelo_row = modelos_msg[modelos_msg["tipo_mensagem"] == tipo_escolhido].iloc[0]
                mensagem_gerada = preencher_variaveis_mensagem(modelo_row["conteudo"], agendamento_atual)
                st.text_area(
                    "Mensagem gerada",
                    value=mensagem_gerada,
                    height=110,
                    key=f"agenda_msg_preview_{int(agendamento_atual['id'])}",
                )
                if telefone_wa:
                    link_wa = f"https://wa.me/{telefone_wa}?text={quote(mensagem_gerada)}"
                    m1, m2 = st.columns(2)
                    if m1.button("Registrar como preparada", use_container_width=True, key=f"agenda_msg_preparada_{int(agendamento_atual['id'])}"):
                        registrar_lembrete_agendamento(int(agendamento_atual["id"]), tipo_escolhido, "WhatsApp Web", mensagem_gerada, "preparada", enviado=False)
                        conn.commit()
                        st.success("Mensagem registrada como preparada.")
                        st.rerun()
                    m2.markdown(
                        f"""
                        <a class="ss-wa-link" href="{link_wa}" target="_blank">
                            Abrir WhatsApp Web
                        </a>
                        """,
                        unsafe_allow_html=True,
                    )
                    if st.button("Registrar como enviada manualmente", use_container_width=True, key=f"agenda_msg_enviada_{int(agendamento_atual['id'])}"):
                        registrar_lembrete_agendamento(int(agendamento_atual["id"]), tipo_escolhido, "WhatsApp Web", mensagem_gerada, "enviada manualmente", enviado=True)
                        conn.commit()
                        st.success("Envio manual registrado.")
                        st.rerun()
                else:
                    st.warning("O paciente não possui telefone válido para WhatsApp.")

                historico = historico_lembretes_agendamento(int(agendamento_atual["id"]))
                if not historico.empty:
                    historico_exibicao = historico[["tipo_lembrete", "canal", "status_envio", "criado_em", "enviado_em"]].copy()
                    historico_exibicao["criado_em"] = historico_exibicao["criado_em"].apply(formatar_data_hora_br_valor)
                    historico_exibicao["enviado_em"] = historico_exibicao["enviado_em"].apply(formatar_data_hora_br_valor)
                    historico_exibicao = historico_exibicao.rename(
                        columns={
                            "tipo_lembrete": "Tipo",
                            "canal": "Canal",
                            "status_envio": "Status",
                            "criado_em": "Criado em",
                            "enviado_em": "Enviado em",
                        }
                    )
                    st.dataframe(historico_exibicao, use_container_width=True, hide_index=True)

            if not st.session_state.get(f"agenda_editar_{int(agendamento_atual['id'])}", False):
                return

        pesquisa_paciente = st.text_input("Pesquisar paciente", value=agendamento_atual["paciente_exibicao"] if agendamento_atual is not None else "", key=f"agenda_pesquisa_{evento_id or 'novo'}")
        pacientes_filtrados = pacientes_df.copy()
        if pesquisa_paciente.strip():
            termo = normalizar_texto(pesquisa_paciente)
            pacientes_filtrados = pacientes_filtrados[
                pacientes_filtrados.apply(
                    lambda row: termo in normalizar_texto(
                        f"{row.get('nome','')} {row.get('prontuario','')} {row.get('telefone','')}"
                    ),
                    axis=1,
                )
            ].copy()
        opcoes_pacientes = []
        for _, paciente in pacientes_filtrados.head(30).iterrows():
            opcoes_pacientes.append(
                (
                    int(paciente["id"]),
                    f"{paciente['nome']} - Prontuário {formatar_prontuario_valor(paciente['prontuario'])}",
                )
            )
        opcoes_pacientes.append(("__novo__", "Não encontrado - Novo"))

        ids_profissionais = profissionais_df["id"].astype(int).tolist()
        ids_tipos = tipos_df["id"].astype(int).tolist() if not tipos_df.empty else []
        ids_procedimentos = procedimentos_df["id"].astype(int).tolist() if not procedimentos_df.empty else []
        paciente_default = int(agendamento_atual["paciente_id"]) if agendamento_atual is not None and pd.notna(agendamento_atual["paciente_id"]) else "__novo__"
        profissional_default = int(agendamento_atual["profissional_id"]) if agendamento_atual is not None and pd.notna(agendamento_atual["profissional_id"]) else (int(profissional_query) if str(profissional_query).isdigit() else (ids_profissionais[0] if ids_profissionais else None))
        tipo_default = int(agendamento_atual["tipo_atendimento_id"]) if agendamento_atual is not None and pd.notna(agendamento_atual["tipo_atendimento_id"]) else (ids_tipos[0] if ids_tipos else None)
        procedimento_default = int(agendamento_atual["procedimento_id"]) if agendamento_atual is not None and pd.notna(agendamento_atual["procedimento_id"]) else (ids_procedimentos[0] if ids_procedimentos else None)
        data_default = data_para_estado_agenda(agendamento_atual["data_ref"] if agendamento_atual is not None else data_ref)
        hora_default = agendamento_atual["hora_inicio_fmt"] if agendamento_atual is not None else (horario_query or "08:00")
        duracao_default = int(agendamento_atual["duracao_real"]) if agendamento_atual is not None else 60
        hora_fim_default = agendamento_atual["hora_fim_fmt"] if agendamento_atual is not None else (hora_fim_query or adicionar_minutos_hora(hora_default, duracao_default))
        lista_horarios = gerar_horarios_intervalo("07:00", "20:00", 15)

        with st.form(f"agenda_form_{evento_id or 'novo'}"):
            paciente_escolha = st.selectbox(
                "Paciente",
                options=[valor for valor, _ in opcoes_pacientes],
                index=[valor for valor, _ in opcoes_pacientes].index(paciente_default) if paciente_default in [valor for valor, _ in opcoes_pacientes] else len(opcoes_pacientes) - 1,
                format_func=lambda valor: next((rotulo for op, rotulo in opcoes_pacientes if op == valor), "Paciente"),
            )

            paciente_row = None
            if paciente_escolha not in (None, "__novo__"):
                filtro = pacientes_df[pacientes_df["id"] == int(paciente_escolha)]
                if not filtro.empty:
                    paciente_row = filtro.iloc[0]

            c1, c2, c3 = st.columns(3)
            nome_paciente = c1.text_input("Nome", value=(agendamento_atual["paciente_exibicao"] if agendamento_atual is not None else "") or (paciente_row["nome"] if paciente_row is not None else ""))
            telefone_snapshot = c2.text_input("Telefone", value=(agendamento_atual["telefone_snapshot"] if agendamento_atual is not None else "") or (paciente_row["telefone"] if paciente_row is not None else ""))
            email_snapshot = c3.text_input("E-mail", value=agendamento_atual["email_snapshot"] if agendamento_atual is not None else "")

            procedimentos_contratados_df = buscar_procedimentos_contratados_em_aberto(int(paciente_escolha)) if paciente_escolha not in (None, "__novo__") else pd.DataFrame()
            opcoes_contrato = ["Manual"]
            mapa_contrato = {"Manual": None}
            if not procedimentos_contratados_df.empty:
                for _, proc in procedimentos_contratados_df.iterrows():
                    rotulo = (
                        f"{proc['procedimento_nome']} | Contrato {proc['contrato_id']} | "
                        f"Sessões {proc['sessoes_total']} | Restantes {proc['sessoes_restantes']}"
                    )
                    opcoes_contrato.append(rotulo)
                    mapa_contrato[rotulo] = proc.to_dict()

            origem_padrao = "Manual"
            if agendamento_atual is not None and pd.notna(agendamento_atual["contrato_id"]) and int(agendamento_atual["contrato_id"] or 0) > 0:
                for rotulo, info in mapa_contrato.items():
                    if info and int(info["contrato_id"]) == int(agendamento_atual["contrato_id"]) and normalizar_texto(info["procedimento_nome"]) == normalizar_texto(agendamento_atual["procedimento_exibicao"]):
                        origem_padrao = rotulo
                        break
            origem_selecionada = st.selectbox(
                "Procedimento do contrato",
                options=opcoes_contrato,
                index=opcoes_contrato.index(origem_padrao) if origem_padrao in opcoes_contrato else 0,
                help="Selecione um procedimento contratado em aberto ou mantenha Manual.",
            )
            contrato_selecionado = mapa_contrato.get(origem_selecionada)

            c4, c5, c6 = st.columns(3)
            profissional_id = c4.selectbox(
                "Profissional",
                options=ids_profissionais,
                index=ids_profissionais.index(profissional_default) if profissional_default in ids_profissionais else 0,
                format_func=lambda valor: profissionais_df.loc[profissionais_df["id"] == valor, "nome"].iloc[0],
            ) if ids_profissionais else None

            opcoes_tipo = ids_tipos + ["__novo__"]
            tipo_id = c5.selectbox(
                "Tipo de atendimento",
                options=opcoes_tipo,
                index=opcoes_tipo.index(tipo_default) if tipo_default in opcoes_tipo else 0 if opcoes_tipo else None,
                format_func=lambda valor: "Adicionar..." if valor == "__novo__" else tipos_df.loc[tipos_df["id"] == valor, "nome"].iloc[0],
            ) if opcoes_tipo else None
            novo_tipo_nome = st.text_input("Novo tipo de atendimento", value="", disabled=(tipo_id != "__novo__"))

            procedimento_id = c6.selectbox(
                "Procedimento",
                options=ids_procedimentos,
                index=ids_procedimentos.index(procedimento_default) if procedimento_default in ids_procedimentos else 0 if ids_procedimentos else None,
                format_func=lambda valor: procedimentos_df.loc[procedimentos_df["id"] == valor, "nome"].iloc[0],
            ) if ids_procedimentos else None

            procedimento_row = None
            contrato_id_selecionado = None
            origem_contrato_salvar = 0
            if contrato_selecionado:
                contrato_id_selecionado = int(contrato_selecionado["contrato_id"])
                origem_contrato_salvar = 1
                if contrato_selecionado.get("procedimento_id"):
                    procedimento_id = int(contrato_selecionado["procedimento_id"])
            if procedimento_id and not procedimentos_df.empty:
                filtro_proc = procedimentos_df[procedimentos_df["id"] == procedimento_id]
                if not filtro_proc.empty:
                    procedimento_row = filtro_proc.iloc[0]

            duracao_sugerida = (
                int(contrato_selecionado["duracao_padrao"])
                if contrato_selecionado
                else int(procedimento_row["duracao_padrao_minutos"]) if procedimento_row is not None else duracao_default
            )
            valor_padrao = float(procedimento_row["valor_padrao"]) if procedimento_row is not None else 0.0

            c7, c8, c9, c10 = st.columns(4)
            data_agendamento = c7.date_input("Data", value=data_default, format="DD/MM/YYYY")
            hora_inicio = c8.selectbox("Horário inicial", options=lista_horarios, index=lista_horarios.index(hora_default) if hora_default in lista_horarios else 4)
            duracao_minutos = c9.selectbox(
                "Duração total",
                options=[15, 30, 45, 60, 75, 90, 120, 150, 180],
                index=[15, 30, 45, 60, 75, 90, 120, 150, 180].index(duracao_sugerida) if duracao_sugerida in [15, 30, 45, 60, 75, 90, 120, 150, 180] else 3,
            )
            status = c10.selectbox(
                "Status",
                options=STATUS_AGENDAMENTO,
                index=STATUS_AGENDAMENTO.index(str(agendamento_atual["status"])) if agendamento_atual is not None and str(agendamento_atual["status"]) in STATUS_AGENDAMENTO else 0,
            )

            hora_fim = hora_fim_query or hora_fim_default or adicionar_minutos_hora(hora_inicio, duracao_minutos)
            if hora_para_minutos(hora_fim) is None or (hora_para_minutos(hora_fim) <= hora_para_minutos(hora_inicio)):
                hora_fim = adicionar_minutos_hora(hora_inicio, duracao_minutos)
            duracao_visual = max((hora_para_minutos(hora_fim) or 0) - (hora_para_minutos(hora_inicio) or 0), 15)

            c11, c12, c13, c14 = st.columns(4)
            c11.text_input("Horário final", value=hora_fim, disabled=True)
            c12.text_input("Valor padrão", value=formatar_moeda_br(valor_padrao), disabled=True)
            c13.text_input("Agendado por", value=str(agendamento_atual["criado_por"]) if agendamento_atual is not None and str(agendamento_atual.get("criado_por", "")).strip() else usuario_logado["usuario"], disabled=True)
            c14.text_input("Agendado em", value=formatar_data_hora_br_valor(agendamento_atual["criado_em"]) if agendamento_atual is not None else formatar_data_hora_br_valor(agora_str()), disabled=True)

            observacoes = st.text_area("Observações", value=agendamento_atual["observacoes"] if agendamento_atual is not None else "", height=90)

            profissional_row = None
            if profissional_id:
                filtro_prof = profissionais_df[profissionais_df["id"] == profissional_id]
                if not filtro_prof.empty:
                    profissional_row = filtro_prof.iloc[0]

            st.markdown("**Seleção visual do horário**")
            renderizar_grade_selecao_horario(data_para_estado_agenda(data_agendamento), profissional_row, agendamento_atual, hora_inicio, hora_fim)

            salvar = st.form_submit_button("Salvar agendamento", use_container_width=True)
            excluir = st.form_submit_button("Excluir agendamento", use_container_width=True) if agendamento_atual is not None else False
            cancelar = st.form_submit_button("Fechar", use_container_width=True)

            if cancelar:
                if agendamento_atual is not None:
                    st.session_state.pop(f"agenda_editar_{int(agendamento_atual['id'])}", None)
                limpar_query_params_agenda()
                st.rerun()

            if salvar:
                paciente_id_salvar = None
                if paciente_escolha not in (None, "__novo__"):
                    paciente_id_salvar = int(paciente_escolha)
                    if paciente_row is not None:
                        nome_paciente = nome_paciente or paciente_row["nome"]
                        telefone_snapshot = telefone_snapshot or paciente_row["telefone"]
                else:
                    if nome_paciente.strip():
                        salvar_paciente_rapido(nome_paciente, telefone_snapshot, email_snapshot)

                if tipo_id == "__novo__":
                    if not novo_tipo_nome.strip():
                        st.error("Informe o nome do novo tipo de atendimento.")
                        return
                    salvar_tipo_atendimento(novo_tipo_nome, "#c5a77a", True, len(tipos_df) + 1)
                    conn.commit()
                    tipo_id = int(cursor.execute("SELECT id FROM tipos_atendimento WHERE nome=? ORDER BY id DESC LIMIT 1", (novo_tipo_nome.strip(),)).fetchone()["id"])

                if not nome_paciente.strip():
                    st.error("Informe o paciente.")
                elif not profissional_id:
                    st.error("Selecione um profissional.")
                elif tem_conflito_agendamento(
                    data_para_estado_agenda(data_agendamento),
                    profissional_id,
                    hora_inicio,
                    hora_fim,
                    excluir_agendamento_id=int(agendamento_atual["id"]) if agendamento_atual is not None else None,
                ):
                    st.error("Há conflito de horário para esse profissional.")
                else:
                    salvar_agendamento_completo(
                        paciente_id=paciente_id_salvar,
                        nome_paciente=nome_paciente,
                        telefone_snapshot=telefone_snapshot,
                        email_snapshot=email_snapshot,
                        profissional_id=profissional_id,
                        tipo_atendimento_id=(tipo_id if tipo_id != "__novo__" else None),
                        procedimento_id=procedimento_id,
                        contrato_id=contrato_id_selecionado,
                        origem_contrato=origem_contrato_salvar,
                        data_agendamento=data_agendamento,
                        hora_inicio=hora_inicio,
                        duracao_minutos=duracao_visual,
                        observacoes=observacoes,
                        status=status,
                        criado_por=(agendamento_atual["criado_por"] if agendamento_atual is not None and str(agendamento_atual.get("criado_por", "")).strip() else usuario_logado["usuario"]),
                        agendamento_id=int(agendamento_atual["id"]) if agendamento_atual is not None else None,
                        procedimento_nome_override=(contrato_selecionado["procedimento_nome"] if contrato_selecionado else ""),
                    )
                    conn.commit()
                    if agendamento_atual is not None:
                        st.session_state.pop(f"agenda_editar_{int(agendamento_atual['id'])}", None)
                    limpar_query_params_agenda()
                    st.success("Agendamento salvo com sucesso.")
                    st.rerun()

            if excluir:
                excluir_agendamento_completo(int(agendamento_atual["id"]))
                conn.commit()
                st.session_state.pop(f"agenda_editar_{int(agendamento_atual['id'])}", None)
                limpar_query_params_agenda()
                st.success("Agendamento excluído.")
                st.rerun()

    modal_agendamento()


def renderizar_cadastros_agenda():
    with st.expander("Configurações da agenda", expanded=False):
        tab_prof, tab_tipos, tab_proc, tab_msg = st.tabs(["Profissionais", "Tipos de atendimento", "Procedimentos", "Mensagens"])

        with tab_prof:
            profissionais_df = obter_profissionais_agenda()
            with st.form("form_profissional_agenda"):
                p1, p2, p3 = st.columns(3)
                nome = p1.text_input("Nome do profissional")
                especialidade = p2.text_input("Especialidade")
                cor = p3.color_picker("Cor", value="#c5a77a")
                p4, p5, p6, p7 = st.columns(4)
                ativo = p4.checkbox("Ativo", value=True)
                ordem = p5.number_input("Ordem", min_value=0, value=int(len(profissionais_df) + 1), step=1)
                hora_inicio = p6.selectbox("Hora inicial", options=gerar_horarios_intervalo("07:00", "20:00", 15), index=4)
                hora_fim = p7.selectbox("Hora final", options=gerar_horarios_intervalo("07:15", "21:00", 15), index=31)
                dias = st.multiselect("Dias de atendimento", options=DIAS_SEMANA_PT, default=DIAS_SEMANA_PT[:5])
                if st.form_submit_button("Salvar profissional", use_container_width=True):
                    try:
                        salvar_profissional_agenda(nome, cor, ativo, especialidade, ordem, dias, hora_inicio, hora_fim)
                        conn.commit()
                        st.success("Profissional salvo.")
                        st.rerun()
                    except ValueError as exc:
                        st.error(str(exc))
            if not profissionais_df.empty:
                st.dataframe(
                    profissionais_df[["nome", "especialidade", "cor", "ativo", "hora_inicio", "hora_fim", "dias_atendimento"]],
                    use_container_width=True,
                    hide_index=True,
                )

        with tab_tipos:
            tipos_df = obter_tipos_agenda()
            with st.form("form_tipo_atendimento"):
                t1, t2, t3 = st.columns(3)
                nome = t1.text_input("Nome do tipo")
                cor = t2.color_picker("Cor do tipo", value="#c5a77a")
                ordem = t3.number_input("Ordem de exibição", min_value=0, value=int(len(tipos_df) + 1), step=1)
                ativo = st.checkbox("Ativo", value=True, key="tipo_ativo_agenda")
                if st.form_submit_button("Salvar tipo", use_container_width=True):
                    try:
                        salvar_tipo_atendimento(nome, cor, ativo, ordem)
                        conn.commit()
                        st.success("Tipo de atendimento salvo.")
                        st.rerun()
                    except ValueError as exc:
                        st.error(str(exc))
            if not tipos_df.empty:
                st.dataframe(tipos_df[["nome", "cor", "ativo", "ordem_exibicao"]], use_container_width=True, hide_index=True)

        with tab_proc:
            procedimentos_df = obter_procedimentos_agenda()
            with st.form("form_procedimento_agenda"):
                r1, r2, r3 = st.columns(3)
                nome = r1.text_input("Nome do procedimento")
                categoria = r2.text_input("Categoria")
                cor = r3.color_picker("Cor opcional", value="#c5a77a")
                r4, r5 = st.columns(2)
                valor = r4.number_input("Valor padrão", min_value=0.0, value=0.0, step=50.0)
                duracao = r5.selectbox("Duração padrão (min)", options=[15, 30, 45, 60, 75, 90, 120, 150, 180], index=3)
                descricao = st.text_area("Descrição", height=80)
                ativo = st.checkbox("Ativo", value=True, key="procedimento_ativo_agenda")
                if st.form_submit_button("Salvar procedimento", use_container_width=True):
                    try:
                        salvar_procedimento_catalogo(nome, categoria, valor, duracao, descricao, cor, ativo)
                        conn.commit()
                        st.success("Procedimento salvo.")
                        st.rerun()
                    except ValueError as exc:
                        st.error(str(exc))
            if not procedimentos_df.empty:
                exibir = procedimentos_df[["nome", "categoria", "valor_padrao", "duracao_padrao_minutos", "ativo"]].copy()
                exibir["valor_padrao"] = exibir["valor_padrao"].apply(formatar_moeda_br)
                st.dataframe(exibir, use_container_width=True, hide_index=True)

        with tab_msg:
            modelos_msg = carregar_modelos_mensagem_agendamento()
            mapa_nomes = {
                "confirmacao_agendamento": "Confirmação de agendamento",
                "lembrete_1_dia_antes": "Lembrete 1 dia antes",
                "lembrete_mesmo_dia": "Lembrete no mesmo dia",
            }
            for _, modelo in modelos_msg.iterrows():
                with st.form(f"form_modelo_msg_{modelo['tipo_mensagem']}"):
                    st.markdown(f"**{mapa_nomes.get(modelo['tipo_mensagem'], modelo['tipo_mensagem'])}**")
                    titulo = st.text_input("Título", value=modelo["titulo"] or "")
                    conteudo = st.text_area(
                        "Mensagem",
                        value=modelo["conteudo"] or "",
                        height=100,
                        help="Variáveis: {paciente}, {data}, {hora}, {profissional}, {clinica}",
                    )
                    if st.form_submit_button("Salvar modelo", use_container_width=True):
                        salvar_modelo_mensagem_agendamento(modelo["tipo_mensagem"], titulo, conteudo)
                        conn.commit()
                        st.success("Modelo salvo.")
                        st.rerun()


def renderizar_agenda_clinica(usuario_logado):
    sincronizar_estado_agenda_por_query()
    if "agenda_data_referencia" not in st.session_state:
        st.session_state["agenda_data_referencia"] = date.today()
    if "agenda_visualizacao" not in st.session_state:
        st.session_state["agenda_visualizacao"] = "Dia"

    data_ref = data_para_estado_agenda(st.session_state.get("agenda_data_referencia"))
    visualizacao = st.session_state.get("agenda_visualizacao", "Dia")
    profissionais_df = obter_profissionais_agenda()
    tipos_df = obter_tipos_agenda()
    procedimentos_df = obter_procedimentos_agenda()
    pacientes_df = carregar_pacientes()

    if "agenda_profissionais_selecionados" not in st.session_state:
        st.session_state["agenda_profissionais_selecionados"] = profissionais_df[profissionais_df["ativo"].fillna(1).astype(int) == 1]["id"].astype(int).tolist()

    st.title("Agenda Clínica")
    esquerda, direita = st.columns([1.2, 4.8], gap="large")
    with esquerda:
        st.markdown("### Filtros")
        data_input = st.date_input("Mini calendário", value=data_ref, format="DD/MM/YYYY")
        data_input = data_para_estado_agenda(data_input)
        if data_input != data_ref:
            st.session_state["agenda_data_referencia"] = data_input
            st.rerun()

        profissionais_ativos = profissionais_df[profissionais_df["ativo"].fillna(1).astype(int) == 1].copy()
        ids_opcoes = profissionais_ativos["id"].astype(int).tolist()
        selecionados = st.multiselect(
            "Profissionais",
            options=ids_opcoes,
            default=[valor for valor in st.session_state.get("agenda_profissionais_selecionados", []) if valor in ids_opcoes] or ids_opcoes,
            format_func=lambda valor: profissionais_ativos.loc[profissionais_ativos["id"] == valor, "nome"].iloc[0],
            key="agenda_profissionais_multiselect",
        )
        st.session_state["agenda_profissionais_selecionados"] = selecionados or ids_opcoes
        html_cores_profissionais(profissionais_ativos, st.session_state["agenda_profissionais_selecionados"])
        renderizar_cadastros_agenda()

    profissionais_ids = st.session_state["agenda_profissionais_selecionados"]
    profissionais_visiveis = profissionais_df[profissionais_df["id"].astype(int).isin(profissionais_ids)].copy()
    if profissionais_visiveis.empty and not profissionais_df.empty:
        profissionais_visiveis = profissionais_df[profissionais_df["ativo"].fillna(1).astype(int) == 1].copy()

    intervalo_inicial = data_ref if visualizacao == "Dia" else inicio_semana(data_ref) if visualizacao == "Semana" else date(data_ref.year, data_ref.month, 1)
    intervalo_final = data_ref if visualizacao == "Dia" else fim_semana(data_ref) if visualizacao == "Semana" else date(data_ref.year, data_ref.month, calendar.monthrange(data_ref.year, data_ref.month)[1])
    agendamentos_df = carregar_agendamentos_clinica(intervalo_inicial, intervalo_final, profissionais_ids)

    with direita:
        b1, b2, b3, b4 = st.columns([1, 1, 1, 2])
        if b1.button("Anterior", use_container_width=True):
            if visualizacao == "Dia":
                st.session_state["agenda_data_referencia"] = data_ref - timedelta(days=1)
            elif visualizacao == "Semana":
                st.session_state["agenda_data_referencia"] = data_ref - timedelta(days=7)
            else:
                primeiro = date(data_ref.year, data_ref.month, 1)
                anterior = primeiro - timedelta(days=1)
                st.session_state["agenda_data_referencia"] = date(anterior.year, anterior.month, 1)
            st.rerun()
        if b2.button("Hoje", use_container_width=True):
            st.session_state["agenda_data_referencia"] = date.today()
            st.rerun()
        if b3.button("Próximo", use_container_width=True):
            if visualizacao == "Dia":
                st.session_state["agenda_data_referencia"] = data_ref + timedelta(days=1)
            elif visualizacao == "Semana":
                st.session_state["agenda_data_referencia"] = data_ref + timedelta(days=7)
            else:
                st.session_state["agenda_data_referencia"] = date(data_ref.year + (1 if data_ref.month == 12 else 0), 1 if data_ref.month == 12 else data_ref.month + 1, 1)
            st.rerun()
        st.radio("Visualização", options=["Dia", "Semana", "Mes"], index=["Dia", "Semana", "Mes"].index(visualizacao), key="agenda_visualizacao", horizontal=True)

        renderizar_formulario_agendamento(usuario_logado, data_ref, profissionais_df, tipos_df, procedimentos_df, pacientes_df)
        if profissionais_visiveis.empty:
            st.info("Cadastre ao menos um profissional ativo para usar a agenda.")
            return
        if visualizacao == "Dia":
            agenda_dia = agendamentos_df[agendamentos_df["data_ref"] == data_ref].copy()
            renderizar_html_interface(gerar_html_grade_dia(data_ref, profissionais_visiveis, agenda_dia))
        elif visualizacao == "Semana":
            renderizar_agenda_semana(data_ref, agendamentos_df)
        else:
            renderizar_agenda_mes(data_ref, agendamentos_df)


def validar_paciente(nome, prontuario):
    erros = []
    if not nome.strip():
        erros.append("Informe o nome do paciente.")
    if not prontuario.strip():
        erros.append("Informe o prontuario.")
    return erros


def validar_dados_paciente(nome, prontuario, cpf, menor, responsavel, cpf_responsavel):
    erros = validar_paciente(nome, prontuario)
    if cpf and not cpf_valido(cpf):
        erros.append("O CPF do paciente e invalido.")
    if menor and not responsavel.strip():
        erros.append("Informe o responsavel legal para paciente menor de idade.")
    if cpf_responsavel and not cpf_valido(cpf_responsavel):
        erros.append("O CPF do responsavel e invalido.")
    return erros


def validar_contrato(procedimentos, valores, entrada, parcelas, vencimento, forma_pagamento, data_pagamento_entrada):
    erros = []
    if not procedimentos:
        erros.append("Informe ao menos um procedimento no contrato.")
    total = sum(valores)
    if entrada < 0:
        erros.append("A entrada nao pode ser negativa.")
    if entrada > total:
        erros.append("A entrada nao pode ser maior que o valor total.")
    if forma_pagamento_a_vista(forma_pagamento):
        if not parse_data_contrato(data_pagamento_entrada):
            erros.append("Informe a data do pagamento no formato DD/MM/AAAA.")
    elif entrada > 0 and not parse_data_contrato(data_pagamento_entrada):
        erros.append("Informe a data do pagamento da entrada no formato DD/MM/AAAA.")
    if not forma_pagamento_a_vista(forma_pagamento) and total - entrada > 0 and int(parcelas or 1) > 0 and not parse_data_contrato(vencimento):
        erros.append("Informe o primeiro vencimento no formato DD/MM/AAAA.")
    return erros, total


def buscar_endereco_por_cep(cep):
    if not cep.strip():
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}

    cep_limpo = "".join(ch for ch in cep if ch.isdigit())
    if len(cep_limpo) != 8:
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}

    try:
        resposta = requests.get(
            f"https://viacep.com.br/ws/{cep_limpo}/json/",
            timeout=5,
        )
        resposta.raise_for_status()
        dados = resposta.json()
        if dados.get("erro"):
            return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}
        return dados
    except requests.RequestException:
        return {"logradouro": "", "bairro": "", "localidade": "", "uf": ""}


def salvar_procedimentos_contrato(contrato_id, procedimentos, valores):
    cursor.execute("DELETE FROM procedimentos_contrato WHERE contrato_id=?", (contrato_id,))
    for procedimento, valor in zip(procedimentos, valores):
        cursor.execute(
            """
            INSERT INTO procedimentos_contrato
            (contrato_id, procedimento, valor)
            VALUES (?, ?, ?)
            """,
            (contrato_id, procedimento, valor),
        )


def sincronizar_financeiro_contrato(contrato_id, paciente_nome, valor_total):
    registro = cursor.execute(
        """
        SELECT id FROM financeiro
        WHERE contrato_id=? AND tipo='Entrada'
        ORDER BY id DESC
        LIMIT 1
        """,
        (contrato_id,),
    ).fetchone()

    if not registro:
        candidatos = cursor.execute(
            """
            SELECT id FROM financeiro
            WHERE contrato_id IS NULL
              AND tipo='Entrada'
              AND descricao=?
              AND origem=?
            ORDER BY id DESC
            """,
            (DESCRICAO_CONTRATO, paciente_nome),
        ).fetchall()
        if len(candidatos) == 1:
            registro = candidatos[0]

    if registro:
        cursor.execute(
            """
            UPDATE financeiro
            SET origem=?, descricao=?, valor=?, data=?, tipo='Entrada', contrato_id=?
            WHERE id=?
            """,
            (
                paciente_nome,
                DESCRICAO_CONTRATO,
                valor_total,
                agora_str(),
                contrato_id,
                registro["id"],
            ),
        )
        return

    cursor.execute(
        """
        INSERT INTO financeiro
        (origem, descricao, valor, data, tipo, contrato_id)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (paciente_nome, DESCRICAO_CONTRATO, valor_total, agora_str(), "Entrada", contrato_id),
    )


def calcular_valores_parcelas(valor_total, entrada, parcelas):
    restante = round(float(valor_total or 0) - float(entrada or 0), 2)
    quantidade = max(int(parcelas or 1), 1)
    if restante <= 0:
        return []

    valor_base = round(restante / quantidade, 2)
    valores = [valor_base] * quantidade
    diferenca = round(restante - sum(valores), 2)
    if valores:
        valores[-1] = round(valores[-1] + diferenca, 2)
    return valores


def montar_texto_pagamento(contrato_row):
    parcelas = int(contrato_row["parcelas"] or 1)
    forma = contrato_row["forma_pagamento"] or ""
    valor_total = float(contrato_row["valor_total"] or 0)
    entrada = float(contrato_row["entrada"] or 0)
    restante = round(valor_total - entrada, 2)
    valor_parcela = round(restante / parcelas, 2) if parcelas > 0 else restante
    data_pagamento = formatar_data_br_valor(
        contrato_row["data_pagamento_entrada"] or contrato_row["data_criacao"] or ""
    )
    vencimento = formatar_data_br_valor(contrato_row["primeiro_vencimento"] or "")

    if forma_pagamento_a_vista(forma):
        return (
            f"VALOR TOTAL R$ {valor_total:.2f}. "
            f"PAGO NO {str(forma).upper()} NO DIA {data_pagamento}."
        )

    partes = [f"VALOR TOTAL R$ {valor_total:.2f}."]
    if entrada > 0:
        partes.append(f"ENTRADA R$ {entrada:.2f} PAGA EM {data_pagamento}.")
    if restante > 0:
        partes.append(
            f"RESTANTE R$ {restante:.2f} EM {parcelas} PARCELAS DE R$ {valor_parcela:.2f} "
            f"NO {str(forma).upper()} COM PRIMEIRO VENCIMENTO EM {vencimento}."
        )
    return " ".join(partes)


def montar_recebiveis_planejados(contrato_id, paciente_id, paciente_nome, prontuario, valor_total, entrada, parcelas, primeiro_vencimento, forma_pagamento):
    data_inicial = parse_data_contrato(primeiro_vencimento)
    if not data_inicial:
        raise ValueError("Informe o primeiro vencimento no formato DD/MM/AAAA para gerar os recebiveis.")

    valores = calcular_valores_parcelas(valor_total, entrada, parcelas)
    recebiveis = []
    for indice, valor_parcela in enumerate(valores, start=1):
        vencimento = adicionar_meses(data_inicial, indice - 1)
        recebiveis.append(
            {
                "contrato_id": int(contrato_id),
                "paciente_id": int(paciente_id),
                "paciente_nome": paciente_nome,
                "prontuario": prontuario,
                "parcela_numero": indice,
                "vencimento": formatar_data_br(vencimento),
                "valor": round(valor_parcela, 2),
                "forma_pagamento": forma_pagamento,
                "status": "Aberto",
                "observacao": "",
            }
        )
    return recebiveis


def carregar_recebiveis_contrato(contrato_id):
    return pd.read_sql(
        "SELECT * FROM recebiveis WHERE contrato_id=? ORDER BY parcela_numero, id",
        conn,
        params=(contrato_id,),
    )


def atualizar_recebivel(recebivel_id, paciente_nome, prontuario, vencimento, valor, forma_pagamento, status, observacao):
    cursor.execute(
        """
        UPDATE recebiveis
        SET paciente_nome=?, prontuario=?, vencimento=?, valor=?, forma_pagamento=?, status=?, observacao=?
        WHERE id=?
        """,
        (
            paciente_nome.strip(),
            prontuario.strip(),
            vencimento.strip(),
            float(valor),
            forma_pagamento,
            status,
            observacao.strip(),
            int(recebivel_id),
        ),
    )


def atualizar_recebiveis_lote_contrato(contrato_id, paciente_nome, prontuario, forma_pagamento, status, observacao, primeiro_vencimento=""):
    recebiveis = cursor.execute(
        """
        SELECT id, parcela_numero, vencimento
        FROM recebiveis
        WHERE contrato_id=?
        ORDER BY parcela_numero, id
        """,
        (int(contrato_id),),
    ).fetchall()

    if not recebiveis:
        return

    data_base = parse_data_contrato(primeiro_vencimento) if primeiro_vencimento.strip() else None
    for row in recebiveis:
        vencimento = row["vencimento"]
        if data_base:
            vencimento = formatar_data_br(adicionar_meses(data_base, int(row["parcela_numero"] or 1) - 1))

        cursor.execute(
            """
            UPDATE recebiveis
            SET paciente_nome=?, prontuario=?, vencimento=?, forma_pagamento=?, status=?, observacao=?
            WHERE id=?
            """,
            (
                paciente_nome.strip(),
                prontuario.strip(),
                vencimento,
                forma_pagamento,
                status,
                observacao.strip(),
                int(row["id"]),
            ),
        )


def atualizar_status_contas_pagar_automaticamente():
    hoje = date.today()
    contas = cursor.execute(
        """
        SELECT id, data_vencimento, status
        FROM contas_pagar
        """
    ).fetchall()

    houve_atualizacao = False
    for conta in contas:
        status_atual = str(conta["status"] or "A vencer").strip() or "A vencer"
        if status_atual in {"Pago", "Suspenso", "Cancelado"}:
            continue

        data_vencimento = parse_data_contrato(conta["data_vencimento"])
        if not data_vencimento:
            continue

        novo_status = "Atrasado" if data_vencimento < hoje else "A vencer"
        if status_atual != novo_status:
            cursor.execute(
                "UPDATE contas_pagar SET status=? WHERE id=?",
                (novo_status, int(conta["id"])),
            )
            houve_atualizacao = True

    if houve_atualizacao:
        conn.commit()


def registrar_movimento_caixa(origem, descricao, valor, tipo, data_movimento, prontuario="", forma_pagamento="", conta_caixa="", observacao="", contrato_id=None, recebivel_id=None):
    data_registro = data_movimento.isoformat() if hasattr(data_movimento, "isoformat") else str(data_movimento or agora_str())
    cursor.execute(
        """
        INSERT INTO financeiro
        (origem, descricao, valor, data, tipo, contrato_id, recebivel_id, prontuario, forma_pagamento, conta_caixa, observacao)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            origem.strip(),
            descricao.strip(),
            float(valor),
            data_registro,
            tipo,
            contrato_id,
            recebivel_id,
            prontuario.strip(),
            forma_pagamento.strip(),
            (conta_caixa or "").strip(),
            observacao.strip(),
        ),
    )
    return cursor.lastrowid


def registrar_saldo_conta(data_saldo, conta, saldo, observacao=""):
    data_registro = data_saldo.isoformat() if hasattr(data_saldo, "isoformat") else str(data_saldo or agora_str())
    cursor.execute(
        """
        INSERT INTO saldos_conta
        (data, conta, saldo, observacao)
        VALUES (?, ?, ?, ?)
        """,
        (
            data_registro,
            conta.strip(),
            float(saldo),
            observacao.strip(),
        ),
    )
    return cursor.lastrowid


def baixar_recebivel_no_caixa(recebivel_id, origem, data_pagamento, observacao="", forma_recebimento="", conta_caixa=""):
    recebivel = cursor.execute(
        """
        SELECT *
        FROM recebiveis
        WHERE id=?
        """,
        (int(recebivel_id),),
    ).fetchone()

    if not recebivel:
        raise ValueError("Recebivel nao encontrado.")
    if (recebivel["status"] or "") == "Pago":
        raise ValueError("Este recebivel ja esta pago.")

    forma_baixa = forma_recebimento.strip() or (recebivel["forma_pagamento"] or "")
    descricao = f"Baixa recebivel parcela {int(recebivel['parcela_numero'] or 0)} - {recebivel['paciente_nome']}"
    financeiro_id = registrar_movimento_caixa(
        origem=origem or "Pagamento",
        descricao=descricao,
        valor=float(recebivel["valor"] or 0),
        tipo="Entrada",
        data_movimento=data_pagamento,
        prontuario=str(recebivel["prontuario"] or ""),
        forma_pagamento=forma_baixa,
        conta_caixa=conta_caixa or forma_baixa,
        observacao=observacao,
        contrato_id=recebivel["contrato_id"],
        recebivel_id=recebivel["id"],
    )

    observacao_atual = str(recebivel["observacao"] or "").strip()
    observacao_nova = observacao.strip()
    if observacao_atual and observacao_nova:
        observacao_final = f"{observacao_atual} | {observacao_nova}"
    else:
        observacao_final = observacao_nova or observacao_atual

    cursor.execute(
        """
        UPDATE recebiveis
        SET status='Pago', data_pagamento=?, observacao=?
        WHERE id=?
        """,
        (
            data_pagamento.strftime("%d/%m/%Y") if hasattr(data_pagamento, "strftime") else str(data_pagamento),
            observacao_final,
            int(recebivel_id),
        ),
    )
    return financeiro_id


def dataframe_para_excel_bytes(df, nome_aba="Dados"):
    if not OPENPYXL_DISPONIVEL:
        return None
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=nome_aba[:31])
    buffer.seek(0)
    return buffer.getvalue()


def formatar_moeda_br(valor):
    try:
        valor_float = float(valor or 0)
    except (TypeError, ValueError):
        valor_float = 0.0
    texto = f"{valor_float:,.2f}"
    return f"R$ {texto.replace(',', 'X').replace('.', ',').replace('X', '.')}"


def html_logo_sidebar():
    if not os.path.exists(LOGO_PATH):
        return ""
    try:
        with open(LOGO_PATH, "rb") as arquivo_logo:
            logo_base64 = base64.b64encode(arquivo_logo.read()).decode("ascii")
        return f'<div class="ss-brand-logo"><img src="data:image/png;base64,{logo_base64}" alt="SoulSul"></div>'
    except Exception:
        return ""


def renderizar_sidebar_premium(menus_disponiveis, menu_atual):
    return None
    itens_menu = []
    for nome in ORDEM_MENU_SIDEBAR:
        if nome not in menus_disponiveis:
            continue
        ativo = " active" if nome == menu_atual else ""
        itens_menu.append(
            f"""
            <a class="ss-nav-link{ativo}" href="?menu={quote(nome)}">
                <span class="ss-nav-icon">{SVG_ICONES_MENU.get(nome, "")}</span>
                <span class="ss-nav-label">{ROTULOS_MENU.get(nome, nome)}</span>
            </a>
            """
        )

    st.sidebar.markdown(
        f"""
        <div class="ss-brand">
            {html_logo_sidebar()}
            <div class="ss-brand-copy">
                <div class="ss-brand-title">SoulSul</div>
                <div class="ss-brand-subtitle">clínica integrada</div>
            </div>
        </div>
        <div class="ss-menu-title">Menu</div>
        <div class="ss-nav">
            {''.join(itens_menu)}
        </div>
        """,
        unsafe_allow_html=True,
    )


def renderizar_sidebar_marca():
    st.sidebar.markdown(
        f"""
        <div class="ss-brand">
            {html_logo_sidebar()}
            <div class="ss-brand-copy">
                <div class="ss-brand-title">SoulSul</div>
                <div class="ss-brand-subtitle">cl&iacute;nica integrada</div>
            </div>
        </div>
        <div class="ss-menu-title">Menu</div>
        """,
        unsafe_allow_html=True,
    )


def selecionar_blocos_visiveis(rotulo, opcoes, padrao, chave):
    return set(
        st.multiselect(
            rotulo,
            options=opcoes,
            default=padrao,
            key=chave,
        )
    )


CONTAS_CAIXA_MODELO = ["CAIXA", "SICOOB", "INFINITEPAY", "C6", "SANTANDER", "PAGSEGURO"]


def identificar_conta_caixa(origem, descricao="", observacao=""):
    texto = normalizar_texto(f"{origem} {descricao} {observacao}")
    if "infinitepay" in texto:
        return "INFINITEPAY"
    if "sicoob" in texto:
        return "SICOOB"
    if re.search(r"\bc6\b", texto):
        return "C6"
    if "santander" in texto:
        return "SANTANDER"
    if "pagseguro" in texto:
        return "PAGSEGURO"
    return "CAIXA"


def saldos_informados_por_conta_ate(data_limite):
    saldos = {conta: 0.0 for conta in CONTAS_CAIXA_MODELO}
    registros = pd.read_sql(
        "SELECT * FROM saldos_conta ORDER BY data ASC, id ASC",
        conn,
    )
    if registros.empty:
        return saldos

    registros["data_ref"] = registros["data"].apply(parse_data_contrato)
    for _, row in registros.iterrows():
        data_ref = row["data_ref"]
        conta = str(row["conta"] or "").strip().upper()
        if not data_ref or conta not in saldos:
            continue
        if data_ref <= data_limite:
            saldos[conta] = float(row["saldo"] or 0)
    return saldos


def montar_resumo_saldos_conta(saldos_dict, data_referencia):
    linhas = []
    data_texto = formatar_data_br(data_referencia) if hasattr(data_referencia, "strftime") else formatar_data_br_valor(data_referencia)
    for conta in CONTAS_CAIXA_MODELO:
        linhas.append(
            {
                "Data": data_texto,
                "Descricao": f"SALDO {conta}",
                "Saldo": formatar_moeda_br(saldos_dict.get(conta, 0)),
                "Banco": conta,
            }
        )
    return pd.DataFrame(linhas)


def estilizar_resumo_saldos(df):
    return (
        df.style
        .apply(lambda _: ["background-color: #C6E0B4"] * len(df.columns), axis=1)
        .set_properties(**{"color": "black"})
    )


def montar_caixa_diario(financeiro_df):
    if financeiro_df.empty:
        return pd.DataFrame(), []

    caixa = financeiro_df.copy()
    caixa["data_date"] = caixa["data"].apply(parse_data_contrato)
    caixa["data_grupo"] = caixa["data_date"].apply(lambda valor: formatar_data_br(valor) if valor else "")
    caixa["data_grupo"] = caixa["data_grupo"].replace("", "Data nao informada")
    caixa["ordem_data"] = caixa["data_date"].apply(lambda valor: valor.toordinal() if valor else -1)
    caixa = caixa.sort_values(["ordem_data", "id"], ascending=[False, False])

    grupos = []
    resumo_linhas = []
    for data_grupo, grupo in caixa.groupby("data_grupo", sort=False):
        entradas = float(grupo.loc[grupo["tipo"] == "Entrada", "valor"].sum())
        saidas = float(grupo.loc[grupo["tipo"] == "Saida", "valor"].sum())
        saldo = entradas - saidas

        detalhamento = grupo[
            ["tipo", "origem", "conta_caixa", "prontuario", "descricao", "forma_pagamento", "observacao", "valor"]
        ].copy()
        detalhamento = detalhamento.rename(
            columns={
                "tipo": "Tipo",
                "origem": "Origem",
                "conta_caixa": "Conta/Banco",
                "prontuario": "Prontuario",
                "descricao": "Descricao",
                "forma_pagamento": "Forma de pagamento",
                "observacao": "Observacao",
                "valor": "Valor",
            }
        )
        detalhamento["Conta/Banco"] = detalhamento["Conta/Banco"].fillna("").replace("None", "")
        detalhamento["Prontuario"] = detalhamento["Prontuario"].fillna("").replace("None", "")
        detalhamento["Forma de pagamento"] = detalhamento["Forma de pagamento"].fillna("").replace("None", "")
        detalhamento["Observacao"] = detalhamento["Observacao"].fillna("").replace("None", "")

        grupos.append(
            {
                "data": data_grupo,
                "entradas": entradas,
                "saidas": saidas,
                "saldo": saldo,
                "detalhamento": detalhamento.reset_index(drop=True),
            }
        )

        resumo_linhas.append(
            {
                "Data": data_grupo,
                "Entradas": entradas,
                "Saidas": saidas,
                "Saldo do dia": saldo,
                "Lancamentos": len(grupo),
            }
        )

    resumo_df = pd.DataFrame(resumo_linhas)
    return resumo_df, grupos


def caixa_diario_para_excel_bytes(financeiro_df):
    if not OPENPYXL_DISPONIVEL:
        return None
    caixa = financeiro_df.copy()
    if caixa.empty:
        return dataframe_para_excel_bytes(pd.DataFrame(columns=["Data", "Descricao", "Entradas", "Saidas", "Banco"]), nome_aba="Caixa")

    caixa["data_date"] = caixa["data"].apply(parse_data_contrato)
    caixa = caixa[caixa["data_date"].notna()].copy()
    if caixa.empty:
        return dataframe_para_excel_bytes(pd.DataFrame(columns=["Data", "Descricao", "Entradas", "Saidas", "Banco"]), nome_aba="Caixa")

    caixa["conta_caixa"] = caixa.apply(
        lambda row: str(row.get("conta_caixa") or "").strip().upper()
        or identificar_conta_caixa(row.get("origem", ""), row.get("descricao", ""), row.get("observacao", "")),
        axis=1,
    )
    caixa = caixa.sort_values(["data_date", "id"], ascending=[True, True])

    wb = Workbook()
    ws = wb.active
    ws.title = "Caixa"

    fonte_titulo = Font(name="Times New Roman", size=20, bold=False, color="000000")
    fonte_cabecalho = Font(name="Times New Roman", size=12, bold=False, color="000000")
    fonte_normal = Font(name="Times New Roman", size=12, color="000000")
    fonte_vermelha = Font(name="Times New Roman", size=12, color="FF0000")
    fonte_preta = Font(name="Times New Roman", size=12, color="000000")
    preenchimento_verde = PatternFill(fill_type="solid", fgColor="C6E0B4")
    borda_fina = Border(
        left=Side(style="thin", color="000000"),
        right=Side(style="thin", color="000000"),
        top=Side(style="thin", color="000000"),
        bottom=Side(style="thin", color="000000"),
    )

    ws.merge_cells("A1:E1")
    ws["A1"] = "CAIXA"
    ws["A1"].font = fonte_titulo
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")

    larguras = {"A": 14, "B": 58, "C": 18, "D": 16, "E": 18}
    for coluna, largura in larguras.items():
        ws.column_dimensions[coluna].width = largura

    linha_atual = 3
    saldo_anterior = {conta: 0.0 for conta in CONTAS_CAIXA_MODELO}

    for data_atual in sorted(caixa["data_date"].dropna().unique().tolist()):
        data_texto = formatar_data_br(data_atual)
        grupo = caixa[caixa["data_date"] == data_atual].copy()
        saldo_anterior = saldos_informados_por_conta_ate(data_atual)

        ws.cell(row=linha_atual, column=1, value="Data").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=2, value="Descrição").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=3, value="Entradas").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=4, value="Saídas").font = fonte_cabecalho
        ws.cell(row=linha_atual, column=5, value="Banco").font = fonte_cabecalho
        for coluna in range(1, 6):
            ws.cell(row=linha_atual, column=coluna).border = borda_fina
        linha_atual += 1

        for conta in CONTAS_CAIXA_MODELO:
            saldo_valor = round(saldo_anterior[conta], 2)
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=f"SALDO {conta}")
            ws.cell(row=linha_atual, column=3, value=saldo_valor)
            for coluna in range(1, 6):
                ws.cell(row=linha_atual, column=coluna).fill = preenchimento_verde
                ws.cell(row=linha_atual, column=coluna).border = borda_fina
                ws.cell(row=linha_atual, column=coluna).font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=3).font = fonte_vermelha if saldo_valor < 0 else fonte_preta
            linha_atual += 1

        linha_atual += 3

        for _, row in grupo.iterrows():
            valor = round(float(row["valor"] or 0), 2)
            conta = row["conta_caixa"]
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=str(row["descricao"] or row["origem"] or "").strip())
            if str(row["tipo"]) == "Entrada":
                ws.cell(row=linha_atual, column=3, value=valor)
                ws.cell(row=linha_atual, column=4, value=None)
                saldo_anterior[conta] += valor
            else:
                ws.cell(row=linha_atual, column=3, value=None)
                ws.cell(row=linha_atual, column=4, value=valor)
                saldo_anterior[conta] -= valor
            ws.cell(row=linha_atual, column=5, value=conta)

            for coluna in range(1, 6):
                ws.cell(row=linha_atual, column=coluna).border = borda_fina
                ws.cell(row=linha_atual, column=coluna).font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=4).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=4).font = fonte_vermelha if ws.cell(row=linha_atual, column=4).value not in (None, 0, 0.0) else fonte_normal
            linha_atual += 1

        linha_atual += 3

        for conta in CONTAS_CAIXA_MODELO:
            saldo_valor = round(saldo_anterior[conta], 2)
            ws.cell(row=linha_atual, column=1, value=data_texto)
            ws.cell(row=linha_atual, column=2, value=f"SALDO {conta}")
            ws.cell(row=linha_atual, column=3, value=saldo_valor)
            for coluna in range(1, 6):
                ws.cell(row=linha_atual, column=coluna).fill = preenchimento_verde
                ws.cell(row=linha_atual, column=coluna).border = borda_fina
                ws.cell(row=linha_atual, column=coluna).font = fonte_normal
            ws.cell(row=linha_atual, column=3).number_format = '"R$" #,##0.00'
            ws.cell(row=linha_atual, column=3).font = fonte_vermelha if saldo_valor < 0 else fonte_preta
            linha_atual += 1

        linha_atual += 3

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def dataframe_para_pdf_bytes(df, titulo="Relatorio"):
    if not REPORTLAB_DISPONIVEL:
        return None

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    largura, altura = A4
    y = altura - 40

    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(40, y, titulo)
    y -= 24

    pdf.setFont("Helvetica", 8)
    colunas = [str(coluna) for coluna in df.columns]
    linhas = [colunas] + df.astype(str).values.tolist()

    for linha in linhas:
        texto = " | ".join(linha)
        while len(texto) > 140:
            pdf.drawString(40, y, texto[:140])
            texto = texto[140:]
            y -= 12
            if y < 40:
                pdf.showPage()
                pdf.setFont("Helvetica", 8)
                y = altura - 40
        pdf.drawString(40, y, texto)
        y -= 12
        if y < 40:
            pdf.showPage()
            pdf.setFont("Helvetica", 8)
            y = altura - 40

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def identificar_coluna(colunas, candidatos):
    mapa = {normalizar_texto(coluna): coluna for coluna in colunas}
    for candidato in candidatos:
        chave = normalizar_texto(candidato)
        if chave in mapa:
            return mapa[chave]
    for coluna in colunas:
        texto = normalizar_texto(coluna)
        if any(normalizar_texto(candidato) in texto for candidato in candidatos):
            return coluna
    return None


def ler_extrato_arquivo(arquivo):
    nome = (arquivo.name or "").lower()
    if nome.endswith(".csv"):
        try:
            return pd.read_csv(arquivo, sep=None, engine="python")
        except Exception:
            arquivo.seek(0)
            return pd.read_csv(arquivo, sep=";", engine="python")
    return pd.read_excel(arquivo)


def reconhecer_extrato(df_origem):
    if df_origem is None or df_origem.empty:
        return pd.DataFrame()

    df = df_origem.copy()
    coluna_data = identificar_coluna(df.columns, ["data", "date", "lancamento", "lançamento", "movimento"])
    coluna_descricao = identificar_coluna(df.columns, ["descricao", "descrição", "historico", "histórico", "detalhe", "complemento"])
    coluna_valor = identificar_coluna(df.columns, ["valor", "amount", "valor (r$)", "lancamento", "lançamento"])
    coluna_credito = identificar_coluna(df.columns, ["credito", "crédito", "entrada"])
    coluna_debito = identificar_coluna(df.columns, ["debito", "débito", "saida", "saída"])

    if not coluna_data or not coluna_descricao or (not coluna_valor and not (coluna_credito or coluna_debito)):
        return pd.DataFrame()

    reconhecido = pd.DataFrame()
    reconhecido["data"] = pd.to_datetime(df[coluna_data], dayfirst=True, errors="coerce")
    reconhecido["descricao"] = df[coluna_descricao].fillna("").astype(str).str.strip()

    if coluna_valor:
        serie_valor = pd.to_numeric(
            df[coluna_valor]
            .astype(str)
            .str.replace(".", "", regex=False)
            .str.replace(",", ".", regex=False),
            errors="coerce",
        )
    else:
        credito = pd.to_numeric(
            df[coluna_credito].fillna(0).astype(str).str.replace(".", "", regex=False).str.replace(",", ".", regex=False),
            errors="coerce",
        ) if coluna_credito else 0
        debito = pd.to_numeric(
            df[coluna_debito].fillna(0).astype(str).str.replace(".", "", regex=False).str.replace(",", ".", regex=False),
            errors="coerce",
        ) if coluna_debito else 0
        serie_valor = credito.fillna(0) - debito.fillna(0)

    reconhecido["valor_original"] = serie_valor
    reconhecido = reconhecido.dropna(subset=["data", "valor_original"])
    reconhecido = reconhecido[reconhecido["descricao"] != ""]
    reconhecido["tipo"] = reconhecido["valor_original"].apply(lambda valor: "Entrada" if float(valor) >= 0 else "Saida")
    reconhecido["valor"] = reconhecido["valor_original"].abs().round(2)
    reconhecido["data_exibicao"] = reconhecido["data"].dt.strftime("%d/%m/%Y")
    return reconhecido[["data", "data_exibicao", "descricao", "valor", "tipo"]].reset_index(drop=True)


def preencher_prontuarios_recebiveis():
    pacientes = carregar_pacientes()
    if pacientes.empty:
        return

    prontuario_por_id = {}
    prontuario_por_nome = {}
    for _, row in pacientes.iterrows():
        prontuario = str(row["prontuario"] or "").strip()
        prontuario_por_id[int(row["id"])] = prontuario
        nome_normalizado = normalizar_texto(row["nome"])
        if nome_normalizado and prontuario and nome_normalizado not in prontuario_por_nome:
            prontuario_por_nome[nome_normalizado] = prontuario

    recebiveis_sem_prontuario = cursor.execute(
        """
        SELECT id, paciente_id, paciente_nome, prontuario
        FROM recebiveis
        """
    ).fetchall()

    houve_atualizacao = False
    for row in recebiveis_sem_prontuario:
        prontuario_atual = str(row["prontuario"] or "").strip()
        if prontuario_atual and prontuario_atual.lower() != "none":
            continue

        prontuario = ""
        if row["paciente_id"] is not None:
            prontuario = prontuario_por_id.get(int(row["paciente_id"]), "")
        if not prontuario:
            prontuario = prontuario_por_nome.get(normalizar_texto(row["paciente_nome"]), "")

        if prontuario:
            cursor.execute(
                "UPDATE recebiveis SET prontuario=? WHERE id=?",
                (prontuario, int(row["id"])),
            )
            houve_atualizacao = True

    if houve_atualizacao:
        conn.commit()


def normalizar_status_recebiveis_por_observacao():
    recebiveis = cursor.execute(
        """
        SELECT id, status, observacao
        FROM recebiveis
        WHERE COALESCE(observacao, '') <> ''
        """
    ).fetchall()

    marcadores_suspensao = [
        "suspens",
        "cancel",
    ]

    houve_atualizacao = False
    for row in recebiveis:
        status_atual = str(row["status"] or "").strip()
        if normalizar_texto(status_atual) == "pago":
            continue

        observacao = normalizar_texto(row["observacao"])
        if not observacao:
            continue

        if any(marcador in observacao for marcador in marcadores_suspensao):
            if status_atual != "Suspenso":
                cursor.execute(
                    "UPDATE recebiveis SET status=? WHERE id=?",
                    ("Suspenso", int(row["id"])),
                )
                houve_atualizacao = True

    if houve_atualizacao:
        conn.commit()


def sincronizar_recebiveis_contrato(contrato_id, paciente_id, paciente_nome, prontuario, valor_total, entrada, parcelas, primeiro_vencimento, forma_pagamento):
    planejados = montar_recebiveis_planejados(
        contrato_id,
        paciente_id,
        paciente_nome,
        prontuario,
        valor_total,
        entrada,
        parcelas,
        primeiro_vencimento,
        forma_pagamento,
    )

    existentes = cursor.execute(
        """
        SELECT paciente_id, paciente_nome, prontuario, parcela_numero, vencimento, valor, forma_pagamento, status, observacao
        FROM recebiveis
        WHERE contrato_id=?
        ORDER BY parcela_numero, id
        """,
        (contrato_id,),
    ).fetchall()

    existentes_normalizados = [
        (
            int(row["paciente_id"] or 0),
            row["paciente_nome"] or "",
            row["prontuario"] or "",
            int(row["parcela_numero"] or 0),
            row["vencimento"] or "",
            round(float(row["valor"] or 0), 2),
            row["forma_pagamento"] or "",
            row["status"] or "Aberto",
            row["observacao"] or "",
        )
        for row in existentes
    ]
    planejados_normalizados = [
        (
            item["paciente_id"],
            item["paciente_nome"],
            item["prontuario"],
            item["parcela_numero"],
            item["vencimento"],
            item["valor"],
            item["forma_pagamento"],
            item["status"],
            item["observacao"],
        )
        for item in planejados
    ]

    if existentes_normalizados == planejados_normalizados:
        return "inalterado"

    cursor.execute("DELETE FROM recebiveis WHERE contrato_id=?", (contrato_id,))
    for item in planejados:
        cursor.execute(
            """
            INSERT INTO recebiveis
            (contrato_id, paciente_id, paciente_nome, prontuario, parcela_numero, vencimento, valor, forma_pagamento, status, observacao, data_criacao)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                item["contrato_id"],
                item["paciente_id"],
                item["paciente_nome"],
                item["prontuario"],
                item["parcela_numero"],
                item["vencimento"],
                item["valor"],
                item["forma_pagamento"],
                item["status"],
                item["observacao"],
                agora_str(),
            ),
        )

    return "criado" if not existentes_normalizados else "atualizado"


def opcoes_pacientes(df_pacientes):
    opcoes = []
    for _, row in df_pacientes.iterrows():
        opcoes.append((int(row["id"]), f"{row['nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])}"))
    return opcoes


def registrar_paciente_recente(paciente_id):
    try:
        paciente_id = int(paciente_id)
    except (TypeError, ValueError):
        return
    recentes = [int(valor) for valor in st.session_state.get("pacientes_recentes", []) if str(valor).isdigit()]
    recentes = [valor for valor in recentes if valor != paciente_id]
    recentes.insert(0, paciente_id)
    st.session_state["pacientes_recentes"] = recentes[:10]


def pacientes_recentes(df_pacientes, limite=6):
    if df_pacientes.empty:
        return df_pacientes
    recentes_ids = [int(valor) for valor in st.session_state.get("pacientes_recentes", []) if str(valor).isdigit()]
    if recentes_ids:
        pacientes_map = {int(row["id"]): row for _, row in df_pacientes.iterrows()}
        linhas = [pacientes_map[paciente_id] for paciente_id in recentes_ids if paciente_id in pacientes_map]
        if linhas:
            return pd.DataFrame(linhas).head(limite)
    return df_pacientes.sort_values("id", ascending=False).head(limite).copy()


def filtrar_pacientes_busca(df_pacientes, termo_busca):
    termo = normalizar_texto(termo_busca)
    if not termo or df_pacientes.empty:
        return df_pacientes.copy()

    def linha_combina(row):
        campos = [
            row.get("nome", ""),
            row.get("apelido", ""),
            row.get("prontuario", ""),
            row.get("cpf", ""),
            row.get("telefone", ""),
            row.get("email", ""),
        ]
        return termo in normalizar_texto(" ".join([str(campo or "") for campo in campos]))

    return df_pacientes[df_pacientes.apply(linha_combina, axis=1)].copy()


def iniciais_paciente(nome):
    partes = [parte for parte in str(nome or "").strip().split() if parte]
    if not partes:
        return "SS"
    if len(partes) == 1:
        return partes[0][:2].upper()
    return (partes[0][0] + partes[-1][0]).upper()


def status_cor_financeiro(status):
    status_norm = normalizar_texto(status)
    if status_norm == "pago":
        return "#027a48", "#eefbf2", "#b7e5c8"
    if status_norm == "atrasado":
        return "#b42318", "#fff2f0", "#f2b8b5"
    if status_norm in {"a vencer", "aberto"}:
        return "#b54708", "#fff7ed", "#f4d2a8"
    if status_norm in {"suspenso", "cancelado"}:
        return "#667085", "#f8f9fb", "#d0d5dd"
    return "#667085", "#f8f9fb", "#d0d5dd"


def status_cor_agendamento_paciente(status):
    status_norm = normalizar_texto(status)
    mapa = {
        "agendado": ("#1d4ed8", "#eff6ff", "#bfdbfe"),
        "confirmado": ("#027a48", "#eefbf2", "#b7e5c8"),
        "em espera": ("#b54708", "#fff7ed", "#f4d2a8"),
        "em atendimento": ("#7c3aed", "#f5f3ff", "#d8b4fe"),
        "atendido": ("#0f766e", "#ecfeff", "#99f6e4"),
        "atrasado": ("#b42318", "#fff2f0", "#f2b8b5"),
        "faltou": ("#b42318", "#fff2f0", "#f2b8b5"),
        "cancelado": ("#667085", "#f8f9fb", "#d0d5dd"),
    }
    return mapa.get(status_norm, ("#667085", "#f8f9fb", "#d0d5dd"))


def carregar_contratos_paciente(paciente_id):
    return pd.read_sql(
        "SELECT * FROM contratos WHERE paciente_id=? ORDER BY COALESCE(data_criacao, '' ) DESC, id DESC",
        conn,
        params=(int(paciente_id),),
    )


def carregar_recebiveis_paciente(paciente_row):
    return pd.read_sql(
        """
        SELECT *
        FROM recebiveis
        WHERE paciente_id = ?
           OR lower(trim(paciente_nome)) = lower(trim(?))
        ORDER BY vencimento, id
        """,
        conn,
        params=(int(paciente_row["id"]), str(paciente_row["nome"] or "")),
    )


def carregar_agendamentos_paciente(paciente_row):
    return pd.read_sql(
        """
        SELECT *
        FROM agendamentos
        WHERE paciente_id = ?
           OR lower(trim(COALESCE(nome_paciente_snapshot, paciente_nome, ''))) = lower(trim(?))
        ORDER BY COALESCE(data_agendamento, data), hora_inicio, id
        """,
        conn,
        params=(int(paciente_row["id"]), str(paciente_row["nome"] or "")),
    )


def proximo_agendamento_paciente(agendamentos_df):
    if agendamentos_df.empty:
        return None
    itens = []
    agora = datetime.now()
    for _, row in agendamentos_df.iterrows():
        data_ref = parse_data_contrato(row.get("data_agendamento") or row.get("data"))
        hora_ref = parse_hora_agendamento(row.get("hora_inicio"))
        if not data_ref or not hora_ref:
            continue
        data_hora = datetime.combine(data_ref, hora_ref)
        if data_hora >= agora and normalizar_texto(row.get("status")) not in {"cancelado", "faltou"}:
            itens.append((data_hora, row))
    if not itens:
        return None
    itens.sort(key=lambda item: item[0])
    return itens[0][1]


def resumo_financeiro_paciente(recebiveis_df):
    if recebiveis_df.empty:
        return {
            "total": 0.0,
            "aberto": 0.0,
            "atrasado": 0.0,
            "pagos": 0.0,
            "quantidade_atrasados": 0,
        }
    df = recebiveis_df.copy()
    df["valor"] = df["valor"].fillna(0).astype(float)
    total = float(df["valor"].sum())
    pagos = float(df[df["status"] == "Pago"]["valor"].sum())
    atrasados = float(df[df["status"] == "Atrasado"]["valor"].sum())
    aberto = float(df[df["status"].isin(["Aberto", "A vencer", "Atrasado"])]["valor"].sum())
    qtd_atrasados = int((df["status"] == "Atrasado").sum())
    return {
        "total": total,
        "aberto": aberto,
        "atrasado": atrasados,
        "pagos": pagos,
        "quantidade_atrasados": qtd_atrasados,
    }


def slug_paciente_arquivos(paciente_row):
    prontuario = formatar_prontuario_valor(paciente_row.get("prontuario"))
    nome = limpar_nome(paciente_row.get("nome", "PACIENTE"))
    base = f"{prontuario}_{nome}".strip("_")
    return base or f"paciente_{int(paciente_row['id'])}"


def pasta_exames_paciente(paciente_row):
    pasta = os.path.join("dados_pacientes", "exames", slug_paciente_arquivos(paciente_row))
    os.makedirs(pasta, exist_ok=True)
    return pasta


def listar_exames_paciente(paciente_row):
    pasta = pasta_exames_paciente(paciente_row)
    arquivos = []
    for nome_arquivo in sorted(os.listdir(pasta), reverse=True):
        caminho = os.path.join(pasta, nome_arquivo)
        if os.path.isfile(caminho):
            arquivos.append(
                {
                    "nome": nome_arquivo,
                    "caminho": caminho,
                    "extensao": os.path.splitext(nome_arquivo)[1].lower(),
                    "modificado_em": datetime.fromtimestamp(os.path.getmtime(caminho)),
                }
            )
    return arquivos


def salvar_uploads_exames_paciente(paciente_row, arquivos_upload):
    pasta = pasta_exames_paciente(paciente_row)
    salvos = 0
    for arquivo in arquivos_upload or []:
        if arquivo is None:
            continue
        nome_limpo = re.sub(r"[^A-Za-z0-9._ -]", "_", arquivo.name)
        destino = os.path.join(pasta, nome_limpo)
        if os.path.exists(destino):
            base, ext = os.path.splitext(nome_limpo)
            destino = os.path.join(pasta, f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}")
        with open(destino, "wb") as arquivo_saida:
            arquivo_saida.write(arquivo.getbuffer())
        salvos += 1
    return salvos


def listar_documentos_paciente(paciente_row):
    if not os.path.isdir(DOCS_DIR):
        return []
    prontuario = normalizar_texto(formatar_prontuario_valor(paciente_row.get("prontuario")))
    nome = normalizar_texto(limpar_nome(paciente_row.get("nome", "")))
    encontrados = []
    for nome_arquivo in sorted(os.listdir(DOCS_DIR), reverse=True):
        caminho = os.path.join(DOCS_DIR, nome_arquivo)
        if not os.path.isfile(caminho):
            continue
        nome_norm = normalizar_texto(nome_arquivo)
        if (prontuario and prontuario in nome_norm) or (nome and nome in nome_norm):
            encontrados.append(
                {
                    "nome": nome_arquivo,
                    "caminho": caminho,
                    "modificado_em": datetime.fromtimestamp(os.path.getmtime(caminho)),
                }
            )
    return encontrados


def salvar_paciente_completo(dados):
    cursor.execute(
        """
        INSERT INTO pacientes
        (
            nome, apelido, sexo, prontuario, cpf, rg, data_nascimento, telefone, email, cep,
            endereco, numero, bairro, cidade, estado, estado_civil, observacoes,
            menor_idade, responsavel, cpf_responsavel
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            dados["nome"].strip(),
            dados["apelido"].strip(),
            dados["sexo"].strip(),
            dados["prontuario"].strip(),
            limpar_cpf(dados["cpf"]),
            dados["rg"].strip(),
            dados["data_nascimento"].strip(),
            dados["telefone"].strip(),
            dados["email"].strip(),
            dados["cep"].strip(),
            dados["endereco"].strip(),
            dados["numero"].strip(),
            dados["bairro"].strip(),
            dados["cidade"].strip(),
            dados["estado"].strip(),
            dados["estado_civil"].strip(),
            dados["observacoes"].strip(),
            int(bool(dados["menor"])),
            dados["responsavel"].strip(),
            limpar_cpf(dados["cpf_responsavel"]),
        ),
    )
    conn.commit()
    return cursor.lastrowid


def atualizar_paciente_completo(paciente_id, dados):
    cursor.execute(
        """
        UPDATE pacientes
        SET nome=?, apelido=?, sexo=?, prontuario=?, cpf=?, rg=?, data_nascimento=?, telefone=?, email=?, cep=?,
            endereco=?, numero=?, bairro=?, cidade=?, estado=?, estado_civil=?, observacoes=?, menor_idade=?,
            responsavel=?, cpf_responsavel=?
        WHERE id=?
        """,
        (
            dados["nome"].strip(),
            dados["apelido"].strip(),
            dados["sexo"].strip(),
            dados["prontuario"].strip(),
            limpar_cpf(dados["cpf"]),
            dados["rg"].strip(),
            dados["data_nascimento"].strip(),
            dados["telefone"].strip(),
            dados["email"].strip(),
            dados["cep"].strip(),
            dados["endereco"].strip(),
            dados["numero"].strip(),
            dados["bairro"].strip(),
            dados["cidade"].strip(),
            dados["estado"].strip(),
            dados["estado_civil"].strip(),
            dados["observacoes"].strip(),
            int(bool(dados["menor"])),
            dados["responsavel"].strip() if dados["menor"] else "",
            limpar_cpf(dados["cpf_responsavel"]) if dados["menor"] else "",
            int(paciente_id),
        ),
    )
    conn.commit()


def preparar_estado_paciente(prefixo, dados, forcar=False):
    for chave, valor in dados.items():
        chave_estado = f"{prefixo}_{chave}"
        if forcar or chave_estado not in st.session_state:
            st.session_state[chave_estado] = valor


def dados_formulario_paciente(prefixo):
    return {
        "nome": st.session_state.get(f"{prefixo}_nome", ""),
        "apelido": st.session_state.get(f"{prefixo}_apelido", ""),
        "sexo": st.session_state.get(f"{prefixo}_sexo", ""),
        "prontuario": st.session_state.get(f"{prefixo}_prontuario", ""),
        "cpf": st.session_state.get(f"{prefixo}_cpf", ""),
        "rg": st.session_state.get(f"{prefixo}_rg", ""),
        "data_nascimento": st.session_state.get(f"{prefixo}_data_nascimento", ""),
        "telefone": st.session_state.get(f"{prefixo}_telefone", ""),
        "email": st.session_state.get(f"{prefixo}_email", ""),
        "cep": st.session_state.get(f"{prefixo}_cep", ""),
        "endereco": st.session_state.get(f"{prefixo}_endereco", ""),
        "numero": st.session_state.get(f"{prefixo}_numero", ""),
        "bairro": st.session_state.get(f"{prefixo}_bairro", ""),
        "cidade": st.session_state.get(f"{prefixo}_cidade", ""),
        "estado": st.session_state.get(f"{prefixo}_estado", ""),
        "estado_civil": st.session_state.get(f"{prefixo}_estado_civil", ""),
        "observacoes": st.session_state.get(f"{prefixo}_observacoes", ""),
        "menor": st.session_state.get(f"{prefixo}_menor", False),
        "responsavel": st.session_state.get(f"{prefixo}_responsavel", ""),
        "cpf_responsavel": st.session_state.get(f"{prefixo}_cpf_responsavel", ""),
    }


def dados_paciente_vazios():
    return {
        "nome": "",
        "apelido": "",
        "sexo": "",
        "prontuario": "",
        "cpf": "",
        "rg": "",
        "data_nascimento": "",
        "telefone": "",
        "email": "",
        "cep": "",
        "endereco": "",
        "numero": "",
        "bairro": "",
        "cidade": "",
        "estado": "",
        "estado_civil": "",
        "observacoes": "",
        "menor": False,
        "responsavel": "",
        "cpf_responsavel": "",
    }


def aplicar_dados_extraidos_paciente_prefixo(prefixo, dados):
    preparar_estado_paciente(
        prefixo,
        {
            "nome": dados.get("nome", ""),
            "apelido": dados.get("apelido", ""),
            "sexo": dados.get("sexo", ""),
            "prontuario": dados.get("prontuario", ""),
            "cpf": dados.get("cpf", ""),
            "rg": dados.get("rg", ""),
            "data_nascimento": dados.get("data_nascimento", ""),
            "telefone": dados.get("telefone", ""),
            "email": dados.get("email", ""),
            "cep": dados.get("cep", ""),
            "endereco": dados.get("endereco", ""),
            "numero": dados.get("numero", ""),
            "bairro": dados.get("bairro", ""),
            "cidade": dados.get("cidade", ""),
            "estado": dados.get("estado", ""),
            "estado_civil": dados.get("estado_civil", ""),
            "observacoes": dados.get("observacoes", ""),
            "menor": dados.get("menor_idade", False),
            "responsavel": dados.get("responsavel", ""),
            "cpf_responsavel": dados.get("cpf_responsavel", ""),
        },
        forcar=True,
    )


def renderizar_campos_paciente(prefixo, dados_iniciais):
    preparar_estado_paciente(prefixo, dados_iniciais)
    cep_valor = st.session_state.get(f"{prefixo}_cep", "")
    dados_cep = buscar_endereco_por_cep(cep_valor) if cep_valor else {}
    if dados_cep:
        if not st.session_state.get(f"{prefixo}_endereco"):
            st.session_state[f"{prefixo}_endereco"] = dados_cep.get("logradouro", "")
        if not st.session_state.get(f"{prefixo}_bairro"):
            st.session_state[f"{prefixo}_bairro"] = dados_cep.get("bairro", "")
        if not st.session_state.get(f"{prefixo}_cidade"):
            st.session_state[f"{prefixo}_cidade"] = dados_cep.get("localidade", "")
        if not st.session_state.get(f"{prefixo}_estado"):
            st.session_state[f"{prefixo}_estado"] = dados_cep.get("uf", "")

    opcoes_sexo = ["", "Feminino", "Masculino", "Outro"]
    mapa_sexo = {normalizar_texto(valor): valor for valor in opcoes_sexo if valor}
    sexo_atual = st.session_state.get(f"{prefixo}_sexo", "")
    st.session_state[f"{prefixo}_sexo"] = mapa_sexo.get(normalizar_texto(sexo_atual), "")

    opcoes_estado_civil = ["", "Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)", "União estável"]
    mapa_estado_civil = {normalizar_texto(valor): valor for valor in opcoes_estado_civil if valor}
    estado_civil_atual = st.session_state.get(f"{prefixo}_estado_civil", "")
    st.session_state[f"{prefixo}_estado_civil"] = mapa_estado_civil.get(
        normalizar_texto(estado_civil_atual),
        "",
    )

    st.markdown('<div class="ss-patient-form-section"><div class="ss-patient-form-title">Dados principais</div>', unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns([2.2, 1.2, 1.1, 1.2])
    c1.text_input("Nome *", key=f"{prefixo}_nome")
    c2.text_input("Apelido", key=f"{prefixo}_apelido")
    c3.selectbox("Sexo", options=opcoes_sexo, key=f"{prefixo}_sexo")
    c4.text_input("Nascimento", key=f"{prefixo}_data_nascimento")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="ss-patient-form-section"><div class="ss-patient-form-title">Contato</div>', unsafe_allow_html=True)
    c5, c6 = st.columns([1.2, 1.4])
    c5.text_input("Telefone", key=f"{prefixo}_telefone")
    c6.text_input("Email", key=f"{prefixo}_email")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="ss-patient-form-section"><div class="ss-patient-form-title">Documentos e cadastro</div>', unsafe_allow_html=True)
    c7, c8, c9, c10 = st.columns([1.2, 1.2, 1.1, 1.2])
    c7.text_input("CPF", key=f"{prefixo}_cpf")
    c8.text_input("RG", key=f"{prefixo}_rg")
    c9.text_input("Prontuário *", key=f"{prefixo}_prontuario")
    c10.selectbox("Estado civil", options=opcoes_estado_civil, key=f"{prefixo}_estado_civil")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="ss-patient-form-section"><div class="ss-patient-form-title">Endereço</div>', unsafe_allow_html=True)
    c11, c12 = st.columns([1, 3])
    c11.text_input("CEP", key=f"{prefixo}_cep")
    c12.text_input("Rua", key=f"{prefixo}_endereco")
    c13, c14, c15, c16 = st.columns([1, 1.4, 1.4, 0.8])
    c13.text_input("Número", key=f"{prefixo}_numero")
    c14.text_input("Bairro", key=f"{prefixo}_bairro")
    c15.text_input("Cidade", key=f"{prefixo}_cidade")
    c16.text_input("Estado", key=f"{prefixo}_estado")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="ss-patient-form-section"><div class="ss-patient-form-title">Complementar</div>', unsafe_allow_html=True)
    st.text_area("Observações", key=f"{prefixo}_observacoes", height=100)
    menor = st.checkbox("Paciente menor de idade", key=f"{prefixo}_menor")
    if menor:
        c17, c18 = st.columns(2)
        c17.text_input("Responsável legal", key=f"{prefixo}_responsavel")
        c18.text_input("CPF do responsável", key=f"{prefixo}_cpf_responsavel")
    st.markdown("</div>", unsafe_allow_html=True)

    return dados_formulario_paciente(prefixo)


def opcoes_contratos(df_contratos, df_pacientes):
    pacientes_por_id = {}
    for _, row in df_pacientes.iterrows():
        pacientes_por_id[int(row["id"])] = {
            "nome": row["nome"],
            "prontuario": row["prontuario"],
        }

    opcoes = []
    for _, row in df_contratos.iterrows():
        contrato_id = int(row["id"])
        paciente_id = int(row["paciente_id"]) if not pd.isna(row["paciente_id"]) else None
        paciente = pacientes_por_id.get(paciente_id, {"nome": "Paciente nao encontrado", "prontuario": ""})
        valor_total = float(row["valor_total"] or 0)
        forma_pagamento = row["forma_pagamento"] or ""
        descricao = (
            f"Contrato {contrato_id} - {paciente['nome']} - "
            f"Prontuario {paciente['prontuario']} - "
            f"R$ {valor_total:.2f} - {forma_pagamento}"
        )
        opcoes.append((contrato_id, descricao))
    return opcoes


def filtrar_opcoes_contratos(contratos_opcoes, termo_busca):
    termo = normalizar_texto(termo_busca)
    if not termo:
        return contratos_opcoes
    return [
        (contrato_id, descricao)
        for contrato_id, descricao in contratos_opcoes
        if termo in normalizar_texto(descricao)
    ]


def normalizar_nome_coluna_importacao(coluna):
    texto = normalizar_texto(coluna)
    texto = texto.replace(" ", "_")
    return texto


def normalizar_forma_pagamento_importacao(valor):
    texto = normalizar_texto(valor)
    if "boleto" in texto:
        return "Boleto"
    if "pix" in texto:
        return "Pix"
    if "debito" in texto:
        return "Debito"
    if "credito" in texto or "cartao" in texto or "cartao de credito" in texto or "a vista" in texto:
        return "Credito" if "credito" in texto or "cartao" in texto else "Dinheiro"
    if "dinheiro" in texto:
        return "Dinheiro"
    return "Boleto" if "parcela" in texto else "Credito"


def valor_float_importacao(valor):
    if pd.isna(valor):
        return 0.0
    if isinstance(valor, (int, float)):
        return round(float(valor), 2)
    texto = str(valor).strip()
    if not texto:
        return 0.0
    texto = texto.replace("R$", "").replace(".", "").replace(",", ".").strip()
    try:
        return round(float(texto), 2)
    except ValueError:
        return 0.0


def texto_importacao(valor):
    if pd.isna(valor):
        return ""
    texto = str(valor).strip()
    return "" if texto.lower() in {"nat", "nan", "none"} else texto


def data_importacao_para_br(valor):
    return formatar_data_br_valor(valor)


def montar_hash_importacao_contrato(prontuario, cpf, data_contrato, valor_total, forma_pagamento, procedimentos):
    base = "|".join(
        [
            str(prontuario or "").strip(),
            limpar_cpf(cpf),
            data_importacao_para_br(data_contrato),
            f"{valor_float_importacao(valor_total):.2f}",
            normalizar_forma_pagamento_importacao(forma_pagamento),
            "|".join([normalizar_texto(proc.get("procedimento", "")) + f":{float(proc.get('valor', 0)):.2f}" for proc in procedimentos]),
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def localizar_ou_criar_paciente_importacao(row):
    prontuario = texto_importacao(row.get("cdg"))
    cpf = limpar_cpf(row.get("cpf"))
    paciente_row = None

    if prontuario:
        paciente_row = cursor.execute("SELECT * FROM pacientes WHERE prontuario=? LIMIT 1", (prontuario,)).fetchone()
    if paciente_row is None and cpf:
        paciente_row = cursor.execute("SELECT * FROM pacientes WHERE cpf=? LIMIT 1", (cpf,)).fetchone()

    nome = texto_importacao(row.get("nome"))
    cep = texto_importacao(row.get("cep"))
    endereco = texto_importacao(row.get("rua"))
    numero = texto_importacao(row.get("numero"))
    complemento = texto_importacao(row.get("complemento"))
    if complemento:
        numero = f"{numero} {complemento}".strip()
    bairro = texto_importacao(row.get("bairro"))
    cidade = texto_importacao(row.get("cidade"))
    estado = texto_importacao(row.get("estado"))
    telefone = texto_importacao(row.get("telefone"))
    nascimento = data_importacao_para_br(row.get("nascimento"))
    menor_nome = texto_importacao(row.get("menor_nome"))
    menor_cpf = limpar_cpf(row.get("menor_cpf"))
    menor_nascimento = data_importacao_para_br(row.get("menor_nascimento"))

    if menor_nome:
        responsavel = nome
        cpf_responsavel = cpf
        nome_paciente = menor_nome
        cpf_paciente = menor_cpf
        nascimento_paciente = menor_nascimento
        menor_idade = 1
    else:
        responsavel = ""
        cpf_responsavel = ""
        nome_paciente = nome
        cpf_paciente = cpf
        nascimento_paciente = nascimento
        menor_idade = 0

    if paciente_row is None:
        cursor.execute(
            """
            INSERT INTO pacientes
            (nome, prontuario, cpf, data_nascimento, telefone, cep, endereco, numero, bairro, cidade, estado, menor_idade, responsavel, cpf_responsavel)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                nome_paciente,
                prontuario,
                cpf_paciente,
                nascimento_paciente,
                telefone,
                cep,
                endereco,
                numero,
                bairro,
                cidade,
                estado,
                menor_idade,
                responsavel,
                cpf_responsavel,
            ),
        )
        return cursor.lastrowid

    cursor.execute(
        """
        UPDATE pacientes
        SET nome=?, prontuario=?, cpf=?, data_nascimento=?, telefone=?, cep=?, endereco=?, numero=?, bairro=?, cidade=?, estado=?, menor_idade=?, responsavel=?, cpf_responsavel=?
        WHERE id=?
        """,
        (
            nome_paciente or paciente_row["nome"],
            prontuario or paciente_row["prontuario"],
            cpf_paciente or paciente_row["cpf"],
            nascimento_paciente or paciente_row["data_nascimento"],
            telefone or paciente_row["telefone"],
            cep or paciente_row["cep"],
            endereco or paciente_row["endereco"],
            numero or paciente_row["numero"],
            bairro or paciente_row["bairro"],
            cidade or paciente_row["cidade"],
            estado or paciente_row["estado"],
            menor_idade,
            responsavel or paciente_row["responsavel"],
            cpf_responsavel or paciente_row["cpf_responsavel"],
            int(paciente_row["id"]),
        ),
    )
    return int(paciente_row["id"])


def extrair_procedimentos_importacao(row):
    procedimentos = []
    for indice in range(1, 12):
        chave_proc = f"procedimento_{indice}"
        chave_valor = f"valor_{indice}"
        procedimento = texto_importacao(row.get(chave_proc))
        valor = valor_float_importacao(row.get(chave_valor))
        if procedimento:
            procedimentos.append({"procedimento": procedimento, "valor": valor})
    return procedimentos


def preparar_planilha_importacao_contratos(arquivo):
    df = pd.read_excel(arquivo, sheet_name="Planilha1")
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    contratos = []
    for _, row in df.iterrows():
        nome = texto_importacao(row.get("nome"))
        prontuario = texto_importacao(row.get("cdg"))
        if not nome or not prontuario:
            continue
        procedimentos = extrair_procedimentos_importacao(row)
        if not procedimentos:
            continue
        forma_pagamento = normalizar_forma_pagamento_importacao(row.get("forma_de_pagamento"))
        valor_total = valor_float_importacao(row.get("total") or row.get("valor"))
        entrada = valor_float_importacao(row.get("valor_de_entrada"))
        parcelas = int(valor_float_importacao(row.get("parcelas")) or 1)
        contratos.append(
            {
                "row": row.to_dict(),
                "nome": nome,
                "prontuario": prontuario,
                "cpf": limpar_cpf(row.get("cpf")),
                "data": data_importacao_para_br(row.get("data")),
                "forma_pagamento": forma_pagamento,
                "valor_total": valor_total,
                "entrada": entrada,
                "parcelas": parcelas if parcelas > 0 else 1,
                "data_entrada": (
                    data_importacao_para_br(row.get("data_da_entrada"))
                    if entrada > 0
                    else ""
                ),
                "primeiro_boleto": data_importacao_para_br(row.get("vencimento_do_1_boleto")),
                "procedimentos": procedimentos,
                "hash_importacao": montar_hash_importacao_contrato(
                    prontuario,
                    row.get("cpf"),
                    row.get("data"),
                    valor_total,
                    forma_pagamento,
                    procedimentos,
                ),
            }
        )
    return contratos


def localizar_contrato_existente_importacao(item, paciente_id):
    contrato_existente = cursor.execute(
        "SELECT * FROM contratos WHERE hash_importacao=? LIMIT 1",
        (item["hash_importacao"],),
    ).fetchone()
    if contrato_existente is not None:
        return contrato_existente

    data_contrato = formatar_data_br_valor(item["data"])
    forma_pagamento = item["forma_pagamento"]
    valor_total = float(item["valor_total"] or 0)

    consultas = []
    if paciente_id:
        consultas.append(
            (
                """
                SELECT *
                FROM contratos
                WHERE paciente_id=?
                  AND abs(valor_total - ?) < 0.01
                  AND forma_pagamento=?
                  AND COALESCE(data_criacao, '')=?
                ORDER BY id DESC
                LIMIT 1
                """,
                (int(paciente_id), valor_total, forma_pagamento, data_contrato),
            )
        )
        consultas.append(
            (
                """
                SELECT *
                FROM contratos
                WHERE paciente_id=?
                  AND abs(valor_total - ?) < 0.01
                  AND forma_pagamento=?
                ORDER BY id DESC
                LIMIT 1
                """,
                (int(paciente_id), valor_total, forma_pagamento),
            )
        )

    for sql, params in consultas:
        contrato_existente = cursor.execute(sql, params).fetchone()
        if contrato_existente is not None:
            return contrato_existente

    return None


def importar_contratos_preparados(contratos_preparados):
    resultado = {"importados": 0, "ignorados": 0, "atualizados": 0, "erros": [], "recebiveis_pendentes": []}
    for item in contratos_preparados:
        try:
            paciente_id = localizar_ou_criar_paciente_importacao(item["row"])
            contrato_existente = localizar_contrato_existente_importacao(item, paciente_id)

            if contrato_existente is not None:
                cursor.execute(
                    """
                    UPDATE contratos
                    SET paciente_id=?, valor_total=?, entrada=?, parcelas=?, primeiro_vencimento=?, data_pagamento_entrada=?, forma_pagamento=?, hash_importacao=?, data_criacao=?
                    WHERE id=?
                    """,
                    (
                        paciente_id,
                        item["valor_total"],
                        item["entrada"],
                        item["parcelas"],
                        item["primeiro_boleto"],
                        item["data_entrada"],
                        item["forma_pagamento"],
                        item["hash_importacao"],
                        formatar_data_br_valor(item["data"]),
                        int(contrato_existente["id"]),
                    ),
                )
                salvar_procedimentos_contrato(
                    int(contrato_existente["id"]),
                    [proc["procedimento"] for proc in item["procedimentos"]],
                    [proc["valor"] for proc in item["procedimentos"]],
                )

                if forma_pagamento_a_vista(item["forma_pagamento"]):
                    cursor.execute("DELETE FROM recebiveis WHERE contrato_id=?", (int(contrato_existente["id"]),))
                else:
                    if parse_data_contrato(item["primeiro_boleto"]):
                        sincronizar_recebiveis_contrato(
                            int(contrato_existente["id"]),
                            paciente_id,
                            item["nome"],
                            item["prontuario"],
                            item["valor_total"],
                            item["entrada"],
                            item["parcelas"],
                            item["primeiro_boleto"],
                            item["forma_pagamento"],
                        )
                    else:
                        resultado["recebiveis_pendentes"].append(
                            f"{item['nome']} / prontuario {item['prontuario']}"
                        )

                resultado["atualizados"] += 1
                continue

            cursor.execute(
                """
                INSERT INTO contratos
                (paciente_id, valor_total, entrada, parcelas, primeiro_vencimento, data_pagamento_entrada, forma_pagamento, hash_importacao, data_criacao)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    paciente_id,
                    item["valor_total"],
                    item["entrada"],
                    item["parcelas"],
                    item["primeiro_boleto"],
                    item["data_entrada"],
                    item["forma_pagamento"],
                    item["hash_importacao"],
                    item["data"],
                ),
            )
            contrato_id = cursor.lastrowid
            salvar_procedimentos_contrato(
                contrato_id,
                [proc["procedimento"] for proc in item["procedimentos"]],
                [proc["valor"] for proc in item["procedimentos"]],
            )
            if not forma_pagamento_a_vista(item["forma_pagamento"]):
                if parse_data_contrato(item["primeiro_boleto"]):
                    sincronizar_recebiveis_contrato(
                        contrato_id,
                        paciente_id,
                        item["nome"],
                        item["prontuario"],
                        item["valor_total"],
                        item["entrada"],
                        item["parcelas"],
                        item["primeiro_boleto"],
                        item["forma_pagamento"],
                    )
                else:
                    resultado["recebiveis_pendentes"].append(
                        f"{item['nome']} / prontuario {item['prontuario']}"
                    )
            resultado["importados"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['nome']} / prontuario {item['prontuario']}: {exc}")
    conn.commit()
    return resultado


def garantir_metas_vendas_iniciais(ano=None):
    ano_referencia = int(ano or date.today().year)
    row = cursor.execute("SELECT id FROM metas_vendas WHERE ano=? LIMIT 1", (ano_referencia,)).fetchone()
    if row is None:
        cursor.execute(
            """
            INSERT INTO metas_vendas
            (ano, meta, supermeta, hipermeta, data_atualizacao)
            VALUES (?, ?, ?, ?, ?)
            """,
            (ano_referencia, 100000.0, 150000.0, 200000.0, agora_str()),
        )
        conn.commit()


def carregar_metas_vendas(ano=None):
    ano_referencia = int(ano or date.today().year)
    garantir_metas_vendas_iniciais(ano_referencia)
    row = cursor.execute("SELECT * FROM metas_vendas WHERE ano=? LIMIT 1", (ano_referencia,)).fetchone()
    return dict(row) if row else {
        "ano": ano_referencia,
        "meta": 100000.0,
        "supermeta": 150000.0,
        "hipermeta": 200000.0,
    }


def salvar_metas_vendas(ano, meta, supermeta, hipermeta):
    garantir_metas_vendas_iniciais(ano)
    cursor.execute(
        """
        UPDATE metas_vendas
        SET meta=?, supermeta=?, hipermeta=?, data_atualizacao=?
        WHERE ano=?
        """,
        (float(meta), float(supermeta), float(hipermeta), agora_str(), int(ano)),
    )


def localizar_paciente_por_nome_importacao(nome, pacientes_df=None):
    nome_texto = texto_importacao(nome)
    if not nome_texto:
        return None
    pacientes = pacientes_df.copy() if pacientes_df is not None else carregar_pacientes()
    if pacientes.empty:
        return None
    if "_nome_norm" not in pacientes.columns:
        pacientes["_nome_norm"] = pacientes["nome"].apply(normalizar_texto)
    nome_norm = normalizar_texto(nome_texto)
    encontrados = pacientes[pacientes["_nome_norm"] == nome_norm]
    if not encontrados.empty:
        return encontrados.iloc[0]

    melhor_indice = None
    melhor_score = 0.0
    for indice, row in pacientes.iterrows():
        score = SequenceMatcher(None, nome_norm, row["_nome_norm"]).ratio()
        if score > melhor_score:
            melhor_score = score
            melhor_indice = indice

    if melhor_indice is not None and melhor_score >= 0.94:
        return pacientes.loc[melhor_indice]
    return None


def montar_hash_importacao_recebivel(nome, vencimento, valor):
    base = "|".join(
        [
            normalizar_texto(nome),
            data_importacao_para_br(vencimento),
            f"{valor_float_importacao(valor):.2f}",
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def normalizar_status_recebivel_importacao(status, pago):
    status_texto = normalizar_texto(status)
    pago_texto = normalizar_texto(pago)
    data_pagamento = data_importacao_para_br(pago) if parse_data_contrato(pago) else ""
    marcadores_pago = {"pago", "sim", "ok", "x", "quitado", "quitada"}
    texto_combinado = " ".join([status_texto, pago_texto]).strip()

    if data_pagamento or status_texto in marcadores_pago or pago_texto in marcadores_pago:
        return "Pago", data_pagamento
    if any(termo in texto_combinado for termo in ["suspens", "cancel", "abatid"]):
        return "Suspenso", ""
    if any(termo in texto_combinado for termo in ["atras", "devendo", "vencid", "inadimpl"]):
        return "Atrasado", ""
    return "Aberto", ""


def montar_observacao_recebivel_importacao(row, status_final):
    partes = []
    cobranca = texto_importacao(row.get("cobranca"))
    telefone = texto_importacao(row.get("telefone_da_cobranca"))
    status_original = texto_importacao(row.get("status"))
    pago_original = texto_importacao(row.get("pago"))

    if cobranca:
        partes.append(f"Cobranca: {cobranca}")
    if telefone:
        partes.append(f"Telefone cobranca: {telefone}")
    if status_original and normalizar_texto(status_original) not in {normalizar_texto(status_final), ""}:
        partes.append(f"Status original: {status_original}")
    if pago_original and not parse_data_contrato(pago_original) and normalizar_texto(pago_original) not in {"", "pago", "sim", "ok", "x"}:
        partes.append(f"Pago original: {pago_original}")
    return " | ".join(partes)


def preparar_planilha_importacao_recebiveis(arquivo):
    df = pd.read_excel(arquivo)
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    pacientes_df = carregar_pacientes()
    if not pacientes_df.empty:
        pacientes_df["_nome_norm"] = pacientes_df["nome"].apply(normalizar_texto)

    recebiveis_preparados = []
    for _, row in df.iterrows():
        nome = texto_importacao(row.get("paciente"))
        vencimento = data_importacao_para_br(row.get("data_venc"))
        valor = valor_float_importacao(row.get("valor"))
        if not nome or not parse_data_contrato(vencimento) or valor <= 0:
            continue

        paciente = localizar_paciente_por_nome_importacao(nome, pacientes_df)
        status_final, data_pagamento = normalizar_status_recebivel_importacao(row.get("status"), row.get("pago"))
        recebiveis_preparados.append(
            {
                "paciente_id": int(paciente["id"]) if paciente is not None else None,
                "paciente_nome": str(paciente["nome"]) if paciente is not None else nome,
                "prontuario": str(paciente["prontuario"] or "") if paciente is not None else "",
                "vencimento": vencimento,
                "valor": valor,
                "forma_pagamento": "Boleto",
                "status": status_final,
                "data_pagamento": data_pagamento,
                "observacao": montar_observacao_recebivel_importacao(row, status_final),
                "hash_importacao": montar_hash_importacao_recebivel(nome, vencimento, valor),
            }
        )
    return recebiveis_preparados


def importar_recebiveis_preparados(recebiveis_preparados):
    resultado = {"inseridos": 0, "atualizados": 0, "erros": []}
    for item in recebiveis_preparados:
        try:
            existente = cursor.execute(
                "SELECT id FROM recebiveis WHERE hash_importacao=? LIMIT 1",
                (item["hash_importacao"],),
            ).fetchone()
            if existente is None:
                existente = cursor.execute(
                    """
                    SELECT id
                    FROM recebiveis
                    WHERE lower(trim(paciente_nome)) = lower(trim(?))
                      AND vencimento = ?
                      AND abs(valor - ?) < 0.01
                    LIMIT 1
                    """,
                    (item["paciente_nome"], item["vencimento"], float(item["valor"])),
                ).fetchone()

            if existente is None:
                cursor.execute(
                    """
                    INSERT INTO recebiveis
                    (contrato_id, paciente_id, paciente_nome, prontuario, parcela_numero, vencimento, valor, forma_pagamento, status, observacao, data_criacao, data_pagamento, hash_importacao)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        None,
                        item["paciente_id"],
                        item["paciente_nome"],
                        item["prontuario"],
                        None,
                        item["vencimento"],
                        item["valor"],
                        item["forma_pagamento"],
                        item["status"],
                        item["observacao"],
                        agora_str(),
                        item["data_pagamento"],
                        item["hash_importacao"],
                    ),
                )
                resultado["inseridos"] += 1
            else:
                cursor.execute(
                    """
                    UPDATE recebiveis
                    SET paciente_id=?, paciente_nome=?, prontuario=?, vencimento=?, valor=?, forma_pagamento=?, status=?, observacao=?, data_pagamento=?, hash_importacao=?
                    WHERE id=?
                    """,
                    (
                        item["paciente_id"],
                        item["paciente_nome"],
                        item["prontuario"],
                        item["vencimento"],
                        item["valor"],
                        item["forma_pagamento"],
                        item["status"],
                        item["observacao"],
                        item["data_pagamento"],
                        item["hash_importacao"],
                        int(existente["id"]),
                    ),
                )
                resultado["atualizados"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['paciente_nome']} / {item['vencimento']}: {exc}")
    conn.commit()
    return resultado


def montar_hash_importacao_conta_pagar(data_vencimento, descricao, fornecedor, valor):
    base = "|".join(
        [
            data_importacao_para_br(data_vencimento),
            normalizar_texto(descricao),
            normalizar_texto(fornecedor),
            f"{valor_float_importacao(valor):.2f}",
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def sugerir_categoria_conta_pagar(descricao, fornecedor):
    texto = normalizar_texto(f"{descricao} {fornecedor}")

    regras = [
        ("Impostos", ["imposto", "tribut", "darf", "darj", "fgts", "inss", "iss", "simples", "irrf", "gps"]),
        ("Funcionarios", ["funcionario", "salario", "folha", "ferias", "rescis", "pro labore", "adiantamento", "decimo", "13"]),
        ("Aluguel", ["aluguel", "locacao", "locacao", "condominio"]),
        ("Tarifas bancarias", ["tarifa", "juros", "multa", "banco", "infinitepay", "pagseguro", "maquininha", "antecipacao"]),
        ("Marketing", ["marketing", "trafego", "google", "meta", "instagram", "facebook", "impulsionamento", "anuncio"]),
        ("Laboratorio", ["laboratorio", "protese", "ortodont", "neodent"]),
        ("Materia-prima", ["material", "materia prima", "dental", "odonto", "insumo"]),
        ("Equipamentos", ["equipamento", "aparelho", "cadeira", "compressor", "raio x", "autoclave"]),
        ("Manutencao", ["manutencao", "conserto", "reparo", "assistencia tecnica"]),
        ("Servicos", ["contador", "contabilidade", "clinicorp", "sistema", "software", "limpeza", "consultoria", "advogado"]),
        ("Comissoes", ["comissao", "percentual"]),
        ("Custo fixo", ["energia", "luz", "agua", "telefone", "internet", "gas"]),
    ]

    for categoria, termos in regras:
        if any(normalizar_texto(termo) in texto for termo in termos):
            return categoria

    if texto.strip():
        return "Fornecedores"
    return "Outros"


def categorizar_contas_pagar_existentes():
    contas = cursor.execute(
        """
        SELECT id, descricao, fornecedor, categoria
        FROM contas_pagar
        """
    ).fetchall()

    houve_atualizacao = False
    for row in contas:
        if str(row["categoria"] or "").strip():
            continue
        categoria_sugerida = sugerir_categoria_conta_pagar(row["descricao"], row["fornecedor"])
        cursor.execute(
            "UPDATE contas_pagar SET categoria=? WHERE id=?",
            (categoria_sugerida, int(row["id"])),
        )
        houve_atualizacao = True

    if houve_atualizacao:
        conn.commit()


def normalizar_status_conta_pagar_importacao(status, pago):
    status_texto = normalizar_texto(status)
    pago_texto = normalizar_texto(pago)
    if parse_data_contrato(pago) or pago_texto in {"sim", "pago", "ok", "x"} or status_texto == "pago":
        return "Pago", data_importacao_para_br(pago) if parse_data_contrato(pago) else ""
    if "atras" in status_texto:
        return "Atrasado", ""
    if any(termo in status_texto for termo in ["cancel", "suspens"]):
        return "Suspenso", ""
    return "A vencer", ""


def preparar_planilha_importacao_contas_pagar(arquivo):
    colunas_esperadas = {"data_de_vencimento", "descricao", "fornecedor", "valor"}
    df = pd.read_excel(arquivo, header=1)
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    if not colunas_esperadas.issubset(set(df.columns)):
        arquivo.seek(0)
        df = pd.read_excel(arquivo)
        df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    contas_preparadas = []
    for _, row in df.iterrows():
        data_vencimento = data_importacao_para_br(row.get("data_de_vencimento"))
        descricao = texto_importacao(row.get("descricao"))
        fornecedor = texto_importacao(row.get("fornecedor"))
        valor = valor_float_importacao(row.get("valor"))
        if not parse_data_contrato(data_vencimento) or (not descricao and not fornecedor) or valor <= 0:
            continue
        status_final, data_pagamento = normalizar_status_conta_pagar_importacao(row.get("status"), row.get("pago"))
        valor_pago = valor_float_importacao(row.get("valor_pago")) or (valor if status_final == "Pago" else 0.0)
        observacao = ""
        if texto_importacao(row.get("status")) and normalizar_texto(row.get("status")) not in {normalizar_texto(status_final), ""}:
            observacao = f"Status original: {texto_importacao(row.get('status'))}"
        contas_preparadas.append(
            {
                "data_vencimento": data_vencimento,
                "descricao": descricao,
                "fornecedor": fornecedor,
                "categoria": sugerir_categoria_conta_pagar(descricao, fornecedor),
                "valor": valor,
                "pago": data_pagamento,
                "valor_pago": valor_pago,
                "status": status_final,
                "observacao": observacao,
                "hash_importacao": montar_hash_importacao_conta_pagar(data_vencimento, descricao, fornecedor, valor),
            }
        )
    return contas_preparadas


def importar_contas_pagar_preparadas(contas_preparadas):
    resultado = {"inseridos": 0, "atualizados": 0, "erros": []}
    for item in contas_preparadas:
        try:
            existente = cursor.execute(
                "SELECT id FROM contas_pagar WHERE hash_importacao=? LIMIT 1",
                (item["hash_importacao"],),
            ).fetchone()
            if existente is None:
                cursor.execute(
                    """
                    INSERT INTO contas_pagar
                    (data_vencimento, descricao, fornecedor, categoria, valor, pago, valor_pago, status, observacao, data_criacao, hash_importacao)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        item["data_vencimento"],
                        item["descricao"],
                        item["fornecedor"],
                        item["categoria"],
                        item["valor"],
                        item["pago"],
                        item["valor_pago"],
                        item["status"],
                        item["observacao"],
                        agora_str(),
                        item["hash_importacao"],
                    ),
                )
                resultado["inseridos"] += 1
            else:
                cursor.execute(
                    """
                    UPDATE contas_pagar
                    SET data_vencimento=?, descricao=?, fornecedor=?, categoria=?, valor=?, pago=?, valor_pago=?, status=?, observacao=?, hash_importacao=?
                    WHERE id=?
                    """,
                    (
                        item["data_vencimento"],
                        item["descricao"],
                        item["fornecedor"],
                        item["categoria"],
                        item["valor"],
                        item["pago"],
                        item["valor_pago"],
                        item["status"],
                        item["observacao"],
                        item["hash_importacao"],
                        int(existente["id"]),
                    ),
                )
                resultado["atualizados"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['fornecedor']} / {item['descricao']}: {exc}")
    conn.commit()
    return resultado


def montar_hash_importacao_venda(data_venda, paciente_nome, valor_total, nf):
    base = "|".join(
        [
            data_importacao_para_br(data_venda),
            normalizar_texto(paciente_nome),
            f"{valor_float_importacao(valor_total):.2f}",
            texto_importacao(nf),
        ]
    )
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def preparar_planilha_importacao_vendas(arquivo):
    planilhas = pd.ExcelFile(arquivo)
    nome_planilha = "Planilha2" if "Planilha2" in planilhas.sheet_names else planilhas.sheet_names[0]
    arquivo.seek(0)
    df = pd.read_excel(arquivo, sheet_name=nome_planilha)
    df.columns = [normalizar_nome_coluna_importacao(coluna) for coluna in df.columns]
    df = df.dropna(how="all")
    vendas_preparadas = []
    for _, row in df.iterrows():
        data_venda = data_importacao_para_br(row.get("data"))
        paciente_nome = texto_importacao(row.get("nome_paciente"))
        valor_total = valor_float_importacao(row.get("valor_total"))
        if not parse_data_contrato(data_venda) or not paciente_nome or valor_total <= 0:
            continue
        valor_a_vista = valor_float_importacao(row.get("a_vista"))
        valor_cartao = valor_float_importacao(row.get("cartao"))
        valor_boleto = valor_float_importacao(row.get("boleto"))
        saldo = valor_float_importacao(row.get("saldo"))
        data_a_pagar = data_importacao_para_br(row.get("data_a_pagar"))
        vendas_preparadas.append(
            {
                "data_venda": data_venda,
                "paciente_nome": paciente_nome,
                "valor_total": valor_total,
                "valor_a_vista": valor_a_vista,
                "valor_cartao": valor_cartao,
                "valor_boleto": valor_boleto,
                "saldo": saldo,
                "data_a_pagar": data_a_pagar,
                "avaliador": texto_importacao(row.get("avaliador")),
                "vendedor": texto_importacao(row.get("vendedor")),
                "nf": texto_importacao(row.get("nf")),
                "hash_importacao": montar_hash_importacao_venda(data_venda, paciente_nome, valor_total, row.get("nf")),
            }
        )
    return vendas_preparadas


def criar_contrato_automatico_venda(item, pacientes_df=None):
    paciente = localizar_paciente_por_nome_importacao(item["paciente_nome"], pacientes_df)
    if paciente is None:
        return None

    forma_pagamento = "Boleto" if item["saldo"] > 0 or item["valor_boleto"] > 0 else ("Credito" if item["valor_cartao"] > 0 else "Dinheiro")
    entrada = round(float(item["valor_a_vista"] or 0) + float(item["valor_cartao"] or 0), 2)
    primeiro_vencimento = item["data_a_pagar"] if not forma_pagamento_a_vista(forma_pagamento) else ""
    data_pagamento_entrada = item["data_venda"]

    contrato_hash = item["hash_importacao"]
    existente = cursor.execute("SELECT id FROM contratos WHERE hash_importacao=? LIMIT 1", (contrato_hash,)).fetchone()
    if existente is not None:
        return int(existente["id"])

    cursor.execute(
        """
        INSERT INTO contratos
        (paciente_id, valor_total, entrada, parcelas, primeiro_vencimento, data_pagamento_entrada, forma_pagamento, hash_importacao, data_criacao)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            int(paciente["id"]),
            item["valor_total"],
            entrada,
            1,
            primeiro_vencimento,
            data_pagamento_entrada,
            forma_pagamento,
            contrato_hash,
            item["data_venda"],
        ),
    )
    contrato_id = cursor.lastrowid
    salvar_procedimentos_contrato(contrato_id, ["Venda importada"], [float(item["valor_total"])])
    sincronizar_financeiro_contrato(contrato_id, item["paciente_nome"], item["valor_total"])
    if not forma_pagamento_a_vista(forma_pagamento) and item["saldo"] > 0:
        sincronizar_recebiveis_contrato(
            contrato_id,
            int(paciente["id"]),
            item["paciente_nome"],
            str(paciente["prontuario"] or ""),
            item["valor_total"],
            entrada,
            1,
            primeiro_vencimento,
            forma_pagamento,
        )
    return contrato_id


def importar_vendas_preparadas(vendas_preparadas, criar_contratos=True):
    resultado = {"inseridas": 0, "ignoradas": 0, "contratos_criados": 0, "erros": []}
    pacientes_df = carregar_pacientes()
    if not pacientes_df.empty:
        pacientes_df["_nome_norm"] = pacientes_df["nome"].apply(normalizar_texto)

    for item in vendas_preparadas:
        if cursor.execute("SELECT id FROM vendas WHERE hash_importacao=? LIMIT 1", (item["hash_importacao"],)).fetchone():
            resultado["ignoradas"] += 1
            continue
        try:
            contrato_id = None
            if criar_contratos:
                contrato_id = criar_contrato_automatico_venda(item, pacientes_df)
                if contrato_id:
                    resultado["contratos_criados"] += 1
            cursor.execute(
                """
                INSERT INTO vendas
                (data_venda, paciente_nome, valor_total, valor_a_vista, valor_cartao, valor_boleto, saldo, data_a_pagar, avaliador, vendedor, nf, contrato_id, hash_importacao, data_criacao)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    item["data_venda"],
                    item["paciente_nome"],
                    item["valor_total"],
                    item["valor_a_vista"],
                    item["valor_cartao"],
                    item["valor_boleto"],
                    item["saldo"],
                    item["data_a_pagar"],
                    item["avaliador"],
                    item["vendedor"],
                    item["nf"],
                    contrato_id,
                    item["hash_importacao"],
                    agora_str(),
                ),
            )
            resultado["inseridas"] += 1
        except Exception as exc:
            resultado["erros"].append(f"{item['paciente_nome']} / {item['data_venda']}: {exc}")
    conn.commit()
    return resultado


def inicializar_banco_e_arquivos():
    inicializar_banco()
    garantir_metas_vendas_iniciais()
    garantir_pasta_documentos()
    normalizar_status_recebiveis_por_observacao()
    categorizar_contas_pagar_existentes()


inicializar_banco_e_arquivos()
aplicar_tema_visual()
os.makedirs(ASSETS_DIR, exist_ok=True)

if "usuario_logado" not in st.session_state:
    st.session_state["usuario_logado"] = None

if not st.session_state["usuario_logado"]:
    renderizar_login()
    st.stop()


def navegar_para_menu(menu_destino, **estado_extra):
    st.session_state["_menu_sidebar_pendente"] = menu_destino
    for chave, valor in estado_extra.items():
        st.session_state[chave] = valor
    st.rerun()


usuario_logado = st.session_state["usuario_logado"]
menus_disponiveis = [menu_nome for menu_nome in MODULOS_SISTEMA if usuario_tem_acesso(usuario_logado, menu_nome)]
menu_pendente = st.session_state.pop("_menu_sidebar_pendente", None)
if menu_pendente in menus_disponiveis:
    st.session_state["menu_sidebar"] = menu_pendente
renderizar_sidebar_marca()
menu = st.sidebar.radio(
    "Menu",
    [nome for nome in ORDEM_MENU_SIDEBAR if nome in menus_disponiveis],
    format_func=lambda nome: ROTULOS_MENU.get(nome, nome),
    label_visibility="collapsed",
    key="menu_sidebar",
)
exibir_feedback_visual()

st.sidebar.markdown('<div class="ss-sidebar-spacer"></div>', unsafe_allow_html=True)
st.sidebar.markdown('<div class="ss-sidebar-footer">', unsafe_allow_html=True)
st.sidebar.markdown(f'<div class="ss-user-meta">Usuário: {usuario_logado["nome"]}</div>', unsafe_allow_html=True)
st.sidebar.markdown(f'<div class="ss-user-meta">Perfil: {usuario_logado["perfil"]}</div>', unsafe_allow_html=True)

with st.sidebar.expander("Minha senha"):
    senha_atual = st.text_input("Senha atual", type="password", key="senha_atual_sidebar")
    nova_senha = st.text_input("Nova senha", type="password", key="nova_senha_sidebar")
    confirmar_nova_senha = st.text_input("Confirmar nova senha", type="password", key="confirmar_nova_senha_sidebar")
    if st.button("Alterar minha senha"):
        usuario_db = cursor.execute("SELECT * FROM usuarios WHERE id=?", (int(usuario_logado["id"]),)).fetchone()
        if usuario_db is None:
            st.sidebar.error("Usuario nao encontrado.")
        elif not verificar_senha(senha_atual, usuario_db["senha_hash"]):
            st.sidebar.error("Senha atual invalida.")
        elif len(nova_senha) < 4:
            st.sidebar.error("A nova senha deve ter pelo menos 4 caracteres.")
        elif nova_senha != confirmar_nova_senha:
            st.sidebar.error("A confirmacao da nova senha nao confere.")
        else:
            redefinir_senha_usuario(usuario_logado["id"], nova_senha)
            conn.commit()
            st.sidebar.success("Senha alterada com sucesso.")

if st.sidebar.button("Sair", use_container_width=True):
    registrar_log_acesso(usuario_logado["id"], usuario_logado["usuario"], "LOGOUT")
    st.session_state["usuario_logado"] = None
    st.rerun()
st.sidebar.markdown("</div>", unsafe_allow_html=True)

if menu == "Agenda":
    renderizar_agenda_clinica(usuario_logado)
    st.stop()

if menu == "Dashboard":
    meses_pt_br = {
        1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
        7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
    }
    ordem_meses = [
        "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
        "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
    ]
    ano_dashboard = st.selectbox(
        "Ano das vendas",
        options=list(range(date.today().year - 2, date.today().year + 3)),
        index=2,
        key="dashboard_ano_visual",
    )
    metas = carregar_metas_vendas(ano_dashboard)
    vendas = pd.read_sql("SELECT * FROM vendas ORDER BY data_venda DESC", conn)
    atualizar_status_contas_pagar_automaticamente()
    contas_pagar_dashboard = pd.read_sql(
        "SELECT * FROM contas_pagar ORDER BY data_vencimento, fornecedor, descricao",
        conn,
    )
    recebiveis_dashboard = pd.read_sql(
        "SELECT * FROM recebiveis ORDER BY vencimento, paciente_nome",
        conn,
    )

    st.markdown(
        f"""
        <div class="ss-dashboard-shell">
            <div class="ss-dashboard-header">
                <div class="ss-search">Buscar paciente, contrato, venda ou vencimento...</div>
                <div class="ss-userbox">Bem-vinda, {usuario_logado["nome"]}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    dashboard_blocos = selecionar_blocos_visiveis(
        "Exibir no dashboard",
        [
            "Indicadores",
            "Evolução mensal",
            "Resumo do dia",
            "Alertas",
            "Calendário de pagamentos",
            "Atividades recentes",
            "Editar venda importada",
        ],
        [
            "Indicadores",
            "Evolução mensal",
            "Resumo do dia",
            "Alertas",
        ],
        "dashboard_blocos_visiveis",
    )

    with st.expander("Metas de vendas", expanded=False):
        m1, m2, m3 = st.columns(3)
        meta = m1.number_input("Meta", min_value=0.0, value=float(metas["meta"] or 100000), key="dash_meta")
        supermeta = m2.number_input("Supermeta", min_value=0.0, value=float(metas["supermeta"] or 150000), key="dash_supermeta")
        hipermeta = m3.number_input("Hipermeta", min_value=0.0, value=float(metas["hipermeta"] or 200000), key="dash_hipermeta")
        if st.button("Salvar metas de vendas", key="dash_salvar_metas"):
            salvar_metas_vendas(ano_dashboard, meta, supermeta, hipermeta)
            conn.commit()
            st.success("Metas atualizadas.")
            st.rerun()

    if vendas.empty:
        st.info("Nenhuma venda importada ainda.")
        st.stop()

    vendas["data_ref"] = vendas["data_venda"].apply(parse_data_contrato)
    vendas = vendas[vendas["data_ref"].notna()].copy()
    vendas["ano"] = vendas["data_ref"].apply(lambda valor: valor.year)
    vendas = vendas[vendas["ano"] == ano_dashboard].copy()

    if vendas.empty:
        st.info("Não há vendas para o ano selecionado.")
        st.stop()

    vendas["mes"] = vendas["data_ref"].apply(lambda valor: valor.month)
    vendas["mes_nome"] = vendas["mes"].map(meses_pt_br)

    filtro_col1, filtro_col2, filtro_col3, filtro_col4 = st.columns(4)
    filtro_pacientes_venda = filtro_col1.multiselect(
        "Pacientes",
        options=sorted([valor for valor in vendas["paciente_nome"].fillna("").unique().tolist() if valor]),
        key="dash_filtro_pacientes",
    )
    filtro_avaliadores_venda = filtro_col2.multiselect(
        "Avaliadores",
        options=sorted([valor for valor in vendas["avaliador"].fillna("").unique().tolist() if valor]),
        key="dash_filtro_avaliadores",
    )
    filtro_vendedores_venda = filtro_col3.multiselect(
        "Vendedores",
        options=sorted([valor for valor in vendas["vendedor"].fillna("").unique().tolist() if valor]),
        key="dash_filtro_vendedores",
    )
    filtro_meses_venda = filtro_col4.multiselect(
        "Meses",
        options=list(range(1, 13)),
        format_func=lambda valor: meses_pt_br[valor],
        key="dash_filtro_meses",
    )

    vendas_filtradas = vendas.copy()
    if filtro_pacientes_venda:
        vendas_filtradas = vendas_filtradas[vendas_filtradas["paciente_nome"].isin(filtro_pacientes_venda)]
    if filtro_avaliadores_venda:
        vendas_filtradas = vendas_filtradas[vendas_filtradas["avaliador"].isin(filtro_avaliadores_venda)]
    if filtro_vendedores_venda:
        vendas_filtradas = vendas_filtradas[vendas_filtradas["vendedor"].isin(filtro_vendedores_venda)]
    if filtro_meses_venda:
        vendas_filtradas = vendas_filtradas[vendas_filtradas["mes"].isin(filtro_meses_venda)]

    if vendas_filtradas.empty:
        st.info("Não há vendas para os filtros selecionados.")
        st.stop()

    hoje = date.today()
    vendas_mes_atual = vendas_filtradas[
        (vendas_filtradas["mes"] == hoje.month) & (vendas_filtradas["ano"] == hoje.year)
    ]
    total_mes = float(vendas_mes_atual["valor_total"].sum()) if not vendas_mes_atual.empty else 0.0
    pacientes_unicos = int(vendas_filtradas["paciente_nome"].fillna("").replace("", pd.NA).dropna().nunique())

    procedimentos_mes = pd.read_sql(
        """
        SELECT COUNT(1) AS total
        FROM procedimentos_contrato pc
        JOIN contratos c ON c.id = pc.contrato_id
        WHERE substr(c.data_criacao, 7, 4) = ?
          AND substr(c.data_criacao, 4, 2) = ?
        """,
        conn,
        params=(str(hoje.year), f"{hoje.month:02d}"),
    )
    total_procedimentos_mes = int(procedimentos_mes.iloc[0]["total"] or 0) if not procedimentos_mes.empty else 0

    pagamentos_hoje = 0.0
    recebimentos_hoje = 0.0
    inadimplencia_total = 0

    alertas = []
    if not contas_pagar_dashboard.empty:
        contas_pagar_dashboard["_data_ref"] = contas_pagar_dashboard["data_vencimento"].apply(parse_data_contrato)
        pagamentos_hoje = float(
            contas_pagar_dashboard[
                contas_pagar_dashboard["data_vencimento"] == formatar_data_br(hoje)
            ]["valor"].fillna(0).astype(float).sum()
        )
        atrasadas = contas_pagar_dashboard[
            (contas_pagar_dashboard["status"] == "Atrasado")
            | (
                contas_pagar_dashboard["_data_ref"].notna()
                & (contas_pagar_dashboard["_data_ref"] < hoje)
                & (~contas_pagar_dashboard["status"].isin(["Pago", "Suspenso", "Cancelado"]))
            )
        ]
        if not atrasadas.empty:
            alertas.append(f"{len(atrasadas)} contas a pagar atrasadas")
    if not recebiveis_dashboard.empty:
        recebiveis_dashboard["_data_ref"] = recebiveis_dashboard["vencimento"].apply(parse_data_contrato)
        recebimentos_hoje = float(
            recebiveis_dashboard[
                recebiveis_dashboard["vencimento"] == formatar_data_br(hoje)
            ]["valor"].fillna(0).astype(float).sum()
        )
        recebiveis_atrasados = recebiveis_dashboard[
            (recebiveis_dashboard["_data_ref"].notna())
            & (recebiveis_dashboard["_data_ref"] < hoje)
            & (~recebiveis_dashboard["status"].isin(["Pago", "Suspenso", "Cancelado"]))
        ]
        inadimplencia_total = len(recebiveis_atrasados)
        if not recebiveis_atrasados.empty:
            alertas.append(f"{len(recebiveis_atrasados)} recebíveis em atraso")

    card1, card2, card3, card4 = st.columns(4)
    card1.markdown(
        f'<div class="ss-card"><div class="ss-card-label">Faturamento do mês</div><div class="ss-card-value">{formatar_moeda_br(total_mes)}</div></div>',
        unsafe_allow_html=True,
    )
    card2.markdown(
        f'<div class="ss-card"><div class="ss-card-label">Pacientes atendidos</div><div class="ss-card-value">{pacientes_unicos}</div></div>',
        unsafe_allow_html=True,
    )
    card3.markdown(
        f'<div class="ss-card"><div class="ss-card-label">Procedimentos do mês</div><div class="ss-card-value">{total_procedimentos_mes}</div></div>',
        unsafe_allow_html=True,
    )
    card4.markdown(
        f'<div class="ss-card"><div class="ss-card-label">Recebimentos de hoje</div><div class="ss-card-value">{formatar_moeda_br(recebimentos_hoje)}</div></div>',
        unsafe_allow_html=True,
    )

    centro, lateral = st.columns([2.25, 1], gap="large")

    with centro:
        st.markdown('<div class="ss-panel">', unsafe_allow_html=True)
        st.subheader("Evolução mensal")
        resumo_mensal = vendas_filtradas.groupby(["mes", "mes_nome"], as_index=False)["valor_total"].sum().sort_values("mes")
        resumo_mensal = resumo_mensal.rename(columns={"mes_nome": "Mes", "valor_total": "Valor total"})
        resumo_mensal["Mes"] = pd.Categorical(resumo_mensal["Mes"], categories=ordem_meses, ordered=True)
        resumo_mensal = resumo_mensal.sort_values("Mes")

        metas_linhas = pd.DataFrame(
            [
                {"Meta": "Meta", "Valor": float(metas["meta"] or 0)},
                {"Meta": "Supermeta", "Valor": float(metas["supermeta"] or 0)},
                {"Meta": "Hipermeta", "Valor": float(metas["hipermeta"] or 0)},
            ]
        )
        barras = alt.Chart(resumo_mensal).mark_bar(cornerRadiusTopLeft=8, cornerRadiusTopRight=8, color="#c5a77a").encode(
            x=alt.X("Mes:N", sort=ordem_meses, title="Mês"),
            y=alt.Y("Valor total:Q", title="Valor"),
            tooltip=[
                alt.Tooltip("Mes:N", title="Mês"),
                alt.Tooltip("Valor total:Q", title="Valor", format=",.2f"),
            ],
        )
        linhas = alt.Chart(metas_linhas).mark_rule(strokeWidth=2).encode(
            y=alt.Y("Valor:Q"),
            color=alt.Color(
                "Meta:N",
                scale=alt.Scale(
                    domain=["Meta", "Supermeta", "Hipermeta"],
                    range=["#111111", "#7b5d34", "#2f855a"],
                ),
                legend=alt.Legend(title="Metas"),
            ),
            tooltip=[
                alt.Tooltip("Meta:N", title="Meta"),
                alt.Tooltip("Valor:Q", title="Valor", format=",.2f"),
            ],
        )
        st.altair_chart((barras + linhas).properties(height=320), use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    with lateral:
        st.markdown('<div class="ss-panel">', unsafe_allow_html=True)
        st.subheader("Resumo financeiro de hoje")
        if True:
            st.markdown(
                f'<div class="ss-list-item"><strong>Pagamentos</strong><br><span style="color:#6b6b6b;">{formatar_moeda_br(pagamentos_hoje)}</span></div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                f'<div class="ss-list-item"><strong>Recebimentos</strong><br><span style="color:#6b6b6b;">{formatar_moeda_br(recebimentos_hoje)}</span></div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                f'<div class="ss-list-item"><strong>Inadimplencias</strong><br><span style="color:#6b6b6b;">{inadimplencia_total}</span></div>',
                unsafe_allow_html=True,
            )
        else:
            for _, row in agenda_hoje.head(6).iterrows():
                horario = formatar_hora_agendamento(row["hora_inicio"]) or "--:--"
                paciente = row["paciente_nome"] or "Paciente não informado"
                procedimento = row["procedimento"] or "Sem procedimento"
                st.markdown(
                    f'<div class="ss-list-item"><strong>{horario}</strong> &nbsp; {paciente}<br><span style="color:#6b6b6b;">{procedimento}</span></div>',
                    unsafe_allow_html=True,
                )
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="ss-panel">', unsafe_allow_html=True)
        st.subheader("Alertas")
        if alertas:
            for alerta in alertas:
                st.warning(alerta)
        else:
            st.success("Nenhum alerta crítico no momento.")
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="ss-panel">', unsafe_allow_html=True)
    st.subheader("Calendário de pagamentos")
    if contas_pagar_dashboard.empty:
        st.info("Nenhum pagamento cadastrado para exibir no calendário.")
    else:
        contas_pagar_dashboard["valor"] = contas_pagar_dashboard["valor"].fillna(0).astype(float)
        contas_pagar_dashboard["_ordem_vencimento"] = pd.to_datetime(
            contas_pagar_dashboard["data_vencimento"],
            format="%d/%m/%Y",
            errors="coerce",
        )
        contas_calendario = contas_pagar_dashboard[contas_pagar_dashboard["_ordem_vencimento"].notna()].copy()
        if contas_calendario.empty:
            st.info("Nenhum vencimento válido para exibir no calendário.")
        else:
            anos_calendario = sorted(set(contas_calendario["_ordem_vencimento"].dt.year.tolist()) | {date.today().year})
            mes_atual = date.today().month
            ano_atual = date.today().year

            cal1, cal2 = st.columns(2)
            ano_pagamentos = cal1.selectbox(
                "Ano do calendário",
                options=anos_calendario,
                index=anos_calendario.index(ano_atual) if ano_atual in anos_calendario else 0,
                key="dashboard_ano_pagamentos_visual",
            )
            mes_pagamentos = cal2.selectbox(
                "Mês do calendário",
                options=list(range(1, 13)),
                index=mes_atual - 1,
                format_func=lambda valor: meses_pt_br[valor],
                key="dashboard_mes_pagamentos_visual",
            )

            contas_mes = contas_calendario[
                (contas_calendario["_ordem_vencimento"].dt.year == ano_pagamentos)
                & (contas_calendario["_ordem_vencimento"].dt.month == mes_pagamentos)
            ].copy()

            cabecalho = st.columns(7)
            for coluna, nome_dia in zip(cabecalho, ["Seg", "Ter", "Qua", "Qui", "Sex", "Sáb", "Dom"]):
                coluna.markdown(f"**{nome_dia}**")

            for semana in calendar.monthcalendar(ano_pagamentos, mes_pagamentos):
                cols = st.columns(7)
                for idx, dia in enumerate(semana):
                    with cols[idx]:
                        if dia == 0:
                            st.write("")
                        else:
                            dia_contas = contas_mes[contas_mes["_ordem_vencimento"].dt.day == dia].copy()
                            total_dia = float(dia_contas["valor"].sum()) if not dia_contas.empty else 0.0
                            qtd_dia = len(dia_contas)
                            qtd_pagos = int((dia_contas["status"] == "Pago").sum()) if not dia_contas.empty else 0
                            qtd_atrasados = int((dia_contas["status"] == "Atrasado").sum()) if not dia_contas.empty else 0
                            eh_hoje = hoje.year == ano_pagamentos and hoje.month == mes_pagamentos and hoje.day == dia
                            borda = "2px solid #111111" if eh_hoje else "1px solid #e9dfd2"
                            fundo = "#ffffff"
                            if qtd_atrasados > 0:
                                fundo = "#fff1f1"
                            elif qtd_dia > 0 and qtd_pagos == qtd_dia:
                                fundo = "#eefbf2"
                            elif qtd_dia > 0:
                                fundo = "#fbf7f1"
                            resumo_status = ""
                            if qtd_atrasados > 0:
                                resumo_status = f"<div style='font-size:11px;color:#b42318;'>Atrasados: {qtd_atrasados}</div>"
                            elif qtd_dia > 0 and qtd_pagos == qtd_dia:
                                resumo_status = "<div style='font-size:11px;color:#027a48;'>Todos pagos</div>"
                            elif qtd_dia > 0:
                                resumo_status = f"<div style='font-size:11px;color:#7b5d34;'>Títulos: {qtd_dia}</div>"
                            st.markdown(
                                f"""
                                <div style="border:{borda};border-radius:14px;padding:8px;min-height:86px;background:{fundo};">
                                    <div style="font-weight:700;font-size:14px;margin-bottom:4px;">{dia:02d}</div>
                                    <div style="font-size:12px;color:#3b3b3b;">{formatar_moeda_br(total_dia) if qtd_dia else "Sem títulos"}</div>
                                    {resumo_status}
                                </div>
                                """,
                                unsafe_allow_html=True,
                            )
    st.markdown("</div>", unsafe_allow_html=True)

    if "Atividades recentes" in dashboard_blocos:
        with st.expander("Atividades recentes", expanded=False):
            detalhe_vendas = vendas_filtradas[
                ["data_venda", "paciente_nome", "valor_total", "avaliador", "vendedor", "nf"]
            ].copy()
            detalhe_vendas = detalhe_vendas.rename(
                columns={
                    "data_venda": "Data",
                    "paciente_nome": "Paciente",
                    "valor_total": "Valor total",
                    "avaliador": "Avaliador",
                    "vendedor": "Vendedor",
                    "nf": "NF",
                }
            )
            detalhe_vendas["Data"] = detalhe_vendas["Data"].map(formatar_data_br_valor)
            detalhe_vendas["Valor total"] = detalhe_vendas["Valor total"].map(formatar_moeda_br)
            st.dataframe(detalhe_vendas.head(15), use_container_width=True, hide_index=True)

    st.markdown("**Editar venda importada**")
    opcoes_vendas = [
        (
            int(row["id"]),
            f"{formatar_data_br_valor(row['data_venda'])} - {row['paciente_nome']} - {formatar_moeda_br(float(row['valor_total'] or 0))}"
        )
        for _, row in vendas_filtradas.sort_values(["data_ref", "paciente_nome"], ascending=[False, True]).iterrows()
    ]
    venda_id = st.selectbox(
        "Venda para editar",
        options=[opcao[0] for opcao in opcoes_vendas],
        format_func=lambda valor: next(rotulo for chave, rotulo in opcoes_vendas if chave == valor),
        key="venda_edicao_id_visual",
    )
    venda_row = vendas_filtradas[vendas_filtradas["id"] == venda_id].iloc[0]

    ve1, ve2 = st.columns(2)
    edit_data_venda = ve1.text_input("Data da venda", formatar_data_br_valor(venda_row["data_venda"] or ""), key=f"edit_data_venda_visual_{venda_id}")
    edit_paciente_venda = ve2.text_input("Paciente", venda_row["paciente_nome"] or "", key=f"edit_paciente_venda_visual_{venda_id}")

    ve3, ve4, ve5, ve6 = st.columns(4)
    edit_valor_total = ve3.number_input("Valor total", min_value=0.0, value=float(venda_row["valor_total"] or 0), key=f"edit_valor_total_venda_visual_{venda_id}")
    edit_valor_a_vista = ve4.number_input("À vista", min_value=0.0, value=float(venda_row["valor_a_vista"] or 0), key=f"edit_valor_avista_venda_visual_{venda_id}")
    edit_valor_cartao = ve5.number_input("Cartão", min_value=0.0, value=float(venda_row["valor_cartao"] or 0), key=f"edit_valor_cartao_venda_visual_{venda_id}")
    edit_valor_boleto = ve6.number_input("Boleto", min_value=0.0, value=float(venda_row["valor_boleto"] or 0), key=f"edit_valor_boleto_venda_visual_{venda_id}")

    ve7, ve8, ve9, ve10 = st.columns(4)
    edit_saldo_venda = ve7.number_input("Saldo", min_value=0.0, value=float(venda_row["saldo"] or 0), key=f"edit_saldo_venda_visual_{venda_id}")
    edit_data_a_pagar = ve8.text_input("Data a pagar", formatar_data_br_valor(venda_row["data_a_pagar"] or ""), key=f"edit_data_a_pagar_venda_visual_{venda_id}")
    edit_avaliador = ve9.text_input("Avaliador", venda_row["avaliador"] or "", key=f"edit_avaliador_venda_visual_{venda_id}")
    edit_vendedor = ve10.text_input("Vendedor", venda_row["vendedor"] or "", key=f"edit_vendedor_venda_visual_{venda_id}")
    edit_nf = st.text_input("NF", venda_row["nf"] or "", key=f"edit_nf_venda_visual_{venda_id}")

    if st.button("Salvar alterações da venda", key=f"salvar_venda_visual_{venda_id}"):
        if not parse_data_contrato(edit_data_venda):
            st.error("Informe a data da venda no formato DD/MM/AAAA.")
        elif not edit_paciente_venda.strip():
            st.error("Informe o nome do paciente.")
        elif float(edit_valor_total) <= 0:
            st.error("Informe um valor total maior que zero.")
        elif edit_data_a_pagar.strip() and not parse_data_contrato(edit_data_a_pagar):
            st.error("Informe a data a pagar no formato DD/MM/AAAA.")
        else:
            cursor.execute(
                """
                UPDATE vendas
                SET data_venda=?, paciente_nome=?, valor_total=?, valor_a_vista=?, valor_cartao=?, valor_boleto=?, saldo=?, data_a_pagar=?, avaliador=?, vendedor=?, nf=?
                WHERE id=?
                """,
                (
                    formatar_data_br_valor(edit_data_venda),
                    edit_paciente_venda.strip(),
                    float(edit_valor_total),
                    float(edit_valor_a_vista),
                    float(edit_valor_cartao),
                    float(edit_valor_boleto),
                    float(edit_saldo_venda),
                    formatar_data_br_valor(edit_data_a_pagar),
                    edit_avaliador.strip(),
                    edit_vendedor.strip(),
                    edit_nf.strip(),
                    int(venda_id),
                ),
            )
            conn.commit()
            st.success("Venda atualizada com sucesso.")
            st.rerun()

    st.stop()
    st.title("Dashboard de Vendas")
    ano_dashboard = st.selectbox(
        "Ano das vendas",
        options=list(range(date.today().year - 2, date.today().year + 3)),
        index=2,
    )
    metas = carregar_metas_vendas(ano_dashboard)
    vendas = pd.read_sql("SELECT * FROM vendas ORDER BY data_venda DESC", conn)
    atualizar_status_contas_pagar_automaticamente()
    contas_pagar_dashboard = pd.read_sql("SELECT * FROM contas_pagar ORDER BY data_vencimento, fornecedor, descricao", conn)

    with st.expander("Metas de vendas", expanded=False):
        m1, m2, m3 = st.columns(3)
        meta = m1.number_input("Meta", min_value=0.0, value=float(metas["meta"] or 100000))
        supermeta = m2.number_input("Supermeta", min_value=0.0, value=float(metas["supermeta"] or 150000))
        hipermeta = m3.number_input("Hipermeta", min_value=0.0, value=float(metas["hipermeta"] or 200000))
        if st.button("Salvar metas de vendas"):
            salvar_metas_vendas(ano_dashboard, meta, supermeta, hipermeta)
            conn.commit()
            st.success("Metas atualizadas.")
            st.rerun()

    if vendas.empty:
        st.info("Nenhuma venda importada ainda.")
    else:
        vendas["data_ref"] = vendas["data_venda"].apply(parse_data_contrato)
        vendas = vendas[vendas["data_ref"].notna()].copy()
        vendas["ano"] = vendas["data_ref"].apply(lambda valor: valor.year)
        vendas = vendas[vendas["ano"] == ano_dashboard].copy()
        if vendas.empty:
            st.info("Nao ha vendas para o ano selecionado.")
        else:
            vendas["mes"] = vendas["data_ref"].apply(lambda valor: valor.month)
            vendas["mes_nome"] = vendas["mes"].map(
                {
                    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                    7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                }
            )
            fv1, fv2, fv3, fv4 = st.columns(4)
            filtro_pacientes_venda = fv1.multiselect(
                "Pacientes",
                options=sorted([valor for valor in vendas["paciente_nome"].fillna("").unique().tolist() if valor]),
            )
            filtro_avaliadores_venda = fv2.multiselect(
                "Avaliadores",
                options=sorted([valor for valor in vendas["avaliador"].fillna("").unique().tolist() if valor]),
            )
            filtro_vendedores_venda = fv3.multiselect(
                "Vendedores",
                options=sorted([valor for valor in vendas["vendedor"].fillna("").unique().tolist() if valor]),
            )
            filtro_meses_venda = fv4.multiselect(
                "Meses",
                options=list(range(1, 13)),
                format_func=lambda valor: {
                    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                    7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                }[valor],
            )

            vendas_filtradas = vendas.copy()
            if filtro_pacientes_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["paciente_nome"].isin(filtro_pacientes_venda)]
            if filtro_avaliadores_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["avaliador"].isin(filtro_avaliadores_venda)]
            if filtro_vendedores_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["vendedor"].isin(filtro_vendedores_venda)]
            if filtro_meses_venda:
                vendas_filtradas = vendas_filtradas[vendas_filtradas["mes"].isin(filtro_meses_venda)]

            if vendas_filtradas.empty:
                st.info("Nao ha vendas para os filtros selecionados.")
                st.stop()

            st.subheader("Calendário de pagamentos")
            if contas_pagar_dashboard.empty:
                st.info("Nenhum pagamento cadastrado para exibir no calendário.")
            else:
                contas_pagar_dashboard["valor"] = contas_pagar_dashboard["valor"].fillna(0).astype(float)
                contas_pagar_dashboard["_ordem_vencimento"] = pd.to_datetime(
                    contas_pagar_dashboard["data_vencimento"],
                    format="%d/%m/%Y",
                    errors="coerce",
                )
                contas_calendario = contas_pagar_dashboard[contas_pagar_dashboard["_ordem_vencimento"].notna()].copy()

                if contas_calendario.empty:
                    st.info("Nenhum vencimento válido para exibir no calendário.")
                else:
                    meses_pt_br = {
                        1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                        7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                    }
                    anos_calendario = sorted(set(contas_calendario["_ordem_vencimento"].dt.year.tolist()) | {date.today().year})
                    mes_atual = date.today().month
                    ano_atual = date.today().year

                    cal1, cal2 = st.columns(2)
                    ano_pagamentos = cal1.selectbox(
                        "Ano do calendário",
                        options=anos_calendario,
                        index=anos_calendario.index(ano_atual) if ano_atual in anos_calendario else 0,
                        key="dashboard_ano_pagamentos",
                    )
                    mes_pagamentos = cal2.selectbox(
                        "Mês do calendário",
                        options=list(range(1, 13)),
                        index=mes_atual - 1,
                        format_func=lambda valor: meses_pt_br[valor],
                        key="dashboard_mes_pagamentos",
                    )

                    contas_mes = contas_calendario[
                        (contas_calendario["_ordem_vencimento"].dt.year == ano_pagamentos)
                        & (contas_calendario["_ordem_vencimento"].dt.month == mes_pagamentos)
                    ].copy()

                    cabecalho = st.columns(7)
                    for coluna, nome_dia in zip(cabecalho, ["Seg", "Ter", "Qua", "Qui", "Sex", "Sáb", "Dom"]):
                        coluna.markdown(f"**{nome_dia}**")

                    hoje = date.today()
                    for semana in calendar.monthcalendar(ano_pagamentos, mes_pagamentos):
                        cols = st.columns(7)
                        for idx, dia in enumerate(semana):
                            with cols[idx]:
                                if dia == 0:
                                    st.write("")
                                else:
                                    dia_contas = contas_mes[contas_mes["_ordem_vencimento"].dt.day == dia].copy()
                                    total_dia = float(dia_contas["valor"].sum()) if not dia_contas.empty else 0.0
                                    qtd_dia = len(dia_contas)
                                    qtd_pagos = int((dia_contas["status"] == "Pago").sum()) if not dia_contas.empty else 0
                                    qtd_atrasados = int((dia_contas["status"] == "Atrasado").sum()) if not dia_contas.empty else 0

                                    eh_hoje = hoje.year == ano_pagamentos and hoje.month == mes_pagamentos and hoje.day == dia
                                    borda = "2px solid #111827" if eh_hoje else "1px solid #d1d5db"
                                    fundo = "#ffffff"
                                    if qtd_atrasados > 0:
                                        fundo = "#fef2f2"
                                    elif qtd_dia > 0 and qtd_pagos == qtd_dia:
                                        fundo = "#f0fdf4"
                                    elif qtd_dia > 0:
                                        fundo = "#fffbeb"

                                    resumo_status = ""
                                    if qtd_atrasados > 0:
                                        resumo_status = f"<div style='font-size:11px;color:#b91c1c;'>Atrasados: {qtd_atrasados}</div>"
                                    elif qtd_dia > 0 and qtd_pagos == qtd_dia:
                                        resumo_status = f"<div style='font-size:11px;color:#15803d;'>Todos pagos</div>"
                                    elif qtd_dia > 0:
                                        resumo_status = f"<div style='font-size:11px;color:#92400e;'>Títulos: {qtd_dia}</div>"

                                    st.markdown(
                                        f"""
                                        <div style="border:{borda};border-radius:10px;padding:6px;min-height:88px;background:{fundo};">
                                            <div style="font-weight:700;font-size:13px;margin-bottom:4px;">{dia:02d}</div>
                                            <div style="font-size:11px;color:#374151;">{formatar_moeda_br(total_dia) if qtd_dia else "Sem títulos"}</div>
                                            {resumo_status}
                                        </div>
                                        """,
                                        unsafe_allow_html=True,
                                    )

            hoje = date.today()
            vendas_mes_atual = vendas_filtradas[(vendas_filtradas["mes"] == hoje.month) & (vendas_filtradas["ano"] == hoje.year)]
            total_ano = float(vendas_filtradas["valor_total"].sum())
            total_mes = float(vendas_mes_atual["valor_total"].sum()) if not vendas_mes_atual.empty else 0.0
            ticket_medio = float(vendas_filtradas["valor_total"].mean()) if not vendas_filtradas.empty else 0.0

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Vendas do ano", formatar_moeda_br(total_ano))
            c2.metric("Vendas do mes", formatar_moeda_br(total_mes))
            c3.metric("Quantidade", str(len(vendas_filtradas)))
            c4.metric("Ticket medio", formatar_moeda_br(ticket_medio))

            c5, c6, c7 = st.columns(3)
            c5.metric("Meta", formatar_moeda_br(float(metas["meta"] or 0)))
            c6.metric("Supermeta", formatar_moeda_br(float(metas["supermeta"] or 0)))
            c7.metric("Hipermeta", formatar_moeda_br(float(metas["hipermeta"] or 0)))

            st.subheader("Progresso das metas")
            metas_progresso = [
                ("Meta", float(metas["meta"] or 0)),
                ("Supermeta", float(metas["supermeta"] or 0)),
                ("Hipermeta", float(metas["hipermeta"] or 0)),
            ]
            for nome_meta, valor_meta in metas_progresso:
                percentual = (total_mes / valor_meta) if valor_meta > 0 else 0
                percentual_exibicao = max(0.0, min(percentual, 1.0))
                diferenca = total_mes - valor_meta
                if diferenca >= 0:
                    status_meta = f"Meta batida. Excedente: {formatar_moeda_br(diferenca)}"
                else:
                    status_meta = f"Faltam {formatar_moeda_br(abs(diferenca))}"

                st.markdown(f"**{nome_meta}**: {formatar_moeda_br(valor_meta)}")
                st.progress(percentual_exibicao)
                st.caption(f"Progresso mensal: {percentual * 100:.1f}% | {status_meta}")

            resumo_mensal = vendas_filtradas.groupby(["mes", "mes_nome"], as_index=False)["valor_total"].sum().sort_values("mes")
            resumo_mensal = resumo_mensal.rename(columns={"mes_nome": "Mes", "valor_total": "Valor total"})
            ordem_meses = [
                "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
            ]
            resumo_mensal["Mes"] = pd.Categorical(resumo_mensal["Mes"], categories=ordem_meses, ordered=True)
            resumo_mensal = resumo_mensal.sort_values("Mes")
            st.subheader("Vendas por mes")
            grafico_vendas = resumo_mensal.copy()
            metas_linhas = pd.DataFrame(
                [
                    {"Meta": "Meta", "Valor": float(metas["meta"] or 0)},
                    {"Meta": "Supermeta", "Valor": float(metas["supermeta"] or 0)},
                    {"Meta": "Hipermeta", "Valor": float(metas["hipermeta"] or 0)},
                ]
            )
            barras = alt.Chart(grafico_vendas).mark_bar(color="#c5a77a").encode(
                x=alt.X("Mes:N", sort=ordem_meses, title="Mês"),
                y=alt.Y("Valor total:Q", title="Valor"),
                tooltip=[
                    alt.Tooltip("Mes:N", title="Mês"),
                    alt.Tooltip("Valor total:Q", title="Valor", format=",.2f"),
                ],
            )
            linhas = alt.Chart(metas_linhas).mark_rule(strokeWidth=2).encode(
                y=alt.Y("Valor:Q"),
                color=alt.Color(
                    "Meta:N",
                    scale=alt.Scale(
                        domain=["Meta", "Supermeta", "Hipermeta"],
                        range=["#2563eb", "#ea580c", "#15803d"],
                    ),
                    legend=alt.Legend(title="Metas"),
                ),
                tooltip=[
                    alt.Tooltip("Meta:N", title="Meta"),
                    alt.Tooltip("Valor:Q", title="Valor", format=",.2f"),
                ],
            )
            st.altair_chart((barras + linhas).properties(height=320), use_container_width=True)
            resumo_exibicao = resumo_mensal.copy()
            resumo_exibicao["Mes"] = resumo_exibicao["Mes"].astype(str)
            resumo_exibicao["Valor total"] = resumo_exibicao["Valor total"].map(formatar_moeda_br)
            st.dataframe(resumo_exibicao[["Mes", "Valor total"]], use_container_width=True, hide_index=True)

            detalhe_vendas = vendas_filtradas[
                ["data_venda", "paciente_nome", "valor_total", "valor_a_vista", "valor_cartao", "valor_boleto", "saldo", "avaliador", "vendedor", "nf"]
            ].copy()
            detalhe_vendas = detalhe_vendas.rename(
                columns={
                    "data_venda": "Data",
                    "paciente_nome": "Paciente",
                    "valor_total": "Valor total",
                    "valor_a_vista": "A vista",
                    "valor_cartao": "Cartao",
                    "valor_boleto": "Boleto",
                    "saldo": "Saldo",
                    "avaliador": "Avaliador",
                    "vendedor": "Vendedor",
                    "nf": "NF",
                }
            )
            for coluna in ["Valor total", "A vista", "Cartao", "Boleto", "Saldo"]:
                detalhe_vendas[coluna] = detalhe_vendas[coluna].map(formatar_moeda_br)
            st.subheader("Vendas importadas")
            st.dataframe(detalhe_vendas, use_container_width=True, hide_index=True)

            st.markdown("**Editar venda importada**")
            opcoes_vendas = [
                (
                    int(row["id"]),
                    f"{formatar_data_br_valor(row['data_venda'])} - {row['paciente_nome']} - {formatar_moeda_br(float(row['valor_total'] or 0))}"
                )
                for _, row in vendas_filtradas.sort_values(["data_ref", "paciente_nome"], ascending=[False, True]).iterrows()
            ]
            venda_id = st.selectbox(
                "Venda para editar",
                options=[opcao[0] for opcao in opcoes_vendas],
                format_func=lambda valor: next(rotulo for chave, rotulo in opcoes_vendas if chave == valor),
                key="venda_edicao_id",
            )
            venda_row = vendas_filtradas[vendas_filtradas["id"] == venda_id].iloc[0]

            ve1, ve2 = st.columns(2)
            edit_data_venda = ve1.text_input(
                "Data da venda",
                formatar_data_br_valor(venda_row["data_venda"] or ""),
                key=f"edit_data_venda_{venda_id}",
            )
            edit_paciente_venda = ve2.text_input(
                "Paciente",
                venda_row["paciente_nome"] or "",
                key=f"edit_paciente_venda_{venda_id}",
            )

            ve3, ve4, ve5, ve6 = st.columns(4)
            edit_valor_total = ve3.number_input(
                "Valor total",
                min_value=0.0,
                value=float(venda_row["valor_total"] or 0),
                key=f"edit_valor_total_venda_{venda_id}",
            )
            edit_valor_a_vista = ve4.number_input(
                "A vista",
                min_value=0.0,
                value=float(venda_row["valor_a_vista"] or 0),
                key=f"edit_valor_avista_venda_{venda_id}",
            )
            edit_valor_cartao = ve5.number_input(
                "Cartao",
                min_value=0.0,
                value=float(venda_row["valor_cartao"] or 0),
                key=f"edit_valor_cartao_venda_{venda_id}",
            )
            edit_valor_boleto = ve6.number_input(
                "Boleto",
                min_value=0.0,
                value=float(venda_row["valor_boleto"] or 0),
                key=f"edit_valor_boleto_venda_{venda_id}",
            )

            ve7, ve8, ve9, ve10 = st.columns(4)
            edit_saldo_venda = ve7.number_input(
                "Saldo",
                min_value=0.0,
                value=float(venda_row["saldo"] or 0),
                key=f"edit_saldo_venda_{venda_id}",
            )
            edit_data_a_pagar = ve8.text_input(
                "Data a pagar",
                formatar_data_br_valor(venda_row["data_a_pagar"] or ""),
                key=f"edit_data_a_pagar_venda_{venda_id}",
            )
            edit_avaliador = ve9.text_input(
                "Avaliador",
                venda_row["avaliador"] or "",
                key=f"edit_avaliador_venda_{venda_id}",
            )
            edit_vendedor = ve10.text_input(
                "Vendedor",
                venda_row["vendedor"] or "",
                key=f"edit_vendedor_venda_{venda_id}",
            )

            edit_nf = st.text_input(
                "NF",
                venda_row["nf"] or "",
                key=f"edit_nf_venda_{venda_id}",
            )

            if st.button("Salvar alteracoes da venda", key=f"salvar_venda_{venda_id}"):
                if not parse_data_contrato(edit_data_venda):
                    st.error("Informe a data da venda no formato DD/MM/AAAA.")
                elif not edit_paciente_venda.strip():
                    st.error("Informe o nome do paciente.")
                elif float(edit_valor_total) <= 0:
                    st.error("Informe um valor total maior que zero.")
                elif edit_data_a_pagar.strip() and not parse_data_contrato(edit_data_a_pagar):
                    st.error("Informe a data a pagar no formato DD/MM/AAAA.")
                else:
                    cursor.execute(
                        """
                        UPDATE vendas
                        SET data_venda=?, paciente_nome=?, valor_total=?, valor_a_vista=?, valor_cartao=?, valor_boleto=?, saldo=?, data_a_pagar=?, avaliador=?, vendedor=?, nf=?
                        WHERE id=?
                        """,
                        (
                            formatar_data_br_valor(edit_data_venda),
                            edit_paciente_venda.strip(),
                            float(edit_valor_total),
                            float(edit_valor_a_vista),
                            float(edit_valor_cartao),
                            float(edit_valor_boleto),
                            float(edit_saldo_venda),
                            formatar_data_br_valor(edit_data_a_pagar),
                            edit_avaliador.strip(),
                            edit_vendedor.strip(),
                            edit_nf.strip(),
                            int(venda_id),
                        ),
                    )
                    conn.commit()
                    st.success("Venda atualizada com sucesso.")
                    st.rerun()


@st.dialog("Novo paciente", width="large")
def modal_novo_paciente():
    if st.session_state.get("novo_paciente_reset", False):
        preparar_estado_paciente(
            "novo_paciente",
            dados_paciente_vazios(),
            forcar=True,
        )
        st.session_state["novo_paciente_reset"] = False

    with st.expander("Importar paciente por PDF", expanded=False):
        pdf_paciente = st.file_uploader(
            "Enviar PDF com os dados do paciente",
            type=["pdf"],
            key="pdf_paciente_upload_modal",
        )
        if pdf_paciente is not None:
            try:
                texto_pdf_paciente = extrair_texto_pdf_upload(pdf_paciente)
                dados_extraidos = extrair_dados_paciente_pdf(texto_pdf_paciente)
                if any(str(valor).strip() for valor in dados_extraidos.values()):
                    st.write("Dados reconhecidos no PDF")
                    st.json(dados_extraidos)
                    if st.button("Usar dados extraídos no cadastro", key="usar_pdf_modal_paciente"):
                        aplicar_dados_extraidos_paciente_prefixo("novo_paciente", dados_extraidos)
                        st.rerun()
                else:
                    st.warning("Não consegui reconhecer dados suficientes nesse PDF.")
            except Exception as exc:
                st.error(f"Falha ao ler o PDF do paciente: {exc}")

    dados = renderizar_campos_paciente(
        "novo_paciente",
        dados_paciente_vazios(),
    )

    if st.button("Salvar paciente", use_container_width=True, type="primary", key="salvar_modal_novo_paciente"):
        erros = validar_dados_paciente(
            dados["nome"],
            dados["prontuario"],
            dados["cpf"],
            dados["menor"],
            dados["responsavel"],
            dados["cpf_responsavel"],
        )
        if erros:
            for erro in erros:
                st.error(erro)
        else:
            with st.spinner("Salvando paciente..."):
                paciente_id = salvar_paciente_completo(dados)
            registrar_paciente_recente(paciente_id)
            registrar_feedback_visual("Paciente cadastrado com sucesso.")
            navegar_para_menu("Editar Paciente", editar_paciente_id=paciente_id)


@st.dialog("Detalhes do orçamento", width="large")
def modal_orcamento_paciente(contrato_id, paciente_nome):
    contrato = cursor.execute("SELECT * FROM contratos WHERE id=?", (int(contrato_id),)).fetchone()
    if contrato is None:
        st.warning("Orçamento não encontrado.")
        return
    procedimentos = carregar_procedimentos(int(contrato_id))
    st.markdown('<div class="ss-contract-status">APROVADO</div>', unsafe_allow_html=True)
    st.markdown(f"### {paciente_nome}")
    col1, col2, col3 = st.columns(3)
    col1.text_input("Forma de pagamento", contrato["forma_pagamento"] or "", disabled=True, key=f"orc_forma_{contrato_id}")
    col2.text_input("Parcelas", str(contrato["parcelas"] or ""), disabled=True, key=f"orc_parcelas_{contrato_id}")
    col3.text_input("Primeiro vencimento", formatar_data_br_valor(contrato["primeiro_vencimento"] or ""), disabled=True, key=f"orc_venc_{contrato_id}")
    st.text_input("Valor total", formatar_moeda_br(float(contrato["valor_total"] or 0)), disabled=True, key=f"orc_total_{contrato_id}")
    st.markdown("**Procedimentos**")
    if procedimentos.empty:
        st.info("Nenhum procedimento vinculado.")
    else:
        procedimentos_exibicao = procedimentos.rename(columns={"procedimento": "Procedimento", "valor": "Valor"}).copy()
        procedimentos_exibicao["Valor"] = procedimentos_exibicao["Valor"].fillna(0).astype(float).map(formatar_moeda_br)
        st.dataframe(procedimentos_exibicao[["Procedimento", "Valor"]], use_container_width=True, hide_index=True)
    st.button("APROVADO", disabled=True, use_container_width=True, key=f"orc_aprovado_{contrato_id}")


def renderizar_pagina_pacientes():
    st.title("Pacientes")
    pacientes = carregar_pacientes()

    topo_esq, topo_dir = st.columns([3, 1.2], gap="large")
    with topo_esq:
        st.markdown('<div class="ss-patient-hero-search">', unsafe_allow_html=True)
        centro_busca = st.columns([0.08, 0.84, 0.08])[1]
        centro_busca.text_input(
            "Buscar paciente",
            value=st.session_state.get("pacientes_busca", ""),
            placeholder="Buscar por nome, apelido, prontuário, telefone ou CPF...",
            label_visibility="collapsed",
            key="pacientes_busca",
        )
        st.markdown(
            '<div class="ss-patient-search-note">Encontre rapidamente a ficha do paciente, acesse os últimos atendidos e entre em um cadastro novo sem navegar por telas operacionais.</div>',
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)
    with topo_dir:
        if st.button("Novo paciente", use_container_width=True, type="primary", key="abrir_modal_novo_paciente"):
            st.session_state["novo_paciente_reset"] = True
            modal_novo_paciente()
        with st.expander("Importar paciente por PDF", expanded=False):
            pdf_paciente = st.file_uploader(
                "Enviar PDF com os dados do paciente",
                type=["pdf"],
                key="pdf_paciente_upload_busca",
            )
            if pdf_paciente is not None:
                try:
                    texto_pdf_paciente = extrair_texto_pdf_upload(pdf_paciente)
                    dados_extraidos = extrair_dados_paciente_pdf(texto_pdf_paciente)
                    if any(str(valor).strip() for valor in dados_extraidos.values()):
                        st.write("Dados reconhecidos no PDF")
                        st.json(dados_extraidos)
                        if st.button("Criar novo paciente com esses dados", key="criar_novo_por_pdf"):
                            preparar_estado_paciente("novo_paciente", dados_paciente_vazios(), forcar=True)
                            st.session_state["novo_paciente_reset"] = False
                            aplicar_dados_extraidos_paciente_prefixo("novo_paciente", dados_extraidos)
                            modal_novo_paciente()
                    else:
                        st.warning("Não consegui reconhecer dados suficientes nesse PDF.")
                except Exception as exc:
                    st.error(f"Falha ao ler o PDF do paciente: {exc}")

    busca = st.session_state.get("pacientes_busca", "")
    pacientes_filtrados = filtrar_pacientes_busca(pacientes, busca)
    recentes_df = pacientes_recentes(pacientes)

    if not recentes_df.empty:
        st.markdown("#### Últimos acessados")
        chips_cols = st.columns(min(4, max(1, len(recentes_df.head(4)))))
        for idx, (_, row) in enumerate(recentes_df.head(4).iterrows()):
            with chips_cols[idx % len(chips_cols)]:
                if st.button(
                    f"{row['nome']} • {formatar_prontuario_valor(row['prontuario'])}",
                    key=f"atalho_paciente_{int(row['id'])}",
                    use_container_width=True,
                ):
                    registrar_paciente_recente(int(row["id"]))
                    navegar_para_menu("Editar Paciente", editar_paciente_id=int(row["id"]))

    st.markdown(f"#### {'Resultados da busca' if busca else 'Pacientes recentes'}")
    pacientes_cards = pacientes_filtrados if busca else recentes_df
    if pacientes_cards.empty:
        st.markdown('<div class="ss-patient-empty">Nenhum paciente encontrado para esse filtro.</div>', unsafe_allow_html=True)
    else:
        cols = st.columns(3, gap="large")
        for idx, (_, row) in enumerate(pacientes_cards.head(9).iterrows()):
            with cols[idx % 3]:
                resumo = []
                if row.get("telefone"):
                    resumo.append(str(row.get("telefone")))
                if row.get("email"):
                    resumo.append(str(row.get("email")))
                resumo_texto = " • ".join(resumo) if resumo else "Sem contato principal informado"
                st.markdown(
                    f"""
                    <div class="ss-patient-grid-card">
                        <div class="ss-patient-grid-title">{row['nome']}</div>
                        <div class="ss-patient-chip-row">
                            <span class="ss-patient-chip">Prontuário {formatar_prontuario_valor(row.get('prontuario')) or '-'}</span>
                            <span class="ss-patient-chip">{formatar_data_br_valor(row.get('data_nascimento')) or 'Nascimento não informado'}</span>
                        </div>
                        <div class="ss-patient-grid-meta">{resumo_texto}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                if st.button("Abrir ficha", key=f"abrir_ficha_paciente_{int(row['id'])}", use_container_width=True):
                    registrar_paciente_recente(int(row["id"]))
                    navegar_para_menu("Editar Paciente", editar_paciente_id=int(row["id"]))

    with st.expander("Lista completa de pacientes", expanded=False):
        if pacientes.empty:
            st.info("Nenhum paciente cadastrado.")
        else:
            pacientes_exibicao = pacientes.copy()
            if "prontuario" in pacientes_exibicao.columns:
                pacientes_exibicao["prontuario"] = pacientes_exibicao["prontuario"].apply(formatar_prontuario_valor)
            pacientes_exibicao = pacientes_exibicao.drop(columns=["id"], errors="ignore")
            st.dataframe(pacientes_exibicao, use_container_width=True, hide_index=True)
            if OPENPYXL_DISPONIVEL:
                excel_pacientes = dataframe_para_excel_bytes(pacientes_exibicao, nome_aba="Pacientes")
                st.download_button(
                    "Baixar pacientes em Excel",
                    data=excel_pacientes,
                    file_name="pacientes.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="baixar_pacientes_excel_busca",
                )


def renderizar_ficha_paciente():
    st.title("Paciente")
    pacientes = carregar_pacientes()
    if pacientes.empty:
        st.warning("Nenhum paciente cadastrado.")
        st.stop()

    paciente_preselecionado = st.session_state.pop("editar_paciente_id", None)
    busca_col, select_col = st.columns([1.8, 1.2], gap="large")
    termo_busca = busca_col.text_input("Buscar paciente", value=st.session_state.get("editar_paciente_busca", ""), key="editar_paciente_busca")
    pacientes_filtrados = filtrar_pacientes_busca(pacientes, termo_busca)
    if pacientes_filtrados.empty:
        st.warning("Nenhum paciente encontrado para essa busca.")
        st.stop()

    pacientes_opcoes = opcoes_pacientes(pacientes_filtrados)
    ids_opcoes = [opcao[0] for opcao in pacientes_opcoes]
    indice_inicial = ids_opcoes.index(paciente_preselecionado) if paciente_preselecionado in ids_opcoes else 0
    paciente_id = select_col.selectbox(
        "Selecionar paciente",
        options=ids_opcoes,
        index=indice_inicial,
        format_func=lambda valor: next(rotulo for chave, rotulo in pacientes_opcoes if chave == valor),
    )
    paciente = pacientes[pacientes["id"] == paciente_id].iloc[0]
    registrar_paciente_recente(int(paciente_id))

    contratos_paciente = carregar_contratos_paciente(paciente_id)
    recebiveis_paciente = carregar_recebiveis_paciente(paciente)
    agendamentos_paciente = carregar_agendamentos_paciente(paciente)
    documentos_paciente = listar_documentos_paciente(paciente)
    exames_paciente = listar_exames_paciente(paciente)
    resumo_fin = resumo_financeiro_paciente(recebiveis_paciente)
    proximo_ag = proximo_agendamento_paciente(agendamentos_paciente)
    telefone_whatsapp = telefone_para_wa(paciente.get("telefone"))
    paciente_nome = paciente["nome"] or "Paciente"
    email_paciente = paciente.get("email", "") or "Email não informado"
    telefone_paciente = paciente.get("telefone", "") or "Telefone não informado"

    st.markdown(
        f"""
        <div class="ss-patient-hero">
            <div class="ss-patient-hero-top">
                <div class="ss-patient-hero-left">
                    <div class="ss-patient-avatar">{iniciais_paciente(paciente_nome)}</div>
                    <div>
                        <div class="ss-patient-name">{paciente_nome}</div>
                        <div class="ss-patient-subline">
                            <span class="ss-patient-status">Ativo</span>
                            <span>{telefone_paciente}</span>
                            <span>{email_paciente}</span>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    acao1, acao2, acao3 = st.columns([1, 1, 1.2])
    if telefone_whatsapp:
        acao1.link_button("WhatsApp", f"https://wa.me/{telefone_whatsapp}", use_container_width=True)
    else:
        acao1.button("WhatsApp indisponível", disabled=True, use_container_width=True, key=f"wa_indisp_{paciente_id}")
    acao2.button("Cadastro ativo", use_container_width=True, disabled=True, key=f"cadastro_ativo_{paciente_id}")
    if acao3.button("Abrir financeiro", use_container_width=True, key=f"abrir_fin_paciente_{paciente_id}"):
        navegar_para_menu("Financeiro", financeiro_foco={
            "paciente_nome": paciente_nome,
            "contrato_id": int(contratos_paciente.iloc[0]["id"]) if not contratos_paciente.empty else None,
        })

    lado_resumo, lado_conteudo = st.columns([0.86, 2.14], gap="large")
    with lado_resumo:
        st.markdown('<div class="ss-patient-sidecard"><div class="ss-patient-sidecard-title">Dados rápidos</div>', unsafe_allow_html=True)
        st.markdown(
            f"""
            <div class="ss-patient-sidecard-value">
                <strong>Prontuário:</strong> {formatar_prontuario_valor(paciente.get('prontuario')) or '-'}<br>
                <strong>CPF:</strong> {paciente.get('cpf') or '-'}<br>
                <strong>Nascimento:</strong> {formatar_data_br_valor(paciente.get('data_nascimento')) or '-'}<br>
                <strong>Sexo:</strong> {paciente.get('sexo') or '-'}<br>
                <strong>Estado civil:</strong> {paciente.get('estado_civil') or '-'}
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="ss-patient-sidecard"><div class="ss-patient-sidecard-title">Próximo agendamento</div>', unsafe_allow_html=True)
        if proximo_ag is None:
            st.markdown('<div class="ss-patient-sidecard-value">Nenhum agendamento futuro.</div>', unsafe_allow_html=True)
        else:
            st.markdown(
                f"""
                <div class="ss-patient-sidecard-value">
                    <strong>{formatar_data_br_valor(proximo_ag.get('data_agendamento') or proximo_ag.get('data'))}</strong><br>
                    {formatar_hora_agendamento(proximo_ag.get('hora_inicio'))} - {formatar_hora_agendamento(proximo_ag.get('hora_fim'))}<br>
                    {proximo_ag.get('profissional') or 'Profissional não informado'}
                </div>
                """,
                unsafe_allow_html=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)

        cor_alerta = "#027a48" if resumo_fin["quantidade_atrasados"] == 0 else "#b42318"
        fundo_alerta = "#eefbf2" if resumo_fin["quantidade_atrasados"] == 0 else "#fff2f0"
        st.markdown(
            f"""
            <div class="ss-patient-sidecard" style="background:{fundo_alerta};border-color:rgba(0,0,0,0.04);">
                <div class="ss-patient-sidecard-title">Alerta financeiro</div>
                <div class="ss-patient-sidecard-value" style="color:{cor_alerta};">
                    <strong>{'Sem atrasos' if resumo_fin['quantidade_atrasados'] == 0 else f"{resumo_fin['quantidade_atrasados']} parcela(s) em atraso"}</strong><br>
                    Em aberto: {formatar_moeda_br(resumo_fin['aberto'])}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with lado_conteudo:
        tab_principal, tab_clinico, tab_documentos, tab_comercial = st.tabs(["Principal", "Clínico", "Documentos", "Comercial"])

        with tab_principal:
            sub_cadastro, sub_financeiro, sub_agendamentos = st.tabs(["Cadastro", "Financeiro", "Agendamentos"])

            with sub_cadastro:
                prefixo = f"editpac_{int(paciente_id)}"
                dados = renderizar_campos_paciente(
                    prefixo,
                    {
                        "nome": paciente.get("nome", "") or "",
                        "apelido": paciente.get("apelido", "") or "",
                        "sexo": paciente.get("sexo", "") or "",
                        "prontuario": formatar_prontuario_valor(paciente.get("prontuario", "") or ""),
                        "cpf": paciente.get("cpf", "") or "",
                        "rg": paciente.get("rg", "") or "",
                        "data_nascimento": paciente.get("data_nascimento", "") or "",
                        "telefone": paciente.get("telefone", "") or "",
                        "email": paciente.get("email", "") or "",
                        "cep": paciente.get("cep", "") or "",
                        "endereco": paciente.get("endereco", "") or "",
                        "numero": paciente.get("numero", "") or "",
                        "bairro": paciente.get("bairro", "") or "",
                        "cidade": paciente.get("cidade", "") or "",
                        "estado": paciente.get("estado", "") or "",
                        "estado_civil": paciente.get("estado_civil", "") or "",
                        "observacoes": paciente.get("observacoes", "") or "",
                        "menor": str(paciente.get("menor_idade", "")) in {"1", "True", "true"},
                        "responsavel": paciente.get("responsavel", "") or "",
                        "cpf_responsavel": paciente.get("cpf_responsavel", "") or "",
                    },
                )
                if st.button("Salvar alterações do paciente", use_container_width=True, type="primary", key=f"salvar_paciente_{paciente_id}"):
                    erros = validar_dados_paciente(
                        dados["nome"],
                        dados["prontuario"],
                        dados["cpf"],
                        dados["menor"],
                        dados["responsavel"],
                        dados["cpf_responsavel"],
                    )
                    if erros:
                        for erro in erros:
                            st.error(erro)
                    else:
                        with st.spinner("Salvando alterações..."):
                            atualizar_paciente_completo(paciente_id, dados)
                        registrar_feedback_visual("Paciente atualizado com sucesso.")
                        st.rerun()

            with sub_financeiro:
                metric1, metric2, metric3 = st.columns(3)
                metric1.metric("Total do paciente", formatar_moeda_br(resumo_fin["total"]))
                metric2.metric("Pago", formatar_moeda_br(resumo_fin["pagos"]))
                metric3.metric("Em aberto", formatar_moeda_br(resumo_fin["aberto"]))

                if recebiveis_paciente.empty:
                    st.markdown('<div class="ss-patient-empty">Nenhum lançamento financeiro vinculado a este paciente.</div>', unsafe_allow_html=True)
                else:
                    for _, row in recebiveis_paciente.sort_values("vencimento").iterrows():
                        cor, fundo, borda = status_cor_financeiro(row.get("status"))
                        st.markdown(
                            f"""
                            <div class="ss-finance-item" style="border-left-color:{cor};background:{fundo};">
                                <div class="ss-finance-head">
                                    <div>
                                        <div class="ss-finance-value">{formatar_moeda_br(row.get('valor') or 0)}</div>
                                        <div class="ss-finance-meta">Vencimento: {formatar_data_br_valor(row.get('vencimento')) or '-'} • {row.get('forma_pagamento') or '-'}</div>
                                    </div>
                                    <span class="ss-status-pill" style="background:{fundo};color:{cor};border-color:{borda};">{row.get('status') or 'Aberto'}</span>
                                </div>
                                <div class="ss-finance-meta">
                                    Parcela: {formatar_parcela_valor(row.get('parcela_numero'))} • Pagamento: {formatar_data_br_valor(row.get('data_pagamento')) or '-'}<br>
                                    {row.get('observacao') or 'Sem observações'}
                                </div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

            with sub_agendamentos:
                if agendamentos_paciente.empty:
                    st.markdown('<div class="ss-patient-empty">Nenhum agendamento encontrado para este paciente.</div>', unsafe_allow_html=True)
                else:
                    agendamentos_exib = agendamentos_paciente.copy()
                    agendamentos_exib["_data_ordenacao"] = agendamentos_exib.apply(
                        lambda row: parse_data_contrato(row.get("data_agendamento") or row.get("data")),
                        axis=1,
                    )
                    agendamentos_exib = agendamentos_exib.sort_values(by=["_data_ordenacao", "hora_inicio"], ascending=[False, False])
                    for _, row in agendamentos_exib.iterrows():
                        cor, fundo, borda = status_cor_agendamento_paciente(row.get("status"))
                        st.markdown(
                            f"""
                            <div class="ss-finance-item" style="border-left-color:{cor};background:{fundo};">
                                <div class="ss-finance-head">
                                    <div>
                                        <div class="ss-finance-value">{formatar_data_br_valor(row.get('data_agendamento') or row.get('data'))}</div>
                                        <div class="ss-finance-meta">{formatar_hora_agendamento(row.get('hora_inicio'))} - {formatar_hora_agendamento(row.get('hora_fim'))} • {row.get('profissional') or '-'}</div>
                                    </div>
                                    <span class="ss-status-pill" style="background:{fundo};color:{cor};border-color:{borda};">{row.get('status') or 'Agendado'}</span>
                                </div>
                                <div class="ss-finance-meta">{row.get('procedimento_nome_snapshot') or row.get('procedimento') or 'Sem procedimento'}<br>{row.get('observacoes') or row.get('observacao') or 'Sem observações'}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

        with tab_clinico:
            sub_plano, sub_odontograma, sub_anamnese, sub_especialidades = st.tabs(["Plano e ficha clínica", "Odontograma", "Anamnese", "Especialidades"])

            with sub_plano:
                if contratos_paciente.empty:
                    st.markdown('<div class="ss-patient-empty">Nenhum plano ou contrato vinculado a este paciente.</div>', unsafe_allow_html=True)
                else:
                    for _, contrato in contratos_paciente.iterrows():
                        procedimentos = carregar_procedimentos(int(contrato["id"]))
                        procedimentos_lista = ", ".join(procedimentos["procedimento"].fillna("").tolist()) if not procedimentos.empty else "Sem procedimentos"
                        st.markdown(
                            f"""
                            <div class="ss-contract-card">
                                <div class="ss-contract-status">PLANO ATIVO</div>
                                <div class="ss-contract-value">{formatar_moeda_br(float(contrato.get('valor_total') or 0))}</div>
                                <div class="ss-finance-meta">Contrato #{int(contrato['id'])} • {contrato.get('forma_pagamento') or '-'} • {formatar_data_br_valor(contrato.get('data_criacao')) or '-'}</div>
                                <div class="ss-finance-meta" style="margin-top:8px;">{procedimentos_lista}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

            with sub_odontograma:
                st.markdown('<div class="ss-patient-empty">Área reservada para o odontograma. O espaço foi preservado com mais respiro visual para futura integração clínica.</div>', unsafe_allow_html=True)

            with sub_anamnese:
                st.markdown(
                    f"""
                    <div class="ss-patient-empty">
                        <strong>Resumo clínico inicial</strong><br><br>
                        {paciente.get('observacoes') or 'Sem anotações clínicas registradas neste paciente.'}
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            with sub_especialidades:
                if contratos_paciente.empty:
                    st.markdown('<div class="ss-patient-empty">Nenhuma especialidade vinculada ainda.</div>', unsafe_allow_html=True)
                else:
                    categorias = []
                    for _, contrato in contratos_paciente.iterrows():
                        procedimentos = carregar_procedimentos(int(contrato["id"]))
                        categorias.extend([str(valor) for valor in procedimentos["procedimento"].fillna("").tolist() if str(valor).strip()])
                    if categorias:
                        st.markdown("##### Procedimentos vinculados")
                        st.markdown(
                            "".join([f'<span class="ss-patient-chip">{valor}</span>' for valor in categorias[:20]]),
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown('<div class="ss-patient-empty">Nenhuma especialidade vinculada ainda.</div>', unsafe_allow_html=True)

        with tab_documentos:
            sub_documentos, sub_exames, sub_recibos = st.tabs(["Documentos", "Exames", "Recibos"])

            with sub_documentos:
                if not documentos_paciente:
                    st.markdown('<div class="ss-patient-empty">Nenhum documento localizado para este paciente.</div>', unsafe_allow_html=True)
                else:
                    for doc in documentos_paciente:
                        st.markdown(
                            f"""
                            <div class="ss-doc-card">
                                <strong>{doc['nome']}</strong><br>
                                <span class="ss-finance-meta">Atualizado em {doc['modificado_em'].strftime('%d/%m/%Y %H:%M')}</span>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )
                        with open(doc["caminho"], "rb") as arquivo_doc:
                            st.download_button(
                                f"Baixar {doc['nome']}",
                                data=arquivo_doc.read(),
                                file_name=doc["nome"],
                                key=f"doc_baixar_{paciente_id}_{doc['nome']}",
                                use_container_width=True,
                            )

            with sub_exames:
                uploads = st.file_uploader(
                    "Adicionar exame",
                    accept_multiple_files=True,
                    type=["png", "jpg", "jpeg", "pdf"],
                    key=f"upload_exame_{paciente_id}",
                )
                if st.button("Salvar exames", key=f"salvar_exames_{paciente_id}", use_container_width=True):
                    with st.spinner("Enviando exames..."):
                        quantidade = salvar_uploads_exames_paciente(paciente, uploads)
                    if quantidade:
                        registrar_feedback_visual(f"{quantidade} exame(s) adicionado(s).")
                        st.rerun()
                    else:
                        st.info("Nenhum novo exame selecionado.")
                exames_paciente = listar_exames_paciente(paciente)
                if not exames_paciente:
                    st.markdown('<div class="ss-patient-empty">Nenhum exame disponível.</div>', unsafe_allow_html=True)
                else:
                    cols_exames = st.columns(3, gap="large")
                    for idx, exame in enumerate(exames_paciente):
                        with cols_exames[idx % 3]:
                            st.markdown(
                                f"""
                                <div class="ss-doc-card">
                                    <strong>{exame['nome']}</strong><br>
                                    <span class="ss-finance-meta">{exame['modificado_em'].strftime('%d/%m/%Y %H:%M')}</span>
                                </div>
                                """,
                                unsafe_allow_html=True,
                            )
                            if exame["extensao"] in {".png", ".jpg", ".jpeg"}:
                                st.image(exame["caminho"], use_container_width=True)
                            with open(exame["caminho"], "rb") as arquivo_exame:
                                st.download_button(
                                    f"Baixar {exame['nome']}",
                                    data=arquivo_exame.read(),
                                    file_name=exame["nome"],
                                    key=f"baixar_exame_{paciente_id}_{idx}",
                                    use_container_width=True,
                                )

            with sub_recibos:
                recibos = recebiveis_paciente[recebiveis_paciente["status"] == "Pago"].copy() if not recebiveis_paciente.empty else pd.DataFrame()
                if recibos.empty:
                    st.markdown('<div class="ss-patient-empty">Nenhum recibo financeiro identificado.</div>', unsafe_allow_html=True)
                else:
                    for _, row in recibos.sort_values("data_pagamento", ascending=False).iterrows():
                        st.markdown(
                            f"""
                            <div class="ss-doc-card">
                                <strong>{formatar_moeda_br(row.get('valor') or 0)}</strong><br>
                                <span class="ss-finance-meta">Pago em {formatar_data_br_valor(row.get('data_pagamento')) or '-'} • {row.get('forma_pagamento') or '-'}</span>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

        with tab_comercial:
            st.subheader("Orçamentos")
            if contratos_paciente.empty:
                st.markdown('<div class="ss-patient-empty">Nenhum orçamento disponível para este paciente.</div>', unsafe_allow_html=True)
            else:
                for _, contrato in contratos_paciente.iterrows():
                    st.markdown(
                        f"""
                        <div class="ss-contract-card">
                            <div class="ss-contract-status">APROVADO</div>
                            <div class="ss-contract-value">{formatar_moeda_br(float(contrato.get('valor_total') or 0))}</div>
                            <div class="ss-finance-meta">Data: {formatar_data_br_valor(contrato.get('data_criacao')) or '-'} • Forma: {contrato.get('forma_pagamento') or '-'}</div>
                            <div class="ss-finance-meta" style="margin-top:8px;">Responsável: não informado</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    if st.button(f"Abrir orçamento #{int(contrato['id'])}", key=f"abrir_orcamento_{int(contrato['id'])}", use_container_width=True):
                        modal_orcamento_paciente(int(contrato["id"]), paciente_nome)

    with st.expander("Lista completa de pacientes", expanded=False):
        pacientes_exibicao = pacientes.copy()
        if "prontuario" in pacientes_exibicao.columns:
            pacientes_exibicao["prontuario"] = pacientes_exibicao["prontuario"].apply(formatar_prontuario_valor)
        pacientes_exibicao = pacientes_exibicao.drop(columns=["id"], errors="ignore")
        st.dataframe(pacientes_exibicao, use_container_width=True, hide_index=True)
        if OPENPYXL_DISPONIVEL:
            excel_pacientes = dataframe_para_excel_bytes(pacientes_exibicao, nome_aba="Pacientes")
            st.download_button(
                "Baixar pacientes em Excel",
                data=excel_pacientes,
                file_name="pacientes.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="baixar_pacientes_excel_edicao_novo",
            )


if menu == "Pacientes":
    renderizar_pagina_pacientes()
    st.stop()

if menu == "Editar Paciente":
    renderizar_ficha_paciente()
    st.stop()

if menu == "Agenda":
    st.stop()

if menu == "Pacientes":
    st.title("Cadastro de Pacientes")
    secoes_pacientes = selecionar_blocos_visiveis(
        "Exibir nesta tela",
        ["Cadastro", "Importação por PDF", "Lista de pacientes"],
        ["Cadastro"],
        "pacientes_secoes_visiveis",
    )

    if "Importação por PDF" in secoes_pacientes:
        with st.expander("Importar paciente por PDF", expanded=False):
            pdf_paciente = st.file_uploader(
                "Enviar PDF com os dados do paciente",
                type=["pdf"],
                key="pdf_paciente_upload",
            )
            if pdf_paciente is not None:
                try:
                    texto_pdf_paciente = extrair_texto_pdf_upload(pdf_paciente)
                    dados_extraidos = extrair_dados_paciente_pdf(texto_pdf_paciente)
                    if any(str(valor).strip() for valor in dados_extraidos.values()):
                        st.write("Dados reconhecidos no PDF")
                        st.json(dados_extraidos)
                        if st.button("Usar dados extraidos no formulario"):
                            aplicar_dados_extraidos_paciente(dados_extraidos)
                            st.success("Formulario preenchido com os dados extraidos do PDF.")
                    else:
                        st.warning("Nao consegui reconhecer dados suficientes nesse PDF.")
                except Exception as exc:
                    st.error(f"Falha ao ler o PDF do paciente: {exc}")

    if "Cadastro" in secoes_pacientes:
        nome = st.text_input("Nome", key="paciente_nome_input")
        prontuario = st.text_input("Prontuario", key="paciente_prontuario_input")
        cpf = st.text_input("CPF", key="paciente_cpf_input")
        nascimento = st.text_input("Data nascimento", key="paciente_nascimento_input")
        telefone = st.text_input("Telefone", key="paciente_telefone_input")
        cep = st.text_input("CEP", key="paciente_cep_input")

        dados_cep = buscar_endereco_por_cep(cep) if cep else {}
        if dados_cep:
            if not st.session_state.get("paciente_endereco_input"):
                st.session_state["paciente_endereco_input"] = dados_cep.get("logradouro", "")
            if not st.session_state.get("paciente_bairro_input"):
                st.session_state["paciente_bairro_input"] = dados_cep.get("bairro", "")
            if not st.session_state.get("paciente_cidade_input"):
                st.session_state["paciente_cidade_input"] = dados_cep.get("localidade", "")
            if not st.session_state.get("paciente_estado_input"):
                st.session_state["paciente_estado_input"] = dados_cep.get("uf", "")

        endereco = st.text_input("Rua", key="paciente_endereco_input")
        numero = st.text_input("Numero", key="paciente_numero_input")
        bairro = st.text_input("Bairro", key="paciente_bairro_input")
        cidade = st.text_input("Cidade", key="paciente_cidade_input")
        estado = st.text_input("Estado", key="paciente_estado_input")

        menor = st.checkbox("Paciente menor de idade", key="paciente_menor_input")
        responsavel = ""
        cpf_responsavel = ""

        if menor:
            responsavel = st.text_input("Responsavel", key="paciente_responsavel_input")
            cpf_responsavel = st.text_input("CPF responsavel", key="paciente_cpf_responsavel_input")

        if st.button("Salvar paciente"):
            erros = validar_dados_paciente(nome, prontuario, cpf, menor, responsavel, cpf_responsavel)
            if erros:
                for erro in erros:
                    st.error(erro)
            else:
                cursor.execute(
                    """
                    INSERT INTO pacientes
                    (
                        nome, prontuario, cpf, data_nascimento, telefone, cep,
                        endereco, numero, bairro, cidade, estado, menor_idade,
                        responsavel, cpf_responsavel
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        nome.strip(),
                        prontuario.strip(),
                        limpar_cpf(cpf),
                        nascimento.strip(),
                        telefone.strip(),
                        cep.strip(),
                        endereco.strip(),
                        numero.strip(),
                        bairro.strip(),
                        cidade.strip(),
                        estado.strip(),
                        int(menor),
                        responsavel.strip(),
                        limpar_cpf(cpf_responsavel),
                    ),
                )
                conn.commit()
                st.success("Paciente cadastrado.")

    if "Lista de pacientes" in secoes_pacientes:
        with st.expander("Pacientes cadastrados", expanded=False):
            pacientes = carregar_pacientes()
            if not pacientes.empty:
                pacientes_exibicao = pacientes.copy()
                if "prontuario" in pacientes_exibicao.columns:
                    pacientes_exibicao["prontuario"] = pacientes_exibicao["prontuario"].apply(formatar_prontuario_valor)
                pacientes_exibicao = pacientes_exibicao.drop(columns=["id"], errors="ignore")
                st.dataframe(pacientes_exibicao, use_container_width=True, hide_index=True)
                if OPENPYXL_DISPONIVEL:
                    excel_pacientes = dataframe_para_excel_bytes(pacientes_exibicao, nome_aba="Pacientes")
                    st.download_button(
                        "Baixar pacientes em Excel",
                        data=excel_pacientes,
                        file_name="pacientes.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="baixar_pacientes_excel_cadastro",
                    )

if menu == "Editar Paciente":
    st.title("Editar paciente")
    pacientes = carregar_pacientes()

    if pacientes.empty:
        st.warning("Nenhum paciente cadastrado.")
        st.stop()

    pacientes_opcoes = opcoes_pacientes(pacientes)
    paciente_preselecionado = st.session_state.pop("editar_paciente_id", None)
    if paciente_preselecionado not in [opcao[0] for opcao in pacientes_opcoes]:
        paciente_preselecionado = None
    paciente_id = st.selectbox(
        "Paciente",
        options=[opcao[0] for opcao in pacientes_opcoes],
        index=[opcao[0] for opcao in pacientes_opcoes].index(paciente_preselecionado) if paciente_preselecionado in [opcao[0] for opcao in pacientes_opcoes] else 0,
        format_func=lambda paciente_id: next(
            label for valor, label in pacientes_opcoes if valor == paciente_id
        ),
    )

    paciente = pacientes[pacientes["id"] == paciente_id].iloc[0]

    nome = st.text_input("Nome", paciente["nome"] or "")
    prontuario = st.text_input("Prontuario", paciente["prontuario"] or "")
    cpf = st.text_input("CPF", paciente["cpf"] or "")
    nascimento = st.text_input("Data nascimento", paciente["data_nascimento"] or "")
    telefone = st.text_input("Telefone", paciente["telefone"] or "")
    cep = st.text_input("CEP", paciente["cep"] or "")

    endereco = st.text_input("Rua", paciente["endereco"] or "")
    numero = st.text_input("Numero", paciente["numero"] or "")
    bairro = st.text_input("Bairro", paciente["bairro"] or "")
    cidade = st.text_input("Cidade", paciente["cidade"] or "")
    estado = st.text_input("Estado", paciente["estado"] or "")

    menor = st.checkbox(
        "Paciente menor de idade",
        value=str(paciente["menor_idade"]) in {"1", "True", "true"},
    )

    responsavel_inicial = paciente["responsavel"] or ""
    cpf_responsavel_inicial = paciente["cpf_responsavel"] or ""
    responsavel = ""
    cpf_responsavel = ""

    if menor:
        responsavel = st.text_input("Responsavel", responsavel_inicial)
        cpf_responsavel = st.text_input("CPF responsavel", cpf_responsavel_inicial)

    if st.button("Salvar alteracoes do paciente"):
        erros = validar_dados_paciente(nome, prontuario, cpf, menor, responsavel, cpf_responsavel)
        if erros:
            for erro in erros:
                st.error(erro)
        else:
            cursor.execute(
                """
                UPDATE pacientes
                SET nome=?, prontuario=?, cpf=?, data_nascimento=?, telefone=?, cep=?,
                    endereco=?, numero=?, bairro=?, cidade=?, estado=?, menor_idade=?,
                    responsavel=?, cpf_responsavel=?
                WHERE id=?
                """,
                (
                    nome.strip(),
                    prontuario.strip(),
                    limpar_cpf(cpf),
                    nascimento.strip(),
                    telefone.strip(),
                    cep.strip(),
                    endereco.strip(),
                    numero.strip(),
                    bairro.strip(),
                    cidade.strip(),
                    estado.strip(),
                    int(menor),
                    responsavel.strip() if menor else "",
                    limpar_cpf(cpf_responsavel) if menor else "",
                    paciente_id,
                ),
            )
            conn.commit()
            st.success("Paciente atualizado.")

    with st.expander("Pacientes cadastrados", expanded=False):
        pacientes_exibicao = pacientes.copy()
        if "prontuario" in pacientes_exibicao.columns:
            pacientes_exibicao["prontuario"] = pacientes_exibicao["prontuario"].apply(formatar_prontuario_valor)
        pacientes_exibicao = pacientes_exibicao.drop(columns=["id"], errors="ignore")
        st.dataframe(pacientes_exibicao, use_container_width=True, hide_index=True)
        if OPENPYXL_DISPONIVEL:
            excel_pacientes = dataframe_para_excel_bytes(pacientes_exibicao, nome_aba="Pacientes")
            st.download_button(
                "Baixar pacientes em Excel",
                data=excel_pacientes,
                file_name="pacientes.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="baixar_pacientes_excel_edicao",
            )

if menu == "Agenda":
    st.title("Agenda")
    pacientes_agenda = carregar_pacientes()
    profissionais_df = carregar_profissionais()
    profissionais_ativos_df = carregar_profissionais(somente_ativos=True)
    horarios_agenda = gerar_horarios_intervalo("08:00", "20:00", 15)
    agendamentos = pd.read_sql("SELECT * FROM agendamentos ORDER BY data, hora_inicio, id", conn)

    if not agendamentos.empty:
        agendamentos["data_ref"] = agendamentos["data"].apply(parse_data_contrato)
        agendamentos["hora_inicio_fmt"] = agendamentos["hora_inicio"].apply(formatar_hora_agendamento)
        agendamentos["hora_fim_fmt"] = agendamentos["hora_fim"].apply(formatar_hora_agendamento)
        agendamentos = agendamentos.sort_values(
            by=["data_ref", "hora_inicio"],
            key=lambda serie: serie.map(ordenar_hora_agendamento) if serie.name == "hora_inicio" else serie,
            na_position="last",
        )

    tab_agenda_dia, tab_agenda_novo, tab_agenda_editar, tab_agenda_profissionais = st.tabs(
        ["Agenda do dia", "Novo agendamento", "Editar agendamento", "Profissionais"]
    )

    with tab_agenda_dia:
        profissionais_existentes = sorted([valor for valor in profissionais_df["nome"].fillna("").unique().tolist() if valor])
        if not profissionais_existentes and not agendamentos.empty:
            profissionais_existentes = sorted([valor for valor in agendamentos["profissional"].fillna("").unique().tolist() if valor])
        ad1, ad2 = st.columns(2)
        data_agenda = ad1.date_input("Data", value=date.today(), key="agenda_data_visualizacao")
        filtro_profissional_agenda = ad2.selectbox(
            "Profissional",
            options=["Todos", *profissionais_existentes],
            key="agenda_profissional_visualizacao",
        )

        agenda_dia = agendamentos.copy() if not agendamentos.empty else pd.DataFrame()
        if not agenda_dia.empty:
            agenda_dia = agenda_dia[agenda_dia["data"] == formatar_data_br(data_agenda)].copy()
            if filtro_profissional_agenda != "Todos":
                agenda_dia = agenda_dia[agenda_dia["profissional"] == filtro_profissional_agenda]
            agenda_dia = agenda_dia.sort_values("hora_inicio", key=lambda serie: serie.map(ordenar_hora_agendamento))

        m1, m2, m3 = st.columns(3)
        m1.metric("Agendamentos", str(len(agenda_dia)))
        m2.metric("Confirmados", str(int((agenda_dia["status"] == "Confirmado").sum()) if not agenda_dia.empty else 0))
        m3.metric("Atendidos", str(int((agenda_dia["status"] == "Atendido").sum()) if not agenda_dia.empty else 0))

        st.subheader("Grade do dia")
        horarios_grade = horarios_agenda[:-1]
        for horario in horarios_grade:
            itens_hora = agenda_dia[agenda_dia["hora_inicio_fmt"] == horario].copy() if not agenda_dia.empty else pd.DataFrame()
            g1, g2 = st.columns([1, 6])
            g1.markdown(f"**{horario}**")
            with g2:
                if itens_hora.empty:
                    st.caption("Livre")
                else:
                    for _, item in itens_hora.iterrows():
                        cor_status = {
                            "Agendado": "#eff6ff",
                            "Confirmado": "#ecfdf5",
                            "Em atendimento": "#fffbeb",
                            "Atendido": "#f0fdf4",
                            "Faltou": "#fef2f2",
                            "Cancelado": "#f3f4f6",
                        }.get(str(item["status"] or ""), "#ffffff")
                        st.markdown(
                            f"""
                            <div style="border:1px solid #d9e2dd;border-radius:14px;padding:10px 12px;margin-bottom:8px;background:{cor_status};">
                                <div style="font-weight:700;">{item['paciente_nome'] or 'Paciente não informado'}</div>
                                <div style="font-size:12px;color:#52635d;">{formatar_hora_agendamento(item['hora_inicio'])} - {formatar_hora_agendamento(item['hora_fim'])} | {item['profissional'] or 'Profissional não informado'}</div>
                                <div style="font-size:12px;margin-top:4px;">{item['procedimento'] or 'Sem procedimento informado'}</div>
                                <div style="font-size:11px;color:#42534d;margin-top:4px;">Status: {item['status'] or 'Agendado'}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

        if not agenda_dia.empty:
            st.subheader("Lista do dia")
            agenda_exibicao = agenda_dia[
                ["hora_inicio_fmt", "hora_fim_fmt", "paciente_nome", "profissional", "procedimento", "status", "observacao"]
            ].rename(
                columns={
                    "hora_inicio_fmt": "Início",
                    "hora_fim_fmt": "Fim",
                    "paciente_nome": "Paciente",
                    "profissional": "Profissional",
                    "procedimento": "Procedimento",
                    "status": "Status",
                    "observacao": "Observação",
                }
            )
            st.dataframe(agenda_exibicao, use_container_width=True, hide_index=True)

    with tab_agenda_novo:
        if pacientes_agenda.empty:
            st.info("Cadastre pacientes antes de criar agendamentos.")
        elif profissionais_ativos_df.empty:
            st.info("Cadastre pelo menos um profissional ativo antes de criar agendamentos.")
        else:
            pacientes_opcoes_agenda = opcoes_pacientes(pacientes_agenda)
            profissionais_opcoes = [valor for valor in profissionais_ativos_df["nome"].fillna("").tolist() if valor]
            n1, n2, n3 = st.columns(3)
            data_novo_agendamento = n1.date_input("Data do agendamento", value=date.today(), key="novo_agendamento_data")
            hora_inicio_novo = n2.selectbox(
                "Hora inicial",
                options=horarios_agenda[:-1],
                index=horarios_agenda[:-1].index("09:00") if "09:00" in horarios_agenda[:-1] else 0,
                key="novo_agendamento_hora_inicio",
            )
            hora_fim_novo = n3.selectbox(
                "Hora final",
                options=horarios_agenda[1:],
                index=horarios_agenda[1:].index("09:15") if "09:15" in horarios_agenda[1:] else 0,
                key="novo_agendamento_hora_fim",
            )

            na1, na2 = st.columns(2)
            paciente_agendamento_id = na1.selectbox(
                "Paciente",
                options=[opcao[0] for opcao in pacientes_opcoes_agenda],
                format_func=lambda paciente_id: next(label for valor, label in pacientes_opcoes_agenda if valor == paciente_id),
                key="novo_agendamento_paciente",
            )
            profissional_novo = na2.selectbox("Profissional", options=profissionais_opcoes, key="novo_agendamento_profissional")
            procedimento_novo = st.text_input("Procedimento", key="novo_agendamento_procedimento")
            status_novo_agendamento = st.selectbox("Status", STATUS_AGENDAMENTO, key="novo_agendamento_status")
            observacao_novo_agendamento = st.text_area("Observação", key="novo_agendamento_observacao")

            if st.button("Salvar agendamento"):
                hora_inicio_obj = parse_hora_agendamento(hora_inicio_novo)
                hora_fim_obj = parse_hora_agendamento(hora_fim_novo)
                if hora_inicio_obj is None or hora_fim_obj is None:
                    st.error("Informe os horários no formato HH:MM.")
                elif hora_fim_obj <= hora_inicio_obj:
                    st.error("A hora final deve ser maior que a hora inicial.")
                elif not profissional_novo.strip():
                    st.error("Selecione um profissional.")
                else:
                    paciente_row = pacientes_agenda[pacientes_agenda["id"] == paciente_agendamento_id].iloc[0]
                    cursor.execute(
                        """
                        INSERT INTO agendamentos
                        (data, hora_inicio, hora_fim, paciente_id, paciente_nome, profissional, procedimento, status, observacao, data_criacao)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            formatar_data_br(data_novo_agendamento),
                            formatar_hora_agendamento(hora_inicio_novo),
                            formatar_hora_agendamento(hora_fim_novo),
                            int(paciente_agendamento_id),
                            paciente_row["nome"],
                            profissional_novo.strip(),
                            procedimento_novo.strip(),
                            status_novo_agendamento,
                            observacao_novo_agendamento.strip(),
                            agora_str(),
                        ),
                    )
                    conn.commit()
                    st.success("Agendamento salvo.")
                    st.rerun()

    with tab_agenda_editar:
        if agendamentos.empty:
            st.info("Nenhum agendamento cadastrado ainda.")
        else:
            profissionais_opcoes = sorted([valor for valor in profissionais_df["nome"].fillna("").tolist() if valor])
            if not profissionais_opcoes:
                profissionais_opcoes = sorted([valor for valor in agendamentos["profissional"].fillna("").tolist() if valor])
            opcoes_agendamento = [
                (
                    int(row["id"]),
                    f"{formatar_data_br_valor(row['data'])} - {formatar_hora_agendamento(row['hora_inicio'])} - {row['paciente_nome']} - {row['profissional']}"
                )
                for _, row in agendamentos.sort_values(["data_ref", "hora_inicio"], ascending=[False, True]).iterrows()
            ]
            agendamento_id = st.selectbox(
                "Agendamento",
                options=[opcao[0] for opcao in opcoes_agendamento],
                format_func=lambda valor: next(rotulo for chave, rotulo in opcoes_agendamento if chave == valor),
                key="agenda_edicao_id",
            )
            agendamento_row = agendamentos[agendamentos["id"] == agendamento_id].iloc[0]

            e1, e2, e3 = st.columns(3)
            edit_data_agendamento = e1.text_input("Data", formatar_data_br_valor(agendamento_row["data"] or ""), key=f"agenda_edit_data_{agendamento_id}")
            edit_hora_inicio = e2.selectbox(
                "Hora inicial",
                options=horarios_agenda[:-1],
                index=horarios_agenda[:-1].index(formatar_hora_agendamento(agendamento_row["hora_inicio"])) if formatar_hora_agendamento(agendamento_row["hora_inicio"]) in horarios_agenda[:-1] else 0,
                key=f"agenda_edit_hora_ini_{agendamento_id}",
            )
            edit_hora_fim = e3.selectbox(
                "Hora final",
                options=horarios_agenda[1:],
                index=horarios_agenda[1:].index(formatar_hora_agendamento(agendamento_row["hora_fim"])) if formatar_hora_agendamento(agendamento_row["hora_fim"]) in horarios_agenda[1:] else 0,
                key=f"agenda_edit_hora_fim_{agendamento_id}",
            )

            e4, e5 = st.columns(2)
            edit_paciente_nome = e4.text_input("Paciente", agendamento_row["paciente_nome"] or "", key=f"agenda_edit_paciente_{agendamento_id}")
            edit_profissional = e5.selectbox(
                "Profissional",
                options=profissionais_opcoes if profissionais_opcoes else [agendamento_row["profissional"] or ""],
                index=(profissionais_opcoes.index(agendamento_row["profissional"]) if agendamento_row["profissional"] in profissionais_opcoes else 0),
                key=f"agenda_edit_profissional_{agendamento_id}",
            )
            edit_procedimento = st.text_input("Procedimento", agendamento_row["procedimento"] or "", key=f"agenda_edit_procedimento_{agendamento_id}")
            edit_status = st.selectbox(
                "Status",
                STATUS_AGENDAMENTO,
                index=STATUS_AGENDAMENTO.index(agendamento_row["status"]) if agendamento_row["status"] in STATUS_AGENDAMENTO else 0,
                key=f"agenda_edit_status_{agendamento_id}",
            )
            edit_observacao = st.text_area("Observação", agendamento_row["observacao"] or "", key=f"agenda_edit_obs_{agendamento_id}")

            col_salvar_agenda, col_excluir_agenda = st.columns(2)
            if col_salvar_agenda.button("Salvar alterações", key=f"agenda_salvar_{agendamento_id}"):
                hora_inicio_obj = parse_hora_agendamento(edit_hora_inicio)
                hora_fim_obj = parse_hora_agendamento(edit_hora_fim)
                if not parse_data_contrato(edit_data_agendamento):
                    st.error("Informe a data no formato DD/MM/AAAA.")
                elif hora_inicio_obj is None or hora_fim_obj is None:
                    st.error("Informe os horários no formato HH:MM.")
                elif hora_fim_obj <= hora_inicio_obj:
                    st.error("A hora final deve ser maior que a hora inicial.")
                elif not edit_paciente_nome.strip():
                    st.error("Informe o paciente.")
                elif not edit_profissional.strip():
                    st.error("Selecione um profissional.")
                else:
                    cursor.execute(
                        """
                        UPDATE agendamentos
                        SET data=?, hora_inicio=?, hora_fim=?, paciente_nome=?, profissional=?, procedimento=?, status=?, observacao=?
                        WHERE id=?
                        """,
                        (
                            formatar_data_br_valor(edit_data_agendamento),
                            formatar_hora_agendamento(edit_hora_inicio),
                            formatar_hora_agendamento(edit_hora_fim),
                            edit_paciente_nome.strip(),
                            edit_profissional.strip(),
                            edit_procedimento.strip(),
                            edit_status,
                            edit_observacao.strip(),
                            int(agendamento_id),
                        ),
                    )
                    conn.commit()
                    st.success("Agendamento atualizado.")
                    st.rerun()

            if col_excluir_agenda.button("Excluir agendamento", key=f"agenda_excluir_{agendamento_id}"):
                cursor.execute("DELETE FROM agendamentos WHERE id=?", (int(agendamento_id),))
                conn.commit()
                st.success("Agendamento excluído.")
                st.rerun()

    with tab_agenda_profissionais:
        st.subheader("Cadastro de profissionais")
        p1, p2 = st.columns(2)
        profissional_nome = p1.text_input("Nome do profissional", key="agenda_profissional_nome")
        profissional_especialidade = p2.text_input("Especialidade", key="agenda_profissional_especialidade")
        profissional_observacao = st.text_area("Observação", key="agenda_profissional_observacao")
        if st.button("Salvar profissional", key="agenda_salvar_profissional"):
            if not profissional_nome.strip():
                st.error("Informe o nome do profissional.")
            else:
                existente = cursor.execute(
                    "SELECT id FROM profissionais WHERE lower(trim(nome)) = lower(trim(?)) LIMIT 1",
                    (profissional_nome.strip(),),
                ).fetchone()
                if existente is None:
                    cursor.execute(
                        """
                        INSERT INTO profissionais (nome, especialidade, ativo, observacao, data_criacao)
                        VALUES (?, ?, 1, ?, ?)
                        """,
                        (
                            profissional_nome.strip(),
                            profissional_especialidade.strip(),
                            profissional_observacao.strip(),
                            agora_str(),
                        ),
                    )
                else:
                    cursor.execute(
                        """
                        UPDATE profissionais
                        SET especialidade=?, observacao=?, ativo=1
                        WHERE id=?
                        """,
                        (
                            profissional_especialidade.strip(),
                            profissional_observacao.strip(),
                            int(existente["id"]),
                        ),
                    )
                conn.commit()
                st.success("Profissional salvo com sucesso.")
                st.rerun()

        st.markdown("**Profissionais cadastrados**")
        if profissionais_df.empty:
            st.info("Nenhum profissional cadastrado ainda.")
        else:
            profissionais_exibicao = profissionais_df.copy()
            profissionais_exibicao["ativo"] = profissionais_exibicao["ativo"].map(lambda valor: "Ativo" if int(valor) == 1 else "Inativo")
            profissionais_exibicao = profissionais_exibicao.rename(
                columns={
                    "nome": "Nome",
                    "especialidade": "Especialidade",
                    "ativo": "Status",
                    "observacao": "Observação",
                }
            )
            st.dataframe(
                profissionais_exibicao[["Nome", "Especialidade", "Status", "Observação"]],
                use_container_width=True,
                hide_index=True,
            )

            opcoes_profissionais = [
                (int(row["id"]), f"{row['nome']} - {row['especialidade'] or 'Sem especialidade'}")
                for _, row in profissionais_df.iterrows()
            ]
            profissional_id_editar = st.selectbox(
                "Profissional para editar",
                options=[opcao[0] for opcao in opcoes_profissionais],
                format_func=lambda valor: next(rotulo for chave, rotulo in opcoes_profissionais if chave == valor),
                key="agenda_profissional_id_editar",
            )
            profissional_row = profissionais_df[profissionais_df["id"] == profissional_id_editar].iloc[0]
            pe1, pe2 = st.columns(2)
            profissional_nome_edit = pe1.text_input("Nome", value=profissional_row["nome"] or "", key=f"agenda_profissional_nome_edit_{profissional_id_editar}")
            profissional_especialidade_edit = pe2.text_input("Especialidade", value=profissional_row["especialidade"] or "", key=f"agenda_profissional_esp_edit_{profissional_id_editar}")
            profissional_ativo_edit = st.checkbox("Profissional ativo", value=int(profissional_row["ativo"] or 0) == 1, key=f"agenda_profissional_ativo_edit_{profissional_id_editar}")
            profissional_observacao_edit = st.text_area("Observação", value=profissional_row["observacao"] or "", key=f"agenda_profissional_obs_edit_{profissional_id_editar}")
            ped1, ped2 = st.columns(2)
            if ped1.button("Salvar alterações do profissional", key=f"agenda_profissional_salvar_edit_{profissional_id_editar}"):
                if not profissional_nome_edit.strip():
                    st.error("Informe o nome do profissional.")
                else:
                    cursor.execute(
                        """
                        UPDATE profissionais
                        SET nome=?, especialidade=?, ativo=?, observacao=?
                        WHERE id=?
                        """,
                        (
                            profissional_nome_edit.strip(),
                            profissional_especialidade_edit.strip(),
                            int(bool(profissional_ativo_edit)),
                            profissional_observacao_edit.strip(),
                            int(profissional_id_editar),
                        ),
                    )
                    conn.commit()
                    st.success("Profissional atualizado.")
                    st.rerun()
            if ped2.button("Excluir profissional", key=f"agenda_profissional_excluir_{profissional_id_editar}"):
                cursor.execute("DELETE FROM profissionais WHERE id=?", (int(profissional_id_editar),))
                conn.commit()
                st.success("Profissional excluído.")
                st.rerun()

if menu == "Contratos":
    st.title("Contratos")
    tab_contrato_novo, tab_contrato_lista = st.tabs(["Novo contrato", "Lista de contratos"])

    with tab_contrato_novo:
        pacientes = carregar_pacientes()

        if pacientes.empty:
            st.warning("Cadastre paciente primeiro.")
            st.stop()

        pacientes_opcoes = opcoes_pacientes(pacientes)
        paciente_id = st.selectbox(
            "Paciente",
            options=[opcao[0] for opcao in pacientes_opcoes],
            format_func=lambda paciente_id: next(
                label for valor, label in pacientes_opcoes if valor == paciente_id
            ),
        )
        paciente_nome = pacientes.loc[pacientes["id"] == paciente_id, "nome"].iloc[0]
        prontuario_paciente = pacientes.loc[pacientes["id"] == paciente_id, "prontuario"].iloc[0]

        forma = st.selectbox("Forma pagamento", FORMAS_PAGAMENTO)
        entrada = st.number_input("Entrada", min_value=0.0, value=0.0)
        data_pagamento_entrada = st.text_input(
            "Data do pagamento" if forma_pagamento_a_vista(forma) else "Data do pagamento da entrada",
            value=formatar_data_br(date.today()),
        )
        if forma_pagamento_a_vista(forma):
            parcelas = 1
            vencimento = ""
            st.caption("Para Pix, Dinheiro, Debito e Credito o sistema considera pagamento a vista.")
        else:
            parcelas = st.number_input("Parcelas", min_value=1, max_value=36, value=1)
            vencimento = st.text_input("Primeiro vencimento")

        procedimentos = []
        valores = []
        for i in range(10):
            col1, col2 = st.columns(2)
            proc = col1.text_input(f"Procedimento {i + 1}")
            valor = col2.number_input(f"Valor {i + 1}", min_value=0.0, value=0.0, key=f"v{i}")
            if proc.strip():
                procedimentos.append(proc.strip())
                valores.append(float(valor))

        total = sum(valores)
        st.write("Valor total:", total)

        if st.button("Salvar contrato"):
            erros, total = validar_contrato(procedimentos, valores, entrada, parcelas, vencimento, forma, data_pagamento_entrada)
            if erros:
                for erro in erros:
                    st.error(erro)
            else:
                cursor.execute(
                    """
                    INSERT INTO contratos
                    (
                        paciente_id, valor_total, entrada, parcelas,
                        primeiro_vencimento, data_pagamento_entrada, forma_pagamento, data_criacao
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        paciente_id,
                        total,
                        entrada,
                        parcelas,
                        vencimento.strip(),
                        formatar_data_br_valor(data_pagamento_entrada),
                        forma,
                        agora_str(),
                    ),
                )
                contrato_id = cursor.lastrowid
                salvar_procedimentos_contrato(contrato_id, procedimentos, valores)
                sincronizar_financeiro_contrato(contrato_id, paciente_nome, total)
                if not forma_pagamento_a_vista(forma):
                    sincronizar_recebiveis_contrato(
                        contrato_id,
                        paciente_id,
                        paciente_nome,
                        prontuario_paciente,
                        total,
                        entrada,
                        parcelas,
                        vencimento.strip(),
                        forma,
                    )
                conn.commit()

                try:
                    arquivo = gerar_documento(conn, contrato_id)
                    st.success(f"Contrato salvo em: {os.path.abspath(arquivo)}")
                except Exception as exc:
                    st.error(f"Contrato salvo, mas houve falha ao gerar o documento: {exc}")

    with tab_contrato_lista:
        contratos_lista = carregar_contratos()
        pacientes_lista = carregar_pacientes()
        if contratos_lista.empty:
            st.info("Nenhum contrato cadastrado ainda.")
        else:
            secoes_contratos = selecionar_blocos_visiveis(
                "Exibir na lista",
                ["Filtros", "Exportação", "Resultado"],
                ["Filtros", "Resultado"],
                "contratos_lista_secoes_visiveis",
            )
            pacientes_base = pacientes_lista[["id", "nome", "prontuario"]].copy() if not pacientes_lista.empty else pd.DataFrame(columns=["id", "nome", "prontuario"])
            pacientes_base = pacientes_base.rename(columns={"id": "paciente_id", "nome": "paciente_nome", "prontuario": "prontuario"})
            contratos_exibicao = contratos_lista.merge(pacientes_base, on="paciente_id", how="left")
            procedimentos_df = pd.read_sql(
                "SELECT contrato_id, procedimento FROM procedimentos_contrato ORDER BY id",
                conn,
            )
            if not procedimentos_df.empty:
                resumo_procedimentos = (
                    procedimentos_df.groupby("contrato_id")["procedimento"]
                    .apply(lambda serie: ", ".join([str(valor).strip() for valor in serie.tolist() if str(valor).strip()]))
                    .reset_index()
                    .rename(columns={"procedimento": "procedimentos"})
                )
                contratos_exibicao = contratos_exibicao.merge(resumo_procedimentos, left_on="id", right_on="contrato_id", how="left")
            else:
                contratos_exibicao["procedimentos"] = ""

            contratos_exibicao["paciente_nome"] = contratos_exibicao["paciente_nome"].fillna("")
            contratos_exibicao["prontuario"] = contratos_exibicao["prontuario"].apply(formatar_prontuario_valor)
            contratos_exibicao["forma_pagamento"] = contratos_exibicao["forma_pagamento"].fillna("")
            contratos_exibicao["data_criacao_bruta"] = pd.to_datetime(contratos_exibicao["data_criacao"], errors="coerce")

            filtro_pacientes_contrato = []
            filtro_prontuarios_contrato = []
            filtro_formas_contrato = []
            filtro_periodo_contrato = ()
            if "Filtros" in secoes_contratos:
                with st.expander("Filtros da lista", expanded=True):
                    fc1, fc2, fc3, fc4 = st.columns(4)
                    filtro_pacientes_contrato = fc1.multiselect(
                        "Pacientes",
                        options=sorted([valor for valor in contratos_exibicao["paciente_nome"].unique().tolist() if valor]),
                    )
                    filtro_prontuarios_contrato = fc2.multiselect(
                        "Prontuarios",
                        options=sorted([valor for valor in contratos_exibicao["prontuario"].unique().tolist() if valor]),
                    )
                    filtro_formas_contrato = fc3.multiselect(
                        "Formas de pagamento",
                        options=sorted([valor for valor in contratos_exibicao["forma_pagamento"].unique().tolist() if valor]),
                    )
                    datas_contrato = [
                        valor.date()
                        for valor in contratos_exibicao["data_criacao_bruta"].dropna().sort_values().unique().tolist()
                    ]
                    data_inicio_contrato = datas_contrato[0] if datas_contrato else None
                    data_fim_contrato = datas_contrato[-1] if datas_contrato else None
                    filtro_periodo_contrato = fc4.date_input(
                        "Periodo de criacao",
                        value=(data_inicio_contrato, data_fim_contrato) if data_inicio_contrato and data_fim_contrato else (),
                    )

            contratos_filtrados = contratos_exibicao.copy()
            if filtro_pacientes_contrato:
                contratos_filtrados = contratos_filtrados[contratos_filtrados["paciente_nome"].isin(filtro_pacientes_contrato)]
            if filtro_prontuarios_contrato:
                contratos_filtrados = contratos_filtrados[contratos_filtrados["prontuario"].isin(filtro_prontuarios_contrato)]
            if filtro_formas_contrato:
                contratos_filtrados = contratos_filtrados[contratos_filtrados["forma_pagamento"].isin(filtro_formas_contrato)]
            if isinstance(filtro_periodo_contrato, tuple) and len(filtro_periodo_contrato) == 2 and filtro_periodo_contrato[0] and filtro_periodo_contrato[1]:
                inicio_contrato = pd.to_datetime(filtro_periodo_contrato[0])
                fim_contrato = pd.to_datetime(filtro_periodo_contrato[1])
                contratos_filtrados = contratos_filtrados[
                    (contratos_filtrados["data_criacao_bruta"] >= inicio_contrato) &
                    (contratos_filtrados["data_criacao_bruta"] <= fim_contrato)
                ]

            rc1, rc2, rc3 = st.columns(3)
            rc1.metric("Contratos filtrados", str(len(contratos_filtrados)))
            rc2.metric("Valor total", formatar_moeda_br(float(contratos_filtrados["valor_total"].fillna(0).sum())))
            rc3.metric("Entradas", formatar_moeda_br(float(contratos_filtrados["entrada"].fillna(0).sum())))

            lista_exportacao = contratos_filtrados[
                ["id", "paciente_nome", "prontuario", "valor_total", "entrada", "parcelas", "primeiro_vencimento", "data_pagamento_entrada", "forma_pagamento", "data_criacao", "procedimentos"]
            ].copy()
            lista_exportacao = lista_exportacao.rename(
                columns={
                    "id": "Contrato",
                    "paciente_nome": "Paciente",
                    "prontuario": "Prontuario",
                    "valor_total": "Valor total",
                    "entrada": "Entrada",
                    "parcelas": "Parcelas",
                    "primeiro_vencimento": "Primeiro vencimento",
                    "data_pagamento_entrada": "Data pagamento",
                    "forma_pagamento": "Forma pagamento",
                    "data_criacao": "Criado em",
                    "procedimentos": "Procedimentos",
                }
            )
            lista_exportacao["Valor total"] = lista_exportacao["Valor total"].map(formatar_moeda_br)
            lista_exportacao["Entrada"] = lista_exportacao["Entrada"].map(formatar_moeda_br)
            lista_exportacao["Primeiro vencimento"] = lista_exportacao["Primeiro vencimento"].apply(formatar_data_br_valor)
            lista_exportacao["Data pagamento"] = lista_exportacao["Data pagamento"].apply(formatar_data_br_valor)
            lista_exportacao["Criado em"] = lista_exportacao["Criado em"].apply(formatar_data_hora_br_valor)

            if "Exportação" in secoes_contratos:
                with st.expander("Exportação", expanded=False):
                    if OPENPYXL_DISPONIVEL:
                        excel_contratos = dataframe_para_excel_bytes(lista_exportacao, nome_aba="Contratos")
                        st.download_button(
                            "Baixar contratos em Excel",
                            data=excel_contratos,
                            file_name="contratos_filtrados.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )
                    else:
                        st.info("Excel indisponivel: instale openpyxl para habilitar esta exportacao.")

            if "Resultado" in secoes_contratos:
                st.dataframe(lista_exportacao, use_container_width=True, hide_index=True)

if menu == "Editar Contrato":
    st.title("Editar contrato")
    contratos = carregar_contratos()

    if contratos.empty:
        st.warning("Nenhum contrato encontrado.")
        st.stop()

    pacientes = carregar_pacientes()
    contratos_opcoes = opcoes_contratos(contratos, pacientes)
    busca_contrato = st.text_input(
        "Buscar contrato",
        help="Digite nome do paciente, prontuario, forma de pagamento ou numero do contrato.",
    )
    contratos_filtrados = filtrar_opcoes_contratos(contratos_opcoes, busca_contrato)

    if not contratos_filtrados:
        st.warning("Nenhum contrato encontrado para a busca informada.")
        st.stop()

    contrato_id = st.selectbox(
        "Contrato",
        options=[opcao[0] for opcao in contratos_filtrados],
        format_func=lambda contrato_id: next(
            label for valor, label in contratos_filtrados if valor == contrato_id
        ),
        help="Pesquise pelo nome do paciente, prontuario ou identificacao do contrato.",
    )
    contrato = contratos[contratos["id"] == contrato_id].iloc[0]

    if pacientes.empty:
        st.warning("Nenhum paciente cadastrado.")
        st.stop()

    pacientes_opcoes = opcoes_pacientes(pacientes)
    paciente_id_atual = int(contrato["paciente_id"])
    if paciente_id_atual not in [opcao[0] for opcao in pacientes_opcoes]:
        st.error("O paciente vinculado a este contrato nao foi encontrado.")
        st.stop()

    paciente_id = st.selectbox(
        "Paciente",
        options=[opcao[0] for opcao in pacientes_opcoes],
        index=[opcao[0] for opcao in pacientes_opcoes].index(paciente_id_atual),
        format_func=lambda paciente_id: next(
            label for valor, label in pacientes_opcoes if valor == paciente_id
        ),
    )
    paciente_nome = pacientes.loc[pacientes["id"] == paciente_id, "nome"].iloc[0]
    prontuario_paciente = pacientes.loc[pacientes["id"] == paciente_id, "prontuario"].iloc[0]

    forma_atual = normalizar_forma_pagamento(contrato["forma_pagamento"])
    forma = st.selectbox(
        "Forma pagamento",
        FORMAS_PAGAMENTO,
        index=FORMAS_PAGAMENTO.index(forma_atual),
    )
    entrada = st.number_input("Entrada", min_value=0.0, value=float(contrato["entrada"] or 0))
    data_pagamento_entrada = st.text_input(
        "Data do pagamento" if forma_pagamento_a_vista(forma) else "Data do pagamento da entrada",
        formatar_data_br_valor(contrato["data_pagamento_entrada"] or contrato["data_criacao"] or ""),
    )
    if forma_pagamento_a_vista(forma):
        parcelas = 1
        vencimento = ""
        st.caption("Para Pix, Dinheiro, Debito e Credito o sistema considera pagamento a vista.")
    else:
        parcelas = st.number_input("Parcelas", min_value=1, max_value=36, value=int(contrato["parcelas"] or 1))
        vencimento = st.text_input("Primeiro vencimento", contrato["primeiro_vencimento"] or "")

    procedimentos_db = carregar_procedimentos(contrato_id)
    procedimentos = []
    valores = []

    for i in range(10):
        proc_nome = ""
        proc_valor = 0.0
        if i < len(procedimentos_db):
            proc_nome = procedimentos_db.iloc[i]["procedimento"]
            proc_valor = float(procedimentos_db.iloc[i]["valor"] or 0)

        col1, col2 = st.columns(2)
        proc = col1.text_input(
            f"Procedimento {i + 1}",
            proc_nome,
            key=f"ep_{contrato_id}_{i}",
        )
        valor = col2.number_input(
            f"Valor {i + 1}",
            min_value=0.0,
            value=proc_valor,
            key=f"ev_{contrato_id}_{i}",
        )
        if proc.strip():
            procedimentos.append(proc.strip())
            valores.append(float(valor))

    total = sum(valores)
    st.write("Valor total:", total)

    if st.button("Salvar alteracoes"):
        erros, total = validar_contrato(procedimentos, valores, entrada, parcelas, vencimento, forma, data_pagamento_entrada)
        if erros:
            for erro in erros:
                st.error(erro)
        else:
            cursor.execute(
                """
                UPDATE contratos
                SET paciente_id=?, valor_total=?, entrada=?, parcelas=?,
                    primeiro_vencimento=?, data_pagamento_entrada=?, forma_pagamento=?
                WHERE id=?
                """,
                (
                    paciente_id,
                    total,
                    entrada,
                    parcelas,
                    vencimento.strip(),
                    formatar_data_br_valor(data_pagamento_entrada),
                    forma,
                    contrato_id,
                ),
            )
            salvar_procedimentos_contrato(contrato_id, procedimentos, valores)
            sincronizar_financeiro_contrato(contrato_id, paciente_nome, total)
            if forma_pagamento_a_vista(forma):
                cursor.execute("DELETE FROM recebiveis WHERE contrato_id=?", (contrato_id,))
                status_recebiveis = "removido"
            else:
                status_recebiveis = sincronizar_recebiveis_contrato(
                    contrato_id,
                    paciente_id,
                    paciente_nome,
                    prontuario_paciente,
                    total,
                    entrada,
                    parcelas,
                    vencimento.strip(),
                    forma,
                )
            conn.commit()
            st.success("Contrato atualizado.")
            if status_recebiveis == "atualizado":
                st.warning("Os recebiveis do contrato foram atualizados porque houve mudanca no cronograma.")

            try:
                arquivo = gerar_documento(conn, contrato_id)
                st.success(f"Contrato atualizado e salvo em: {os.path.abspath(arquivo)}")
            except Exception as exc:
                st.error(f"Contrato atualizado, mas houve falha ao gerar o documento: {exc}")

if menu == "Importacoes":
    st.title("Importacoes")
    tab_imp_contratos, tab_imp_recebiveis, tab_imp_pagar, tab_imp_vendas = st.tabs(
        ["Contratos", "Recebiveis", "A pagar", "Vendas"]
    )

    with tab_imp_contratos:
        st.subheader("Importar contratos por Excel")
        st.caption("Use a planilha de contratos para trazer pacientes, contratos, procedimentos e recebiveis para o sistema.")

        arquivo_importacao = st.file_uploader(
            "Enviar planilha de contratos",
            type=["xlsx", "xls"],
            key="arquivo_importacao_contratos",
        )

        if arquivo_importacao is not None:
            try:
                contratos_preparados = preparar_planilha_importacao_contratos(arquivo_importacao)
                if not contratos_preparados:
                    st.warning("Nao encontrei contratos validos para importar nessa planilha.")
                else:
                    preview_df = pd.DataFrame(
                        [
                            {
                                "Paciente": item["nome"],
                                "Prontuario": item["prontuario"],
                                "Data": item["data"],
                                "Forma": item["forma_pagamento"],
                                "Valor total": formatar_moeda_br(item["valor_total"]),
                                "Entrada": formatar_moeda_br(item["entrada"]),
                                "Parcelas": item["parcelas"],
                                "Procedimentos": len(item["procedimentos"]),
                            }
                            for item in contratos_preparados
                        ]
                    )
                    st.success(f"{len(contratos_preparados)} contratos encontrados para importacao.")
                    st.dataframe(preview_df, use_container_width=True, hide_index=True)

                    if st.button("Importar contratos da planilha"):
                        resultado_importacao = importar_contratos_preparados(contratos_preparados)
                        st.success(
                            f"Importacao concluida. Importados: {resultado_importacao['importados']} | Atualizados: {resultado_importacao['atualizados']} | Ignorados: {resultado_importacao['ignorados']}"
                        )
                        if resultado_importacao["recebiveis_pendentes"]:
                            st.warning("Alguns contratos foram importados sem gerar recebiveis porque nao havia primeiro vencimento informado:")
                            for pendencia in resultado_importacao["recebiveis_pendentes"]:
                                st.write(f"- {pendencia}")
                        if resultado_importacao["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado_importacao["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha: {exc}")

    with tab_imp_recebiveis:
        st.subheader("Importar recebiveis por Excel")
        st.caption("Atualiza recebiveis existentes por paciente + vencimento + valor, ou cria recebiveis avulsos quando nao encontrar correspondencia.")
        arquivo_recebiveis = st.file_uploader(
            "Enviar planilha de recebiveis",
            type=["xlsx", "xls"],
            key="arquivo_importacao_recebiveis",
        )
        if arquivo_recebiveis is not None:
            try:
                recebiveis_preparados = preparar_planilha_importacao_recebiveis(arquivo_recebiveis)
                if not recebiveis_preparados:
                    st.warning("Nao encontrei recebiveis validos para importar nessa planilha.")
                else:
                    preview_recebiveis = pd.DataFrame(recebiveis_preparados).rename(
                        columns={
                            "paciente_nome": "Paciente",
                            "prontuario": "Prontuario",
                            "vencimento": "Vencimento",
                            "valor": "Valor",
                            "forma_pagamento": "Forma pagamento",
                            "status": "Status",
                            "data_pagamento": "Data pagamento",
                            "observacao": "Observacao",
                        }
                    )
                    preview_recebiveis["Valor"] = preview_recebiveis["Valor"].map(formatar_moeda_br)
                    st.success(f"{len(recebiveis_preparados)} recebiveis encontrados para importacao.")
                    st.dataframe(
                        preview_recebiveis[["Paciente", "Prontuario", "Vencimento", "Valor", "Forma pagamento", "Status", "Data pagamento", "Observacao"]],
                        use_container_width=True,
                        hide_index=True,
                    )
                    if st.button("Importar recebiveis da planilha"):
                        resultado = importar_recebiveis_preparados(recebiveis_preparados)
                        st.success(
                            f"Importacao concluida. Inseridos: {resultado['inseridos']} | Atualizados: {resultado['atualizados']}"
                        )
                        if resultado["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha de recebiveis: {exc}")

    with tab_imp_pagar:
        st.subheader("Importar contas a pagar por Excel")
        st.caption("Use a planilha do contas a pagar para trazer vencimentos futuros e titulos ja pagos para o sistema.")
        arquivo_pagar = st.file_uploader(
            "Enviar planilha de contas a pagar",
            type=["xlsx", "xls"],
            key="arquivo_importacao_pagar",
        )
        if arquivo_pagar is not None:
            try:
                contas_preparadas = preparar_planilha_importacao_contas_pagar(arquivo_pagar)
                if not contas_preparadas:
                    st.warning("Nao encontrei contas validas para importar nessa planilha.")
                else:
                    preview_pagar = pd.DataFrame(contas_preparadas).rename(
                        columns={
                            "data_vencimento": "Vencimento",
                            "descricao": "Descricao",
                            "fornecedor": "Fornecedor",
                            "valor": "Valor",
                            "pago": "Data pagamento",
                            "valor_pago": "Valor pago",
                            "status": "Status",
                            "observacao": "Observacao",
                        }
                    )
                    preview_pagar["Valor"] = preview_pagar["Valor"].map(formatar_moeda_br)
                    preview_pagar["Valor pago"] = preview_pagar["Valor pago"].map(formatar_moeda_br)
                    st.success(f"{len(contas_preparadas)} contas encontradas para importacao.")
                    st.dataframe(
                        preview_pagar[["Vencimento", "Descricao", "Fornecedor", "Valor", "Data pagamento", "Valor pago", "Status", "Observacao"]],
                        use_container_width=True,
                        hide_index=True,
                    )
                    if st.button("Importar contas a pagar da planilha"):
                        resultado = importar_contas_pagar_preparadas(contas_preparadas)
                        st.success(
                            f"Importacao concluida. Inseridas: {resultado['inseridos']} | Atualizadas: {resultado['atualizados']}"
                        )
                        if resultado["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha do contas a pagar: {exc}")

    with tab_imp_vendas:
        st.subheader("Importar vendas por Excel")
        st.caption("Registra as vendas no dashboard e tenta criar contratos automaticamente quando o paciente ja existe no cadastro.")
        arquivo_vendas = st.file_uploader(
            "Enviar planilha de vendas",
            type=["xlsx", "xls"],
            key="arquivo_importacao_vendas",
        )
        criar_contratos_automaticamente = st.checkbox(
            "Tentar criar contratos automaticamente para pacientes ja cadastrados",
            value=True,
            key="criar_contratos_vendas",
        )
        if arquivo_vendas is not None:
            try:
                vendas_preparadas = preparar_planilha_importacao_vendas(arquivo_vendas)
                if not vendas_preparadas:
                    st.warning("Nao encontrei vendas validas para importar nessa planilha.")
                else:
                    preview_vendas = pd.DataFrame(vendas_preparadas).rename(
                        columns={
                            "data_venda": "Data",
                            "paciente_nome": "Paciente",
                            "valor_total": "Valor total",
                            "valor_a_vista": "A vista",
                            "valor_cartao": "Cartao",
                            "valor_boleto": "Boleto",
                            "saldo": "Saldo",
                            "data_a_pagar": "Data a pagar",
                            "avaliador": "Avaliador",
                            "vendedor": "Vendedor",
                            "nf": "NF",
                        }
                    )
                    for coluna in ["Valor total", "A vista", "Cartao", "Boleto", "Saldo"]:
                        preview_vendas[coluna] = preview_vendas[coluna].map(formatar_moeda_br)
                    st.success(f"{len(vendas_preparadas)} vendas encontradas para importacao.")
                    st.dataframe(
                        preview_vendas[["Data", "Paciente", "Valor total", "A vista", "Cartao", "Boleto", "Saldo", "Data a pagar", "Avaliador", "Vendedor", "NF"]],
                        use_container_width=True,
                        hide_index=True,
                    )
                    if st.button("Importar vendas da planilha"):
                        resultado = importar_vendas_preparadas(vendas_preparadas, criar_contratos_automaticamente)
                        st.success(
                            f"Importacao concluida. Vendas inseridas: {resultado['inseridas']} | Ignoradas: {resultado['ignoradas']} | Contratos criados: {resultado['contratos_criados']}"
                        )
                        if resultado["erros"]:
                            st.warning("Alguns registros tiveram erro:")
                            for erro in resultado["erros"]:
                                st.write(f"- {erro}")
            except Exception as exc:
                st.error(f"Falha ao processar a planilha de vendas: {exc}")

if menu == "Usuarios":
    st.title("Usuarios e Acessos")

    if usuario_logado["perfil"] != "Administrador":
        st.error("Apenas administradores podem gerenciar usuarios.")
        st.stop()

    tab_novo_usuario, tab_gerenciar_usuario, tab_logs_acesso = st.tabs(["Novo usuario", "Gerenciar usuarios", "Logs de acesso"])

    with tab_novo_usuario:
        st.subheader("Criar usuario")
        nu1, nu2 = st.columns(2)
        novo_nome = nu1.text_input("Nome do usuario", key="novo_usuario_nome")
        novo_usuario = nu2.text_input("Login", key="novo_usuario_login")
        nu3, nu4 = st.columns(2)
        novo_perfil = nu3.selectbox("Perfil", ["Administrador", "Usuario"], key="novo_usuario_perfil")
        nova_senha = nu4.text_input("Senha inicial", type="password", key="novo_usuario_senha")
        st.markdown("**Permissoes do novo usuario**")
        np1, np2, np3 = st.columns(3)
        novo_acesso_dashboard = np1.checkbox("Dashboard", value=True, key="novo_acesso_dashboard")
        novo_acesso_pacientes = np1.checkbox("Pacientes", value=True, key="novo_acesso_pacientes")
        novo_acesso_contratos = np2.checkbox("Contratos", value=True, key="novo_acesso_contratos")
        novo_acesso_financeiro = np2.checkbox("Financeiro", value=(novo_perfil == "Administrador"), key="novo_acesso_financeiro")
        novo_acesso_usuarios = np3.checkbox("Usuarios", value=(novo_perfil == "Administrador"), key="novo_acesso_usuarios")

        if st.button("Criar usuario"):
            if not novo_nome.strip():
                st.error("Informe o nome do usuario.")
            elif not novo_usuario.strip():
                st.error("Informe o login do usuario.")
            elif len(nova_senha) < 4:
                st.error("A senha inicial deve ter pelo menos 4 caracteres.")
            elif usuario_existe(novo_usuario):
                st.error("Ja existe um usuario com esse login.")
            else:
                criar_usuario(novo_nome, novo_usuario, nova_senha, novo_perfil)
                cursor.execute(
                    """
                    UPDATE usuarios
                    SET acesso_dashboard=?, acesso_pacientes=?, acesso_contratos=?, acesso_financeiro=?, acesso_usuarios=?
                    WHERE lower(usuario)=lower(?)
                    """,
                    (
                        int(bool(novo_acesso_dashboard)),
                        int(bool(novo_acesso_pacientes)),
                        int(bool(novo_acesso_contratos)),
                        int(bool(novo_acesso_financeiro)),
                        int(bool(novo_acesso_usuarios)),
                        novo_usuario.strip(),
                    ),
                )
                conn.commit()
                st.success("Usuario criado com sucesso.")
                st.rerun()

    with tab_gerenciar_usuario:
        st.subheader("Gerenciar usuarios")
        usuarios_df = carregar_usuarios()
        if usuarios_df.empty:
            st.info("Nenhum usuario cadastrado.")
        else:
            opcoes_usuarios = [
                (
                    int(row["id"]),
                    f"{row['nome']} - {row['usuario']} - {row['perfil']} - {'Ativo' if int(row['ativo']) == 1 else 'Inativo'}",
                )
                for _, row in usuarios_df.iterrows()
            ]
            usuario_id_edicao = st.selectbox(
                "Usuario",
                options=[opcao[0] for opcao in opcoes_usuarios],
                format_func=lambda valor: next(label for chave, label in opcoes_usuarios if chave == valor),
                key="usuario_id_edicao",
            )
            usuario_edicao = usuarios_df[usuarios_df["id"] == usuario_id_edicao].iloc[0]

            gu1, gu2 = st.columns(2)
            editar_nome = gu1.text_input("Nome", value=usuario_edicao["nome"] or "", key="editar_usuario_nome")
            editar_login = gu2.text_input("Login", value=usuario_edicao["usuario"] or "", key="editar_usuario_login")
            gu3, gu4 = st.columns(2)
            editar_perfil = gu3.selectbox(
                "Perfil",
                ["Administrador", "Usuario"],
                index=["Administrador", "Usuario"].index(usuario_edicao["perfil"] or "Usuario"),
                key="editar_usuario_perfil",
            )
            editar_ativo = gu4.checkbox(
                "Usuario ativo",
                value=int(usuario_edicao["ativo"] or 0) == 1,
                key="editar_usuario_ativo",
            )
            st.markdown("**Permissoes por modulo**")
            ep1, ep2, ep3 = st.columns(3)
            editar_acesso_dashboard = ep1.checkbox("Dashboard", value=int(usuario_edicao["acesso_dashboard"] or 0) == 1, key="editar_acesso_dashboard")
            editar_acesso_pacientes = ep1.checkbox("Pacientes", value=int(usuario_edicao["acesso_pacientes"] or 0) == 1, key="editar_acesso_pacientes")
            editar_acesso_contratos = ep2.checkbox("Contratos", value=int(usuario_edicao["acesso_contratos"] or 0) == 1, key="editar_acesso_contratos")
            editar_acesso_financeiro = ep2.checkbox("Financeiro", value=int(usuario_edicao["acesso_financeiro"] or 0) == 1, key="editar_acesso_financeiro")
            editar_acesso_usuarios = ep3.checkbox("Usuarios", value=int(usuario_edicao["acesso_usuarios"] or 0) == 1, key="editar_acesso_usuarios")

            if st.button("Salvar dados do usuario"):
                if not editar_nome.strip():
                    st.error("Informe o nome do usuario.")
                elif not editar_login.strip():
                    st.error("Informe o login do usuario.")
                elif usuario_existe(editar_login, ignorar_id=usuario_id_edicao):
                    st.error("Ja existe outro usuario com esse login.")
                else:
                    atualizar_usuario_admin(
                        usuario_id_edicao,
                        editar_nome,
                        editar_login,
                        editar_perfil,
                        editar_ativo,
                        editar_acesso_dashboard,
                        editar_acesso_pacientes,
                        editar_acesso_contratos,
                        editar_acesso_financeiro,
                        editar_acesso_usuarios,
                    )
                    conn.commit()
                    st.success("Usuario atualizado.")
                    st.rerun()

            st.markdown("**Redefinir senha**")
            nova_senha_usuario = st.text_input("Nova senha", type="password", key="redefinir_senha_usuario")
            if st.button("Redefinir senha do usuario"):
                if len(nova_senha_usuario) < 4:
                    st.error("A nova senha deve ter pelo menos 4 caracteres.")
                else:
                    redefinir_senha_usuario(usuario_id_edicao, nova_senha_usuario)
                    conn.commit()
                    st.success("Senha redefinida com sucesso.")

            usuarios_exibicao = usuarios_df.copy()
            usuarios_exibicao["ativo"] = usuarios_exibicao["ativo"].apply(lambda valor: "Ativo" if int(valor or 0) == 1 else "Inativo")
            usuarios_exibicao["data_criacao"] = usuarios_exibicao["data_criacao"].apply(formatar_data_br_valor)
            usuarios_exibicao = usuarios_exibicao.rename(
                columns={
                    "nome": "Nome",
                    "usuario": "Login",
                    "perfil": "Perfil",
                    "ativo": "Status",
                    "data_criacao": "Criado em",
                }
            )[["Nome", "Login", "Perfil", "Status", "Criado em"]]
            st.markdown("**Usuarios cadastrados**")
            st.dataframe(usuarios_exibicao, use_container_width=True, hide_index=True)

    with tab_logs_acesso:
        st.subheader("Logs de acesso")
        logs_df = pd.read_sql("SELECT usuario, evento, data_hora FROM logs_acesso ORDER BY id DESC LIMIT 200", conn)
        if logs_df.empty:
            st.info("Nenhum log de acesso registrado ainda.")
        else:
            logs_df["data_hora"] = logs_df["data_hora"].apply(formatar_data_hora_br_valor)
            logs_df = logs_df.rename(columns={"usuario": "Usuario", "evento": "Evento", "data_hora": "Data"})
            st.dataframe(logs_df, use_container_width=True, hide_index=True)

if menu == "Financeiro":
    st.title("Controle Financeiro")
    preencher_prontuarios_recebiveis()
    atualizar_status_contas_pagar_automaticamente()
    financeiro = pd.read_sql("SELECT * FROM financeiro ORDER BY data DESC", conn)
    saldos_conta = pd.read_sql("SELECT * FROM saldos_conta ORDER BY data DESC, id DESC", conn)
    contas_pagar = pd.read_sql("SELECT * FROM contas_pagar ORDER BY data_vencimento, fornecedor, descricao", conn)
    recebiveis = pd.read_sql(
        "SELECT * FROM recebiveis ORDER BY vencimento, paciente_nome, parcela_numero",
        conn,
    )
    foco_financeiro = st.session_state.pop("financeiro_foco", None)
    if foco_financeiro:
        texto_foco = []
        if foco_financeiro.get("paciente_nome"):
            texto_foco.append(f"Paciente: {foco_financeiro['paciente_nome']}")
        if foco_financeiro.get("contrato_id"):
            texto_foco.append(f"Contrato: {int(foco_financeiro['contrato_id'])}")
        st.info("Foco financeiro aberto a partir da agenda. " + " | ".join(texto_foco))
    tab_caixa, tab_visao, tab_individual, tab_lote, tab_pagar, tab_calendario_pagar, tab_novo_pagar = st.tabs(
        ["Caixa", "Visao de Recebiveis", "Editar Individual", "Editar em Lote", "Contas a Pagar", "Calendario", "Nova Divida"]
    )

    if not financeiro.empty:
        financeiro["prontuario"] = financeiro["prontuario"].fillna("").replace("None", "")
        financeiro["forma_pagamento"] = financeiro["forma_pagamento"].fillna("").replace("None", "")
        financeiro["conta_caixa"] = financeiro["conta_caixa"].fillna("").replace("None", "")
        financeiro["observacao"] = financeiro["observacao"].fillna("").replace("None", "")
        financeiro["data_exibicao"] = financeiro["data"].apply(formatar_data_br_valor)

    caixa_real = financeiro.copy()
    if not caixa_real.empty:
        caixa_real = caixa_real[
            ~(
                (caixa_real["descricao"].fillna("") == DESCRICAO_CONTRATO)
                & (caixa_real["recebivel_id"].isna())
            )
        ].copy()

    if not recebiveis.empty:
        recebiveis["valor"] = recebiveis["valor"].astype(float)
        recebiveis["prontuario"] = recebiveis["prontuario"].fillna("").replace("None", "")
        recebiveis["observacao"] = recebiveis["observacao"].fillna("").replace("None", "")
        recebiveis["data_pagamento"] = recebiveis["data_pagamento"].fillna("").replace("None", "")
        recebiveis["data_pagamento"] = recebiveis["data_pagamento"].apply(formatar_data_br_valor)
        recebiveis["_ordem_vencimento"] = pd.to_datetime(recebiveis["vencimento"], format="%d/%m/%Y", errors="coerce")
        recebiveis = recebiveis.sort_values(["_ordem_vencimento", "paciente_nome", "parcela_numero"], na_position="last")

    if not contas_pagar.empty:
        contas_pagar["valor"] = contas_pagar["valor"].astype(float)
        contas_pagar["valor_pago"] = contas_pagar["valor_pago"].fillna(0).astype(float)
        contas_pagar["pago"] = contas_pagar["pago"].fillna("").replace("None", "").apply(formatar_data_br_valor)
        contas_pagar["observacao"] = contas_pagar["observacao"].fillna("").replace("None", "")
        contas_pagar["_ordem_vencimento"] = pd.to_datetime(contas_pagar["data_vencimento"], format="%d/%m/%Y", errors="coerce")
        contas_pagar = contas_pagar.sort_values(["_ordem_vencimento", "fornecedor", "descricao"], na_position="last")

    with tab_caixa:
        st.subheader("Lancamento manual")
        blocos_caixa = selecionar_blocos_visiveis(
            "Exibir em caixa",
            ["Lancamento manual", "Saldos do dia anterior", "Baixa de recebivel", "Livro-caixa", "Importar extrato"],
            ["Lancamento manual", "Baixa de recebivel", "Livro-caixa"],
            "financeiro_blocos_caixa",
        )
        if "Lancamento manual" not in blocos_caixa:
            st.caption("Use os botões acima para mostrar apenas os blocos que deseja ver nesta aba.")
        if "Lancamento manual" in blocos_caixa:
            cx1, cx2, cx3, cx4 = st.columns(4)
            caixa_data = cx1.date_input("Data do movimento", value=date.today(), key="caixa_data")
            caixa_tipo = cx2.selectbox("Tipo", ["Entrada", "Saida"], key="caixa_tipo")
            caixa_forma = cx3.selectbox("Forma pagamento", ["", *FORMAS_PAGAMENTO], key="caixa_forma")
            caixa_conta = cx4.selectbox("Conta/Banco", CONTAS_CAIXA_MODELO, key="caixa_conta")

            cx5, cx6 = st.columns(2)
            caixa_origem = cx5.text_input("Origem", help="Ex.: banco, pagamento, fornecedor.", key="caixa_origem")
            caixa_prontuario = cx6.text_input("Prontuario", key="caixa_prontuario")

            cx7, cx8 = st.columns(2)
            caixa_descricao = cx7.text_input("Descricao", key="caixa_descricao")
            caixa_valor = cx8.number_input("Valor", min_value=0.0, value=0.0, key="caixa_valor")
            caixa_observacao = st.text_area("Observacao", key="caixa_observacao")

            if st.button("Registrar no caixa"):
                if not caixa_descricao.strip():
                    st.error("Informe a descricao do movimento.")
                elif caixa_valor <= 0:
                    st.error("Informe um valor maior que zero.")
                else:
                    registrar_movimento_caixa(
                        origem=caixa_origem or "Caixa manual",
                        descricao=caixa_descricao,
                        valor=caixa_valor,
                        tipo=caixa_tipo,
                        data_movimento=caixa_data,
                        prontuario=caixa_prontuario,
                        forma_pagamento=caixa_forma,
                        conta_caixa=caixa_conta,
                        observacao=caixa_observacao,
                    )
                    conn.commit()
                    st.success("Movimento registrado no caixa.")
                    st.rerun()

        if "Saldos do dia anterior" in blocos_caixa:
            with st.expander("Saldos do dia anterior", expanded=False):
                st.caption("Use este bloco para registrar os saldos de abertura do dia anterior em cada conta.")
                saldo_data = st.date_input(
                    "Data de referencia dos saldos",
                    value=date.today() - timedelta(days=1),
                    key="saldo_conta_data_anterior",
                )
                saldo_observacao = st.text_input("Observacao dos saldos", key="saldo_conta_observacao")

                col_saldos_1, col_saldos_2, col_saldos_3 = st.columns(3)
                colunas_saldos = [col_saldos_1, col_saldos_2, col_saldos_3]
                saldos_para_registro = {}
                for indice, conta in enumerate(CONTAS_CAIXA_MODELO):
                    coluna = colunas_saldos[indice % len(colunas_saldos)]
                    saldos_para_registro[conta] = coluna.number_input(
                        f"Saldo {conta}",
                        value=0.0,
                        key=f"saldo_{conta.lower()}_anterior",
                    )

                if st.button("Registrar saldos do dia anterior"):
                    for conta, saldo_valor in saldos_para_registro.items():
                        registrar_saldo_conta(
                            data_saldo=saldo_data,
                            conta=conta,
                            saldo=saldo_valor,
                            observacao=saldo_observacao,
                        )
                    conn.commit()
                    st.success("Saldos do dia anterior registrados.")
                    st.rerun()

                if not saldos_conta.empty:
                    st.markdown("**Ultimos saldos informados**")
                    saldos_exibicao = saldos_conta.copy()
                    saldos_exibicao["data"] = saldos_exibicao["data"].apply(formatar_data_br_valor)
                    saldos_exibicao["saldo"] = saldos_exibicao["saldo"].map(formatar_moeda_br)
                    saldos_exibicao = saldos_exibicao.rename(
                        columns={
                            "data": "Data",
                            "conta": "Conta",
                            "saldo": "Saldo",
                            "observacao": "Observacao",
                        }
                    )[["Data", "Conta", "Saldo", "Observacao"]]
                    st.dataframe(saldos_exibicao.head(20), use_container_width=True, hide_index=True)

        if "Baixa de recebivel" in blocos_caixa:
            with st.expander("Baixa de recebivel", expanded=True):
                if recebiveis.empty:
                    st.info("Nao ha recebiveis cadastrados.")
                else:
                    recebiveis_abertos = recebiveis[recebiveis["status"].isin(["Aberto", "Atrasado"])].copy()
                    if recebiveis_abertos.empty:
                        st.info("Nao ha recebiveis em aberto para baixa.")
                    else:
                        filtro_baixa = st.text_input(
                            "Buscar recebivel para baixa",
                            help="Pesquise por paciente, prontuario, parcela ou vencimento.",
                            key="busca_baixa_recebivel",
                        )
                        if filtro_baixa.strip():
                            termo_baixa = normalizar_texto(filtro_baixa)
                            recebiveis_abertos = recebiveis_abertos[
                                recebiveis_abertos.apply(
                                    lambda row: termo_baixa in normalizar_texto(
                                        f"{row['paciente_nome']} {row['prontuario']} {row['parcela_numero']} {row['vencimento']}"
                                    ),
                                    axis=1,
                                )
                            ]

                        if recebiveis_abertos.empty:
                            st.warning("Nenhum recebivel encontrado para a busca informada.")
                        else:
                            opcoes_baixa = [
                                (
                                    int(row["id"]),
                                    f"{row['paciente_nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])} - Parcela {formatar_parcela_valor(row['parcela_numero'])} - {row['vencimento']} - {formatar_moeda_br(row['valor'])}"
                                )
                                for _, row in recebiveis_abertos.iterrows()
                            ]
                            recebivel_baixa_id = st.selectbox(
                                "Recebivel para dar baixa",
                                options=[opcao[0] for opcao in opcoes_baixa],
                                format_func=lambda recebivel_id: next(
                                    label for valor, label in opcoes_baixa if valor == recebivel_id
                                ),
                                key="recebivel_baixa_id",
                            )
                            recebivel_baixa = recebiveis_abertos[recebiveis_abertos["id"] == recebivel_baixa_id].iloc[0]

                            bx1, bx2, bx3, bx4 = st.columns(4)
                            baixa_data = bx1.date_input("Data do pagamento", value=date.today(), key="baixa_data")
                            baixa_origem = bx2.text_input("Origem da entrada", value="Pagamento", key="baixa_origem")
                            baixa_forma = bx3.selectbox(
                                "Forma de recebimento",
                                FORMAS_PAGAMENTO,
                                index=FORMAS_PAGAMENTO.index(normalizar_forma_pagamento(recebivel_baixa["forma_pagamento"])),
                                key="baixa_forma",
                            )
                            baixa_conta = bx4.selectbox("Conta/Banco da entrada", CONTAS_CAIXA_MODELO, key="baixa_conta")
                            baixa_observacao = st.text_area("Observacao da baixa", key="baixa_observacao")

                            if st.button("Dar baixa no recebivel"):
                                try:
                                    baixar_recebivel_no_caixa(
                                        recebivel_id=recebivel_baixa_id,
                                        origem=baixa_origem,
                                        data_pagamento=baixa_data,
                                        observacao=baixa_observacao,
                                        forma_recebimento=baixa_forma,
                                        conta_caixa=baixa_conta,
                                    )
                                    conn.commit()
                                    st.success("Recebivel baixado no caixa e marcado como pago.")
                                    st.rerun()
                                except ValueError as exc:
                                    st.error(str(exc))

        if "Livro-caixa" in blocos_caixa:
            with st.expander("Livro-caixa", expanded=True):
                st.caption("O caixa mostra apenas movimentacoes efetivas que passaram pela conta ou foram registradas manualmente. Previsoes e contratos futuros nao entram aqui.")
                if caixa_real.empty:
                    st.info("Nenhum lancamento no caixa ainda.")
                else:
                    resumo_caixa_df, grupos_caixa = montar_caixa_diario(caixa_real)

                    st.markdown("**Resumo diario**")
                    resumo_exibicao = resumo_caixa_df.copy()
                    for coluna in ["Entradas", "Saidas", "Saldo do dia"]:
                        resumo_exibicao[coluna] = resumo_exibicao[coluna].map(formatar_moeda_br)
                    st.dataframe(resumo_exibicao, use_container_width=True, hide_index=True)

                    st.markdown("**Lancamentos por dia**")
                    for grupo in grupos_caixa:
                        data_referencia = parse_data_contrato(grupo["data"])
                        saldos_abertura = saldos_informados_por_conta_ate(data_referencia) if data_referencia else {conta: 0.0 for conta in CONTAS_CAIXA_MODELO}
                        saldos_fechamento = saldos_abertura.copy()
                        detalhamento_para_saldo = caixa_real.copy()
                        detalhamento_para_saldo["data_date"] = detalhamento_para_saldo["data"].apply(parse_data_contrato)
                        detalhamento_para_saldo["conta_resolvida"] = detalhamento_para_saldo.apply(
                            lambda row: str(row.get("conta_caixa") or "").strip().upper()
                            or identificar_conta_caixa(row.get("origem", ""), row.get("descricao", ""), row.get("observacao", "")),
                            axis=1,
                        )
                        if data_referencia:
                            movimentos_dia = detalhamento_para_saldo[detalhamento_para_saldo["data_date"] == data_referencia]
                            for _, mov in movimentos_dia.iterrows():
                                conta = mov["conta_resolvida"] if mov["conta_resolvida"] in CONTAS_CAIXA_MODELO else "CAIXA"
                                valor = float(mov["valor"] or 0)
                                if str(mov["tipo"]) == "Entrada":
                                    saldos_fechamento[conta] += valor
                                else:
                                    saldos_fechamento[conta] -= valor

                        st.markdown(f"### {grupo['data']}")
                        m1, m2, m3, m4 = st.columns(4)
                        m1.metric("Entradas do dia", formatar_moeda_br(grupo["entradas"]))
                        m2.metric("Saidas do dia", formatar_moeda_br(grupo["saidas"]))
                        m3.metric("Saldo do dia", formatar_moeda_br(grupo["saldo"]))
                        m4.metric("Lancamentos", len(grupo["detalhamento"]))

                        st.markdown("**Saldos de abertura**")
                        resumo_abertura_df = montar_resumo_saldos_conta(saldos_abertura, data_referencia)
                        st.dataframe(
                            estilizar_resumo_saldos(resumo_abertura_df),
                            use_container_width=True,
                            hide_index=True,
                        )

                        detalhamento_exibicao = grupo["detalhamento"].copy()
                        detalhamento_exibicao["Valor"] = detalhamento_exibicao["Valor"].map(formatar_moeda_br)
                        st.dataframe(detalhamento_exibicao, use_container_width=True, hide_index=True)

                        st.markdown("**Saldos de fechamento**")
                        resumo_fechamento_df = montar_resumo_saldos_conta(saldos_fechamento, data_referencia)
                        st.dataframe(
                            estilizar_resumo_saldos(resumo_fechamento_df),
                            use_container_width=True,
                            hide_index=True,
                        )

                    exp1, exp2 = st.columns(2)
                    if OPENPYXL_DISPONIVEL:
                        exp1.download_button(
                            "Exportar caixa para Excel",
                            data=caixa_diario_para_excel_bytes(caixa_real),
                            file_name=f"caixa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )
                    else:
                        exp1.info("Excel indisponivel: instale openpyxl para habilitar esta exportacao.")
                    if REPORTLAB_DISPONIVEL:
                        caixa_pdf_df = caixa_real[
                            ["data_exibicao", "tipo", "origem", "conta_caixa", "prontuario", "descricao", "valor", "forma_pagamento", "observacao"]
                        ].copy()
                        caixa_pdf_df = caixa_pdf_df.rename(
                            columns={
                                "data_exibicao": "Data",
                                "tipo": "Tipo",
                                "origem": "Origem",
                                "conta_caixa": "Conta/Banco",
                                "prontuario": "Prontuario",
                                "descricao": "Descricao",
                                "valor": "Valor",
                                "forma_pagamento": "Forma de pagamento",
                                "observacao": "Observacao",
                            }
                        )
                        caixa_pdf_df["Valor"] = caixa_pdf_df["Valor"].map(formatar_moeda_br)
                        exp2.download_button(
                            "Exportar caixa para PDF",
                            data=dataframe_para_pdf_bytes(caixa_pdf_df, titulo="Livro-caixa"),
                            file_name=f"caixa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                        )
                    else:
                        exp2.info("PDF indisponivel: instale reportlab para habilitar esta exportacao.")

        if "Importar extrato" in blocos_caixa:
            with st.expander("Importar extrato bancario", expanded=False):
                st.caption("Reconhecimento automatico disponivel para CSV e Excel (.xlsx/.xls). PDF bancario fica como proxima etapa.")
                arquivo_extrato = st.file_uploader(
                    "Enviar extrato",
                    type=["csv", "xlsx", "xls"],
                    key="arquivo_extrato_caixa",
                )
                if arquivo_extrato is not None:
                    try:
                        extrato_origem = ler_extrato_arquivo(arquivo_extrato)
                        extrato_reconhecido = reconhecer_extrato(extrato_origem)
                        if extrato_reconhecido.empty:
                            st.error("Nao foi possivel reconhecer automaticamente esse extrato. Verifique se o arquivo tem colunas de data, descricao e valor.")
                        else:
                            st.success(f"{len(extrato_reconhecido)} transacoes reconhecidas.")
                            preview_extrato = extrato_reconhecido.rename(
                                columns={
                                    "data_exibicao": "Data",
                                    "descricao": "Descricao",
                                    "valor": "Valor",
                                    "tipo": "Tipo",
                                }
                            )[["Data", "Descricao", "Valor", "Tipo"]].copy()
                            preview_extrato["Valor"] = preview_extrato["Valor"].map(formatar_moeda_br)
                            st.dataframe(preview_extrato, use_container_width=True, hide_index=True)

                            ie1, ie2, ie3 = st.columns(3)
                            origem_extrato = ie1.text_input("Origem para importacao", value="Extrato bancario", key="origem_import_extrato")
                            prontuario_extrato = ie2.text_input("Prontuario padrao (opcional)", key="prontuario_import_extrato")
                            forma_extrato = ie3.selectbox("Forma pagamento padrao", ["", *FORMAS_PAGAMENTO], key="forma_import_extrato")
                            observacao_extrato = st.text_area("Observacao padrao da importacao", key="obs_import_extrato")

                            if st.button("Importar extrato para o caixa"):
                                for _, row in extrato_reconhecido.iterrows():
                                    registrar_movimento_caixa(
                                        origem=origem_extrato or "Extrato bancario",
                                        descricao=row["descricao"],
                                        valor=float(row["valor"]),
                                        tipo=row["tipo"],
                                        data_movimento=row["data"].date() if hasattr(row["data"], "date") else row["data"],
                                        prontuario=prontuario_extrato,
                                        forma_pagamento=forma_extrato,
                                        observacao=observacao_extrato,
                                    )
                                conn.commit()
                                st.success("Extrato importado para o caixa.")
                                st.rerun()
                    except Exception as exc:
                        st.error(f"Falha ao ler o extrato: {exc}")

    with tab_visao:
        st.subheader("Recebiveis")
        blocos_recebiveis = selecionar_blocos_visiveis(
            "Exibir em recebíveis",
            ["Filtros", "Resumo", "Agenda mensal", "Detalhamento"],
            ["Filtros", "Resumo", "Agenda mensal"],
            "financeiro_blocos_recebiveis",
        )
        filtro_nomes = []
        filtro_prontuarios = []
        filtro_status = []
        filtro_formas = []
        filtro_periodo = ()
        mapa_meses = {
            1: "Janeiro", 2: "Fevereiro", 3: "Marco", 4: "Abril", 5: "Maio", 6: "Junho",
            7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
        }
        filtro_anos_recebiveis = []
        filtro_meses_recebiveis = []
        if "Filtros" in blocos_recebiveis:
            with st.expander("Filtros de recebíveis", expanded=True):
                col_filtro_1, col_filtro_2 = st.columns(2)
                filtro_nomes = col_filtro_1.multiselect(
                    "Pacientes",
                    options=sorted(recebiveis["paciente_nome"].fillna("").unique().tolist()),
                )
                filtro_prontuarios = col_filtro_2.multiselect(
                    "Prontuarios",
                    options=sorted([valor for valor in recebiveis["prontuario"].fillna("").unique().tolist() if valor]),
                )

                col_filtro_4, col_filtro_5, col_filtro_6 = st.columns(3)
                opcoes_status = sorted(recebiveis["status"].fillna("Aberto").unique().tolist())
                filtro_status = col_filtro_4.multiselect("Status", options=opcoes_status)
                opcoes_forma = sorted(recebiveis["forma_pagamento"].fillna("").unique().tolist())
                filtro_formas = col_filtro_5.multiselect("Formas de pagamento", options=opcoes_forma)
                datas_disponiveis = [
                    valor.date()
                    for valor in recebiveis["_ordem_vencimento"].dropna().sort_values().unique().tolist()
                ]
                data_inicio = datas_disponiveis[0] if datas_disponiveis else None
                data_fim = datas_disponiveis[-1] if datas_disponiveis else None
                filtro_periodo = col_filtro_6.date_input(
                    "Periodo de vencimento",
                    value=(data_inicio, data_fim) if data_inicio and data_fim else (),
                )
                col_filtro_7, col_filtro_8 = st.columns(2)
                anos_disponiveis_recebiveis = sorted(
                    recebiveis["_ordem_vencimento"].dropna().dt.year.unique().tolist()
                ) if not recebiveis.empty else []
                filtro_anos_recebiveis = col_filtro_7.multiselect(
                    "Anos",
                    options=anos_disponiveis_recebiveis,
                )
                filtro_meses_recebiveis = col_filtro_8.multiselect(
                    "Meses",
                    options=list(mapa_meses.keys()),
                    format_func=lambda valor: mapa_meses.get(valor, str(valor)),
                )

        recebiveis_filtrados = recebiveis.copy()
        if filtro_nomes:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["paciente_nome"].isin(filtro_nomes)]
        if filtro_prontuarios:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["prontuario"].isin(filtro_prontuarios)]
        if filtro_status:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["status"].isin(filtro_status)]
        if filtro_formas:
            recebiveis_filtrados = recebiveis_filtrados[recebiveis_filtrados["forma_pagamento"].isin(filtro_formas)]
        if isinstance(filtro_periodo, tuple) and len(filtro_periodo) == 2 and filtro_periodo[0] and filtro_periodo[1]:
            inicio = pd.to_datetime(filtro_periodo[0])
            fim = pd.to_datetime(filtro_periodo[1])
            recebiveis_filtrados = recebiveis_filtrados[
                (recebiveis_filtrados["_ordem_vencimento"] >= inicio) &
                (recebiveis_filtrados["_ordem_vencimento"] <= fim)
            ]
        if filtro_anos_recebiveis:
            recebiveis_filtrados = recebiveis_filtrados[
                recebiveis_filtrados["_ordem_vencimento"].dt.year.isin(filtro_anos_recebiveis)
            ]
        if filtro_meses_recebiveis:
            recebiveis_filtrados = recebiveis_filtrados[
                recebiveis_filtrados["_ordem_vencimento"].dt.month.isin(filtro_meses_recebiveis)
            ]

        if "Resumo" in blocos_recebiveis:
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Total filtrado", formatar_moeda_br(recebiveis_filtrados["valor"].sum()))
            c2.metric("Parcelas filtradas", str(len(recebiveis_filtrados)))
            c3.metric("Vencimentos unicos", str(recebiveis_filtrados["vencimento"].nunique()))
            c4.metric("Pacientes unicos", str(recebiveis_filtrados["paciente_nome"].nunique()))

        if not recebiveis_filtrados.empty:
            recebiveis_excel = recebiveis_filtrados[
                ["paciente_nome", "prontuario", "parcela_numero", "vencimento", "data_pagamento", "valor", "forma_pagamento", "status", "observacao"]
            ].copy()
            recebiveis_excel = recebiveis_excel.rename(
                columns={
                    "paciente_nome": "Paciente",
                    "prontuario": "Prontuario",
                    "parcela_numero": "Parcela",
                    "vencimento": "Vencimento",
                    "data_pagamento": "Data do pagamento",
                    "valor": "Valor",
                    "forma_pagamento": "Forma pagamento",
                    "status": "Status",
                    "observacao": "Observacao",
                }
            )
            recebiveis_excel["Prontuario"] = recebiveis_excel["Prontuario"].apply(formatar_prontuario_valor)
            recebiveis_excel["Parcela"] = recebiveis_excel["Parcela"].apply(formatar_parcela_valor)
            recebiveis_excel["Data do pagamento"] = recebiveis_excel["Data do pagamento"].apply(formatar_data_br_valor)
            if OPENPYXL_DISPONIVEL:
                excel_recebiveis = dataframe_para_excel_bytes(recebiveis_excel, nome_aba="Recebiveis")
                st.download_button(
                    "Baixar recebíveis em Excel",
                    data=excel_recebiveis,
                    file_name="recebiveis_filtrados.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_recebiveis_excel",
                )
            else:
                st.info("Excel indisponível: instale openpyxl para habilitar esta exportação.")

            agenda_mensal = recebiveis_filtrados.copy()
            agenda_mensal["ano"] = agenda_mensal["_ordem_vencimento"].dt.year
            agenda_mensal["mes"] = agenda_mensal["_ordem_vencimento"].dt.month
            agenda_mensal["mes_nome"] = agenda_mensal["mes"].map(mapa_meses)
            resumo_por_mes = (
                agenda_mensal.groupby(["ano", "mes", "mes_nome"], as_index=False)
                .agg(
                    total_previsto=("valor", "sum"),
                    titulos=("id", "count"),
                )
                .sort_values(["ano", "mes"], ascending=[True, True])
            )
            if "Agenda mensal" in blocos_recebiveis:
                with st.expander("Agenda mensal de vencimentos", expanded=False):
                    for _, resumo_mes in resumo_por_mes.iterrows():
                        ano_agenda = int(resumo_mes["ano"])
                        mes_agenda = int(resumo_mes["mes"])
                        mes_nome = resumo_mes["mes_nome"]
                        total_mes = float(resumo_mes["total_previsto"] or 0)
                        titulos_mes = int(resumo_mes["titulos"] or 0)
                        itens_mes = agenda_mensal[
                            (agenda_mensal["ano"] == ano_agenda) & (agenda_mensal["mes"] == mes_agenda)
                        ].copy()
                        resumo_dias = (
                            itens_mes.groupby("vencimento", as_index=False)
                            .agg(
                                total_previsto=("valor", "sum"),
                                titulos=("id", "count"),
                            )
                        )
                        resumo_dias["_ordem_vencimento"] = pd.to_datetime(
                            resumo_dias["vencimento"],
                            format="%d/%m/%Y",
                            errors="coerce",
                        )
                        resumo_dias = resumo_dias.sort_values("_ordem_vencimento", na_position="last").drop(columns=["_ordem_vencimento"])
                        resumo_dias["total_previsto"] = resumo_dias["total_previsto"].map(formatar_moeda_br)
                        resumo_dias = resumo_dias.rename(
                            columns={
                                "vencimento": "Vencimento",
                                "total_previsto": "Total previsto",
                                "titulos": "Titulos",
                            }
                        )

                        with st.expander(f"{mes_nome}/{ano_agenda}  |  Total previsto: {formatar_moeda_br(total_mes)}  |  Titulos: {titulos_mes}", expanded=False):
                            st.dataframe(resumo_dias, use_container_width=True, hide_index=True)

            detalhe = recebiveis_filtrados[
                ["id", "paciente_nome", "prontuario", "parcela_numero", "vencimento", "data_pagamento", "valor", "forma_pagamento", "status", "observacao"]
            ].copy()
            detalhe = detalhe.rename(
                columns={
                    "id": "ID",
                    "paciente_nome": "Paciente",
                    "prontuario": "Prontuario",
                    "parcela_numero": "Parcela",
                    "vencimento": "Vencimento",
                    "data_pagamento": "Data do pagamento",
                    "valor": "Valor",
                    "forma_pagamento": "Forma pagamento",
                    "status": "Status",
                    "observacao": "Observacao",
                }
            )
            detalhe = detalhe.reset_index(drop=True)
            detalhe["Valor"] = detalhe["Valor"].map(lambda valor: f"R$ {float(valor):.2f}")
            if "Detalhamento" in blocos_recebiveis:
                with st.expander("Detalhamento dos recebiveis", expanded=False):
                    st.dataframe(detalhe, use_container_width=True, hide_index=True)

    with tab_individual:
        st.subheader("Editar recebivel individual")
        if recebiveis.empty:
            st.info("Nao ha recebiveis para editar.")
        else:
            opcoes_recebiveis = [
                (
                    int(row["id"]),
                    f"{row['paciente_nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])} - Parcela {formatar_parcela_valor(row['parcela_numero'])} - {row['vencimento']}"
                )
                for _, row in recebiveis.iterrows()
            ]
            recebivel_id = st.selectbox(
                "Recebivel",
                options=[opcao[0] for opcao in opcoes_recebiveis],
                format_func=lambda recebivel_id: next(
                    label for valor, label in opcoes_recebiveis if valor == recebivel_id
                ),
            )
            recebivel = recebiveis[recebiveis["id"] == recebivel_id].iloc[0]

            edit_col_1, edit_col_2 = st.columns(2)
            edit_nome = edit_col_1.text_input("Nome do paciente", recebivel["paciente_nome"] or "", key=f"rec_nome_{recebivel_id}")
            edit_prontuario = edit_col_2.text_input("Prontuario", recebivel["prontuario"] or "", key=f"rec_pront_{recebivel_id}")

            edit_col_3, edit_col_4, edit_col_5 = st.columns(3)
            edit_vencimento = edit_col_3.text_input("Vencimento", recebivel["vencimento"] or "", key=f"rec_venc_{recebivel_id}")
            edit_valor = edit_col_4.number_input(
                "Valor",
                min_value=0.0,
                value=float(recebivel["valor"] or 0),
                key=f"rec_valor_{recebivel_id}",
            )
            edit_forma = edit_col_5.selectbox(
                "Forma pagamento",
                FORMAS_PAGAMENTO,
                index=FORMAS_PAGAMENTO.index(normalizar_forma_pagamento(recebivel["forma_pagamento"])),
                key=f"rec_forma_{recebivel_id}",
            )
            edit_observacao = st.text_area(
                "Observacao",
                recebivel["observacao"] or "",
                key=f"rec_obs_{recebivel_id}",
            )
            edit_status = st.selectbox(
                "Status do recebivel",
                ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"],
                index=["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"].index((recebivel["status"] or "Aberto") if (recebivel["status"] or "Aberto") in ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"] else "Aberto"),
                key=f"rec_status_{recebivel_id}",
            )

            if st.button("Salvar alteracoes do recebivel"):
                if not edit_nome.strip():
                    st.error("Informe o nome do paciente.")
                elif not edit_prontuario.strip():
                    st.error("Informe o prontuario.")
                elif not parse_data_contrato(edit_vencimento):
                    st.error("Informe o vencimento no formato DD/MM/AAAA.")
                else:
                    atualizar_recebivel(
                        recebivel_id,
                        edit_nome,
                        edit_prontuario,
                        edit_vencimento,
                        edit_valor,
                        edit_forma,
                        edit_status,
                        edit_observacao,
                    )
                    conn.commit()
                    st.success("Recebivel atualizado.")

    with tab_lote:
        st.subheader("Editar recebiveis em lote")
        if recebiveis.empty:
            st.info("Nao ha recebiveis para editar em lote.")
        else:
            recebiveis_com_contrato = recebiveis[recebiveis["contrato_id"].notna()].copy()
            recebiveis_avulsos = recebiveis[recebiveis["contrato_id"].isna()].copy()

            if not recebiveis_avulsos.empty:
                st.caption("Recebiveis sem contrato vinculado continuam no controle financeiro e podem ser alterados na aba Editar Individual.")

            if recebiveis_com_contrato.empty:
                st.info("Nao ha lotes vinculados a contrato para editar em lote.")
            else:
                lotes_recebiveis = (
                    recebiveis_com_contrato.groupby(["contrato_id", "paciente_nome", "prontuario"], as_index=False)
                    .agg(
                        primeiro_vencimento=("vencimento", "first"),
                        quantidade_parcelas=("id", "count"),
                    )
                    .sort_values(["paciente_nome", "prontuario", "primeiro_vencimento"])
                )
                opcoes_lotes_recebiveis = [
                    (
                        int(row["contrato_id"]),
                        f"{row['paciente_nome']} - Prontuario {formatar_prontuario_valor(row['prontuario'])} - {int(row['quantidade_parcelas'])} parcelas - inicio {row['primeiro_vencimento']}"
                    )
                    for _, row in lotes_recebiveis.iterrows()
                ]
                contrato_lote_id = st.selectbox(
                    "Lote para editar",
                    options=[opcao[0] for opcao in opcoes_lotes_recebiveis],
                    format_func=lambda contrato_id: next(
                        label for valor, label in opcoes_lotes_recebiveis if valor == contrato_id
                    ),
                )
                recebiveis_lote = recebiveis_com_contrato[recebiveis_com_contrato["contrato_id"] == contrato_lote_id].copy()
                if recebiveis_lote.empty:
                    st.warning("Nao encontrei recebiveis para este lote. Selecione outro contrato.")
                else:
                    recebivel_base = recebiveis_lote.iloc[0]

                    lote_col_1, lote_col_2 = st.columns(2)
                    lote_nome = lote_col_1.text_input("Nome do paciente", recebivel_base["paciente_nome"] or "", key=f"lote_nome_{contrato_lote_id}")
                    lote_prontuario = lote_col_2.text_input("Prontuario", recebivel_base["prontuario"] or "", key=f"lote_pront_{contrato_lote_id}")

                    lote_col_3, lote_col_4, lote_col_5 = st.columns(3)
                    lote_primeiro_vencimento = lote_col_3.text_input(
                        "Novo primeiro vencimento (opcional)",
                        recebivel_base["vencimento"] or "",
                        key=f"lote_venc_{contrato_lote_id}",
                    )
                    lote_forma = lote_col_4.selectbox(
                        "Forma pagamento",
                        FORMAS_PAGAMENTO,
                        index=FORMAS_PAGAMENTO.index(normalizar_forma_pagamento(recebivel_base["forma_pagamento"])),
                        key=f"lote_forma_{contrato_lote_id}",
                    )
                    lote_status = lote_col_5.selectbox(
                        "Status",
                        ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"],
                        index=["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"].index((recebivel_base["status"] or "Aberto") if (recebivel_base["status"] or "Aberto") in ["Aberto", "Pago", "Atrasado", "Suspenso", "Cancelado"] else "Aberto"),
                        key=f"lote_status_{contrato_lote_id}",
                    )
                    lote_observacao = st.text_area(
                        "Observacao para todos os recebiveis deste contrato",
                        recebivel_base["observacao"] or "",
                        key=f"lote_obs_{contrato_lote_id}",
                    )

                    st.write("Recebiveis deste contrato")
                    st.dataframe(
                        recebiveis_lote[
                            ["paciente_nome", "prontuario", "parcela_numero", "vencimento", "valor", "forma_pagamento", "status"]
                        ].rename(
                            columns={
                                "paciente_nome": "Paciente",
                                "prontuario": "Prontuario",
                                "parcela_numero": "Parcela",
                                "vencimento": "Vencimento",
                                "valor": "Valor",
                                "forma_pagamento": "Forma pagamento",
                                "status": "Status",
                            }
                        ),
                        use_container_width=True,
                    )

                    if st.button("Salvar alteracoes em lote"):
                        if not lote_nome.strip():
                            st.error("Informe o nome do paciente.")
                        elif not lote_prontuario.strip():
                            st.error("Informe o prontuario.")
                        elif lote_primeiro_vencimento.strip() and not parse_data_contrato(lote_primeiro_vencimento):
                            st.error("Informe o primeiro vencimento no formato DD/MM/AAAA.")
                        else:
                            atualizar_recebiveis_lote_contrato(
                                contrato_lote_id,
                                lote_nome,
                                lote_prontuario,
                                lote_forma,
                                lote_status,
                                lote_observacao,
                                lote_primeiro_vencimento,
                            )
                            conn.commit()
                            st.success("Recebiveis do contrato atualizados em lote.")

    with tab_pagar:
        st.subheader("Contas a pagar")
        if contas_pagar.empty:
            st.info("Nao ha contas a pagar cadastradas.")
        else:
            blocos_pagar = selecionar_blocos_visiveis(
                "Exibir em contas a pagar",
                ["Filtros e resumo", "Detalhamento", "Atualização rápida"],
                ["Filtros e resumo", "Detalhamento"],
                "financeiro_blocos_pagar",
            )
            filtro_fornecedor = []
            filtro_status_pagar = []
            filtro_categoria_pagar = []
            filtro_periodo_pagar = ()
            if "Filtros e resumo" in blocos_pagar:
                with st.expander("Filtros e resumo", expanded=True):
                    fp1, fp2, fp3 = st.columns(3)
                    filtro_fornecedor = fp1.multiselect(
                        "Fornecedores",
                        options=sorted([valor for valor in contas_pagar["fornecedor"].fillna("").unique().tolist() if valor]),
                    )
                    filtro_status_pagar = fp2.multiselect(
                        "Status",
                        options=sorted([valor for valor in contas_pagar["status"].fillna("A vencer").unique().tolist() if valor]),
                    )
                    filtro_categoria_pagar = fp3.multiselect(
                        "Categorias",
                        options=sorted([valor for valor in contas_pagar["categoria"].fillna("").unique().tolist() if valor]),
                    )
                    datas_pagar = [
                        valor.date()
                        for valor in contas_pagar["_ordem_vencimento"].dropna().sort_values().unique().tolist()
                    ]
                    pagar_inicio = datas_pagar[0] if datas_pagar else None
                    pagar_fim = datas_pagar[-1] if datas_pagar else None
                    filtro_periodo_pagar = st.date_input(
                        "Periodo de vencimento",
                        value=(pagar_inicio, pagar_fim) if pagar_inicio and pagar_fim else (),
                    )

            pagar_filtrado = contas_pagar.copy()
            if filtro_fornecedor:
                pagar_filtrado = pagar_filtrado[pagar_filtrado["fornecedor"].isin(filtro_fornecedor)]
            if filtro_status_pagar:
                pagar_filtrado = pagar_filtrado[pagar_filtrado["status"].isin(filtro_status_pagar)]
            if filtro_categoria_pagar:
                pagar_filtrado = pagar_filtrado[pagar_filtrado["categoria"].isin(filtro_categoria_pagar)]
            if isinstance(filtro_periodo_pagar, tuple) and len(filtro_periodo_pagar) == 2 and filtro_periodo_pagar[0] and filtro_periodo_pagar[1]:
                inicio = pd.to_datetime(filtro_periodo_pagar[0])
                fim = pd.to_datetime(filtro_periodo_pagar[1])
                pagar_filtrado = pagar_filtrado[
                    (pagar_filtrado["_ordem_vencimento"] >= inicio) &
                    (pagar_filtrado["_ordem_vencimento"] <= fim)
                ]

            if "Filtros e resumo" in blocos_pagar:
                p1, p2, p3, p4 = st.columns(4)
                p1.metric("Total filtrado", formatar_moeda_br(float(pagar_filtrado["valor"].sum())))
                p2.metric("Titulos", str(len(pagar_filtrado)))
                p3.metric("Fornecedores", str(pagar_filtrado["fornecedor"].nunique()))
                p4.metric("Pagos", formatar_moeda_br(float(pagar_filtrado["valor_pago"].sum())))

            detalhe_pagar = pagar_filtrado[
                ["data_vencimento", "descricao", "fornecedor", "categoria", "valor", "pago", "valor_pago", "status", "observacao"]
            ].copy()
            detalhe_pagar = detalhe_pagar.rename(
                columns={
                    "data_vencimento": "Vencimento",
                    "descricao": "Descricao",
                    "fornecedor": "Fornecedor",
                    "categoria": "Categoria",
                    "valor": "Valor",
                    "pago": "Data pagamento",
                    "valor_pago": "Valor pago",
                    "status": "Status",
                    "observacao": "Observacao",
                }
            )
            detalhe_pagar["Valor"] = detalhe_pagar["Valor"].map(formatar_moeda_br)
            detalhe_pagar["Valor pago"] = detalhe_pagar["Valor pago"].map(formatar_moeda_br)
            if "Detalhamento" in blocos_pagar:
                with st.expander("Detalhamento", expanded=False):
                    st.dataframe(detalhe_pagar, use_container_width=True, hide_index=True)

            opcoes_pagar = [
                (int(row["id"]), f"{row['data_vencimento']} - {row['fornecedor']} - {row['descricao']} - {formatar_moeda_br(row['valor'])}")
                for _, row in pagar_filtrado.iterrows()
            ]
            if "Atualização rápida" in blocos_pagar and opcoes_pagar:
                st.markdown("**Atualizacao rapida**")
                conta_pagar_id = st.selectbox(
                    "Titulo para atualizar",
                    options=[opcao[0] for opcao in opcoes_pagar],
                    format_func=lambda valor: next(rotulo for chave, rotulo in opcoes_pagar if chave == valor),
                )
                conta_pagar_row = contas_pagar[contas_pagar["id"] == conta_pagar_id].iloc[0]
                ap1, ap2 = st.columns(2)
                novo_titulo_pagar = ap1.text_input(
                    "Titulo / descricao",
                    conta_pagar_row["descricao"] or "",
                    key=f"descricao_pagar_{conta_pagar_id}",
                )
                novo_fornecedor_pagar = ap2.text_input(
                    "Fornecedor",
                    conta_pagar_row["fornecedor"] or "",
                    key=f"fornecedor_pagar_{conta_pagar_id}",
                )

                ap_cat = st.selectbox(
                    "Categoria",
                    CATEGORIAS_CONTAS_PAGAR,
                    index=CATEGORIAS_CONTAS_PAGAR.index(conta_pagar_row["categoria"]) if conta_pagar_row["categoria"] in CATEGORIAS_CONTAS_PAGAR else CATEGORIAS_CONTAS_PAGAR.index("Outros"),
                    key=f"categoria_pagar_{conta_pagar_id}",
                )

                ap3, ap4, ap5 = st.columns(3)
                nova_data_vencimento = ap3.text_input(
                    "Vencimento",
                    formatar_data_br_valor(conta_pagar_row["data_vencimento"] or ""),
                    key=f"vencimento_pagar_{conta_pagar_id}",
                )
                novo_valor_titulo = ap4.number_input(
                    "Valor do titulo",
                    min_value=0.0,
                    value=float(conta_pagar_row["valor"] or 0),
                    key=f"valor_titulo_pagar_{conta_pagar_id}",
                )
                novo_status_pagar = ap5.selectbox(
                    "Status",
                    STATUS_CONTAS_PAGAR,
                    index=STATUS_CONTAS_PAGAR.index(conta_pagar_row["status"]) if conta_pagar_row["status"] in STATUS_CONTAS_PAGAR else 0,
                    key=f"status_pagar_{conta_pagar_id}",
                )

                ap6, ap7 = st.columns(2)
                nova_data_pagamento = ap6.text_input(
                    "Data do pagamento",
                    formatar_data_br_valor(conta_pagar_row["pago"] or ""),
                    key=f"data_pagamento_pagar_{conta_pagar_id}",
                )
                novo_valor_pago = ap7.number_input(
                    "Valor pago",
                    min_value=0.0,
                    value=float(conta_pagar_row["valor_pago"] or 0),
                    key=f"valor_pago_pagar_{conta_pagar_id}",
                )

                nova_obs_pagar = st.text_area(
                    "Observacao",
                    conta_pagar_row["observacao"] or "",
                    key=f"obs_pagar_{conta_pagar_id}",
                )
                if st.button("Salvar conta a pagar"):
                    if not novo_titulo_pagar.strip():
                        st.error("Informe o titulo da conta.")
                    elif not novo_fornecedor_pagar.strip():
                        st.error("Informe o fornecedor.")
                    elif not parse_data_contrato(nova_data_vencimento):
                        st.error("Informe o vencimento no formato DD/MM/AAAA.")
                    elif nova_data_pagamento.strip() and not parse_data_contrato(nova_data_pagamento):
                        st.error("Informe a data do pagamento no formato DD/MM/AAAA.")
                    else:
                        cursor.execute(
                            """
                            UPDATE contas_pagar
                            SET data_vencimento=?, descricao=?, fornecedor=?, categoria=?, valor=?, status=?, pago=?, valor_pago=?, observacao=?
                            WHERE id=?
                            """,
                            (
                                formatar_data_br_valor(nova_data_vencimento),
                                novo_titulo_pagar.strip(),
                                novo_fornecedor_pagar.strip(),
                                ap_cat,
                                float(novo_valor_titulo),
                                novo_status_pagar,
                                formatar_data_br_valor(nova_data_pagamento),
                                float(novo_valor_pago),
                                nova_obs_pagar.strip(),
                                int(conta_pagar_id),
                            ),
                        )
                        conn.commit()
                        st.success("Conta a pagar atualizada.")
                        st.rerun()

    with tab_calendario_pagar:
        st.subheader("Calendario de vencimentos")
        if contas_pagar.empty:
            st.info("Nao ha contas a pagar para exibir no calendario.")
        else:
            contas_validas = contas_pagar[contas_pagar["_ordem_vencimento"].notna()].copy()
            if contas_validas.empty:
                st.info("Nao ha vencimentos validos para exibir no calendario.")
            else:
                anos_disponiveis = sorted(contas_validas["_ordem_vencimento"].dt.year.unique().tolist())
                ano_atual = date.today().year
                mes_atual = date.today().month
                ano_cal = st.selectbox(
                    "Ano",
                    options=anos_disponiveis,
                    index=anos_disponiveis.index(ano_atual) if ano_atual in anos_disponiveis else len(anos_disponiveis) - 1,
                    key="ano_calendario_pagar",
                )
                meses_disponiveis_ano = sorted(
                    contas_validas[contas_validas["_ordem_vencimento"].dt.year == ano_cal]["_ordem_vencimento"].dt.month.unique().tolist()
                )
                if not meses_disponiveis_ano:
                    meses_disponiveis_ano = [mes_atual]
                mes_cal = st.selectbox(
                    "Mes",
                    options=meses_disponiveis_ano,
                    index=meses_disponiveis_ano.index(mes_atual) if mes_atual in meses_disponiveis_ano else 0,
                    format_func=lambda valor: {
                        1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
                        7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
                    }[valor],
                    key="mes_calendario_pagar",
                )
                contas_mes = contas_validas[
                    (contas_validas["_ordem_vencimento"].dt.year == ano_cal) &
                    (contas_validas["_ordem_vencimento"].dt.month == mes_cal)
                ].copy()
                st.caption("Pagamentos agrupados por dia de vencimento.")
                cabecalho = st.columns(7)
                for coluna, nome_dia in zip(cabecalho, ["Seg", "Ter", "Qua", "Qui", "Sex", "Sab", "Dom"]):
                    coluna.markdown(f"**{nome_dia}**")

                hoje = date.today()
                for semana in calendar.monthcalendar(ano_cal, mes_cal):
                    cols = st.columns(7)
                    for idx, dia in enumerate(semana):
                        with cols[idx]:
                            if dia == 0:
                                st.write("")
                            else:
                                eh_hoje = hoje.year == ano_cal and hoje.month == mes_cal and hoje.day == dia
                                contas_dia = contas_mes[contas_mes["_ordem_vencimento"].dt.day == dia]
                                statuses_dia = set(contas_dia["status"].fillna("").tolist()) if not contas_dia.empty else set()
                                borda_dia = "2px solid #1f2937" if eh_hoje else "1px solid #e5e7eb"
                                if "Atrasado" in statuses_dia:
                                    fundo_dia = "#fef2f2"
                                elif "Pago" in statuses_dia and statuses_dia.issubset({"Pago"}):
                                    fundo_dia = "#f0fdf4"
                                else:
                                    fundo_dia = "#f8fafc" if eh_hoje else "#ffffff"
                                st.markdown(
                                    f"<div style='border:{borda_dia};border-radius:10px;padding:8px;background:{fundo_dia};'>"
                                    f"<div style='font-weight:700;margin-bottom:8px;'>{dia:02d}/{mes_cal:02d}/{ano_cal}</div>",
                                    unsafe_allow_html=True,
                                )
                                if contas_dia.empty:
                                    st.caption("Sem vencimentos")
                                else:
                                    for _, conta in contas_dia.iterrows():
                                        status_conta = str(conta["status"] or "")
                                        if status_conta == "Pago":
                                            cor = "#15803d"
                                            fundo = "#dcfce7"
                                            borda = "#86efac"
                                        elif status_conta == "Atrasado":
                                            cor = "#b91c1c"
                                            fundo = "#fee2e2"
                                            borda = "#fca5a5"
                                        else:
                                            cor = "#111827"
                                            fundo = "#ffffff"
                                            borda = "#d1d5db"
                                        st.markdown(
                                            f"<div style='border:1px solid {borda};background:{fundo};padding:6px;margin-bottom:6px;border-radius:6px;'>"
                                            f"<div style='font-size:12px;color:{cor};'><strong>{conta['fornecedor']}</strong></div>"
                                            f"<div style='font-size:12px;'>{conta['descricao']}</div>"
                                            f"<div style='font-size:12px;'><strong>{formatar_moeda_br(conta['valor'])}</strong></div>"
                                            f"<div style='font-size:11px;color:#666;'>{conta['status']}</div>"
                                            f"</div>",
                                            unsafe_allow_html=True,
                                        )
                                st.markdown("</div>", unsafe_allow_html=True)

    with tab_novo_pagar:
        st.subheader("Nova divida")
        np1, np2 = st.columns(2)
        nova_data_venc = np1.text_input("Data de vencimento", value=formatar_data_br(date.today()))
        novo_fornecedor = np2.text_input("Fornecedor")
        nova_descricao = st.text_input("Descricao")
        nova_categoria = st.selectbox("Categoria", CATEGORIAS_CONTAS_PAGAR, index=CATEGORIAS_CONTAS_PAGAR.index("Outros"))
        np3, np4, np5 = st.columns(3)
        novo_valor_pagar = np3.number_input("Valor", min_value=0.0, value=0.0)
        nova_quantidade_parcelas = np4.number_input("Quantidade de parcelas", min_value=1, value=1, step=1)
        novo_status_pagar = np5.selectbox("Status", STATUS_CONTAS_PAGAR, index=0)
        np6, np7 = st.columns(2)
        nova_data_pago = np6.text_input("Data de pagamento", value="")
        novo_valor_pago = np7.number_input("Valor pago", min_value=0.0, value=0.0, key="novo_valor_pago_manual")
        nova_obs_divida = st.text_area("Observacao")
        if st.button("Adicionar divida"):
            if not parse_data_contrato(nova_data_venc):
                st.error("Informe o vencimento no formato DD/MM/AAAA.")
            elif nova_data_pago.strip() and not parse_data_contrato(nova_data_pago):
                st.error("Informe a data de pagamento no formato DD/MM/AAAA.")
            elif not nova_descricao.strip() and not novo_fornecedor.strip():
                st.error("Informe pelo menos a descricao ou o fornecedor.")
            elif float(novo_valor_pagar) <= 0:
                st.error("Informe um valor maior que zero.")
            else:
                data_base_parcela = parse_data_contrato(nova_data_venc)
                quantidade_parcelas = int(nova_quantidade_parcelas or 1)
                valor_parcela = round(float(novo_valor_pagar) / quantidade_parcelas, 2)
                total_distribuido = 0.0

                for indice_parcela in range(quantidade_parcelas):
                    valor_atual = valor_parcela
                    if indice_parcela == quantidade_parcelas - 1:
                        valor_atual = round(float(novo_valor_pagar) - total_distribuido, 2)
                    total_distribuido += valor_atual

                    data_parcela = adicionar_meses(data_base_parcela, indice_parcela)
                    descricao_parcela = nova_descricao.strip()
                    if quantidade_parcelas > 1:
                        descricao_parcela = f"{descricao_parcela} - Parcela {indice_parcela + 1}/{quantidade_parcelas}".strip(" -")

                    cursor.execute(
                        """
                        INSERT INTO contas_pagar
                        (data_vencimento, descricao, fornecedor, categoria, valor, pago, valor_pago, status, observacao, data_criacao, hash_importacao)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            formatar_data_br(data_parcela),
                            descricao_parcela,
                            novo_fornecedor.strip(),
                            nova_categoria,
                            valor_atual,
                            formatar_data_br_valor(nova_data_pago) if indice_parcela == 0 else "",
                            float(novo_valor_pago) if indice_parcela == 0 else 0.0,
                            novo_status_pagar,
                            nova_obs_divida.strip(),
                            agora_str(),
                            montar_hash_importacao_conta_pagar(
                                formatar_data_br(data_parcela),
                                descricao_parcela,
                                novo_fornecedor,
                                valor_atual,
                            ),
                        ),
                    )
                conn.commit()
                st.success("Divida adicionada.")
                st.rerun()
